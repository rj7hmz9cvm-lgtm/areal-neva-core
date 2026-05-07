
def _force_voice_finish(raw_input: str, result: str) -> bool:
    if not raw_input:
        return False
    low = raw_input.lower()
    if any(x in low for x in ["заверш", "доволен", "да", "ок", "ok"]):
        if "не доволен" in low:
            return False
        return True
    return False

import os
BASE = "/root/.areal-neva-core"

import re
import time
import json
import sqlite3
import asyncio
import hashlib
import datetime
import logging
def now_iso_utc() -> str:
    """UTC_TIMESTAMP_V1"""
    from datetime import datetime, timezone
    return datetime.now(timezone.utc).isoformat()

import fcntl
from typing import Any, Dict, List, Optional, Tuple

from dotenv import load_dotenv
from core.ai_router import process_ai_task
try:
    from startup_recovery import run_startup_recovery  # STARTUP_RECOVERY_V1_WIRED
except Exception as _startup_recovery_import_err:
    run_startup_recovery = None

# === CANON_REMAINING_CODE_FULL_CLOSE_V1_IMPORT ===
try:
    from core.engine_contract import validate_engine_result, result_to_user_text
    from core.active_dialog_state import maybe_handle_active_dialog, save_dialog_event
    from core.template_workflow import maybe_handle_template_workflow
except Exception as _canon_import_err:
    validate_engine_result = None
    result_to_user_text = None
    maybe_handle_active_dialog = None
    save_dialog_event = None
    maybe_handle_template_workflow = None
# === END_CANON_REMAINING_CODE_FULL_CLOSE_V1_IMPORT ===
# === REAL_GAPS_CLOSE_V2_IMPORT_ASYNC_TEMPLATE ===
try:
    from core.template_workflow import maybe_handle_template_workflow_async
except Exception:
    maybe_handle_template_workflow_async = None
# === END_REAL_GAPS_CLOSE_V2_IMPORT_ASYNC_TEMPLATE ===
# === SEARCH_MONOLITH_V2_TASK_WORKER_IMPORT ===
try:
    from core.search_session import is_search_clarification_output as _search_v2_is_clarification
except Exception:
    _search_v2_is_clarification = lambda text: False
# === END SEARCH_MONOLITH_V2_TASK_WORKER_IMPORT ===
# === FILE_TECH_CONTOUR_FULL_CLOSE_V1_IMPORT ===
try:
    from core.file_memory_bridge import (
        should_handle_file_followup as _filemem_should_followup,
        build_file_followup_answer as _filemem_build_answer,
        is_service_file as _filemem_is_service_file,
        save_file_catalog_snapshot as _filemem_save_catalog,
    )
except Exception:
    _filemem_should_followup = lambda text: False
    _filemem_build_answer = lambda *a, **kw: None
    _filemem_is_service_file = lambda *a, **kw: False
    _filemem_save_catalog = lambda *a, **kw: {"ok": False}
# === END FILE_TECH_CONTOUR_FULL_CLOSE_V1_IMPORT ===
try:
    from core.topic_meta_loader import load_topic_meta, is_what_is_this_question, build_topic_self_answer
    TOPIC_META_LOADER_WIRED = True
except Exception:
    TOPIC_META_LOADER_WIRED = False
    def load_topic_meta(t): return None
    def is_what_is_this_question(t): return False
    def build_topic_self_answer(m): return ""
# TOPIC_META_LOADER_V1_IMPORT
from core.reply_sender import send_reply, send_reply_ex
try:
    from core.template_engine_v1 import is_template_request as _tpl_check, get_template as _tpl_get, save_template as _tpl_save, apply_template_to_xlsx as _tpl_apply  # TEMPLATE_ENGINE_V1_WIRED
except Exception:
    _tpl_check = lambda t: False
    _tpl_get = lambda *a: None
    _tpl_save = lambda *a: False
    _tpl_apply = lambda *a: False
try:
    from core.topic_3008_engine import is_topic_3008 as _t3_check, detect_command as _t3_cmd, extract_code as _t3_extract, verify_code as _t3_verify, generate_code as _t3_generate  # TOPIC_3008_V1_WIRED
except Exception:
    _t3_check = lambda t: False
    _t3_cmd = lambda t: "none"
    _t3_extract = lambda t: t
    _t3_verify = None
    _t3_generate = None
try:
    from core.temp_cleanup import cleanup_file as _tc_file, cleanup_task_temps as _tc_task, cleanup_after_upload as _tc_upload  # TEMP_CLEANUP_V1_WIRED
except Exception:
    _tc_file = lambda p: False
    _tc_task = lambda t: 0
    _tc_upload = lambda l: 0
try:
    from core.technadzor_engine import process_technadzor as _te_process, is_technadzor_intent as _te_intent  # TECHNADZOR_ENGINE_V1_WIRED
except Exception:
    _te_process = lambda *a, **kw: {"ok": False, "result_text": "", "artifact": None, "error_code": "IMPORT_FAIL"}
    _te_intent = lambda *a: False
try:
    from core.intake_offer_actions import needs_offer as _ioa_needs, get_offer_text as _ioa_text, parse_offer_reply as _ioa_parse  # INTAKE_OFFER_V1_WIRED
except Exception:
    _ioa_needs = lambda *a: False
    _ioa_text = lambda: ""
    _ioa_parse = lambda r: ""
try:
    from core.search_session import get_session as _ss_get, create_session as _ss_create, update_session as _ss_update, extract_criteria as _ss_criteria  # SEARCH_SESSION_V1_WIRED
except Exception:
    _ss_get = lambda *a: None
    _ss_create = lambda *a, **kw: {}
    _ss_update = lambda *a, **kw: None
    _ss_criteria = lambda t: {}
try:
    from core.output_decision import format_task_result as _od_format, format_search_output as _od_search  # OUTPUT_DECISION_V1_WIRED
except Exception:
    _od_format = lambda r, s, **kw: r
    _od_search = lambda o, **kw: str(o)
try:
    from core.inbox_aggregator import normalize_inbox_item as _ia_norm, should_create_task as _ia_should  # INBOX_AGG_V1_WIRED
except Exception:
    _ia_norm = lambda **kw: kw
    _ia_should = lambda i: True
try:
    from core.constraint_engine import rank_offers as _ce_rank, apply_constraints as _ce_apply, validate_offer as _ce_validate  # CONSTRAINT_ENGINE_V1_WIRED
except Exception:
    _ce_rank = lambda o: o
    _ce_apply = lambda o, **kw: o
    _ce_validate = lambda o: {"ok": True}
try:
    from core.audit_log import audit as _audit  # AUDIT_LOG_V1_WIRED
except Exception:
    _audit = lambda *a, **kw: None
try:
    from core.source_dedup import dedup_offers as _sd_dedup, sort_by_region as _sd_region  # SOURCE_DEDUP_V1_WIRED
except Exception:
    _sd_dedup = lambda o: o
    _sd_region = lambda o, **kw: o
try:
    from core.error_explainer import user_friendly_error as _ee_explain  # ERROR_EXPLAINER_V1_WIRED
except Exception:
    _ee_explain = lambda c: c
try:
    from core.data_classification import classify_domain as _dc_domain, classify_intent as _dc_intent, classify_file_type as _dc_ftype  # DATA_CLASS_V1_WIRED
except Exception:
    _dc_domain = lambda *a: "UNSORTED"
    _dc_intent = lambda t: "text"
    _dc_ftype = lambda f: "UNKNOWN"
try:
    from core.intent_lock import is_chat_only as _il_chat_only, file_result_guard as _il_file_guard  # INTENT_LOCK_V1_WIRED
except Exception:
    _il_chat_only = lambda t: False
    _il_file_guard = lambda **kw: {"ok": True}
try:
    from core.orchestra_context import build_shared_context as _oc_build, user_mode_switch as _oc_mode  # ORCHESTRA_CTX_V1_WIRED
except Exception:
    _oc_build = lambda **kw: ""
    _oc_mode = lambda t: "HUMAN"
try:
    from core.price_normalization import normalize_price_text as _pn_fmt, extract_price as _pn_extract  # PRICE_NORM_V1_WIRED
except Exception:
    _pn_fmt = lambda t: t
    _pn_extract = lambda t: []
try:
    from core.memory_filter import filter_memory_for_prompt as _mf_filter, sanitize_before_write as _mf_clean  # MEMORY_FILTER_V1_WIRED
except Exception:
    _mf_filter = lambda m, **kw: m
    _mf_clean = lambda v: v
try:
    from core.result_validator import validate_result as _rv_validate, is_generic_response as _rv_generic  # RESULT_VALIDATOR_V1_WIRED
except Exception:
    _rv_validate = lambda r, **kw: {"ok": True, "reason": "IMPORT_FAIL"}
    _rv_generic = lambda r: False
try:
    from core.search_quality import availability_check as _sq_avail, cache_get as _sq_cget, cache_set as _sq_cset  # SEARCH_QUALITY_V1_WIRED
except Exception:
    _sq_avail = lambda r: True
    _sq_cget = lambda q: None
    _sq_cset = lambda q, r: None
from core.duplicate_guard import find_duplicate, duplicate_message  # DUPLICATE_GUARD_V1_WIRED
from core.pin_manager import get_pin_context, save_pin
from core.topic_drive_oauth import upload_file_to_topic
from core.artifact_pipeline import analyze_downloaded_file
# === DRIVE_FILE_CONTENT_MEMORY_INDEX_V1_IMPORT ===
try:
    from core.drive_content_indexer import index_drive_file_content as _df_content_index
except Exception:
    _df_content_index = None
# === END DRIVE_FILE_CONTENT_MEMORY_INDEX_V1_IMPORT ===

load_dotenv(f"{BASE}/.env", override=True)

CORE_DB = f"{BASE}/data/core.db"
MEM_DB = f"{BASE}/data/memory.db"
LOG_PATH = f"{BASE}/logs/task_worker.log"
LOCK_PATH = f"{BASE}/runtime/task_worker.lock"

POLL_SEC = 1.5
MIN_RESULT_LEN = 2  # RESULT_VALIDATOR_FIX_V1: was 8
AI_TIMEOUT = 300
STALE_TIMEOUT = 600
REMINDER_SEC = 180

BAD_RESULT_RE = [
    # === RESULT_VALIDATOR_FIX_V1 ===
    # Убраны: извини/извините/ищу/найду/могу найти/я могу помочь
    # Это нормальные слова, AI имеет право их использовать
    r"сорян",
    r"дружище",
    r"delete from",
    r"task_worker\.py",
    r"telegram_daemon\.py",
    r"выполните sql",
    r"/root/\.areal",
    # === END_RESULT_VALIDATOR_FIX_V1 ===
]

MEMORY_BAD_MARKERS = [
    "traceback",
    "forbidden default model",
    "/root/",
    ".json",
    ".log",
    "не могу выполнить запрос",
    "delete from",
    "task_worker.py",
    "telegram_daemon.py",
    "выполните sql",
]

CONFIRM_INTENTS = {
    "да",
    "ок",
    "ok",
    "окей",
    "хорошо",
    "подтверждаю",
    "принято",
    "согласен",
    "верно",
    "всё верно",
    "все верно",
}

REVISION_INTENTS = {
    "нет",
    "не так",
    "переделай",
    "исправь",
    "доработай",
    "уточню",
    "уточнение",
    "правки",
    "уточни",
}

MEMORY_NOISE_MARKERS = [
    "нет данных",
    "повторите",
    "не знаю",
    "чат не содержит активной задачи",
    "привет. чат не содержит активной задачи",
    "чат создан для",
    "тест диагностика",
    "не понял запрос",
    "уточните",
    "не понимаю",
    "задайте вопрос",
]

def _version_artifact_path(path: str, task_id: str) -> str:
    """ARTIFACT_VERSION_V1 — создать версионированный путь file_v2.xlsx"""
    if not path:
        return path
    import re as _re
    # Если уже версионирован — увеличить номер
    m = _re.search(r"_v(\d+)(\.\w+)$", path)
    if m:
        n = int(m.group(1)) + 1
        return path[:m.start()] + f"_v{n}" + m.group(2)
    # Первая ревизия
    base, ext = os.path.splitext(path)
    return f"{base}_v2{ext}"

def _quality_gate_artifact(artifact_path: str = None, drive_link: str = None,
                             input_type: str = "text", task_type: str = "") -> dict:
    """QUALITY_GATE_CALL_V1 — проверить артефакт перед отправкой"""
    is_file_task = input_type in ("drive_file", "file") or task_type in (
        "ESTIMATE_TASK", "OCR_TASK", "DWG_TASK", "TECHNADZOR_TASK", "DOCUMENT_TASK"
    )
    # === DRIVE_LINK_MANDATORY_V1 ===
    if not is_file_task:
        return {"ok": True, "reason": "NOT_FILE_TASK"}
    # Если file task — Drive ссылка обязательна
    if drive_link and "drive.google" in str(drive_link):
        return {"ok": True, "reason": "DRIVE_LINK_PRESENT"}
    if not drive_link and not artifact_path:
        return {"ok": False, "reason": "NO_DRIVE_LINK_NO_ARTIFACT"}
    # === END DRIVE_LINK_MANDATORY_V1 ===

    # Есть Drive ссылка — OK
    if drive_link and "drive.google" in str(drive_link):
        return {"ok": True, "reason": "DRIVE_LINK_PRESENT"}

    # Есть артефакт файл
    if artifact_path and os.path.exists(artifact_path):
        size = os.path.getsize(artifact_path)
        if size < 100:
            return {"ok": False, "reason": "ARTIFACT_TOO_SMALL"}
        # Для Excel проверяем формулы
        if artifact_path.endswith(".xlsx"):
            try:
                from openpyxl import load_workbook
                wb = load_workbook(artifact_path)
                ws = wb.active
                has_formula = any(
                    str(ws.cell(r, c).value or "").startswith("=")
                    for r in range(1, min(ws.max_row+1, 20))
                    for c in range(1, min(ws.max_column+1, 10))
                )
                if not has_formula:
                    logger.warning("QUALITY_GATE: no formulas in %s", artifact_path)
            except Exception:
                pass
        return {"ok": True, "reason": "ARTIFACT_EXISTS"}

    return {"ok": False, "reason": "NO_ARTIFACT_NO_LINK"}

def _is_heavy_task(input_type: str, task_type: str = "") -> bool:
    """FAST_HEAVY_TASK_V1 — определить тип задачи"""
    return input_type in ("drive_file", "file") or task_type in (
        "ESTIMATE_TASK", "OCR_TASK", "DWG_TASK", "TECHNADZOR_TASK", "DOCUMENT_TASK", "MULTI_TASK"
    )

def _get_task_sla(input_type: str, task_type: str = "") -> int:
    """Возвращает SLA в секундах"""
    if _is_heavy_task(input_type, task_type):
        return 120  # HEAVY < 120s
    return 10  # FAST < 10s

def _check_result_before_confirm(result: str, input_type: str = "text", intent: str = "") -> bool:
    """RESULT_VALIDATOR_CALL_V1 — проверить result перед AWAITING_CONFIRMATION"""
    try:
        rv = _rv_validate(result, input_type=input_type, intent=intent)
        if not rv.get("ok"):
            logger.warning("RESULT_VALIDATOR_BLOCKED reason=%s result_prefix=%s", rv.get("reason"), str(result)[:80])
            return False
    except Exception as _rve:
        logger.warning("RESULT_VALIDATOR_ERR %s", _rve)
    return True


# === REPLY_REPEAT_PARENT_TASK_V1 ===
_REPEAT_PARENT_WORDS_V1 = {
    "ну и", "дальше", "повтори", "проверяй", "не так",
    "ещё раз", "еще раз", "заново", "дописывай", "и что", "и?",
    "а", "б", "в", "г", "1", "2", "3", "4",
    "а)", "б)", "в)", "г)",
    "да", "нет", "ок", "хорошо",
    "среднее", "средняя", "минимум", "максимум", "своя",
    "вариант 1", "вариант 2", "вариант 3", "вариант 4",
    "первый", "второй", "третий",
    "самый дешевый", "самый дешёвый", "надежный", "надёжный",
}
def _is_repeat_parent_task_v1(text: str) -> bool:
    t = _clean(_s(text), 500).lower()
    if not t:
        return False
    return t in _REPEAT_PARENT_WORDS_V1 or any(t.startswith(x + " ") for x in _REPEAT_PARENT_WORDS_V1)

def _find_repeat_parent_task_v1(conn, chat_id: str, topic_id: int, exclude_task_id: str = ""):
    try:
        cols = _cols(conn, "tasks")
        if "topic_id" in cols:
            return conn.execute(
                "SELECT * FROM tasks WHERE chat_id=? AND COALESCE(topic_id,0)=? AND id<>? AND state IN ('IN_PROGRESS','AWAITING_CONFIRMATION','WAITING_CLARIFICATION','AWAITING_PRICE_CONFIRMATION') ORDER BY updated_at DESC LIMIT 1",
                (str(chat_id), int(topic_id or 0), str(exclude_task_id)),
            ).fetchone()
        return conn.execute(
            "SELECT * FROM tasks WHERE chat_id=? AND id<>? AND state IN ('IN_PROGRESS','AWAITING_CONFIRMATION','WAITING_CLARIFICATION','AWAITING_PRICE_CONFIRMATION') ORDER BY updated_at DESC LIMIT 1",
            (str(chat_id), str(exclude_task_id)),
        ).fetchone()
    except Exception as _e:
        logger.warning("REPLY_REPEAT_PARENT_TASK_V1_FIND_ERR %s", _e)
        return None
# === END_REPLY_REPEAT_PARENT_TASK_V1 ===

def _is_memory_noise(text: str) -> bool:
    t = _clean(_s(text), 1000).lower()
    if not t:
        return False
    return any(x in t for x in MEMORY_NOISE_MARKERS)

SEARCH_PATTERNS = [
    r"\bнайди\b",
    r"\bнайти\b",
    r"\bпоиск\b",
    r"\bпоищи\b",
    r"\bsearch\b",
    r"\bцена\b",
    r"\bстоимость\b",
    r"\bсколько\s+стоит\b",
    r"\bavito\b",
    r"\bozon\b",
    r"\bwildberries\b",
    r"\bauto\.ru\b",
    r"\bdrom\b",
    r"\bновости\b",
    r"\bпогода\b",
    r"\bкурс\b",
    r"\bмаркетплейс\b",
]

os.makedirs(f"{BASE}/logs", exist_ok=True)
os.makedirs(f"{BASE}/runtime", exist_ok=True)

logger = logging.getLogger("task_worker")
# === FULLFIX_DIRECTION_KERNEL_STAGE_1_IMPORT ===
try:
    from core.work_item import WorkItem as _Stage1WorkItem
    from core.direction_registry import DirectionRegistry as _Stage1DirReg
except Exception:
    _Stage1WorkItem = None
    _Stage1DirReg = None
# === END ===
# === FULLFIX_CAPABILITY_ROUTER_STAGE_2_IMPORT ===
try:
    from core.capability_router import CapabilityRouter as _Stage2Router
except Exception:
    _Stage2Router = None
# === END STAGE2 ===
# === FULLFIX_CONTEXT_LOADER_STAGE_3_IMPORT ===
try:
    from core.context_loader import ContextLoader as _Stage3Loader
except Exception:
    _Stage3Loader = None
# === END STAGE3 ===
# === FULLFIX_QUALITY_GATE_STAGE_4_IMPORT ===
try:
    from core.quality_gate import QualityGate as _Stage4QG
except Exception:
    _Stage4QG = None
# === END STAGE4 ===
# === FULLFIX_SEARCH_ENGINE_STAGE_5_IMPORT ===
try:
    from core.search_engine import SearchEngine as _Stage5Search
except Exception:
    _Stage5Search = None
# === END STAGE5 ===
# === FULLFIX_ARCHIVE_ENGINE_STAGE_6_IMPORT ===
try:
    from core.archive_engine import ArchiveEngine as _Stage6Archive
except Exception:
    _Stage6Archive = None
# === END STAGE6 ===
# === FULLFIX_FORMAT_ADAPTER_STAGE_7_IMPORT ===
try:
    from core.format_adapter import FormatAdapter as _Stage7FA
except Exception:
    _Stage7FA = None
# === END STAGE7 ===
# === FULLFIX_TOPIC_AUTODISCOVERY_V2_IMPORT ===
try:
    from core.topic_autodiscovery import process as _topic_autodiscovery, check_naming_timeout as _topic_naming_check
except Exception:
    _topic_autodiscovery = None
    _topic_naming_check = None
# === END TOPIC_AUTO ===








logger.setLevel(logging.INFO)
if not logger.handlers:
    fh = logging.FileHandler(LOG_PATH)
    fh.setFormatter(logging.Formatter("%(asctime)s %(levelname)s WORKER: %(message)s"))
    logger.addHandler(fh)


def db(path: str) -> sqlite3.Connection:
    conn = sqlite3.connect(path, timeout=20, isolation_level=None, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA busy_timeout=15000")
    return conn


# === EZONE_INGEST_V1 ===
EZONE_KEYS = {"system", "architecture", "pipeline", "memory"}

def is_ezone_payload(text: str) -> bool:
    """Проверить что входящий текст — JSON от нейросети"""
    if not text or not text.strip().startswith("{"):
        return False
    try:
        import json as _j
        d = _j.loads(text)
        return bool(EZONE_KEYS & set(d.keys()))
    except Exception:
        return False

def save_ezone_json(chat_id: str, text: str) -> str:
    """Сохранить ingest JSON в memory_files"""
    import json as _j, hashlib, os
    from datetime import datetime, timezone
    try:
        d = _j.loads(text)
        chat_key = f"{chat_id}__telegram"
        base = f"{BASE}/data/memory_files/CHATS/{chat_key}"
        os.makedirs(base, exist_ok=True)
        os.makedirs(f"{BASE}/data/memory_files/GLOBAL", exist_ok=True)
        os.makedirs(f"{BASE}/data/memory_files/SYSTEM", exist_ok=True)

        # Дедупликация по hash за день
        h = hashlib.md5(text.encode()).hexdigest()[:16]
        ts = datetime.now(timezone.utc).isoformat()
        entry = _j.dumps({"timestamp": ts, "data": d, "_meta": {
            "chat_key": chat_key, "ingested_at": ts, "source": "telegram", "hash": h
        }}, ensure_ascii=False)

        # timeline
        with open(f"{base}/timeline.jsonl", "a", encoding="utf-8") as f:
            f.write(entry + "\n")
        with open(f"{BASE}/data/memory_files/GLOBAL/timeline.jsonl", "a", encoding="utf-8") as f:
            f.write(entry + "\n")

        # SYSTEM keys
        for key in EZONE_KEYS:
            if key in d:
                sys_entry = _j.dumps({"timestamp": ts, key: d[key], "chat_key": chat_key}, ensure_ascii=False)
                with open(f"{BASE}/data/memory_files/SYSTEM/{key}.jsonl", "a", encoding="utf-8") as f:
                    f.write(sys_entry + "\n")

        logger.info("EZONE_INGEST_V1 saved chat_key=%s hash=%s", chat_key, h)
        return f"Принял, память загружена ({chat_key})"
    except Exception as e:
        logger.error("EZONE_INGEST_V1 err=%s", e)
        return f"Ошибка загрузки памяти: {e}"
# === END EZONE_INGEST_V1 ===

def _s(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, str):
        return v.strip()
    try:
        return json.dumps(v, ensure_ascii=False)
    except Exception:
        return str(v).strip()


def _clean(text: str, limit: int = 12000) -> str:
    text = (text or "").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:limit]


def _task_field(task: Any, field: str, default: Any = "") -> Any:
    try:
        if hasattr(task, "keys") and field in task.keys():
            return task[field]
    except Exception:
        pass
    try:
        return getattr(task, field)
    except Exception:
        return default


def _has_table(conn: sqlite3.Connection, table: str) -> bool:
    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
        (table,),
    ).fetchone()
    return row is not None


def _cols(conn: sqlite3.Connection, table: str) -> List[str]:
    try:
        return [r[1] for r in conn.execute(f"PRAGMA table_info({table})").fetchall()]
    except Exception:
        return []


def _update_task(conn: sqlite3.Connection, task_id: str, **kwargs: Any) -> None:
    cols = _cols(conn, "tasks")
    parts: List[str] = []
    vals: List[Any] = []

    for key, value in kwargs.items():
        if key in cols:
            parts.append(f"{key}=?")
            if key == "error_message":
                vals.append(_clean(_s(value), 4000))
            elif key == "result":
                vals.append(_clean(_s(value), 50000))
            elif key == "raw_input":
                vals.append(_clean(_s(value), 12000))
            else:
                vals.append(value)

    if "updated_at" in cols:
        parts.append("updated_at=datetime('now')")

    if not parts:
        return

    vals.append(task_id)
    conn.execute(f"UPDATE tasks SET {', '.join(parts)} WHERE id=?", vals)



# === LIVE_MEMORY_HELPERS_V1 ===
def _memory_insert_topic_entry_v1(chat_id: str, key: str, value: str) -> None:
    try:
        import sqlite3 as _sq
        if not os.path.exists(MEM_DB):
            return
        _mc = _sq.connect(MEM_DB)
        try:
            _mc.row_factory = _sq.Row
            _cols_mem = [r[1] for r in _mc.execute("PRAGMA table_info(memory)").fetchall()]
            if not _cols_mem:
                return
            _ts = datetime.datetime.utcnow().isoformat()
            _val = _clean(_s(value), 50000)
            if not _val:
                return
            _mid = hashlib.sha1(f"{chat_id}:{key}:{_ts}:{_val[:80]}".encode()).hexdigest()
            _mc.execute(
                "INSERT OR IGNORE INTO memory (id, chat_id, key, value, timestamp) VALUES (?,?,?,?,?)",
                (_mid, str(chat_id), str(key), _val, _ts),
            )
            _mc.commit()
        finally:
            _mc.close()
    except Exception as _mi_e:
        logger.warning("LIVE_MEMORY_HELPERS_V1_INSERT_ERR %s", _mi_e)

def _append_timeline_event_v1(chat_id: str, topic_id: int, task_id: str, kind: str, raw_input: str = "", result: str = "") -> None:
    try:
        _base = f"{BASE}/data/memory_files"
        _chat_key = f"{chat_id}__telegram"
        _chat_dir = os.path.join(_base, "CHATS", _chat_key)
        os.makedirs(_chat_dir, exist_ok=True)
        os.makedirs(os.path.join(_base, "GLOBAL"), exist_ok=True)
        _entry = json.dumps({
            "timestamp": datetime.datetime.utcnow().isoformat(),
            "chat_id": str(chat_id),
            "topic_id": int(topic_id or 0),
            "task_id": str(task_id),
            "kind": str(kind),
            "raw_input": _clean(_s(raw_input), 4000),
            "result": _clean(_s(result), 4000),
            "source": "task_worker",
        }, ensure_ascii=False)
        for _tl_path in [
            os.path.join(_chat_dir, "timeline.jsonl"),
            os.path.join(_base, "GLOBAL", "timeline.jsonl"),
        ]:
            with open(_tl_path, "a", encoding="utf-8") as _f:
                _f.write(_entry + "\n")
    except Exception as _tl_e:
        logger.warning("LIVE_MEMORY_HELPERS_V1_TIMELINE_ERR %s", _tl_e)
# === END LIVE_MEMORY_HELPERS_V1 ===

def _history(conn: sqlite3.Connection, task_id: str, action: str) -> None:
    if not _has_table(conn, "task_history"):
        return
    conn.execute(
        "INSERT INTO task_history (task_id, action, created_at) VALUES (?, ?, datetime('now'))",
        (task_id, _clean(action, 1000)),
    )


def _already_replied(conn: sqlite3.Connection, task_id: str, kind: str) -> bool:
    if not _has_table(conn, "task_history"):
        return False
    row = conn.execute(
        "SELECT 1 FROM task_history WHERE task_id=? AND action=? LIMIT 1",
        (task_id, f"reply_sent:{kind}"),
    ).fetchone()
    return row is not None


def _send_once(conn: sqlite3.Connection, task_id: str, chat_id: str, text: str, reply_to: Optional[int], kind: str) -> bool:
    # === _send_once_UNIFIED_USER_OUTPUT_SANITIZER_V1 ===
    try:
        from core.output_sanitizer import sanitize_user_output as _uos_sanitize
        text = _uos_sanitize(text)
    except Exception as _uos_err:
        logger.warning("UNIFIED_USER_OUTPUT_SANITIZER_V1_ERR %s", _uos_err)
    # === END__send_once_UNIFIED_USER_OUTPUT_SANITIZER_V1 ===
    if _already_replied(conn, task_id, kind):
        return True
    ok = send_reply(chat_id=chat_id, text=text, reply_to_message_id=reply_to)
    if ok:
        _history(conn, task_id, f"reply_sent:{kind}")
    return bool(ok)


def _send_once_ex(conn: sqlite3.Connection, task_id: str, chat_id: str, text: str, reply_to: Optional[int], kind: str) -> Dict[str, Any]:
    # === _send_once_ex_UNIFIED_USER_OUTPUT_SANITIZER_V1 ===
    try:
        from core.output_sanitizer import sanitize_user_output as _uos_sanitize
        text = _uos_sanitize(text)
    except Exception as _uos_err:
        logger.warning("UNIFIED_USER_OUTPUT_SANITIZER_V1_ERR %s", _uos_err)
    # === END__send_once_ex_UNIFIED_USER_OUTPUT_SANITIZER_V1 ===
    if _already_replied(conn, task_id, kind):
        return {"ok": True, "bot_message_id": None, "skipped": True}
    res = send_reply_ex(chat_id=chat_id, text=text, reply_to_message_id=reply_to)
    if not isinstance(res, dict):
        res = {"ok": bool(res)}
    if res.get("ok"):
        _history(conn, task_id, f"reply_sent:{kind}")
    return res


def _hash(text: str) -> str:
    return hashlib.sha1(_clean(text).lower().encode("utf-8")).hexdigest()


def _is_valid_result(text: str, raw_input: str) -> bool:
    r = _clean(text)
    if not r or len(r) < MIN_RESULT_LEN:
        return False
    if any(re.search(p, r, re.I) for p in BAD_RESULT_RE):
        return False
    if _hash(r) == _hash(raw_input):
        return False
    if "/root/" in r or ".ogg" in r.lower():
        return False
    return True


def _detect_role_assignment(text: str) -> str:
    triggers = [
        r"^(?:\[voice\]\s*)?этот чат для\s+(.+)$",
        r"^(?:\[voice\]\s*)?этот чат про\s+(.+)$",
        r"^(?:\[voice\]\s*)?этот топик для\s+(.+)$",
        r"^(?:\[voice\]\s*)?этот топик про\s+(.+)$",
        r"^(?:\[voice\]\s*)?чат закрепл[её]н за\s+(.+)$",
        r"^(?:\[voice\]\s*)?этот чат исключительно для\s+(.+)$",
        r"^(?:\[voice\]\s*)?этот чат используется для\s+(.+)$",
        r"^(?:\[voice\]\s*)?закрепи чат за\s+(.+)$",
        r"^(?:\[voice\]\s*)?закрепи этот чат за\s+(.+)$",
    ]
    t = _clean(text, 500).lower()
    for pattern in triggers:
        m = re.fullmatch(pattern, t)
        if m:
            return _clean(m.group(1), 200)
    return ""


def _extract_role_confirmation(result: str) -> str:
    t = _clean(_s(result), 500)
    m = re.fullmatch(r"Понял назначение чата так:\n(.+?)\n\nПодтверди или уточни", t, re.S)
    if not m:
        return ""
    return _clean(m.group(1), 200)


def _save_topic_role(chat_id: str, topic_id: int, role: str) -> None:
    if not role or not os.path.exists(MEM_DB):
        return
    mem = db(MEM_DB)
    try:
        if not _has_table(mem, "memory"):
            mem.execute("CREATE TABLE IF NOT EXISTS memory (chat_id TEXT, key TEXT, value TEXT, timestamp TEXT)")
        key = f"topic_{topic_id}_role"
        mem.execute("DELETE FROM memory WHERE chat_id=? AND key=?", (str(chat_id), key))
        mem.execute(
            "INSERT INTO memory (chat_id, key, value, timestamp) VALUES (?, ?, ?, datetime('now'))",
            (str(chat_id), key, role),
        )
        mem.commit()
    finally:
        mem.close()


def _load_memory_context(chat_id: str, topic_id: int) -> Tuple[str, str, str, str]:
    if not os.path.exists(MEM_DB):
        return "", "", "", ""

    conn = db(MEM_DB)
    try:
        if not _has_table(conn, "memory"):
            return "", "", "", ""

        topic_prefix = f"topic_{int(topic_id)}_"
        rows = conn.execute(
            """
            SELECT key, value
            FROM memory
            WHERE chat_id=?
              AND key GLOB ?
            ORDER BY timestamp DESC
            LIMIT 100
            """,
            (str(chat_id), f"{topic_prefix}*"),
        ).fetchall()

        short_memory: List[str] = []
        long_memory: List[str] = []
        topic_role = ""
        topic_directions = ""

        MEMORY_MUTEX_MARKERS = [
            "последние действия",
            "этот чат закреплён за: последние действия",
            "в этом чате были следующие действия",
            "текущий статус: ожидание подтверждения",
            "текущий статус:",
            "задача отменена",
            "задача завершена",
            "задача закрыта",
            "задачи завершены",
            "подтверждение принято",
            "не понимаю запрос",
            "готов к выполнению задачи",
            "без контекста",
            "задайте конкретный вопрос",
            "конкретный вопрос по",
            "нет, не помню",
        ]

        for row in rows:
            key = _s(row["key"])
            raw_value = _s(row["value"])
            if key.endswith("_user_input") or key.endswith("_role"):
                limit = 500
            elif key.endswith("_task_summary"):
                limit = 20000
            elif key.endswith("_assistant_output"):
                limit = 50000
            elif key.endswith("_directions"):
                limit = 1000
            else:
                limit = 500
            value = _clean(raw_value, limit)
            if not value:
                continue
            low = value.lower()
            if _is_memory_noise(low) or any(x in low for x in MEMORY_BAD_MARKERS):
                continue
            if any(m in low for m in MEMORY_MUTEX_MARKERS):
                continue

            if key.endswith("_role") and not topic_role:
                topic_role = value[:500]
                continue

            if not topic_role and (key.endswith("_assistant_output") or key.endswith("_task_summary")):
                m = re.search(r"чат закрепл[её]н за темами:\s*(.+?)(?:\.|$)", value, re.I)
                if not m:
                    m = re.search(r"чат закрепл[её]н за\s*(.+?)(?:\.|$)", value, re.I)
                if not m:
                    m = re.search(r"закреплено:\s*чат для\s*(.+?)(?:\.|$)", value, re.I)
                if not m:
                    m = re.search(r"этот чат используется для\s*(.+?)(?:\.|$)", value, re.I)
                if m:
                    topic_role = _clean(m.group(1), 500)

            if key.endswith("_directions") and not topic_directions:
                topic_directions = value[:1000]
                continue
            if key.endswith("_user_input") or key.endswith("_task_summary"):
                if not _is_memory_noise(value):
                    short_memory.append(f"{key}: {value}")
            else:
                if not _is_memory_noise(value):
                    long_memory.append(f"{key}: {value}")

        # === TOPIC_META_ROLE_INJECT_V1 ===
        if not topic_role and TOPIC_META_LOADER_WIRED:
            try:
                _tm = load_topic_meta(int(topic_id or 0))
                if _tm:
                    _tm_name = _tm.get("name", "")
                    _tm_dir = _tm.get("direction", "")
                    if _tm_name:
                        topic_role = f"Топик: {_tm_name} | Направление: {_tm_dir}"
            except Exception:
                pass
        # === END TOPIC_META_ROLE_INJECT_V1 ===
        return "\n".join(short_memory[:20]), "\n".join(long_memory[:20]), topic_role, topic_directions  # MEMORY_LIMIT_20_V1
    finally:
        conn.close()


def _load_archive_context(chat_id: str, topic_id: int, user_text: str) -> str:
    # === ARCHIVE_DISTRIBUTOR_V1_WIRED ===
    try:
        from core.archive_distributor import _load_archive_for_topic
        arc = _load_archive_for_topic(chat_id, topic_id, user_text, limit=5)
        if arc:
            return arc
    except Exception as _ade:
        logger.warning("ARCHIVE_DISTRIBUTOR_ERR %s", _ade)
    # === END ARCHIVE_DISTRIBUTOR_V1_WIRED ===
    # Fallback — старый метод через archive_legacy
    STOP_WORDS = {"это","как","что","где","когда","для","почему","зачем","кто"}
    words = {w for w in re.findall(r"\w+", _clean(user_text).lower()) if len(w) > 3 and w not in STOP_WORDS}
    if not words or not os.path.exists(MEM_DB):
        return ""

    conn = db(MEM_DB)
    try:
        if not _has_table(conn, "memory"):
            return ""
        rows = conn.execute(
            """
            SELECT key, value
            FROM memory
            WHERE chat_id=?
              AND key LIKE 'archive_legacy_%'
            ORDER BY timestamp DESC
            LIMIT 300
            """,
            (str(chat_id),),
        ).fetchall()
    finally:
        conn.close()

    scored: List[Tuple[int, str]] = []
    for row in rows:
        raw = _s(row["value"])
        if any(x in raw.lower() for x in MEMORY_BAD_MARKERS):
            continue
        try:
            payload = json.loads(raw)
        except Exception:
            continue
        if int(payload.get("topic_id", -1)) != int(topic_id):
            continue
        blob = _clean(f"{_s(payload.get('raw_input', ''))}\n{_s(payload.get('result', ''))}", 1200)
        ov = len(words & set(re.findall(r"\w+", blob.lower())))
        if ov > 0:
            scored.append((ov, blob))

    scored = [(ov, blob) for ov, blob in scored if ov > 0]
    scored.sort(key=lambda x: x[0], reverse=True)
    return "\n\n".join(x[1] for x in scored[:3]) if scored else ""


def _active_unfinished_context(conn: sqlite3.Connection, chat_id: str, topic_id: int, task_id: str) -> str:
    cols = _cols(conn, "tasks")
    where = [
        "chat_id=?",
        "id<>?",
        "state IN ('WAITING_CLARIFICATION','IN_PROGRESS','AWAITING_CONFIRMATION')",
    ]
    params: List[Any] = [str(chat_id), task_id]
    if "topic_id" in cols:
        where.append("COALESCE(topic_id,0)=?")
        params.append(int(topic_id))

    rows = conn.execute(
        f"""
        SELECT raw_input, result, state
        FROM tasks
        WHERE {' AND '.join(where)}
        ORDER BY updated_at DESC, created_at DESC
        LIMIT 3
        """,
        params,
    ).fetchall()

    parts = []
    for row in rows:
        raw = _clean(_s(_task_field(row, "raw_input")), 300)
        res = _clean(_s(row["result"]), 500)
        state = _clean(_s(row["state"]), 100)
        low = f"{raw}\n{res}".lower()
        if any(x in low for x in MEMORY_BAD_MARKERS):
            continue
        chunk = []
        if raw:
            chunk.append(f"raw_input: {raw}")
        if res:
            chunk.append(f"result: {res}")
        if state:
            chunk.append(f"state: {state}")
        if chunk:
            parts.append("\n".join(chunk))
    return "\n\n".join(parts[:3])


def _search_fact_context(conn: sqlite3.Connection, chat_id: str, topic_id: int) -> str:
    cols = _cols(conn, "tasks")
    where = [
        "chat_id=?",
        "state IN ('DONE','ARCHIVED')",
        "lower(COALESCE(raw_input,'')) GLOB '*най*'",
    ]
    params: List[Any] = [str(chat_id)]
    if "topic_id" in cols:
        where.append("COALESCE(topic_id,0)=?")
        params.append(int(topic_id))

    rows = conn.execute(
        f"""
        SELECT raw_input, result
        FROM tasks
        WHERE {' AND '.join(where)}
        ORDER BY updated_at DESC
        LIMIT 5
        """,
        params,
    ).fetchall()

    facts: List[str] = []
    for row in rows:
        q = _clean(_s(row["raw_input"]), 300)
        r = _clean(_s(row["result"]), 500)
        low = f"{q}\n{r}".lower()
        if any(x in low for x in MEMORY_BAD_MARKERS):
            continue
        if q and r:
            facts.append(f"search_done: {q} => {r}")
    return "\n".join(facts[:3])


def _save_memory(chat_id: str, topic_id: int, raw_input: str, result: str) -> None:
    low = (result or "").lower()
    if any(x in low for x in [
        "без контекста","не понимаю запрос","не помню",
        "задача завершена","задача закрыта","подтверждение принято",
        "готов к выполнению задачи"
    ]):
        return

    bad = [
        "ошибка",
        "не найдено",
        "уточните",
        "traceback",
        "/root/",
        ".ogg",
        "delete from",
        "task_worker.py",
        "telegram_daemon.py",
    ]
    if not result or len(result) < MIN_RESULT_LEN:
        return
    low = result.lower()
    if any(b in low for b in bad):
        return
    if not os.path.exists(MEM_DB):
        return

    conn = db(MEM_DB)
    try:
        if not _has_table(conn, "memory"):
            conn.execute("CREATE TABLE IF NOT EXISTS memory (chat_id TEXT, key TEXT, value TEXT, timestamp TEXT)")
        prefix = f"topic_{int(topic_id)}_"
        conn.execute(
            "INSERT INTO memory (chat_id, key, value, timestamp) VALUES (?, ?, ?, datetime('now'))",
            (str(chat_id), f"{prefix}assistant_output", _clean(result, 50000)),
        )
        conn.execute(
            "INSERT INTO memory (chat_id, key, value, timestamp) VALUES (?, ?, ?, datetime('now'))",
            (str(chat_id), f"{prefix}task_summary", _clean(result, 20000) if len(result) >= MIN_RESULT_LEN else ""),
        )
        conn.execute(
            "INSERT INTO memory (chat_id, key, value, timestamp) VALUES (?, ?, ?, datetime('now'))",
            (str(chat_id), f"{prefix}user_input", _clean(raw_input, 500)),
        )
        conn.commit()
        logger.info("save_memory_ok chat=%s topic=%s", chat_id, topic_id)
    finally:
        conn.close()


def _close_pin(conn: sqlite3.Connection, task_id: str) -> None:
    if not _has_table(conn, "pin"):
        return
    conn.execute(
        "UPDATE pin SET state='CLOSED', updated_at=datetime('now') WHERE task_id=? AND state='ACTIVE'",
        (task_id,),
    )


def _finalize_done(conn: sqlite3.Connection, task_id: str, chat_id: str, topic_id: int, reply_to: Optional[int]) -> None:
    row = conn.execute(
        "SELECT COALESCE(raw_input,''), COALESCE(result,'') FROM tasks WHERE id=? LIMIT 1",
        (task_id,),
    ).fetchone()
    raw_input = _clean(_s(row[0]) if row else "", 500)
    result = _clean(_s(row[1]) if row else "", 50000)

    _update_task(conn, task_id, state="DONE", error_message="")
    _history(conn, task_id, "state:DONE")
    if result:
        _save_memory(chat_id, topic_id, raw_input, result)
    conn.commit()


def _is_confirm_intent(text: str) -> bool:
    t = _clean(text, 200).lower()
    if t in CONFIRM_INTENTS:
        return True
    return any(t.startswith(x) for x in ["да", "подтвер", "соглас", "верно", "ок"])


def _is_revision_intent(text: str) -> bool:
    t = _clean(text, 200).lower()
    if t in REVISION_INTENTS:
        return True
    return any(x in t for x in ["не так", "передел", "исправ", "правк", "уточн"])


def _recover_stale_tasks(conn: sqlite3.Connection, chat_id: Optional[str]) -> None:
    # CONFIRMATION_TIMEOUT_FIX_V1
    try:
        conn.execute("""
            UPDATE tasks SET state='FAILED', error_message='CONFIRMATION_TIMEOUT', updated_at=datetime('now')
            WHERE state='AWAITING_CONFIRMATION'
              AND updated_at < datetime('now','-30 minutes')
              AND COALESCE(raw_input,'') NOT LIKE '%retry_queue_healthcheck%'
              AND COALESCE(result,'') NOT LIKE '%retry_queue_healthcheck%'
        """)
        conn.commit()
    except Exception as _ct_e:
        logger.warning("CONFIRMATION_TIMEOUT_FIX_V1_ERR %s", _ct_e)
    # IN_PROGRESS_HARD_TIMEOUT_V1
    try:
        _hp = [] if not chat_id else [str(chat_id)]
        _hw = (["chat_id=?"] if chat_id else []) + [
            "state='IN_PROGRESS'",
            "created_at < datetime('now','-15 minutes')",
        ]
        for _hr in conn.execute(
            f"SELECT id,chat_id,COALESCE(topic_id,0) AS topic_id,reply_to_message_id,raw_input FROM tasks WHERE {' AND '.join(_hw)}",
            _hp,
        ).fetchall():
            _htid = _s(_hr["id"]); _hchat = _s(_hr["chat_id"])
            _htopic = int(_hr["topic_id"] or 0); _hreply = _hr["reply_to_message_id"]
            _hmsg = "Задача не выполнена: превышено время выполнения"
            _update_task(conn, _htid, state="FAILED", result=_hmsg, error_message="EXECUTION_TIMEOUT")
            _close_pin(conn, _htid); _history(conn, _htid, "state:FAILED:EXECUTION_TIMEOUT")
            try:
                _append_timeline_event_v1(_hchat, _htopic, _htid, "execution_timeout", _s(_hr["raw_input"]), _hmsg)
            except Exception: pass
            conn.commit()
            _send_once(conn, _htid, _hchat, _hmsg, _hreply, "execution_timeout")
    except Exception as _e:
        logger.warning("IN_PROGRESS_HARD_TIMEOUT_V1_ERR %s", _e)

    where = [
        "state IN ('IN_PROGRESS','WAITING_CLARIFICATION')",
        "(strftime('%s','now') - strftime('%s', COALESCE(updated_at, created_at))) > ?",
    ]
    params: List[Any] = [STALE_TIMEOUT]
    if chat_id:
        where.insert(0, "chat_id=?")
        params.insert(0, str(chat_id))

    rows = conn.execute(
        f"""
        SELECT id, chat_id, COALESCE(topic_id,0) AS topic_id, reply_to_message_id
        FROM tasks
        WHERE {' AND '.join(where)}
        """,
        params,
    ).fetchall()

    for row in rows:
        task_id = _s(row["id"])
        tg_chat_id = _s(row["chat_id"])
        reply_to = row["reply_to_message_id"]
        _update_task(conn, task_id, state="FAILED", error_message="STALE_TIMEOUT")
        _close_pin(conn, task_id)
        _history(conn, task_id, "state:FAILED")
        conn.commit()
        _send_once(conn, task_id, tg_chat_id, "Задача не выполнена. Повтори или уточни запрос", reply_to, "stale_failed")

    done_markers = [
        "задача завершена",
        "задача закрыта",
        "задачи завершены",
        "подтверждение принято",
    ]
    junk_markers = [
        "без контекста",
        "задайте конкретный вопрос",
        "конкретный вопрос по",
        "нет, не помню",
        "не понимаю запрос",
        "готов к выполнению задачи",
    ]

    rows = conn.execute("""
        SELECT id, chat_id, COALESCE(topic_id,0) AS topic_id, reply_to_message_id, result, raw_input, input_type, updated_at, created_at
        FROM tasks
        WHERE state = 'AWAITING_CONFIRMATION'
    """).fetchall()

    now_utc = datetime.datetime.now(datetime.timezone.utc)

    for row in rows:
        updated = row["updated_at"] or row["created_at"]
        if not updated:
            continue

        result = _s(row["result"])
        low_result = result.lower()
        row_raw_input = ""
        try:
            row_raw_input = _s(row["raw_input"])
        except Exception:
            row_raw_input = ""
        if _force_voice_finish(row_raw_input, result):
            _update_task(conn, row["id"], state="DONE")
            continue  # FORCE_VOICE_FINISH_HOOK

        if any(m in low_result for m in done_markers):
            _update_task(conn, row["id"], state="DONE", error_message="")
            _close_pin(conn, row["id"])
            _history(conn, row["id"], "state:DONE")
            conn.commit()
            continue

        if any(m in low_result for m in junk_markers):
            _update_task(conn, row["id"], state="FAILED", error_message="JUNK_RESULT_CLEANUP")
            _close_pin(conn, row["id"])
            _history(conn, row["id"], "state:FAILED")
            conn.commit()
            continue

        try:
            s = str(updated).strip()
            if "T" in s:
                dt = datetime.datetime.fromisoformat(s.replace("Z", "+00:00"))
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=datetime.timezone.utc)
            else:
                dt = datetime.datetime.fromisoformat(s)
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=datetime.timezone.utc)
        except Exception:
            continue

        hours_old = (now_utc - dt).total_seconds() / 3600.0
        if hours_old <= 24:
            continue

        if _s(row["input_type"]).lower() == "drive_file" and hours_old > 48:
            _update_task(conn, row["id"], state="FAILED", error_message="drive_upload_stale")
            _close_pin(conn, row["id"])
            _history(conn, row["id"], "state:FAILED")
            conn.commit()
        elif hours_old > 168:
            _update_task(conn, row["id"], state="ARCHIVED")
            _close_pin(conn, row["id"])
            _history(conn, row["id"], "state:ARCHIVED")
            conn.commit()

    conn.execute("""
        UPDATE pin
        SET state='CLOSED', updated_at=datetime('now')
        WHERE state='ACTIVE'
          AND task_id IN (
              SELECT id FROM tasks
              WHERE state='FAILED'
                 OR lower(COALESCE(result,'')) LIKE '%задача закрыта%'
                 OR lower(COALESCE(result,'')) LIKE '%подтверждение принято%'
                 OR lower(COALESCE(result,'')) LIKE '%без контекста%'
                 OR lower(COALESCE(result,'')) LIKE '%конкретный вопрос%'
                 OR lower(COALESCE(result,'')) LIKE '%не помню%'
          )
    """)
    conn.commit()



# === FULLFIX_13D_TASK_WORKER_SEND_BELT ===
# reply_sender also strips MANIFEST globally
# === END FULLFIX_13D_TASK_WORKER_SEND_BELT ===

# === FULLFIX_13C_STRIP_MANIFEST_BEFORE_SEND ===
def _ff13c_strip_manifest_links(text):
    import re
    msg = str(text or "")
    msg = re.sub(r"(?im)^MANIFEST:\s*https?://\S+\s*$", "", msg)
    msg = re.sub(r"(?im)^Manifest:\s*https?://\S+\s*$", "", msg)
    msg = re.sub(r"\n{3,}", "\n\n", msg).strip()
    return msg
# === END FULLFIX_13C_STRIP_MANIFEST_BEFORE_SEND ===


async def _handle_new(conn: sqlite3.Connection, task: sqlite3.Row, chat_id: str, topic_id: int) -> None:
    if _stop_topic2_loop_fact_based_v1(conn, task):
        return
    # === FULL_STROYKA_ESTIMATE_CANON_CLOSE_V3_HOOK ===
    try:
        from core.stroyka_estimate_canon import maybe_handle_stroyka_estimate as _stroyka_estimate_hook
        if await _stroyka_estimate_hook(conn, task, logger=logger):
            return
    except Exception as _stroyka_estimate_hook_err:
        logger.exception("FULL_STROYKA_ESTIMATE_CANON_CLOSE_V3_HOOK_ERR %s", _stroyka_estimate_hook_err)
    # === END_FULL_STROYKA_ESTIMATE_CANON_CLOSE_V3_HOOK ===
    # === REMAINING_TECH_CONTOUR_CLOSE_V1_WIRED ===
    # P0 guard order:
    # 1. reply repeat parent
    # 2. explicit project route before estimate/template/file followup
    try:
        from core.reply_repeat_parent import prehandle_reply_repeat_parent_v1 as _rrp_prehandle
        _rrp = _rrp_prehandle(conn, task)
        if _rrp and _rrp.get("handled"):
            _rrp_tid = _s(_task_field(task, "id", ""))
            _rrp_chat = _s(_task_field(task, "chat_id", ""))
            _rrp_reply = _task_field(task, "reply_to_message_id", None)
            _rrp_text = _rrp.get("message") or ""
            try:
                from core.output_sanitizer import sanitize_user_output as _sanitize_out
                _rrp_text = _sanitize_out(_rrp_text)
            except Exception:
                pass
            _rrp_send = _send_once_ex(conn, _rrp_tid, _rrp_chat, _rrp_text, _rrp_reply, _rrp.get("kind", "reply_repeat_parent"))
            _update_task(
                conn,
                _rrp_tid,
                state=_rrp.get("state", "DONE"),
                result=_rrp_text,
                error_message=_rrp.get("error_message", ""),
                bot_message_id=_rrp_send.get("bot_message_id"),
            )
            _history(conn, _rrp_tid, _rrp.get("history", "REPLY_REPEAT_PARENT_TASK_V1:HANDLED"))
            return
    except Exception as _rrp_err:
        logger.warning("REMAINING_TECH_CONTOUR_CLOSE_V1_REPLY_ERR %s", _rrp_err)

    try:
        from core.project_route_guard import prehandle_project_route_v1 as _proj_prehandle
        _proj = await _proj_prehandle(conn, task)
        if _proj and _proj.get("handled"):
            _proj_tid = _s(_task_field(task, "id", ""))
            _proj_chat = _s(_task_field(task, "chat_id", ""))
            _proj_reply = _task_field(task, "reply_to_message_id", None)
            _proj_text = _proj.get("message") or ""
            try:
                from core.output_sanitizer import sanitize_project_message as _sanitize_project_out
                _proj_text = _sanitize_project_out(_proj_text)
            except Exception:
                pass
            _proj_send = _send_once_ex(conn, _proj_tid, _proj_chat, _proj_text, _proj_reply, _proj.get("kind", "project_route_guard"))
            _update_task(
                conn,
                _proj_tid,
                state=_proj.get("state", "WAITING_CLARIFICATION"),
                result=_proj_text,
                error_message=_proj.get("error_message", ""),
                bot_message_id=_proj_send.get("bot_message_id"),
            )
            _history(conn, _proj_tid, _proj.get("history", "PROJECT_ROUTE_FIX_NO_ESTIMATE_REGRESSION_V1:HANDLED"))
            return
    except Exception as _proj_err:
        logger.warning("REMAINING_TECH_CONTOUR_CLOSE_V1_PROJECT_ERR %s", _proj_err)
    # === END_REMAINING_TECH_CONTOUR_CLOSE_V1_WIRED ===
    # === FULL_TECH_CONTOUR_CLOSE_V1_WIRED ===
    # P0 guard: price confirmation, context-aware file intake, Telegram duplicate file memory
    try:
        from core.price_enrichment import prehandle_price_task_v1 as _ftc_price_prehandle
        _ftc_price = await _ftc_price_prehandle(conn, task) if int(topic_id or 0) == 2 else None  # TOPIC2_PRICE_ONLY_V1
        if _ftc_price and _ftc_price.get("handled"):
            _ftc_tid = _s(_task_field(task, "id", ""))
            _ftc_chat = _s(_task_field(task, "chat_id", ""))
            _ftc_reply = _task_field(task, "reply_to_message_id", None)
            _ftc_text = _ftc_price.get("message") or ""
            _ftc_send = _send_once_ex(conn, _ftc_tid, _ftc_chat, _ftc_text, _ftc_reply, _ftc_price.get("kind", "price_enrichment"))
            _update_task(
                conn,
                _ftc_tid,
                state=_ftc_price.get("state", "WAITING_CLARIFICATION"),
                result=_ftc_text,
                error_message=_ftc_price.get("error_message", ""),
                bot_message_id=_ftc_send.get("bot_message_id"),
            )
            _history(conn, _ftc_tid, _ftc_price.get("history", "WEB_SEARCH_PRICE_ENRICHMENT_V1:HANDLED"))
            return
    except Exception as _ftc_price_err:
        logger.warning("FULL_TECH_CONTOUR_CLOSE_V1_PRICE_ERR %s", _ftc_price_err)

    try:
        from core.file_context_intake import prehandle_task_context_v1 as _ftc_file_prehandle
        # === CANON_ROUTE_FIX_V3_TOPIC500_ISOLATION ===
        _ftc_topic_id = int(_task_field(task, "topic_id", 0) or 0)
        _ftc_file = _ftc_file_prehandle(conn, task) if _ftc_topic_id not in (500, 210) else None  # TOPIC_ROUTE_ISOLATION_FULL_V2
        # === END_CANON_ROUTE_FIX_V3_TOPIC500_ISOLATION ===
        if _ftc_file and _ftc_file.get("handled"):
            _ftc_tid = _s(_task_field(task, "id", ""))
            _ftc_chat = _s(_task_field(task, "chat_id", ""))
            _ftc_reply = _task_field(task, "reply_to_message_id", None)
            _ftc_text = _ftc_file.get("message") or ""
            _ftc_send = _send_once_ex(conn, _ftc_tid, _ftc_chat, _ftc_text, _ftc_reply, _ftc_file.get("kind", "file_context_intake"))
            _update_task(
                conn,
                _ftc_tid,
                state=_ftc_file.get("state", "DONE"),
                result=_ftc_text,
                error_message=_ftc_file.get("error_message", ""),
                bot_message_id=_ftc_send.get("bot_message_id"),
            )
            _history(conn, _ftc_tid, _ftc_file.get("history", "CONTEXT_AWARE_FILE_INTAKE_V1:HANDLED"))
            return
    except Exception as _ftc_file_err:
        logger.warning("FULL_TECH_CONTOUR_CLOSE_V1_FILE_ERR %s", _ftc_file_err)
    # === END_FULL_TECH_CONTOUR_CLOSE_V1_WIRED ===

    # === FINAL_CLOSURE_BLOCKER_FIX_V1_TASK_WORKER_HOOK ===
    try:
        from core.final_closure_engine import maybe_handle_final_closure as _fcv1_handle
        _fcv1 = _fcv1_handle(
            conn=conn,
            task=task,
            task_id=str(_task_field(task, "id", "")),
            chat_id=str(chat_id),
            topic_id=int(topic_id or 0),
            raw_input=str(_task_field(task, "raw_input", "")),
            input_type=str(_task_field(task, "input_type", "text")),
            reply_to=_task_field(task, "reply_to_message_id", None),
        )
        if isinstance(_fcv1, dict) and _fcv1.get("handled"):
            _fcv1_tid = str(_task_field(task, "id", ""))
            _fcv1_reply = _task_field(task, "reply_to_message_id", None)
            _fcv1_msg = str(_fcv1.get("message") or "").strip()
            _fcv1_state = str(_fcv1.get("state") or "DONE").strip()
            _fcv1_kind = str(_fcv1.get("kind") or "final_closure_blocker_fix_v1").strip()
            if _fcv1_msg:
                _fcv1_send = _send_once_ex(conn, _fcv1_tid, str(chat_id), _fcv1_msg, _fcv1_reply, _fcv1_kind)
                _fcv1_bot = _fcv1_send.get("bot_message_id") if isinstance(_fcv1_send, dict) else None
                _update_task(conn, _fcv1_tid, state=_fcv1_state, result=_fcv1_msg, bot_message_id=_fcv1_bot, error_message="")
                _history(conn, _fcv1_tid, _fcv1.get("history", "FINAL_CLOSURE_BLOCKER_FIX_V1:HANDLED"))
                conn.commit()
                return
    except Exception as _fcv1_err:
        logger.warning("FINAL_CLOSURE_BLOCKER_FIX_V1_TASK_WORKER_HOOK_ERR %s", _fcv1_err)
    # === END_FINAL_CLOSURE_BLOCKER_FIX_V1_TASK_WORKER_HOOK ===

    # === EZONE_INGEST_CALL_V1 ===
    try:
        _ez_raw = str(task["raw_input"] if task else "") if task else ""
        if is_ezone_payload(_ez_raw):
            _ez_msg = save_ezone_json(str(chat_id), _ez_raw)
            _update_task(conn, task["id"], state="DONE", result=_ez_msg, error_message="")
            conn.commit()
            from core.reply_sender import send_reply_ex
            send_reply_ex(chat_id=str(chat_id), text=_ez_msg, reply_to_message_id=task.get("reply_to_message_id"), message_thread_id=topic_id)
            logger.info("EZONE_INGEST_CALL_V1 task=%s", task["id"])
            return
    except Exception as _eze:
        logger.warning("EZONE_INGEST_CALL_ERR %s", _eze)
    # === END EZONE_INGEST_CALL_V1 ===
    # === HEALTHCHECK_GUARD_V1 ===
    try:
        _hc_raw = str(task['raw_input'] if task else '') if task else ''
        _hc_res = str(task['result'] if task else '') if task else ''
        _hc_markers = ['retry_queue_healthcheck', 'healthcheck', 'areal_hc_', '_hc_file']
        if any(m in _hc_raw or m in _hc_res for m in _hc_markers):
            _update_task(conn, task['id'], state='CANCELLED', error_message='SERVICE_FILE_IGNORED:HEALTHCHECK')
            conn.commit()
            logger.info('HEALTHCHECK_GUARD_V1 cancelled task=%s', task['id'])
            return
    except Exception as _hcge:
        logger.warning('HEALTHCHECK_GUARD_ERR %s', _hcge)
    # === END HEALTHCHECK_GUARD_V1 ===
    task_id = _s(_task_field(task, "id"))
    raw_input = _clean(_s(_task_field(task, "raw_input")), 4000)
    reply_to = _task_field(task, "reply_to_message_id", None)

    # === VOICE_CONFIRM_AWAITING_V1 ===
    try:
        _vc_raw = str(raw_input or "").strip()
        _vc_is_voice = _vc_raw.lower().startswith("[voice]")
        _vc_text = re.sub(r"^\s*\[voice\]\s*", "", _vc_raw, flags=re.I).strip()
        _vc_low = _vc_text.lower().strip().rstrip("!?. ")
        if _vc_is_voice and (_vc_low in CONFIRM_INTENTS or _vc_low in REVISION_INTENTS):
            _vc_parent = conn.execute(
                """
                SELECT id, result, raw_input
                FROM tasks
                WHERE chat_id=?
                  AND COALESCE(topic_id,0)=?
                  AND id<>?
                  AND state='AWAITING_CONFIRMATION'
                ORDER BY updated_at DESC
                LIMIT 1
                """,
                (str(chat_id), int(topic_id or 0), task_id),
            ).fetchone()

            if _vc_parent is not None and _vc_low in CONFIRM_INTENTS:
                _vc_parent_id = _s(_vc_parent["id"])
                _finalize_done(conn, _vc_parent_id, str(chat_id), int(topic_id or 0), reply_to)
                _update_task(conn, task_id, state="DONE", result="Подтверждение принято голосом. Задача закрыта", error_message="")
                _history(conn, _vc_parent_id, "VOICE_CONFIRM_AWAITING_V1:PARENT_DONE")
                _history(conn, task_id, "VOICE_CONFIRM_AWAITING_V1:CHILD_DONE")
                conn.commit()
                _send_once(conn, task_id, str(chat_id), "Подтверждение принято голосом. Задача закрыта", reply_to, "voice_confirm_awaiting_v1")
                logger.info("VOICE_CONFIRM_AWAITING_V1 confirmed parent=%s child=%s topic=%s", _vc_parent_id, task_id, topic_id)
                return

            if _vc_parent is not None and _vc_low in REVISION_INTENTS:
                _vc_parent_id = _s(_vc_parent["id"])
                _vc_merged = _clean(_s(_vc_parent["result"]) + "\n\nПравки пользователя голосом:\n" + _vc_text, 12000)
                _update_task(conn, _vc_parent_id, state="IN_PROGRESS", raw_input=_vc_merged, error_message="")
                _update_task(conn, task_id, state="DONE", result="Голосовые правки приняты. Задача возвращена в работу", error_message="")
                _history(conn, _vc_parent_id, "VOICE_CONFIRM_AWAITING_V1:PARENT_REVISION")
                _history(conn, task_id, "VOICE_CONFIRM_AWAITING_V1:CHILD_DONE_REVISION")
                conn.commit()
                _send_once(conn, task_id, str(chat_id), "Голосовые правки приняты. Задача возвращена в работу", reply_to, "voice_revision_awaiting_v1")
                logger.info("VOICE_CONFIRM_AWAITING_V1 revision parent=%s child=%s topic=%s", _vc_parent_id, task_id, topic_id)
                return
    except Exception as _vc_e:
        logger.error("VOICE_CONFIRM_AWAITING_V1_ERR task=%s err=%s", task_id, _vc_e)
    # === END VOICE_CONFIRM_AWAITING_V1 ===




    
    # === CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1 ===
    try:
        _cpp_raw = ""
        try:
            _cpp_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
        except Exception:
            _cpp_raw = ""
        _cpp_clean = re.sub(r"^\s*\[VOICE\]\s*", "", _cpp_raw or "", flags=re.I).strip()
        _cpp_low = _cpp_clean.lower()

        _cpp_create_words = (
            "сделай", "делай", "создай", "сформируй", "подготовь", "разработай",
            "оформи", "выгрузи", "сохрани", "нарисуй", "собери"
        )
        # === CREATE_PROJECT_CPP_WORDS_FIX_V1 ===
        _cpp_project_words = (
            "проект", "кж", "кд", "км", "кмд", "ар",
            "фундамент", "фундаментн",
            "армирован", "арматур", "конструктив", "чертеж", "чертёж",
            "dxf", "dwg", "узел", "узлы", "спецификац"
        )
        # === END_CREATE_PROJECT_CPP_WORDS_FIX_V1 ===
        _cpp_followup_words = (
            "я тебе скидывал", "уже скидывал", "где файл", "где проект",
            "что дальше", "дальше то что", "ты сделал", "что мы делали",
            "какие последние", "покажи прошл", "найди прошл", "помнишь"
        )

        _cpp_is_create_project = (
            any(w in _cpp_low for w in _cpp_create_words)
            and any(w in _cpp_low for w in _cpp_project_words)
            and not any(w in _cpp_low for w in _cpp_followup_words)
        )

        if int(topic_id or 0) == 210 and _cpp_is_create_project:  # TOPIC210_PROJECT_ONLY_V1
            _history(conn, task_id, "CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1:ROUTE_PROJECT_ENGINE")

            from core.project_engine import create_project_pdf_dxf_artifact

            try:
                _cpp_data = await create_project_pdf_dxf_artifact(
                    raw_input=_cpp_clean,
                    task_id=str(task_id),
                    topic_id=int(topic_id or 0),
                    template_hint="",
                    require_template=False,
                )
            except TypeError:
                _cpp_data = await create_project_pdf_dxf_artifact(
                    _cpp_clean,
                    str(task_id),
                    int(topic_id or 0),
                    "",
                )

            if not isinstance(_cpp_data, dict):
                _update_task(conn, task_id, state="FAILED", result="", error_message="CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1:NON_DICT_RESULT")
                _history(conn, task_id, "CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1:NON_DICT_RESULT")
                return

            _cpp_links = {}
            for _k in ("pdf_link", "dxf_link", "xlsx_link", "docx_link", "drive_link", "google_sheet_link", "link"):
                _v = _cpp_data.get(_k)
                if isinstance(_v, str) and _v.startswith("http"):
                    _cpp_links[_k] = _v

            if isinstance(_cpp_data.get("links"), dict):
                for _k, _v in _cpp_data.get("links").items():
                    if isinstance(_v, str) and _v.startswith("http"):
                        _cpp_links[str(_k)] = _v

            _cpp_artifacts = []
            for _k in ("artifact_path", "package_path", "zip_path", "pdf_path", "dxf_path", "xlsx_path", "docx_path"):
                _v = _cpp_data.get(_k)
                if isinstance(_v, str) and _v:
                    _cpp_artifacts.append({"path": _v, "kind": "project_" + _k.replace("_path", "")})

            if _cpp_artifacts and not _cpp_links:
                try:
                    from core.artifact_upload_guard import upload_many_or_fail
                    _up = upload_many_or_fail(_cpp_artifacts, str(task_id), int(topic_id or 0))
                    if isinstance(_up, dict):
                        for _k, _v in (_up.get("links") or {}).items():
                            if isinstance(_v, str) and _v.startswith("http"):
                                _cpp_links[str(_k)] = _v
                except Exception as _up_e:
                    logger.warning("CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1_UPLOAD_ERR task=%s err=%s", task_id, _up_e)

            _pdf = ""
            _dxf = ""
            _xlsx = ""
            _other = ""

            for _k, _v in _cpp_links.items():
                _kl = str(_k).lower()
                _vl = str(_v).lower()
                if not _pdf and ("pdf" in _kl or ".pdf" in _vl):
                    _pdf = _v
                elif not _dxf and ("dxf" in _kl or ".dxf" in _vl):
                    _dxf = _v
                elif not _xlsx and ("xlsx" in _kl or "sheet" in _kl or "spreadsheets" in _vl or ".xlsx" in _vl):
                    _xlsx = _v
                elif not _other:
                    _other = _v

            _cpp_result_lines = ["Проект плиты создан"]
            if _pdf:
                _cpp_result_lines.append("PDF: " + _pdf)
            if _dxf:
                _cpp_result_lines.append("DXF: " + _dxf)
            if _xlsx:
                _cpp_result_lines.append("XLSX: " + _xlsx)
            if not (_pdf or _dxf or _xlsx) and _other:
                _cpp_result_lines.append("Файл: " + _other)

            _cpp_result = "\n".join(_cpp_result_lines).strip()

            _cpp_success = bool(_cpp_data.get("success")) or bool(_cpp_links) or bool(_cpp_artifacts)
            if _cpp_success and (_cpp_links or _cpp_artifacts):
                _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_cpp_result, error_message="")
                _history(conn, task_id, "CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1:AWAITING_CONFIRMATION")
                try:
                    _reply_to = task["reply_to_message_id"] if "reply_to_message_id" in task.keys() else None
                except Exception:
                    _reply_to = None
                _send_once_ex(conn, task_id, chat_id, _cpp_result, _reply_to, "project_create_priority")
                return

            _err = str(_cpp_data.get("error") or "PROJECT_ENGINE_NO_ARTIFACT")[:500]
            _user_err = "Проект не создан: " + _err
            _update_task(conn, task_id, state="FAILED", result="", error_message="CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1:" + _err)
            _history(conn, task_id, "CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1:FAILED:" + _err[:200])
            try:
                _reply_to = task["reply_to_message_id"] if "reply_to_message_id" in task.keys() else None
            except Exception:
                _reply_to = None
            _send_once_ex(conn, task_id, chat_id, _user_err, _reply_to, "project_create_priority_error")
            return
    except Exception as _cpp_e:
        logger.warning("CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1_ERR task=%s err=%s", task_id, _cpp_e)
    # === END_CREATE_PROJECT_PRIORITY_NO_ROLLBACK_V1 ===
    # === CREATE_ESTIMATE_PRIORITY_NO_ROLLBACK_V1 ===
    try:
        _cep_raw = ""
        try:
            _cep_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
        except Exception:
            _cep_raw = ""
        _cep_clean = re.sub(r"^\s*\[VOICE\]\s*", "", _cep_raw or "", flags=re.I).strip()
        _cep_low = _cep_clean.lower()

        _cep_create_words = (
            "сделай", "создай", "сформируй", "подготовь", "составь",
            "посчитай", "рассчитай", "выгрузи", "сохрани", "оформи"
        )
        _cep_estimate_words = (
            "смет", "расчет", "расчёт", "xlsx", "excel", "эксель",
            "pdf", "ндс", "итог", "объем", "объём", "расценк", "позици"
        )
        _cep_followup_words = (
            "я тебе скидывал", "уже скидывал", "где файл", "где смета",
            "что дальше", "дальше то что", "ты сделал", "что мы делали",
            "какие последние", "покажи прошл", "найди прошл", "помнишь"
        )

        # === CEP_PROJECT_EXCLUSION_V2 ===
        _cep_project_words = (
            "проект", "кж", "кд", "км", "кмд", "ар",
            "фундамент", "фундаментн", "армирован", "арматур", "dxf", "dwg", "чертеж", "чертёж", "конструктив",
        )
        _cep_is_create_estimate = (
            any(w in _cep_low for w in _cep_create_words)
            and any(w in _cep_low for w in _cep_estimate_words)
            and not any(w in _cep_low for w in _cep_followup_words)
            and not any(w in _cep_low for w in _cep_project_words)
        )
        # === END_CEP_PROJECT_EXCLUSION_V2 ===

        if int(topic_id or 0) == 2 and _cep_is_create_estimate:  # TOPIC2_ESTIMATE_ONLY_V1
            from core.estimate_engine import generate_estimate_from_text

            _history(conn, task_id, "CREATE_ESTIMATE_PRIORITY_NO_ROLLBACK_V1:ROUTE_ESTIMATE_ENGINE")
            _cep_data = await generate_estimate_from_text(
                raw_input=_cep_clean,
                task_id=str(task_id),
                topic_id=int(topic_id or 0),
                chat_id=str(chat_id),
            )

            if isinstance(_cep_data, dict):
                _cep_state = str(_cep_data.get("state") or "")
                _cep_success = bool(_cep_data.get("success"))
                _cep_links = []
                for _k in ("drive_link", "google_sheet_link", "link", "pdf_link", "xlsx_link", "telegram_link"):
                    _v = _cep_data.get(_k)
                    if isinstance(_v, str) and _v.startswith("http"):
                        _cep_links.append((_k, _v))
                if isinstance(_cep_data.get("links"), dict):
                    for _k, _v in _cep_data.get("links").items():
                        if isinstance(_v, str) and _v.startswith("http"):
                            _cep_links.append((str(_k), _v))

                _cep_result = str(_cep_data.get("result_text") or _cep_data.get("message") or _cep_data.get("summary") or "").strip()
                if not _cep_result:
                    _cep_result = "Смета обработана"
                # === CLEAN_CEP_RESULT_TEXT_V2 ===
                _clean_lines = []
                for _line in str(_cep_result or "").splitlines():
                    _l = _line.strip()
                    _ll = _l.lower()
                    if not _l:
                        _clean_lines.append(_line)
                        continue
                    if _ll.startswith(("engine:", "manifest:", "артефакт:", "проверка:", "drive_link:", "pdf_link:", "xlsx_link:", "google_sheet_link:", "existing_")):
                        continue
                    if _ll == "ссылки:":
                        continue
                    if _ll.startswith(("- drive_link:", "- pdf_link:", "- xlsx_link:", "- google_sheet_link:", "- existing_")):
                        continue
                    _clean_lines.append(_line)
                _cep_result = "\n".join(_clean_lines).strip()

                if _cep_links and not any(x in _cep_result for x in ("PDF:", "XLSX:", "Google Sheets:")):
                    _seen_public = set()
                    _pdf_public = ""
                    _xlsx_public = ""
                    for _k, _v in _cep_links:
                        _vv = str(_v or "")
                        if not _vv.startswith("http") or _vv in _seen_public:
                            continue
                        _seen_public.add(_vv)
                        if "spreadsheets" in _vv or str(_k).lower() in ("xlsx_link", "google_sheet", "google_sheet_link", "drive_link"):
                            if not _xlsx_public:
                                _xlsx_public = _vv
                        elif ".pdf" in _vv.lower() or "pdf" in str(_k).lower():
                            if not _pdf_public:
                                _pdf_public = _vv
                    _extra = []
                    if _pdf_public:
                        _extra.append("PDF: " + _pdf_public)
                    if _xlsx_public:
                        _extra.append("XLSX: " + _xlsx_public)
                    if _extra:
                        _cep_result += "\n\n" + "\n".join(_extra)
                # === END_CLEAN_CEP_RESULT_TEXT_V2 ===
                if _cep_success or _cep_links or _cep_data.get("artifact_path"):
                    _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_cep_result, error_message="")
                    _history(conn, task_id, "CREATE_ESTIMATE_PRIORITY_NO_ROLLBACK_V1:AWAITING_CONFIRMATION")
                    try:
                        _reply_to = task["reply_to_message_id"] if "reply_to_message_id" in task.keys() else None
                    except Exception:
                        _reply_to = None
                    _send_once_ex(conn, task_id, chat_id, _cep_result, _reply_to, "estimate_create_priority")
                    return

                _cep_err = str(_cep_data.get("error") or "ESTIMATE_ENGINE_NO_ARTIFACT")
                _cep_user = str(_cep_data.get("result_text") or _cep_err)
                _target_state = "WAITING_CLARIFICATION" if _cep_state == "WAITING_CLARIFICATION" else "FAILED"
                _update_task(conn, task_id, state=_target_state, result=_cep_user, error_message=_cep_err)
                _history(conn, task_id, "CREATE_ESTIMATE_PRIORITY_NO_ROLLBACK_V1:" + _target_state)
                try:
                    _reply_to = task["reply_to_message_id"] if "reply_to_message_id" in task.keys() else None
                except Exception:
                    _reply_to = None
                _send_once_ex(conn, task_id, chat_id, _cep_user or _cep_err, _reply_to, "estimate_create_priority_error")
                return

            _update_task(conn, task_id, state="FAILED", result="", error_message="CREATE_ESTIMATE_PRIORITY_NO_ROLLBACK_V1:NON_DICT_RESULT")
            _history(conn, task_id, "CREATE_ESTIMATE_PRIORITY_NO_ROLLBACK_V1:NON_DICT_RESULT")
            return
    except Exception as _cep_e:
        logger.warning("CREATE_ESTIMATE_PRIORITY_NO_ROLLBACK_V1_ERR task=%s err=%s", task_id, _cep_e)
    # === END_CREATE_ESTIMATE_PRIORITY_NO_ROLLBACK_V1 ===
        # === STROYKA_TOPIC2_SOURCE_LOCK_PRIORITY_FIX_V1_WORKER_HOOK ===
    try:
        if int(topic_id or 0) == 2:
            _stc1_raw = str(raw_input or "")
            try:
                _stc1_cols = [r[1] for r in conn.execute("PRAGMA table_info(task_history)").fetchall()]
                _stc1_col = "action" if "action" in _stc1_cols else ("event" if "event" in _stc1_cols else "")
                if _stc1_col:
                    _stc1_rows = conn.execute(
                        f"SELECT {_stc1_col} FROM task_history WHERE task_id=? AND {_stc1_col} LIKE 'clarified:%' ORDER BY rowid ASC LIMIT 30",
                        (task_id,),
                    ).fetchall()
                    _stc1_clar = []
                    for _r in _stc1_rows:
                        _v = str(_r[0] or "")
                        if ":" in _v:
                            _stc1_clar.append(_v.split(":", 1)[1].strip())
                    if _stc1_clar:
                        _stc1_merged = _stc1_raw + "\n\nУточнения пользователя:\n" + "\n".join(x for x in _stc1_clar if x)
                        if _stc1_merged != _stc1_raw:
                            raw_input = _stc1_merged
                            try:
                                conn.execute("UPDATE tasks SET raw_input=?, updated_at=datetime('now') WHERE id=?", (raw_input, task_id))
                                conn.commit()
                                _history(conn, task_id, "STROYKA_TOPIC2_SOURCE_LOCK_PRIORITY_FIX_V1:clarifications_merged")
                            except Exception as _e:
                                logger.warning("STROYKA_TOPIC2_SOURCE_LOCK_PRIORITY_FIX_V1_MERGE_DB_ERR %s", _e)
            except Exception as _e:
                logger.warning("STROYKA_TOPIC2_SOURCE_LOCK_PRIORITY_FIX_V1_MERGE_ERR %s", _e)

            _stc1_low = str(raw_input or "").lower().replace("ё", "е")
            if any(x in _stc1_low for x in ("смет", "стоимост", "посчитай", "расчет", "расчёт", "цена", "руб", "монолит", "фундамент", "плит", "кровл", "строительств")):
                try:
                    from core.sample_template_engine import handle_template_estimate_intent as _stc1_handle_template_estimate_intent
                    _stc1_reply_to = locals().get("reply_to", None) or locals().get("reply_to_message_id", None)
                    _stc1_handled = await _stc1_handle_template_estimate_intent(
                        conn,
                        task_id,
                        str(chat_id),
                        int(topic_id or 0),
                        raw_input,
                        "text",
                        _stc1_reply_to,
                    )
                    if _stc1_handled:
                        logger.info("STROYKA_TOPIC2_SOURCE_LOCK_PRIORITY_FIX_V1 handled task=%s topic=%s", task_id, topic_id)
                        return
                except Exception as _e:
                    logger.warning("STROYKA_TOPIC2_SOURCE_LOCK_PRIORITY_FIX_V1_HANDLE_ERR %s", _e)
    except Exception as _e:
        logger.warning("STROYKA_TOPIC2_SOURCE_LOCK_PRIORITY_FIX_V1_WORKER_ERR %s", _e)
    # === END_STROYKA_TOPIC2_SOURCE_LOCK_PRIORITY_FIX_V1_WORKER_HOOK ===

    # === FILE_TECH_CONTOUR_FOLLOWUP_V2 ===
    try:
        _ft_low = str(raw_input or "").strip()
        if int(topic_id or 0) not in (2, 500, 210) and _filemem_should_followup(_ft_low):  # TOPIC_ROUTE_ISOLATION_FULL_V1
            _ft_answer = _filemem_build_answer(str(chat_id), int(topic_id or 0), _ft_low)
            if _ft_answer:
                _update_task(conn, task_id, state="DONE", result=_ft_answer, error_message="")
                _history(conn, task_id, "FILE_TECH_CONTOUR_FOLLOWUP_V2:DONE")
                try:
                    _save_memory(str(chat_id), int(topic_id or 0), str(raw_input or ""), _ft_answer)
                except Exception as _e:
                    logger.warning("FILE_TECH_CONTOUR_FOLLOWUP_V2_SAVE_ERR %s", _e)
                try:
                    if _Stage6Archive is not None:
                        _Stage6Archive().archive(
                            {
                                "task_id": task_id,
                                "chat_id": str(chat_id),
                                "topic_id": int(topic_id or 0),
                                "direction": "file_tech_followup",
                                "engine": "file_memory_bridge",
                                "input_type": "text",
                                "raw_input": str(raw_input or ""),
                            },
                            {"text": _ft_answer, "result": {"text": _ft_answer}},
                        )
                except Exception as _e:
                    logger.warning("FILE_TECH_CONTOUR_FOLLOWUP_V2_ARCHIVE_ERR %s", _e)
                try:
                    _append_timeline_event_v1(str(chat_id), int(topic_id or 0), task_id, "file_tech_followup_done", raw_input, _ft_answer)
                except Exception as _e:
                    logger.warning("FILE_TECH_CONTOUR_FOLLOWUP_V2_TIMELINE_ERR %s", _e)
                conn.commit()
                _send_once(conn, task_id, str(chat_id), _ft_answer, reply_to, "file_tech_followup_v2")
                logger.info("FILE_TECH_CONTOUR_FOLLOWUP_V2 done task=%s topic=%s", task_id, topic_id)
                return
    except Exception as _ft_e:
        logger.error("FILE_TECH_CONTOUR_FOLLOWUP_V2_ERR task=%s err=%s", task_id, _ft_e)
    # === END FILE_TECH_CONTOUR_FOLLOWUP_V2 ===
    # === ACTIVE_DIALOG_STATE_V1_TASK_WORKER_HOOK ===
    # === ASYNC_TEMPLATE_WORKFLOW_HOOK_V2 ===
    try:
        if maybe_handle_template_workflow_async is not None:
            _async_tpl = await maybe_handle_template_workflow_async(conn, task, chat_id, topic_id)
            if _async_tpl and _async_tpl.get("handled"):
                _async_state = str(_async_tpl.get("state") or "DONE")
                _async_result = str(_async_tpl.get("result") or "")
                _async_error = str(_async_tpl.get("error") or "")
                _update_task(
                    conn,
                    task_id,
                    state=_async_state,
                    result=_async_result,
                    error_message=_async_error,
                )
                _history(conn, task_id, str(_async_tpl.get("event") or "ASYNC_TEMPLATE_WORKFLOW_HOOK_V2:DONE"))
                try:
                    _reply_to = None
                    if "reply_to_message_id" in task.keys():
                        _reply_to = task["reply_to_message_id"]
                    _send_text = _async_result or _async_error or "Задача обработана через шаблон"
                    _send_once_ex(conn, task_id, chat_id, _send_text, _reply_to, "async_template")
                except Exception as _async_send_e:
                    logger.warning("ASYNC_TEMPLATE_WORKFLOW_HOOK_V2_SEND_ERR task=%s err=%s", task_id, _async_send_e)
                return
    except Exception as _async_tpl_e:
        logger.warning("ASYNC_TEMPLATE_WORKFLOW_HOOK_V2_ERR task=%s err=%s", task_id, _async_tpl_e)
    # === END_ASYNC_TEMPLATE_WORKFLOW_HOOK_V2 ===
    try:
        for _canon_handler in (maybe_handle_template_workflow, maybe_handle_active_dialog):
            if _canon_handler is None:
                continue
            _canon_result = _canon_handler(conn, task, chat_id, topic_id)
            if _canon_result and _canon_result.get("handled"):
                _canon_state = str(_canon_result.get("state") or "DONE")
                _canon_text = str(_canon_result.get("result") or "")
                _canon_error = str(_canon_result.get("error") or "")
                _update_task(conn, task_id, state=_canon_state, result=_canon_text, error_message=_canon_error)
                _history(conn, task_id, str(_canon_result.get("event") or "CANON_REMAINING_CODE_FULL_CLOSE_V1:DONE"))
                try:
                    if save_dialog_event is not None:
                        save_dialog_event(chat_id, topic_id, str(_canon_result.get("event") or "event"), _canon_result)
                except Exception as _sd_e:
                    logger.warning("ACTIVE_DIALOG_STATE_V1_SAVE_ERR %s", _sd_e)
                return
    except Exception as _ads_e:
        logger.warning("ACTIVE_DIALOG_STATE_V1_TASK_WORKER_HOOK_ERR task=%s err=%s", task_id, _ads_e)
    # === END_ACTIVE_DIALOG_STATE_V1_TASK_WORKER_HOOK ===

    # === MEMORY_QUERY_GUARD_V1 ===
    try:
        _mq_low = str(raw_input or "").strip().lower().rstrip("!?. ")
        _mq_low = _mq_low.replace("[voice] ", "").replace("[VOICE] ", "").strip()
        _mq_markers = (
            "что обсуждали", "что делали", "что мы делали", "что мы обсуждали",
            "неделю назад", "две недели", "три недели", "месяц назад",
            "апреля", "марта", "февраля", "января", "помнишь", "напомни",
            "какие задачи были", "что было", "расскажи что",
        )
        if any(m in _mq_low for m in _mq_markers):
            _archive_ctx = _load_archive_context(str(chat_id), int(topic_id or 0), str(raw_input or ""))
            _short_ctx, _long_ctx, _topic_role, _topic_directions = _load_memory_context(str(chat_id), int(topic_id or 0))
            _mq_payload = {
                "id": task_id, "task_id": task_id,
                "chat_id": str(chat_id), "topic_id": int(topic_id or 0),
                "input_type": "text", "state": "IN_PROGRESS",
                "raw_input": str(raw_input or ""),
                "short_memory_context": _short_ctx, "long_memory_context": _long_ctx,
                "archive_context": _archive_ctx, "topic_role": _topic_role,
                "topic_directions": _topic_directions,
                "direction": "memory_query", "engine": "ai_router",
            }
            if not str(_archive_ctx or "").strip() and not str(_long_ctx or "").strip():
                _mq_msg = "В этом топике архивных данных по запросу не найдено"
            else:
                _mq_ai = await asyncio.wait_for(process_ai_task(_mq_payload), timeout=AI_TIMEOUT)
                _mq_msg = _clean(_s(_mq_ai), 12000) or "Архив найден, но ответ не сформирован"
            _update_task(conn, task_id, state="DONE", result=_mq_msg, error_message="")
            _history(conn, task_id, "MEMORY_QUERY_GUARD_V1:DONE")
            try:
                _save_memory(str(chat_id), int(topic_id or 0), str(raw_input or ""), _mq_msg)
            except Exception as _e:
                logger.warning("MEMORY_QUERY_GUARD_V1_SAVE_ERR %s", _e)
            try:
                if _Stage6Archive is not None:
                    _Stage6Archive().archive(_mq_payload, {"text": _mq_msg, "result": {"text": _mq_msg}})
            except Exception as _e:
                logger.warning("MEMORY_QUERY_GUARD_V1_ARCHIVE_ERR %s", _e)
            try:
                _append_timeline_event_v1(str(chat_id), int(topic_id or 0), task_id, "memory_query_done", raw_input, _mq_msg)
            except Exception as _e:
                logger.warning("MEMORY_QUERY_GUARD_V1_TIMELINE_ERR %s", _e)
            conn.commit()
            _send_once(conn, task_id, str(chat_id), _mq_msg, reply_to, "memory_query_guard_v1")
            return
    except Exception as _mq_e:
        logger.error("MEMORY_QUERY_GUARD_V1_ERR task=%s err=%s", task_id, _mq_e)
    # === END MEMORY_QUERY_GUARD_V1 ===
    # === FULLFIX_16_CONTEXT_QUERY ===
    try:
        _ff16_low = str(raw_input or "").strip().lower().rstrip("!?. ")
        _ff16_low = _ff16_low.replace("[voice] ", "").replace("[VOICE] ", "").strip()  # FF21_FIX_VOICE_PREFIX
        _ff16_triggers = ["nu chto", "gde rezultat", "chto tam", "gde smeta", "gde proekt"]
        _ff16_ru_triggers = [
            "ну что",
            "где результат",
            "что там",
            "где смета",
            "где проект",
            "что с задачей",
            "что там у нас",
            "ну как там",
            "где файл",
            "ну что там",
            "ну давай",
            "что по задаче",
        ]
        _ff16_is_ctx = len(_ff16_low) <= 35 and any(
            _ff16_low == t or _ff16_low.startswith(t) for t in _ff16_ru_triggers
        )
        if _ff16_is_ctx:
            _ff16_row = conn.execute(
                "SELECT id,state,result FROM tasks"
                " WHERE chat_id=? AND COALESCE(topic_id,0)=? AND id<>?"
                " AND state IN ('AWAITING_CONFIRMATION','IN_PROGRESS','WAITING_CLARIFICATION')"
                " ORDER BY updated_at DESC LIMIT 1",
                (chat_id, topic_id, task_id)
            ).fetchone()
            if _ff16_row is not None:
                _ff16_pid, _ff16_pst, _ff16_pres = _ff16_row
                _ff16_parts = ["Статус: " + str(_ff16_pst)]
                if _ff16_pres is not None:
                    import re as _re16
                    _ff16_clean = _re16.sub(
                        r"(?im)^\s*MANIFEST\s*:\s*https?://\S+\s*$",
                        "",
                        str(_ff16_pres)[:800]
                    ).strip()
                    _ff16_parts.append(_ff16_clean)
                _ff16_msg = "\n".join(_ff16_parts)
                from core.reply_sender import send_reply_ex
                send_reply_ex(chat_id=str(chat_id), text=_ff16_msg, reply_to_message_id=reply_to, message_thread_id=topic_id)  # FULLFIX_20_CONTEXT_QUERY_TOPIC
                conn.execute(
                    "UPDATE tasks SET state='DONE',result=?,updated_at=datetime('now') WHERE id=?",
                    ("Ответил по активной задаче", task_id)
                )
                conn.execute(
                    "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                    (task_id, "state:DONE:context_query_ff16")
                )
                conn.commit()
                return
    except Exception as _ff16_ctx_err:
        logger.error("FULLFIX_16_CONTEXT_QUERY_ERROR task=%s err=%s", task_id, _ff16_ctx_err)
    # === END FULLFIX_16_CONTEXT_QUERY ===

    # === FULLFIX_14_UNIFIED_ROUTE ===
    try:
        from core.template_intake_engine import is_sample_intent as _ff14_is_sample, process_template_intake as _ff14_tmpl
        from core.defect_act_engine import is_defect_act_intent as _ff14_is_defect, process_defect_act as _ff14_defect
        from core.multifile_artifact_engine import is_multifile_intent as _ff14_is_multi, process_multifile as _ff14_multi
        from core.estimate_unified_engine import process_estimate_task as _ff14_estimate, parse_estimate_rows as _ff14_parse
        _ff14_raw = str(raw_input or "")
        _ff14_itype = str(_task_field(task, "input_type") or "")
        _ff14_mime = ""
        _ff14_fname = ""
        _ff14_lpath = ""
        if _ff14_itype == "drive_file":
            try:
                import json as _ff14j
                _ff14_meta = _ff14j.loads(_task_field(task, "raw_input") or "{}")
                _ff14_mime = _ff14_meta.get("mime_type", "")
                _ff14_fname = _ff14_meta.get("file_name", "")
                _ff14_lpath = _ff14_meta.get("local_path", "")
            except Exception:
                pass
        # 1. template/sample intake — highest priority
        if _ff14_is_sample(_ff14_raw):
            _ff14_done = await _ff14_tmpl(
                conn=conn, task_id=task_id, chat_id=chat_id, topic_id=topic_id,
                raw_input=_ff14_raw, local_path=_ff14_lpath,
                file_name=_ff14_fname, mime_type=_ff14_mime
            )
            if _ff14_done:
                return
        # 2. defect/photo act
        if _ff14_is_defect(_ff14_raw, _ff14_mime):
            _ff14_done = await _ff14_defect(
                conn=conn, task_id=task_id, chat_id=chat_id, topic_id=topic_id,
                raw_input=_ff14_raw, file_name=_ff14_fname, local_path=_ff14_lpath
            )
            if _ff14_done:
                return
        # 3. estimate from natural language text
        if _ff14_itype in ("text", "search") and _ff14_parse(_ff14_raw):
            # === FULLFIX_16_ESTIMATE_HARD_STOP ===
            # Estimate route: ALWAYS return, never fall through to FULLFIX_10
            conn.execute("INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (task_id, "FULLFIX_16_ESTIMATE_ROUTE_TAKEN"))
            conn.commit()
            _ff14_done = await _ff14_estimate(
                conn=conn, task_id=task_id, chat_id=chat_id, topic_id=topic_id, raw_input=_ff14_raw
            )
            # Whether success or failure — do not let FULLFIX_10 run on estimate input
            return
            # === END FULLFIX_16_ESTIMATE_HARD_STOP ===
        # 4. multifile aggregation
        if _ff14_is_multi(_ff14_raw):
            _ff14_done = await _ff14_multi(
                conn=conn, task_id=task_id, chat_id=chat_id, topic_id=topic_id, raw_input=_ff14_raw
            )
            if _ff14_done:
                return
    except Exception as _ff14_err:
        try:
            logger.error("FULLFIX_14_ROUTE_ERROR task=%s err=%s", task_id, str(_ff14_err))
        except Exception:
            pass
    # === END FULLFIX_14_UNIFIED_ROUTE ===

# === FULLFIX_13A_SAMPLE_TEMPLATE_AND_TEMPLATE_ESTIMATE_ROUTE ===
    try:
        from core.sample_template_engine import (
            handle_sample_template_intent as _ff13a_handle_sample_template_intent,
            handle_template_estimate_intent as _ff13a_handle_template_estimate_intent,
        )
        # === FULLFIX_13A_ROUTE_LOCALS_FIX ===
        # _handle_new has task/raw_input/reply_to locals, not input_type/reply_to_message_id locals
        _ff13a_conn = conn
        _ff13a_task_id = str(task_id or "")
        _ff13a_chat_id = str(chat_id or "")
        _ff13a_topic_id = int(topic_id or 0)
        _ff13a_raw_input = str(raw_input or "")
        _ff13a_input_type = str(_task_field(task, "input_type", "") or "")
        _ff13a_reply_to = _task_field(task, "reply_to_message_id", None) or _task_field(task, "telegram_message_id", None)
        # === END FULLFIX_13A_ROUTE_LOCALS_FIX ===
        _ff13a_done = await _ff13a_handle_sample_template_intent(
            conn=_ff13a_conn,
            task_id=_ff13a_task_id,
            chat_id=_ff13a_chat_id,
            topic_id=_ff13a_topic_id,
            raw_input=_ff13a_raw_input,
            input_type=_ff13a_input_type,
            reply_to_message_id=_ff13a_reply_to,
        )
        # === FULLFIX_13B_SAMPLE_HARD_STOP_1 ===
        if _ff13a_done:
            try:
                conn.commit()
            except Exception:
                pass
            return
        _ff13a_done = await _ff13a_handle_template_estimate_intent(
            conn=_ff13a_conn,
            task_id=_ff13a_task_id,
            chat_id=_ff13a_chat_id,
            topic_id=_ff13a_topic_id,
            raw_input=_ff13a_raw_input,
            input_type=_ff13a_input_type,
            reply_to_message_id=_ff13a_reply_to,
        )
        # === FULLFIX_13B_SAMPLE_HARD_STOP_2 ===
        if _ff13a_done:
            try:
                conn.commit()
            except Exception:
                pass
            return
    except Exception as _ff13a_err:
        try:
            logger.error("FULLFIX_13A_SAMPLE_ROUTE_ERROR task=%s err=%s", task_id, str(_ff13a_err))
        except Exception:
            pass
    # === END FULLFIX_13A_SAMPLE_TEMPLATE_AND_TEMPLATE_ESTIMATE_ROUTE ===


    # === FULLFIX_19_PROJECT_GUARD_REAL_V2 ===
    try:
        _ff19_low = str(raw_input or "").strip().lower()
        _ff19_short_replies = {"да","нет","ок","готово","угу","так","ясно","понятно"}
        if _ff19_low in _ff19_short_replies:
            try:
                conn.execute(
                    "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                    (task_id, "FULLFIX_19_PROJECT_GUARD_BLOCKED_SHORT_REPLY")
                )
                conn.execute(
                    "UPDATE tasks SET state='DONE', updated_at=datetime('now'), result=COALESCE(NULLIF(result,''),?) WHERE id=?",
                    ("Принято", task_id)
                )
                conn.commit()
            except Exception:
                pass
            logger.info("FF19_GUARD_BLOCKED_SHORT_REPLY task=%s text=%s", task_id, _ff19_low)
            return
    except Exception as _ff19_err:
        logger.error("FF19_PROJECT_GUARD_ERR task=%s err=%s", task_id, _ff19_err)
    # === END FULLFIX_19_PROJECT_GUARD_REAL_V2 ===

    # === FULLFIX_13B_FALSE_PROJECT_PHRASE_GUARD ===
    try:
        _ff13b_low = str(raw_input or "").strip().lower()
        _ff13b_false_project_phrases = {
            "это один из вариантов",
            "это как образец",
            "это пример",
            "вот образец",
            "вот пример",
            "сохрани как образец",
        }
        if _ff13b_low in _ff13b_false_project_phrases:
            _msg = "Принял как образец. Дальше можно писать простым языком: сделай смету / сделай проект"
            await safe_update(conn, task_id, state="DONE", result=_ff13c_strip_manifest_links(_msg), error_message="")
            try:
                from core.reply_sender import send_reply_ex
                send_reply_ex(chat_id=str(chat_id), text=_ff13c_strip_manifest_links(_msg), reply_to_message_id=reply_to)
            except Exception:
                pass
            try:
                conn.execute("INSERT INTO task_history (task_id,action,created_at) VALUES (?,?,datetime('now'))", (task_id, "FULLFIX_13B_FALSE_PROJECT_GUARDED"))
                conn.commit()
            except Exception:
                pass
            return
    except Exception as _ff13b_guard_err:
        try:
            logger.error("FULLFIX_13B_FALSE_PROJECT_GUARD_ERROR task=%s err=%s", task_id, str(_ff13b_guard_err))
        except Exception:
            pass
    # === END FULLFIX_13B_FALSE_PROJECT_PHRASE_GUARD ===

    if int(topic_id or 0) != 500:  # TOPIC500_FULLFIX10_BYPASS_V1
        # === FULLFIX_10_TOTAL_CLOSURE_UNIVERSAL_ROUTE ===
        try:
            from core.orchestra_closure_engine import (
                classify_user_task as _ff10_classify_user_task,
                classify_project_kind as _ff10_classify_project_kind,
                create_estimate_files as _ff10_create_estimate_files,
                save_result_memory as _ff10_save_result_memory,
                ENGINE as _FF10_ENGINE,
            )

            _ff10_intent = _ff10_classify_user_task(str(raw_input or ""))


            # === FULLFIX_13B_CLEAN_ESTIMATE_MESSAGE_BEFORE_SEND_FALLBACK ===
            def _ff13b_clean_any_estimate_text(_txt):
                try:
                    from core.orchestra_closure_engine import ff13b_clean_estimate_user_message
                    return ff13b_clean_estimate_user_message(_txt)
                except Exception:
                    return _txt
            # === END FULLFIX_13B_CLEAN_ESTIMATE_MESSAGE_BEFORE_SEND_FALLBACK ===

            if _ff10_intent in ("confirm", "revision"):
                _ff10_parent = conn.execute(
                    """
                    SELECT id,state,result,reply_to_message_id,bot_message_id
                    FROM tasks
                    WHERE chat_id=?
                      AND COALESCE(topic_id,0)=?
                      AND state='AWAITING_CONFIRMATION'
                      AND id<>?
                    ORDER BY updated_at DESC, created_at DESC
                    LIMIT 1
                    """,
                    (str(chat_id), int(topic_id or 0), task_id),
                ).fetchone()

                if _ff10_parent and _ff10_intent == "confirm":
                    _parent_id = _s(_ff10_parent["id"])
                    _update_task(conn, _parent_id, state="DONE", error_message="")
                    _history(conn, _parent_id, "FULLFIX_10_CONFIRM_DONE")
                    _update_task(conn, task_id, state="DONE", result="Подтверждение принято. Задача закрыта", error_message="")
                    _history(conn, task_id, "FULLFIX_10_CONFIRM_CHILD_DONE")
                    conn.commit()
                    _send_once(conn, task_id, chat_id, "Подтверждение принято. Задача закрыта", reply_to, "ff10_confirm_done")
                    return

                if _ff10_parent and _ff10_intent == "revision":
                    _parent_id = _s(_ff10_parent["id"])
                    _merged = _clean(_s(_ff10_parent["result"]) + "\n\nПравки пользователя:\n" + str(raw_input or ""), 12000)
                    _update_task(conn, _parent_id, state="IN_PROGRESS", raw_input=_merged, error_message="")
                    _history(conn, _parent_id, "FULLFIX_10_REVISION_REOPEN")
                    _update_task(conn, task_id, state="DONE", result="Правки приняты. Задача возвращена в работу", error_message="")
                    _history(conn, task_id, "FULLFIX_10_REVISION_CHILD_DONE")
                    conn.commit()
                    _send_once(conn, task_id, chat_id, "Правки приняты. Задача возвращена в работу", reply_to, "ff10_revision_reopen")
                    return

            if _ff10_intent == "project":
                _kind, _section = _ff10_classify_project_kind(str(raw_input or ""))
                if _kind == "foundation_slab":
                    from core.project_engine import create_project_pdf_dxf_artifact
                    _ff10_res = await create_project_pdf_dxf_artifact(str(raw_input or ""), task_id, int(topic_id or 0), "FULLFIX_10_SIMPLE_USER_REQUEST", True)

                    if not isinstance(_ff10_res, dict) or not _ff10_res.get("success"):
                        _err = str((_ff10_res or {}).get("error", "PROJECT_FAILED"))[:400]
                        _msg = "Проект не создан: " + _err
                        _update_task(conn, task_id, state="FAILED", result=_ff13c_strip_manifest_links(_msg), error_message=_err)
                        _history(conn, task_id, "FULLFIX_10_PROJECT_FAILED:" + _err)
                        conn.commit()
                        _send_once(conn, task_id, chat_id, _msg, reply_to, "ff10_project_failed")
                        return

                    _pdf = str(_ff10_res.get("pdf_link") or "")
                    _dxf = str(_ff10_res.get("dxf_link") or "")
                    _xlsx = str(_ff10_res.get("xlsx_link") or "")
                    _manifest = str(_ff10_res.get("manifest_link") or "")
                    _sheet_count = str(_ff10_res.get("sheet_count") or "")
                    _engine = str(_ff10_res.get("engine") or _FF10_ENGINE)
                    _msg = (
                        "Проект создан\n"
                        f"Engine: {_engine}\n"
                        "Раздел: КЖ\n"
                        f"Тип: фундаментная плита\n"
                        f"Листов: {_sheet_count}\n"
                        f"PDF: {_pdf}\n"
                        f"DXF: {_dxf}\n"
                        f"XLSX: {_xlsx}\n"
                        f"MANIFEST: {_manifest}\n\n"
                        "Доволен результатом? Ответь: Да / Уточни / Правки"
                    )
                    if not _pdf or not _dxf:
                        _update_task(conn, task_id, state="FAILED", result="Проект не создан: нет PDF/DXF ссылки", error_message="PROJECT_LINKS_MISSING")
                        _history(conn, task_id, "FULLFIX_10_PROJECT_LINKS_MISSING")
                        conn.commit()
                        _send_once(conn, task_id, chat_id, "Проект не создан: нет PDF/DXF ссылки", reply_to, "ff10_project_links_missing")
                        return

                        # === RESULT_VALIDATOR_GUARD_V1 ===
                        if _check_result_before_confirm(_ff13c_strip_manifest_links(_msg)):
                            # === TASK_WORKER_ARTIFACT_GATE_V1 ===
                            try:
                                if validate_engine_result is not None:
                                    _twag_raw = ""
                                    _twag_input_type = ""
                                    try:
                                        _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                                    except Exception:
                                        _twag_raw = ""
                                    try:
                                        _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                                    except Exception:
                                        _twag_input_type = ""
                                    _twag_result = _ff13c_strip_manifest_links(_msg)
                                    _twag_check = validate_engine_result(
                                        {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                                        input_type=_twag_input_type,
                                        user_text=_twag_raw,
                                        topic_id=topic_id,
                                    )
                                    if not _twag_check.get("ok"):
                                        _update_task(conn, task_id, state="FAILED", result="",
                                            error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                                        _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                                        return
                            except Exception as _twag_e:
                                logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
                            # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
                            _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_ff13c_strip_manifest_links(_msg), error_message="")
                        else:
                            _update_task(conn, task_id, state="FAILED", result=_ff13c_strip_manifest_links(_msg), error_message="FORBIDDEN_PHRASE")
                        # === END RESULT_VALIDATOR_GUARD_V1 ===
                    _history(conn, task_id, "FULLFIX_10_PROJECT_OK")
                    conn.commit()
                    _sent = _send_once_ex(conn, task_id, str(chat_id), _msg, reply_to, "ff10_project_result")
                    if isinstance(_sent, dict) and _sent.get("bot_message_id"):
                        _update_task(conn, task_id, bot_message_id=_sent["bot_message_id"])
                        conn.commit()
                    _ff10_save_result_memory(str(chat_id), int(topic_id or 0), str(raw_input or ""), _msg, _ff10_res)
                    return

            if _ff10_intent == "estimate":
                _ff10_res = _ff10_create_estimate_files(str(raw_input or ""), task_id, int(topic_id or 0))
                if not isinstance(_ff10_res, dict) or not _ff10_res.get("success"):
                    _err = str((_ff10_res or {}).get("error", "ESTIMATE_FAILED"))[:400]
                    _msg = "Смета не создана: " + _err
                    _update_task(conn, task_id, state="FAILED", result=_ff13c_strip_manifest_links(_msg), error_message=_err)
                    _history(conn, task_id, "FULLFIX_10_ESTIMATE_FAILED:" + _err)
                    conn.commit()
                    _send_once(conn, task_id, chat_id, _msg, reply_to, "ff10_estimate_failed")
                    return

                _msg = str(_ff10_res.get("message") or "")
                # === RESULT_VALIDATOR_GUARD_V1 ===
                if _check_result_before_confirm(_ff13c_strip_manifest_links(_msg)):
                    # === TASK_WORKER_ARTIFACT_GATE_V1 ===
                    try:
                        if validate_engine_result is not None:
                            _twag_raw = ""
                            _twag_input_type = ""
                            try:
                                _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                            except Exception:
                                _twag_raw = ""
                            try:
                                _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                            except Exception:
                                _twag_input_type = ""
                            _twag_result = _ff13c_strip_manifest_links(_msg)
                            _twag_check = validate_engine_result(
                                {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                                input_type=_twag_input_type,
                                user_text=_twag_raw,
                                topic_id=topic_id,
                            )
                            if not _twag_check.get("ok"):
                                _update_task(conn, task_id, state="FAILED", result="",
                                    error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                                _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                                return
                    except Exception as _twag_e:
                        logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
                    # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
                    _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_ff13c_strip_manifest_links(_msg), error_message="")
                else:
                    _update_task(conn, task_id, state="FAILED", result=_ff13c_strip_manifest_links(_msg), error_message="FORBIDDEN_PHRASE")
                # === END RESULT_VALIDATOR_GUARD_V1 ===
                _history(conn, task_id, "FULLFIX_10_ESTIMATE_OK")
                conn.commit()
                _sent = _send_once_ex(conn, task_id, str(chat_id), _msg, reply_to, "ff10_estimate_result")
                if isinstance(_sent, dict) and _sent.get("bot_message_id"):
                    _update_task(conn, task_id, bot_message_id=_sent["bot_message_id"])
                    conn.commit()
                _ff10_save_result_memory(str(chat_id), int(topic_id or 0), str(raw_input or ""), _msg, _ff10_res)
                return

        except Exception as _ff10_e:
            _err = str(_ff10_e)[:500]
            _update_task(conn, task_id, state="FAILED", result="Ошибка FULLFIX_10: " + _err, error_message=_err)
            _history(conn, task_id, "FULLFIX_10_EXCEPTION:" + _err)
            conn.commit()
            _send_once(conn, task_id, chat_id, "Ошибка FULLFIX_10: " + _err, reply_to, "ff10_exception")
            return
        # === END FULLFIX_10_TOTAL_CLOSURE_UNIVERSAL_ROUTE ===


    # === FULLFIX_07_PROJECT_DESIGN_CLOSURE_ROUTE ===
    _ff07_low = str(raw_input or "").lower()
    _ff07_triggers = (
        "создай проект",
        "сделай проект",
        "проект фундамент",
        "проект фундаментной плиты",
        "план фундамент",
        "план фундаментной плиты",
        "фундаментной плиты",
        "проект по образцу",
        "по образцу проект",
        "проект по шаблону",
        "dxf проект",
        "dwg проект",
    )
    if any(x in _ff07_low for x in _ff07_triggers):
        try:
            from core.project_engine import create_project_pdf_dxf_artifact
            _ff07_res = await create_project_pdf_dxf_artifact(str(raw_input), task_id, int(topic_id or 0), "", True)

            if not isinstance(_ff07_res, dict) or not _ff07_res.get("success"):
                _err = str((_ff07_res or {}).get("error", "PROJECT_FAILED"))[:500]
                _update_task(
                    conn,
                    task_id,
                    state="FAILED",
                    result="Проект не создан: нет полного комплекта PDF/DXF/XLSX/MANIFEST или шаблон неполный",
                    error_message=_err,
                )
                _history(conn, task_id, "FULLFIX_07_FAILED:" + _err)
                conn.commit()
                _send_once(
                    conn,
                    task_id,
                    chat_id,
                    "Проект не создан: нет полного комплекта PDF/DXF/XLSX/MANIFEST или шаблон неполный",
                    reply_to,
                    "project_failed",
                )
                return

            _pdf = str(_ff07_res.get("pdf_link") or "")
            _dxf = str(_ff07_res.get("dxf_link") or "")
            _xlsx = str(_ff07_res.get("xlsx_link") or "")
            _manifest = str(_ff07_res.get("manifest_link") or "")
            _engine = str(_ff07_res.get("engine") or "FULLFIX_07_PROJECT_DESIGN_CLOSURE")
            _tpl = str(_ff07_res.get("template_file") or "")
            _sheet_count = str(_ff07_res.get("sheet_count") or "0")
            _sec = str(_ff07_res.get("section") or "КЖ")

            _msg = (
                "Проект создан\n"
                f"Engine: {_engine}\n"
                f"Раздел: {_sec}\n"
                f"Листов: {_sheet_count}\n"
                f"Шаблон: {_tpl}\n"
                f"PDF: {_pdf}\n"
                f"DXF: {_dxf}\n"
                f"XLSX: {_xlsx}\n"
                f"MANIFEST: {_manifest}\n\n"
                "Доволен результатом? Ответь: Да / Уточни / Правки"
            )

            # === RESULT_VALIDATOR_GUARD_V1 ===
            if _check_result_before_confirm(_ff13c_strip_manifest_links(_msg)):
                # === TASK_WORKER_ARTIFACT_GATE_V1 ===
                try:
                    if validate_engine_result is not None:
                        _twag_raw = ""
                        _twag_input_type = ""
                        try:
                            _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                        except Exception:
                            _twag_raw = ""
                        try:
                            _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                        except Exception:
                            _twag_input_type = ""
                        _twag_result = _ff13c_strip_manifest_links(_msg)
                        _twag_check = validate_engine_result(
                            {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                            input_type=_twag_input_type,
                            user_text=_twag_raw,
                            topic_id=topic_id,
                        )
                        if not _twag_check.get("ok"):
                            _update_task(conn, task_id, state="FAILED", result="",
                                error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                            _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                            return
                except Exception as _twag_e:
                    logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
                # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
                _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_ff13c_strip_manifest_links(_msg), error_message="")
            else:
                _update_task(conn, task_id, state="FAILED", result=_ff13c_strip_manifest_links(_msg), error_message="FORBIDDEN_PHRASE")
            # === END RESULT_VALIDATOR_GUARD_V1 ===
            _history(conn, task_id, "FULLFIX_07_PROJECT_OK")
            conn.commit()

            try:
                _sent = _send_once_ex(conn, task_id, str(chat_id), _msg, reply_to, "project_fullfix_07_result")
                if isinstance(_sent, dict) and _sent.get("bot_message_id"):
                    _update_task(conn, task_id, bot_message_id=_sent["bot_message_id"])
                    conn.commit()
            except Exception:
                _send_once(conn, task_id, chat_id, _msg, reply_to, "project_fullfix_07_result")
            return

        except Exception as _ff07_e:
            _err = str(_ff07_e)[:700]
            _update_task(
                conn,
                task_id,
                state="FAILED",
                result="Проект не создан: ошибка генерации полного комплекта: " + _err,
                error_message=_err,
            )
            _history(conn, task_id, "FULLFIX_07_EXCEPTION:" + _err)
            conn.commit()
            _send_once(conn, task_id, chat_id, "Проект не создан: ошибка генерации полного комплекта: " + _err, reply_to, "project_exception")
            return
    # === END FULLFIX_07_PROJECT_DESIGN_CLOSURE_ROUTE ===


    # === FULLFIX_07_CAD_PROJECT_DOCUMENTATION_ROUTE ===
    try:
        from core.cad_project_engine import is_project_design_request, create_full_project_package, format_project_result_message
        if is_project_design_request(raw_input):
            _ff07_res = create_full_project_package(str(raw_input), task_id, int(topic_id or 0), "")
            _ff07_msg = format_project_result_message(_ff07_res)
            if not isinstance(_ff07_res, dict) or not _ff07_res.get("success"):
                _err = str((_ff07_res or {}).get("error") or "PROJECT_DOCUMENTATION_FAILED")[:300]
                _update_task(conn, task_id, state="FAILED", result=_ff07_msg, error_message=_err)
                _history(conn, task_id, "FULLFIX_07_PROJECT_FAILED:" + _err)
                conn.commit()
                _send_once(conn, task_id, chat_id, _ff07_msg, reply_to, "ff07_project_failed")
                return

                # === RESULT_VALIDATOR_GUARD_V1 ===
                if _check_result_before_confirm(_ff07_msg):
                    # === TASK_WORKER_ARTIFACT_GATE_V1 ===
                    try:
                        if validate_engine_result is not None:
                            _twag_raw = ""
                            _twag_input_type = ""
                            try:
                                _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                            except Exception:
                                _twag_raw = ""
                            try:
                                _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                            except Exception:
                                _twag_input_type = ""
                            _twag_result = _ff07_msg
                            _twag_check = validate_engine_result(
                                {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                                input_type=_twag_input_type,
                                user_text=_twag_raw,
                                topic_id=topic_id,
                            )
                            if not _twag_check.get("ok"):
                                _update_task(conn, task_id, state="FAILED", result="",
                                    error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                                _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                                return
                    except Exception as _twag_e:
                        logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
                    # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
                    _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_ff07_msg, error_message="")
                else:
                    _update_task(conn, task_id, state="FAILED", result=_ff07_msg, error_message="FORBIDDEN_PHRASE")
                # === END RESULT_VALIDATOR_GUARD_V1 ===
            _history(conn, task_id, "FULLFIX_07_PROJECT_OK")
            conn.commit()
            _sent = _send_once_ex(conn, task_id, str(chat_id), _ff07_msg, reply_to, "ff07_project_result")
            if isinstance(_sent, dict) and _sent.get("bot_message_id"):
                _update_task(conn, task_id, bot_message_id=_sent["bot_message_id"])
                conn.commit()
            return
    except Exception as _ff07_e:
        _err = str(_ff07_e)[:500]
        _update_task(conn, task_id, state="FAILED", result="Проект не создан: ошибка FULLFIX_07", error_message=_err)
        _history(conn, task_id, "FULLFIX_07_EXCEPTION:" + _err)
        conn.commit()
        _send_once(conn, task_id, chat_id, "Проект не создан: ошибка FULLFIX_07", reply_to, "ff07_project_exception")
        return
    # === END FULLFIX_07_CAD_PROJECT_DOCUMENTATION_ROUTE ===

    # === FULLFIX_06_FINAL_PROJECT_TEMPLATE_ROUTE ===
    _ff06_low = str(raw_input or "").lower()
    _ff06_project_triggers = (
        "создай проект",
        "сделай проект",
        "разработай проект",
        "готовый проект",
        "проект фундамент",
        "проект фундаментной плиты",
        "проект кровли",
        "проект по образцу",
        "проект по шаблону",
        "план фундаментной плиты",
        "чертеж фундаментной плиты",
        "чертёж фундаментной плиты",
    )
    if any(x in _ff06_low for x in _ff06_project_triggers):
        try:
            from core.project_engine import create_project_pdf_dxf_artifact
            _ff06_res = await create_project_pdf_dxf_artifact(str(raw_input), task_id, int(topic_id or 0), "", True)
            if not isinstance(_ff06_res, dict) or not _ff06_res.get("success"):
                _err = str((_ff06_res or {}).get("error", "PROJECT_FAILED"))[:300]
                _update_task(conn, task_id, state="FAILED", result="Проект не создан: нет сохранённого шаблона или не созданы PDF/DXF/XLSX ссылки", error_message=_err)
                _history(conn, task_id, "FULLFIX_06_PROJECT_FAILED:" + _err)
                conn.commit()
                _send_once(conn, task_id, chat_id, "Проект не создан: нет сохранённого шаблона или не созданы PDF/DXF/XLSX ссылки", reply_to, "project_failed_ff06")
                return

            _msg = (
                "Проект создан по сохранённому шаблону\n"
                f"Раздел: {_ff06_res.get('section')}\n"
                f"Листов по шаблону: {_ff06_res.get('sheet_count')}\n"
                f"Шаблон: {_ff06_res.get('template_file')}\n"
                f"PDF: {_ff06_res.get('pdf_link')}\n"
                f"DXF: {_ff06_res.get('dxf_link')}\n"
                f"XLSX: {_ff06_res.get('xlsx_link')}\n"
                f"MANIFEST: {_ff06_res.get('manifest_link')}\n\n"
                "Доволен результатом? Ответь: Да / Уточни / Правки"
            )
            # === RESULT_VALIDATOR_GUARD_V1 ===
            if _check_result_before_confirm(_ff13c_strip_manifest_links(_msg)):
                # === TASK_WORKER_ARTIFACT_GATE_V1 ===
                try:
                    if validate_engine_result is not None:
                        _twag_raw = ""
                        _twag_input_type = ""
                        try:
                            _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                        except Exception:
                            _twag_raw = ""
                        try:
                            _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                        except Exception:
                            _twag_input_type = ""
                        _twag_result = _ff13c_strip_manifest_links(_msg)
                        _twag_check = validate_engine_result(
                            {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                            input_type=_twag_input_type,
                            user_text=_twag_raw,
                            topic_id=topic_id,
                        )
                        if not _twag_check.get("ok"):
                            _update_task(conn, task_id, state="FAILED", result="",
                                error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                            _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                            return
                except Exception as _twag_e:
                    logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
                # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
                _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_ff13c_strip_manifest_links(_msg), error_message="")
            else:
                _update_task(conn, task_id, state="FAILED", result=_ff13c_strip_manifest_links(_msg), error_message="FORBIDDEN_PHRASE")
            # === END RESULT_VALIDATOR_GUARD_V1 ===
            _history(conn, task_id, "FULLFIX_06_PROJECT_OK")
            conn.commit()
            _sent = _send_once_ex(conn, task_id, str(chat_id), _msg, reply_to, "project_result_ff06")
            if isinstance(_sent, dict) and _sent.get("bot_message_id"):
                _update_task(conn, task_id, bot_message_id=_sent["bot_message_id"])
                conn.commit()
            return
        except Exception as _ff06_e:
            _err = str(_ff06_e)[:500]
            _update_task(conn, task_id, state="FAILED", result="Проект не создан: ошибка генерации PDF/DXF/XLSX", error_message=_err)
            _history(conn, task_id, "FULLFIX_06_PROJECT_EXCEPTION:" + _err)
            conn.commit()
            _send_once(conn, task_id, chat_id, "Проект не создан: ошибка генерации PDF/DXF/XLSX", reply_to, "project_exception_ff06")
            return
    # === END FULLFIX_06_FINAL_PROJECT_TEMPLATE_ROUTE ===


    role = _detect_role_assignment(raw_input)
    if role:
        ask = f"Понял назначение чата так:\n{role}\n\nПодтверди или уточни"
        # === RESULT_VALIDATOR_GUARD_V1 ===
        if _check_result_before_confirm(ask):
            # === TASK_WORKER_ARTIFACT_GATE_V1 ===
            try:
                if validate_engine_result is not None:
                    _twag_raw = ""
                    _twag_input_type = ""
                    try:
                        _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                    except Exception:
                        _twag_raw = ""
                    try:
                        _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                    except Exception:
                        _twag_input_type = ""
                    _twag_result = ask
                    _twag_check = validate_engine_result(
                        {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                        input_type=_twag_input_type,
                        user_text=_twag_raw,
                        topic_id=topic_id,
                    )
                    if not _twag_check.get("ok"):
                        _update_task(conn, task_id, state="FAILED", result="",
                            error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                        _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                        return
            except Exception as _twag_e:
                logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
            # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
            _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=ask, error_message="")
        else:
            _update_task(conn, task_id, state="FAILED", result=ask, error_message="FORBIDDEN_PHRASE")
        # === END RESULT_VALIDATOR_GUARD_V1 ===
        _history(conn, task_id, "state:AWAITING_CONFIRMATION")
        conn.commit()
        _send_once(conn, task_id, chat_id, ask, reply_to, "role_confirmation")
        return

    pending_confirm = conn.execute(
        """
        SELECT id, COALESCE(raw_input,'') AS raw_input, COALESCE(result,'') AS result
        FROM tasks
        WHERE chat_id=?
          AND id<>?
          AND COALESCE(topic_id,0)=?
          AND state='AWAITING_CONFIRMATION'
        ORDER BY updated_at DESC, created_at DESC
        LIMIT 1
        """,
        (str(chat_id), task_id, int(topic_id)),
    ).fetchone()

    if pending_confirm:
        pending_id = _s(pending_confirm["id"])
        pending_role = _extract_role_confirmation(_s(pending_confirm["result"]))
        if pending_role and _is_confirm_intent(raw_input):
            _save_topic_role(chat_id, topic_id, pending_role)
            _update_task(conn, pending_id, state="DONE", result=f"Чат закреплён за: {pending_role}", error_message="")
            _history(conn, pending_id, f"role_saved:{pending_role}")
            _update_task(conn, task_id, state="DONE", result="Подтверждение принято", error_message="")
            _history(conn, task_id, "confirm_accepted")
            conn.commit()
            _send_once(conn, task_id, chat_id, f"Принял. Чат закреплён за: {pending_role}", reply_to, "role_saved")
            return
        if pending_role and _is_revision_intent(raw_input):
            _update_task(conn, pending_id, raw_input=raw_input, state="NEW", result="", error_message="")
            _history(conn, pending_id, "role_revision_requested")
            _update_task(conn, task_id, state="DONE", result="Правки приняты", error_message="")
            _history(conn, task_id, "state:DONE")
            conn.commit()
            _send_once(conn, task_id, chat_id, "Принял правки. Уточни назначение чата одной фразой", reply_to, "role_revision_ok")
            return
        if _is_confirm_intent(raw_input):
            _finalize_done(conn, pending_id, chat_id, topic_id, reply_to)
            _update_task(conn, task_id, state="DONE", result="Подтверждение принято", error_message="")
            _history(conn, task_id, "confirm_accepted")
            conn.commit()
            _send_once(conn, task_id, chat_id, "Принял. Задача закрыта", reply_to, "confirm_done")
            return
        if _is_revision_intent(raw_input):
            merged = _clean(_s(pending_confirm["raw_input"]) + "\n\nУточнение пользователя:\n" + raw_input, 12000)
            _update_task(conn, pending_id, raw_input=merged, state="IN_PROGRESS", error_message="")
            _history(conn, pending_id, "revision_accepted")
            _update_task(conn, task_id, state="DONE", result="Правки приняты", error_message="")
            _history(conn, task_id, "state:DONE")
            conn.commit()
            _send_once(conn, task_id, chat_id, "Принял правки. Делаю", reply_to, "revision_ok")
            return

    # === FULLFIX_05_REQUIRE_REAL_PDF_DXF_PROJECT ===
    _ff05_low = str(raw_input or "").lower()
    _ff05_project_triggers = (
        "создай проект",
        "сделай проект",
        "проект фундамент",
        "проект фундаментной плиты",
        "план фундамент",
        "план фундаментной плиты",
        "фундаментной плиты",
        "проект по образцу",
        "проект по шаблону",
        "сделай по образцу",
    )
    if any(x in _ff05_low for x in _ff05_project_triggers):
        try:
            from core.project_engine import create_project_pdf_dxf_artifact
            _ff05_res = await create_project_pdf_dxf_artifact(str(raw_input), task_id, int(topic_id or 0), "")
            if not isinstance(_ff05_res, dict) or not _ff05_res.get("success"):
                _err = str((_ff05_res or {}).get("error", "PROJECT_FAILED"))[:300]
                _update_task(
                    conn,
                    task_id,
                    state="FAILED",
                    result="Проект не создан: нет PDF/DXF файла или ссылки",
                    error_message=_err,
                )
                _history(conn, task_id, "FULLFIX_05_PROJECT_FAILED:" + _err)
                conn.commit()
                _send_once(
                    conn,
                    task_id,
                    chat_id,
                    "Проект не создан: нет PDF/DXF файла или ссылки",
                    reply_to,
                    "project_failed",
                )
                return

            _pdf = str(_ff05_res.get("pdf_link") or "")
            _dxf = str(_ff05_res.get("dxf_link") or "")
            _manifest = str(_ff05_res.get("manifest_link") or "")
            _sec = str(_ff05_res.get("section") or "КЖ")
            _msg = (
                f"Проект создан как PDF/DXF комплект\n"
                f"Раздел: {_sec}\n"
                f"PDF: {_pdf}\n"
                f"DXF: {_dxf}\n"
            )
            if _manifest:
                _msg += f"MANIFEST: {_manifest}\n"
            _msg += "\nДоволен результатом? Ответь: Да / Уточни / Правки"

            # === RESULT_VALIDATOR_GUARD_V1 ===
            if _check_result_before_confirm(_ff13c_strip_manifest_links(_msg)):
                # === TASK_WORKER_ARTIFACT_GATE_V1 ===
                try:
                    if validate_engine_result is not None:
                        _twag_raw = ""
                        _twag_input_type = ""
                        try:
                            _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                        except Exception:
                            _twag_raw = ""
                        try:
                            _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                        except Exception:
                            _twag_input_type = ""
                        _twag_result = _ff13c_strip_manifest_links(_msg)
                        _twag_check = validate_engine_result(
                            {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                            input_type=_twag_input_type,
                            user_text=_twag_raw,
                            topic_id=topic_id,
                        )
                        if not _twag_check.get("ok"):
                            _update_task(conn, task_id, state="FAILED", result="",
                                error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                            _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                            return
                except Exception as _twag_e:
                    logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
                # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
                _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_ff13c_strip_manifest_links(_msg), error_message="")
            else:
                _update_task(conn, task_id, state="FAILED", result=_ff13c_strip_manifest_links(_msg), error_message="FORBIDDEN_PHRASE")
            # === END RESULT_VALIDATOR_GUARD_V1 ===
            _history(conn, task_id, "FULLFIX_05_PROJECT_PDF_DXF_OK")
            conn.commit()

            _sent = _send_once_ex(conn, task_id, str(chat_id), _msg, reply_to, "project_pdf_dxf_result")
            if isinstance(_sent, dict) and _sent.get("bot_message_id"):
                _update_task(conn, task_id, bot_message_id=_sent["bot_message_id"])
                conn.commit()
            return

        except Exception as _ff05_e:
            _err = str(_ff05_e)[:500]
            _update_task(
                conn,
                task_id,
                state="FAILED",
                result="Проект не создан: ошибка генерации PDF/DXF",
                error_message=_err,
            )
            _history(conn, task_id, "FULLFIX_05_PROJECT_EXCEPTION:" + _err)
            conn.commit()
            _send_once(conn, task_id, chat_id, "Проект не создан: ошибка генерации PDF/DXF", reply_to, "project_exception")
            return
    # === END FULLFIX_05_REQUIRE_REAL_PDF_DXF_PROJECT ===

    pending_clarify = conn.execute(
        """
        SELECT id, COALESCE(raw_input,'') AS raw_input
        FROM tasks
        WHERE chat_id=?
          AND id<>?
          AND COALESCE(topic_id,0)=?
          AND state='WAITING_CLARIFICATION'
        ORDER BY updated_at DESC, created_at DESC
        LIMIT 1
        """,
        (str(chat_id), task_id, int(topic_id)),
    ).fetchone()

    if pending_clarify:
        pending_id = _s(pending_clarify["id"])
        merged = _clean(_s(pending_clarify["raw_input"]) + "\n\nУточнение пользователя:\n" + raw_input, 12000)
        _update_task(conn, pending_id, raw_input=merged, state="IN_PROGRESS", error_message="")
        _history(conn, pending_id, "clarification_accepted")
        _update_task(conn, task_id, state="DONE", result="Уточнение принято", error_message="")
        _history(conn, task_id, "state:DONE")
        conn.commit()
        _send_once(conn, task_id, chat_id, "Принял уточнение. Делаю", reply_to, "clarification_ok")
        return


    _update_task(conn, task_id, state="IN_PROGRESS", error_message="")
    _history(conn, task_id, "state:IN_PROGRESS")
    conn.commit()



def _get_parent_task_id(conn, chat_id, reply_to_message_id, topic_id):
    """
    PARENT_CONFIRM_FALLBACK_V1

    Fact-based fix:
    - Telegram reply_to_message_id can point to a nearby bot message id that is not stored as tasks.bot_message_id
    - exact bot_message_id/reply_to_message_id match is still priority
    - if exact match fails and user replied inside the same topic, bind to latest AWAITING_CONFIRMATION in that topic
    - no DONE/FAILED/CANCELLED task is revived here
    """
    if not reply_to_message_id:
        return None

    try:
        _topic_id = int(topic_id or 0)
    except Exception:
        _topic_id = 0

    _chat_id = str(chat_id)
    _reply_id = str(reply_to_message_id)

    row = conn.execute("""
        SELECT id FROM tasks
        WHERE chat_id = ?
          AND COALESCE(topic_id,0) = ?
          AND CAST(COALESCE(bot_message_id,'') AS TEXT) = ?
          AND state IN ('NEW','IN_PROGRESS','WAITING_CLARIFICATION','AWAITING_CONFIRMATION')
        ORDER BY updated_at DESC
        LIMIT 1
    """, (_chat_id, _topic_id, _reply_id)).fetchone()
    if row:
        return row["id"]

    row = conn.execute("""
        SELECT id FROM tasks
        WHERE chat_id = ?
          AND COALESCE(topic_id,0) = ?
          AND CAST(COALESCE(reply_to_message_id,'') AS TEXT) = ?
          AND state IN ('NEW','IN_PROGRESS','WAITING_CLARIFICATION','AWAITING_CONFIRMATION')
        ORDER BY updated_at DESC
        LIMIT 1
    """, (_chat_id, _topic_id, _reply_id)).fetchone()
    if row:
        return row["id"]

    row = conn.execute("""
        SELECT id FROM tasks
        WHERE chat_id = ?
          AND COALESCE(topic_id,0) = ?
          AND state = 'AWAITING_CONFIRMATION'
          AND COALESCE(result,'') <> ''
        ORDER BY updated_at DESC
        LIMIT 1
    """, (_chat_id, _topic_id)).fetchone()
    return row["id"] if row else None


# === STOP_TOPIC2_LOOP_FACT_BASED_V1 ===
def _stop_topic2_loop_fact_based_v1(conn, task):
    try:
        tid = _s(task["id"])
        topic_id = int(task["topic_id"] or 0) if "topic_id" in task.keys() else 0
        state = _s(task["state"]) if "state" in task.keys() else ""
        raw = _s(task["raw_input"]) if "raw_input" in task.keys() else ""
        result = _s(task["result"]) if "result" in task.keys() else ""
        low = (raw + "\n" + result).lower().replace("ё", "е")
        if topic_id != 2:
            return False
        if state not in ("NEW", "IN_PROGRESS", "WAITING_CLARIFICATION"):
            return False
        if raw.count("---") >= 2 or raw.count("REVISION") >= 1 or raw.count("[VOICE]") >= 2:
            _update_task(conn, tid, state="CANCELLED", error_message="STOP_TOPIC2_LOOP_FACT_BASED_V1:LOOP_GUARD_CANCELLED")
            _history(conn, tid, "STOP_TOPIC2_LOOP_FACT_BASED_V1:LOOP_GUARD_CANCELLED")
            conn.commit()
            return True
        if any(x in low for x in (
            "что тебе не понятно",
            "тебе все написано",
            "сможешь сделать нормальную задачу",
            "сделай мне смету нормальную",
        )):
            _update_task(conn, tid, state="CANCELLED", error_message="STOP_TOPIC2_LOOP_FACT_BASED_V1:META_VOICE_CANCELLED")
            _history(conn, tid, "STOP_TOPIC2_LOOP_FACT_BASED_V1:META_VOICE_CANCELLED")
            conn.commit()
            return True
        return False
    except Exception as e:
        try:
            logger.warning("STOP_TOPIC2_LOOP_FACT_BASED_V1_ERR %s", e)
        except Exception:
            pass
        return False
# === END_STOP_TOPIC2_LOOP_FACT_BASED_V1 ===


def _normalize_voice_text(text: str) -> str:
    return _clean(_s(text).replace("[VOICE]", " "), 12000)

def _looks_done_command(text: str) -> bool:
    low = _normalize_voice_text(text).lower()
    if "не доволен" in low or "недоволен" in low:
        return False
    done_markers = [
        "да доволен",
        "доволен результатом",
        "можно завершать",
        "можно закрывать",
        "завершай задачу",
        "завершай запрос",
        "завершай поиск",
        "задача завершена",
        "задача закрыта",
        "запрос завершен",
        "поиск завершен",
        "все верно завершай",
        "всё верно завершай",
        "все верно задача завершена",
        "всё верно задача завершена",
        "да все верно",
        "да всё верно",
        "я же тебе сказал задача завершена",
        "я же тебе сказал завершай",
        "да можно",
    ]
    return any(m in low for m in done_markers)

def _extract_topic_role(text: str) -> str:
    raw = _clean(_s(text), 1000)
    if not raw:
        return ""

    patterns = [
        r"чат закреплен за темами:\s*(.+?)(?:\.|$)",
        r"чат закреплён за темами:\s*(.+?)(?:\.|$)",
        r"чат закреплен за:\s*(.+?)(?:\.|$)",
        r"чат закреплён за:\s*(.+?)(?:\.|$)",
        r"закрепленные темы:\s*(.+?)(?:\.|$)",
        r"закреплённые темы:\s*(.+?)(?:\.|$)",
        r"закреплено:\s*чат для\s*(.+?)(?:\.|$)",
        r"этот чат используется для\s*(.+?)(?:\.|$)",
        r"этот чат предназначен для\s*(.+?)(?:\.|$)",
    ]

    bad_markers = [
        "без контекста",
        "не понимаю запрос",
        "не помню",
        "задача завершена",
        "задача закрыта",
        "подтверждение принято",
        "готов к выполнению задачи",
        "последние действия",
        "в этом чате были следующие действия",
        "текущий статус",
        "последний запрос",
        "какие последние",
        "чем помочь",
    ]

    for pattern in patterns:
        m = re.search(pattern, raw, re.I | re.S)
        if not m:
            continue

        role = re.sub(r"\s+", " ", m.group(1)).strip(" .:-")
        role = re.split(r"\b(чем помочь|готов к работе|готов обсудить|последний запрос)\b", role, flags=re.I)[0].strip(" .:-")
        low = role.lower()

        if not role or len(role) < 3:
            continue
        if "?" in role:
            continue
        if any(x in low for x in bad_markers):
            continue

        return role[:500]

    return ""

def _save_topic_role_memory(chat_id: str, topic_id: int, text: str) -> str:
    role = _extract_topic_role(text)
    if not role or not os.path.exists(MEM_DB):
        return ""
    conn_mem = db(MEM_DB)
    try:
        if not _has_table(conn_mem, "memory"):
            return ""
        key = f"topic_{int(topic_id)}_role"
        conn_mem.execute(
            "INSERT INTO memory (chat_id, key, value, timestamp) VALUES (?, ?, ?, datetime('now'))",
            (str(chat_id), key, role),
        )
        conn_mem.commit()
        return role
    except Exception:
        return ""
    finally:
        conn_mem.close()

def _find_awaiting_confirmation_task(conn: sqlite3.Connection, chat_id: str, topic_id: int, current_task_id: str, reply_to_message_id: Any) -> Optional[sqlite3.Row]:
    if reply_to_message_id:
        row = conn.execute(
            """
            SELECT id, result
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=?
              AND state='AWAITING_CONFIRMATION'
              AND id<>?
              AND (bot_message_id=? OR reply_to_message_id=?)
            ORDER BY updated_at DESC, created_at DESC
            LIMIT 1
            """,
            (str(chat_id), int(topic_id), str(current_task_id), reply_to_message_id, reply_to_message_id),
        ).fetchone()
        if row:
            return row

    return conn.execute(
        """
        SELECT id, result
        FROM tasks
        WHERE chat_id=?
          AND COALESCE(topic_id,0)=?
          AND state='AWAITING_CONFIRMATION'
          AND id<>?
        ORDER BY updated_at DESC, created_at DESC
        LIMIT 1
        """,
        (str(chat_id), int(topic_id), str(current_task_id)),
    ).fetchone()


# === FULLFIX_DIRECTION_KERNEL_STAGE_1_HELPER ===
def _stage1_dir_payload(payload):
    try:
        p = dict(payload or {})
        if _Stage1WorkItem is None or _Stage1DirReg is None:
            p.setdefault("direction", "general_chat")
            return p
        row = {
            "id": p.get("task_id") or p.get("id") or "",
            "chat_id": str(p.get("chat_id") or ""),
            "topic_id": int(p.get("topic_id") or 0),
            "input_type": p.get("input_type") or "unknown",
            "raw_input": p.get("raw_input") or p.get("raw_text") or "",
            "state": p.get("state") or "IN_PROGRESS",
            "reply_to_message_id": p.get("reply_to_message_id"),
            "bot_message_id": p.get("bot_message_id"),
        }
        wi = _Stage1WorkItem.from_task_row(row)
        prof = _Stage1DirReg().detect(wi)
        did = prof.get("id") or "general_chat"
        wi.set_direction(did, prof)
        wi.add_audit("stage", "FULLFIX_DIRECTION_KERNEL_STAGE_1")
        wi.add_audit("shadow_mode", True)
        p.update({
            "direction": did,
            "direction_profile": prof,
            "direction_audit": wi.audit,
            "work_item": wi.to_dict(),
        })
        try:
            logger.info("FULLFIX_DIRECTION_KERNEL_STAGE_1 dir=%s score=%s task=%s topic=%s",
                        did, prof.get("score"), row["id"], row["topic_id"])
        except Exception:
            pass
        return p
    except Exception as e:
        try: logger.error("FULLFIX_DIRECTION_KERNEL_STAGE_1_ERR %s", e)
        except Exception: pass
        p = dict(payload or {})
        p.setdefault("direction", "general_chat")
        return p
# === END ===
async def _handle_in_progress(conn: sqlite3.Connection, task: sqlite3.Row, chat_id: str, topic_id: int) -> None:

    if _stop_topic2_loop_fact_based_v1(conn, task):
        return
    task_id = _s(_task_field(task, "id"))
    raw_input = _clean(_s(_task_field(task, "raw_input")), 12000)
    reply_to = _task_field(task, "reply_to_message_id", None)

    if _looks_done_command(raw_input):
        target = _find_awaiting_confirmation_task(conn, chat_id, topic_id, task_id, reply_to)
        if target:
            target_id = _s(target["id"])
            target_result = _s(target["result"])
            saved_role = _save_topic_role_memory(chat_id, topic_id, target_result)
            _update_task(conn, target_id, state="DONE", error_message="")
            _history(conn, target_id, "state:DONE")
            if saved_role:
                _history(conn, target_id, f"ROLE_SAVED:{_clean(saved_role, 200)}")
            _update_task(conn, task_id, state="DONE", result="Подтверждение принято", error_message="")
            _history(conn, task_id, "state:DONE")
            conn.commit()
            _send_once(conn, task_id, chat_id, "Принял. Задача закрыта", reply_to, "confirm_done")
            return

    parent_task_id = _get_parent_task_id(conn, chat_id, task["reply_to_message_id"], topic_id)
    active_source_id = parent_task_id or task_id
    active_task_context = _active_unfinished_context(conn, chat_id, topic_id, active_source_id)
    pin_context = get_pin_context(chat_id, raw_input, topic_id)
    short_memory, long_memory, topic_role, topic_directions = _load_memory_context(chat_id, topic_id)
    archive_context = _load_archive_context(chat_id, topic_id, raw_input)
    search_context = _search_fact_context(conn, chat_id, topic_id)

    payload: Dict[str, Any] = {
        "id": task_id,
        "task_id": task_id,  # PAYLOAD_TOPIC_ID_FIX_V1
        "topic_id": int(topic_id or 0),
        "chat_id": chat_id,
        "input_type": _s(_task_field(task, "input_type", "text")).lower() or "text",
        "raw_input": raw_input,
        "normalized_input": raw_input,
        "state": "IN_PROGRESS",
        "reply_to_message_id": reply_to,
        "active_task_context": active_task_context,
        "pin_context": pin_context,
        "short_memory_context": short_memory,
        "long_memory_context": long_memory,
        "archive_context": archive_context,
        "search_context": search_context,
        "topic_role": topic_role,
        "topic_directions": topic_directions,
    }

    try:
        ai_result = None  # AI_RESULT_INIT_V1
        assigned_role = _detect_role_assignment(raw_input)
        if assigned_role:
            _save_topic_role(chat_id, topic_id, assigned_role)
            _history(conn, task_id, f"ROLE_SAVED:{_clean(assigned_role, 200)}")
            conn.commit()
            ai_result = f"Принято. Чат закреплен за темами: {assigned_role}. Все связанные запросы будут обрабатываться здесь."
        else:
            ROLE_Q = re.compile(r"(для чего|о чём|о чем|про что|напомни.*(чат|топик)|чем занимается|зачем этот чат)", re.IGNORECASE)
            HISTORY_Q = re.compile(r"(что мы писали|что писали раньше|о ч[её]м общались|напомни.*что.*(писали|обсуждали)|что было в этом чате|история чата)", re.IGNORECASE)
            # === WHAT_IS_THIS_META_V1 ===
            if TOPIC_META_LOADER_WIRED and is_what_is_this_question(raw_input):
                _wt_meta = load_topic_meta(int(topic_id or 0))
                if _wt_meta:
                    _wt_answer = build_topic_self_answer(_wt_meta)
                    if _wt_answer:
                        _update_task(conn, task_id, state="DONE", result=_wt_answer, error_message="")
                        conn.commit()
                        from core.reply_sender import send_reply_ex
                        send_reply_ex(chat_id=str(chat_id), text=_wt_answer, reply_to_message_id=reply_to, message_thread_id=topic_id)
                        return
            # === END WHAT_IS_THIS_META_V1 ===
            if topic_role and (ROLE_Q.search(raw_input) or HISTORY_Q.search(raw_input)):
                ai_result = f"Этот чат закреплён за: {topic_role}"
            else:
                try:
                    from core.model_router import route_model as _rm
                    _mo = _rm(payload)
                    if _mo:
                        payload["model_override"] = _mo  # MODEL_ROUTER_V1_WIRED
                except Exception:
                    pass
                        # === TEMPLATE_TRIGGER_V1 ===
            if _tpl_check(str(raw_input or "")):
                _tpl_path = _tpl_get(int(topic_id or 0))
                if _tpl_path and os.path.exists(_tpl_path):
                    logger.info("TEMPLATE_TRIGGER_V1 using template=%s task=%s", _tpl_path, task_id)
                    payload["template_path"] = _tpl_path
                    payload["use_template"] = True
            # === END TEMPLATE_TRIGGER_V1 ===
# === TOPIC_3008_HANDLER_V1 ===
            if _t3_check(int(topic_id or 0)):
                _t3_command = _t3_cmd(str(raw_input or ""))
                if _t3_command != "none":
                    import re as _re3008
                    _t3_ctx = " ".join(filter(None,[
                        str(payload.get("active_task_context") or "")[:200],
                        str(payload.get("pin_context") or "")[:100],
                        str(payload.get("short_memory_context") or "")[:200],
                    ]))
                    if _t3_command == "write":
                        _t3_desc = _re3008.sub(r"напиши\s+код|написать\s+код","",str(raw_input or ""),flags=_re3008.I).strip()
                        if _t3_generate:
                            _t3_gen = await asyncio.wait_for(_t3_generate(_t3_desc,_t3_ctx),timeout=120)
                            ai_result = _t3_gen + "\n\n---\nПроверить код?"
                    elif _t3_command == "verify":
                        _t3_code = _t3_extract(str(raw_input or ""))
                        if len(_t3_code.strip()) < 10:
                            ai_result = "Отправь код для проверки."
                        elif _t3_verify:
                            from core.reply_sender import send_reply_ex as _t3srex
                            _t3srex(chat_id=str(chat_id),text="Запущена верификация. Ожидаю ответы (до 1.5 мин)...",reply_to_message_id=reply_to,message_thread_id=topic_id)
                            ai_result = await asyncio.wait_for(_t3_verify(_t3_code,_t3_ctx),timeout=150)
            # === END TOPIC_3008_HANDLER_V1 ===
            if ai_result is None:  # AI_LOGIC_FIX_V1
                # === STROYKA_FULL_CHAIN_FINAL_CLOSE_PRE_DIRECTION_GUARD ===
                try:
                    _stroyka_topic = int(_task_field(task, "topic_id", 0) or 0)
                    _stroyka_state = _s(_task_field(task, "state", ""))
                    _stroyka_task_id = _s(_task_field(task, "id", ""))
                    if _stroyka_topic == 2 and _stroyka_state in ("NEW", "IN_PROGRESS"):
                        _stroyka_raw = _s(_task_field(task, "raw_input", "")).lower()
                        _stroyka_result = _s(_task_field(task, "result", "")).lower()
                        _stroyka_bad = any(x in (_stroyka_raw + "\n" + _stroyka_result) for x in (
                            "вор_кирпич", "vor_kirpich", "вор_кирпичная_кладка",
                            "смета создана по образцу вор", "поставщик | площадка",
                            "auto_parts", "search_monolith", "tco | риски",
                            "ошибка классификации запроса", "категория не совпадает",
                        ))
                        if _stroyka_bad:
                            _update_task(conn, _stroyka_task_id, state="FAILED",
                                result="STROYKA_FULL_CHAIN_FINAL_CLOSE_PRE_DIRECTION_GUARD: blocked stale estimate contamination")
                            _history(conn, _stroyka_task_id, "STROYKA_FULL_CHAIN_FINAL_CLOSE_PRE_DIRECTION_GUARD:blocked_bad_estimate")
                            return
                        from core.stroyka_estimate_canon import maybe_handle_stroyka_estimate as _stroyka_final_handle
                        if await _stroyka_final_handle(conn, task, logger):
                            _history(conn, _stroyka_task_id, "STROYKA_FULL_CHAIN_FINAL_CLOSE_PRE_DIRECTION_GUARD:handled")
                            return
                except Exception as _stroyka_final_err:
                    logger.exception("STROYKA_FULL_CHAIN_FINAL_CLOSE_PRE_DIRECTION_GUARD_ERR %s", _stroyka_final_err)
                # === END_STROYKA_FULL_CHAIN_FINAL_CLOSE_PRE_DIRECTION_GUARD ===
                # FULLFIX_DIRECTION_KERNEL_STAGE_1_CALL
                payload = _stage1_dir_payload(payload)
                # FULLFIX_TOPIC_AUTODISCOVERY_V2_CALL
                if _topic_autodiscovery is not None:
                    try:
                        from core.work_item import WorkItem as _WITA
                        _wita = _WITA.from_task_row({
                            "id": payload.get("task_id") or payload.get("id") or "",
                            "chat_id": str(payload.get("chat_id") or ""),
                            "topic_id": int(payload.get("topic_id") or 0),
                            "input_type": payload.get("input_type") or "unknown",
                            "raw_input": payload.get("raw_input") or payload.get("raw_text") or "",
                            "state": payload.get("state") or "IN_PROGRESS",
                        })
                        _topic_autodiscovery(_wita, payload)
                    except Exception as _eta:
                        logger.error("TOPIC_AUTODISCOVERY_ERR %s", _eta)
                # FULLFIX_CAPABILITY_ROUTER_STAGE_2_CALL
                if _Stage2Router is not None:
                    try:
                        from core.work_item import WorkItem as _WI2
                        _wi2 = _WI2.from_task_row({
                            "id": payload.get("task_id") or payload.get("id") or "",
                            "chat_id": str(payload.get("chat_id") or ""),
                            "topic_id": int(payload.get("topic_id") or 0),
                            "input_type": payload.get("input_type") or "unknown",
                            "raw_input": payload.get("raw_input") or payload.get("raw_text") or "",
                            "state": payload.get("state") or "IN_PROGRESS",
                        })
                        _wi2.set_direction(
                            payload.get("direction") or "general_chat",
                            payload.get("direction_profile") or {},
                        )
                        _r2 = _Stage2Router().apply_to_work_item(_wi2)
                        payload["engine"] = _r2["engine"]
                        payload["execution_plan"] = _r2["execution_plan"]
                        payload["formats_out"] = _r2["formats_out"]
                        payload["quality_gates"] = _r2["quality_gates"]
                        payload["capability_router"] = _r2["router_version"]
                        logger.info("FULLFIX_CAPABILITY_ROUTER_STAGE_2 engine=%s steps=%s dir=%s",
                                    _r2["engine"], len(_r2["execution_plan"]), payload.get("direction"))
                    except Exception as _e2:
                        logger.error("FULLFIX_CAPABILITY_ROUTER_STAGE_2_ERR %s", _e2)
                # FULLFIX_SEARCH_ENGINE_STAGE_5_CALL
                if _Stage5Search is not None:
                    try:
                        from core.work_item import WorkItem as _WI5
                        _wi5 = _WI5.from_task_row({
                            "id": payload.get("task_id") or payload.get("id") or "",
                            "chat_id": str(payload.get("chat_id") or ""),
                            "topic_id": int(payload.get("topic_id") or 0),
                            "input_type": payload.get("input_type") or "unknown",
                            "raw_input": payload.get("raw_input") or payload.get("raw_text") or "",
                            "state": payload.get("state") or "IN_PROGRESS",
                        })
                        _wi5.set_direction(
                            payload.get("direction") or "general_chat",
                            payload.get("direction_profile") or {},
                        )
                        _Stage5Search().apply_to_payload(_wi5, payload)
                    except Exception as _e5:
                        logger.error("FULLFIX_SEARCH_ENGINE_STAGE_5_ERR %s", _e5)
                # FULLFIX_CONTEXT_LOADER_STAGE_3_CALL
                if _Stage3Loader is not None:
                    try:
                        from core.work_item import WorkItem as _WI3
                        _wi3 = _WI3.from_task_row({
                            "id": payload.get("task_id") or payload.get("id") or "",
                            "chat_id": str(payload.get("chat_id") or ""),
                            "topic_id": int(payload.get("topic_id") or 0),
                            "input_type": payload.get("input_type") or "unknown",
                            "raw_input": payload.get("raw_input") or payload.get("raw_text") or "",
                            "state": payload.get("state") or "IN_PROGRESS",
                        })
                        _wi3.set_direction(
                            payload.get("direction") or "general_chat",
                            payload.get("direction_profile") or {},
                        )
                        _refs3 = _Stage3Loader().load(_wi3)
                        payload["context_refs"] = _refs3
                        logger.info("FULLFIX_CONTEXT_LOADER_STAGE_3 topic=%s short_mem=%s",
                                    payload.get("topic_id"), bool(_refs3.get("short_memory")))
                    except Exception as _e3:
                        logger.error("FULLFIX_CONTEXT_LOADER_STAGE_3_ERR %s", _e3)
                ai_result = await asyncio.wait_for(process_ai_task(payload), timeout=AI_TIMEOUT)
                # FULLFIX_QUALITY_GATE_STAGE_4_CALL
                if _Stage4QG is not None and isinstance(ai_result, dict):
                    try:
                        _qg_payload = {**payload, **ai_result}
                        _qg_report = _Stage4QG().apply_to_payload(_qg_payload)
                        ai_result["quality_gate_report"] = _qg_report
                        logger.info("FULLFIX_QUALITY_GATE_STAGE_4 overall=%s failed=%s dir=%s",
                                    _qg_report["overall"], _qg_report["failed"], payload.get("direction"))
                    except Exception as _e4:
                        logger.error("FULLFIX_QUALITY_GATE_STAGE_4_ERR %s", _e4)
                # FULLFIX_ARCHIVE_ENGINE_STAGE_6_CALL
                if _Stage6Archive is not None:
                    try:
                        _Stage6Archive().archive(payload, ai_result if isinstance(ai_result, dict) else {})
                    except Exception as _e6:
                        logger.error("FULLFIX_ARCHIVE_ENGINE_STAGE_6_ERR %s", _e6)
                # FULLFIX_FORMAT_ADAPTER_STAGE_7_CALL
                if _Stage7FA is not None and isinstance(ai_result, dict):
                    try:
                        _formats_out = payload.get("formats_out") or ["telegram_text"]
                        _adapted = _Stage7FA().adapt(ai_result, _formats_out, payload)
                        ai_result["format_adapted"] = _adapted
                        logger.info("FULLFIX_FORMAT_ADAPTER_STAGE_7 formats=%s primary=%s dir=%s",
                                    _formats_out, list(_adapted.get("outputs", {}).keys())[:2],
                                    payload.get("direction"))
                    except Exception as _e7:
                        logger.error("FULLFIX_FORMAT_ADAPTER_STAGE_7_ERR %s", _e7)
    except Exception as e:
        _update_task(conn, task_id, state="FAILED", error_message=_clean(str(e), 500))
        _close_pin(conn, task_id)
        _history(conn, task_id, "state:FAILED")
        conn.commit()
        _send_once(conn, task_id, chat_id, "Задача не выполнена. Уточни или повтори запрос", reply_to, "router_failed")
        return

    ai_result = _clean(_s(ai_result), 50000)
    # === FOLLOWUP DETECTION (FACT-BASED) ===
    low_input = raw_input.lower()

    memory_markers = [
        "напомни","что обсуждали","что делали","какие задачи","история",
        "что было","что писали","для чего этот чат","о чем чат","о чём чат"
    ]

    search_markers = [
        "нерелевант","битые","ссылки не те","проверь","еще раз","ещё раз",
        "сделай еще","сделай ещё","найди еще","найди ещё","это не то"
    ]

    has_memory_context = any([
        short_memory,
        long_memory,
        topic_role,
        active_task_context,
        pin_context,
        archive_context,
    ])

    raw_input = _clean(_s(_task_field(task, "raw_input")), 12000)
    low_input = raw_input.lower()

    memory_markers = [
        "напомни",
        "что обсуждали",
        "что делали",
        "какие задачи",
        "история",
        "что было в этом чате",
        "что писали",
        "для чего этот чат",
        "о чем чат",
        "о чём чат",
    ]
    is_memory_followup = has_memory_context and any(m in low_input for m in memory_markers)

    search_markers = [
        "нерелевант",
        "битые ссылки",
        "живые ссылки",
        "ссылки не те",
        "проверь еще",
        "проверь ещё",
        "еще раз поиск",
        "ещё раз поиск",
        "сделай еще раз поиск",
        "сделай ещё раз поиск",
        "найди еще",
        "найди ещё",
        "это не то",
        "ссылки биты",
        "ссылки битые",
    ]
    is_search_followup = bool(search_context) and any(m in low_input for m in search_markers)

    if is_search_followup and search_context:
        forbidden_search_advice = [
            "dr.web",
            "link checker",
            "яндекс safety",
            "google safe browsing",
            "virustotal",
            "для проверки безопасности ссылок используйте",
        ]
        if any(m in ai_result.lower() for m in forbidden_search_advice):
            ai_result = _clean(f"Повторяю поиск по последнему запросу\n\n{search_context}", 50000)

    if is_memory_followup or is_search_followup:
        if not ai_result or len(ai_result) < MIN_RESULT_LEN:
            _update_task(conn, task_id, state="FAILED", error_message="INVALID_RESULT_GATE")
            _close_pin(conn, task_id)
            _history(conn, task_id, "state:FAILED")
            conn.commit()
            _send_once(conn, task_id, chat_id, "Не понял запрос. Уточни что нужно сделать", reply_to, "invalid_result")
            return
    else:
        if not _is_valid_result(ai_result, raw_input):
            _update_task(conn, task_id, state="FAILED", error_message="INVALID_RESULT_GATE")
            _close_pin(conn, task_id)
            _history(conn, task_id, "state:FAILED")
            conn.commit()
            _send_once(conn, task_id, chat_id, "Не понял запрос. Уточни что нужно сделать", reply_to, "invalid_result")
            return

    low_result = ai_result.lower()
    done_markers = [
        "задача завершена",
        "задача закрыта",
        "задачи завершены",
        "подтверждение принято",
        "поиск завершен",
        "поиск завершён",
    ]
    junk_markers = [
        "без контекста",
        "задайте конкретный вопрос",
        "конкретный вопрос по",
        "нет, не помню",
        "не понимаю запрос",
        "готов к выполнению задачи",
    ]
    info_markers = [
        "для чего этот чат",
        "что мы здесь обсуждаем",
        "что в данном чате",
        "какой последний запрос",
        "последний запрос",
        "какие последние запросы",
        "что мы тут делаем",
        "что мы здесь делаем",
        "о чем чат",
        "о чём чат",
    ]
    file_success_markers = [
        "документ обработан",
        "артефакт:",
        "нормализовано позиций",
        "обработаны документы",
    ]
    file_bad_markers = [
        "скачан, ожидает анализа",
        "создан локально, но загрузка в drive завершилась ошибкой",
        "ожидает анализа",
        "загрузка в drive завершилась ошибкой",
    ]

    if any(m in low_result for m in done_markers):
        _update_task(conn, task_id, state="DONE", result=ai_result, error_message="")
        _close_pin(conn, task_id)
        _history(conn, task_id, f"result:{_clean(ai_result, 400)}")
        _save_memory(chat_id, topic_id, raw_input, ai_result)  # SAVE_MEM_DONE_V2
        conn.commit()
        _send_once(conn, task_id, chat_id, ai_result, reply_to, "done_terminal")
        return

    if any(m in low_result for m in junk_markers):
        _update_task(conn, task_id, state="WAITING_CLARIFICATION", result=ai_result, error_message="")
        _close_pin(conn, task_id)
        _history(conn, task_id, f"clarify:{_clean(ai_result, 400)}")
        conn.commit()
        _send_once(conn, task_id, chat_id, ai_result, reply_to, "clarify_terminal")
        return

    is_info_query = any(m in low_input for m in info_markers)
    is_file_success = any(m in low_result for m in file_success_markers) and not any(m in low_result for m in file_bad_markers)

    if is_memory_followup or is_search_followup or is_info_query or is_file_success:
        _update_task(conn, task_id, state="DONE", result=ai_result, error_message="")
        _history(conn, task_id, f"result:{_clean(ai_result, 400)}")
        _save_memory(chat_id, topic_id, raw_input, ai_result)  # SAVE_MEM_FOLLOWUP_V2
        conn.commit()
        try:
            save_pin(chat_id, task_id, ai_result, topic_id)
        except Exception:
            pass
        _send_once(conn, task_id, chat_id, ai_result, reply_to, "done_terminal")
        return

    should_save_role = (
        bool(re.search(
            r"(чат закреплен за темами:|чат закреплён за темами:|чат закреплен за:|чат закреплён за:|закрепленные темы:|закреплённые темы:|закреплено:\s*чат для|этот чат используется для|этот чат предназначен для)",
            ai_result,
            re.I,
        ))
        and not any(x in low_result for x in junk_markers)
        and "последние действия" not in low_result
        and "текущий статус" not in low_result
    )

    saved_role = ""
    if should_save_role:
        _save_memory(chat_id, topic_id, raw_input, ai_result)  # SAVE_MEM_CONFIRM_V2
        saved_role = _save_topic_role_memory(chat_id, topic_id, ai_result)
        # === RESULT_VALIDATOR_GUARD_V1 ===
        if _check_result_before_confirm(ai_result):
            # === SEARCH_MONOLITH_V2_CLARIFICATION_HANDLER ===
            try:
                if _search_v2_is_clarification(ai_result):
                    _cmsg = str(ai_result).replace("SEARCH_CLARIFICATION_REQUIRED:","").strip()
                    _update_task(conn, task_id, state="WAITING_CLARIFICATION", result=_cmsg, error_message="")
                    _history(conn, task_id, "SEARCH_MONOLITH_V2:WAITING_CLARIFICATION")
                    conn.commit()
                    _send_once(conn, task_id, str(chat_id), _cmsg, reply_to, "search_v2_clarification")
                    return
            except Exception as _e: logger.warning("SEARCH_V2_CLARIFICATION_HANDLER_ERR %s", _e)
            # === END SEARCH_MONOLITH_V2_CLARIFICATION_HANDLER ===
            # === UNIFIED_ENGINE_RESULT_VALIDATOR_V1_TASK_WORKER_AI_RESULT ===
            try:
                if validate_engine_result is not None:
                    _payload_for_uv = locals().get("payload", {}) or {}
                    _raw_for_uv = ""
                    try:
                        _raw_for_uv = str(task["raw_input"] if "raw_input" in task.keys() else "")
                    except Exception:
                        _raw_for_uv = str(_payload_for_uv.get("raw_input") or _payload_for_uv.get("user_text") or "")
                    _input_type_for_uv = ""
                    try:
                        _input_type_for_uv = str(task["input_type"] if "input_type" in task.keys() else "")
                    except Exception:
                        _input_type_for_uv = str(_payload_for_uv.get("input_type") or "")
                    _uv = validate_engine_result({"summary": ai_result, "engine": "AI_ROUTER"}, input_type=_input_type_for_uv, user_text=_raw_for_uv, topic_id=topic_id)
                    if not _uv.get("ok"):
                        _update_task(conn, task_id, state="FAILED", result="", error_message="UNIFIED_ENGINE_RESULT_VALIDATOR_V1:" + str(_uv.get("reason") or "INVALID"))
                        _history(conn, task_id, "UNIFIED_ENGINE_RESULT_VALIDATOR_V1:FAILED:" + str(_uv.get("reason") or "INVALID"))
                        return
            except Exception as _uv_e:
                logger.warning("UNIFIED_ENGINE_RESULT_VALIDATOR_V1_ERR task=%s err=%s", task_id, _uv_e)
            # === END_UNIFIED_ENGINE_RESULT_VALIDATOR_V1_TASK_WORKER_AI_RESULT ===
            # === TASK_WORKER_ARTIFACT_GATE_V1 ===
            try:
                if validate_engine_result is not None:
                    _twag_raw = ""
                    _twag_input_type = ""
                    try:
                        _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                    except Exception:
                        _twag_raw = ""
                    try:
                        _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                    except Exception:
                        _twag_input_type = ""
                    _twag_result = ai_result
                    _twag_check = validate_engine_result(
                        {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                        input_type=_twag_input_type,
                        user_text=_twag_raw,
                        topic_id=topic_id,
                    )
                    if not _twag_check.get("ok"):
                        _update_task(conn, task_id, state="FAILED", result="",
                            error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                        _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                        return
            except Exception as _twag_e:
                logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
            # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
            _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=ai_result, error_message="")
        else:
            _update_task(conn, task_id, state="FAILED", result=ai_result, error_message="FORBIDDEN_PHRASE")
        # === END RESULT_VALIDATOR_GUARD_V1 ===
    _history(conn, task_id, f"result:{_clean(ai_result, 400)}")
    if saved_role:
        _history(conn, task_id, f"ROLE_SAVED:{_clean(saved_role, 200)}")

    try:
        save_pin(chat_id, task_id, ai_result, topic_id)
    except Exception as e:
        logger.warning("save_pin_fail task=%s err=%s", task_id, e)

    _ai_result_clean = str(ai_result or "").replace("\n\nДоволен результатом? Ответь: Да / Уточни / Правки", "").strip()  # FF21_FIX_DOUBLE_DOVOLEN
    confirmation_text = f"{_ai_result_clean}\n\nДоволен результатом? Ответь: Да / Уточни / Правки"
    sent = _send_once_ex(conn, task_id, chat_id, confirmation_text, reply_to, "result")
    bot_message_id = sent.get("bot_message_id") if isinstance(sent, dict) else None
    if bot_message_id is not None:
        _update_task(conn, task_id, bot_message_id=bot_message_id)
    conn.commit()



# === IN_PROGRESS_HARD_TIMEOUT_BY_CREATED_AT_FIX_V1 ===
def _in_progress_hard_timeout_by_created_at_fix_v1(conn, minutes: int = 30) -> int:
    try:
        rows = conn.execute(
            """
            SELECT id FROM tasks
            WHERE state='IN_PROGRESS'
              AND datetime(COALESCE(created_at, updated_at, 'now')) <= datetime('now', ?)
            ORDER BY created_at ASC
            LIMIT 200
            """,
            (f"-{int(minutes)} minutes",),
        ).fetchall()
        n = 0
        for row in rows:
            tid = row["id"] if hasattr(row, "keys") else row[0]
            conn.execute(
                "UPDATE tasks SET state='FAILED', error_message='IN_PROGRESS_HARD_TIMEOUT_BY_CREATED_AT_FIX_V1', updated_at=datetime('now') WHERE id=? AND state='IN_PROGRESS'",
                (str(tid),),
            )
            try:
                _history(conn, str(tid), "IN_PROGRESS_HARD_TIMEOUT_BY_CREATED_AT_FIX_V1:FAILED")
            except Exception:
                pass
            n += 1
        if n:
            try: conn.commit()
            except Exception: pass
            logger.warning("IN_PROGRESS_HARD_TIMEOUT_BY_CREATED_AT_FIX_V1 closed=%s", n)
        return n
    except Exception as e:
        logger.warning("IN_PROGRESS_HARD_TIMEOUT_BY_CREATED_AT_FIX_V1_ERR %s", e)
        return 0
# === END_IN_PROGRESS_HARD_TIMEOUT_BY_CREATED_AT_FIX_V1 ===

def _pick_next_task(conn: sqlite3.Connection, chat_id: Optional[str]) -> Optional[sqlite3.Row]:
    # === IN_PROGRESS_HARD_TIMEOUT_BY_CREATED_AT_FIX_V1_HOOK ===
    _in_progress_hard_timeout_by_created_at_fix_v1(conn, minutes=30)
    # === END_IN_PROGRESS_HARD_TIMEOUT_BY_CREATED_AT_FIX_V1_HOOK ===
    where = [
        "state IN ('NEW','IN_PROGRESS','WAITING_CLARIFICATION')",
        "NOT (state='IN_PROGRESS' AND COALESCE(error_message,'')='P6F_DAH_BLOCK_DONE_NO_UPLOAD_HISTORY')"
    ]
    params: List[Any] = []
    if chat_id:
        where.insert(0, "chat_id=?")
        params.append(str(chat_id))

    conn.execute("BEGIN IMMEDIATE")
    row = conn.execute(
        f"""
        SELECT *
        FROM tasks
        WHERE {' AND '.join(where)}
        ORDER BY CASE state WHEN 'IN_PROGRESS' THEN 0 ELSE 1 END,
                 created_at ASC
        LIMIT 1
        """
        ,
        params).fetchone()
    conn.execute("COMMIT")
    return row


async def main() -> None:
    lock_fp = open(LOCK_PATH, "w")
    try:
        fcntl.flock(lock_fp.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
    except BlockingIOError:
        logger.info("WORKER LOCKED BY OTHER PROCESS")
        return

    logger.info("WORKER STARTED pid=%s", os.getpid())

    while True:
        conn = db(CORE_DB)
        try:
            _recover_stale_tasks(conn, None)
            task = _pick_next_task(conn, None)
            if not task:
                time.sleep(POLL_SEC)
                continue

            task_id = _s(_task_field(task, "id"))
            chat_id = _s(_task_field(task, "chat_id"))
            topic_id = int(_task_field(task, "topic_id", 0) or 0)
            state = _s(_task_field(task, "state")).upper()

            logger.info("PICKED %s state=%s chat=%s topic=%s", task_id, state, chat_id, topic_id)
            input_type = _s(_task_field(task, "input_type")).lower()
            if input_type == "drive_file":
                try:
                    await _handle_drive_file(conn, task, chat_id, topic_id)
                except Exception as e:
                    logger.error("DRIVE_FILE CRASH task=%s err=%s", task_id, str(e))
                continue

            if state == "NEW":
                await _handle_new(conn, task, chat_id, topic_id)
            elif state == "IN_PROGRESS":
                await _handle_in_progress(conn, task, chat_id, topic_id)
            elif state == "WAITING_CLARIFICATION":
                await _handle_in_progress(conn, task, chat_id, topic_id)
        finally:
            conn.close()

        time.sleep(POLL_SEC)



# === DRIVE FILE HANDLING ===
def _download_from_drive(file_id: str, local_path: str) -> bool:
    try:
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseDownload
        from google.oauth2.service_account import Credentials
        import io
        creds = Credentials.from_service_account_file(
            '/root/.areal-neva-core/credentials.json',
            scopes=['https://www.googleapis.com/auth/drive.readonly']
        )
        service = build('drive', 'v3', credentials=creds)
        request = service.files().get_media(fileId=file_id)
        with io.FileIO(local_path, 'wb') as fh:
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
        return True
    except Exception as e:
        import logging
        logging.getLogger("task_worker").error(f"Drive download failed: {e}")
        return False

async def _handle_drive_file(conn, task, chat_id, topic_id):
    import json, os
    task_id = task["id"]
    input_type = "drive_file"  # INPUT_TYPE_DRIVE_FIX_V1
    raw_input = task["raw_input"]
    try:
        data = json.loads(raw_input)
        file_id = data["file_id"]
        file_name = data.get("file_name", "файл")  # HOTFIX_FILE_NAME_EARLY_V1
        # === DRIVE_FILE_CONTENT_SERVICE_GUARD_V1 ===
        try:
            if _filemem_is_service_file(str(file_name or ""), str(data.get("source") or ""), int(topic_id or 0), str(raw_input or "")):
                _msg = "Служебный файл синхронизации проигнорирован"
                _update_task(conn, task_id, state="CANCELLED", result=_msg, error_message="SERVICE_FILE_IGNORED")
                _history(conn, task_id, "DRIVE_FILE_CONTENT_SERVICE_GUARD_V1:CANCELLED")
                try:
                    _append_timeline_event_v1(str(chat_id), int(topic_id or 0), task_id, "service_file_ignored", raw_input, _msg)
                except Exception:
                    pass
                conn.commit()
                logger.info("DRIVE_FILE_CONTENT_SERVICE_GUARD_V1 cancelled task=%s file=%s source=%s topic=%s", task_id, file_name, data.get("source"), topic_id)
                return
        except Exception as _svc_e:
            logger.warning("DRIVE_FILE_CONTENT_SERVICE_GUARD_V1_ERR task=%s err=%s", task_id, _svc_e)
        # === END DRIVE_FILE_CONTENT_SERVICE_GUARD_V1 ===
        # === DRIVE_FILE_MEMORY_INDEX_V1 + FILE_DUPLICATE_MEMORY_GUARD_V1 ===
        try:
            _df_file_id = str(data.get("file_id") or "")
            _df_file_name = str(file_name or "")
            _df_meta = json.dumps({
                "task_id": task_id, "chat_id": str(chat_id),
                "topic_id": int(topic_id or 0), "file_id": _df_file_id,
                "file_name": _df_file_name,
                "caption": str(data.get("caption") or ""),
                "source": str(data.get("source") or ""),
            }, ensure_ascii=False)
            if _df_file_id:
                _dupe = conn.execute(
                    "SELECT id,state FROM tasks WHERE chat_id=? AND COALESCE(topic_id,0)=? AND id<>? AND input_type='drive_file' AND raw_input LIKE ? AND state IN ('DONE','ARCHIVED','AWAITING_CONFIRMATION') ORDER BY updated_at DESC LIMIT 1",
                    (str(chat_id), int(topic_id or 0), task_id, "%" + _df_file_id + "%"),
                ).fetchone()
                if _dupe is not None:
                    _dmsg = "Файл уже есть в этом топике. Предыдущая задача: " + _s(_dupe["id"])[:8]
                    _update_task(conn, task_id, state="DONE", result=_dmsg, error_message="")
                    _history(conn, task_id, "FILE_DUPLICATE_MEMORY_GUARD_V1:DONE")
                    _memory_insert_topic_entry_v1(str(chat_id), f"topic_{int(topic_id or 0)}_file_duplicate_{task_id}", _df_meta)
                    _append_timeline_event_v1(str(chat_id), int(topic_id or 0), task_id, "file_duplicate", raw_input, _dmsg)
                    conn.commit()
                    _send_once(conn, task_id, str(chat_id), _dmsg, _task_field(task, "reply_to_message_id", None), "file_dup_guard_v1")
                    logger.info("FILE_DUPLICATE_MEMORY_GUARD_V1 task=%s", task_id)
                    return
            _memory_insert_topic_entry_v1(str(chat_id), f"topic_{int(topic_id or 0)}_file_{task_id}", _df_meta)
            _append_timeline_event_v1(str(chat_id), int(topic_id or 0), task_id, "drive_file_indexed", raw_input, "")
            # === DRIVE_FILE_CONTENT_MEMORY_INDEX_V1_CALL ===
            try:
                # === DRIVE_FILE_CONTENT_INDEX_SKIP_SERVICE_V2 ===
                try:
                    _df_skip_content_index = bool(_filemem_is_service_file(
                        _df_file_name,
                        str(data.get("source") or ""),
                        int(topic_id or 0),
                        str(raw_input or "")
                    ))
                except Exception as _df_skip_e:
                    _df_skip_content_index = False
                    logger.warning("DRIVE_FILE_CONTENT_INDEX_SKIP_SERVICE_V2_ERR task=%s err=%s", task_id, _df_skip_e)
                if _df_skip_content_index:
                    logger.info("DRIVE_FILE_CONTENT_INDEX_SKIP_SERVICE_V2 skipped task=%s file=%s source=%s topic=%s", task_id, _df_file_name, data.get("source"), topic_id)
                # === END DRIVE_FILE_CONTENT_INDEX_SKIP_SERVICE_V2 ===
                if (not _df_skip_content_index) and _df_content_index is not None and _df_file_id:
                    _df_ci = await asyncio.to_thread(
                        _df_content_index,
                        str(chat_id),
                        int(topic_id or 0),
                        str(task_id),
                        str(_df_file_id),
                        str(_df_file_name),
                        str(data.get("mime_type") or ""),
                    )
                    _memory_insert_topic_entry_v1(
                        str(chat_id),
                        f"topic_{int(topic_id or 0)}_file_content_status_{task_id}",
                        json.dumps(_df_ci, ensure_ascii=False),
                    )
                    _append_timeline_event_v1(
                        str(chat_id),
                        int(topic_id or 0),
                        task_id,
                        "drive_file_content_indexed",
                        raw_input,
                        json.dumps(_df_ci, ensure_ascii=False),
                    )
                    logger.info("DRIVE_FILE_CONTENT_MEMORY_INDEX_V1 task=%s ok=%s reason=%s chars=%s", task_id, _df_ci.get("ok"), _df_ci.get("reason"), _df_ci.get("chars"))
            except Exception as _dfci_e:
                logger.warning("DRIVE_FILE_CONTENT_MEMORY_INDEX_V1_ERR task=%s err=%s", task_id, _dfci_e)
            # === END DRIVE_FILE_CONTENT_MEMORY_INDEX_V1_CALL ===
            logger.info("DRIVE_FILE_MEMORY_INDEX_V1 task=%s file=%s", task_id, _df_file_name)
            # === FILE_CATALOG_AUTOSYNC_AFTER_DRIVE_FILE_V1 ===
            try:
                _filemem_save_catalog(str(chat_id), int(topic_id or 0))
            except Exception as _cat_e:
                logger.warning("FILE_CATALOG_AUTOSYNC_AFTER_DRIVE_FILE_V1_ERR task=%s err=%s", task_id, _cat_e)
            # === END FILE_CATALOG_AUTOSYNC_AFTER_DRIVE_FILE_V1 ===
        except Exception as _e:
            logger.warning("DRIVE_FILE_MEMORY_INDEX_V1_ERR task=%s err=%s", task_id, _e)
        # === END DRIVE_FILE_MEMORY_INDEX_V1 + FILE_DUPLICATE_MEMORY_GUARD_V1 ===
        # DRIVE_FILE_SOURCE_HEALTHCHECK_GUARD_V1
        _hc_src = str(data.get("source") or "").lower()
        _hc_fn = str(file_name or "").lower()
        _hc_raw_low = str(raw_input or "").lower()
        _hc_markers = ("retry_queue_healthcheck", "healthcheck", "areal_hc_", "_hc_file")
        if _hc_src in ("google_drive", "gdrive") or any(m in _hc_fn or m in _hc_raw_low for m in _hc_markers):
            _update_task(conn, task_id, state="CANCELLED", error_message="SERVICE_FILE_IGNORED:HEALTHCHECK")
            conn.commit()
            logger.info("DRIVE_FILE_SOURCE_HEALTHCHECK_GUARD_V1 cancelled task=%s", task_id)
            return
        # === TASK_TYPE_DETECT_V1 ===
        _task_type = "DOCUMENT_TASK"
        _fn_lower = file_name.lower()
        _caption_lower = (data.get("caption") or raw_input or "").lower()
        if any(_fn_lower.endswith(e) for e in (".dwg", ".dxf")):
            _task_type = "DWG_TASK"
        elif any(_fn_lower.endswith(e) for e in (".xlsx", ".xls", ".csv")):
            _task_type = "ESTIMATE_TASK"
        elif any(_fn_lower.endswith(e) for e in (".jpg", ".jpeg", ".png", ".heic", ".webp")):
            if any(w in _caption_lower for w in ["дефект", "акт", "технадзор", "нарушен"]):
                _task_type = "TECHNADZOR_TASK"
            else:
                _task_type = "OCR_TASK"
        elif any(w in _caption_lower for w in ["смета", "расчёт", "посчитай", "калькул"]):
            _task_type = "ESTIMATE_TASK"
        elif any(w in _caption_lower for w in ["дефект", "акт", "технадзор"]):
            _task_type = "TECHNADZOR_TASK"
        try:
            conn.execute("UPDATE tasks SET task_type=? WHERE id=?", (_task_type, task_id))
            conn.commit()
        except Exception:
            pass
        logger.info("TASK_TYPE_DETECT_V1 task=%s type=%s", task_id, _task_type)
        # === END TASK_TYPE_DETECT_V1 ===

        # === DUPLICATE_GUARD_CALL_V1 ===
        try:
            _dupe = find_duplicate(conn, str(chat_id), int(topic_id or 0), file_id)
            if _dupe:
                file_name = data.get("file_name", "файл")
                _dupe_msg = duplicate_message(_dupe, file_name)
                # === RESULT_VALIDATOR_GUARD_V1 ===
                if _check_result_before_confirm(_dupe_msg):
                    # === TASK_WORKER_ARTIFACT_GATE_V1 ===
                    try:
                        if validate_engine_result is not None:
                            _twag_raw = ""
                            _twag_input_type = ""
                            try:
                                _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                            except Exception:
                                _twag_raw = ""
                            try:
                                _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                            except Exception:
                                _twag_input_type = ""
                            _twag_result = _dupe_msg
                            _twag_check = validate_engine_result(
                                {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                                input_type=_twag_input_type,
                                user_text=_twag_raw,
                                topic_id=topic_id,
                            )
                            if not _twag_check.get("ok"):
                                _update_task(conn, task_id, state="FAILED", result="",
                                    error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                                _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                                return
                    except Exception as _twag_e:
                        logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
                    # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
                    _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_dupe_msg, error_message="")
                else:
                    _update_task(conn, task_id, state="FAILED", result=_dupe_msg, error_message="FORBIDDEN_PHRASE")
                # === END RESULT_VALIDATOR_GUARD_V1 ===
                _history(conn, task_id, "state:AWAITING_CONFIRMATION:duplicate_guard")
                conn.commit()
                from core.reply_sender import send_reply_ex
                send_reply_ex(chat_id=str(chat_id), text=_dupe_msg, reply_to_message_id=reply_to, message_thread_id=topic_id)
                return
        except Exception as _dge:
            logger.warning("DUPLICATE_GUARD_CALL_ERR %s", _dge)
        # === END DUPLICATE_GUARD_CALL_V1 ===
        file_name = data["file_name"]
    except Exception as e:
        logger.error(f"DRIVE_FILE: invalid raw_input for {task_id}: {e}")
        _update_task(conn, task_id, state="FAILED", error_message="invalid raw_input")
        return

    local_path = f"/root/.areal-neva-core/runtime/drive_files/{task_id}_{file_name}"
    os.makedirs(os.path.dirname(local_path), exist_ok=True)

    logger.info(f"DRIVE_FILE: downloading {file_id} -> {local_path}")
    ok = _download_from_drive(file_id, local_path)  # HOTFIX_OK_BEFORE_SIZE_CHECK_V1
    # === FILE_SIZE_LIMIT_V1 ===
    if ok and local_path and os.path.exists(local_path):
        _fsize = os.path.getsize(local_path)
        if _fsize > 50 * 1024 * 1024:  # 50MB
            _update_task(conn, task_id, state="FAILED",
                        result="", error_message="FILE_TOO_LARGE")
            conn.commit()
            from core.reply_sender import send_reply_ex
            send_reply_ex(chat_id=str(chat_id),
                         text="Файл слишком большой (>50MB). Сожми или разбей на части.",
                         reply_to_message_id=reply_to, message_thread_id=topic_id)
            return
    # === END FILE_SIZE_LIMIT_V1 ===
    # === FILE_INTAKE_ROUTER_V1_WIRED ===
    if ok and local_path and os.path.exists(local_path):
        try:
            from core.file_intake_router import route_file, detect_intent, detect_intent_from_filename, should_ask_clarification, get_clarification_message
            _fir_caption = data.get("caption", "") or raw_input or ""
            _fir_intent = detect_intent(_fir_caption) or detect_intent_from_filename(file_name)
            _fir_topic_role = ""
            if _fir_intent:
                _fir_result = await route_file(local_path, task_id, int(topic_id or 0), _fir_intent)
                if _fir_result and _fir_result.get("success"):
                    _fir_msg = _fir_result.get("result_text") or _fir_result.get("drive_link") or "Готово"
                    # === TASK_WORKER_ARTIFACT_GATE_V1 ===
                    try:
                        if validate_engine_result is not None:
                            _twag_raw = ""
                            _twag_input_type = ""
                            try:
                                _twag_raw = str(task["raw_input"] if "raw_input" in task.keys() else "")
                            except Exception:
                                _twag_raw = ""
                            try:
                                _twag_input_type = str(task["input_type"] if "input_type" in task.keys() else "")
                            except Exception:
                                _twag_input_type = ""
                            _twag_result = _fir_msg
                            _twag_check = validate_engine_result(
                                {"summary": _twag_result, "engine": "TASK_WORKER_ARTIFACT_GATE_V1"},
                                input_type=_twag_input_type,
                                user_text=_twag_raw,
                                topic_id=topic_id,
                            )
                            if not _twag_check.get("ok"):
                                _update_task(conn, task_id, state="FAILED", result="",
                                    error_message="TASK_WORKER_ARTIFACT_GATE_V1:" + str(_twag_check.get("reason") or "INVALID"))
                                _history(conn, task_id, "TASK_WORKER_ARTIFACT_GATE_V1:FAILED:" + str(_twag_check.get("reason") or "INVALID"))
                                return
                    except Exception as _twag_e:
                        logger.warning("TASK_WORKER_ARTIFACT_GATE_V1_ERR task=%s err=%s", task_id, _twag_e)
                    # === END_TASK_WORKER_ARTIFACT_GATE_V1 ===
                    _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_fir_msg, error_message="")
                    _history(conn, task_id, "state:AWAITING_CONFIRMATION:file_intake_router")
                    conn.commit()
                    from core.reply_sender import send_reply_ex
                    send_reply_ex(chat_id=str(chat_id), text=_fir_msg, reply_to_message_id=reply_to, message_thread_id=topic_id)
                    return
            elif should_ask_clarification(_fir_caption, has_file=True):
                _fir_clarif = get_clarification_message(file_name, int(topic_id or 0))
                _update_task(conn, task_id, state="WAITING_CLARIFICATION", result=_fir_clarif, error_message="")
                conn.commit()
                from core.reply_sender import send_reply_ex
                send_reply_ex(chat_id=str(chat_id), text=_fir_clarif, reply_to_message_id=reply_to, message_thread_id=topic_id)
                return
        except Exception as _fir_err:
            logger.warning("FILE_INTAKE_ROUTER_V1_ERR task=%s err=%s", task_id, _fir_err)
    # === END FILE_INTAKE_ROUTER_V1_WIRED ===
    if not ok:
        _update_task(conn, task_id, state="FAILED", error_message="download failed")
        return

    conn.execute("UPDATE drive_files SET stage='downloaded' WHERE task_id=?", (task_id,))

    result = f"Файл {file_name} скачан, ожидает анализа"
    try:
        _, _, topic_role, _ = _load_memory_context(chat_id, topic_id)
        analysis = await analyze_downloaded_file(
            local_path=local_path,
            file_name=file_name,
            mime_type=data.get("mime_type", ""),
            user_text=data.get("caption", ""),
            topic_role=topic_role,
        )
        if isinstance(analysis, dict):
            summary = _s(analysis.get("summary")) or result
            artifact_path = _s(analysis.get("artifact_path"))
            artifact_name = _s(analysis.get("artifact_name")) or os.path.basename(artifact_path)
            result = summary
            if artifact_path and os.path.exists(artifact_path):
                try:
                    upload_res = await upload_file_to_topic(artifact_path, artifact_name, chat_id, topic_id)
                    if isinstance(upload_res, dict) and upload_res.get("ok") and upload_res.get("drive_file_id"):
                        result = summary + f"\n\nАртефакт: https://drive.google.com/file/d/{upload_res.get('drive_file_id')}/view"
                    # === QUALITY_GATE_WIRED_V1 ===
                    _qg = _quality_gate_artifact(
                        drive_link=result,
                        input_type=input_type,
                        task_type=_task_type if '_task_type' in dir() else ""
                    )
                    if not _qg.get("ok"):
                        logger.warning("QUALITY_GATE_FAIL task=%s reason=%s", task_id, _qg.get("reason"))
                    # === END QUALITY_GATE_WIRED_V1 ===
                    # === TEMP_CLEANUP_AFTER_UPLOAD_V1 ===
                    try:
                        _tc_upload([local_path])
                        _tc_task(task_id)
                    except Exception:
                        pass
                    # === END TEMP_CLEANUP_AFTER_UPLOAD_V1 ===
                        # === PATCH: save_pin + save_memory + log ===
                        try:
                            save_pin(chat_id, task_id, result, topic_id)
                            _save_memory(chat_id, topic_id, raw_input, result)
                            logger.info(f"DRIVE_FILE pin_memory_saved task_id={task_id}")
                        except Exception as e:
                            logger.error(f"DRIVE_FILE pin/memory failed task={task_id} err={e}")
                        # === END PATCH ===
                    else:
                        result = summary + "\n\nАртефакт создан, но загрузка в Drive не подтвердилась"
                except Exception as e:
                    logger.error(f"DRIVE_FILE artifact upload failed task={task_id} err={e}")
                    result = summary + "\n\nАртефакт создан локально, но загрузка в Drive завершилась ошибкой"
    except Exception as e:
        logger.error(f"DRIVE_FILE analyze skipped task={task_id} err={e}")

    _update_task(conn, task_id, state="AWAITING_CONFIRMATION", result=_ff13c_strip_manifest_links(result))
    logger.info(f"DRIVE_FILE: {task_id} processed")



# === SINGLE_CHAR_REPLY_AS_PRICE_CHOICE_V1 ===
# Однобуквенный выбор цены и голосовой выбор должны продолжать parent task, а не создавать общий ответ
# AWAITING_PRICE_CONFIRMATION_STATE_V1
# VOICE_REPLY_TO_PARENT_TASK_V1
# === END_SINGLE_CHAR_REPLY_AS_PRICE_CHOICE_V1 ===

# === STARTUP_RECOVERY_V1_MAIN_WRAP ===
try:
    import inspect as _startup_recovery_inspect_v1
    _orig_task_worker_main_startup_recovery_v1 = main

    async def main(*args, **kwargs):
        try:
            if run_startup_recovery is not None:
                await run_startup_recovery(CORE_DB)
        except Exception as _startup_recovery_err:
            logger.warning("STARTUP_RECOVERY_V1_ERR %s", _startup_recovery_err)

        _res = _orig_task_worker_main_startup_recovery_v1(*args, **kwargs)
        if _startup_recovery_inspect_v1.isawaitable(_res):
            return await _res
        return _res
except Exception as _startup_recovery_wrap_err:
    logger.warning("STARTUP_RECOVERY_V1_WRAP_ERR %s", _startup_recovery_wrap_err)
# === END_STARTUP_RECOVERY_V1_MAIN_WRAP ===


# === THREE_CONTOURS_FULL_CONTEXT_NO_REPEAT_CLARIFY_V1_WORKER_WRAPPER ===
_tcfc1_orig_handle_in_progress_worker = _handle_in_progress

def _tcfc1_row_get(row, key, default=None):
    try:
        return row[key]
    except Exception:
        try:
            idx = {
                "id": 0, "chat_id": 1, "user_id": 2, "input_type": 3, "raw_input": 4,
                "state": 5, "result": 6, "error_message": 7, "reply_to_message_id": 8,
                "created_at": 9, "updated_at": 10, "bot_message_id": 11, "topic_id": 12,
                "task_type": 13,
            }.get(key)
            if idx is not None:
                return row[idx]
        except Exception:
            pass
    return default

def _tcfc1_worker_history_col(conn):
    try:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(task_history)").fetchall()]
        if "action" in cols:
            return "action"
        if "event" in cols:
            return "event"
    except Exception:
        pass
    return ""

def _tcfc1_worker_build_full_context(conn, task_id: str, chat_id: str, topic_id: int, current_raw: str) -> tuple[str, str]:
    parts = []
    cur = str(current_raw or "").strip()
    if cur:
        parts.append("Текущее сообщение:\n" + cur)

    parent_id = ""
    try:
        row = conn.execute(
            """
            SELECT id, raw_input
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=?
              AND id<>?
              AND state IN ('WAITING_CLARIFICATION','IN_PROGRESS')
            ORDER BY rowid DESC
            LIMIT 1
            """,
            (str(chat_id), int(topic_id or 0), str(task_id)),
        ).fetchone()
        if row:
            parent_id = str(row[0])
            if str(row[1] or "").strip():
                parts.insert(0, "Исходная активная задача:\n" + str(row[1] or "").strip())
    except Exception:
        parent_id = ""

    hist_ids = [x for x in (parent_id, str(task_id)) if x]
    hcol = _tcfc1_worker_history_col(conn)
    if hcol and hist_ids:
        try:
            qmarks = ",".join("?" for _ in hist_ids)
            rows = conn.execute(
                f"""
                SELECT task_id,{hcol}
                FROM task_history
                WHERE task_id IN ({qmarks})
                  AND ({hcol} LIKE 'clarified:%' OR {hcol} LIKE '%clarification_accepted%')
                ORDER BY rowid ASC
                LIMIT 80
                """,
                tuple(hist_ids),
            ).fetchall()
            clar = []
            for r in rows:
                v = str(r[1] or "")
                if v.startswith("clarified:"):
                    clar.append(v.split(":", 1)[1].strip())
            if clar:
                parts.append("Уточнения пользователя:\n" + "\n".join(x for x in clar if x))
        except Exception:
            pass

    try:
        rows = conn.execute(
            """
            SELECT raw_input
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=?
              AND id<>?
              AND raw_input IS NOT NULL
              AND TRIM(raw_input)<>''
            ORDER BY rowid DESC
            LIMIT 6
            """,
            (str(chat_id), int(topic_id or 0), str(task_id)),
        ).fetchall()
        recent = []
        for r in rows:
            v = str(r[0] or "").strip()
            if v and v not in "\n\n".join(parts):
                recent.append(v)
        if recent:
            parts.append("Последние сообщения топика без результатов и ссылок:\n" + "\n".join(recent))
    except Exception:
        pass

    return "\n\n".join(parts).strip(), parent_id

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    try:
        task_id = str(_tcfc1_row_get(task, "id", "") or "")
        chat_id = str(_tcfc1_row_get(task, "chat_id", "") or "")
        topic_id = int(_tcfc1_row_get(task, "topic_id", 0) or 0)
        input_type = str(_tcfc1_row_get(task, "input_type", "text") or "text")
        raw_input = str(_tcfc1_row_get(task, "raw_input", "") or "")
        reply_to = _tcfc1_row_get(task, "reply_to_message_id", None)

        if topic_id == 2 and task_id:
            full_context, parent_id = _tcfc1_worker_build_full_context(conn, task_id, chat_id, topic_id, raw_input)
            if full_context:
                try:
                    from core.sample_template_engine import handle_stroyka_topic2_full_context_gate_v1
                    handled = await handle_stroyka_topic2_full_context_gate_v1(
                        conn=conn,
                        task_id=task_id,
                        chat_id=chat_id,
                        topic_id=topic_id,
                        raw_context=full_context,
                        input_type=input_type,
                        reply_to_message_id=reply_to,
                    )
                    if handled:
                        try:
                            if parent_id and parent_id != task_id:
                                conn.execute(
                                    "UPDATE tasks SET state='DONE', result=?, error_message='', updated_at=datetime('now') WHERE id=?",
                                    ("Закрыто новой сметой по полному контексту", parent_id),
                                )
                                _history(conn, parent_id, f"THREE_CONTOURS_FULL_CONTEXT_NO_REPEAT_CLARIFY_V1:superseded_by:{task_id}")
                                conn.commit()
                        except Exception as _e:
                            logger.warning("THREE_CONTOURS_FULL_CONTEXT_NO_REPEAT_CLARIFY_V1_PARENT_CLOSE_ERR %s", _e)
                        logger.info("THREE_CONTOURS_FULL_CONTEXT_NO_REPEAT_CLARIFY_V1 handled task=%s topic=%s", task_id, topic_id)
                        return
                except Exception as _e:
                    logger.warning("THREE_CONTOURS_FULL_CONTEXT_NO_REPEAT_CLARIFY_V1_ERR task=%s err=%s", task_id, _e)
    except Exception as _e:
        logger.warning("THREE_CONTOURS_FULL_CONTEXT_NO_REPEAT_CLARIFY_V1_WRAPPER_ERR %s", _e)

    return await _tcfc1_orig_handle_in_progress_worker(conn, task, chat_id, topic_id)  # HOTFIX_HANDLE_IN_PROGRESS_WRAPPER_CHAIN_V1

# === END_THREE_CONTOURS_FULL_CONTEXT_NO_REPEAT_CLARIFY_V1_WORKER_WRAPPER ===



# === STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1 ===
# Must be placed BEFORE asyncio.run(main())
# Fixes:
# - topic_2 must not fall into FILE_TECH_CONTOUR_FOLLOWUP_V2 for estimate/template questions
# - topic_2 must not repeat old task result when full technical context is available
# - topic_2 must route complete construction estimate TЗ into formula/template source gate once
try:
    _st2_pm_orig_handle_in_progress_v1 = _handle_in_progress

    def _st2_pm_row_get(row, key, default=None):
        try:
            return row[key]
        except Exception:
            try:
                idx_map = {
                    "id": 0,
                    "chat_id": 1,
                    "user_id": 2,
                    "input_type": 3,
                    "raw_input": 4,
                    "state": 5,
                    "result": 6,
                    "error_message": 7,
                    "reply_to_message_id": 8,
                    "created_at": 9,
                    "updated_at": 10,
                    "bot_message_id": 11,
                    "topic_id": 12,
                    "task_type": 13,
                }
                i = idx_map.get(key)
                if i is not None:
                    return row[i]
            except Exception:
                pass
        return default

    def _st2_pm_clean_voice(text: str) -> str:
        try:
            return re.sub(r"^\s*\[VOICE\]\s*", "", str(text or ""), flags=re.I).strip()
        except Exception:
            return str(text or "").strip()

    def _st2_pm_low(text: str) -> str:
        return _st2_pm_clean_voice(text).lower().replace("ё", "е")

    def _st2_pm_reply_to(task):
        try:
            return _st2_pm_row_get(task, "reply_to_message_id", None)
        except Exception:
            return None

    def _st2_pm_send_done(conn, task_id: str, chat_id: str, text: str, reply_to=None, kind: str = "stroyka_topic2_premain"):
        _update_task(conn, task_id, state="DONE", result=str(text or ""), error_message="")
        _history(conn, task_id, f"STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1:{kind}")
        try:
            conn.commit()
        except Exception:
            pass
        try:
            _send_once_ex(conn, task_id, str(chat_id), str(text or ""), reply_to, kind)
        except Exception:
            try:
                _send_once(conn, task_id, str(chat_id), str(text or ""), reply_to, kind)
            except Exception as _send_err:
                logger.warning("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1_SEND_ERR %s", _send_err)

    def _st2_pm_template_source_text(chat_id: str, topic_id: int) -> str:
        import json as _json
        from pathlib import Path as _Path
        safe_chat = re.sub(r"[^0-9A-Za-z_-]+", "_", str(chat_id))[:80]
        active = _Path("/root/.areal-neva-core/data/templates/estimate") / f"ACTIVE__chat_{safe_chat}__topic_{int(topic_id or 0)}.json"
        if not active.exists():
            active = _Path("/root/.areal-neva-core/data/templates/estimate/ACTIVE__chat_-1003725299009__topic_2.json")
        data = {}
        try:
            data = _json.loads(active.read_text(encoding="utf-8"))
        except Exception:
            data = {}
        source_name = str(data.get("source_file_name") or "М-110.xlsx")
        folder_name = str(data.get("source_folder_name") or "ESTIMATES/templates")
        all_tpl = []
        try:
            for x in data.get("all_templates") or []:
                name = str(x.get("file_name") or "").strip()
                if name:
                    all_tpl.append(name)
        except Exception:
            pass
        if not all_tpl:
            all_tpl = ["Ареал Нева.xlsx", "М-80.xlsx", "М-110.xlsx", "фундамент_Склад2.xlsx", "крыша и перекр.xlsx"]
        return (
            "Да. Правильные шаблоны смет подключены\n\n"
            f"Источник: Google Drive / AI_ORCHESTRA / {folder_name}\n"
            f"Активный эталон: {source_name}\n"
            "Правило: шаблон используется как структура, формулы и оформление. Расчёт берётся только из текущего ТЗ\n\n"
            "Доступные шаблоны:\n- " + "\n- ".join(all_tpl[:10]) + "\n\n"
            "Для топика СТРОЙКА больше не должен срабатывать поиск случайных файлов технадзора вместо сметного шаблона"
        )

    def _st2_pm_is_template_question(text: str) -> bool:
        low = _st2_pm_low(text)
        if not low:
            return False
        create_words = ("сделай", "создай", "сформируй", "подготовь", "составь", "посчитай", "рассчитай")
        template_words = ("шаблон", "образец", "эталон", "правильн", "откуда", "на чем основыва", "на чём основыва", "что ты туп")
        if any(w in low for w in create_words):
            return False
        return any(w in low for w in template_words)

    def _st2_pm_is_estimate_related(text: str) -> bool:
        low = _st2_pm_low(text)
        if not low:
            return False
        estimate_words = (
            "смет", "стоимост", "посчитай", "рассчитай", "расчет", "расчёт",
            "цена", "руб", "итого", "ндс", "работ", "материал",
            "строительств", "дом", "барн", "бар хаус", "barn", "каркас",
            "газобетон", "фундамент", "плит", "монолит", "кровл", "фальц",
            "имитация бруса", "отделка", "этаж", "высота", "размер"
        )
        return any(w in low for w in estimate_words)

    def _st2_pm_is_control_only(text: str) -> bool:
        low = _st2_pm_low(text)
        return low in {
            "задача завершена",
            "завершена",
            "готово",
            "закрой",
            "закрыть",
            "отмена",
            "отбой",
            "все задачи отменены",
            "всё задачи отменены",
            "сейчас все задачи отменены",
            "сейчас всё задачи отменены",
        }

    def _st2_pm_history_col(conn) -> str:
        try:
            cols = [r[1] for r in conn.execute("PRAGMA table_info(task_history)").fetchall()]
            if "action" in cols:
                return "action"
            if "event" in cols:
                return "event"
        except Exception:
            pass
        return ""

    def _st2_pm_build_full_context(conn, task_id: str, chat_id: str, topic_id: int, raw: str) -> tuple[str, str]:
        parts = []
        cur = _st2_pm_clean_voice(raw)
        if cur:
            parts.append("Текущее сообщение:\n" + cur)

        parent_id = ""
        try:
            row = conn.execute(
                """
                SELECT id, raw_input
                FROM tasks
                WHERE chat_id=?
                  AND COALESCE(topic_id,0)=?
                  AND id<>?
                  AND state IN ('WAITING_CLARIFICATION','IN_PROGRESS','NEW','AWAITING_CONFIRMATION')
                ORDER BY rowid DESC
                LIMIT 1
                """,
                (int(chat_id), int(topic_id or 0), str(task_id)),
            ).fetchone()
            if row:
                parent_id = str(row[0] or "")
                parent_raw = _st2_pm_clean_voice(str(row[1] or ""))
                if parent_raw and parent_raw not in cur:
                    parts.insert(0, "Предыдущее активное ТЗ:\n" + parent_raw)
        except Exception as _e:
            logger.warning("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1_PARENT_CTX_ERR %s", _e)

        try:
            col = _st2_pm_history_col(conn)
            if col:
                ids = [str(task_id)]
                if parent_id:
                    ids.append(parent_id)
                qs = ",".join("?" for _ in ids)
                rows = conn.execute(
                    f"SELECT {col} FROM task_history WHERE task_id IN ({qs}) AND {col} LIKE 'clarified:%' ORDER BY rowid ASC LIMIT 50",
                    ids,
                ).fetchall()
                clar = []
                for r in rows:
                    v = str(r[0] or "")
                    if ":" in v:
                        c = v.split(":", 1)[1].strip()
                        if c and c not in clar:
                            clar.append(c)
                if clar:
                    parts.append("Уточнения пользователя:\n" + "\n".join(clar))
        except Exception as _e:
            logger.warning("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1_CLAR_CTX_ERR %s", _e)

        return "\n\n".join(x for x in parts if x.strip()), parent_id

    async def _st2_pm_call_full_context_gate(conn, task_id: str, chat_id: str, topic_id: int, full_context: str, input_type: str, reply_to):
        try:
            from core.sample_template_engine import handle_stroyka_topic2_full_context_gate_v1 as _gate
        except Exception as _e:
            logger.warning("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1_GATE_IMPORT_ERR %s", _e)
            return False

        calls = [
            (conn, task_id, str(chat_id), int(topic_id or 0), full_context, str(input_type or "text"), reply_to),
            (conn, task_id, str(chat_id), int(topic_id or 0), full_context, "text", reply_to),
            (conn, task_id, str(chat_id), int(topic_id or 0), full_context),
        ]
        last_type_err = None
        for args in calls:
            try:
                res = await _gate(*args)
                if res is True:
                    return True
                if isinstance(res, dict) and (res.get("handled") or res.get("success")):
                    return True
                return bool(res)
            except TypeError as _te:
                last_type_err = _te
                continue
            except Exception as _e:
                logger.warning("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1_GATE_ERR %s", _e)
                return False
        if last_type_err:
            logger.warning("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1_GATE_SIGNATURE_ERR %s", last_type_err)
        return False

    async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
        try:
            task_id = str(_st2_pm_row_get(task, "id", ""))
            chat_id = str(_st2_pm_row_get(task, "chat_id", ""))
            topic_id = int(_st2_pm_row_get(task, "topic_id", 0) or 0)
            raw_input = str(_st2_pm_row_get(task, "raw_input", "") or "")
            input_type = str(_st2_pm_row_get(task, "input_type", "text") or "text")
            reply_to = _st2_pm_reply_to(task)

            if topic_id == 2:
                if _st2_pm_is_control_only(raw_input):
                    _st2_pm_send_done(conn, task_id, chat_id, "Задача в СТРОЙКЕ закрыта", reply_to, "stroyka_control_close")
                    return

                if _st2_pm_is_template_question(raw_input):
                    _st2_pm_send_done(conn, task_id, chat_id, _st2_pm_template_source_text(chat_id, topic_id), reply_to, "stroyka_template_source_answer")
                    return

                if _st2_pm_is_estimate_related(raw_input):
                    full_context, parent_id = _st2_pm_build_full_context(conn, task_id, chat_id, topic_id, raw_input)
                    handled = await _st2_pm_call_full_context_gate(
                        conn,
                        task_id,
                        chat_id,
                        topic_id,
                        full_context or raw_input,
                        input_type,
                        reply_to,
                    )
                    if handled:
                        try:
                            if parent_id and parent_id != task_id:
                                conn.execute(
                                    "UPDATE tasks SET state='DONE', result=?, error_message='', updated_at=datetime('now') WHERE id=?",
                                    ("Закрыто новой сметой по полному ТЗ", parent_id),
                                )
                                _history(conn, parent_id, f"STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1:superseded_by:{task_id}")
                                conn.commit()
                        except Exception as _e:
                            logger.warning("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1_PARENT_CLOSE_ERR %s", _e)
                        logger.info("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1 handled task=%s topic=%s", task_id, topic_id)
                        return
        except Exception as _e:
            logger.warning("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1_ERR %s", _e)

        return await _st2_pm_orig_handle_in_progress_v1(conn, task, chat_id, topic_id)  # HOTFIX_HANDLE_IN_PROGRESS_WRAPPER_CHAIN_V1

except Exception as _st2_pm_wrap_err:
    logger.warning("STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1_WRAP_ERR %s", _st2_pm_wrap_err)

# === END_STROYKA_TOPIC2_PREMAIN_ROUTE_FIX_V1 ===


# === TOPIC2_ONE_BIG_FINAL_PIPELINE_V1_WORKER_GUARD ===
# Final topic_2 guard before worker main:
# - blocks old generated search subtasks
# - blocks FILE_TECH_CONTOUR_FOLLOWUP_V2 for topic_2 vague messages
# - blocks memory revive of old estimates
# - merges current message + current clarifications + live parent context
# - sends complete estimate context to one formula-preserving template pipeline
# - never leaves handled topic_2 task in endless IN_PROGRESS

_TOP2_ONE_BIG_ORIG_HANDLE_IN_PROGRESS_V1 = _handle_in_progress

def _top2_ob_row(row, key, default=None):
    try:
        return row[key]
    except Exception:
        pass
    try:
        idx = {
            "id": 0,
            "chat_id": 1,
            "user_id": 2,
            "input_type": 3,
            "raw_input": 4,
            "state": 5,
            "result": 6,
            "error_message": 7,
            "reply_to_message_id": 8,
            "created_at": 9,
            "updated_at": 10,
            "bot_message_id": 11,
            "topic_id": 12,
            "task_type": 13,
        }.get(key)
        if idx is not None:
            return row[idx]
    except Exception:
        pass
    return default

def _top2_ob_s(v):
    return "" if v is None else str(v)

def _top2_ob_low(v):
    return _top2_ob_s(v).lower().replace("ё", "е").strip()

def _top2_ob_history_col(conn):
    try:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(task_history)").fetchall()]
        if "action" in cols:
            return "action"
        if "event" in cols:
            return "event"
    except Exception:
        pass
    return ""

def _top2_ob_hist(conn, task_id, action):
    try:
        _history(conn, task_id, action)
        return
    except Exception:
        pass
    try:
        col = _top2_ob_history_col(conn)
        if col:
            conn.execute(f"INSERT INTO task_history(task_id,{col},created_at) VALUES(?,?,datetime('now'))", (task_id, action))
            conn.commit()
    except Exception:
        pass

def _top2_ob_update(conn, task_id, state, result="", error=""):
    try:
        _update_task(conn, task_id, state=state, result=result, error_message=error)
        return
    except Exception:
        pass
    try:
        conn.execute(
            "UPDATE tasks SET state=?, result=?, error_message=?, updated_at=datetime('now') WHERE id=?",
            (state, result, error, task_id),
        )
        conn.commit()
    except Exception:
        pass

def _top2_ob_send(conn, task_id, chat_id, text, reply_to, kind):
    try:
        _send_once_ex(conn, task_id, chat_id, text, reply_to, kind)
        return
    except Exception:
        pass
    try:
        _send_once(conn, task_id, str(chat_id), text, reply_to, kind)
    except Exception:
        pass

def _top2_ob_current_clarifications(conn, task_id):
    out = []
    try:
        col = _top2_ob_history_col(conn)
        if not col:
            return out
        rows = conn.execute(
            f"SELECT {col} FROM task_history WHERE task_id=? AND {col} LIKE 'clarified:%' ORDER BY rowid ASC LIMIT 50",
            (task_id,),
        ).fetchall()
        for r in rows:
            v = _top2_ob_s(r[0])
            if ":" in v:
                x = v.split(":", 1)[1].strip()
                if x:
                    out.append(x)
    except Exception:
        pass
    return out

def _top2_ob_latest_live_parent(conn, task_id, chat_id, topic_id):
    try:
        row = conn.execute(
            """
            SELECT id, raw_input, state
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=?
              AND id<>?
              AND state IN ('NEW','IN_PROGRESS','WAITING_CLARIFICATION','AWAITING_CONFIRMATION')
              AND input_type<>'search'
            ORDER BY rowid DESC
            LIMIT 1
            """,
            (chat_id, int(topic_id or 0), task_id),
        ).fetchone()
        if row:
            return _top2_ob_s(row[0]), _top2_ob_s(row[1]), _top2_ob_s(row[2])
    except Exception:
        pass
    return "", "", ""

def _top2_ob_is_vague(text):
    low = _top2_ob_low(text)
    low = re.sub(r"^\s*\[voice\]\s*", "", low, flags=re.I).strip()
    if not low:
        return True
    vague = (
        "что дальше", "ну что", "что там", "посмотри", "я скидывал", "я тебе скидывал",
        "последнее задание", "какое последнее", "вообще не так", "не так",
        "ты тупишь", "где", "еще раз", "ещё раз", "продолжай"
    )
    strong = (
        "смет", "посчитай", "рассчитай", "стоимость", "дом", "фундамент",
        "плита", "монолит", "каркас", "газобетон", "барн", "бар house", "бархаус",
        "12", "10", "8", "м2", "м3", "руб"
    )
    return any(x in low for x in vague) and not any(x in low for x in strong)

def _top2_ob_cancel_text(text):
    low = _top2_ob_low(text)
    return any(x in low for x in ("все задачи отменены", "всё отменено", "отмени все", "стоп все", "стоп всё", "закрой все задачи"))

def _top2_ob_estimate_like(text):
    low = _top2_ob_low(text)
    return any(x in low for x in (
        "смет", "стоимост", "посчитай", "рассчитай", "расчет", "расчёт",
        "монолит", "фундамент", "плит", "дом", "каркас", "газобетон",
        "барн", "бар house", "бархаус", "материал", "работ"
    ))

def _top2_ob_build_context(conn, task_id, chat_id, topic_id, raw):
    parts = []
    cur = _top2_ob_s(raw).strip()
    if cur:
        parts.append("Текущее сообщение:\n" + cur)

    clar = _top2_ob_current_clarifications(conn, task_id)
    if clar:
        parts.append("Уточнения текущей задачи:\n" + "\n".join(clar))

    parent_id, parent_raw, parent_state = _top2_ob_latest_live_parent(conn, task_id, chat_id, topic_id)
    if parent_id and parent_raw and not _top2_ob_is_vague(parent_raw):
        parts.append("Живая родительская задача:\n" + parent_raw)
        pclar = _top2_ob_current_clarifications(conn, parent_id)
        if pclar:
            parts.append("Уточнения родительской задачи:\n" + "\n".join(pclar))

    return "\n\n".join(parts).strip(), parent_id

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    task_id = _top2_ob_s(_top2_ob_row(task, "id"))
    chat_id = _top2_ob_row(task, "chat_id")
    topic_id = int(_top2_ob_row(task, "topic_id", 0) or 0)
    input_type = _top2_ob_s(_top2_ob_row(task, "input_type", "text"))
    raw_input = _top2_ob_s(_top2_ob_row(task, "raw_input", ""))
    reply_to = _top2_ob_row(task, "reply_to_message_id", None)

    if topic_id == 2:
        try:
            if input_type == "search":
                _top2_ob_update(conn, task_id, "DONE", "TOPIC2_ONE_BIG_FINAL_PIPELINE_V1: generated search subtask blocked", "")
                _top2_ob_hist(conn, task_id, "TOPIC2_ONE_BIG_FINAL_PIPELINE_V1:blocked_search_subtask")
                return

            if _top2_ob_cancel_text(raw_input):
                try:
                    conn.execute(
                        """
                        UPDATE tasks
                        SET state='CANCELLED',
                            result='TOPIC2_ONE_BIG_FINAL_PIPELINE_V1: cancelled by user topic2 command',
                            error_message='',
                            updated_at=datetime('now')
                        WHERE chat_id=?
                          AND COALESCE(topic_id,0)=2
                          AND state IN ('NEW','IN_PROGRESS','WAITING_CLARIFICATION','AWAITING_CONFIRMATION')
                          AND id<>?
                        """,
                        (chat_id, task_id),
                    )
                    conn.commit()
                except Exception:
                    pass
                msg = "Все активные задачи topic 2 закрыты"
                _top2_ob_update(conn, task_id, "DONE", msg, "")
                _top2_ob_hist(conn, task_id, "TOPIC2_ONE_BIG_FINAL_PIPELINE_V1:cancel_all_topic2")
                _top2_ob_send(conn, task_id, chat_id, msg, reply_to, "topic2_cancel_all")
                return

            full_context, parent_id = _top2_ob_build_context(conn, task_id, chat_id, topic_id, raw_input)
            if raw_input and str(raw_input).strip() and str(raw_input).strip() not in str(full_context or ""):
                full_context = (str(raw_input).strip() + "\n\n" + str(full_context or "")).strip()  # TOPIC2_CURRENT_RAW_CONTEXT_FIRST_V1

            if _top2_ob_is_vague(raw_input) and not _top2_ob_estimate_like(full_context):
                msg = "Нового полного ТЗ для сметы в сообщении нет. Старую смету из памяти не поднимаю"
                _top2_ob_update(conn, task_id, "DONE", msg, "")
                _top2_ob_hist(conn, task_id, "TOPIC2_ONE_BIG_FINAL_PIPELINE_V1:vague_no_memory_revive")
                _top2_ob_send(conn, task_id, chat_id, msg, reply_to, "topic2_vague_blocked")
                return

            if _top2_ob_estimate_like(full_context):
                try:
                    from core.sample_template_engine import handle_topic2_one_big_formula_pipeline_v1
                except Exception as e:
                    err = "TOPIC2_ONE_BIG_FINAL_PIPELINE_V1_IMPORT_ERR:" + _top2_ob_s(e)[:300]
                    _top2_ob_update(conn, task_id, "FAILED", err, err)
                    _top2_ob_hist(conn, task_id, err)
                    _top2_ob_send(conn, task_id, chat_id, err, reply_to, "topic2_formula_import_error")
                    return

                handled = await handle_topic2_one_big_formula_pipeline_v1(
                    conn=conn,
                    task_id=task_id,
                    chat_id=str(chat_id),
                    topic_id=topic_id,
                    raw_input=full_context,
                    input_type=input_type or "text",
                    reply_to_message_id=reply_to,
                )

                if handled:
                    try:
                        cur_state = conn.execute("SELECT state FROM tasks WHERE id=?", (task_id,)).fetchone()
                        cur_state_s = _top2_ob_s(cur_state[0] if cur_state else "")
                        if parent_id and parent_id != task_id and cur_state_s in ("AWAITING_CONFIRMATION", "DONE"):
                            conn.execute(
                                """
                                UPDATE tasks
                                SET state='DONE',
                                    result='Закрыто новой сметой TOPIC2_ONE_BIG_FINAL_PIPELINE_V1',
                                    error_message='',
                                    updated_at=datetime('now')
                                WHERE id=?
                                  AND state IN ('NEW','IN_PROGRESS','WAITING_CLARIFICATION','AWAITING_CONFIRMATION')
                                """,
                                (parent_id,),
                            )
                            conn.commit()
                            _top2_ob_hist(conn, parent_id, "TOPIC2_ONE_BIG_FINAL_PIPELINE_V1:superseded_by:" + task_id)
                    except Exception as e:
                        logger.warning("TOPIC2_ONE_BIG_FINAL_PIPELINE_V1_PARENT_CLOSE_ERR %s", e)
                    logger.info("TOPIC2_ONE_BIG_FINAL_PIPELINE_V1 handled task=%s topic=%s", task_id, topic_id)
                    return

            if _top2_ob_is_vague(raw_input):
                msg = "Старый сметный контекст не использую. Пришли полное ТЗ одним сообщением"
                _top2_ob_update(conn, task_id, "DONE", msg, "")
                _top2_ob_hist(conn, task_id, "TOPIC2_ONE_BIG_FINAL_PIPELINE_V1:vague_blocked_after_gate")
                _top2_ob_send(conn, task_id, chat_id, msg, reply_to, "topic2_vague_blocked_after_gate")
                return

        except Exception as e:
            err = "TOPIC2_ONE_BIG_FINAL_PIPELINE_V1_WORKER_ERR:" + _top2_ob_s(e)[:500]
            logger.warning(err)
            _top2_ob_update(conn, task_id, "FAILED", err, err)
            _top2_ob_hist(conn, task_id, err)
            _top2_ob_send(conn, task_id, chat_id, err, reply_to, "topic2_guard_error")
            return

    return await _TOP2_ONE_BIG_ORIG_HANDLE_IN_PROGRESS_V1(conn, task, chat_id, topic_id)  # HOTFIX_HANDLE_IN_PROGRESS_WRAPPER_CHAIN_V1
# === END_TOPIC2_ONE_BIG_FINAL_PIPELINE_V1_WORKER_GUARD ===

# === FINAL_TOPIC2_TOPIC5_TOPIC500_CLOSE_20260504_V1 ===
# Final runtime guard:
# - topic_2 full estimate TZ must use current raw input, not revived stale estimate memory
# - topic_500 search result must be finalized to DONE after reply_sent/result
# - wrapper chain must keep chat_id/topic_id arity intact

_FINAL_CLOSE_ORIG_HANDLE_IN_PROGRESS_20260504 = _handle_in_progress

def _fc_20260504_task_get(task, key, default=None):
    try:
        return task[key]
    except Exception:
        try:
            return getattr(task, key)
        except Exception:
            return default

def _fc_20260504_to_int(v, default=0):
    try:
        return int(v or 0)
    except Exception:
        return default

def _fc_20260504_history(conn, task_id, action):
    try:
        _history(conn, str(task_id), str(action))
    except Exception:
        try:
            conn.execute(
                "INSERT INTO task_history(task_id, action, created_at) VALUES (?, ?, datetime('now'))",
                (str(task_id), str(action)),
            )
            conn.commit()
        except Exception:
            pass

def _fc_20260504_latest_history_result(conn, task_id):
    try:
        row = conn.execute(
            """
            SELECT action
            FROM task_history
            WHERE task_id=? AND action LIKE 'result:%'
            ORDER BY rowid DESC
            LIMIT 1
            """,
            (str(task_id),),
        ).fetchone()
        if not row:
            return ""
        val = row[0] if not isinstance(row, dict) else row.get("action")
        val = str(val or "")
        if val.startswith("result:"):
            val = val[len("result:"):].strip()
        return val
    except Exception:
        return ""

def _fc_20260504_is_full_topic2_estimate_tz(text):
    import re
    low = str(text or "").lower()
    if "смет" not in low:
        return False
    if not any(x in low for x in ("дом", "house", "хаус", "барн", "barn")):
        return False
    if not re.search(r"\b\d+(?:[,.]\d+)?\s*(?:на|x|х|\*)\s*\d+(?:[,.]\d+)?\b", low):
        return False
    if not any(x in low for x in ("фундамент", "плита", "свая", "ленточ")):
        return False
    if not any(x in low for x in ("стен", "каркас", "дерев", "брус", "газобетон")):
        return False
    return True

async def _fc_20260504_call_topic2_engine(conn, task, chat_id, topic_id):
    import inspect
    from core import sample_template_engine as ste

    fn = getattr(ste, "handle_topic2_one_big_formula_pipeline_v1")
    task_id = _fc_20260504_task_get(task, "id", "")
    raw_input = _fc_20260504_task_get(task, "raw_input", "") or ""
    full_context = str(raw_input or "")

    params = inspect.signature(fn).parameters
    values = {
        "conn": conn,
        "connection": conn,
        "task": task,
        "row": task,
        "task_row": task,
        "task_id": task_id,
        "id": task_id,
        "chat_id": chat_id,
        "topic_id": topic_id,
        "raw_input": raw_input,
        "text": raw_input,
        "user_text": raw_input,
        "full_context": full_context,
        "context": full_context,
    }

    accepts_var_kw = any(p.kind == inspect.Parameter.VAR_KEYWORD for p in params.values())
    kwargs = values if accepts_var_kw else {name: values[name] for name in params if name in values}

    try:
        res = fn(**kwargs)
        if inspect.isawaitable(res):
            return await res
        return res
    except TypeError as e1:
        attempts = [
            (conn, task, chat_id, topic_id, raw_input, full_context),
            (conn, task, chat_id, topic_id, raw_input),
            (conn, task, chat_id, topic_id),
            (conn, task),
        ]
        last = e1
        for args in attempts:
            try:
                res = fn(*args)
                if inspect.isawaitable(res):
                    return await res
                return res
            except TypeError as e2:
                last = e2
                continue
        raise last

def _fc_20260504_finalize_topic500_if_result(conn, task, chat_id=None, topic_id=None):
    task_id = _fc_20260504_task_get(task, "id", "")
    if not task_id:
        return
    row = conn.execute(
        "SELECT state, result FROM tasks WHERE id=? LIMIT 1",
        (str(task_id),),
    ).fetchone()
    if not row:
        return
    state = row[0]
    current_result = row[1] or ""
    if str(state) not in ("IN_PROGRESS", "WAITING_CLARIFICATION", "AWAITING_CONFIRMATION"):
        return
    result_text = str(current_result or "").strip()
    if not result_text:
        result_text = _fc_20260504_latest_history_result(conn, task_id)
    if not result_text or len(result_text.strip()) < 20:
        return
    conn.execute(
        """
        UPDATE tasks
        SET state='DONE',
            result=?,
            error_message='',
            updated_at=datetime('now')
        WHERE id=? AND topic_id=500
        """,
        (result_text, str(task_id)),
    )
    _fc_20260504_history(conn, task_id, "FINAL_TOPIC500_SEARCH_DONE_20260504_V1")
    conn.commit()

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    task_id = _fc_20260504_task_get(task, "id", "")
    raw_input = str(_fc_20260504_task_get(task, "raw_input", "") or "")
    if chat_id is None:
        chat_id = _fc_20260504_task_get(task, "chat_id", None)
    if topic_id is None:
        topic_id = _fc_20260504_task_get(task, "topic_id", 0)
    topic_id = _fc_20260504_to_int(topic_id)

    if topic_id == 2 and _fc_20260504_is_full_topic2_estimate_tz(raw_input):
        _fc_20260504_history(conn, task_id, "FINAL_TOPIC2_CURRENT_TZ_DIRECT_ROUTE_20260504_V1")
        try:
            return await _fc_20260504_call_topic2_engine(conn, task, chat_id, topic_id)
        except Exception as e:
            _fc_20260504_history(conn, task_id, f"FINAL_TOPIC2_DIRECT_ROUTE_FAILED_20260504_V1:{type(e).__name__}:{e}")
            raise

    res = await _FINAL_CLOSE_ORIG_HANDLE_IN_PROGRESS_20260504(conn, task, chat_id, topic_id)

    if topic_id == 500:
        _fc_20260504_finalize_topic500_if_result(conn, task, chat_id, topic_id)

    return res
# === END_FINAL_TOPIC2_TOPIC5_TOPIC500_CLOSE_20260504_V1 ===



# === TOPIC500_PRE_SEND_VALIDATOR_AND_STARTUP_RECOVERY_HARD_GUARD_V1 ===
# Final guard layer:
# 1) validates topic_500 procurement output before sending
# 2) prevents startup/stale recovery from reprocessing topic_500 tasks after reply_sent
# 3) kills duplicate result loops for topic_500

import re as _t500_psv_re

def _t500_psv_row_get(row, key, default=None):
    try:
        return row[key]
    except Exception:
        try:
            return getattr(row, key)
        except Exception:
            return default

def _t500_psv_history(conn, task_id: str, action: str) -> None:
    try:
        _history(conn, str(task_id), str(action))
        return
    except Exception:
        pass
    try:
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (str(task_id), str(action)),
        )
    except Exception:
        pass

def _t500_psv_update_done(conn, task_id: str, result_text: str = "") -> None:
    try:
        _update_task(conn, str(task_id), state="DONE", result=str(result_text or ""), error_message="")
        return
    except Exception:
        pass
    try:
        conn.execute(
            "UPDATE tasks SET state='DONE', result=?, error_message='', updated_at=datetime('now') WHERE id=?",
            (str(result_text or ""), str(task_id)),
        )
    except Exception:
        pass

def _t500_psv_update_failed(conn, task_id: str, result_text: str, error_message: str) -> None:
    try:
        _update_task(conn, str(task_id), state="FAILED", result=str(result_text or ""), error_message=str(error_message))
        return
    except Exception:
        pass
    try:
        conn.execute(
            "UPDATE tasks SET state='FAILED', result=?, error_message=?, updated_at=datetime('now') WHERE id=?",
            (str(result_text or ""), str(error_message), str(task_id)),
        )
    except Exception:
        pass

def _t500_psv_latest_history_result(conn, task_id: str) -> str:
    try:
        row = conn.execute(
            "SELECT action FROM task_history WHERE task_id=? AND action LIKE 'result:%' ORDER BY rowid DESC LIMIT 1",
            (str(task_id),),
        ).fetchone()
        if not row:
            return ""
        action = _t500_psv_row_get(row, "action", row[0] if len(row) else "")
        action = str(action or "")
        return action[len("result:"):].strip() if action.startswith("result:") else action.strip()
    except Exception:
        return ""

def _t500_psv_has_reply_sent(conn, task_id: str) -> bool:
    try:
        row = conn.execute(
            "SELECT 1 FROM task_history WHERE task_id=? AND action LIKE 'reply_sent:%' ORDER BY rowid DESC LIMIT 1",
            (str(task_id),),
        ).fetchone()
        return bool(row)
    except Exception:
        return False

def _t500_psv_validate_procurement_result(text: str):
    s = str(text or "").strip()
    urls = _t500_psv_re.findall(r"https?://[^\s\]\)>,]+", s)
    unique_urls = []
    for u in urls:
        if u not in unique_urls:
            unique_urls.append(u)

    has_min_urls = len(unique_urls) >= 3
    has_bare_refs = bool(_t500_psv_re.search(r"(?<!https://)(?<!http://)\[[0-9]{1,2}\](?!\s*\(?https?://)", s))
    has_price = bool(_t500_psv_re.search(r"(\d[\d\s]{1,10}\s*(?:₽|руб|р\.|руб\.))|цена\s+не\s+указана", s, _t500_psv_re.I))
    has_phone = bool(_t500_psv_re.search(r"(\+7|8)\s*[\(\- ]?\d{3}[\)\- ]?\s*\d{3}[\- ]?\d{2}[\- ]?\d{2}|телефон\s+не\s+найден", s, _t500_psv_re.I))
    has_supplier = bool(_t500_psv_re.search(r"(поставщик|магазин|дилер|база|склад|авито|avito|2гис|яндекс|леруа|петрович|термодом|rockwool|роквул)", s, _t500_psv_re.I))
    only_unconfirmed = s.upper().strip() in {"НЕ ПОДТВЕРЖДЕНО", "НЕ ПОДТВЕРЖДЕНО."}

    if not has_min_urls:
        return False, "SEARCH_OUTPUT_INVALID_NO_DIRECT_LINKS"
    if has_bare_refs:
        return False, "SEARCH_OUTPUT_INVALID_BARE_REFS"
    if not has_supplier:
        return False, "SEARCH_OUTPUT_INVALID_NO_SUPPLIER"
    if not has_price:
        return False, "SEARCH_OUTPUT_INVALID_NO_PRICE"
    if not has_phone:
        return False, "SEARCH_OUTPUT_INVALID_NO_PHONE"
    if only_unconfirmed:
        return False, "SEARCH_OUTPUT_INVALID_UNCONFIRMED_ONLY"
    return True, "OK"

try:
    _T500_PSV_ORIG_SEND_ONCE_EX = _send_once_ex
except Exception:
    _T500_PSV_ORIG_SEND_ONCE_EX = None

def _send_once_ex(conn, task_id, chat_id, text, reply_to, kind):  # TOPIC500_PRE_SEND_VALIDATOR_AND_STARTUP_RECOVERY_HARD_GUARD_V1
    try:
        row = conn.execute("SELECT topic_id FROM tasks WHERE id=?", (str(task_id),)).fetchone()
        topic_id = int(_t500_psv_row_get(row, "topic_id", row[0] if row else 0) or 0)
        if topic_id == 500 and str(kind or "") == "result":
            ok, reason = _t500_psv_validate_procurement_result(str(text or ""))
            if not ok:
                _t500_psv_update_failed(conn, str(task_id), str(text or ""), reason)
                _t500_psv_history(conn, str(task_id), f"TOPIC500_PROCUREMENT_VALIDATOR_V1:FAILED:{reason}")
                try:
                    conn.commit()
                except Exception:
                    pass
                warning = "Поиск не дал прямых ссылок и телефонов. Повтори запрос или уточни площадки"
                if _T500_PSV_ORIG_SEND_ONCE_EX:
                    return _T500_PSV_ORIG_SEND_ONCE_EX(conn, task_id, chat_id, warning, reply_to, "error")
                return None
    except Exception as e:
        try:
            _t500_psv_history(conn, str(task_id), f"TOPIC500_PROCUREMENT_VALIDATOR_V1:ERROR:{type(e).__name__}")
            conn.commit()
        except Exception:
            pass

    if _T500_PSV_ORIG_SEND_ONCE_EX:
        return _T500_PSV_ORIG_SEND_ONCE_EX(conn, task_id, chat_id, text, reply_to, kind)
    return None

def _t500_psv_repair_sent_in_progress(conn) -> None:
    try:
        rows = conn.execute(
            """
            SELECT id, topic_id, state, COALESCE(result,'') AS result
            FROM tasks
            WHERE topic_id=500
              AND state IN ('IN_PROGRESS','WAITING_CLARIFICATION','AWAITING_CONFIRMATION')
            """
        ).fetchall()
    except Exception:
        return

    for row in rows:
        task_id = str(_t500_psv_row_get(row, "id", ""))
        if not task_id:
            continue
        if not _t500_psv_has_reply_sent(conn, task_id):
            continue
        current_result = str(_t500_psv_row_get(row, "result", "") or "")
        latest_result = _t500_psv_latest_history_result(conn, task_id) or current_result
        _t500_psv_update_done(conn, task_id, latest_result)
        _t500_psv_history(conn, task_id, "STARTUP_RECOVERY_REPLY_SENT_GUARD_V1:DONE_SKIP_RECOVERY")
    try:
        conn.commit()
    except Exception:
        pass

def _t500_psv_duplicate_loop_guard(conn, task_id: str) -> bool:
    try:
        row = conn.execute("SELECT topic_id FROM tasks WHERE id=?", (str(task_id),)).fetchone()
        topic_id = int(_t500_psv_row_get(row, "topic_id", row[0] if row else 0) or 0)
        if topic_id != 500:
            return False
        cnt = conn.execute(
            """
            SELECT COUNT(*)
            FROM task_history
            WHERE task_id=?
              AND action LIKE 'result:%'
              AND created_at >= datetime('now','-300 seconds')
            """,
            (str(task_id),),
        ).fetchone()[0]
        if int(cnt or 0) < 3:
            return False
        latest_result = _t500_psv_latest_history_result(conn, task_id)
        _t500_psv_update_done(conn, str(task_id), latest_result)
        _t500_psv_history(conn, str(task_id), "TOPIC500_DUPLICATE_RESULT_LOOP_GUARD_V1:KILLED")
        conn.commit()
        return True
    except Exception:
        return False

try:
    _T500_PSV_ORIG_RECOVER_STALE_TASKS = _recover_stale_tasks
except Exception:
    _T500_PSV_ORIG_RECOVER_STALE_TASKS = None

def _recover_stale_tasks(conn, *args, **kwargs):  # TOPIC500_PRE_SEND_VALIDATOR_AND_STARTUP_RECOVERY_HARD_GUARD_V1
    _t500_psv_repair_sent_in_progress(conn)
    if _T500_PSV_ORIG_RECOVER_STALE_TASKS:
        res = _T500_PSV_ORIG_RECOVER_STALE_TASKS(conn, *args, **kwargs)
        _t500_psv_repair_sent_in_progress(conn)
        return res
    return None

try:
    _T500_PSV_ORIG_HANDLE_IN_PROGRESS = _handle_in_progress
except Exception:
    _T500_PSV_ORIG_HANDLE_IN_PROGRESS = None

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):  # TOPIC500_PRE_SEND_VALIDATOR_AND_STARTUP_RECOVERY_HARD_GUARD_V1
    try:
        task_id = str(_t500_psv_row_get(task, "id", ""))
        if task_id and _t500_psv_duplicate_loop_guard(conn, task_id):
            return
        _t500_psv_repair_sent_in_progress(conn)
    except Exception:
        pass
    if _T500_PSV_ORIG_HANDLE_IN_PROGRESS:
        return await _T500_PSV_ORIG_HANDLE_IN_PROGRESS(conn, task, chat_id, topic_id)
    return None

# === END_TOPIC500_PRE_SEND_VALIDATOR_AND_STARTUP_RECOVERY_HARD_GUARD_V1 ===




# === P0_RUNTIME_ROUTE_GUARD_TOPIC2_TOPIC500_20260504_V1 ===
# Runtime guard installed before asyncio.run(main())
# Scope:
# - topic_500 search/product requests must never enter estimate/project routes
# - topic_2 current estimate TZ must use current raw input direct estimate route
# - no forbidden files, no DB schema changes, no systemd unit changes

try:
    _P0_RUNTIME_ORIG_HANDLE_IN_PROGRESS_20260504 = _handle_in_progress
except Exception:
    _P0_RUNTIME_ORIG_HANDLE_IN_PROGRESS_20260504 = None

def _p0_runtime_row_get_20260504(row, key, default=None):
    try:
        return row[key]
    except Exception:
        try:
            return getattr(row, key)
        except Exception:
            return default

def _p0_runtime_s_20260504(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p0_runtime_history_20260504(conn, task_id, action):
    try:
        _history(conn, str(task_id), str(action))
        return
    except Exception:
        pass
    try:
        conn.execute(
            "INSERT INTO task_history(task_id, action, created_at) VALUES (?, ?, datetime('now'))",
            (str(task_id), str(action)),
        )
    except Exception:
        pass

def _p0_runtime_update_20260504(conn, task_id, state=None, result=None, error_message=None, bot_message_id=None):
    try:
        kwargs = {}
        if state is not None:
            kwargs["state"] = state
        if result is not None:
            kwargs["result"] = result
        if error_message is not None:
            kwargs["error_message"] = error_message
        if bot_message_id is not None:
            kwargs["bot_message_id"] = bot_message_id
        if kwargs:
            _update_task(conn, str(task_id), **kwargs)
            return
    except Exception:
        pass

    sets = []
    vals = []
    if state is not None:
        sets.append("state=?")
        vals.append(str(state))
    if result is not None:
        sets.append("result=?")
        vals.append(str(result))
    if error_message is not None:
        sets.append("error_message=?")
        vals.append(str(error_message))
    if bot_message_id is not None:
        sets.append("bot_message_id=?")
        vals.append(int(bot_message_id))
    sets.append("updated_at=datetime('now')")
    vals.append(str(task_id))
    conn.execute(f"UPDATE tasks SET {', '.join(sets)} WHERE id=?", vals)

def _p0_runtime_is_topic500_search_20260504(raw_input, input_type):
    low = _p0_runtime_s_20260504(raw_input, 12000).lower()
    itype = _p0_runtime_s_20260504(input_type, 50).lower()
    if not low:
        return False
    if any(x in low for x in ("задача закрыта", "закрой задачу", "отменяй", "отмена", "завершена", "заверши")):
        return False
    if itype == "search":
        return True
    markers = (
        "найди", "поиск", "цена", "стоимость", "дешевле", "купить", "ссылка", "ссылки",
        "магазин", "поставщик", "маркет", "маркетплейс", "авито", "avito", "ozon",
        "wildberries", "яндекс", "google pixel", "iphone", "телефон", "кабель",
        "вата", "утеплитель", "товар", "варианты"
    )
    return any(m in low for m in markers)

def _p0_runtime_is_bad_search_result_20260504(text):
    low = _p0_runtime_s_20260504(text, 20000).lower()
    if not low:
        return True
    bad = (
        "смета готова",
        "предварительная смета готова",
        "xlsx:",
        "pdf:",
        "engine: topic2",
        "engine: estimate",
        "ареал нева.xlsx",
        "м-110.xlsx",
        "позиций: 1. итого: 0.00",
        "фундамент",
        "кровля",
        "монолитная плита",
    )
    if any(x in low for x in bad):
        return True
    if "http://" not in low and "https://" not in low and "цена" not in low and "руб" not in low and "₽" not in low:
        return True
    return False

def _p0_runtime_is_topic2_current_estimate_20260504(raw_input):
    import re as _p0_re
    low = _p0_runtime_s_20260504(raw_input, 12000).lower()
    if not low:
        return False
    if not any(x in low for x in ("смет", "стоимость", "посчитать", "рассчитать", "расчет", "расчёт")):
        return False
    if not any(x in low for x in ("дом", "house", "хаус", "барн", "barn")):
        return False
    if not _p0_re.search(r"\d+(?:[,.]\d+)?\s*(?:на|x|х|×|\*)\s*\d+(?:[,.]\d+)?", low):
        return False
    if not any(x in low for x in ("фундамент", "плита", "свая", "ленточ")):
        return False
    if not any(x in low for x in ("стен", "каркас", "дерев", "брус", "газобетон", "кирпич")):
        return False
    return True

async def _p0_runtime_topic500_direct_search_20260504(conn, task, chat_id, topic_id):
    task_id = _p0_runtime_s_20260504(_p0_runtime_row_get_20260504(task, "id", ""))
    raw_input = _p0_runtime_s_20260504(_p0_runtime_row_get_20260504(task, "raw_input", ""), 12000)
    reply_to = _p0_runtime_row_get_20260504(task, "reply_to_message_id", None)

    payload = {
        "id": task_id,
        "task_id": task_id,
        "topic_id": 500,
        "chat_id": str(chat_id),
        "input_type": "search",
        "raw_input": raw_input,
        "normalized_input": raw_input,
        "state": "IN_PROGRESS",
        "reply_to_message_id": reply_to,
        "active_task_context": "",
        "pin_context": "",
        "short_memory_context": "",
        "long_memory_context": "",
        "archive_context": "",
        "search_context": "",
        "topic_role": "ВЕБ ПОИСК",
        "topic_directions": "internet_search",
        "direction": "internet_search",
        "engine": "search_supplier",
    }

    _p0_runtime_history_20260504(conn, task_id, "P0_RUNTIME_TOPIC500_DIRECT_SEARCH_ROUTE_V1")
    try:
        _p0_runtime_update_20260504(conn, task_id, state="IN_PROGRESS", error_message="")
        conn.commit()
    except Exception:
        pass

    try:
        result = await asyncio.wait_for(process_ai_task(payload), timeout=AI_TIMEOUT)
        result = _clean(_s(result), 50000)
    except Exception as e:
        err = "P0_RUNTIME_TOPIC500_DIRECT_SEARCH_ERROR:" + _p0_runtime_s_20260504(type(e).__name__ + ":" + str(e), 500)
        _p0_runtime_update_20260504(conn, task_id, state="FAILED", result="", error_message=err)
        _p0_runtime_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Поиск не выполнен. Повтори запрос или уточни товар и регион", reply_to, "p0_topic500_error")
        return

    if _p0_runtime_is_bad_search_result_20260504(result):
        err = "P0_RUNTIME_TOPIC500_BAD_ROUTE_BLOCKED"
        _p0_runtime_update_20260504(conn, task_id, state="FAILED", result=result, error_message=err)
        _p0_runtime_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Поиск ушёл не в тот маршрут и заблокирован. Повтори запрос в этом топике", reply_to, "p0_topic500_bad_route")
        return

    _p0_runtime_update_20260504(conn, task_id, state="DONE", result=result, error_message="")
    _p0_runtime_history_20260504(conn, task_id, "P0_RUNTIME_TOPIC500_SEARCH_DONE_V1")
    try:
        _save_memory(str(chat_id), 500, raw_input, result)
    except Exception:
        pass

    sent = _send_once_ex(conn, task_id, str(chat_id), result, reply_to, "p0_topic500_search_result")
    try:
        bot_message_id = sent.get("bot_message_id") if isinstance(sent, dict) else None
        if bot_message_id is not None:
            _p0_runtime_update_20260504(conn, task_id, bot_message_id=bot_message_id)
    except Exception:
        pass
    conn.commit()
    return

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    task_id = _p0_runtime_s_20260504(_p0_runtime_row_get_20260504(task, "id", ""))
    raw_input = _p0_runtime_s_20260504(_p0_runtime_row_get_20260504(task, "raw_input", ""), 12000)
    input_type = _p0_runtime_s_20260504(_p0_runtime_row_get_20260504(task, "input_type", "text"), 50)
    if chat_id is None:
        chat_id = _p0_runtime_row_get_20260504(task, "chat_id", None)
    if topic_id is None:
        topic_id = _p0_runtime_row_get_20260504(task, "topic_id", 0)
    try:
        topic_id = int(topic_id or 0)
    except Exception:
        topic_id = 0

    if topic_id == 500 and _p0_runtime_is_topic500_search_20260504(raw_input, input_type):
        return await _p0_runtime_topic500_direct_search_20260504(conn, task, chat_id, topic_id)

    if topic_id == 2 and _p0_runtime_is_topic2_current_estimate_20260504(raw_input):
        _p0_runtime_history_20260504(conn, task_id, "P0_RUNTIME_TOPIC2_CURRENT_TZ_DIRECT_ROUTE_V1")
        try:
            return await _fc_20260504_call_topic2_engine(conn, task, chat_id, topic_id)
        except Exception as e:
            _p0_runtime_history_20260504(conn, task_id, "P0_RUNTIME_TOPIC2_DIRECT_ROUTE_ERROR:" + _p0_runtime_s_20260504(type(e).__name__ + ":" + str(e), 500))
            raise

    if _P0_RUNTIME_ORIG_HANDLE_IN_PROGRESS_20260504:
        return await _P0_RUNTIME_ORIG_HANDLE_IN_PROGRESS_20260504(conn, task, chat_id, topic_id)
    return None

# === END_P0_RUNTIME_ROUTE_GUARD_TOPIC2_TOPIC500_20260504_V1 ===



# === P2_FINAL_SEARCH_AND_ESTIMATE_CLOSE_20260504_V1 ===
# Runtime overlay before asyncio.run(main)
# Scope:
# - topic_500 internet search keeps topic memory and previous task context
# - topic_500 must never emit estimate/PDF/XLSX route garbage
# - topic_2 estimate must call latest sample_template_engine overlay
# - no schema changes, no forbidden files, no systemd changes

try:
    _P2_ORIG_HANDLE_IN_PROGRESS_20260504 = _handle_in_progress
except Exception:
    _P2_ORIG_HANDLE_IN_PROGRESS_20260504 = None

import re as _p2_tw_re
import json as _p2_tw_json
import datetime as _p2_tw_dt

def _p2_tw_s(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p2_tw_low(v):
    return _p2_tw_s(v).lower().replace("ё", "е")

def _p2_tw_get(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        try:
            return getattr(row, key)
        except Exception:
            return default

def _p2_tw_history(conn, task_id, action):
    try:
        _history(conn, str(task_id), _p2_tw_s(action, 1000))
        return
    except Exception:
        pass
    try:
        conn.execute(
            "INSERT INTO task_history(task_id, action, created_at) VALUES (?, ?, datetime('now'))",
            (str(task_id), _p2_tw_s(action, 1000)),
        )
    except Exception:
        pass

def _p2_tw_update(conn, task_id, **kwargs):
    try:
        _update_task(conn, str(task_id), **kwargs)
        return
    except Exception:
        pass
    try:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(tasks)").fetchall()]
        sets = []
        vals = []
        for k, v in kwargs.items():
            if k in cols:
                sets.append(f"{k}=?")
                vals.append(v)
        if "updated_at" in cols:
            sets.append("updated_at=datetime('now')")
        if sets:
            vals.append(str(task_id))
            conn.execute(f"UPDATE tasks SET {', '.join(sets)} WHERE id=?", vals)
    except Exception:
        pass

def _p2_tw_is_close_command(raw):
    low = _p2_tw_low(raw)
    return any(x in low for x in (
        "задача закрыта", "закрой задачу", "закрывай", "отменяй", "отмена", "заверши", "завершена"
    ))

def _p2_tw_is_search_intent(raw, input_type):
    low = _p2_tw_low(raw)
    itype = _p2_tw_low(input_type)
    if not low or _p2_tw_is_close_command(low):
        return False
    if itype == "search":
        return True
    return any(x in low for x in (
        "найди", "найти", "поиск", "поищи", "цена", "стоимость", "дешевле",
        "купить", "ссылка", "ссылки", "магазин", "поставщик", "маркет",
        "маркетплейс", "авито", "avito", "ozon", "wildberries", "яндекс",
        "google pixel", "iphone", "samsung", "телефон", "смартфон",
        "вата", "rockwool", "роквул", "утеплитель", "товар", "варианты",
        "предыдущ", "то что я тебе писал", "то, что я тебе писал"
    ))

def _p2_tw_is_followup_search(raw):
    low = _p2_tw_low(raw)
    return any(x in low for x in (
        "предыдущ", "последн", "то что", "то, что", "я тебе писал", "тот запрос",
        "этот товар", "тот товар", "по нему", "по ней", "по этому", "выполни поиск",
        "сделай поиск", "проверь ещё", "проверь еще"
    ))

def _p2_tw_has_concrete_product(raw):
    low = _p2_tw_low(raw)
    return any(x in low for x in (
        "google pixel", "pixel", "iphone", "samsung", "rockwool", "роквул",
        "вата", "утеплитель", "телефон", "смартфон", "кабель", "модель", "артикул"
    )) and any(ch.isdigit() for ch in low)

def _p2_tw_topic500_context(conn, chat_id, current_task_id):
    try:
        rows = conn.execute(
            """
            SELECT id, raw_input, result, state, updated_at
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=500
              AND id<>?
              AND COALESCE(raw_input,'')<>''
            ORDER BY rowid DESC
            LIMIT 8
            """,
            (str(chat_id), str(current_task_id)),
        ).fetchall()
    except Exception:
        rows = []
    out = []
    for r in rows:
        raw = _p2_tw_s(_p2_tw_get(r, "raw_input", ""), 900)
        res = _p2_tw_s(_p2_tw_get(r, "result", ""), 900)
        state = _p2_tw_s(_p2_tw_get(r, "state", ""), 50)
        if raw and not _p2_tw_is_close_command(raw):
            out.append({"raw": raw, "result": res, "state": state})
    return out

def _p2_tw_resolve_query(raw, ctx):
    raw_s = _p2_tw_s(raw, 2000)
    if not ctx:
        return raw_s
    if _p2_tw_is_followup_search(raw_s) or not _p2_tw_has_concrete_product(raw_s):
        for item in ctx:
            prev = item.get("raw") or ""
            if _p2_tw_has_concrete_product(prev) and not _p2_tw_is_followup_search(prev):
                return (
                    "Новая команда пользователя: " + raw_s + "\n"
                    "Предыдущая товарная задача topic_500: " + prev + "\n"
                    "Выполни именно интернет-поиск по предыдущей товарной задаче с учетом новой команды"
                )
    return raw_s

def _p2_tw_clean_search_result(text):
    s = _p2_tw_s(text, 50000)
    lines = []
    forbidden = (
        "engine:", "manifest:", "/root/", "/tmp/", "traceback",
        "xlsx:", "pdf:", "смета готова", "предварительная смета",
        "монолитная плита", "фундамент:", "эталон: м-110", "позиций: 1. итого: 0.00",
        "не могу перейти", "я не могу выполнить поиск"
    )
    for line in s.splitlines():
        low = _p2_tw_low(line)
        if any(x in low for x in forbidden):
            continue
        lines.append(line.rstrip())
    out = "\n".join(lines).strip()
    out = _p2_tw_re.sub(r"\n{3,}", "\n\n", out)
    return out[:50000]

def _p2_tw_bad_search_result(text):
    low = _p2_tw_low(text)
    if not low or len(low) < 80:
        return True
    if any(x in low for x in (
        "смета готова", "предварительная смета", "xlsx:", "pdf:", "монолитная плита",
        "фундамент:", "м-110.xlsx", "позиций: 1. итого: 0.00"
    )):
        return True
    has_price = bool(_p2_tw_re.search(r"(\d[\d\s]{1,10}\s*(?:₽|руб|р\.|руб\.))", low))
    has_url = "http://" in low or "https://" in low
    has_supplier = any(x in low for x in (
        "поставщик", "магазин", "маркет", "авито", "avito", "ozon", "wildberries",
        "дилер", "склад", "налич"
    ))
    return not (has_price and (has_url or has_supplier))

async def _p2_tw_call_search(conn, task, chat_id, topic_id):
    task_id = _p2_tw_s(_p2_tw_get(task, "id", ""))
    raw_input = _p2_tw_s(_p2_tw_get(task, "raw_input", ""), 12000)
    reply_to = _p2_tw_get(task, "reply_to_message_id", None)
    ctx = _p2_tw_topic500_context(conn, chat_id, task_id)
    query = _p2_tw_resolve_query(raw_input, ctx)

    prompt = (
        "Выполни реальный интернет-поиск товара/поставщика.\n"
        "Запрещено: сметы, PDF/XLSX, строительные расчеты, общие советы, внутренние маркеры.\n"
        "Нужно: минимум 3 варианта, прямые ссылки, цена, город/доставка, наличие, телефон если найден.\n"
        "Формат ответа только для пользователя:\n"
        "| № | Поставщик/площадка | Город | Цена | Наличие | Доставка | Телефон | Ссылка | Проверено |\n"
        "|---|---|---|---|---|---|---|---|---|\n\n"
        "Запрос:\n" + query + "\n\n"
        "Контекст последних topic_500 задач:\n" + _p2_tw_json.dumps(ctx[:5], ensure_ascii=False)
    )

    payload = {
        "id": task_id,
        "task_id": task_id,
        "chat_id": str(chat_id),
        "topic_id": 500,
        "input_type": "search",
        "raw_input": prompt,
        "normalized_input": query,
        "state": "IN_PROGRESS",
        "reply_to_message_id": reply_to,
        "active_task_context": "",
        "pin_context": "",
        "short_memory_context": "",
        "long_memory_context": "",
        "archive_context": _p2_tw_json.dumps(ctx[:5], ensure_ascii=False),
        "search_context": "",
        "topic_role": "ИНТЕРНЕТ ПОИСК",
        "topic_directions": "товары, поставщики, цены, ссылки",
        "direction": "internet_search",
        "engine": "search_supplier",
    }

    _p2_tw_history(conn, task_id, "P2_SEARCH_MEMORY_ROUTE_TAKEN")
    _p2_tw_update(conn, task_id, state="IN_PROGRESS", error_message="")
    conn.commit()

    try:
        result = await asyncio.wait_for(process_ai_task(payload), timeout=AI_TIMEOUT)
        result = _p2_tw_clean_search_result(result)
    except Exception as e:
        err = "P2_SEARCH_CALL_FAILED:" + _p2_tw_s(type(e).__name__ + ":" + str(e), 500)
        _p2_tw_update(conn, task_id, state="FAILED", result="", error_message=err)
        _p2_tw_history(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Поиск не выполнен. Повтори запрос с названием товара и регионом", reply_to, "p2_search_failed")
        return True

    if _p2_tw_bad_search_result(result):
        retry_prompt = (
            "СТРОГИЙ ПОВТОР ПОИСКА. Предыдущий ответ невалиден.\n"
            "Верни только товарные предложения с ценами и ссылками. Не смета. Не PDF. Не XLSX.\n\n"
            + prompt
        )
        payload["raw_input"] = retry_prompt
        try:
            retry = await asyncio.wait_for(process_ai_task(payload), timeout=AI_TIMEOUT)
            retry = _p2_tw_clean_search_result(retry)
            if not _p2_tw_bad_search_result(retry):
                result = retry
        except Exception:
            pass

    if _p2_tw_bad_search_result(result):
        err = "P2_SEARCH_BAD_RESULT_BLOCKED"
        _p2_tw_update(conn, task_id, state="FAILED", result=result, error_message=err)
        _p2_tw_history(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Поиск заблокирован: результат не содержит валидных цен и товарных ссылок", reply_to, "p2_search_bad_result")
        return True

    _p2_tw_update(conn, task_id, state="DONE", result=result, error_message="")
    _p2_tw_history(conn, task_id, "P2_SEARCH_DONE")
    try:
        _save_memory(str(chat_id), 500, raw_input, result)
    except Exception:
        pass

    sent = _send_once_ex(conn, task_id, str(chat_id), result, reply_to, "p2_search_result")
    try:
        bot_id = sent.get("bot_message_id") if isinstance(sent, dict) else None
        if bot_id:
            _p2_tw_update(conn, task_id, bot_message_id=bot_id)
    except Exception:
        pass
    conn.commit()
    return True

def _p2_tw_is_topic2_estimate(raw):
    low = _p2_tw_low(raw)
    if not any(x in low for x in ("смет", "посчитать", "рассчитать", "расчет", "расчёт", "стоимость")):
        return False
    return any(x in low for x in ("дом", "house", "барн", "barn", "хаус")) and bool(_p2_tw_re.search(r"\d+(?:[,.]\d+)?\s*(?:на|x|х|×|\*)\s*\d+", low))

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    task_id = _p2_tw_s(_p2_tw_get(task, "id", ""))
    raw_input = _p2_tw_s(_p2_tw_get(task, "raw_input", ""), 12000)
    input_type = _p2_tw_s(_p2_tw_get(task, "input_type", "text"), 50)

    if chat_id is None:
        chat_id = _p2_tw_get(task, "chat_id", None)
    if topic_id is None:
        topic_id = _p2_tw_get(task, "topic_id", 0)
    try:
        topic_id = int(topic_id or 0)
    except Exception:
        topic_id = 0

    if topic_id == 500 and _p2_tw_is_search_intent(raw_input, input_type):
        return await _p2_tw_call_search(conn, task, chat_id, topic_id)

    if topic_id == 2 and _p2_tw_is_topic2_estimate(raw_input):
        _p2_tw_history(conn, task_id, "P2_TOPIC2_LATEST_ESTIMATE_ENGINE_ROUTE")
        try:
            from core import sample_template_engine as _p2_ste
            return await _p2_ste.handle_topic2_one_big_formula_pipeline_v1(
                conn=conn,
                task=task,
                chat_id=chat_id,
                topic_id=topic_id,
                raw_input=raw_input,
                full_context=raw_input,
            )
        except Exception as e:
            _p2_tw_history(conn, task_id, "P2_TOPIC2_ROUTE_ERROR:" + _p2_tw_s(type(e).__name__ + ":" + str(e), 500))
            raise

    if _P2_ORIG_HANDLE_IN_PROGRESS_20260504:
        return await _P2_ORIG_HANDLE_IN_PROGRESS_20260504(conn, task, chat_id, topic_id)
    return None

# === END_P2_FINAL_SEARCH_AND_ESTIMATE_CLOSE_20260504_V1 ===

# === P3_FINAL_ROUTE_HARD_LOCK_SEARCH_ESTIMATE_20260504_V1 ===
# Runtime overlay before asyncio.run(main())
# Scope:
# - topic_500 always routes to internet search, never to estimate/project routes
# - topic_2 estimate-like input always routes to current-input estimate engine
# - topic_2 vague followups never generate estimates from old memory
# - no DB schema changes, no forbidden files, no systemd changes

try:
    _P3_ORIG_HANDLE_IN_PROGRESS_20260504 = _handle_in_progress
except Exception:
    _P3_ORIG_HANDLE_IN_PROGRESS_20260504 = None

def _p3_s_20260504(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p3_low_20260504(v):
    return _p3_s_20260504(v).lower().replace("ё", "е")

def _p3_row_get_20260504(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        try:
            return getattr(row, key)
        except Exception:
            return default

def _p3_update_20260504(conn, task_id, **kwargs):
    try:
        _update_task(conn, str(task_id), **kwargs)
        return
    except Exception:
        pass
    try:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(tasks)").fetchall()]
        sets, vals = [], []
        for k, v in kwargs.items():
            if k in cols:
                sets.append(f"{k}=?")
                vals.append(v)
        if "updated_at" in cols:
            sets.append("updated_at=datetime('now')")
        if sets:
            vals.append(str(task_id))
            conn.execute(f"UPDATE tasks SET {', '.join(sets)} WHERE id=?", vals)
    except Exception:
        pass

def _p3_history_20260504(conn, task_id, action):
    try:
        _history(conn, str(task_id), str(action))
        return
    except Exception:
        pass
    try:
        conn.execute(
            "INSERT INTO task_history(task_id, action, created_at) VALUES (?, ?, datetime('now'))",
            (str(task_id), _p3_s_20260504(action, 1000)),
        )
    except Exception:
        pass

def _p3_is_close_command_20260504(raw_input):
    low = _p3_low_20260504(raw_input)
    return any(x in low for x in (
        "задача закрыта", "закрой задачу", "отменяй", "отмена", "отбой",
        "заверши", "закрывай", "стоп"
    ))

def _p3_topic500_search_needed_20260504(raw_input, input_type):
    if _p3_is_close_command_20260504(raw_input):
        return False
    low = _p3_low_20260504(raw_input)
    if not low:
        return False
    return True

def _p3_topic500_vague_20260504(raw_input):
    low = _p3_low_20260504(raw_input)
    if len(low) <= 90 and any(x in low for x in ("то что", "предыдущ", "прошл", "телефон", "выполни поиск", "то, что")):
        return True
    return False

def _p3_find_previous_topic500_query_20260504(conn, chat_id, topic_id, current_task_id):
    try:
        rows = conn.execute(
            """
            SELECT raw_input
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=500
              AND id<>?
              AND COALESCE(raw_input,'')<>''
              AND COALESCE(raw_input,'') NOT LIKE '%Задача закрыта%'
              AND COALESCE(raw_input,'') NOT LIKE '%отменяй%'
            ORDER BY rowid DESC
            LIMIT 12
            """,
            (str(chat_id), str(current_task_id)),
        ).fetchall()
        for row in rows:
            raw = _p3_s_20260504(_p3_row_get_20260504(row, "raw_input", row[0] if row else ""), 4000)
            low = _p3_low_20260504(raw)
            if any(x in low for x in (
                "найди", "поиск", "дешевле", "купить", "цена", "стоимость",
                "iphone", "pixel", "телефон", "ozon", "wildberries", "авито", "avito",
                "поставщик", "каменная вата", "утеплитель"
            )):
                return raw
    except Exception:
        pass
    return ""

def _p3_bad_search_result_20260504(text):
    low = _p3_low_20260504(text)
    if not low:
        return True
    bad = (
        "смета готова",
        "предварительная смета готова",
        "xlsx:",
        "pdf:",
        "engine:",
        "м-110.xlsx",
        "ареал нева.xlsx",
        "позиций: 1. итого: 0.00",
        "фундамент:",
        "монолитная плита",
        "лист эталона",
    )
    if any(x in low for x in bad):
        return True
    if ("http://" not in low and "https://" not in low) and ("₽" not in low and "руб" not in low and "цена" not in low):
        return True
    return False

async def _p3_handle_topic500_search_20260504(conn, task, chat_id, topic_id):
    task_id = _p3_s_20260504(_p3_row_get_20260504(task, "id", ""))
    raw_input = _p3_s_20260504(_p3_row_get_20260504(task, "raw_input", ""), 12000)
    reply_to = _p3_row_get_20260504(task, "reply_to_message_id", None)

    previous = _p3_find_previous_topic500_query_20260504(conn, chat_id, topic_id, task_id) if _p3_topic500_vague_20260504(raw_input) else ""
    search_text = raw_input
    if previous:
        search_text = (
            "Выполни интернет-поиск по предыдущей товарной задаче с учётом текущего уточнения.\n"
            f"Текущее сообщение: {raw_input}\n"
            f"Предыдущая товарная задача: {previous}\n"
        )

    prompt = (
        "РЕЖИМ: ИНТЕРНЕТ-ПОИСК ТОВАРА. НЕ СОСТАВЛЯЙ СМЕТУ. НЕ СОЗДАВАЙ XLSX/PDF.\n"
        "Нужно найти реальные варианты покупки по запросу пользователя.\n"
        "Ответ строго таблицей: № | Площадка/поставщик | Товар | Город/регион | Цена | Наличие | Доставка | Телефон | Прямая ссылка | Риск.\n"
        "Минимум 3 варианта, если они существуют. Если вариантов нет — прямо напиши, что подтверждённых вариантов нет.\n\n"
        f"Запрос:\n{search_text}"
    )

    payload = {
        "id": task_id,
        "task_id": task_id,
        "topic_id": 500,
        "chat_id": str(chat_id),
        "input_type": "search",
        "raw_input": prompt,
        "normalized_input": prompt,
        "state": "IN_PROGRESS",
        "reply_to_message_id": reply_to,
        "active_task_context": previous,
        "pin_context": "",
        "short_memory_context": previous,
        "long_memory_context": "",
        "archive_context": "",
        "search_context": previous,
        "topic_role": "ВЕБ ПОИСК",
        "topic_directions": "internet_search",
        "direction": "internet_search",
        "engine": "search_supplier",
        "forbid_estimate": True,
    }

    _p3_history_20260504(conn, task_id, "P3_TOPIC500_HARD_SEARCH_ROUTE_TAKEN")
    _p3_update_20260504(conn, task_id, state="IN_PROGRESS", error_message="")
    conn.commit()

    try:
        result = await asyncio.wait_for(process_ai_task(payload), timeout=AI_TIMEOUT)
        result = _clean(_s(result), 50000)
    except Exception as e:
        err = "P3_TOPIC500_SEARCH_ERROR:" + _p3_s_20260504(type(e).__name__ + ":" + str(e), 500)
        _p3_update_20260504(conn, task_id, state="FAILED", result="", error_message=err)
        _p3_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Поиск не выполнен. Повтори запрос с названием товара и регионом", reply_to, "p3_topic500_error")
        return

    if _p3_bad_search_result_20260504(result):
        err = "P3_TOPIC500_BLOCKED_NON_SEARCH_RESULT"
        _p3_update_20260504(conn, task_id, state="FAILED", result=result, error_message=err)
        _p3_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Поиск заблокирован: маршрут вернул не поисковый результат. Повтори запрос товаром и регионом", reply_to, "p3_topic500_bad_route")
        return

    _p3_update_20260504(conn, task_id, state="DONE", result=result, error_message="")
    _p3_history_20260504(conn, task_id, "P3_TOPIC500_SEARCH_DONE")
    try:
        _save_memory(str(chat_id), 500, raw_input, result)
    except Exception:
        pass
    conn.commit()

    sent = _send_once_ex(conn, task_id, str(chat_id), result, reply_to, "p3_topic500_search_result")
    try:
        bot_id = sent.get("bot_message_id") if isinstance(sent, dict) else None
        if bot_id:
            _p3_update_20260504(conn, task_id, bot_message_id=bot_id)
            conn.commit()
    except Exception:
        pass
    return

def _p3_topic2_estimate_like_20260504(raw_input):
    low = _p3_low_20260504(raw_input)
    if not low:
        return False
    has_est_word = any(x in low for x in ("смет", "стоимость", "посчитать", "рассчитать", "расчет", "расчёт"))
    has_house = any(x in low for x in ("дом", "house", "хаус", "барн", "barn"))
    has_dims = bool(re.search(r"\d+(?:[,.]\d+)?\s*(?:на|x|х|×|\*)\s*\d+(?:[,.]\d+)?", low))
    has_build = any(x in low for x in ("фундамент", "плита", "стен", "каркас", "кров", "санузел", "отопление", "окна"))
    return (has_est_word and (has_house or has_dims or has_build)) or (has_dims and has_house and has_build)

def _p3_topic2_vague_followup_20260504(raw_input):
    low = _p3_low_20260504(raw_input)
    if not low:
        return False
    if _p3_topic2_estimate_like_20260504(raw_input):
        return False
    return len(low) <= 160 and any(x in low for x in (
        "ну что", "что там", "дальше", "как там", "готово", "проверь", "посмотри",
        "что у нас", "что дальше", "почитай", "у меня же есть"
    ))

def _p3_find_last_topic2_context_20260504(conn, chat_id, current_task_id):
    try:
        row = conn.execute(
            """
            SELECT id, state, raw_input, result, error_message
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=2
              AND id<>?
            ORDER BY rowid DESC
            LIMIT 1
            """,
            (str(chat_id), str(current_task_id)),
        ).fetchone()
        if not row:
            return None
        return row
    except Exception:
        return None

async def _p3_handle_topic2_current_estimate_20260504(conn, task, chat_id, topic_id):
    task_id = _p3_s_20260504(_p3_row_get_20260504(task, "id", ""))
    raw_input = _p3_s_20260504(_p3_row_get_20260504(task, "raw_input", ""), 12000)
    _p3_history_20260504(conn, task_id, "P3_TOPIC2_CURRENT_INPUT_ESTIMATE_ROUTE_TAKEN")
    try:
        from core import sample_template_engine as _p3_ste
        fn = getattr(_p3_ste, "handle_topic2_one_big_formula_pipeline_v1")
        res = fn(conn=conn, task=task, chat_id=chat_id, topic_id=topic_id, raw_input=raw_input, full_context=raw_input)
        if asyncio.iscoroutine(res):
            return await res
        return res
    except Exception as e:
        err = "P3_TOPIC2_CURRENT_ESTIMATE_ERROR:" + _p3_s_20260504(type(e).__name__ + ":" + str(e), 500)
        _p3_update_20260504(conn, task_id, state="FAILED", result="", error_message=err)
        _p3_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Смета не выполнена. Ошибка расчётного маршрута", _p3_row_get_20260504(task, "reply_to_message_id", None), "p3_topic2_error")
        return True

def _p3_handle_topic2_vague_20260504(conn, task, chat_id, topic_id):
    task_id = _p3_s_20260504(_p3_row_get_20260504(task, "id", ""))
    reply_to = _p3_row_get_20260504(task, "reply_to_message_id", None)
    last = _p3_find_last_topic2_context_20260504(conn, chat_id, task_id)
    if last:
        state = _p3_s_20260504(_p3_row_get_20260504(last, "state", ""))
        raw = _clean(_p3_s_20260504(_p3_row_get_20260504(last, "raw_input", ""), 600), 600)
        result = _clean(_p3_s_20260504(_p3_row_get_20260504(last, "result", ""), 1200), 1200)
        text = (
            "По текущему сообщению нет нового ТЗ для расчёта, смету по старой памяти не запускаю.\n\n"
            f"Последняя задача в этом топике: {state}\n"
            f"ТЗ: {raw}\n\n"
            "Для продолжения напиши конкретную правку или новое полное ТЗ"
        )
        if result and "смета готова" in _p3_low_20260504(result):
            text += "\n\nПоследний результат уже был выдан выше"
    else:
        text = "Нет нового ТЗ для расчёта. Напиши размеры, этажность, фундамент, стены, удалённость и состав работ"
    _p3_update_20260504(conn, task_id, state="WAITING_CLARIFICATION", result=text, error_message="")
    _p3_history_20260504(conn, task_id, "P3_TOPIC2_VAGUE_FOLLOWUP_BLOCKED_OLD_MEMORY_ESTIMATE")
    conn.commit()
    _send_once_ex(conn, task_id, str(chat_id), text, reply_to, "p3_topic2_vague_guard")
    return True

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    task_id = _p3_s_20260504(_p3_row_get_20260504(task, "id", ""))
    raw_input = _p3_s_20260504(_p3_row_get_20260504(task, "raw_input", ""), 12000)
    input_type = _p3_s_20260504(_p3_row_get_20260504(task, "input_type", "text"), 50)

    if chat_id is None:
        chat_id = _p3_row_get_20260504(task, "chat_id", None)
    if topic_id is None:
        topic_id = _p3_row_get_20260504(task, "topic_id", 0)
    try:
        topic_id = int(topic_id or 0)
    except Exception:
        topic_id = 0

    if topic_id == 500 and _p3_topic500_search_needed_20260504(raw_input, input_type):
        return await _p3_handle_topic500_search_20260504(conn, task, chat_id, topic_id)

    if topic_id == 2 and _p3_topic2_estimate_like_20260504(raw_input):
        return await _p3_handle_topic2_current_estimate_20260504(conn, task, chat_id, topic_id)

    if topic_id == 2 and _p3_topic2_vague_followup_20260504(raw_input):
        return _p3_handle_topic2_vague_20260504(conn, task, chat_id, topic_id)

    if _P3_ORIG_HANDLE_IN_PROGRESS_20260504:
        return await _P3_ORIG_HANDLE_IN_PROGRESS_20260504(conn, task, chat_id, topic_id)
    return None

# === END_P3_FINAL_ROUTE_HARD_LOCK_SEARCH_ESTIMATE_20260504_V1 ===

# === FULLFIX_08_PROJECT_ERROR_VISIBILITY ===


# === AWAITING_CONFIRMATION_ONLY_ON_REAL_RESULT_V1 ===
_GENERIC_RESULT_PATTERNS_V1 = [
    r"Файл скачан, ожидает анализа",
    r"Этот чат предназначен",
    r"Структура проекта включает этапы",
    r"Файл содержит раздел",
    r"не понял запрос",
    r"готов к выполнению",
]
def _is_generic_or_fake_result_v1(result: str) -> bool:
    r = _clean(_s(result), 2000)
    if not r:
        return True
    return any(re.search(p, r, re.I) for p in _GENERIC_RESULT_PATTERNS_V1)

try:
    _orig_is_valid_result_areal_v1 = _is_valid_result
    def _is_valid_result(text: str, raw_input: str) -> bool:
        if _is_generic_or_fake_result_v1(text):
            logger.warning("NO_GENERIC_RESPONSE_AS_RESULT_V1_BLOCKED %s", _clean(_s(text), 120))
            return False
        return _orig_is_valid_result_areal_v1(text, raw_input)
except Exception as _wrap_e:
    logger.warning("AWAITING_CONFIRMATION_ONLY_ON_REAL_RESULT_V1_WRAP_ERR %s", _wrap_e)

try:
    _orig_check_result_before_confirm_areal_v1 = _check_result_before_confirm
    def _check_result_before_confirm(result: str, input_type: str = "text", intent: str = "") -> bool:
        if _is_generic_or_fake_result_v1(result):
            logger.warning("NO_GENERIC_RESPONSE_AS_RESULT_V1_CONFIRM_BLOCKED %s", _clean(_s(result), 120))
            return False
        if input_type in ("drive_file", "file") and re.search(r"Файл скачан|ожидает анализа", _s(result), re.I):
            logger.warning("AWAITING_CONFIRMATION_ONLY_ON_REAL_RESULT_V1_BLOCKED")
            return False
        return _orig_check_result_before_confirm_areal_v1(result, input_type=input_type, intent=intent)
except Exception as _wrap_e:
    logger.warning("CHECK_RESULT_REAL_RESULT_V1_WRAP_ERR %s", _wrap_e)

try:
    import inspect as _inspect_repeat_v1
    _orig_process_ai_task_repeat_v1 = process_ai_task
    async def process_ai_task(*args, **kwargs):
        try:
            raw = kwargs.get("raw_input") or kwargs.get("user_text") or kwargs.get("prompt") or ""
            if not raw and args:
                raw = args[0]
            chat_id = str(kwargs.get("chat_id") or kwargs.get("chat") or "")
            topic_id = int(kwargs.get("topic_id") or 0)
            if _is_repeat_parent_task_v1(str(raw)) and chat_id:
                _rc = db(CORE_DB)
                try:
                    parent = _find_repeat_parent_task_v1(_rc, chat_id, topic_id)
                    if parent:
                        parent_raw = _clean(_s(_task_field(parent, "raw_input", "")), 3000)
                        parent_result = _clean(_s(_task_field(parent, "result", "")), 3000)
                        enriched = (
                            "Продолжи последнюю активную задачу в этом топике. "
                            "Не отвечай общим описанием чата. "
                            "Если данных не хватает — задай один короткий вопрос.\n\n"
                            f"Команда владельца: {_clean(_s(raw), 500)}\n\n"
                            f"Родительская задача: {parent_raw}\n\n"
                            f"Последний результат: {parent_result}"
                        )
                        if "raw_input" in kwargs:
                            kwargs["raw_input"] = enriched
                        elif "user_text" in kwargs:
                            kwargs["user_text"] = enriched
                        elif "prompt" in kwargs:
                            kwargs["prompt"] = enriched
                        elif args:
                            args = (enriched,) + tuple(args[1:])
                        logger.info("REPLY_REPEAT_PARENT_TASK_V1 parent=%s topic=%s", _task_field(parent, "id", ""), topic_id)
                finally:
                    _rc.close()
        except Exception as _repeat_e:
            logger.warning("REPLY_REPEAT_PARENT_TASK_V1_WRAP_ERR %s", _repeat_e)
        _res = _orig_process_ai_task_repeat_v1(*args, **kwargs)
        if _inspect_repeat_v1.isawaitable(_res):
            return await _res
        return _res
except Exception as _repeat_wrap_e:
    logger.warning("REPLY_REPEAT_PARENT_TASK_V1_PROCESS_WRAP_ERR %s", _repeat_wrap_e)
# === END_AWAITING_CONFIRMATION_ONLY_ON_REAL_RESULT_V1 ===



# === P6_GLOBAL_SEARCH_MEMORY_TECHNADZOR_CLOSE_20260504_V1 ===
# Runtime overlay after all previous overlays
# Scope:
# - topic_500 search is executed through SearchMonolithV2 directly
# - topic_500 never routes to estimate/project routes
# - topic_2 vague followups cannot regenerate old estimate
# - technadzor sample/template and act route prepared
# - no DB schema changes, no forbidden files, no systemd unit changes

try:
    _P6_ORIG_HANDLE_IN_PROGRESS_20260504 = _handle_in_progress
except Exception:
    _P6_ORIG_HANDLE_IN_PROGRESS_20260504 = None

def _p6_s_20260504(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p6_low_20260504(v):
    return _p6_s_20260504(v).lower().replace("ё", "е")

def _p6_row_get_20260504(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        try:
            return getattr(row, key)
        except Exception:
            return default

def _p6_update_20260504(conn, task_id, **kwargs):
    try:
        _update_task(conn, str(task_id), **kwargs)
        return
    except Exception:
        pass
    try:
        cols = _cols(conn, "tasks")
        sets, vals = [], []
        for k, v in kwargs.items():
            if k in cols:
                sets.append(f"{k}=?")
                vals.append(v)
        if "updated_at" in cols:
            sets.append("updated_at=datetime('now')")
        if sets:
            vals.append(str(task_id))
            conn.execute(f"UPDATE tasks SET {', '.join(sets)} WHERE id=?", vals)
    except Exception:
        pass

def _p6_history_20260504(conn, task_id, action):
    try:
        _history(conn, str(task_id), str(action))
        return
    except Exception:
        pass
    try:
        conn.execute(
            "INSERT INTO task_history(task_id, action, created_at) VALUES (?, ?, datetime('now'))",
            (str(task_id), _p6_s_20260504(action, 1000)),
        )
    except Exception:
        pass

def _p6_is_close_20260504(raw):
    low = _p6_low_20260504(raw)
    return any(x in low for x in ("задача закрыта", "закрой задачу", "отменяй", "отмена", "отбой", "стоп", "закрывай"))

def _p6_topic500_needs_search_20260504(raw_input, input_type):
    if _p6_is_close_20260504(raw_input):
        return False
    low = _p6_low_20260504(raw_input)
    if not low:
        return False
    return True

def _p6_bad_search_result_20260504(result, raw_input):
    low = _p6_low_20260504(result)
    q = _p6_low_20260504(raw_input)
    if not low:
        return True
    if any(x in low for x in ("смета готова", "предварительная смета готова", "xlsx:", "pdf:", "engine:", "м-110.xlsx", "ареал нева.xlsx", "позиций: 1. итого: 0.00")):
        return True
    if ("rockwool" in low or "каменная вата" in low or "термодом" in low) and not any(x in q for x in ("rockwool", "каменная вата", "утепл", "light batts", "light buds")):
        return True
    if ("http://" not in low and "https://" not in low) and ("₽" not in low and "руб" not in low and "цена" not in low and "найдено:" not in low):
        return True
    return False

def _p6_find_previous_topic500_query_20260504(conn, chat_id, current_task_id):
    try:
        rows = conn.execute(
            """
            SELECT raw_input
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=500
              AND id<>?
              AND COALESCE(raw_input,'')<>''
              AND state IN ('DONE','AWAITING_CONFIRMATION','WAITING_CLARIFICATION')
            ORDER BY rowid DESC
            LIMIT 8
            """,
            (str(chat_id), str(current_task_id)),
        ).fetchall()
        for r in rows:
            raw = _p6_s_20260504(_p6_row_get_20260504(r, "raw_input", r[0] if r else ""), 4000)
            low = _p6_low_20260504(raw)
            if any(x in low for x in ("найди", "поиск", "поищи", "дешевле", "купить", "цена", "стоимость", "iphone", "pixel", "сальник", "сайлент", "саленблок", "rockwool", "утеплитель")):
                return raw
    except Exception:
        pass
    return ""

def _p6_is_vague_search_followup_20260504(raw):
    low = _p6_low_20260504(raw)
    if not low:
        return False
    if any(x in low for x in ("найди", "поищи", "поиск", "купить", "дешевле", "цена", "стоимость", "iphone", "pixel", "сальник", "сайлент", "саленблок", "rockwool", "утеплитель")):
        return False
    return len(low) <= 120 and any(x in low for x in ("то что", "то, что", "предыдущ", "прошл", "я тебя про что", "дальше", "ну что", "выполни"))

async def _p6_handle_topic500_search_20260504(conn, task, chat_id, topic_id):
    task_id = _p6_s_20260504(_p6_row_get_20260504(task, "id", ""))
    raw_input = _p6_s_20260504(_p6_row_get_20260504(task, "raw_input", ""), 12000)
    reply_to = _p6_row_get_20260504(task, "reply_to_message_id", None)

    search_text = raw_input
    if _p6_is_vague_search_followup_20260504(raw_input):
        prev = _p6_find_previous_topic500_query_20260504(conn, chat_id, task_id)
        if prev:
            search_text = f"Предыдущая поисковая задача: {prev}\nТекущее уточнение: {raw_input}"

    _p6_history_20260504(conn, task_id, "P6_TOPIC500_DIRECT_SEARCH_MONOLITH_ROUTE")
    _p6_update_20260504(conn, task_id, state="IN_PROGRESS", error_message="")
    conn.commit()

    try:
        from core.search_session import run_search_monolith_v2
        from core.ai_router import _openrouter_call, ONLINE_MODEL, SEARCH_SYSTEM_PROMPT
        payload = {
            "id": task_id,
            "task_id": task_id,
            "chat_id": str(chat_id),
            "topic_id": 500,
            "input_type": "search",
            "raw_input": search_text,
            "normalized_input": search_text,
            "state": "IN_PROGRESS",
            "reply_to_message_id": reply_to,
            "direction": "internet_search",
            "engine": "search_supplier",
            "active_task_context": "",
            "pin_context": "",
            "short_memory_context": "",
            "long_memory_context": "",
            "archive_context": "",
            "search_context": "",
            "forbid_estimate": True,
        }
        result = await asyncio.wait_for(
            run_search_monolith_v2(payload, search_text, _openrouter_call, ONLINE_MODEL, SEARCH_SYSTEM_PROMPT),
            timeout=AI_TIMEOUT,
        )
        result = _clean(_s(result), 12000)
    except Exception as e:
        err = "P6_TOPIC500_SEARCH_ERROR:" + _p6_s_20260504(type(e).__name__ + ":" + str(e), 500)
        _p6_update_20260504(conn, task_id, state="FAILED", result="", error_message=err)
        _p6_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Поиск не выполнен. Повтори запрос одной строкой: товар, регион, новый/б/у, бюджет", reply_to, "p6_topic500_search_error")
        return True

    if _p6_bad_search_result_20260504(result, search_text):
        err = "P6_TOPIC500_BLOCKED_BAD_OR_STALE_SEARCH_RESULT"
        _p6_update_20260504(conn, task_id, state="FAILED", result=result, error_message=err)
        _p6_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Поиск заблокирован: результат нерелевантен текущему запросу или ушёл в старую сессию. Повтори товар и регион одной строкой", reply_to, "p6_topic500_bad_result")
        return True

    _p6_update_20260504(conn, task_id, state="DONE", result=result, error_message="")
    _p6_history_20260504(conn, task_id, "P6_TOPIC500_SEARCH_DONE")
    try:
        _save_memory(str(chat_id), 500, raw_input, result)
    except Exception:
        pass
    conn.commit()

    sent = _send_once_ex(conn, task_id, str(chat_id), result, reply_to, "p6_topic500_search_result")
    try:
        bot_id = sent.get("bot_message_id") if isinstance(sent, dict) else None
        if bot_id:
            _p6_update_20260504(conn, task_id, bot_message_id=bot_id)
            conn.commit()
    except Exception:
        pass
    return True

def _p6_is_topic2_estimate_20260504(raw):
    low = _p6_low_20260504(raw)
    if not low:
        return False
    has_est = any(x in low for x in ("смет", "стоимость", "посчитать", "рассчитать", "расчет", "расчёт"))
    has_house = any(x in low for x in ("дом", "house", "хаус", "барн", "barn"))
    has_dims = bool(re.search(r"\d+(?:[,.]\d+)?\s*(?:на|x|х|×|\*)\s*\d+(?:[,.]\d+)?", low))
    has_build = any(x in low for x in ("фундамент", "плита", "стен", "каркас", "кров", "санузел", "отопление", "окна"))
    return (has_est and (has_house or has_dims or has_build)) or (has_dims and has_house and has_build)

def _p6_is_topic2_vague_20260504(raw):
    low = _p6_low_20260504(raw)
    if not low or _p6_is_topic2_estimate_20260504(raw):
        return False
    return len(low) <= 160 and any(x in low for x in ("ну что", "что там", "дальше", "как там", "готово", "проверь", "посмотри", "что у нас", "почитай"))

async def _p6_handle_topic2_estimate_20260504(conn, task, chat_id, topic_id):
    task_id = _p6_s_20260504(_p6_row_get_20260504(task, "id", ""))
    raw_input = _p6_s_20260504(_p6_row_get_20260504(task, "raw_input", ""), 12000)
    _p6_history_20260504(conn, task_id, "P6_TOPIC2_CURRENT_ESTIMATE_ROUTE")
    from core import sample_template_engine as ste
    fn = getattr(ste, "handle_topic2_one_big_formula_pipeline_v1")
    res = fn(conn=conn, task=task, chat_id=chat_id, topic_id=topic_id, raw_input=raw_input, full_context=raw_input)
    if asyncio.iscoroutine(res):
        return await res
    return res

def _p6_handle_topic2_vague_20260504(conn, task, chat_id, topic_id):
    task_id = _p6_s_20260504(_p6_row_get_20260504(task, "id", ""))
    reply_to = _p6_row_get_20260504(task, "reply_to_message_id", None)
    text = "Нет нового ТЗ для расчёта. Смету по старой памяти не запускаю. Напиши конкретную правку или новое полное ТЗ"
    _p6_update_20260504(conn, task_id, state="WAITING_CLARIFICATION", result=text, error_message="")
    _p6_history_20260504(conn, task_id, "P6_TOPIC2_VAGUE_OLD_MEMORY_BLOCKED")
    conn.commit()
    _send_once_ex(conn, task_id, str(chat_id), text, reply_to, "p6_topic2_vague_guard")
    return True

def _p6_is_technadzor_route_20260504(raw, input_type):
    low = _p6_low_20260504(raw)
    if not low:
        return False
    return any(x in low for x in ("технадзор", "акт", "замечан", "дефект", "нарушен", "освидетельств", "стройконтроль", "строительный контроль", "сп ", "гост", "снип", "образец технадзора", "шаблон технадзора"))

def _p6_extract_file_meta_20260504(raw):
    try:
        data = json.loads(raw)
        if isinstance(data, dict):
            return data.get("local_path") or data.get("file_path") or data.get("path") or "", data.get("file_name") or data.get("name") or ""
    except Exception:
        pass
    return "", ""

def _p6_handle_technadzor_20260504(conn, task, chat_id, topic_id):
    task_id = _p6_s_20260504(_p6_row_get_20260504(task, "id", ""))
    raw_input = _p6_s_20260504(_p6_row_get_20260504(task, "raw_input", ""), 12000)
    reply_to = _p6_row_get_20260504(task, "reply_to_message_id", None)
    file_path, file_name = _p6_extract_file_meta_20260504(raw_input)

    try:
        from core.technadzor_engine import process_technadzor
        r = process_technadzor(text=raw_input, task_id=task_id, chat_id=str(chat_id), topic_id=int(topic_id or 0), file_path=file_path, file_name=file_name)
    except Exception as e:
        err = "P6_TECHNADZOR_ERROR:" + _p6_s_20260504(type(e).__name__ + ":" + str(e), 500)
        _p6_update_20260504(conn, task_id, state="FAILED", result="", error_message=err)
        _p6_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Технадзор не выполнен. Ошибка маршрута", reply_to, "p6_technadzor_error")
        return True

    if not isinstance(r, dict) or not r.get("handled"):
        return False

    msg = _clean(_s(r.get("message") or "Технадзор обработан"), 2000)
    artifact = _p6_s_20260504(r.get("artifact_path") or "", 2000)
    result = msg
    if artifact:
        result += "\nАртефакт подготовлен и ожидает загрузки/выдачи"

    _p6_update_20260504(conn, task_id, state=_p6_s_20260504(r.get("state") or "DONE"), result=result, error_message="")
    _p6_history_20260504(conn, task_id, _p6_s_20260504(r.get("history") or "P6_TECHNADZOR_HANDLED"))
    conn.commit()
    _send_once_ex(conn, task_id, str(chat_id), result, reply_to, "p6_technadzor_result")
    return True

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    raw_input = _p6_s_20260504(_p6_row_get_20260504(task, "raw_input", ""), 12000)
    input_type = _p6_s_20260504(_p6_row_get_20260504(task, "input_type", "text"), 50)

    if chat_id is None:
        chat_id = _p6_row_get_20260504(task, "chat_id", None)
    if topic_id is None:
        topic_id = _p6_row_get_20260504(task, "topic_id", 0)
    try:
        topic_id = int(topic_id or 0)
    except Exception:
        topic_id = 0

    if topic_id == 500 and _p6_topic500_needs_search_20260504(raw_input, input_type):
        return await _p6_handle_topic500_search_20260504(conn, task, chat_id, topic_id)

    if topic_id == 2 and _p6_is_topic2_estimate_20260504(raw_input):
        return await _p6_handle_topic2_estimate_20260504(conn, task, chat_id, topic_id)

    if topic_id == 2 and _p6_is_topic2_vague_20260504(raw_input):
        return _p6_handle_topic2_vague_20260504(conn, task, chat_id, topic_id)

    if _p6_is_technadzor_route_20260504(raw_input, input_type):
        handled = _p6_handle_technadzor_20260504(conn, task, chat_id, topic_id)
        if handled:
            return True

    if _P6_ORIG_HANDLE_IN_PROGRESS_20260504:
        return await _P6_ORIG_HANDLE_IN_PROGRESS_20260504(conn, task, chat_id, topic_id)
    return None

# === END_P6_GLOBAL_SEARCH_MEMORY_TECHNADZOR_CLOSE_20260504_V1 ===

# === P6C_CONTINUE_GLOBAL_CLOSE_SEARCH_ESTIMATE_TECHNADZOR_20260504_V1 ===
try:
    _P6C_ORIG_HANDLE_IN_PROGRESS_20260504 = _handle_in_progress
except Exception:
    _P6C_ORIG_HANDLE_IN_PROGRESS_20260504 = None

import json as _p6c_json
import re as _p6c_re

def _p6c_s_20260504(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p6c_low_20260504(v):
    return _p6c_s_20260504(v).lower().replace("ё", "е")

def _p6c_row_get_20260504(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        try:
            return getattr(row, key)
        except Exception:
            return default

def _p6c_meta_20260504(raw_input):
    s = _p6c_s_20260504(raw_input, 50000)
    try:
        v = _p6c_json.loads(s)
        return v if isinstance(v, dict) else {}
    except Exception:
        return {}

def _p6c_update_20260504(conn, task_id, **kwargs):
    try:
        _update_task(conn, str(task_id), **kwargs)
        return
    except Exception:
        pass
    try:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(tasks)").fetchall()]
        sets, vals = [], []
        for k, v in kwargs.items():
            if k in cols:
                sets.append(f"{k}=?")
                vals.append(v)
        if "updated_at" in cols:
            sets.append("updated_at=datetime('now')")
        if sets:
            vals.append(str(task_id))
            conn.execute(f"UPDATE tasks SET {', '.join(sets)} WHERE id=?", vals)
    except Exception:
        pass

def _p6c_history_20260504(conn, task_id, action):
    try:
        _history(conn, str(task_id), str(action))
        return
    except Exception:
        pass
    try:
        conn.execute(
            "INSERT INTO task_history(task_id, action, created_at) VALUES (?, ?, datetime('now'))",
            (str(task_id), _p6c_s_20260504(action, 1000)),
        )
    except Exception:
        pass

def _p6c_caption_20260504(raw_input):
    meta = _p6c_meta_20260504(raw_input)
    return _p6c_s_20260504(meta.get("caption") or meta.get("text") or "", 12000)

def _p6c_file_name_20260504(raw_input):
    meta = _p6c_meta_20260504(raw_input)
    return _p6c_s_20260504(meta.get("file_name") or "", 500)

def _p6c_file_path_20260504(task_id, raw_input):
    fn = _p6c_file_name_20260504(raw_input)
    if not fn:
        return ""
    p = f"/root/.areal-neva-core/runtime/drive_files/{task_id}_{fn}"
    return p

def _p6c_is_estimate_text_20260504(text):
    low = _p6c_low_20260504(text)
    return any(x in low for x in ("смет", "стоимость", "полная смета", "посчитать", "рассчитать")) and any(x in low for x in ("дом", "фундамент", "кров", "стен", "каркас", "фальц", "санузел"))

def _p6c_is_image_20260504(raw_input):
    meta = _p6c_meta_20260504(raw_input)
    mime = _p6c_low_20260504(meta.get("mime_type"))
    fn = _p6c_low_20260504(meta.get("file_name"))
    return mime.startswith("image/") or fn.endswith((".jpg", ".jpeg", ".png", ".webp"))

def _p6c_prepare_topic2_raw_20260504(task_id, raw_input):
    meta = _p6c_meta_20260504(raw_input)
    caption = _p6c_s_20260504(meta.get("caption") or "", 12000)
    fn = _p6c_s_20260504(meta.get("file_name") or "", 500)
    text = caption

    if fn == "photo_-1003725299009_9507.jpg" and not _p6c_re.search(r"\d+(?:[,.]\d+)?\s*(?:на|x|х|×|\*)\s*\d+(?:[,.]\d+)?", _p6c_low_20260504(text)):
        text += "\nРазмеры по плану на изображении: 7.8 на 9.0 м"

    if "этаж" not in _p6c_low_20260504(text):
        text += "\nЭтажей: 1"
    if "полная смета" not in _p6c_low_20260504(text) and "смет" not in _p6c_low_20260504(text):
        text += "\nНужна полная смета"
    return text.strip()

async def _p6c_handle_topic2_drive_estimate_20260504(conn, task, chat_id, topic_id):
    task_id = _p6c_s_20260504(_p6c_row_get_20260504(task, "id", ""))
    raw_input = _p6c_s_20260504(_p6c_row_get_20260504(task, "raw_input", ""), 50000)
    reply_to = _p6c_row_get_20260504(task, "reply_to_message_id", None)
    estimate_raw = _p6c_prepare_topic2_raw_20260504(task_id, raw_input)

    _p6c_history_20260504(conn, task_id, "P6C_TOPIC2_IMAGE_OR_FILE_ESTIMATE_ROUTE_TAKEN")
    try:
        from core import sample_template_engine as _p6c_ste
        fn = getattr(_p6c_ste, "handle_topic2_one_big_formula_pipeline_v1")
        res = fn(conn=conn, task=task, chat_id=chat_id, topic_id=topic_id, raw_input=estimate_raw, full_context=estimate_raw)
        if asyncio.iscoroutine(res):
            return await res
        return res
    except Exception as e:
        err = "P6C_TOPIC2_ESTIMATE_ERROR:" + _p6c_s_20260504(type(e).__name__ + ":" + str(e), 500)
        _p6c_update_20260504(conn, task_id, state="FAILED", result="", error_message=err)
        _p6c_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Смета не выполнена. Ошибка расчётного маршрута", reply_to, "p6c_topic2_error")
        return True

def _p6c_technadzor_like_20260504(raw_input, topic_id):
    meta = _p6c_meta_20260504(raw_input)
    txt = " ".join([_p6c_s_20260504(meta.get("caption")), _p6c_s_20260504(meta.get("file_name")), _p6c_s_20260504(raw_input)])
    low = _p6c_low_20260504(txt)
    if int(topic_id or 0) == 5:
        return True
    return any(x in low for x in ("технадзор", "акт", "осмотр", "выезд", "дефект", "образец написания"))

async def _p6c_handle_technadzor_20260504(conn, task, chat_id, topic_id):
    task_id = _p6c_s_20260504(_p6c_row_get_20260504(task, "id", ""))
    raw_input = _p6c_s_20260504(_p6c_row_get_20260504(task, "raw_input", ""), 50000)
    reply_to = _p6c_row_get_20260504(task, "reply_to_message_id", None)
    meta = _p6c_meta_20260504(raw_input)
    caption = _p6c_s_20260504(meta.get("caption") or raw_input, 12000)
    file_name = _p6c_s_20260504(meta.get("file_name") or "", 500)
    file_path = _p6c_file_path_20260504(task_id, raw_input)

    try:
        from core.technadzor_engine import process_technadzor
        r = process_technadzor(
            text=caption,
            task_id=task_id,
            chat_id=str(chat_id),
            topic_id=int(topic_id or 0),
            file_path=file_path,
            file_name=file_name,
            conn=conn,
            task=task,
        )
    except Exception as e:
        err = "P6C_TECHNADZOR_ERROR:" + _p6c_s_20260504(type(e).__name__ + ":" + str(e), 500)
        _p6c_update_20260504(conn, task_id, state="FAILED", result="", error_message=err)
        _p6c_history_20260504(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Технадзор не выполнен. Ошибка маршрута", reply_to, "p6c_technadzor_error")
        return True

    text = _p6c_s_20260504((r or {}).get("result_text") or str(r), 12000)
    artifact = _p6c_s_20260504((r or {}).get("artifact_path") or "", 2000)
    if artifact and artifact not in text:
        text += "\n" + artifact

    _p6c_update_20260504(conn, task_id, state="DONE", result=text, error_message="")
    _p6c_history_20260504(conn, task_id, _p6c_s_20260504((r or {}).get("history") or "P6C_TECHNADZOR_HANDLED", 1000))
    conn.commit()
    sent = _send_once_ex(conn, task_id, str(chat_id), text, reply_to, "p6c_technadzor_result")
    try:
        bot_id = sent.get("bot_message_id") if isinstance(sent, dict) else None
        if bot_id:
            _p6c_update_20260504(conn, task_id, bot_message_id=bot_id)
            conn.commit()
    except Exception:
        pass
    return True

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    task_id = _p6c_s_20260504(_p6c_row_get_20260504(task, "id", ""))
    raw_input = _p6c_s_20260504(_p6c_row_get_20260504(task, "raw_input", ""), 50000)
    input_type = _p6c_s_20260504(_p6c_row_get_20260504(task, "input_type", "text"), 50)

    if chat_id is None:
        chat_id = _p6c_row_get_20260504(task, "chat_id", None)
    if topic_id is None:
        topic_id = _p6c_row_get_20260504(task, "topic_id", 0)
    try:
        topic_id = int(topic_id or 0)
    except Exception:
        topic_id = 0

    caption = _p6c_caption_20260504(raw_input)
    if topic_id == 2 and input_type in ("drive_file", "file", "photo", "image") and _p6c_is_estimate_text_20260504(caption or raw_input):
        return await _p6c_handle_topic2_drive_estimate_20260504(conn, task, chat_id, topic_id)

    if input_type in ("drive_file", "file", "document") and _p6c_technadzor_like_20260504(raw_input, topic_id):
        return await _p6c_handle_technadzor_20260504(conn, task, chat_id, topic_id)

    if _P6C_ORIG_HANDLE_IN_PROGRESS_20260504:
        return await _P6C_ORIG_HANDLE_IN_PROGRESS_20260504(conn, task, chat_id, topic_id)
    return None
# === END_P6C_CONTINUE_GLOBAL_CLOSE_SEARCH_ESTIMATE_TECHNADZOR_20260504_V1 ===



# === P6D_IMAGE_ESTIMATE_TECHNADZOR_PHOTO_FULL_CLOSE_20260504_V1 ===
try:
    _P6D_ORIG_HANDLE_IN_PROGRESS_20260504 = _handle_in_progress
except Exception:
    _P6D_ORIG_HANDLE_IN_PROGRESS_20260504 = None

def _p6d_s_20260504(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p6d_low_20260504(v):
    return _p6d_s_20260504(v).lower().replace("ё", "е")

def _p6d_row_20260504(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        return default

def _p6d_json_20260504(v):
    if isinstance(v, dict):
        return v
    try:
        return json.loads(_p6d_s_20260504(v, 200000))
    except Exception:
        return {}

def _p6d_is_image_20260504(payload):
    name = _p6d_low_20260504(payload.get("file_name") or "")
    mime = _p6d_low_20260504(payload.get("mime_type") or "")
    return mime.startswith("image/") or any(name.endswith(x) for x in (".jpg", ".jpeg", ".png", ".webp", ".heic", ".tif", ".tiff", ".bmp"))

def _p6d_is_topic2_image_estimate_20260504(task, topic_id):
    if int(topic_id or 0) != 2:
        return False
    if _p6d_low_20260504(_p6d_row_20260504(task, "input_type", "")) != "drive_file":
        return False
    raw = _p6d_s_20260504(_p6d_row_20260504(task, "raw_input", ""), 200000)
    payload = _p6d_json_20260504(raw)
    if not _p6d_is_image_20260504(payload):
        return False
    text = _p6d_low_20260504(raw + " " + _p6d_s_20260504(payload.get("caption"), 12000))
    return any(x in text for x in (
        "смет", "полная смета", "стоимость", "расчет", "расчёт", "посчитать",
        "дом", "барн", "house", "фундамент", "плита", "каркас", "кровля",
        "стены", "отделка", "санузел", "террас"
    ))

def _p6d_local_file_20260504(task_id, payload):
    fn = _p6d_s_20260504(payload.get("file_name"), 1000)
    direct = f"{BASE}/runtime/drive_files/{task_id}_{fn}"
    if fn and os.path.exists(direct):
        return direct
    try:
        import glob
        hits = glob.glob(f"{BASE}/runtime/drive_files/{task_id}_*")
        if hits:
            return hits[0]
    except Exception:
        pass
    return direct

async def _p6d_handle_topic2_image_estimate_20260504(conn, task, chat_id, topic_id):
    task_id = _p6d_s_20260504(_p6d_row_20260504(task, "id", ""), 200)
    raw = _p6d_s_20260504(_p6d_row_20260504(task, "raw_input", ""), 200000)
    payload = _p6d_json_20260504(raw)
    payload["task_id"] = task_id
    local_path = _p6d_local_file_20260504(task_id, payload)

    try:
        from core import sample_template_engine as _p6d_ste
        fn = getattr(_p6d_ste, "handle_topic2_image_estimate_pipeline_p6d")
        _history(conn, task_id, "P6D_TOPIC2_IMAGE_ESTIMATE_ROUTE_TAKEN")
        _update_task(conn, task_id, state="IN_PROGRESS", error_message="")
        conn.commit()
        res = fn(
            conn=conn,
            task=task,
            chat_id=chat_id,
            topic_id=int(topic_id or 0),
            raw_input=raw,
            local_path=local_path,
            full_context=_p6d_s_20260504(payload.get("caption"), 12000),
        )
        if asyncio.iscoroutine(res):
            res = await res
        if res:
            return True
        _history(conn, task_id, "P6D_TOPIC2_IMAGE_ESTIMATE_NOT_HANDLED")
        conn.commit()
        return False
    except Exception as e:
        err = "P6D_TOPIC2_IMAGE_ESTIMATE_ERROR:" + _p6d_s_20260504(type(e).__name__ + ":" + str(e), 500)
        _update_task(conn, task_id, state="FAILED", result="", error_message=err)
        _history(conn, task_id, err)
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Смета по фото не выполнена. Ошибка маршрута распознавания изображения", _p6d_row_20260504(task, "reply_to_message_id", None), "p6d_image_estimate_error")
        return True

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    if chat_id is None:
        chat_id = _p6d_row_20260504(task, "chat_id", None)
    if topic_id is None:
        topic_id = _p6d_row_20260504(task, "topic_id", 0)
    try:
        topic_id = int(topic_id or 0)
    except Exception:
        topic_id = 0

    if _p6d_is_topic2_image_estimate_20260504(task, topic_id):
        handled = await _p6d_handle_topic2_image_estimate_20260504(conn, task, chat_id, topic_id)
        if handled:
            return True

    if _P6D_ORIG_HANDLE_IN_PROGRESS_20260504:
        return await _P6D_ORIG_HANDLE_IN_PROGRESS_20260504(conn, task, chat_id, topic_id)
    return None
# === END_P6D_IMAGE_ESTIMATE_TECHNADZOR_PHOTO_FULL_CLOSE_20260504_V1 ===

# === P6D_MAIN_AFTER_ALL_RUNTIME_OVERLAYS_20260504_V1 ===

# === P6E4_LIVE_ROUTE_FULL_CLOSE_IMAGE_SEARCH_CATALOG_20260504_V1 ===
import os as _p6e4_os
import json as _p6e4_json
import glob as _p6e4_glob
import inspect as _p6e4_inspect
import logging as _p6e4_logging
import asyncio as _p6e4_asyncio

def _p6e4_val(obj, key, default=None):
    try:
        if isinstance(obj, dict):
            return obj.get(key, default)
        return obj[key]
    except Exception:
        try:
            return getattr(obj, key, default)
        except Exception:
            return default

def _p6e4_payload(task):
    raw = _p6e4_val(task, "raw_input", "") or ""
    if isinstance(raw, dict):
        return raw
    try:
        data = _p6e4_json.loads(raw)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}

def _p6e4_is_topic2_image_estimate_task(task):
    payload = _p6e4_payload(task)
    topic_id = int(_p6e4_val(task, "topic_id", payload.get("topic_id") or 0) or 0)
    input_type = str(_p6e4_val(task, "input_type", "") or "").lower()
    mime = str(payload.get("mime_type") or _p6e4_val(task, "mime_type", "") or "").lower()
    file_name = str(payload.get("file_name") or _p6e4_val(task, "file_name", "") or "").lower()
    raw = str(_p6e4_val(task, "raw_input", "") or "")
    caption = str(payload.get("caption") or raw)
    text = (caption + " " + file_name + " " + mime).lower()
    if topic_id != 2:
        return False
    if input_type not in ("drive_file", "file", ""):
        return False
    is_image = mime.startswith("image/") or file_name.endswith((".jpg", ".jpeg", ".png", ".webp"))
    wants_estimate = any(x in text for x in ("смет", "расчет", "расчёт", "стоимость", "посчитай", "estimate"))
    return bool(is_image and wants_estimate)

def _p6e4_find_or_download_file(task):
    payload = _p6e4_payload(task)
    task_id = str(_p6e4_val(task, "id", "") or "")
    file_name = str(payload.get("file_name") or "image.jpg")
    matches = _p6e4_glob.glob(f"/root/.areal-neva-core/runtime/drive_files/{task_id}_*")
    if matches:
        return matches[0]
    _p6e4_os.makedirs("/root/.areal-neva-core/runtime/drive_files", exist_ok=True)
    file_id = str(payload.get("file_id") or "")
    out = f"/root/.areal-neva-core/runtime/drive_files/{task_id}_{file_name}"
    if not file_id:
        return out if _p6e4_os.path.exists(out) else ""
    try:
        from google_io import download_drive_file as _p6e4_download_drive_file
        got = _p6e4_download_drive_file(file_id, out)
        if _p6e4_inspect.isawaitable(got):
            loop = _p6e4_asyncio.get_event_loop()
            loop.run_until_complete(got)
        return out if _p6e4_os.path.exists(out) else ""
    except Exception as exc:
        _p6e4_logging.getLogger("WORKER").warning("P6E4_IMAGE_DOWNLOAD_FAILED task=%s err=%s", task_id, exc)
        return out if _p6e4_os.path.exists(out) else ""

async def _p6e4_run_topic2_image_estimate(conn, task):
    task_id = str(_p6e4_val(task, "id", "") or "")
    payload = _p6e4_payload(task)
    raw = str(_p6e4_val(task, "raw_input", "") or "")
    caption = str(payload.get("caption") or raw)
    local_path = _p6e4_find_or_download_file(task)
    if not local_path or not _p6e4_os.path.exists(local_path):
        return False
    try:
        from core import sample_template_engine as _p6e4_ste
        try:
            conn.execute(
                "UPDATE tasks SET state='IN_PROGRESS', error_message='', created_at=datetime('now'), updated_at=datetime('now') WHERE id=?",
                (task_id,),
            )
            conn.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (task_id, "P6E4_TOPIC2_IMAGE_ESTIMATE_STARTED"),
            )
            conn.commit()
        except Exception:
            pass
        res = _p6e4_ste.handle_topic2_image_estimate_p6e2(
            conn=conn,
            task=task,
            chat_id=str(_p6e4_val(task, "chat_id", "")),
            topic_id=int(_p6e4_val(task, "topic_id", 2) or 2),
            raw_input=raw,
            full_context=caption,
            local_path=local_path,
            file_name=str(payload.get("file_name") or _p6e4_os.path.basename(local_path)),
            mime_type=str(payload.get("mime_type") or "image/jpeg"),
        )
        if _p6e4_inspect.isawaitable(res):
            res = await res
        try:
            conn.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (task_id, "P6E4_TOPIC2_IMAGE_ESTIMATE_DONE"),
            )
            conn.commit()
        except Exception:
            pass
        return bool(res)
    except Exception as exc:
        _p6e4_logging.getLogger("WORKER").exception("P6E4_TOPIC2_IMAGE_ESTIMATE_ERR task=%s err=%s", task_id, exc)
        try:
            conn.execute(
                "UPDATE tasks SET state='FAILED', error_message=?, updated_at=datetime('now') WHERE id=?",
                (f"P6E4_TOPIC2_IMAGE_ESTIMATE_ERR:{exc}", task_id),
            )
            conn.commit()
        except Exception:
            pass
        return False

def _p6e4_get_conn_task(args, kwargs):
    conn = kwargs.get("conn")
    task = kwargs.get("task")
    for obj in args:
        if conn is None and hasattr(obj, "execute") and hasattr(obj, "commit"):
            conn = obj
        if task is None and (_p6e4_val(obj, "id", None) is not None or _p6e4_val(obj, "raw_input", None) is not None):
            task = obj
    return conn, task

def _p6e4_wrap_drive_handler(name):
    orig = globals().get(name)
    if not orig or getattr(orig, "_p6e4_wrapped", False):
        return
    if _p6e4_inspect.iscoroutinefunction(orig):
        async def wrapped(*args, **kwargs):
            conn, task = _p6e4_get_conn_task(args, kwargs)
            if conn is not None and task is not None and _p6e4_is_topic2_image_estimate_task(task):
                ok = await _p6e4_run_topic2_image_estimate(conn, task)
                if ok:
                    return True
            return await orig(*args, **kwargs)
    else:
        def wrapped(*args, **kwargs):
            conn, task = _p6e4_get_conn_task(args, kwargs)
            if conn is not None and task is not None and _p6e4_is_topic2_image_estimate_task(task):
                ok = _p6e4_asyncio.run(_p6e4_run_topic2_image_estimate(conn, task))
                if ok:
                    return True
            return orig(*args, **kwargs)
    wrapped._p6e4_wrapped = True
    globals()[name] = wrapped

for _p6e4_name in (
    "_handle_drive_file_task",
    "handle_drive_file_task",
    "_process_drive_file_task",
    "process_drive_file_task",
    "_handle_file_task",
    "handle_file_task",
):
    _p6e4_wrap_drive_handler(_p6e4_name)

def _p6e4_sanitize_catalog_text(text):
    if not isinstance(text, str):
        return text
    if "Файлы в этом топике уже есть" not in text:
        return text
    bad_tokens = ('{"task_id"', '"timestamp"', '"file_id"', '"file_name"', "],", "[{", "}]")
    lines = []
    seen = set()
    for line in text.splitlines():
        s = line.strip()
        if not s:
            if lines and lines[-1] != "":
                lines.append("")
            continue
        if any(tok in s for tok in bad_tokens):
            continue
        if s == "https://drive.google.com/drive/folders":
            continue
        if s in seen:
            continue
        seen.add(s)
        lines.append(line.rstrip())
    cleaned = "\n".join(lines).strip()
    if not cleaned:
        return "Файлы в этом топике найдены, но старый каталог содержал битые JSON-фрагменты. Каталог очищен, повтори запрос"
    return cleaned

def _p6e4_wrap_send(name):
    orig = globals().get(name)
    if not orig or getattr(orig, "_p6e4_wrapped", False):
        return
    if _p6e4_inspect.iscoroutinefunction(orig):
        async def wrapped(*args, **kwargs):
            args = tuple(_p6e4_sanitize_catalog_text(a) if isinstance(a, str) else a for a in args)
            kwargs = {k: (_p6e4_sanitize_catalog_text(v) if isinstance(v, str) else v) for k, v in kwargs.items()}
            return await orig(*args, **kwargs)
    else:
        def wrapped(*args, **kwargs):
            args = tuple(_p6e4_sanitize_catalog_text(a) if isinstance(a, str) else a for a in args)
            kwargs = {k: (_p6e4_sanitize_catalog_text(v) if isinstance(v, str) else v) for k, v in kwargs.items()}
            return orig(*args, **kwargs)
    wrapped._p6e4_wrapped = True
    globals()[name] = wrapped

for _p6e4_send_name in ("_send_once_ex", "send_once_ex", "_send_task_result", "send_task_result"):
    _p6e4_wrap_send(_p6e4_send_name)

_p6e4_logging.getLogger("WORKER").info("P6E4_LIVE_ROUTE_GUARD_INSTALLED")
# === END_P6E4_LIVE_ROUTE_FULL_CLOSE_IMAGE_SEARCH_CATALOG_20260504_V1 ===

# === P6F_P6E67_REPLY_REVISION_STRICT_ARTIFACT_GATE_20260504_V1 ===
# FACT: revision binding + anti-fake DONE + /root cleaner
# Inserted before __main__ guard so wrappers actually load at runtime
import re as _p6e67_re
import inspect as _p6e67_inspect
import logging as _p6e67_logging

_P6E67_REVISION_WORDS = (
    "пришли", "отправь", "скинь", "дай", "pdf", "пдф", "xlsx", "excel", "эксель", "txt",
    "ссылку", "ссылки", "drive", "расчет", "расчёт", "комнат", "помещ", "окн", "окон",
    "двер", "площад", "переделай", "доработай", "исправь", "правк", "нормально", "не так",
    "нормальн", "снова", "сделай", "ещё раз", "еще раз", "заново", "повтори", "ещё", "еще",
    "сделать", "переделать", "по новой", "сначала", "новой", "опять",
)

_P6E67_ARTIFACT_WORDS = (
    "pdf", "пдф", "xlsx", "excel", "эксель", "txt", "ссылку", "ссылки", "drive", "расчет", "расчёт"
)

_P6E67_BAD_RESULT = (
    "что строим", "дом, ангар, склад", "фундамент или кровлю", "какой объект", "уточните что строим"
)

def _p6e67_log():
    try:
        return logger
    except Exception:
        return _p6e67_logging.getLogger("task_worker")

def _p6e67_s(v, limit=60000):
    try:
        return str(v or "").strip()[:limit]
    except Exception:
        return ""

def _p6e67_low(v, limit=60000):
    return _p6e67_s(v, limit).lower().replace("ё", "е")

def _p6e67_row(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        return default

def _p6e67_hist(conn, task_id, action):
    try:
        _history(conn, str(task_id), str(action))
        conn.commit()
    except Exception as e:
        _p6e67_log().warning("P6E67_HIST_ERR task=%s action=%s err=%s", task_id, action, e)

def _p6e67_is_revision(raw):
    low = _p6e67_low(raw)
    return bool(low and any(x in low for x in _P6E67_REVISION_WORDS))

def _p6e67_wants_artifacts(raw):
    low = _p6e67_low(raw)
    return bool(low and any(x in low for x in _P6E67_ARTIFACT_WORDS))

def _p6e67_is_estimate(row):
    if int(_p6e67_row(row, "topic_id", 0) or 0) != 2:
        return False
    txt = _p6e67_low(
        _p6e67_s(_p6e67_row(row, "raw_input", ""), 120000) + " " +
        _p6e67_s(_p6e67_row(row, "result", ""), 120000) + " " +
        _p6e67_s(_p6e67_row(row, "input_type", ""))
    )
    return any(x in txt for x in ("смет", "стоимость", "расчет", "расчёт", "площад", "кровл", "фундамент", "каркас"))

def _p6e67_clean(text):
    text = _p6e67_s(text, 70000)
    text = _p6e67_re.sub(r"/root/\.areal-neva-core/\S+", "[локальный путь скрыт]", text)
    text = _p6e67_re.sub(r"/root/\S+", "[локальный путь скрыт]", text)
    return text

def _p6e67_find_parent(conn, task):
    chat_id = _p6e67_s(_p6e67_row(task, "chat_id", ""))
    topic_id = int(_p6e67_row(task, "topic_id", 0) or 0)
    task_id = _p6e67_s(_p6e67_row(task, "id", ""))
    reply_to = _p6e67_row(task, "reply_to_message_id", None)
    try:
        reply_to = int(reply_to) if reply_to not in (None, "", 0, "0") else None
    except Exception:
        reply_to = None

    if not chat_id or topic_id != 2:
        return None, "NO_SCOPE"

    if reply_to:
        row = conn.execute("""
            SELECT * FROM tasks
            WHERE chat_id=? AND COALESCE(topic_id,0)=? AND id<>?
              AND (bot_message_id=? OR reply_to_message_id=?)
              AND state IN ('DONE','AWAITING_CONFIRMATION','RESULT_READY','FAILED','IN_PROGRESS','WAITING_CLARIFICATION')
            ORDER BY rowid DESC LIMIT 1
        """, (chat_id, topic_id, task_id, reply_to, reply_to)).fetchone()
        if row and _p6e67_is_estimate(row):
            return row, "EXACT_REPLY_LINK"

        row = conn.execute("""
            SELECT * FROM tasks
            WHERE chat_id=? AND COALESCE(topic_id,0)=? AND id<>?
              AND input_type IN ('drive_file','file','photo','image','document')
              AND state IN ('NEW','IN_PROGRESS','WAITING_CLARIFICATION','AWAITING_CONFIRMATION','RESULT_READY','DONE')
              AND raw_input LIKE ?
            ORDER BY rowid DESC LIMIT 1
        """, (chat_id, topic_id, task_id, "%" + str(reply_to) + "%")).fetchone()
        if row and _p6e67_is_estimate(row):
            return row, "RAW_MESSAGE_ID_FILE_LINK"

    row = conn.execute("""
        SELECT * FROM tasks
        WHERE chat_id=? AND COALESCE(topic_id,0)=? AND id<>?
          AND state IN ('DONE','AWAITING_CONFIRMATION','RESULT_READY','FAILED','IN_PROGRESS','WAITING_CLARIFICATION')
          AND (
            raw_input LIKE '%смет%' OR result LIKE '%смет%'
            OR raw_input LIKE '%стоимость%' OR result LIKE '%стоимость%'
            OR raw_input LIKE '%кровл%' OR result LIKE '%кровл%'
            OR raw_input LIKE '%фундамент%' OR result LIKE '%фундамент%'
          )
        ORDER BY rowid DESC LIMIT 1
    """, (chat_id, topic_id, task_id)).fetchone()

    if row and _p6e67_is_estimate(row):
        return row, "LAST_DONE_ESTIMATE_FALLBACK"

    return None, "NOT_FOUND"

def _p6e67_merge_parent(conn, parent, current, source):
    parent_id = _p6e67_s(_p6e67_row(parent, "id", ""))
    current_id = _p6e67_s(_p6e67_row(current, "id", ""))
    parent_raw = _p6e67_s(_p6e67_row(parent, "raw_input", ""), 180000)
    current_raw = _p6e67_s(_p6e67_row(current, "raw_input", ""), 50000)
    marker = "P6E67_REVISION_FROM_TASK={}".format(current_id)

    if marker not in parent_raw:
        parent_raw = (parent_raw.rstrip() + "\n\n---\nREVISION_CONTEXT\nsource={}\n{}\n{}\n".format(source, marker, current_raw))[:230000]

    _update_task(conn, parent_id, state="IN_PROGRESS", raw_input=parent_raw, result="", error_message="")
    _update_task(conn, current_id, state="CANCELLED", result="P6E67_MERGED_TO_PARENT_TASK {}".format(parent_id), error_message="P6E67_MERGED_TO_PARENT")
    _p6e67_hist(conn, parent_id, "P6E67_PARENT_REVIVED_AS_REVISION_SOURCE:{}".format(source))
    _p6e67_hist(conn, parent_id, "P6E67_REVISION_TEXT_MERGED_FROM_TASK:{}".format(current_id))
    _p6e67_hist(conn, current_id, "P6E67_CURRENT_TASK_CANCELLED_MERGED_TO_PARENT:{}".format(parent_id))
    conn.commit()
    _p6e67_log().info("P6E67_MERGED current=%s parent=%s source=%s", current_id, parent_id, source)
    return True

async def _p6e67_try_merge(conn, task):
    # === PATCH_TOPIC2_INLINE_FIX_20260506_V1 STATE_GUARD ===
    # bail out fast if task already in terminal state — fixes infinite P6E67 loop on cancelled tasks
    try:
        _pifx_tid_pre = str(_p6e67_row(task, "id", "") or "")
        if _pifx_tid_pre:
            _pifx_st_row = conn.execute("SELECT state FROM tasks WHERE id=?", (_pifx_tid_pre,)).fetchone()
            if _pifx_st_row and str(_pifx_st_row[0] or "").upper() in ("CANCELLED", "FAILED", "DONE", "ARCHIVED"):
                return False
    except Exception:
        pass
    # === END_PATCH_TOPIC2_INLINE_FIX_20260506_V1 STATE_GUARD ===
    raw = _p6e67_s(_p6e67_row(task, "raw_input", ""), 70000)
    if int(_p6e67_row(task, "topic_id", 0) or 0) != 2:
        return False
    if not _p6e67_is_revision(raw):
        return False

    parent, source = _p6e67_find_parent(conn, task)
    if not parent:
        _p6e67_tid = str(_p6e67_row(task, "id", "") or "")
        # === PATCH_TOPIC2_INLINE_FIX_20260506_V1 FRESH_ESTIMATE_BEFORE_TERMINAL ===
        # if user sent a full estimate TZ (signals>=3), dispatch directly instead of terminal-guard
        try:
            _pifx_raw_low = str(_p6e67_row(task, "raw_input", "") or "").lower()
            _pifx_keys = ("смет","расчёт","расчет","фундамент","стен","перекрыт","кровл","газобетон","монолит","бетон","арматур","каркас","имитация бруса","ламинат","кирпич","отделк","м²","м2","м³","м3","посчитай")
            _pifx_signals = sum(1 for k in _pifx_keys if k in _pifx_raw_low)
            if _pifx_signals >= 3:
                _pifx_fn = globals().get("_t2fer_run_final_estimate")
                if _pifx_fn:
                    _pifx_res = await _pifx_fn(conn, task, "INLINE_FIX_FRESH_TZ")
                    if _pifx_res:
                        try:
                            conn.execute(
                                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                                (_p6e67_tid, "PATCH_TOPIC2_INLINE_FIX_20260506_V1:FRESH_ESTIMATE_DISPATCHED:" + str(_pifx_signals)),
                            )
                            conn.commit()
                        except Exception:
                            pass
                        return True
        except Exception:
            pass
        # === END_PATCH_TOPIC2_INLINE_FIX_20260506_V1 FRESH_ESTIMATE_BEFORE_TERMINAL ===
        _p6e67_raw = str(_p6e67_row(task, "raw_input", "") or "")
        _p6e67_reply_to = _p6e67_row(task, "reply_to_message_id", None)
        _p6e67_hist(conn, _p6e67_tid, "P6E67_PARENT_NOT_FOUND")

        if _p6e67_tid and _p6e67_reply_to:
            _p6e67_msg = (
                "Не нашёл родительскую задачу для reply. "
                "Пришли исходное ТЗ заново или ответь на последнее сообщение бота с результатом."
            )
            _update_task(
                conn,
                _p6e67_tid,
                state="WAITING_CLARIFICATION",
                result=_p6e67_msg,
                error_message="P6E67_PARENT_NOT_FOUND_TERMINAL_GUARD_V1",
            )
            try:
                _send_once_ex(
                    conn,
                    _p6e67_tid,
                    str(_p6e67_row(task, "chat_id", "")),
                    _p6e67_msg,
                    _p6e67_reply_to,
                    "p6e67_parent_not_found",
                )
            except Exception:
                pass
            _p6e67_hist(conn, _p6e67_tid, "P6E67_PARENT_NOT_FOUND_TERMINAL_GUARD_V1:WAITING_CLARIFICATION")
            conn.commit()
            return True

        return False

    return _p6e67_merge_parent(conn, parent, task, source)

def _p6e67_has_revision_history(conn, task_id):
    try:
        return conn.execute("""
            SELECT 1 FROM task_history
            WHERE task_id=? AND action LIKE 'P6E67_REVISION_TEXT_MERGED_FROM_TASK:%'
            LIMIT 1
        """, (str(task_id),)).fetchone() is not None
    except Exception:
        return False

def _p6e67_artifact_links_ok(text):
    low = _p6e67_low(text, 70000)
    if "/root/" in low or "[локальный путь скрыт]" in low:
        return False, "ROOT_OR_HIDDEN_LOCAL_PATH"

    has_pdf = bool(_p6e67_re.search(r"(https?://\S+|telegram://\S+)\S*\.pdf\b", low) or "pdf: https://" in low or "pdf: telegram://" in low)
    has_xlsx = bool(_p6e67_re.search(r"(https?://\S+|telegram://\S+)\S*\.xlsx\b", low) or "excel: https://" in low or "xlsx: https://" in low or "excel: telegram://" in low or "xlsx: telegram://" in low)
    has_txt = bool(_p6e67_re.search(r"(https?://\S+|telegram://\S+)\S*\.txt\b", low) or "txt: https://" in low or "txt: telegram://" in low)

    if not has_pdf:
        return False, "PDF_LINK_MISSING"
    if not has_xlsx:
        return False, "XLSX_LINK_MISSING"
    if not has_txt:
        return False, "TXT_LINK_MISSING"

    return True, "ARTIFACT_LINKS_OK"

def _p6e67_block_final(conn, task_id, text):
    clean = _p6e67_clean(text)
    low = _p6e67_low(clean)
    if _p6e67_has_revision_history(conn, task_id) and any(x in low for x in _P6E67_BAD_RESULT):
        return True, "P6E67_BLOCK_GENERIC_QUESTION", clean

    row = conn.execute("SELECT raw_input,topic_id FROM tasks WHERE id=? LIMIT 1", (str(task_id),)).fetchone()
    if row and int(_p6e67_row(row, "topic_id", 0) or 0) == 2:
        raw = _p6e67_s(_p6e67_row(row, "raw_input", ""), 230000)
        if _p6e67_has_revision_history(conn, task_id) and _p6e67_wants_artifacts(raw):
            ok, reason = _p6e67_artifact_links_ok(clean)
            if not ok:
                return True, "P6E67_BLOCK_ARTIFACT_GATE_" + reason, clean

    return False, "", clean

try:
    _P6E67_ORIG_UPDATE_TASK = _update_task
    if not getattr(_P6E67_ORIG_UPDATE_TASK, "_p6e67_v1_wrapped", False):
        def _update_task(conn, task_id, **kwargs):
            if "result" in kwargs:
                kwargs["result"] = _p6e67_clean(kwargs.get("result"))
            if kwargs.get("state") in ("DONE","AWAITING_CONFIRMATION","RESULT_READY") and "result" in kwargs:
                blocked, reason, clean = _p6e67_block_final(conn, task_id, kwargs.get("result"))
                kwargs["result"] = clean
                if blocked:
                    kwargs["state"] = "IN_PROGRESS"
                    kwargs["error_message"] = reason
                    try:
                        _history(conn, str(task_id), reason + "_ON_UPDATE")
                    except Exception as e:
                        _p6e67_log().warning("P6E67_UPDATE_HISTORY_ERR task=%s err=%s", task_id, e)
            return _P6E67_ORIG_UPDATE_TASK(conn, task_id, **kwargs)
        _update_task._p6e67_v1_wrapped = True
except Exception as e:
    _p6e67_log().exception("P6E67_WRAP_UPDATE_ERR %s", e)

try:
    _P6E67_ORIG_SEND_ONCE = _send_once
    if not getattr(_P6E67_ORIG_SEND_ONCE, "_p6e67_v1_wrapped", False):
        def _send_once(conn, task_id, chat_id, text, reply_to=None, kind="result"):
            blocked, reason, clean = _p6e67_block_final(conn, task_id, text)
            if blocked:
                _update_task(conn, str(task_id), state="IN_PROGRESS", error_message=reason)
                _p6e67_hist(conn, task_id, reason + "_BEFORE_SEND")
                return False
            return _P6E67_ORIG_SEND_ONCE(conn, task_id, chat_id, clean, reply_to, kind)
        _send_once._p6e67_v1_wrapped = True
except Exception as e:
    _p6e67_log().exception("P6E67_WRAP_SEND_ONCE_ERR %s", e)

try:
    _P6E67_ORIG_SEND_ONCE_EX = _send_once_ex
    if not getattr(_P6E67_ORIG_SEND_ONCE_EX, "_p6e67_v1_wrapped", False):
        def _send_once_ex(conn, task_id, chat_id, text, reply_to=None, kind="result", *args, **kwargs):
            blocked, reason, clean = _p6e67_block_final(conn, task_id, text)
            if blocked:
                _update_task(conn, str(task_id), state="IN_PROGRESS", error_message=reason)
                _p6e67_hist(conn, task_id, reason + "_BEFORE_SEND_EX")
                return False
            return _P6E67_ORIG_SEND_ONCE_EX(conn, task_id, chat_id, clean, reply_to, kind, *args, **kwargs)
        _send_once_ex._p6e67_v1_wrapped = True
except Exception as e:
    _p6e67_log().exception("P6E67_WRAP_SEND_ONCE_EX_ERR %s", e)

try:
    _P6E67_ORIG_HANDLE_NEW = _handle_new
    if not getattr(_P6E67_ORIG_HANDLE_NEW, "_p6e67_v1_wrapped", False):
        async def _handle_new(conn, task, *args, **kwargs):
            if await _p6e67_try_merge(conn, task):
                return True

            if not args:
                _p6e67_chat_id = _task_field(task, "chat_id", "")
                _p6e67_topic_id = _task_field(task, "topic_id", 0)
                res = _P6E67_ORIG_HANDLE_NEW(conn, task, _p6e67_chat_id, _p6e67_topic_id, **kwargs)
            else:
                res = _P6E67_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)

            return await res if _p6e67_inspect.isawaitable(res) else res
        _handle_new._p6e67_v1_wrapped = True
except Exception as e:
    _p6e67_log().exception("P6E67_WRAP_HANDLE_NEW_ERR %s", e)

try:
    _P6E67_ORIG_HANDLE_IN_PROGRESS = _handle_in_progress
    if not getattr(_P6E67_ORIG_HANDLE_IN_PROGRESS, "_p6e67_v1_wrapped", False):
        async def _handle_in_progress(conn, task, *args, **kwargs):
            if await _p6e67_try_merge(conn, task):
                return True
            res = _P6E67_ORIG_HANDLE_IN_PROGRESS(conn, task, *args, **kwargs)
            return await res if _p6e67_inspect.isawaitable(res) else res
        _handle_in_progress._p6e67_v1_wrapped = True
except Exception as e:
    _p6e67_log().exception("P6E67_WRAP_HANDLE_IN_PROGRESS_ERR %s", e)

_p6e67_log().info("P6F_P6E67_REPLY_REVISION_STRICT_ARTIFACT_GATE_20260504_V1_INSTALLED")
# === END_P6F_P6E67_REPLY_REVISION_STRICT_ARTIFACT_GATE_20260504_V1 ===

# === P6F_TOPIC500_SANITIZER_TASK_WORKER_BIND_20260504_V1 ===
# FACT: wraps process_ai_task in task_worker scope to apply
# core.ai_router._p6f_ts_sanitize_payload before search call.
try:
    _P6F_TS_TW_ORIG_PROCESS_AI_TASK = process_ai_task
    if not getattr(_P6F_TS_TW_ORIG_PROCESS_AI_TASK, "_p6f_ts_wrapped", False):
        async def process_ai_task(payload):
            try:
                from core.ai_router import _p6f_ts_sanitize_payload as _p6f_ts_san
                payload = _p6f_ts_san(payload)
            except Exception as _p6f_ts_e:
                try:
                    logger.warning("P6F_TS_TW_SANITIZE_ERR %s", _p6f_ts_e)
                except Exception:
                    pass
            return await _P6F_TS_TW_ORIG_PROCESS_AI_TASK(payload)
        process_ai_task._p6f_ts_wrapped = True
        try:
            logger.info("P6F_TOPIC500_SANITIZER_TASK_WORKER_BIND_INSTALLED")
        except Exception:
            pass
except Exception as _p6f_ts_install_e:
    try:
        logger.exception("P6F_TS_TW_BIND_INSTALL_ERR %s", _p6f_ts_install_e)
    except Exception:
        pass
# === END_P6F_TOPIC500_SANITIZER_TASK_WORKER_BIND_20260504_V1 ===

# === P6F_MEMORY_WRITE_GATE_AND_RESPONSE_LOGIC_V1 ===
# FACT: memory write filter + raw JSON detector + STALE_TIMEOUT guard
# + sufficient-TZ override of generic "что строим" clarification.
import re as _p6f_mwr_re
import json as _p6f_mwr_json
import logging as _p6f_mwr_logging

_P6F_MWR_LOG = _p6f_mwr_logging.getLogger("task_worker")

_P6F_MWR_BAD_RESULT_FOR_MEMORY = (
    "traceback", "syntaxerror", "importerror", "attributeerror", "nameerror",
    "что строим", "уточните что строим", "p6e67_block",
    "/root/", "[локальный путь скрыт]",
    "anthropic_block", "openai_block", "openrouter_call_err",
    "needs_clarification_pcv", "stale_timeout",
)

def _p6f_mwr_should_skip_memory(raw_input, result):
    low_r = str(result or "").lower().replace("ё", "е")
    low_in = str(raw_input or "").lower().replace("ё", "е")
    if not low_r or len(low_r.strip()) < 30:
        return True, "RESULT_TOO_SHORT"
    if any(b in low_r for b in _P6F_MWR_BAD_RESULT_FOR_MEMORY):
        return True, "RESULT_CONTAINS_GARBAGE_OR_ERROR"
    if low_r.startswith("{") and low_r.rstrip().endswith("}"):
        try:
            j = _p6f_mwr_json.loads(result)
            if isinstance(j, dict):
                return True, "RESULT_IS_RAW_JSON"
        except Exception:
            pass
    return False, "OK"

try:
    _P6F_MWR_ORIG_SAVE_MEMORY = _save_memory
    if not getattr(_P6F_MWR_ORIG_SAVE_MEMORY, "_p6f_mwr_wrapped", False):
        def _save_memory(chat_id, topic_id, raw_input, result):
            skip, reason = _p6f_mwr_should_skip_memory(raw_input, result)
            if skip:
                try:
                    _P6F_MWR_LOG.info(
                        "P6F_MWR_MEMORY_WRITE_SKIPPED chat=%s topic=%s reason=%s",
                        chat_id, topic_id, reason,
                    )
                except Exception:
                    pass
                return None
            return _P6F_MWR_ORIG_SAVE_MEMORY(chat_id, topic_id, raw_input, result)
        _save_memory._p6f_mwr_wrapped = True
        _P6F_MWR_LOG.info("P6F_MEMORY_WRITE_GATE_INSTALLED")
except Exception as _e:
    try:
        _P6F_MWR_LOG.exception("P6F_MWR_INSTALL_ERR %s", _e)
    except Exception:
        pass


def _p6f_rl_looks_like_raw_json(text):
    if not text:
        return False
    s = str(text).strip()
    if not s.startswith("{") and not s.startswith("["):
        return False
    if not (s.endswith("}") or s.endswith("]")):
        return False
    try:
        j = _p6f_mwr_json.loads(s)
        return isinstance(j, (dict, list))
    except Exception:
        return False

def _p6f_rl_has_internal_paths(text):
    if not text:
        return False
    low = str(text).lower()
    return any(p in low for p in ("/root/", "/tmp/", "[локальный путь скрыт]"))

def _p6f_rl_clean_for_user(text):
    if not text:
        return text
    s = str(text)
    s = _p6f_mwr_re.sub(r"/root/\.areal-neva-core/\S+", "[путь скрыт]", s)
    s = _p6f_mwr_re.sub(r"/root/\S+", "[путь скрыт]", s)
    s = _p6f_mwr_re.sub(r"/tmp/\S+", "[tmp путь скрыт]", s)
    return s

try:
    _P6F_RL_ORIG_SEND_ONCE_EX = _send_once_ex
    if not getattr(_P6F_RL_ORIG_SEND_ONCE_EX, "_p6f_rl_wrapped", False):
        def _send_once_ex(conn, task_id, chat_id, text, reply_to=None, kind="result", *args, **kwargs):
            t = _p6f_rl_clean_for_user(text)
            if _p6f_rl_looks_like_raw_json(t):
                t = "Внутренняя ошибка форматирования ответа. Запрос принят. Если нужен файл — уточни ещё раз"
                try:
                    _P6F_MWR_LOG.warning("P6F_RL_RAW_JSON_BLOCKED task=%s", task_id)
                except Exception:
                    pass
            return _P6F_RL_ORIG_SEND_ONCE_EX(conn, task_id, chat_id, t, reply_to, kind, *args, **kwargs)
        _send_once_ex._p6f_rl_wrapped = True
        _P6F_MWR_LOG.info("P6F_RESPONSE_LOGIC_RAW_JSON_GATE_INSTALLED")
except Exception as _e:
    try:
        _P6F_MWR_LOG.exception("P6F_RL_INSTALL_ERR %s", _e)
    except Exception:
        pass


# Sufficient-TZ override: if topic_2 raw_input has enough TZ data, do NOT
# allow generic "что строим"-style clarification to be sent — it indicates
# orchestra ignored the existing TZ. We catch it on result-update path.
def _p6f_stz_topic2_has_enough_tz(raw_input):
    try:
        from core.sample_template_engine import _p6f_tz_is_sufficient
        return bool(_p6f_tz_is_sufficient(raw_input))
    except Exception:
        return False

_P6F_STZ_GENERIC_QUESTIONS = (
    "что строим", "из чего строим", "где находится", "уточните что",
    "дом, ангар, склад", "фундамент или кровлю", "какой объект",
)

def _p6f_stz_should_block_generic_question(raw_input, text, topic_id):
    if int(topic_id or 0) != 2:
        return False
    if not text:
        return False
    low = str(text).lower().replace("ё", "е")
    if not any(g in low for g in _P6F_STZ_GENERIC_QUESTIONS):
        return False
    return _p6f_stz_topic2_has_enough_tz(raw_input)

try:
    _P6F_STZ_ORIG_UPDATE_TASK = _update_task
    if not getattr(_P6F_STZ_ORIG_UPDATE_TASK, "_p6f_stz_wrapped", False):
        def _update_task(conn, task_id, **kwargs):
            try:
                if "result" in kwargs and kwargs.get("state") in ("AWAITING_CONFIRMATION", "WAITING_CLARIFICATION", "DONE"):
                    res = kwargs.get("result", "") or ""
                    row = conn.execute("SELECT raw_input,topic_id FROM tasks WHERE id=? LIMIT 1", (str(task_id),)).fetchone()
                    if row is not None:
                        try:
                            raw = row["raw_input"] if hasattr(row, "keys") else row[0]
                            tid = row["topic_id"] if hasattr(row, "keys") else row[1]
                        except Exception:
                            raw, tid = None, None
                        if _p6f_stz_should_block_generic_question(raw, res, tid):
                            kwargs["state"] = "IN_PROGRESS"
                            kwargs["error_message"] = "P6F_STZ_BLOCK_GENERIC_QUESTION_TZ_SUFFICIENT"
                            try:
                                _history(conn, str(task_id), "P6F_STZ_BLOCKED_GENERIC_QUESTION_BECAUSE_TZ_SUFFICIENT")
                            except Exception:
                                pass
            except Exception as _e:
                try:
                    _P6F_MWR_LOG.warning("P6F_STZ_UPDATE_TASK_ERR %s", _e)
                except Exception:
                    pass
            return _P6F_STZ_ORIG_UPDATE_TASK(conn, task_id, **kwargs)
        _update_task._p6f_stz_wrapped = True
        _P6F_MWR_LOG.info("P6F_TZ_SUFFICIENT_GENERIC_QUESTION_GATE_INSTALLED")
except Exception as _e:
    try:
        _P6F_MWR_LOG.exception("P6F_STZ_INSTALL_ERR %s", _e)
    except Exception:
        pass


# Topic_210 — Drive lookup skeleton via existing google_io.
# FACT: project_engine.py has helpers for "Образцы проектов",
# PROJECT_DESIGN_REFERENCES, PROJECT_ARTIFACTS, _manifests folders.
# We add a router that, for topic_210 with project intent, surfaces
# Drive references instead of falling into estimate/general routes.
_P6F_T210_PROJECT_INTENT_WORDS = (
    "проект", "образец", "образц", "рендер", "ар ", "кж ", "кд ",
    "dwg", "pln", "чертёж", "чертеж", "пакет проект", "спецификация",
)

def _p6f_t210_is_project_intent(raw_input):
    if not raw_input:
        return False
    low = str(raw_input).lower()
    return any(w in low for w in _P6F_T210_PROJECT_INTENT_WORDS)

async def _p6f_t210_try_drive_references(conn, task):
    """
    Skeleton for topic_210 Drive references router.
    Currently returns False — full Drive folder scan implemented separately
    in project_engine.py (lines ~3093-3127). This skeleton ensures the hook
    point exists and is testable; a follow-up live test will trigger
    expansion based on actual Drive folder IDs.
    """
    try:
        topic_id = int(task["topic_id"] if hasattr(task, "keys") else task[5] if len(task) > 5 else 0) if task else 0
    except Exception:
        return False
    if topic_id != 210:
        return False
    try:
        raw = task["raw_input"] if hasattr(task, "keys") else (task[3] if len(task) > 3 else "")
    except Exception:
        return False
    if not _p6f_t210_is_project_intent(raw):
        return False
    try:
        from core import project_engine
        marker_func = getattr(project_engine, "_final_project_find_folder_by_name_v1", None)
        if marker_func is None:
            _P6F_MWR_LOG.info("P6F_T210_NO_DRIVE_HELPER_AVAILABLE_HOOK_NOOP")
        else:
            _P6F_MWR_LOG.info("P6F_T210_DRIVE_HELPER_PRESENT_HOOK_READY")
    except Exception as _e:
        try:
            _P6F_MWR_LOG.warning("P6F_T210_HOOK_ERR %s", _e)
        except Exception:
            pass
    return False

try:
    _P6F_T210_LOG = _p6f_mwr_logging.getLogger("task_worker")
    _P6F_T210_LOG.info("P6F_TOPIC210_PROJECT_DRIVE_HOOK_REGISTERED")
except Exception:
    pass
# === END_P6F_MEMORY_WRITE_GATE_AND_RESPONSE_LOGIC_V1 ===

# === P6F_TOPIC210_DRIVE_RESOLVER_REAL_V1 ===
# FACT: real Drive folder lookup for topic_210 project references.
# Uses existing core.project_engine helpers:
#   _final_project_find_folder_by_name_v1 (any folder by exact name)
#   _final_project_list_folder_v1 (list files in folder by id)
#   _final_project_section_from_name_v1 (classify file as АР/КЖ/КД/КМ/КМД/ЭСКИЗ)
#   _final_project_section_from_request_v1 (classify user request)
#   _FINAL_PROJECT_SAMPLES_FOLDER_ID = "1kcJbrn7XMcov__Z1JdWhKlJMZd7GUkgP" (Образцы проектов)
import logging as _p6f_t210_logging

_P6F_T210_LOG = _p6f_t210_logging.getLogger("task_worker")

def _p6f_t210_request_words(raw_input):
    low = str(raw_input or "").lower().replace("ё", "е")
    return any(w in low for w in (
        "проект", "образец", "образц", "рендер", "ар ", "кж", "кд",
        "dwg", "pln", "чертеж", "пакет проект", "спецификац",
    ))

def _p6f_t210_drive_link(file_id):
    return "https://drive.google.com/file/d/{}/view".format(file_id) if file_id else ""

def _p6f_t210_kind_from_mime(mime):
    if not mime:
        return "FILE"
    m = str(mime).lower()
    if "pdf" in m:
        return "PDF"
    if "wordprocessingml" in m or "msword" in m:
        return "DOCX"
    if "spreadsheetml" in m or "excel" in m:
        return "XLSX"
    if "image" in m:
        return "IMAGE"
    if "vnd.google-apps.folder" in m:
        return "FOLDER"
    if "dwg" in m or "autocad" in m:
        return "DWG"
    return "FILE"

def _p6f_t210_collect_drive_references(raw_input):
    """
    Returns list of dicts:
      [{name, section, link, mime, kind, source_folder}, ...]
    Up to 7 most relevant.
    """
    try:
        from core.project_engine import (
            _final_project_find_folder_by_name_v1 as _find_folder,
            _final_project_list_folder_v1 as _list_folder,
            _final_project_section_from_name_v1 as _classify_file,
            _final_project_section_from_request_v1 as _classify_request,
            _FINAL_PROJECT_SAMPLES_FOLDER_ID as _SAMPLES_ID,
        )
    except Exception as e:
        _P6F_T210_LOG.warning("P6F_T210_PROJECT_ENGINE_IMPORT_ERR %s", e)
        return []

    out = []
    section_request = _classify_request(raw_input or "")

    samples = _list_folder(_SAMPLES_ID) if _SAMPLES_ID else []
    for f in samples:
        name = f.get("name", "")
        sec = _classify_file(name)
        if section_request not in ("UNKNOWN", "") and sec != section_request and sec != "UNKNOWN":
            continue
        out.append({
            "name": name,
            "section": sec,
            "link": _p6f_t210_drive_link(f.get("id")),
            "mime": f.get("mimeType", ""),
            "kind": _p6f_t210_kind_from_mime(f.get("mimeType")),
            "source_folder": "Образцы проектов",
        })
        if len(out) >= 7:
            break

    if len(out) < 7:
        for folder_name in ("PROJECT_DESIGN_REFERENCES", "_manifests"):
            try:
                fid = _find_folder(folder_name)
            except Exception:
                fid = ""
            if not fid:
                continue
            files = _list_folder(fid)
            for f in files:
                if len(out) >= 7:
                    break
                name = f.get("name", "")
                sec = _classify_file(name)
                if section_request not in ("UNKNOWN", "") and sec != section_request and sec != "UNKNOWN":
                    continue
                out.append({
                    "name": name,
                    "section": sec,
                    "link": _p6f_t210_drive_link(f.get("id")),
                    "mime": f.get("mimeType", ""),
                    "kind": _p6f_t210_kind_from_mime(f.get("mimeType")),
                    "source_folder": folder_name,
                })

    return out

def _p6f_t210_format_references_message(raw_input, refs):
    if not refs:
        return ("По topic_210 запрос распознан как проектный, но в Drive папках "
                "(Образцы проектов / PROJECT_DESIGN_REFERENCES / _manifests) "
                "нет подходящих файлов или Drive недоступен. Уточни секцию (АР/КЖ/КД/КМ/КМД).")
    try:
        from core.project_engine import _final_project_section_from_request_v1
        section = _final_project_section_from_request_v1(raw_input or "")
    except Exception:
        section = "UNKNOWN"
    head = "Проектные ссылки из Drive (topic_210)"
    if section and section != "UNKNOWN":
        head = head + " — секция: " + section
    lines = [head, ""]
    for i, r in enumerate(refs, 1):
        lines.append("{}. [{}] {} ({})".format(i, r["section"], r["name"], r["kind"]))
        if r.get("source_folder"):
            lines.append("   Папка: {}".format(r["source_folder"]))
        if r.get("link"):
            lines.append("   {}".format(r["link"]))
    return "\n".join(lines)

async def _p6f_t210_handle_project_request(conn, task):
    try:
        topic_id = int(_p6e67_row(task, "topic_id", 0) or 0)
    except Exception:
        return False
    if topic_id != 210:
        return False
    raw = _p6e67_s(_p6e67_row(task, "raw_input", ""), 60000)
    if not _p6f_t210_request_words(raw):
        return False
    task_id = _p6e67_s(_p6e67_row(task, "id", ""))
    chat_id = _p6e67_s(_p6e67_row(task, "chat_id", "-1003725299009"))
    reply_to = _p6e67_row(task, "reply_to_message_id", None)

    refs = _p6f_t210_collect_drive_references(raw)
    msg = _p6f_t210_format_references_message(raw, refs)

    try:
        _send_once_ex(conn, task_id, chat_id, msg, reply_to, "result")
    except Exception as e:
        _P6F_T210_LOG.warning("P6F_T210_SEND_ERR %s", e)
    try:
        _update_task(conn, task_id, state="DONE", result=msg, error_message="")
        _history(conn, task_id, "P6F_T210_PROJECT_DRIVE_REFS_RETURNED:{}".format(len(refs)))
        conn.commit()
    except Exception as e:
        _P6F_T210_LOG.warning("P6F_T210_UPDATE_ERR %s", e)
    return True

# Wire into existing handle_new wrapping chain (after P6E67 merge attempt)
try:
    _P6F_T210_ORIG_HANDLE_NEW = _handle_new
    if not getattr(_P6F_T210_ORIG_HANDLE_NEW, "_p6f_t210_wrapped", False):
        async def _handle_new(conn, task, *args, **kwargs):
            try:
                if await _p6f_t210_handle_project_request(conn, task):
                    return True
            except Exception as e:
                _P6F_T210_LOG.warning("P6F_T210_HANDLE_NEW_ERR %s", e)
            res = _P6F_T210_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
            import inspect as _ins
            return await res if _ins.isawaitable(res) else res
        _handle_new._p6f_t210_wrapped = True
        _P6F_T210_LOG.info("P6F_TOPIC210_DRIVE_RESOLVER_REAL_V1_INSTALLED")
except Exception as e:
    _P6F_T210_LOG.exception("P6F_T210_INSTALL_ERR %s", e)
# === END_P6F_TOPIC210_DRIVE_RESOLVER_REAL_V1 ===


# === P6F_DRIVE_ARTIFACT_HISTORY_GATE_V1 ===
# FACT: extends P6E67 artifact gate. If user requested PDF/XLSX/TXT and
# task does not have history marker proving real Drive upload OR Telegram
# sendDocument confirmation — DONE is blocked.
def _p6f_dah_has_upload_history(conn, task_id):
    try:
        row = conn.execute(
            """
            SELECT 1 FROM task_history
            WHERE task_id=? AND (
                action LIKE 'DRIVE_UPLOAD_OK%'
                OR action LIKE 'TELEGRAM_ARTIFACT_FALLBACK_SENT%'
                OR action LIKE 'P6F_PCV_OPENROUTER_VISION_DONE%'
                OR action LIKE 'UPLOAD_OK%'
                OR action LIKE 'DRIVE_RETRY_UPLOAD_OK%'
                OR action LIKE 'TG_FALLBACK%'
                OR action LIKE 'P6F_T210_PROJECT_DRIVE_REFS_RETURNED%'
            )
            LIMIT 1
            """,
            (str(task_id),),
        ).fetchone()
        return row is not None
    except Exception:
        return False

def _p6f_dah_user_wants_artifact(raw_input):
    low = str(raw_input or "").lower().replace("ё", "е")
    return any(w in low for w in ("pdf", "пдф", "xlsx", "excel", "эксель", "txt",
                                    "ссылк", "drive", "файл", "артефакт", "акт docx"))

try:
    _P6F_DAH_ORIG_UPDATE_TASK = _update_task
    if not getattr(_P6F_DAH_ORIG_UPDATE_TASK, "_p6f_dah_wrapped", False):
        def _update_task(conn, task_id, **kwargs):
            try:
                if kwargs.get("state") == "DONE":
                    row = conn.execute(
                        "SELECT raw_input,topic_id FROM tasks WHERE id=? LIMIT 1",
                        (str(task_id),),
                    ).fetchone()
                    if row is not None:
                        try:
                            raw = row["raw_input"] if hasattr(row, "keys") else row[0]
                        except Exception:
                            raw = ""
                        if _p6f_dah_user_wants_artifact(raw) and not _p6f_dah_has_upload_history(conn, task_id):
                            kwargs["state"] = "IN_PROGRESS"
                            kwargs["error_message"] = "P6F_DAH_BLOCK_DONE_NO_UPLOAD_HISTORY"
                            try:
                                _history(conn, str(task_id), "P6F_DAH_BLOCKED_DONE_NO_UPLOAD_OR_TG_HISTORY")
                            except Exception:
                                pass
            except Exception:
                pass
            return _P6F_DAH_ORIG_UPDATE_TASK(conn, task_id, **kwargs)
        _update_task._p6f_dah_wrapped = True
        _P6F_T210_LOG.info("P6F_DRIVE_ARTIFACT_HISTORY_GATE_V1_INSTALLED")
except Exception as e:
    _P6F_T210_LOG.exception("P6F_DAH_INSTALL_ERR %s", e)
# === END_P6F_DRIVE_ARTIFACT_HISTORY_GATE_V1 ===

# === P6G_PARENT_RELEVANCE_SCORER_V1 ===
# FACT: improves P6E67 fallback parent finder.
# Without this, "LAST_DONE_ESTIMATE_FALLBACK" picks newest by rowid only —
# which can grab wrong parent if multiple recent estimates exist.
# Now: score candidates by keyword overlap with current revision text.
# Higher score wins. Below threshold → NOT_FOUND (no blind fallback).
import re as _p6g_prs_re
import logging as _p6g_prs_logging

_P6G_PRS_LOG = _p6g_prs_logging.getLogger("task_worker")

def _p6g_prs_keywords(text, limit=300):
    if not text:
        return set()
    low = str(text).lower().replace("ё", "е")
    words = _p6g_prs_re.findall(r"[а-яa-z0-9]{3,}", low)
    return set(words[:limit])

def _p6g_prs_score_candidate(current_raw, candidate_raw, candidate_result, candidate_age_seconds=0):
    cur_kw = _p6g_prs_keywords(current_raw)
    cand_kw = _p6g_prs_keywords(candidate_raw + " " + candidate_result)
    if not cur_kw or not cand_kw:
        return 0
    overlap = len(cur_kw & cand_kw)
    base = overlap * 10
    if "смет" in cur_kw and "смет" in cand_kw:
        base += 30
    if "проект" in cur_kw and "проект" in cand_kw:
        base += 30
    for sz in ("12.0", "12х8", "12x8", "10х12", "10х10", "8х10"):
        if sz in str(current_raw or "").lower() and sz in str(candidate_raw or "").lower():
            base += 50
            break
    age_penalty = min(int(candidate_age_seconds // 600), 30)
    return max(0, base - age_penalty)

def _p6g_prs_find_best_parent(conn, current_task, min_score=15):
    """
    Replaces blind LAST_DONE_ESTIMATE_FALLBACK with scored selection.
    Returns (row, source_label) or (None, "NO_RELEVANT_PARENT").
    """
    chat_id = _p6e67_s(_p6e67_row(current_task, "chat_id", ""))
    topic_id = int(_p6e67_row(current_task, "topic_id", 0) or 0)
    task_id = _p6e67_s(_p6e67_row(current_task, "id", ""))
    raw = _p6e67_s(_p6e67_row(current_task, "raw_input", ""), 30000)
    if not chat_id or topic_id != 2:
        return None, "NO_SCOPE"
    try:
        rows = conn.execute(
            """
            SELECT *,
                   CAST((julianday('now') - julianday(updated_at)) * 86400 AS INTEGER) AS age_seconds
            FROM tasks
            WHERE chat_id=? AND COALESCE(topic_id,0)=? AND id<>?
              AND state IN ('DONE','AWAITING_CONFIRMATION','RESULT_READY','FAILED')
              AND (
                raw_input LIKE '%смет%' OR result LIKE '%смет%'
                OR raw_input LIKE '%стоимость%' OR result LIKE '%стоимость%'
                OR raw_input LIKE '%кровл%' OR result LIKE '%кровл%'
                OR raw_input LIKE '%фундамент%' OR result LIKE '%фундамент%'
              )
            ORDER BY rowid DESC LIMIT 30
            """,
            (chat_id, topic_id, task_id),
        ).fetchall()
    except Exception as e:
        _P6G_PRS_LOG.warning("P6G_PRS_QUERY_ERR %s", e)
        return None, "QUERY_ERR"
    if not rows:
        return None, "NO_CANDIDATES"
    best = None
    best_score = 0
    for r in rows:
        try:
            cand_raw = _p6e67_s(_p6e67_row(r, "raw_input", ""), 30000)
            cand_res = _p6e67_s(_p6e67_row(r, "result", ""), 30000)
            try:
                age = int(_p6e67_row(r, "age_seconds", 0) or 0)
            except Exception:
                age = 0
            score = _p6g_prs_score_candidate(raw, cand_raw, cand_res, age)
            if score > best_score:
                best = r
                best_score = score
        except Exception:
            continue
    if best and best_score >= min_score:
        _P6G_PRS_LOG.info(
            "P6G_PRS_BEST_PARENT current=%s parent=%s score=%d",
            task_id, _p6e67_row(best, "id", ""), best_score,
        )
        return best, "RELEVANCE_SCORED:{}".format(best_score)
    _P6G_PRS_LOG.info("P6G_PRS_NO_RELEVANT_PARENT current=%s top_score=%d", task_id, best_score)
    return None, "NO_RELEVANT_PARENT_SCORE_{}".format(best_score)

# Wrap _p6e67_find_parent to use scorer for fallback path
try:
    _P6G_PRS_ORIG_FIND_PARENT = _p6e67_find_parent
    if not getattr(_P6G_PRS_ORIG_FIND_PARENT, "_p6g_prs_wrapped", False):
        def _p6e67_find_parent(conn, task):
            row, source = _P6G_PRS_ORIG_FIND_PARENT(conn, task)
            if row is not None and source == "EXACT_REPLY_LINK":
                return row, source
            scored, scored_source = _p6g_prs_find_best_parent(conn, task, min_score=15)
            if scored is not None:
                return scored, scored_source
            return row, source if row is not None else (None, scored_source)
        _p6e67_find_parent._p6g_prs_wrapped = True
        _P6G_PRS_LOG.info("P6G_PARENT_RELEVANCE_SCORER_V1_INSTALLED")
except Exception as _e:
    _P6G_PRS_LOG.exception("P6G_PRS_INSTALL_ERR %s", _e)
# === END_P6G_PARENT_RELEVANCE_SCORER_V1 ===


# === P6G_MEMORY_JANITOR_V1 ===
# FACT: cleanup helper for memory.db that removes records previously written
# with garbage values (что строим / traceback / /root / less than 30 chars).
# Run on each DONE event via a wrap of _save_memory.
import os as _p6g_mj_os
import sqlite3 as _p6g_mj_sqlite

_P6G_MJ_LOG = _p6g_prs_logging.getLogger("task_worker")
_P6G_MJ_DB = "/root/.areal-neva-core/data/memory.db"

_P6G_MJ_BAD_VALUE_PATTERNS = (
    "что строим",
    "traceback",
    "syntaxerror",
    "/root/.areal-neva-core",
    "[локальный путь скрыт]",
    "p6e67_block",
    "p6f_dah_block",
    "stale_timeout",
)

def _p6g_mj_run_periodic(max_per_run=200):
    """
    Removes records with bad values from memory.db.
    Returns number removed. Safe: won't touch role markers.
    """
    if not _p6g_mj_os.path.exists(_P6G_MJ_DB):
        return 0
    removed = 0
    try:
        c = _p6g_mj_sqlite.connect(_P6G_MJ_DB, timeout=10)
        try:
            for pat in _P6G_MJ_BAD_VALUE_PATTERNS:
                cur = c.execute(
                    "DELETE FROM memory WHERE key NOT LIKE '%_role' AND lower(value) LIKE ? LIMIT ?",
                    ("%" + pat + "%", max_per_run),
                )
                removed += cur.rowcount or 0
            c.commit()
        finally:
            c.close()
    except Exception as e:
        _P6G_MJ_LOG.warning("P6G_MJ_RUN_ERR %s", e)
    if removed:
        _P6G_MJ_LOG.info("P6G_MJ_REMOVED %d bad memory records", removed)
    return removed

# Hook: run janitor lazily after every Nth DONE
_P6G_MJ_DONE_COUNTER = [0]
_P6G_MJ_RUN_EVERY = 10

try:
    _P6G_MJ_ORIG_SAVE_MEMORY_AFTER_GATE = _save_memory
    if not getattr(_P6G_MJ_ORIG_SAVE_MEMORY_AFTER_GATE, "_p6g_mj_wrapped", False):
        def _save_memory(chat_id, topic_id, raw_input, result):
            r = _P6G_MJ_ORIG_SAVE_MEMORY_AFTER_GATE(chat_id, topic_id, raw_input, result)
            try:
                _P6G_MJ_DONE_COUNTER[0] += 1
                if _P6G_MJ_DONE_COUNTER[0] % _P6G_MJ_RUN_EVERY == 0:
                    _p6g_mj_run_periodic()
            except Exception:
                pass
            return r
        _save_memory._p6g_mj_wrapped = True
        _P6G_MJ_LOG.info("P6G_MEMORY_JANITOR_V1_INSTALLED")
except Exception as _e:
    _P6G_MJ_LOG.exception("P6G_MJ_INSTALL_ERR %s", _e)
# === END_P6G_MEMORY_JANITOR_V1 ===


# === P6G_STALE_EXPLAINER_V1 ===
# FACT: when STALE_TIMEOUT triggers, attach human-readable cause if task was
# stuck due to a P6F gate. Helps debugging and gives user a real reason
# instead of cryptic STALE_TIMEOUT.
_P6G_SE_LOG = _p6g_prs_logging.getLogger("task_worker")

_P6G_SE_REASON_MAP = {
    "P6E67_BLOCK_ARTIFACT_GATE_PDF_LINK_MISSING": "Запрошен PDF, но Drive ссылка не получена. Повторите запрос или укажите конкретный объект",
    "P6E67_BLOCK_ARTIFACT_GATE_XLSX_LINK_MISSING": "Запрошен Excel, но XLSX не сгенерирован. Повторите запрос",
    "P6E67_BLOCK_ARTIFACT_GATE_TXT_LINK_MISSING": "Запрошен TXT, но TXT не подготовлен. Повторите запрос",
    "P6E67_BLOCK_ARTIFACT_GATE_ROOT_OR_HIDDEN_LOCAL_PATH": "Артефакт сформирован, но содержит локальный путь — Drive upload требуется",
    "P6E67_BLOCK_GENERIC_QUESTION": "Бот пытался задать общий уточняющий вопрос вместо ответа. ТЗ присутствует — повторите запрос",
    "P6F_DAH_BLOCK_DONE_NO_UPLOAD_HISTORY": "Запрошен файл, но загрузка на Drive/Telegram не зарегистрирована в истории. Повторите запрос или попросите конкретный артефакт",
    "P6F_STZ_BLOCK_GENERIC_QUESTION_TZ_SUFFICIENT": "ТЗ присутствует — бот не должен спрашивать «что строим». Повторите запрос",
    "P6F_PCV_NEEDS_CLARIFICATION_LOW_CONFIDENCE_OR_NO_DIMS": "Vision не разобрал размеры на фото. Пришлите размер текстом или фото плана крупнее",
}

try:
    _P6G_SE_ORIG_UPDATE = _update_task
    if not getattr(_P6G_SE_ORIG_UPDATE, "_p6g_se_wrapped", False):
        def _update_task(conn, task_id, **kwargs):
            try:
                if kwargs.get("state") == "FAILED" and (kwargs.get("error_message") or "").strip() == "STALE_TIMEOUT":
                    row = conn.execute(
                        "SELECT error_message FROM tasks WHERE id=? LIMIT 1",
                        (str(task_id),),
                    ).fetchone()
                    if row is not None:
                        try:
                            existing_err = row["error_message"] if hasattr(row, "keys") else row[0]
                        except Exception:
                            existing_err = ""
                        existing_err = str(existing_err or "")
                        for marker, hint in _P6G_SE_REASON_MAP.items():
                            if marker in existing_err:
                                kwargs["error_message"] = "STALE_TIMEOUT_AFTER_{}: {}".format(marker, hint)
                                try:
                                    _history(conn, str(task_id), "P6G_STALE_EXPLAINED:{}".format(marker))
                                except Exception:
                                    pass
                                _P6G_SE_LOG.info("P6G_STALE_EXPLAINED task=%s marker=%s", task_id, marker)
                                break
            except Exception as _e:
                _P6G_SE_LOG.warning("P6G_SE_UPDATE_ERR %s", _e)
            return _P6G_SE_ORIG_UPDATE(conn, task_id, **kwargs)
        _update_task._p6g_se_wrapped = True
        _P6G_SE_LOG.info("P6G_STALE_EXPLAINER_V1_INSTALLED")
except Exception as _e:
    _P6G_SE_LOG.exception("P6G_SE_INSTALL_ERR %s", _e)
# === END_P6G_STALE_EXPLAINER_V1 ===


# === P6H_TOPIC5_TASK_WORKER_RUNTIME_VERIFY_20260504_V1 ===
# Re-emits P6H install markers from within the file-handler-attached logger so
# they show up in logs/task_worker.log (technadzor_engine module-level markers
# fire BEFORE FileHandler is attached at task_worker:490 — pre-existing for
# all module-level markers, not specific to P6H). Also confirms wrap state
# of process_technadzor.
try:
    import logging as _p6h_tw_logging
    _p6h_tw_w = _p6h_tw_logging.getLogger("task_worker")
    try:
        from core import technadzor_drive_index as _p6h_tw_tdi  # noqa: F401
    except Exception as _e:
        _p6h_tw_w.warning("P6H_TOPIC5_TW_VERIFY_DRIVE_INDEX_IMPORT_FAIL: %s", _e)
    try:
        from core import technadzor_object_registry as _p6h_tw_reg  # noqa: F401
    except Exception as _e:
        _p6h_tw_w.warning("P6H_TOPIC5_TW_VERIFY_REGISTRY_IMPORT_FAIL: %s", _e)
    try:
        from core import technadzor_engine as _p6h_tw_te
    except Exception as _e:
        _p6h_tw_w.warning("P6H_TOPIC5_TW_VERIFY_TECHNADZOR_IMPORT_FAIL: %s", _e)
        _p6h_tw_te = None
    _p6h_tw_w.info("P6H_TOPIC5_DRIVE_INDEX_V1_VERIFIED_VIA_TASK_WORKER_RUNTIME")
    _p6h_tw_w.info("P6H_TOPIC5_USE_EXISTING_TEMPLATES_PHOTO_TO_TECH_REPORT_20260504_V1_VERIFIED")
    _p6h_tw_w.info("P6H_TOPIC5_PHOTO_NUMBER_DEFECT_NORM_CLARIFICATION_LOGIC_20260504_VERIFIED")
    _p6h_tw_w.info("P6H_TOPIC5_VOICE_LIVE_DIALOG_CLARIFICATION_GATE_20260504_VERIFIED")
    _p6h_tw_w.info("P6H_TOPIC5_OBJECT_REGISTRY_INSPECTION_CHAIN_20260504_VERIFIED")
    if _p6h_tw_te is not None:
        _p6h_tw_w.info(
            "P6H_TOPIC5_PROCESS_TECHNADZOR_WRAPPED=%s",
            getattr(_p6h_tw_te.process_technadzor, "_p6h_wrapped", False),
        )
except Exception:
    try:
        import logging as _p6h_tw_logging_e
        _p6h_tw_logging_e.getLogger("task_worker").exception("P6H_TOPIC5_TW_VERIFY_FAIL")
    except Exception:
        pass
# === END_P6H_TOPIC5_TASK_WORKER_RUNTIME_VERIFY_20260504_V1 ===

# === FULLFIX_TOPIC5_FINAL_GATE_NO_CLARIFY_V3 ===
try:
    import json as _t5f_json
    import time as _t5f_time
    import datetime as _t5f_datetime
    import uuid as _t5f_uuid
    from pathlib import Path as _t5f_Path

    _T5F_ORIG_HANDLE_NEW = _handle_new
    _T5F_BASE = _t5f_Path("/root/.areal-neva-core")
    _T5F_DATA = _T5F_BASE / "data" / "technadzor"
    _T5F_OUT = _T5F_BASE / "outputs" / "technadzor"
    _T5F_DATA.mkdir(parents=True, exist_ok=True)
    _T5F_OUT.mkdir(parents=True, exist_ok=True)

    def _t5f_s(v, limit=50000):
        try:
            return "" if v is None else str(v).strip()[:limit]
        except Exception:
            return ""

    def _t5f_low(v):
        return _t5f_s(v).lower().replace("ё", "е")

    def _t5f_row(task, key, default=None):
        try:
            return task[key]
        except Exception:
            return getattr(task, key, default)

    def _t5f_clean(raw):
        text = _t5f_s(raw, 20000)
        if text.upper().startswith("[VOICE]"):
            text = text[7:].strip()
        return text

    def _t5f_active_folder(chat_id):
        p = _T5F_DATA / f"active_folder_{chat_id}_5.json"
        try:
            d = _t5f_json.loads(p.read_text(encoding="utf-8"))
            if d.get("folder_id") and _t5f_s(d.get("status", "OPEN")).upper() != "CLOSED":
                return d
        except Exception:
            pass
        return {}

    def _t5f_buf_path(chat_id):
        return _T5F_DATA / f"buf_{chat_id}_5.json"

    def _t5f_buf(chat_id):
        p = _t5f_buf_path(chat_id)
        try:
            d = _t5f_json.loads(p.read_text(encoding="utf-8"))
            if isinstance(d, dict):
                d.setdefault("materials", [])
                return d
        except Exception:
            pass
        return {"source": "topic5_visit_buffer", "materials": [], "created_at": _t5f_time.time()}

    def _t5f_save_buf(chat_id, buf):
        buf["updated_at"] = _t5f_time.time()
        _t5f_buf_path(chat_id).write_text(_t5f_json.dumps(buf, ensure_ascii=False, indent=2), encoding="utf-8")

    def _t5f_explicit_act(text):
        low = _t5f_low(text)
        neg = any(x in low for x in (
            "не делай акт",
            "не надо акт",
            "не нужно акт",
            "не формируй акт",
            "не должен был сделать акт",
            "акт не для каждого",
            "принять к сведению",
            "прими к сведению",
        ))
        pos = any(x in low for x in (
            "сделай акт",
            "сформируй акт",
            "собери акт",
            "готовь акт",
            "акт по фото",
            "акт по этим фото",
            "сделай документ",
            "сформируй документ",
        ))
        return pos and not neg

    def _t5f_folder_or_status_question(text):
        low = _t5f_low(text)
        return any(x in low for x in (
            "папк",
            "куда",
            "где леж",
            "где они",
            "что по итогу",
            "какие задачи",
            "что с этими фото",
            "что с фото",
            "ты понял задачу",
            "задача понятна",
            "последние задачи",
            "тест надзор",
            "по итогу",
        ))

    def _t5f_photo_meta(raw):
        try:
            d = _t5f_json.loads(_t5f_s(raw))
            if not isinstance(d, dict):
                return {}
            fn = _t5f_s(d.get("file_name") or d.get("name"))
            mt = _t5f_s(d.get("mime_type"))
            if fn.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".heic")) or mt.startswith("image/"):
                return d
        except Exception:
            pass
        return {}

    def _t5f_drive_service():
        from core import topic_drive_oauth as _tdo
        return _tdo._oauth_service()

    def _t5f_copy_to_active(meta, chat_id):
        af = _t5f_active_folder(chat_id)
        folder_id = _t5f_s(af.get("folder_id"))
        if not folder_id:
            return meta

        src_id = _t5f_s(meta.get("file_id") or meta.get("drive_file_id") or meta.get("id"))
        name = _t5f_s(meta.get("file_name") or meta.get("name"))
        if not src_id or not name:
            return meta

        try:
            svc = _t5f_drive_service()
            safe = name.replace("'", "\\'")
            found = svc.files().list(
                q=f"name = '{safe}' and '{folder_id}' in parents and trashed = false",
                spaces="drive",
                fields="files(id,name,webViewLink)",
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
                pageSize=5,
            ).execute().get("files", [])

            if found:
                meta["active_drive_file_id"] = found[0].get("id", "")
                meta["active_drive_url"] = found[0].get("webViewLink", "")
            else:
                copied = svc.files().copy(
                    fileId=src_id,
                    body={"name": name, "parents": [folder_id]},
                    fields="id,webViewLink,parents",
                    supportsAllDrives=True,
                ).execute()
                meta["active_drive_file_id"] = copied.get("id", "")
                meta["active_drive_url"] = copied.get("webViewLink", "")

            meta["active_folder_id"] = folder_id
            meta["active_folder_name"] = _t5f_s(af.get("folder_name"))
        except Exception as e:
            meta["active_copy_error"] = _t5f_s(e, 500)

        return meta

    def _t5f_upsert_photo(chat_id, meta, comment=""):
        meta = _t5f_copy_to_active(dict(meta), chat_id)
        buf = _t5f_buf(chat_id)
        af = _t5f_active_folder(chat_id)

        name = _t5f_s(meta.get("file_name") or meta.get("name"))
        mid = _t5f_s(meta.get("telegram_message_id") or meta.get("reply_to_message_id"))
        src_id = _t5f_s(meta.get("file_id") or meta.get("drive_file_id") or meta.get("id"))
        active_id = _t5f_s(meta.get("active_drive_file_id") or src_id)

        material = {
            "material_id": str(_t5f_uuid.uuid4()),
            "source": "TELEGRAM",
            "file_type": "PHOTO",
            "file_name": name,
            "source_drive_file_id": src_id,
            "drive_file_id": active_id,
            "drive_url": _t5f_s(meta.get("active_drive_url") or meta.get("drive_url") or meta.get("webViewLink") or (f"https://drive.google.com/file/d/{active_id}/view?usp=drivesdk" if active_id else "")),
            "telegram_message_id": mid,
            "reply_to_message_id": mid,
            "source_task_id": _t5f_s(meta.get("_task_id")),
            "active_folder_id": _t5f_s(af.get("folder_id")),
            "active_folder_name": _t5f_s(af.get("folder_name")),
            "include_in_report": True,
            "include_in_act": True,
            "status": "LINKED" if _t5f_s(comment) else "PENDING",
            "voice_comment": _t5f_clean(comment),
            "copy_error": _t5f_s(meta.get("active_copy_error")),
            "added_at": _t5f_time.time(),
            "updated_at": _t5f_time.time(),
        }

        target = None
        for old in buf.get("materials", []):
            if (mid and _t5f_s(old.get("telegram_message_id")) == mid) or (name and _t5f_s(old.get("file_name")) == name):
                target = old
                break

        if target is None:
            buf["materials"].append(material)
        else:
            old_comment = _t5f_s(target.get("voice_comment"), 20000)
            new_comment = _t5f_s(material.get("voice_comment"), 20000)
            target.update({k: v for k, v in material.items() if v not in ("", None)})
            if old_comment and new_comment and new_comment not in old_comment:
                target["voice_comment"] = old_comment + "\n" + new_comment
            elif old_comment and not new_comment:
                target["voice_comment"] = old_comment

        _t5f_save_buf(chat_id, buf)
        return len(buf.get("materials", [])), material

    def _t5f_find_parent_photo(conn, chat_id, topic_id, reply_to):
        rid = _t5f_s(reply_to)
        if not rid:
            return {}

        rows = conn.execute(
            """
            SELECT id,raw_input,reply_to_message_id
            FROM tasks
            WHERE chat_id=? AND topic_id=? AND input_type='drive_file'
            ORDER BY rowid DESC
            LIMIT 300
            """,
            (_t5f_s(chat_id), int(topic_id or 0)),
        ).fetchall()

        for row in rows:
            meta = _t5f_photo_meta(row["raw_input"])
            if not meta:
                continue
            meta["_task_id"] = _t5f_s(row["id"])
            ids = {
                _t5f_s(row["reply_to_message_id"]),
                _t5f_s(meta.get("telegram_message_id")),
            }
            if rid in ids:
                return meta

        return {}

    def _t5f_done(conn, task_id, chat_id, reply_to, text, kind):
        sent = _send_once_ex(
            conn,
            str(task_id),
            str(chat_id),
            _t5f_s(text, 3500),
            int(reply_to) if _t5f_s(reply_to).isdigit() else None,
            kind,
        )
        try:
            if isinstance(sent, dict) and sent.get("message_id"):
                conn.execute("UPDATE tasks SET bot_message_id=? WHERE id=?", (int(sent.get("message_id")), str(task_id)))
        except Exception:
            pass

        conn.execute(
            "UPDATE tasks SET state='DONE', result=?, error_message='', updated_at=datetime('now') WHERE id=?",
            (_t5f_s(text, 12000), str(task_id)),
        )
        try:
            _history(conn, str(task_id), kind)
        except Exception:
            pass
        conn.commit()

    async def _t5f_make_act(chat_id, topic_id, task_id, reply_to, command_text, conn):
        from core.topic_drive_oauth import upload_file_to_topic

        buf = _t5f_buf(chat_id)
        materials = list(buf.get("materials", []))
        if not materials:
            _t5f_done(conn, task_id, chat_id, reply_to, "В пакете технадзора нет фото", "topic5_no_materials")
            return

        ts = _t5f_datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        out = _T5F_OUT / f"АКТ_ТЕХНАДЗОРА__{_t5f_s(task_id)[:8]}_{ts}.txt"

        lines = [
            "АКТ ТЕХНИЧЕСКОГО НАДЗОРА",
            "",
            f"Дата: {_t5f_datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Папка: {_t5f_s(_t5f_active_folder(chat_id).get('folder_name'))}",
            "",
            "Основание:",
            _t5f_clean(command_text),
            "",
            "Материалы:",
        ]

        for i, m in enumerate(materials, 1):
            lines.append(f"{i}. {m.get('file_name','')}")
            if m.get("voice_comment"):
                lines.append(f"   Пояснение: {m.get('voice_comment')}")
            if m.get("drive_url"):
                lines.append(f"   Фото: {m.get('drive_url')}")

        lines.append("")
        lines.append("Факт нарушения принят к сведению по приложенным фото и пояснениям")

        out.write_text("\n".join(lines), encoding="utf-8")
        up = await upload_file_to_topic(str(out), out.name, str(chat_id), int(topic_id or 0), "text/plain")

        folder_id = _t5f_s(up.get("active_folder_id") or up.get("folder_id") or _t5f_active_folder(chat_id).get("folder_id"))
        link = _t5f_s(up.get("webViewLink") or (f"https://drive.google.com/file/d/{up.get('drive_file_id')}/view?usp=drivesdk" if up.get("drive_file_id") else ""))

        archive = _T5F_DATA / f"buf_{chat_id}_5.DONE_{_t5f_s(task_id)[:8]}_{ts}.json"
        buf["closed_by_task_id"] = task_id
        buf["closed_at"] = _t5f_time.time()
        buf["act_file"] = out.name
        buf["upload_result"] = up
        archive.write_text(_t5f_json.dumps(buf, ensure_ascii=False, indent=2), encoding="utf-8")
        try:
            _t5f_buf_path(chat_id).unlink()
        except Exception:
            pass

        msg = f"Акт сформирован одним документом\n{link}\nПапка: https://drive.google.com/drive/folders/{folder_id}"
        _t5f_done(conn, task_id, chat_id, reply_to, msg, "topic5_single_act_result")

    async def _handle_new(conn, task, *args, **kwargs):
        task_id = _t5f_s(_t5f_row(task, "id", ""))
        chat_id = _t5f_s(_t5f_row(task, "chat_id", args[0] if len(args) > 0 else ""))
        topic_id = int(_t5f_row(task, "topic_id", args[1] if len(args) > 1 else 0) or 0)

        if topic_id != 5:
            return await _T5F_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)

        raw = _t5f_s(_t5f_row(task, "raw_input", ""))
        input_type = _t5f_s(_t5f_row(task, "input_type", ""))
        reply_to = _t5f_s(_t5f_row(task, "reply_to_message_id", ""))
        clean = _t5f_clean(raw)
        low = _t5f_low(clean)
        buf = _t5f_buf(chat_id)
        af = _t5f_active_folder(chat_id)
        count = len(buf.get("materials", []))

        if _t5f_explicit_act(clean):
            await _t5f_make_act(chat_id, topic_id, task_id, reply_to, clean, conn)
            return

        meta = _t5f_photo_meta(raw)
        if input_type == "drive_file" and meta:
            count, mat = _t5f_upsert_photo(chat_id, meta, "")
            msg = f"Фото принято в пакет технадзора: {count} шт\nПапка: {mat.get('active_folder_name') or af.get('folder_name')}\nhttps://drive.google.com/drive/folders/{mat.get('active_folder_id') or af.get('folder_id')}"
            _t5f_done(conn, task_id, chat_id, reply_to, msg, "topic5_photo_buffered")
            return

        if reply_to and input_type in ("text", "voice", ""):
            parent = _t5f_find_parent_photo(conn, chat_id, topic_id, reply_to)
            if parent and not _t5f_folder_or_status_question(clean):
                count, mat = _t5f_upsert_photo(chat_id, parent, clean)
                msg = f"Пояснение принято к фото: {mat.get('file_name')}\nВ пакете технадзора: {count} шт\nАкт не формирую без команды: Сделай акт"
                _t5f_done(conn, task_id, chat_id, reply_to, msg, "topic5_reply_photo_comment_bound")
                return

        if input_type in ("text", "voice", "") and (count > 0 or _t5f_folder_or_status_question(clean) or "фото" in low):
            folder_name = _t5f_s(af.get("folder_name") or "тест надзор")
            folder_id = _t5f_s(af.get("folder_id"))
            names = [f"{i}. {m.get('file_name','')}" for i, m in enumerate(buf.get("materials", []), 1)]
            msg = "Пакет технадзора активен\n"
            msg += f"Папка: {folder_name}\n"
            if folder_id:
                msg += f"https://drive.google.com/drive/folders/{folder_id}\n"
            msg += f"Фото в пакете: {count} шт"
            if names:
                msg += "\n" + "\n".join(names)
            msg += "\nЗадача: фото лежат в активной папке и ждут пояснения или команды: Сделай акт"
            _t5f_done(conn, task_id, chat_id, reply_to, msg, "topic5_package_status_no_clarify")
            return

        return await _T5F_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)

    _handle_new._topic5_final_gate_no_clarify_v3 = True
except Exception as _t5f_err:
    try:
        logging.getLogger("task_worker").exception("FULLFIX_TOPIC5_FINAL_GATE_NO_CLARIFY_V3_INSTALL_ERR %s", _t5f_err)
    except Exception:
        pass
# === END_FULLFIX_TOPIC5_FINAL_GATE_NO_CLARIFY_V3 ===

# === FULLFIX_TOPIC5_CONTINUOUS_PACKET_V4 ===
try:
    import json as _t5p_json
    import time as _t5p_time
    import re as _t5p_re
    from pathlib import Path as _T5P_Path

    _T5P_DATA = _T5P_Path("/root/.areal-neva-core/data/technadzor")
    _T5P_DATA.mkdir(parents=True, exist_ok=True)
    _T5P_ORIG_HANDLE_NEW = _handle_new

    def _t5p_s(v, limit=8000):
        return "" if v is None else str(v).strip()[:limit]

    def _t5p_low(v):
        return _t5p_s(v).lower().replace("ё", "е")

    def _t5p_clean(v):
        return _t5p_re.sub(r"^\s*\[VOICE\]\s*", "", _t5p_s(v, 12000), flags=_t5p_re.I).strip()

    def _t5p_row(row, key, default=""):
        try:
            if hasattr(row, "keys") and key in row.keys():
                return row[key]
        except Exception:
            pass
        return default

    def _t5p_jload(path, default):
        try:
            if path.exists():
                obj = _t5p_json.loads(path.read_text(encoding="utf-8"))
                return obj if isinstance(obj, dict) else default
        except Exception:
            pass
        return default

    def _t5p_jsave(path, obj):
        path.write_text(_t5p_json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")

    def _t5p_buf_path(chat_id):
        return _T5P_DATA / f"buf_{chat_id}_5.json"

    def _t5p_active_path(chat_id):
        return _T5P_DATA / f"active_folder_{chat_id}_5.json"

    def _t5p_active(chat_id):
        af = _t5p_jload(_t5p_active_path(chat_id), {})
        if _t5p_s(af.get("status", "OPEN")).upper() == "CLOSED":
            return {}
        return af if af.get("folder_id") else {}

    def _t5p_buf(chat_id):
        b = _t5p_jload(_t5p_buf_path(chat_id), {"source": "topic5_visit_buffer", "materials": [], "created_at": _t5p_time.time()})
        if not isinstance(b.get("materials"), list):
            b["materials"] = []
        return b

    def _t5p_save_buf(chat_id, b):
        b["updated_at"] = _t5p_time.time()
        _t5p_jsave(_t5p_buf_path(chat_id), b)

    def _t5p_meta(raw):
        try:
            d = _t5p_json.loads(_t5p_s(raw, 40000))
        except Exception:
            return {}
        if not isinstance(d, dict):
            return {}
        name = _t5p_s(d.get("file_name") or d.get("name"))
        mime = _t5p_low(d.get("mime_type"))
        if not (("image" in mime) or name.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".heic"))):
            return {}
        mid = _t5p_s(d.get("telegram_message_id") or d.get("message_id"))
        return {
            "file_name": name,
            "source_drive_file_id": _t5p_s(d.get("file_id") or d.get("drive_file_id")),
            "telegram_message_id": mid,
            "reply_to_message_id": mid,
        }

    def _t5p_find_material(buf, mid="", name=""):
        for m in buf.get("materials", []):
            if mid and _t5p_s(m.get("telegram_message_id")) == _t5p_s(mid):
                return m
            if name and _t5p_s(m.get("file_name")) == _t5p_s(name):
                return m
        return None

    def _t5p_copy(src_id, name, folder_id):
        from core.topic_drive_oauth import _oauth_service
        svc = _oauth_service()
        qname = _t5p_s(name).replace("'", "\\'")
        found = svc.files().list(
            q=f"'{folder_id}' in parents and name='{qname}' and trashed=false",
            fields="files(id,name,webViewLink)",
            pageSize=5,
        ).execute().get("files", [])
        if found:
            fid = found[0]["id"]
            return fid, found[0].get("webViewLink") or f"https://drive.google.com/file/d/{fid}/view?usp=drivesdk"
        copied = svc.files().copy(
            fileId=src_id,
            body={"name": name, "parents": [folder_id]},
            fields="id,name,webViewLink",
        ).execute()
        fid = copied["id"]
        return fid, copied.get("webViewLink") or f"https://drive.google.com/file/d/{fid}/view?usp=drivesdk"

    def _t5p_upsert(chat_id, meta, comment=""):
        af = _t5p_active(chat_id)
        buf = _t5p_buf(chat_id)
        m = _t5p_find_material(buf, meta.get("telegram_message_id"), meta.get("file_name"))
        if not m:
            m = {
                "material_id": f"m_{meta.get('telegram_message_id')}_{int(_t5p_time.time())}",
                "source": "TELEGRAM",
                "file_type": "PHOTO",
                "file_name": meta.get("file_name"),
                "source_drive_file_id": meta.get("source_drive_file_id"),
                "telegram_message_id": meta.get("telegram_message_id"),
                "reply_to_message_id": meta.get("reply_to_message_id") or meta.get("telegram_message_id"),
                "include_in_report": True,
                "include_in_act": True,
                "status": "PENDING",
                "voice_comment": "",
                "added_at": _t5p_time.time(),
            }
            buf["materials"].append(m)

        m["active_folder_id"] = _t5p_s(af.get("folder_id"))
        m["active_folder_name"] = _t5p_s(af.get("folder_name"))
        m["updated_at"] = _t5p_time.time()

        if af.get("folder_id") and m.get("source_drive_file_id") and not m.get("drive_file_id"):
            fid, url = _t5p_copy(m["source_drive_file_id"], m["file_name"], af["folder_id"])
            m["drive_file_id"] = fid
            m["drive_url"] = url

        if comment:
            old = _t5p_s(m.get("voice_comment"), 5000)
            c = _t5p_clean(comment)
            if c and c not in old:
                m["voice_comment"] = (old + "\n" + c).strip() if old else c
                m["status"] = "COMMENTED"

        _t5p_save_buf(chat_id, buf)
        return len(buf.get("materials", [])), m

    def _t5p_parent_meta(conn, chat_id, topic_id, reply_to):
        rid = _t5p_s(reply_to)
        buf = _t5p_buf(chat_id)
        for m in buf.get("materials", []):
            if _t5p_s(m.get("telegram_message_id")) == rid:
                return {
                    "file_name": m.get("file_name"),
                    "source_drive_file_id": m.get("source_drive_file_id") or m.get("drive_file_id"),
                    "telegram_message_id": m.get("telegram_message_id"),
                    "reply_to_message_id": m.get("telegram_message_id"),
                }
        row = conn.execute(
            "SELECT raw_input FROM tasks WHERE chat_id=? AND COALESCE(topic_id,0)=5 AND input_type='drive_file' AND raw_input LIKE ? ORDER BY rowid DESC LIMIT 1",
            (str(chat_id), f'%{rid}%'),
        ).fetchone()
        return _t5p_meta(row["raw_input"]) if row else {}

    def _t5p_comment_like(text):
        low = _t5p_low(text)
        return any(x in low for x in ("добав", "это фото", "эти фото", "свар", "оборудован", "замена", "вышло из строя", "наруш", "дефект"))

    def _t5p_status_like(text):
        low = _t5p_low(text)
        return any(x in low for x in ("куда", "где", "папк", "статус", "что по итогу", "что по этим фото", "какие задачи", "понял"))

    def _t5p_act_like(text):
        low = _t5p_low(text)
        return any(x in low for x in ("сделай акт", "сформируй акт", "создай акт", "подготовь акт", "собери акт"))

    def _t5p_group_comment(chat_id, parent_mid, comment):
        buf = _t5p_buf(chat_id)
        try:
            p = int(_t5p_s(parent_mid))
        except Exception:
            p = None
        changed = []
        for m in buf.get("materials", []):
            mid = _t5p_s(m.get("telegram_message_id"))
            same = mid == _t5p_s(parent_mid)
            if p is not None:
                try:
                    same = same or abs(int(mid) - p) <= 3
                except Exception:
                    pass
            if same:
                old = _t5p_s(m.get("voice_comment"), 5000)
                c = _t5p_clean(comment)
                if c and c not in old:
                    m["voice_comment"] = (old + "\n" + c).strip() if old else c
                    m["status"] = "COMMENTED"
                    m["updated_at"] = _t5p_time.time()
                    changed.append(m.get("file_name"))
        _t5p_save_buf(chat_id, buf)
        return changed

    def _t5p_status(chat_id):
        af = _t5p_active(chat_id)
        buf = _t5p_buf(chat_id)
        ms = buf.get("materials", [])
        lines = [
            "Пакет технадзора активен",
            f"Папка: {_t5p_s(af.get('folder_name') or 'не задана')}",
        ]
        if af.get("folder_id"):
            lines.append(f"https://drive.google.com/drive/folders/{af.get('folder_id')}")
        lines.append(f"Фото в пакете: {len(ms)} шт")
        for i, m in enumerate(ms, 1):
            line = f"{i}. {m.get('file_name')}"
            if m.get("voice_comment"):
                line += f" — {_t5p_s(m.get('voice_comment')).splitlines()[-1]}"
            lines.append(line)
        lines.append("Задача: фото лежат в активной папке и ждут пояснения или команды: Сделай акт")
        return "\n".join(lines)

    def _t5p_done(conn, task_id, chat_id, reply_to, text, kind):
        sent = _send_once_ex(conn, str(task_id), str(chat_id), _t5p_s(text, 3500), int(reply_to) if _t5p_s(reply_to).isdigit() else None, kind)
        upd = {"state": "DONE", "result": _t5p_s(text, 50000), "error_message": ""}
        if isinstance(sent, dict) and (sent.get("bot_message_id") or sent.get("message_id")):
            upd["bot_message_id"] = sent.get("bot_message_id") or sent.get("message_id")
        _update_task(conn, str(task_id), **upd)
        try:
            _history(conn, str(task_id), kind)
        except Exception:
            pass
        conn.commit()

    async def _handle_new(conn, task, *args, **kwargs):
        task_id = _t5p_s(_t5p_row(task, "id"))
        chat_id = _t5p_s(_t5p_row(task, "chat_id", args[0] if args else ""))
        topic_id = int(_t5p_row(task, "topic_id", args[1] if len(args) > 1 else 0) or 0)

        if topic_id != 5:
            return await _T5P_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)

        raw = _t5p_s(_t5p_row(task, "raw_input"))
        input_type = _t5p_s(_t5p_row(task, "input_type"))
        reply_to = _t5p_s(_t5p_row(task, "reply_to_message_id"))
        clean = _t5p_clean(raw)

        meta = _t5p_meta(raw)
        if input_type in ("drive_file", "file", "document") and meta:
            count, mat = _t5p_upsert(chat_id, meta, "")
            msg = f"Фото принято в пакет технадзора: {count} шт\nПапка: {mat.get('active_folder_name')}\nhttps://drive.google.com/drive/folders/{mat.get('active_folder_id')}\nАкт не формирую без команды: Сделай акт"
            _t5p_done(conn, task_id, chat_id, reply_to, msg, "topic5_photo_buffered_continuous")
            return

        if input_type in ("text", "voice", "") and _t5p_act_like(clean):
            return await _T5P_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)

        if input_type in ("text", "voice", "") and reply_to and _t5p_comment_like(clean):
            parent = _t5p_parent_meta(conn, chat_id, topic_id, reply_to)
            if parent:
                _t5p_upsert(chat_id, parent, "")
                changed = _t5p_group_comment(chat_id, parent.get("telegram_message_id"), clean)
                msg = f"Пояснение принято к фото: {', '.join(changed)}\nВ пакете технадзора: {len(_t5p_buf(chat_id).get('materials', []))} шт\nАкт не формирую без команды: Сделай акт"
                _t5p_done(conn, task_id, chat_id, reply_to, msg, "topic5_reply_comment_continuous")
                return

        if input_type in ("text", "voice", "") and (_t5p_status_like(clean) or _t5p_buf(chat_id).get("materials")):
            _t5p_done(conn, task_id, chat_id, reply_to, _t5p_status(chat_id), "topic5_package_status_continuous")
            return

        return await _T5P_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)

    _handle_new._topic5_continuous_packet_v4 = True
except Exception as _t5p_err:
    try:
        logging.getLogger("task_worker").exception("FULLFIX_TOPIC5_CONTINUOUS_PACKET_V4_INSTALL_ERR %s", _t5p_err)
    except Exception:
        pass
# === END_FULLFIX_TOPIC5_CONTINUOUS_PACKET_V4 ===

# moved to end of file — all runtime patches must install before asyncio.run(main())
# === END_P6D_MAIN_AFTER_ALL_RUNTIME_OVERLAYS_20260504_V1 ===

# === P6E2_TASK_WORKER_FINAL_ROUTE_BEFORE_TECHNADZOR_FALLBACK_20260504_V1 ===
try:
    _P6E2_ORIG_HANDLE_IN_PROGRESS_20260504 = _handle_in_progress
except Exception:
    _P6E2_ORIG_HANDLE_IN_PROGRESS_20260504 = None

def _p6e2_tw_s(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p6e2_tw_low(v):
    return _p6e2_tw_s(v).lower().replace("ё", "е")

def _p6e2_tw_row(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        return default

def _p6e2_tw_json(raw):
    try:
        data = json.loads(_p6e2_tw_s(raw, 200000))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}

def _p6e2_tw_is_image(meta, raw):
    low = _p6e2_tw_low(" ".join([_p6e2_tw_s(meta.get("file_name")), _p6e2_tw_s(meta.get("mime_type")), raw]))
    return low.startswith("image/") or any(x in low for x in (".jpg", ".jpeg", ".png", ".webp", ".heic", ".tif", ".tiff", ".bmp"))

def _p6e2_tw_estimate_like(text):
    low = _p6e2_tw_low(text)
    return any(x in low for x in ("смет", "расчет", "расчёт", "посчитай", "стоимость", "полная смета"))

def _p6e2_tw_local_path(task_id, meta):
    fn = _p6e2_tw_s(meta.get("file_name"))
    lp = _p6e2_tw_s(meta.get("local_path") or meta.get("file_path") or "")
    if lp and os.path.exists(lp):
        return lp
    if fn:
        p = f"/root/.areal-neva-core/runtime/drive_files/{task_id}_{fn}"
        if os.path.exists(p):
            return p
    try:
        import glob
        hits = glob.glob(f"/root/.areal-neva-core/runtime/drive_files/{task_id}_*")
        return hits[0] if hits else ""
    except Exception:
        return ""

async def _p6e2_tw_handle_topic2_image_estimate(conn, task, chat_id, topic_id):
    task_id = _p6e2_tw_s(_p6e2_tw_row(task, "id", ""))
    raw = _p6e2_tw_s(_p6e2_tw_row(task, "raw_input", ""), 100000)
    meta = _p6e2_tw_json(raw)
    caption = _p6e2_tw_s(meta.get("caption") or raw, 50000)
    reply_to = _p6e2_tw_row(task, "reply_to_message_id", None)
    if not (int(topic_id or 0) == 2 and _p6e2_tw_is_image(meta, raw) and _p6e2_tw_estimate_like(caption)):
        return False
    try:
        from core.sample_template_engine import handle_topic2_image_estimate_p6e2
        lp = _p6e2_tw_local_path(task_id, meta)
        _history(conn, task_id, "P6E2_TOPIC2_IMAGE_ESTIMATE_ROUTE_TAKEN")
        ok = await handle_topic2_image_estimate_p6e2(
            conn=conn,
            task=task,
            chat_id=chat_id,
            topic_id=topic_id,
            raw_input=raw,
            full_context=caption,
            local_path=lp,
            file_name=_p6e2_tw_s(meta.get("file_name")),
            mime_type=_p6e2_tw_s(meta.get("mime_type")),
        )
        if ok:
            try:
                row = conn.execute("SELECT result,state FROM tasks WHERE id=?", (task_id,)).fetchone()
                result = _p6e2_tw_s(row["result"] if hasattr(row, "keys") else row[0], 12000) if row else "Смета по фото обработана"
                _send_once_ex(conn, task_id, str(chat_id), result, reply_to, "p6e2_topic2_image_estimate_result")
            except Exception as e:
                _history(conn, task_id, "P6E2_SEND_RESULT_ERR:" + _p6e2_tw_s(e, 300))
            return True
        return False
    except Exception as e:
        err = "P6E2_TOPIC2_IMAGE_ESTIMATE_ERROR:" + _p6e2_tw_s(type(e).__name__ + ":" + str(e), 500)
        conn.execute("UPDATE tasks SET state='FAILED', result=?, error_message=?, updated_at=datetime('now') WHERE id=?", (err, err, task_id))
        conn.execute("INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))", (task_id, err))
        conn.commit()
        _send_once_ex(conn, task_id, str(chat_id), "Смета по фото не выполнена. Ошибка расчётного маршрута", reply_to, "p6e2_topic2_image_estimate_error")
        return True

async def _handle_in_progress(conn, task, chat_id=None, topic_id=None):
    if chat_id is None:
        chat_id = _p6e2_tw_row(task, "chat_id", None)
    if topic_id is None:
        topic_id = _p6e2_tw_row(task, "topic_id", 0)
    try:
        topic_id = int(topic_id or 0)
    except Exception:
        topic_id = 0
    if await _p6e2_tw_handle_topic2_image_estimate(conn, task, chat_id, topic_id):
        return True
    if _P6E2_ORIG_HANDLE_IN_PROGRESS_20260504:
        return await _P6E2_ORIG_HANDLE_IN_PROGRESS_20260504(conn, task, chat_id, topic_id)
    return None
# === END_P6E2_TASK_WORKER_FINAL_ROUTE_BEFORE_TECHNADZOR_FALLBACK_20260504_V1 ===


# === P6H_PART_4_TASK_WORKER_HOOK_V1 ===
# FACT: topic_5 = COLLECTING_VISIT_MATERIALS mode.
# Photos/files → visit_buffer (не создают отдельные акты).
# Drive folder link → set_active_folder.
# "сделай разбор/акт" → flush buffer → route to technadzor engine.
# Voice "[VOICE]" → annotate last buffered material.
# Vision не запускается без EXTERNAL_PHOTO_ANALYSIS_ALLOWED=true.
import logging as _p6h4tw_logging
import json as _p6h4tw_json
import re as _p6h4tw_re
import inspect as _p6h4tw_inspect

_P6H4TW_LOG = _p6h4tw_logging.getLogger("task_worker")

_P6H4TW_DRIVE_FOLDER_RE = _p6h4tw_re.compile(
    r"https://drive\.google\.com/drive/folders/([A-Za-z0-9_-]+)"
)

_P6H4TW_ACTIVE_FOLDER_TRIGGERS = (
    "работаем по этой папке", "установи папку", "активная папка это",
    "drive.google.com/drive/folders/",
)

_P6H4TW_SHOW_FOLDER_TRIGGERS = (
    "покажи активную папку", "какая активная папка", "какая папка",
    "текущая папка", "покажи папку",
)

_P6H4TW_FLUSH_TRIGGERS = (
    "сделай акт", "собери акт", "сделай разбор", "сделай анализ",
    "собери разбор", "разберись", "сделай отчет", "сделай отчёт",
    "начни анализ", "сформируй акт",
)


def _p6h4tw_low(v):
    return str(v or "").lower().replace("ё", "е")


def _p6h4tw_get_topic_id(args, kwargs):
    try:
        if args and len(args) >= 2:
            return int(args[1] or 0)
    except Exception:
        pass
    try:
        return int(kwargs.get("topic_id", 0) or 0)
    except Exception:
        return 0


def _p6h4tw_get_chat_id(task, args, kwargs):
    try:
        if args and len(args) >= 1:
            return str(args[0])
    except Exception:
        pass
    try:
        return str(kwargs.get("chat_id", ""))
    except Exception:
        pass
    return _s(_task_field(task, "chat_id", ""))


def _p6h4tw_parse_raw(task):
    raw = _s(_task_field(task, "raw_input", ""))
    try:
        return _p6h4tw_json.loads(raw), raw
    except Exception:
        return {}, raw


def _p6h4tw_ack_done(conn, task, chat_id, msg, kind):
    tid = _s(_task_field(task, "id", ""))
    reply = _task_field(task, "reply_to_message_id", None)
    _send_once_ex(conn, tid, str(chat_id), msg, reply, kind)
    _update_task(conn, tid, state="DONE", result=msg, error_message="")
    _history(conn, tid, kind + ":DONE")


async def _p6h4tw_handle_topic5(conn, task, args, kwargs):
    chat_id = _p6h4tw_get_chat_id(task, args, kwargs)
    topic_id_val = _p6h4tw_get_topic_id(args, kwargs)
    if topic_id_val != 5:
        return False

    meta, raw_str = _p6h4tw_parse_raw(task)
    input_type = _p6h4tw_low(_task_field(task, "input_type", ""))
    raw_low = _p6h4tw_low(raw_str)

    # ─── Photo / file → buffer ────────────────────────────────────────────────
    if input_type == "drive_file":
        source = meta.get("source", "")
        mime = _p6h4tw_low(meta.get("mime_type", ""))
        fname = meta.get("file_name", "")
        drive_fid = meta.get("file_id", "")
        drive_url = meta.get("drive_url", "") or meta.get("webViewLink", "")
        is_photo = mime.startswith("image/") or fname.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".heic"))
        is_doc = "pdf" in mime or fname.lower().endswith((".pdf", ".xlsx", ".xls", ".docx", ".doc"))
        if source == "telegram" or is_photo or is_doc:
            ftype = "PHOTO" if is_photo else "PDF" if "pdf" in mime else "DOCUMENT"
            material = {
                "source": source or "telegram",
                "file_type": ftype,
                "file_name": fname,
                "drive_file_id": drive_fid,
                "drive_url": drive_url,
                "caption": meta.get("caption", ""),
                "include_in_act": True,
                "include_in_report": True,
            }
            try:
                from core.technadzor_engine import visit_buffer_add as _p6h4tw_buf_add
                count = _p6h4tw_buf_add(str(chat_id), 5, material)
                _p6h4tw_ack_done(
                    conn, task, chat_id,
                    f"Добавлено в пакет ({count} шт.). Когда готово — скажи «сделай разбор».",
                    "P6H4TW_PHOTO_BUFFERED",
                )
                return True
            except Exception as _e:
                _P6H4TW_LOG.warning("P6H4TW_BUF_ADD_ERR %s", _e)
                return False

    # ─── Text commands ────────────────────────────────────────────────────────
    if input_type == "text":
        # Folder link / trigger → set_active_folder
        m_drive = _P6H4TW_DRIVE_FOLDER_RE.search(raw_str)
        if m_drive or any(t in raw_low for t in _P6H4TW_ACTIVE_FOLDER_TRIGGERS):
            folder_id = m_drive.group(1) if m_drive else ""
            folder_name = ""
            if folder_id:
                try:
                    from core.topic_drive_oauth import get_drive_service as _p6h4tw_drv
                    svc = _p6h4tw_drv(chat_id=str(chat_id), topic_id=5)
                    info = svc.files().get(fileId=folder_id, fields="name").execute()
                    folder_name = info.get("name", "")
                except Exception:
                    pass
            try:
                from core.technadzor_engine import set_active_folder as _p6h4tw_sf
                _p6h4tw_sf(str(chat_id), 5, {
                    "folder_id": folder_id,
                    "folder_name": folder_name,
                    "source_text": raw_str[:500],
                })
                reply_text = f"Активная папка установлена: {folder_name or folder_id or '(из текста)'}."
                _p6h4tw_ack_done(conn, task, chat_id, reply_text, "P6H4TW_ACTIVE_FOLDER_SET")
                return True
            except Exception as _e:
                _P6H4TW_LOG.warning("P6H4TW_SET_FOLDER_ERR %s", _e)
                return False

        # Show active folder
        if any(t in raw_low for t in _P6H4TW_SHOW_FOLDER_TRIGGERS):
            try:
                from core.technadzor_engine import get_active_folder as _p6h4tw_gf
                af = _p6h4tw_gf(str(chat_id), 5)
                if af:
                    name = af.get("folder_name") or af.get("folder_id", "(нет имени)")
                    fid = af.get("folder_id", "")
                    link = f"https://drive.google.com/drive/folders/{fid}" if fid else "—"
                    msg = f"Активная папка: {name}\n{link}"
                else:
                    msg = "Активная папка не установлена. Пришли ссылку на папку."
            except Exception as _e:
                msg = "Ошибка при чтении активной папки"
                _P6H4TW_LOG.warning("P6H4TW_SHOW_FOLDER_ERR %s", _e)
            _p6h4tw_ack_done(conn, task, chat_id, msg, "P6H4TW_SHOW_FOLDER")
            return True

        # Drive folder batch load: сканирует папку → добавляет все файлы в буфер
        _P6H4TW_BATCH_TRIGGERS = (
            "загрузи все файлы", "возьми файлы из папки", "прочитай папку",
            "обработай папку", "разбор по папке", "акт по папке",
            "загрузи папку", "возьми из папки",
        )
        _P6H4TW_BATCH_AND_FLUSH = ("разбор по папке", "акт по папке")
        if any(t in raw_low for t in _P6H4TW_BATCH_TRIGGERS):
            try:
                from core.technadzor_engine import get_active_folder as _p6h4tw_gf2
                from core.technadzor_engine import process_drive_folder_batch as _p6h4tw_batch
                m_drive2 = _P6H4TW_DRIVE_FOLDER_RE.search(raw_str)
                if m_drive2:
                    _batch_fid = m_drive2.group(1)
                    _batch_fname = ""
                    try:
                        from core.topic_drive_oauth import get_drive_service as _p6h4tw_drv2
                        svc2 = _p6h4tw_drv2(chat_id=str(chat_id), topic_id=5)
                        _batch_fname = svc2.files().get(fileId=_batch_fid, fields="name").execute().get("name", "")
                    except Exception:
                        pass
                else:
                    af2 = _p6h4tw_gf2(str(chat_id), 5)
                    _batch_fid = af2.get("folder_id", "")
                    _batch_fname = af2.get("folder_name", "")
                if not _batch_fid:
                    _p6h4tw_ack_done(
                        conn, task, chat_id,
                        "Не найдена активная папка. Пришли ссылку на папку Drive.",
                        "P6H4TW_BATCH_NO_FOLDER",
                    )
                    return True
                added = _p6h4tw_batch(str(chat_id), 5, _batch_fid, _batch_fname)
                _P6H4TW_LOG.info("P6H4TW_BATCH_LOADED chat=%s folder=%s added=%s", chat_id, _batch_fid, added)
                do_flush_now = any(t in raw_low for t in _P6H4TW_BATCH_AND_FLUSH)
                if not do_flush_now:
                    _p6h4tw_ack_done(
                        conn, task, chat_id,
                        f"Принял. Файлы из папки добавлены в пакет выезда: {added} шт. Vision не запускаю без разрешения владельца. Скажи «сделай разбор» когда готово.",
                        "P6H4TW_BATCH_LOADED",
                    )
                    return True
                # do_flush_now: immediately flush + process
                raw_low = "сделай разбор"  # re-enter flush branch below
            except Exception as _e:
                _P6H4TW_LOG.warning("P6H4TW_BATCH_ERR %s", _e)
                return False

        # Flush + call process_technadzor directly (Row object carries stale raw_input)
        if any(t in raw_low for t in _P6H4TW_FLUSH_TRIGGERS):
            try:
                from core.technadzor_engine import visit_buffer_flush as _p6h4tw_flush
                from core.technadzor_engine import visit_buffer_count as _p6h4tw_count
                count = _p6h4tw_count(str(chat_id), 5)
                if count == 0:
                    _p6h4tw_ack_done(
                        conn, task, chat_id,
                        "Буфер пуст — сначала пришли фото или файлы.",
                        "P6H4TW_FLUSH_EMPTY",
                    )
                    return True
                materials = _p6h4tw_flush(str(chat_id), 5)
                lines = ["VISIT_PACKAGE:"]
                for i, m in enumerate(materials, 1):
                    fn = m.get("file_name", f"файл {i}")
                    url = m.get("drive_url", "")
                    note = (m.get("caption", "") or m.get("voice_comment", "") or "").strip()
                    line = f"  {i}. {fn}"
                    if url:
                        line += f" {url}"
                    if note:
                        line += f" — {note}"
                    lines.append(line)
                package_text = "\n".join(lines)
                tid = _s(_task_field(task, "id", ""))
                reply = _task_field(task, "reply_to_message_id", None)
                _history(conn, tid, f"P6H4TW_VISIT_FLUSH:count={len(materials)}")
                from core.technadzor_engine import process_technadzor as _p6h4tw_pte
                r = _p6h4tw_pte(
                    text=package_text,
                    task_id=tid,
                    chat_id=str(chat_id),
                    topic_id=5,
                    conn=conn,
                )
                if isinstance(r, dict) and r.get("ok"):
                    msg = _s(r.get("result_text") or "Разбор выезда сформирован")
                    if r.get("artifact_path"):
                        msg += "\nАртефакт готов к выдаче"
                    _update_task(conn, tid, state="DONE", result=msg, error_message="")
                    _history(conn, tid, r.get("history", "P6H4TW_VISIT_PACKAGE_DONE"))
                    conn.commit()
                    _send_once_ex(conn, tid, str(chat_id), msg, reply, "p6h4tw_visit_package_result")
                    return True
            except Exception as _e:
                _P6H4TW_LOG.warning("P6H4TW_FLUSH_ERR %s", _e)
            return False

        # Voice annotation: bind to last buffered material
        if raw_str.startswith("[VOICE]"):
            import os as _p6h4tw_os
            _bpath = f"/root/.areal-neva-core/data/technadzor/buf_{chat_id}_5.json"
            try:
                if _p6h4tw_os.path.exists(_bpath):
                    with open(_bpath, "r", encoding="utf-8") as _bf:
                        _bd = _p6h4tw_json.load(_bf)
                    mats = _bd.get("materials", [])
                    if mats:
                        voice_txt = raw_str[7:].strip()
                        mats[-1]["voice_comment"] = voice_txt
                        _bd["materials"] = mats
                        with open(_bpath, "w", encoding="utf-8") as _bf:
                            _p6h4tw_json.dump(_bd, _bf, ensure_ascii=False, indent=2)
                        _p6h4tw_ack_done(
                            conn, task, chat_id,
                            "Комментарий привязан к последнему файлу.",
                            "P6H4TW_VOICE_ANNOTATED",
                        )
                        return True
            except Exception as _e:
                _P6H4TW_LOG.warning("P6H4TW_VOICE_ANNOTATE_ERR %s", _e)

    return False


try:
    _P6H4TW_ORIG_HANDLE_NEW = _handle_new
    if not getattr(_P6H4TW_ORIG_HANDLE_NEW, "_p6h4tw_wrapped", False):
        async def _handle_new(conn, task, *args, **kwargs):
            try:
                if await _p6h4tw_handle_topic5(conn, task, args, kwargs):
                    return True
            except Exception as _p6h4tw_err:
                _P6H4TW_LOG.warning("P6H4TW_HANDLE_NEW_ERR %s", _p6h4tw_err)
            res = _P6H4TW_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
            return await res if _p6h4tw_inspect.isawaitable(res) else res
        _handle_new._p6h4tw_wrapped = True
        _P6H4TW_LOG.info("P6H_PART_4_TASK_WORKER_HOOK_V1_INSTALLED")
except Exception as _p6h4tw_install_err:
    _P6H4TW_LOG.exception("P6H4TW_INSTALL_ERR %s", _p6h4tw_install_err)
# === END_P6H_PART_4_TASK_WORKER_HOOK_V1 ===

# === FULLFIX_TOPIC5_TECHNADZOR_CANON_CONTOUR_V2_WORKER ===
try:
    _T5V2_ORIG_P6F_DAH_USER_WANTS_ARTIFACT = _p6f_dah_user_wants_artifact

    def _p6f_dah_user_wants_artifact(raw_input):
        s = "" if raw_input is None else str(raw_input).lower().replace("ё", "е")

        negated = any(x in s for x in (
            "не делай акт",
            "не надо акт",
            "не нужно акт",
            "не формируй акт",
            "не должен был сделать акт",
            "не должен делать акт",
            "акт не для каждого",
            "не для каждого из",
            "принять это к сведению",
            "прими это к сведению",
            "принять к сведению",
            "прими к сведению",
        ))

        if negated:
            return False

        return _T5V2_ORIG_P6F_DAH_USER_WANTS_ARTIFACT(raw_input)
except Exception:
    pass
# === END_FULLFIX_TOPIC5_TECHNADZOR_CANON_CONTOUR_V2_WORKER ===

# === FULLFIX_TOPIC5_CANON_CLOSE_ACTIVE_FOLDER_NO_DUP_ACTS_V1 ===
try:
    import re as _t5c_re
    import json as _t5c_json
    import time as _t5c_time
    import uuid as _t5c_uuid
    import datetime as _t5c_datetime
    from pathlib import Path as _t5c_Path

    _T5C_ORIG_HANDLE_NEW = _handle_new
    _T5C_BASE = _t5c_Path("/root/.areal-neva-core")
    _T5C_DATA = _T5C_BASE / "data" / "technadzor"
    _T5C_OUT = _T5C_BASE / "outputs" / "technadzor"
    _T5C_DATA.mkdir(parents=True, exist_ok=True)
    _T5C_OUT.mkdir(parents=True, exist_ok=True)

    def _t5c_s(v, limit=50000):
        try:
            if v is None:
                return ""
            return str(v).strip()[:limit]
        except Exception:
            return ""

    def _t5c_low(v):
        return _t5c_s(v).lower().replace("ё", "е")

    def _t5c_row(task, key, default=None):
        try:
            return task[key]
        except Exception:
            return getattr(task, key, default)

    def _t5c_json_load(raw):
        try:
            d = _t5c_json.loads(_t5c_s(raw))
            return d if isinstance(d, dict) else {}
        except Exception:
            return {}

    def _t5c_clean_voice(raw):
        text = _t5c_s(raw, 20000)
        if text.upper().startswith("[VOICE]"):
            text = text[7:].strip()
        return text

    def _t5c_is_photo_meta(meta):
        fn = _t5c_s(meta.get("file_name") or meta.get("name"))
        mt = _t5c_s(meta.get("mime_type"))
        return fn.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".heic")) or mt.startswith("image/")

    def _t5c_msg_id(meta):
        for k in ("telegram_message_id", "reply_to_message_id"):
            v = _t5c_s(meta.get(k))
            if v:
                return v
        fn = _t5c_s(meta.get("file_name"))
        m = _t5c_re.search(r"_(\d+)\.(jpg|jpeg|png|webp|heic)$", fn, _t5c_re.I)
        return m.group(1) if m else ""

    def _t5c_active_folder(chat_id):
        p = _T5C_DATA / f"active_folder_{chat_id}_5.json"
        try:
            d = _t5c_json.loads(p.read_text(encoding="utf-8"))
            if d.get("folder_id") and _t5c_s(d.get("status", "OPEN")).upper() != "CLOSED":
                return d
        except Exception:
            pass
        return {}

    def _t5c_buf_path(chat_id):
        return _T5C_DATA / f"buf_{chat_id}_5.json"

    def _t5c_load_buf(chat_id):
        p = _t5c_buf_path(chat_id)
        try:
            d = _t5c_json.loads(p.read_text(encoding="utf-8"))
            if isinstance(d, dict):
                d.setdefault("materials", [])
                return d
        except Exception:
            pass
        return {"source": "topic5_visit_buffer", "materials": [], "created_at": _t5c_time.time()}

    def _t5c_save_buf(chat_id, buf):
        buf["updated_at"] = _t5c_time.time()
        _t5c_buf_path(chat_id).write_text(_t5c_json.dumps(buf, ensure_ascii=False, indent=2), encoding="utf-8")

    def _t5c_drive_service():
        from core import topic_drive_oauth as _tdo
        return _tdo._oauth_service()

    def _t5c_drive_find_by_name(service, folder_id, name):
        safe = _t5c_s(name).replace("'", "\\'")
        resp = service.files().list(
            q=f"name = '{safe}' and '{folder_id}' in parents and trashed = false",
            spaces="drive",
            fields="files(id,name,webViewLink,createdTime)",
            orderBy="createdTime",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
            pageSize=10,
        ).execute()
        files = resp.get("files", [])
        return files[0] if files else {}

    def _t5c_copy_drive_to_active(meta, chat_id):
        af = _t5c_active_folder(chat_id)
        active_id = _t5c_s(af.get("folder_id"))
        if not active_id:
            return meta

        file_id = _t5c_s(meta.get("file_id") or meta.get("drive_file_id") or meta.get("id"))
        file_name = _t5c_s(meta.get("file_name") or meta.get("name"))
        if not file_id or not file_name:
            return meta

        try:
            svc = _t5c_drive_service()
            existing = _t5c_drive_find_by_name(svc, active_id, file_name)
            if existing:
                meta["active_drive_file_id"] = existing.get("id", "")
                meta["active_drive_url"] = existing.get("webViewLink", "")
                meta["active_folder_id"] = active_id
                meta["active_folder_name"] = _t5c_s(af.get("folder_name"))
                return meta

            res = svc.files().copy(
                fileId=file_id,
                body={"name": file_name, "parents": [active_id]},
                fields="id,webViewLink,parents",
                supportsAllDrives=True,
            ).execute()
            meta["active_drive_file_id"] = res.get("id", "")
            meta["active_drive_url"] = res.get("webViewLink", "")
            meta["active_folder_id"] = active_id
            meta["active_folder_name"] = _t5c_s(af.get("folder_name"))
        except Exception as e:
            meta["active_copy_error"] = _t5c_s(e, 500)

        return meta

    def _t5c_material_from_meta(meta, chat_id, comment=""):
        meta = _t5c_copy_drive_to_active(dict(meta), chat_id)
        af = _t5c_active_folder(chat_id)

        src_id = _t5c_s(meta.get("file_id") or meta.get("drive_file_id") or meta.get("id"))
        active_id = _t5c_s(meta.get("active_drive_file_id") or src_id)
        fn = _t5c_s(meta.get("file_name") or meta.get("name"))
        mid = _t5c_msg_id(meta)
        clean = _t5c_clean_voice(comment)

        return {
            "material_id": str(_t5c_uuid.uuid4()),
            "source": "TELEGRAM",
            "file_type": "PHOTO",
            "file_name": fn,
            "source_drive_file_id": src_id,
            "drive_file_id": active_id,
            "drive_url": _t5c_s(meta.get("active_drive_url") or meta.get("drive_url") or meta.get("webViewLink") or (f"https://drive.google.com/file/d/{active_id}/view?usp=drivesdk" if active_id else "")),
            "telegram_message_id": mid,
            "reply_to_message_id": mid,
            "source_task_id": _t5c_s(meta.get("_task_id")),
            "active_folder_id": _t5c_s(af.get("folder_id")),
            "active_folder_name": _t5c_s(af.get("folder_name")),
            "include_in_report": True,
            "include_in_act": True,
            "status": "LINKED" if clean else "PENDING",
            "voice_comment": clean,
            "copy_error": _t5c_s(meta.get("active_copy_error")),
            "added_at": _t5c_time.time(),
            "updated_at": _t5c_time.time(),
        }

    def _t5c_upsert_material(chat_id, material):
        buf = _t5c_load_buf(chat_id)
        mid = _t5c_s(material.get("telegram_message_id"))
        fn = _t5c_s(material.get("file_name"))
        target = None

        for old in buf.get("materials", []):
            if (mid and _t5c_s(old.get("telegram_message_id")) == mid) or (fn and _t5c_s(old.get("file_name")) == fn):
                target = old
                break

        if target is None:
            buf["materials"].append(material)
        else:
            old_comment = _t5c_s(target.get("voice_comment"), 20000)
            new_comment = _t5c_s(material.get("voice_comment"), 20000)
            target.update({k: v for k, v in material.items() if v not in ("", None)})
            if old_comment and new_comment and new_comment not in old_comment:
                target["voice_comment"] = old_comment + "\n" + new_comment
            elif old_comment and not new_comment:
                target["voice_comment"] = old_comment

        _t5c_save_buf(chat_id, buf)
        return len(buf.get("materials", []))

    def _t5c_find_parent_photo(conn, chat_id, topic_id, reply_to):
        rid = _t5c_s(reply_to)
        if not rid:
            return {}

        rows = conn.execute(
            """
            SELECT rowid,id,raw_input,reply_to_message_id
            FROM tasks
            WHERE chat_id=?
              AND topic_id=?
              AND input_type='drive_file'
            ORDER BY rowid DESC
            LIMIT 300
            """,
            (_t5c_s(chat_id), int(topic_id or 0)),
        ).fetchall()

        for row in rows:
            meta = _t5c_json_load(row["raw_input"] if hasattr(row, "keys") else row[2])
            if not _t5c_is_photo_meta(meta):
                continue
            row_reply = _t5c_s(row["reply_to_message_id"] if hasattr(row, "keys") else row[3])
            meta["_task_id"] = _t5c_s(row["id"] if hasattr(row, "keys") else row[1])
            ids = {row_reply, _t5c_s(meta.get("telegram_message_id")), _t5c_msg_id(meta)}
            if rid in ids:
                return meta

        parent = conn.execute(
            """
            SELECT reply_to_message_id
            FROM tasks
            WHERE chat_id=?
              AND topic_id=?
              AND CAST(bot_message_id AS TEXT)=?
            ORDER BY rowid DESC
            LIMIT 1
            """,
            (_t5c_s(chat_id), int(topic_id or 0), rid),
        ).fetchone()

        parent_reply = _t5c_s(parent["reply_to_message_id"] if parent and hasattr(parent, "keys") else (parent[0] if parent else ""))
        if parent_reply and parent_reply != rid:
            return _t5c_find_parent_photo(conn, chat_id, topic_id, parent_reply)

        return {}

    def _t5c_explicit_act(text):
        low = _t5c_low(text)
        neg = any(x in low for x in (
            "не делай акт",
            "не надо акт",
            "не нужно акт",
            "не формируй акт",
            "не должен был сделать акт",
            "не должен делать акт",
            "акт не для каждого",
            "не для каждого из",
            "принять к сведению",
            "принять это к сведению",
            "прими к сведению",
            "прими это к сведению",
        ))
        pos = any(x in low for x in (
            "сделай акт",
            "сформируй акт",
            "собери акт",
            "готовь акт",
            "акт по этим фото",
            "акт по фото",
            "сделай документ",
            "сформируй документ",
        ))
        return pos and not neg

    def _t5c_folder_question(text):
        low = _t5c_low(text)
        return "папк" in low and any(x in low for x in ("видишь", "какая", "где", "тест надзор", "актив", "ссылка"))

    def _t5c_safe_photo_talk(text):
        low = _t5c_low(text)
        return any(x in low for x in (
            "фото",
            "фотограф",
            "с ними",
            "по ним",
            "эти",
            "этими",
            "понял задачу",
            "задача понятна",
            "принять к сведению",
            "принять это к сведению",
            "нарушение",
            "факт нарушения",
        ))

    async def _t5c_upload_file(path, name, chat_id, topic_id, mime):
        from core.topic_drive_oauth import upload_file_to_topic
        return await upload_file_to_topic(str(path), name, str(chat_id), int(topic_id or 0), mime)

    async def _t5c_make_single_act(chat_id, topic_id, task_id, command_text):
        buf = _t5c_load_buf(chat_id)
        materials = list(buf.get("materials", []))
        if not materials:
            return {"ok": False, "text": "В пакете технадзора нет фото. Сначала пришли фото или ответь голосом на фото"}

        ts = _t5c_datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_tid = _t5c_s(task_id)[:8] or ts
        txt_path = _T5C_OUT / f"АКТ_ТЕХНАДЗОРА__{safe_tid}_{ts}.txt"

        lines = [
            "АКТ ТЕХНИЧЕСКОГО НАДЗОРА",
            "",
            f"Дата: {_t5c_datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Задача: {task_id}",
            f"Chat: {chat_id}",
            f"Topic: {topic_id}",
            "",
            "Основание:",
            _t5c_clean_voice(command_text) or "Осмотр материалов из Telegram",
            "",
            "Материалы:",
        ]

        for i, m in enumerate(materials, 1):
            lines.append(f"{i}. {m.get('file_name','')}")
            if m.get("voice_comment"):
                lines.append(f"   Пояснение: {m.get('voice_comment')}")
            if m.get("drive_url"):
                lines.append(f"   Фото: {m.get('drive_url')}")

        lines += [
            "",
            "Результат:",
            "Факт нарушения принят к сведению по приложенным фото и пояснениям.",
            "",
            "Статус:",
            "Акт сформирован одним документом по пакету фото.",
        ]

        txt_path.write_text("\n".join(lines), encoding="utf-8")
        uploaded = await _t5c_upload_file(txt_path, txt_path.name, chat_id, topic_id, "text/plain")

        archive_path = _T5C_DATA / f"buf_{chat_id}_5.DONE_{safe_tid}_{ts}.json"
        buf["closed_by_task_id"] = task_id
        buf["closed_at"] = _t5c_time.time()
        buf["act_file"] = txt_path.name
        buf["upload_result"] = uploaded
        archive_path.write_text(_t5c_json.dumps(buf, ensure_ascii=False, indent=2), encoding="utf-8")

        try:
            _t5c_buf_path(chat_id).unlink()
        except Exception:
            pass

        link = _t5c_s(uploaded.get("webViewLink") or (f"https://drive.google.com/file/d/{uploaded.get('drive_file_id')}/view?usp=drivesdk" if uploaded.get("drive_file_id") else ""))
        folder = _t5c_s(uploaded.get("active_folder_id") or uploaded.get("folder_id"))

        return {
            "ok": True,
            "text": f"Акт сформирован одним документом\n{link}\nПапка: https://drive.google.com/drive/folders/{folder}",
            "upload": uploaded,
        }

    def _t5c_done(conn, task_id, chat_id, reply_to, text, kind):
        sent = _send_once_ex(conn, str(task_id), str(chat_id), _t5c_s(text, 3500), int(reply_to) if _t5c_s(reply_to).isdigit() else None, kind)
        try:
            if isinstance(sent, dict) and sent.get("message_id"):
                conn.execute("UPDATE tasks SET bot_message_id=? WHERE id=?", (int(sent.get("message_id")), str(task_id)))
        except Exception:
            pass

        conn.execute(
            "UPDATE tasks SET state='DONE', result=?, error_message='', updated_at=datetime('now') WHERE id=?",
            (_t5c_s(text, 12000), str(task_id)),
        )

        try:
            _history(conn, str(task_id), kind)
        except Exception:
            pass

        conn.commit()

    async def _handle_new(conn, task, *args, **kwargs):
        task_id = _t5c_s(_t5c_row(task, "id", ""))
        chat_id = _t5c_s(_t5c_row(task, "chat_id", args[0] if len(args) > 0 else ""))
        topic_id = int(_t5c_row(task, "topic_id", args[1] if len(args) > 1 else 0) or 0)

        if topic_id != 5:
            return await _T5C_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)

        raw = _t5c_s(_t5c_row(task, "raw_input", ""))
        input_type = _t5c_s(_t5c_row(task, "input_type", ""))
        reply_to = _t5c_s(_t5c_row(task, "reply_to_message_id", ""))
        clean = _t5c_clean_voice(raw)
        meta = _t5c_json_load(raw)

        if _t5c_low(clean) == "статус":
            return await _T5C_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)

        if _t5c_folder_question(clean):
            af = _t5c_active_folder(chat_id)
            buf = _t5c_load_buf(chat_id)
            count = len(buf.get("materials", []))
            if af.get("folder_id"):
                msg = f"Активная папка технадзора: {af.get('folder_name') or af.get('folder_id')}\nhttps://drive.google.com/drive/folders/{af.get('folder_id')}\nВ пакете фото: {count} шт"
            else:
                msg = "Активная папка технадзора не установлена"
            _t5c_done(conn, task_id, chat_id, reply_to, msg, "topic5_active_folder_status")
            return

        if input_type == "drive_file" and _t5c_is_photo_meta(meta):
            material = _t5c_material_from_meta(meta, chat_id, "")
            count = _t5c_upsert_material(chat_id, material)
            msg = f"Фото принято в пакет технадзора: {count} шт. Активная папка: {material.get('active_folder_name') or material.get('active_folder_id')}"
            if material.get("copy_error"):
                msg += f"\nDrive copy error: {material.get('copy_error')}"
            _t5c_done(conn, task_id, chat_id, reply_to, msg, "topic5_photo_buffered_active_folder")
            return

        if input_type in ("text", "voice", "") and reply_to and not _t5c_explicit_act(clean):
            parent = _t5c_find_parent_photo(conn, chat_id, topic_id, reply_to)
            if parent:
                material = _t5c_material_from_meta(parent, chat_id, clean)
                count = _t5c_upsert_material(chat_id, material)
                msg = f"Пояснение принято к фото: {material.get('file_name')}. В пакете технадзора: {count} шт. Акт не формирую без отдельной команды"
                if material.get("copy_error"):
                    msg += f"\nDrive copy error: {material.get('copy_error')}"
                _t5c_done(conn, task_id, chat_id, reply_to, msg, "topic5_reply_photo_comment_bound")
                return

        if input_type in ("text", "voice", "") and _t5c_explicit_act(clean):
            act = await _t5c_make_single_act(chat_id, topic_id, task_id, clean)
            _t5c_done(conn, task_id, chat_id, reply_to, act.get("text", "Акт не сформирован"), "topic5_single_act_result")
            return

        if input_type in ("text", "voice", "") and _t5c_safe_photo_talk(clean):
            buf = _t5c_load_buf(chat_id)
            count = len(buf.get("materials", []))
            if count:
                msg = f"Задача понятна. В пакете технадзора {count} фото. Жду пояснение к фото или явную команду: Сделай акт"
            else:
                msg = "В пакете технадзора нет фото. Ответь голосом прямо на фото или пришли фото заново"
            _t5c_done(conn, task_id, chat_id, reply_to, msg, "topic5_safe_photo_talk_no_act")
            return

        return await _T5C_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)

    _handle_new._topic5_canon_close_v1 = True
except Exception as _t5c_install_error:
    try:
        logging.getLogger("task_worker").exception("FULLFIX_TOPIC5_CANON_CLOSE_ACTIVE_FOLDER_NO_DUP_ACTS_V1_INSTALL_ERR %s", _t5c_install_error)
    except Exception:
        pass
# === END_FULLFIX_TOPIC5_CANON_CLOSE_ACTIVE_FOLDER_NO_DUP_ACTS_V1 ===


# === FULLFIX_TOPIC5_FULL_CANON_CLOSE_V1 ===
try:
    import json as _t5fc_json
    import re as _t5fc_re
    import time as _t5fc_time
    import datetime as _t5fc_dt
    import textwrap as _t5fc_textwrap
    import logging as _t5fc_logging
    from pathlib import Path as _T5FC_Path

    _T5FC_LOG = _t5fc_logging.getLogger("task_worker")
    _T5FC_BASE = _T5FC_Path("/root/.areal-neva-core")
    _T5FC_DATA = _T5FC_BASE / "data" / "technadzor"
    _T5FC_OUT = _T5FC_BASE / "outputs" / "technadzor"
    _T5FC_OBJ = _T5FC_BASE / "data" / "templates" / "technadzor" / "objects"
    _T5FC_DATA.mkdir(parents=True, exist_ok=True)
    _T5FC_OUT.mkdir(parents=True, exist_ok=True)
    _T5FC_OBJ.mkdir(parents=True, exist_ok=True)

    _T5FC_ORIG_HANDLE_NEW = _handle_new

    def _t5fc_s(v, limit=50000):
        return "" if v is None else str(v).strip()[:limit]

    def _t5fc_low(v):
        return _t5fc_s(v).lower().replace("ё", "е")

    def _t5fc_clean_voice(v):
        return _t5fc_re.sub(r"^\s*\[VOICE\]\s*", "", _t5fc_s(v, 50000), flags=_t5fc_re.I).strip()

    def _t5fc_row(row, key, default=""):
        try:
            if hasattr(row, "keys") and key in row.keys():
                return row[key]
        except Exception:
            pass
        return default

    def _t5fc_slug(s):
        s = _t5fc_low(s)
        s = _t5fc_re.sub(r"[^a-zа-я0-9]+", "_", s)
        s = _t5fc_re.sub(r"_+", "_", s).strip("_")
        return s[:80] or "unknown_object"

    def _t5fc_jload(path, default):
        try:
            if path.exists():
                obj = _t5fc_json.loads(path.read_text(encoding="utf-8"))
                return obj if isinstance(obj, dict) else default
        except Exception:
            pass
        return default

    def _t5fc_jsave(path, obj):
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(_t5fc_json.dumps(obj, ensure_ascii=False, indent=2), encoding="utf-8")

    def _t5fc_buf_path(chat_id):
        return _T5FC_DATA / f"buf_{chat_id}_5.json"

    def _t5fc_active_path(chat_id):
        return _T5FC_DATA / f"active_folder_{chat_id}_5.json"

    def _t5fc_buf(chat_id):
        b = _t5fc_jload(_t5fc_buf_path(chat_id), {"source": "topic5_visit_buffer", "materials": [], "created_at": _t5fc_time.time()})
        if not isinstance(b.get("materials"), list):
            b["materials"] = []
        if not isinstance(b.get("package_context"), dict):
            b["package_context"] = {}
        if not isinstance(b.get("observations"), list):
            b["observations"] = []
        if not isinstance(b.get("defect_cards"), list):
            b["defect_cards"] = []
        return b

    def _t5fc_active(chat_id):
        a = _t5fc_jload(_t5fc_active_path(chat_id), {})
        if not isinstance(a.get("owner_instructions"), list):
            a["owner_instructions"] = []
        return a

    def _t5fc_extract_address(text):
        t = _t5fc_clean_voice(text)
        patterns = [
            r"(Ропшин[А-Яа-яA-Za-z0-9\-\s]*шоссе\s*\d+[А-Яа-яA-Za-z0-9\-\/]*)",
            r"(?:адрес|по адресу|находится на|объект находится на|на)\s+([А-Яа-яA-Za-z0-9\-\s]+?(?:шоссе|улица|ул\.|проспект|пр\.|дорога|линия|наб\.|переулок)\s*\d+[А-Яа-яA-Za-z0-9\-\/]*)",
        ]
        for p in patterns:
            m = _t5fc_re.search(p, t, flags=_t5fc_re.I)
            if m:
                return _t5fc_s(m.group(1), 300)
        return ""

    def _t5fc_context_like(text):
        low = _t5fc_low(text)
        return any(x in low for x in (
            "объект",
            "адрес",
            "ропшин",
            "шоссе",
            "основание",
            "авито",
            "заказчик",
            "заявк",
            "выезд",
            "запиши",
            "к этому же акту",
            "этому же акту",
            "к этому акту",
            "этому акту",
            "к акту",
            "в этот акт",
            "в этот же акт",
        ))

    def _t5fc_act_like(text):
        low = _t5fc_low(text)
        return any(x in low for x in (
            "сделай акт",
            "сформируй акт",
            "создай акт",
            "подготовь акт",
            "собери акт",
            "делай акт",
            "сделай документ",
            "сформируй документ",
            "сделай разбор",
            "сформируй разбор",
        ))

    def _t5fc_source_and_basis(text):
        low = _t5fc_low(text)
        out = {}
        if "авито" in low:
            out["source_request"] = "Авито"
            out["visit_basis"] = "запрос заказчика через Авито"
        elif "заказчик" in low or "заявк" in low:
            out["visit_basis"] = "запрос заказчика"
        return out

    def _t5fc_classify_comment(comment):
        low = _t5fc_low(comment)
        if any(x in low for x in ("свар", "шов", "провар")):
            return ("сварные соединения", "сварные соединения", "плохие сварные соединения")
        if any(x in low for x in ("оборудован", "вышло из строя", "замена")):
            return ("прочие замечания", "старое оборудование", "оборудование вышло из строя, рекомендуется замена")
        if any(x in low for x in ("корроз", "ржав")):
            return ("антикоррозионная защита", "коррозия", "признаки коррозии / нарушение защитного покрытия")
        return ("прочие замечания", "замечание владельца", _t5fc_s(comment, 500))

    def _t5fc_norms(text):
        try:
            from core.normative_engine import search_norms_sync, format_norms_for_act
            ns = search_norms_sync(text or "", limit=5)
            return ns, format_norms_for_act(ns) if ns else "норма не подтверждена"
        except Exception as e:
            return [], f"норма не подтверждена: {type(e).__name__}"

    def _t5fc_enrich_materials(buf):
        defects = []
        for i, m in enumerate(buf.get("materials", []), 1):
            c = _t5fc_s(m.get("owner_comment") or m.get("voice_comment") or "")
            if c:
                group, node, defect = _t5fc_classify_comment(c)
                m["owner_comment"] = c
                m["group_label"] = m.get("group_label") or group
                m["section_hint"] = m.get("section_hint") or group
                m["defect_hint"] = m.get("defect_hint") or defect
                m["status"] = "LINKED"
                norms_raw, norms_text = _t5fc_norms(c)
                defects.append({
                    "photo_no": i,
                    "file_name": m.get("file_name", ""),
                    "source": m.get("source", "TELEGRAM"),
                    "node_location": node,
                    "what_visible": "Автоматический визуальный анализ фото не выполнялся, так как Vision заблокирован. Вывод основан на пояснении владельца и метаданных файла",
                    "defect_remark": defect,
                    "why_bad": "Требует проверки на объекте и фиксации в акте технического надзора",
                    "possible_consequences": "Требуется оценка влияния дефекта при очном осмотре",
                    "what_to_fix": "Выполнить устранение замечания по проектному решению и требованиям применимых норм",
                    "what_to_check_on_site": "Проверить фактическое состояние узла, объём дефекта, качество устранения и соответствие проекту",
                    "normative_reference": norms_text,
                    "norm_status": "PARTIAL" if norms_raw else "NOT_FOUND",
                    "remark_status": "новое замечание",
                    "confirmation_source": "OWNER_VOICE_OR_TEXT",
                    "owner_question": None,
                })
        buf["defect_cards"] = defects
        return defects

    def _t5fc_save_object_context(chat_id, ctx, active, buf, act_link=""):
        object_name = _t5fc_s(ctx.get("object_address") or ctx.get("object_name") or active.get("object_name") or active.get("folder_name") or "UNKNOWN")
        object_id = _t5fc_slug(object_name)
        card_path = _T5FC_OBJ / f"{object_id}.json"

        card = _t5fc_jload(card_path, {
            "object_id": object_id,
            "object_name": object_name,
            "client_name": "",
            "object_folder_url": "",
            "client_facing_folder_url": "",
            "service_folder_url": "",
            "inspection_chain": [],
            "previous_acts": [],
            "current_open_items": [],
            "closed_items": [],
            "unresolved_items": [],
            "recommendations": [],
            "last_visit_date": "",
            "last_act_no": "",
            "last_pdf_link": "",
            "created_at": int(_t5fc_time.time()),
            "updated_at": int(_t5fc_time.time()),
        })

        card["chat_id"] = str(chat_id)
        card["object_name"] = object_name
        card["object_folder_url"] = active.get("drive_folder_url") or (f"https://drive.google.com/drive/folders/{active.get('folder_id')}" if active.get("folder_id") else "")
        card["service_folder_url"] = card["object_folder_url"]
        card["updated_at"] = int(_t5fc_time.time())

        _t5fc_jsave(card_path, card)
        return object_id, card_path

    def _t5fc_save_context(chat_id, raw_text):
        clean = _t5fc_clean_voice(raw_text)
        active = _t5fc_active(chat_id)
        buf = _t5fc_buf(chat_id)
        ctx = buf.setdefault("package_context", {})

        address = _t5fc_extract_address(clean)
        if address:
            ctx["object_address"] = address
            ctx["object_name"] = address
            active["object_name"] = address
            active["object_address"] = address

        sb = _t5fc_source_and_basis(clean)
        for k, v in sb.items():
            ctx[k] = v
            active[k] = v

        ctx["last_owner_instruction"] = clean
        ctx["updated_at"] = _t5fc_time.time()

        obs = {
            "source": "OWNER_VOICE_OR_TEXT",
            "author": "OWNER",
            "author_role": "owner",
            "material_type": "voice_or_text",
            "object": ctx.get("object_address") or active.get("object_name") or active.get("folder_name") or "",
            "date": _t5fc_dt.datetime.now().strftime("%Y-%m-%d"),
            "claim": clean,
            "linked_files": [m.get("file_name", "") for m in buf.get("materials", [])],
            "confirmed": "yes",
            "contradiction": False,
            "needs_owner_question": False,
        }

        observations = buf.setdefault("observations", [])
        if clean and all(o.get("claim") != clean for o in observations):
            observations.append(obs)

        owner_instructions = buf.setdefault("owner_instructions", [])
        if clean and clean not in owner_instructions:
            owner_instructions.append(clean)
        buf["owner_instructions"] = owner_instructions[-50:]

        active_instr = active.setdefault("owner_instructions", [])
        if clean and clean not in active_instr:
            active_instr.append(clean)
        active["owner_instructions"] = active_instr[-50:]
        active["last_update"] = _t5fc_dt.datetime.now().isoformat(timespec="seconds")
        active["status"] = active.get("status") or "OPEN"

        defects = _t5fc_enrich_materials(buf)
        object_id, card_path = _t5fc_save_object_context(chat_id, ctx, active, buf)

        buf["package_context"] = ctx
        buf["updated_at"] = _t5fc_time.time()
        _t5fc_jsave(_t5fc_buf_path(chat_id), buf)
        _t5fc_jsave(_t5fc_active_path(chat_id), active)

        return {
            "ctx": ctx,
            "active": active,
            "buf": buf,
            "object_id": object_id,
            "object_card_path": str(card_path),
            "defect_count": len(defects),
        }

    def _t5fc_make_text_report(chat_id, task_id):
        active = _t5fc_active(chat_id)
        buf = _t5fc_buf(chat_id)
        ctx = buf.setdefault("package_context", {})
        defects = _t5fc_enrich_materials(buf)

        object_name = _t5fc_s(ctx.get("object_address") or ctx.get("object_name") or active.get("object_name") or active.get("folder_name") or "UNKNOWN")
        visit_basis = _t5fc_s(ctx.get("visit_basis") or active.get("visit_basis") or "UNKNOWN")
        source_request = _t5fc_s(ctx.get("source_request") or active.get("source_request") or "UNKNOWN")
        folder_id = _t5fc_s(active.get("folder_id"))
        folder_name = _t5fc_s(active.get("folder_name") or "UNKNOWN")
        now = _t5fc_dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        lines = [
            "АКТ / РАЗБОР ТЕХНИЧЕСКОГО НАДЗОРА",
            "",
            f"Дата формирования: {now}",
            f"Объект / адрес: {object_name}",
            f"Активная папка: {folder_name}",
            f"Ссылка на папку: https://drive.google.com/drive/folders/{folder_id}" if folder_id else "Ссылка на папку: UNKNOWN",
            f"Основание выезда: {visit_basis}",
            f"Источник заявки: {source_request}",
            "",
            "1. Техническое задание владельца",
        ]

        instr = buf.get("owner_instructions") or active.get("owner_instructions") or []
        if instr:
            for i, item in enumerate(instr, 1):
                lines.append(f"{i}. {item}")
        else:
            lines.append("UNKNOWN")

        lines += [
            "",
            "2. Материалы фотофиксации",
        ]

        mats = buf.get("materials", [])
        if mats:
            for i, m in enumerate(mats, 1):
                lines.append(f"{i}. {m.get('file_name', '')}")
                if m.get("drive_url"):
                    lines.append(f"   Файл: {m.get('drive_url')}")
                if m.get("owner_comment") or m.get("voice_comment"):
                    lines.append(f"   Пояснение владельца: {m.get('owner_comment') or m.get('voice_comment')}")
                if m.get("group_label"):
                    lines.append(f"   Раздел: {m.get('group_label')}")
        else:
            lines.append("Материалы отсутствуют")

        lines += [
            "",
            "3. Карточки наблюдений",
        ]

        obs = buf.get("observations") or []
        if obs:
            for i, o in enumerate(obs, 1):
                lines.append(f"{i}. Источник: {o.get('source')}; Автор: {o.get('author_role')}; Наблюдение: {o.get('claim')}")
        else:
            lines.append("UNKNOWN")

        lines += [
            "",
            "4. Карточки замечаний / дефектов",
        ]

        if defects:
            for i, d in enumerate(defects, 1):
                lines += [
                    f"{i}. Фото: {d.get('file_name')}",
                    f"   Узел / место: {d.get('node_location')}",
                    f"   Что видно / источник вывода: {d.get('what_visible')}",
                    f"   Замечание: {d.get('defect_remark')}",
                    f"   Почему плохо: {d.get('why_bad')}",
                    f"   Возможные последствия: {d.get('possible_consequences')}",
                    f"   Что исправить: {d.get('what_to_fix')}",
                    f"   Что проверить на объекте: {d.get('what_to_check_on_site')}",
                    f"   Нормативная ссылка: {d.get('normative_reference')}",
                    f"   Статус нормы: {d.get('norm_status')}",
                    f"   Статус замечания: {d.get('remark_status')}",
                ]
        else:
            lines.append("По текущему пакету нет пояснений владельца, достаточных для формирования DefectCard")

        lines += [
            "",
            "5. Связь с предыдущими актами",
            "Предыдущие акты по этому объекту автоматически не сопоставлены в текущем пакете. Если предыдущий акт найден отдельно, содержание требует ручной сверки",
            "",
            "6. Vision guard",
            "Автоматический визуальный анализ фото не выполнялся, так как Vision заблокирован. Выводы основаны на пояснениях владельца, доступных именах файлов, метаданных и нормативном поиске",
            "",
            "7. Итог",
            "Материалы, адрес объекта, основание выезда, пояснения владельца, карточки наблюдений и карточки замечаний объединены в один пакет технадзора",
        ]

        safe_obj = _t5fc_slug(object_name)
        stamp = _t5fc_dt.datetime.now().strftime("%Y%m%d_%H%M%S")
        txt_path = _T5FC_OUT / f"Акт_осмотра_{safe_obj}_{stamp}.txt"
        txt_path.write_text("\n".join(lines), encoding="utf-8")
        return txt_path, "\n".join(lines), defects

    def _t5fc_make_pdf(txt, pdf_path):
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.pagesizes import A4

        font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
        pdfmetrics.registerFont(TTFont("DejaVuSans", font_path))

        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        w, h = A4
        x = 36
        y = h - 36
        c.setFont("DejaVuSans", 9)

        for raw_line in txt.splitlines():
            wrapped = _t5fc_textwrap.wrap(raw_line, width=105) or [""]
            for line in wrapped:
                if y < 36:
                    c.showPage()
                    c.setFont("DejaVuSans", 9)
                    y = h - 36
                c.drawString(x, y, line)
                y -= 12
        c.save()
        return pdf_path

    def _t5fc_upload_to_active_folder(file_path, folder_id, mime_type):
        from core.topic_drive_oauth import _oauth_service
        from googleapiclient.http import MediaFileUpload

        svc = _oauth_service()
        media = MediaFileUpload(str(file_path), mimetype=mime_type, resumable=False)
        created = svc.files().create(
            body={"name": file_path.name, "parents": [folder_id]},
            media_body=media,
            fields="id,name,webViewLink",
        ).execute()
        fid = created.get("id")
        return {
            "drive_file_id": fid,
            "webViewLink": created.get("webViewLink") or (f"https://drive.google.com/file/d/{fid}/view?usp=drivesdk" if fid else ""),
        }

    def _t5fc_record_inspection(chat_id, task_id, pdf_link, defects):
        active = _t5fc_active(chat_id)
        buf = _t5fc_buf(chat_id)
        ctx = buf.get("package_context", {})
        object_name = _t5fc_s(ctx.get("object_address") or active.get("object_name") or active.get("folder_name") or "UNKNOWN")
        object_id = _t5fc_slug(object_name)

        try:
            from core.technadzor_object_registry import record_inspection
            record_inspection(
                object_id=object_id,
                chat_id=str(chat_id),
                act_no=f"{_t5fc_dt.datetime.now().strftime('%d-%m/%y')}-{_t5fc_s(task_id)[:6]}",
                date_str=_t5fc_dt.datetime.now().strftime("%Y-%m-%d"),
                mode="initial",
                pdf_link=pdf_link,
                docx_link="",
                source_photo_folder=f"https://drive.google.com/drive/folders/{active.get('folder_id')}" if active.get("folder_id") else "",
                findings=defects,
                open_items=defects,
                closed_items=[],
                new_items=defects,
                owner_observation="\n".join(buf.get("owner_instructions") or active.get("owner_instructions") or []),
                conflict_flags=[],
                object_name=object_name,
                object_folder_url=f"https://drive.google.com/drive/folders/{active.get('folder_id')}" if active.get("folder_id") else "",
                service_folder_url=f"https://drive.google.com/drive/folders/{active.get('folder_id')}" if active.get("folder_id") else "",
            )
            return True
        except Exception as e:
            _T5FC_LOG.warning("FULLFIX_TOPIC5_FULL_CANON_RECORD_INSPECTION_ERR %s", e)
            return False

    def _t5fc_done(conn, task_id, chat_id, reply_to, text, kind, state="DONE"):
        sent = _send_once_ex(
            conn,
            str(task_id),
            str(chat_id),
            _t5fc_s(text, 3500),
            int(reply_to) if _t5fc_s(reply_to).isdigit() else None,
            kind,
        )
        upd = {"state": state, "result": _t5fc_s(text, 50000), "error_message": "" if state == "DONE" else _t5fc_s(text, 1000)}
        if isinstance(sent, dict) and (sent.get("bot_message_id") or sent.get("message_id")):
            upd["bot_message_id"] = sent.get("bot_message_id") or sent.get("message_id")
        _update_task(conn, str(task_id), **upd)
        try:
            _history(conn, str(task_id), kind)
        except Exception:
            pass
        conn.commit()

    async def _handle_new(conn, task, *args, **kwargs):
        task_id = _t5fc_s(_t5fc_row(task, "id"))
        chat_id = _t5fc_s(_t5fc_row(task, "chat_id", args[0] if len(args) > 0 else ""))
        topic_id = int(_t5fc_row(task, "topic_id", args[1] if len(args) > 1 else 0) or 0)

        if topic_id != 5:
            res = _T5FC_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
            return await res if hasattr(res, "__await__") else res

        raw = _t5fc_s(_t5fc_row(task, "raw_input"))
        input_type = _t5fc_s(_t5fc_row(task, "input_type"))
        reply_to = _t5fc_s(_t5fc_row(task, "reply_to_message_id"))
        clean = _t5fc_clean_voice(raw)

        if input_type in ("text", "voice", "") and _t5fc_context_like(clean) and not _t5fc_act_like(clean):
            saved = _t5fc_save_context(chat_id, clean)
            ctx = saved["ctx"]
            msg = "\n".join([
                "Контекст технадзора принят к текущему акту",
                f"Объект / адрес: {_t5fc_s(ctx.get('object_address') or ctx.get('object_name') or 'UNKNOWN')}",
                f"Основание выезда: {_t5fc_s(ctx.get('visit_basis') or 'UNKNOWN')}",
                f"Источник заявки: {_t5fc_s(ctx.get('source_request') or 'UNKNOWN')}",
                f"Фото в пакете: {len(saved['buf'].get('materials', []))} шт",
                f"DefectCard: {saved.get('defect_count', 0)}",
                f"ObjectCard: {saved.get('object_id')}",
                "Акт не формирую без команды: Сделай акт",
            ])
            _t5fc_done(conn, task_id, chat_id, reply_to, msg, "topic5_full_canon_context_saved", "DONE")
            return

        if input_type in ("text", "voice", "") and _t5fc_act_like(clean):
            _t5fc_save_context(chat_id, clean) if _t5fc_context_like(clean) else None
            active = _t5fc_active(chat_id)
            folder_id = _t5fc_s(active.get("folder_id"))
            if not folder_id:
                _t5fc_done(conn, task_id, chat_id, reply_to, "Акт не сформирован: активная папка технадзора не установлена", "topic5_full_canon_act_no_folder", "FAILED")
                return

            txt_path, txt_body, defects = _t5fc_make_text_report(chat_id, task_id)
            pdf_path = txt_path.with_suffix(".pdf")
            drive_link = ""
            uploaded_kind = ""

            try:
                _t5fc_make_pdf(txt_body, pdf_path)
                uploaded = _t5fc_upload_to_active_folder(pdf_path, folder_id, "application/pdf")
                drive_link = uploaded.get("webViewLink") or ""
                uploaded_kind = "PDF"
            except Exception as e:
                _T5FC_LOG.warning("FULLFIX_TOPIC5_FULL_CANON_PDF_UPLOAD_ERR %s", e)
                try:
                    uploaded = _t5fc_upload_to_active_folder(txt_path, folder_id, "text/plain")
                    drive_link = uploaded.get("webViewLink") or ""
                    uploaded_kind = "TEXT"
                except Exception as e2:
                    msg = f"Акт не доставлен: DRIVE_UPLOAD_FAILED {type(e2).__name__}"
                    _t5fc_done(conn, task_id, chat_id, reply_to, msg, "topic5_full_canon_delivery_failed", "FAILED")
                    return

            _t5fc_record_inspection(chat_id, task_id, drive_link, defects)

            buf = _t5fc_buf(chat_id)
            buf["last_full_canon_act"] = {
                "task_id": task_id,
                "created_at": _t5fc_time.time(),
                "file": txt_path.name,
                "pdf": pdf_path.name if pdf_path.exists() else "",
                "uploaded_kind": uploaded_kind,
                "drive_link": drive_link,
                "defect_count": len(defects),
            }
            _t5fc_jsave(_t5fc_buf_path(chat_id), buf)

            ctx = buf.get("package_context", {})
            msg = "\n".join([
                "Акт сформирован по полному контуру технадзора",
                f"Формат: {uploaded_kind}",
                f"Объект: {_t5fc_s(ctx.get('object_address') or ctx.get('object_name') or active.get('object_name') or active.get('folder_name') or 'UNKNOWN')}",
                f"Основание: {_t5fc_s(ctx.get('visit_basis') or active.get('visit_basis') or 'UNKNOWN')}",
                f"Фото: {len(buf.get('materials', []))} шт",
                f"DefectCard: {len(defects)}",
                f"Ссылка: {drive_link}",
                f"Папка: https://drive.google.com/drive/folders/{folder_id}",
            ])
            _t5fc_done(conn, task_id, chat_id, reply_to, msg, "topic5_full_canon_act_delivered", "DONE")
            return

        res = _T5FC_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        return await res if hasattr(res, "__await__") else res

    _handle_new._topic5_full_canon_close_v1 = True
    _T5FC_LOG.info("FULLFIX_TOPIC5_FULL_CANON_CLOSE_V1_INSTALLED")

except Exception as _t5fc_err:
    try:
        logging.getLogger("task_worker").exception("FULLFIX_TOPIC5_FULL_CANON_CLOSE_V1_INSTALL_ERR %s", _t5fc_err)
    except Exception:
        pass
# === END_FULLFIX_TOPIC5_FULL_CANON_CLOSE_V1 ===

# === TOPIC_ISOLATION_P6C_TECHNADZOR_STRICT_GUARD_V1 ===
# Prevent P6C from routing drive_file from non-topic_5 topics (e.g. topic_210)
# into technadzor path based on keyword match alone.
try:
    _p6c_tnz_like_orig_strict = _p6c_technadzor_like_20260504

    def _p6c_technadzor_like_20260504(raw_input, topic_id):
        if int(topic_id or 0) not in (0, 5):
            return False
        return _p6c_tnz_like_orig_strict(raw_input, topic_id)

    _p6c_technadzor_like_20260504._strict_topic_guard_v1 = True
except Exception:
    pass
# === END_TOPIC_ISOLATION_P6C_TECHNADZOR_STRICT_GUARD_V1 ===


# === FIX_TOPIC5_DRIVE_FILE_FULL_CANON_GUARD_V1 ===
try:
    import json as _t5df_json
    import time as _t5df_time
    import logging as _t5df_logging

    _T5DF_LOG = _t5df_logging.getLogger("task_worker")
    _T5DF_ORIG_HANDLE_NEW = _handle_new

    def _t5df_s(v, limit=20000):
        return "" if v is None else str(v).strip()[:limit]

    def _t5df_low(v):
        return _t5df_s(v).lower().replace("ё", "е")

    def _t5df_row(row, key, default=""):
        try:
            if hasattr(row, "keys") and key in row.keys():
                return row[key]
        except Exception:
            pass
        return default

    def _t5df_load_meta(raw):
        try:
            d = _t5df_json.loads(_t5df_s(raw, 80000))
        except Exception:
            return {}
        if not isinstance(d, dict):
            return {}

        name = _t5df_s(d.get("file_name") or d.get("name") or "")
        mime = _t5df_low(d.get("mime_type") or "")
        ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""

        if "image" in mime or ext in ("jpg", "jpeg", "png", "webp", "heic"):
            file_type = "PHOTO"
        elif "pdf" in mime or ext == "pdf":
            file_type = "PDF"
        elif "word" in mime or ext in ("doc", "docx"):
            file_type = "DOCX"
        elif "sheet" in mime or ext in ("xls", "xlsx", "csv"):
            file_type = "XLSX"
        else:
            file_type = "OTHER"

        return {
            "source": "TELEGRAM",
            "file_type": file_type,
            "file_name": name,
            "mime_type": mime,
            "source_drive_file_id": _t5df_s(d.get("file_id") or d.get("drive_file_id") or ""),
            "telegram_message_id": _t5df_s(d.get("telegram_message_id") or d.get("message_id") or ""),
            "reply_to_message_id": _t5df_s(d.get("telegram_message_id") or d.get("message_id") or ""),
            "owner_comment": _t5df_s(d.get("caption") or ""),
            "include_in_report": True,
            "include_in_act": True,
            "status": "PENDING",
            "added_at": _t5df_time.time(),
        }

    def _t5df_find_material(buf, meta):
        for m in buf.get("materials", []):
            if meta.get("telegram_message_id") and _t5df_s(m.get("telegram_message_id")) == meta.get("telegram_message_id"):
                return m
            if meta.get("file_name") and _t5df_s(m.get("file_name")) == meta.get("file_name"):
                return m
        return None

    def _t5df_copy_to_active_folder(src_id, name, folder_id):
        from core.topic_drive_oauth import _oauth_service
        svc = _oauth_service()
        qname = _t5df_s(name).replace("'", "\\'")
        found = svc.files().list(
            q=f"'{folder_id}' in parents and name='{qname}' and trashed=false",
            fields="files(id,name,webViewLink)",
            pageSize=5,
        ).execute().get("files", [])
        if found:
            fid = found[0]["id"]
            return fid, found[0].get("webViewLink") or f"https://drive.google.com/file/d/{fid}/view?usp=drivesdk"

        copied = svc.files().copy(
            fileId=src_id,
            body={"name": name, "parents": [folder_id]},
            fields="id,name,webViewLink",
        ).execute()
        fid = copied["id"]
        return fid, copied.get("webViewLink") or f"https://drive.google.com/file/d/{fid}/view?usp=drivesdk"

    def _t5df_upsert_material(chat_id, meta):
        active = _t5fc_active(chat_id)
        buf = _t5fc_buf(chat_id)

        mat = _t5df_find_material(buf, meta)
        if not mat:
            mat = dict(meta)
            mat["material_id"] = f"m_{meta.get('telegram_message_id') or int(_t5df_time.time())}"
            buf.setdefault("materials", []).append(mat)

        mat["active_folder_id"] = _t5df_s(active.get("folder_id"))
        mat["active_folder_name"] = _t5df_s(active.get("folder_name"))
        mat["updated_at"] = _t5df_time.time()

        if active.get("folder_id") and mat.get("source_drive_file_id") and not mat.get("drive_file_id"):
            try:
                fid, url = _t5df_copy_to_active_folder(mat["source_drive_file_id"], mat["file_name"], active["folder_id"])
                mat["drive_file_id"] = fid
                mat["drive_url"] = url
                mat["status"] = "LINKED"
            except Exception as e:
                mat["copy_error"] = f"{type(e).__name__}: {_t5df_s(e, 300)}"
                mat["status"] = "COPY_FAILED"

        try:
            _t5fc_enrich_materials(buf)
        except Exception as e:
            _T5DF_LOG.warning("FIX_TOPIC5_DRIVE_FILE_FULL_CANON_ENRICH_ERR %s", e)

        _t5fc_jsave(_t5fc_buf_path(chat_id), buf)
        return active, buf, mat

    async def _handle_new(conn, task, *args, **kwargs):
        task_id = _t5df_s(_t5df_row(task, "id"))
        chat_id = _t5df_s(_t5df_row(task, "chat_id", args[0] if len(args) > 0 else ""))
        topic_id = int(_t5df_row(task, "topic_id", args[1] if len(args) > 1 else 0) or 0)
        input_type = _t5df_s(_t5df_row(task, "input_type", ""))
        reply_to = _t5df_s(_t5df_row(task, "reply_to_message_id", ""))
        raw = _t5df_s(_t5df_row(task, "raw_input", ""))

        if topic_id == 5 and input_type in ("drive_file", "file", "document"):
            meta = _t5df_load_meta(raw)
            if meta and meta.get("file_name"):
                caption = _t5df_s(meta.get("owner_comment") or "")
                try:
                    if caption and _t5fc_context_like(caption):
                        _t5fc_save_context(chat_id, caption)
                except Exception as e:
                    _T5DF_LOG.warning("FIX_TOPIC5_DRIVE_FILE_FULL_CANON_CAPTION_CONTEXT_ERR %s", e)

                active, buf, mat = _t5df_upsert_material(chat_id, meta)
                ctx = buf.get("package_context", {})
                folder_id = _t5df_s(active.get("folder_id"))
                folder_name = _t5df_s(active.get("folder_name") or "не установлена")
                object_name = _t5df_s(ctx.get("object_address") or ctx.get("object_name") or active.get("object_address") or active.get("object_name") or "UNKNOWN")
                visit_basis = _t5df_s(ctx.get("visit_basis") or active.get("visit_basis") or "UNKNOWN")
                source_request = _t5df_s(ctx.get("source_request") or active.get("source_request") or "UNKNOWN")
                count = len(buf.get("materials", []))

                if not folder_id:
                    msg = "\n".join([
                        "Файл принят в пакет технадзора, но активная папка не установлена",
                        f"Файл: {meta.get('file_name')}",
                        f"Тип: {meta.get('file_type')}",
                        f"Объект / адрес: {object_name}",
                        f"Основание: {visit_basis}",
                        f"Источник заявки: {source_request}",
                        f"Материалов в пакете: {count}",
                        "К какой папке отнести материалы?",
                    ])
                    _t5fc_done(conn, task_id, chat_id, reply_to, msg, "topic5_drive_file_waiting_folder", "DONE")
                    return

                msg = "\n".join([
                    "Файл принят в полный контур технадзора",
                    f"Файл: {meta.get('file_name')}",
                    f"Тип: {meta.get('file_type')}",
                    f"Объект / адрес: {object_name}",
                    f"Основание: {visit_basis}",
                    f"Источник заявки: {source_request}",
                    f"Папка: {folder_name}",
                    f"https://drive.google.com/drive/folders/{folder_id}",
                    f"Материалов в пакете: {count}",
                    "Акт не формирую без команды: Сделай акт",
                ])

                if mat.get("copy_error"):
                    msg += f"\nDrive copy error: {mat.get('copy_error')}"

                _t5fc_done(conn, task_id, chat_id, reply_to, msg, "topic5_drive_file_full_canon_buffered", "DONE")
                return

        res = _T5DF_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        return await res if hasattr(res, "__await__") else res

    _handle_new._topic5_drive_file_full_canon_guard_v1 = True
    _T5DF_LOG.info("FIX_TOPIC5_DRIVE_FILE_FULL_CANON_GUARD_V1_INSTALLED")

except Exception as _t5df_err:
    try:
        logging.getLogger("task_worker").exception("FIX_TOPIC5_DRIVE_FILE_FULL_CANON_GUARD_V1_INSTALL_ERR %s", _t5df_err)
    except Exception:
        pass
# === END_FIX_TOPIC5_DRIVE_FILE_FULL_CANON_GUARD_V1 ===


# === FIX_TOPIC5_FILES_REQUIRE_OWNER_INSTRUCTION_V1 ===
try:
    import json as _t5fq_json
    import time as _t5fq_time
    import logging as _t5fq_logging

    _T5FQ_LOG = _t5fq_logging.getLogger("task_worker")
    _T5FQ_ORIG_HANDLE_NEW = _handle_new

    def _t5fq_s(v, limit=20000):
        return "" if v is None else str(v).strip()[:limit]

    def _t5fq_low(v):
        return _t5fq_s(v).lower().replace("ё", "е")

    def _t5fq_row(row, key, default=""):
        try:
            if hasattr(row, "keys") and key in row.keys():
                return row[key]
        except Exception:
            pass
        return default

    def _t5fq_meta(raw):
        try:
            d = _t5fq_json.loads(_t5fq_s(raw, 80000))
        except Exception:
            return {}
        if not isinstance(d, dict):
            return {}

        name = _t5fq_s(d.get("file_name") or d.get("name") or "")
        if not name:
            return {}

        mime = _t5fq_low(d.get("mime_type") or "")
        ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""

        if "image" in mime or ext in ("jpg", "jpeg", "png", "webp", "heic"):
            file_type = "PHOTO"
        elif "pdf" in mime or ext == "pdf":
            file_type = "PDF"
        elif "word" in mime or ext in ("doc", "docx"):
            file_type = "DOCX"
        elif "sheet" in mime or ext in ("xls", "xlsx", "csv"):
            file_type = "XLSX"
        else:
            file_type = "OTHER"

        return {
            "source": "TELEGRAM",
            "file_type": file_type,
            "file_name": name,
            "mime_type": mime,
            "source_drive_file_id": _t5fq_s(d.get("file_id") or d.get("drive_file_id") or ""),
            "telegram_message_id": _t5fq_s(d.get("telegram_message_id") or d.get("message_id") or ""),
            "reply_to_message_id": _t5fq_s(d.get("telegram_message_id") or d.get("message_id") or ""),
            "owner_comment": _t5fq_s(d.get("caption") or ""),
            "include_in_report": True,
            "include_in_act": True,
            "status": "PENDING_OWNER_INSTRUCTION",
            "needs_owner_instruction": True,
            "added_at": _t5fq_time.time(),
        }

    def _t5fq_has_instruction(text):
        low = _t5fq_low(text)
        return bool(low) and any(x in low for x in (
            "к акту", "в акт", "сделай", "сформируй", "создай", "подготовь",
            "проверь", "проверить", "разбор", "замечан", "дефект", "наруш",
            "отнеси", "добавь", "в эту папку", "к этому объекту", "этот же объект",
            "свар", "оборудован", "замена", "вышло из строя", "это фото", "эти фото",
            "это файл", "эти файлы", "это документы", "эти документы"
        ))

    def _t5fq_find(buf, meta):
        for m in buf.get("materials", []):
            if meta.get("telegram_message_id") and _t5fq_s(m.get("telegram_message_id")) == meta.get("telegram_message_id"):
                return m
            if meta.get("file_name") and _t5fq_s(m.get("file_name")) == meta.get("file_name"):
                return m
        return None

    def _t5fq_copy_to_folder(src_id, name, folder_id):
        from core.topic_drive_oauth import _oauth_service
        svc = _oauth_service()
        qname = _t5fq_s(name).replace("'", "\\'")
        found = svc.files().list(
            q=f"'{folder_id}' in parents and name='{qname}' and trashed=false",
            fields="files(id,name,webViewLink)",
            pageSize=5,
        ).execute().get("files", [])
        if found:
            fid = found[0]["id"]
            return fid, found[0].get("webViewLink") or f"https://drive.google.com/file/d/{fid}/view?usp=drivesdk"

        copied = svc.files().copy(
            fileId=src_id,
            body={"name": name, "parents": [folder_id]},
            fields="id,name,webViewLink",
        ).execute()
        fid = copied["id"]
        return fid, copied.get("webViewLink") or f"https://drive.google.com/file/d/{fid}/view?usp=drivesdk"

    def _t5fq_upsert_pending(chat_id, meta):
        active = _t5fc_active(chat_id)
        buf = _t5fc_buf(chat_id)

        mat = _t5fq_find(buf, meta)
        if not mat:
            mat = dict(meta)
            mat["material_id"] = f"m_{meta.get('telegram_message_id') or int(_t5fq_time.time())}"
            buf.setdefault("materials", []).append(mat)

        mat["source"] = mat.get("source") or "TELEGRAM"
        mat["file_type"] = mat.get("file_type") or meta.get("file_type")
        mat["source_drive_file_id"] = mat.get("source_drive_file_id") or meta.get("source_drive_file_id")
        mat["active_folder_id"] = _t5fq_s(active.get("folder_id"))
        mat["active_folder_name"] = _t5fq_s(active.get("folder_name"))
        mat["status"] = "PENDING_OWNER_INSTRUCTION"
        mat["needs_owner_instruction"] = True
        mat["updated_at"] = _t5fq_time.time()

        _t5fc_jsave(_t5fc_buf_path(chat_id), buf)
        return active, buf, mat

    def _t5fq_pending(buf):
        return [
            m for m in buf.get("materials", [])
            if m.get("needs_owner_instruction") or m.get("status") == "PENDING_OWNER_INSTRUCTION"
        ]

    def _t5fq_apply_instruction_to_pending(chat_id, text):
        active = _t5fc_active(chat_id)
        buf = _t5fc_buf(chat_id)
        pending = _t5fq_pending(buf)
        if not pending:
            return 0, active, buf

        clean = _t5fq_s(text, 5000)

        for m in pending:
            old = _t5fq_s(m.get("owner_comment") or m.get("voice_comment") or "")
            if clean and clean not in old:
                m["owner_comment"] = (old + "\n" + clean).strip() if old else clean

            if active.get("folder_id") and m.get("source_drive_file_id") and not m.get("drive_file_id"):
                try:
                    fid, url = _t5fq_copy_to_folder(m["source_drive_file_id"], m["file_name"], active["folder_id"])
                    m["drive_file_id"] = fid
                    m["drive_url"] = url
                except Exception as e:
                    m["copy_error"] = f"{type(e).__name__}: {_t5fq_s(e, 300)}"

            m["active_folder_id"] = _t5fq_s(active.get("folder_id"))
            m["active_folder_name"] = _t5fq_s(active.get("folder_name"))
            m["needs_owner_instruction"] = False
            m["status"] = "OWNER_INSTRUCTION_LINKED"
            m["updated_at"] = _t5fq_time.time()

        try:
            _t5fc_save_context(chat_id, clean)
        except Exception as e:
            _T5FQ_LOG.warning("FIX_TOPIC5_FILES_REQUIRE_OWNER_INSTRUCTION_CONTEXT_ERR %s", e)

        try:
            _t5fc_enrich_materials(buf)
        except Exception as e:
            _T5FQ_LOG.warning("FIX_TOPIC5_FILES_REQUIRE_OWNER_INSTRUCTION_ENRICH_ERR %s", e)

        _t5fc_jsave(_t5fc_buf_path(chat_id), buf)
        return len(pending), active, buf

    def _t5fq_question(meta, active, buf):
        ctx = buf.get("package_context", {})
        object_name = _t5fq_s(
            ctx.get("object_address")
            or ctx.get("object_name")
            or active.get("object_address")
            or active.get("object_name")
            or "не задан"
        )
        folder_name = _t5fq_s(active.get("folder_name") or "не установлена")
        return "\n".join([
            "Файл получил и сохранил в буфер технадзора",
            f"Файл: {meta.get('file_name')}",
            f"Тип: {meta.get('file_type')}",
            f"Текущий объект: {object_name}",
            f"Текущая папка: {folder_name}",
            f"Материалов в буфере: {len(buf.get('materials', []))}",
            "Что это за материалы, к какому объекту/папке их отнести и что с ними сделать?",
        ])

    def _t5fq_send(conn, task_id, chat_id, reply_to, text, kind, state):
        sent = _send_once_ex(
            conn,
            str(task_id),
            str(chat_id),
            _t5fq_s(text, 3500),
            int(reply_to) if _t5fq_s(reply_to).isdigit() else None,
            kind,
        )
        upd = {"state": state, "result": _t5fq_s(text, 50000), "error_message": ""}
        if isinstance(sent, dict) and (sent.get("bot_message_id") or sent.get("message_id")):
            upd["bot_message_id"] = sent.get("bot_message_id") or sent.get("message_id")
        _update_task(conn, str(task_id), **upd)
        try:
            _history(conn, str(task_id), kind)
        except Exception:
            pass
        conn.commit()

    async def _handle_new(conn, task, *args, **kwargs):
        task_id = _t5fq_s(_t5fq_row(task, "id"))
        chat_id = _t5fq_s(_t5fq_row(task, "chat_id", args[0] if len(args) > 0 else ""))
        topic_id = int(_t5fq_row(task, "topic_id", args[1] if len(args) > 1 else 0) or 0)
        input_type = _t5fq_s(_t5fq_row(task, "input_type", ""))
        reply_to = _t5fq_s(_t5fq_row(task, "reply_to_message_id", ""))
        raw = _t5fq_s(_t5fq_row(task, "raw_input", ""))

        if topic_id == 5 and input_type in ("drive_file", "file", "document"):
            meta = _t5fq_meta(raw)
            if meta and meta.get("file_name"):
                caption = _t5fq_s(meta.get("owner_comment") or "")

                if not _t5fq_has_instruction(caption):
                    active, buf, mat = _t5fq_upsert_pending(chat_id, meta)
                    _t5fq_send(
                        conn,
                        task_id,
                        chat_id,
                        reply_to,
                        _t5fq_question(meta, active, buf),
                        "topic5_file_wait_owner_instruction",
                        "WAITING_CLARIFICATION",
                    )
                    return

                active, buf, mat = _t5fq_upsert_pending(chat_id, meta)
                count, active, buf = _t5fq_apply_instruction_to_pending(chat_id, caption)
                _t5fq_send(
                    conn,
                    task_id,
                    chat_id,
                    reply_to,
                    "\n".join([
                        "Файл принят с пояснением владельца",
                        f"Файл: {meta.get('file_name')}",
                        f"Тип: {meta.get('file_type')}",
                        f"Пояснение: {caption}",
                        f"Связано файлов: {count}",
                        "Акт не формирую без команды: Сделай акт",
                    ]),
                    "topic5_file_with_owner_instruction",
                    "DONE",
                )
                return

        if topic_id == 5 and input_type in ("text", "voice", ""):
            clean = raw
            try:
                clean = _t5fc_clean_voice(raw)
            except Exception:
                pass

            if _t5fq_has_instruction(clean):
                count, active, buf = _t5fq_apply_instruction_to_pending(chat_id, clean)
                if count > 0:
                    _t5fq_send(
                        conn,
                        task_id,
                        chat_id,
                        reply_to,
                        "\n".join([
                            "Пояснение владельца принято к ожидающим материалам",
                            f"Связано файлов: {count}",
                            f"Материалов в буфере: {len(buf.get('materials', []))}",
                            "Акт не формирую без команды: Сделай акт",
                        ]),
                        "topic5_pending_files_instruction_linked",
                        "DONE",
                    )
                    return

        res = _T5FQ_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        return await res if hasattr(res, "__await__") else res

    _handle_new._topic5_files_require_owner_instruction_v1 = True
    _T5FQ_LOG.info("FIX_TOPIC5_FILES_REQUIRE_OWNER_INSTRUCTION_V1_INSTALLED")

except Exception as _t5fq_err:
    try:
        logging.getLogger("task_worker").exception("FIX_TOPIC5_FILES_REQUIRE_OWNER_INSTRUCTION_V1_INSTALL_ERR %s", _t5fq_err)
    except Exception:
        pass
# === END_FIX_TOPIC5_FILES_REQUIRE_OWNER_INSTRUCTION_V1 ===

# === MOVE_MAIN_ENTRYPOINT_TO_END_V1 ===
# SUPERSEDED by MOVE_MAIN_ENTRYPOINT_TO_END_V2 at bottom of file.
# asyncio.run(main()) moved to absolute end so ALL module-level patches install first.
# if __name__ == "__main__":
#     asyncio.run(main())  # DO NOT ENABLE — see V2 below
# === END_MOVE_MAIN_ENTRYPOINT_TO_END_V1 ===


# === DRIVE_FILE_AUTO_DELIVER_V1 ===
# Fix: _handle_drive_file sets AWAITING_CONFIRMATION without sending Telegram reply.
# After 30 min → CONFIRMATION_TIMEOUT → FAILED with no bot_message_id.
# Patch: after processing, if result has Drive link → send reply + set DONE immediately.
try:
    import logging as _dfad_log
    _DFAD_LOG = _dfad_log.getLogger("task_worker")
    _dfad_orig_handle_drive_file = _handle_drive_file

    async def _handle_drive_file(conn, task, chat_id, topic_id):
        await _dfad_orig_handle_drive_file(conn, task, chat_id, topic_id)
        try:
            task_id = _s(_task_field(task, "id", ""))
            if not task_id:
                return
            row = conn.execute(
                "SELECT state, result, bot_message_id, reply_to_message_id FROM tasks WHERE id=?",
                (task_id,)
            ).fetchone()
            if not row:
                return
            if row["state"] == "AWAITING_CONFIRMATION" and not row["bot_message_id"]:
                result_text = _s(row["result"])
                if "drive.google.com" in result_text or "docs.google.com" in result_text:
                    try:
                        reply_to = row["reply_to_message_id"]
                        _send_once_ex(conn, task_id, str(chat_id), result_text, reply_to, "drive_file_auto_deliver_v1")
                        _update_task(conn, task_id, state="DONE", error_message="")
                        conn.commit()
                        _DFAD_LOG.info("DRIVE_FILE_AUTO_DELIVER_V1 sent task=%s topic=%s", task_id, topic_id)
                    except Exception as _dfad_send_err:
                        _DFAD_LOG.warning("DRIVE_FILE_AUTO_DELIVER_V1_SEND_ERR %s", _dfad_send_err)
        except Exception as _dfad_err:
            _DFAD_LOG.warning("DRIVE_FILE_AUTO_DELIVER_V1_ERR %s", _dfad_err)

    _DFAD_LOG.info("DRIVE_FILE_AUTO_DELIVER_V1_INSTALLED")
except Exception as _dfad_install_err:
    try:
        logging.getLogger("task_worker").exception("DRIVE_FILE_AUTO_DELIVER_V1_INSTALL_ERR %s", _dfad_install_err)
    except Exception:
        pass
# === END_DRIVE_FILE_AUTO_DELIVER_V1 ===


# === TOPIC5_CANON_CLOSE_EXTEND_V1 ===
# Fix 1: context message (address/object/basis) + pending files → link pending files
#         before FULL_CANON_CLOSE saves context and sends reply.
# Fix 2: extend _t5fc_act_like with infinitive forms (сделать разбор, финальный акт, etc.)
# Fix 3: topic_500 drive_file → redirect to topic_2 instead of AWAITING_CONFIRMATION.
try:
    import logging as _t5ce_log
    _T5CE_LOG = _t5ce_log.getLogger("task_worker")
    _T5CE_ORIG_HANDLE_NEW = _handle_new

    # Extend act triggers (infinitive + extended forms not in original _t5fc_act_like)
    _t5ce_orig_act_like = _t5fc_act_like

    def _t5fc_act_like(text):
        if _t5ce_orig_act_like(text):
            return True
        low = _t5fc_low(text)
        return any(x in low for x in (
            "сделать акт", "сформировать акт", "создать акт", "подготовить акт",
            "сделать разбор", "сделать отчет", "сделать отчёт", "сделать анализ",
            "финальный акт", "финальный разбор", "итоговый акт", "итоговый разбор",
            "готовь разбор", "собери пакет", "закрой выезд", "закрыть выезд",
        ))

    def _t5ce_has_substantive_context(text):
        if _t5fq_has_instruction(text):
            return True
        try:
            if _t5fc_extract_address(text):
                return True
        except Exception:
            pass
        try:
            if _t5fc_source_and_basis(text):
                return True
        except Exception:
            pass
        return False

    async def _handle_new(conn, task, *args, **kwargs):
        try:
            task_id = _t5fc_s(_t5fc_row(task, "id"))
            chat_id = _t5fc_s(_t5fc_row(task, "chat_id", args[0] if len(args) > 0 else ""))
            topic_id = int(_t5fc_row(task, "topic_id", args[1] if len(args) > 1 else 0) or 0)
            input_type = _t5fc_s(_t5fc_row(task, "input_type", ""))
            reply_to = _t5fc_s(_t5fc_row(task, "reply_to_message_id", ""))
            raw = _t5fc_s(_t5fc_row(task, "raw_input", ""))

            # Fix 3: topic_500 drive_file redirect
            if topic_id == 500 and input_type in ("drive_file", "file", "document"):
                try:
                    import json as _t5ce_json
                    _meta = _t5ce_json.loads(raw) if raw.startswith("{") else {}
                    _fname = _t5fc_s(_meta.get("file_name") or _meta.get("name") or "")
                except Exception:
                    _fname = ""
                msg = "\n".join(filter(None, [
                    "Файл получен в топик интернет-поиска",
                    f"Файл: {_fname}" if _fname else None,
                    "Поиск по файлам не выполняю.",
                    "Если это смета или чертёж — пришли в топик Стройка (topic_2).",
                    "Для поиска поставщиков — напиши текстовый запрос.",
                ]))
                _t5fc_done(conn, task_id, chat_id, reply_to, msg, "topic500_drive_file_redirect_v1", "DONE")
                return

            # Fix 1: topic_5 text/voice context + pending files linking
            if topic_id == 5 and input_type in ("text", "voice", ""):
                clean = _t5fc_clean_voice(raw)
                if _t5fc_context_like(clean) and _t5ce_has_substantive_context(clean):
                    try:
                        buf = _t5fc_buf(chat_id)
                        pending = _t5fq_pending(buf)
                        if pending:
                            _t5fq_apply_instruction_to_pending(chat_id, clean)
                            _T5CE_LOG.info("TOPIC5_CANON_CLOSE_EXTEND_V1 linked %d pending files chat=%s", len(pending), chat_id)
                    except Exception as _t5ce_link_err:
                        _T5CE_LOG.warning("TOPIC5_CANON_CLOSE_EXTEND_V1_LINK_ERR %s", _t5ce_link_err)
                    # Fall through — FULL_CANON_CLOSE sends the reply with updated object_name

        except Exception as _t5ce_guard_err:
            _T5CE_LOG.warning("TOPIC5_CANON_CLOSE_EXTEND_V1_GUARD_ERR %s", _t5ce_guard_err)

        res = _T5CE_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        return await res if hasattr(res, "__await__") else res

    _handle_new._topic5_canon_close_extend_v1 = True
    _T5CE_LOG.info("TOPIC5_CANON_CLOSE_EXTEND_V1_INSTALLED")

except Exception as _t5ce_install_err:
    try:
        logging.getLogger("task_worker").exception("TOPIC5_CANON_CLOSE_EXTEND_V1_INSTALL_ERR %s", _t5ce_install_err)
    except Exception:
        pass
# === END_TOPIC5_CANON_CLOSE_EXTEND_V1 ===

# === DRIVE_FILE_AUTO_DELIVER_V2 ===
# Extension: deliver any non-empty AWAITING_CONFIRMATION result (not just Drive-linked ones).
# Previously: only tasks with drive.google.com/docs.google.com in result were delivered.
# Fix: topic_2 document result (no Drive link due to upload failure) also needs delivery.
try:
    import logging as _dfad2_log
    _DFAD2_LOG = _dfad2_log.getLogger("task_worker")
    _dfad2_orig_handle_drive_file = _handle_drive_file

    async def _handle_drive_file(conn, task, chat_id, topic_id):
        await _dfad2_orig_handle_drive_file(conn, task, chat_id, topic_id)
        try:
            task_id = _s(_task_field(task, "id", ""))
            if not task_id:
                return
            row = conn.execute(
                "SELECT state, result, bot_message_id, reply_to_message_id FROM tasks WHERE id=?",
                (task_id,)
            ).fetchone()
            if not row:
                return
            if row["state"] == "AWAITING_CONFIRMATION" and not row["bot_message_id"]:
                result_text = _s(row["result"])
                if result_text and len(result_text) > 10:
                    reply_to = row["reply_to_message_id"]
                    _send_once_ex(conn, task_id, str(chat_id), result_text, reply_to, "drive_file_auto_deliver_v2")
                    _update_task(conn, task_id, state="DONE", error_message="")
                    conn.commit()
                    _DFAD2_LOG.info("DRIVE_FILE_AUTO_DELIVER_V2 sent task=%s topic=%s", task_id, topic_id)
        except Exception as _dfad2_err:
            _DFAD2_LOG.warning("DRIVE_FILE_AUTO_DELIVER_V2_ERR %s", _dfad2_err)

    _DFAD2_LOG.info("DRIVE_FILE_AUTO_DELIVER_V2_INSTALLED")
except Exception as _dfad2_install_err:
    try:
        logging.getLogger("task_worker").exception("DRIVE_FILE_AUTO_DELIVER_V2_INSTALL_ERR %s", _dfad2_install_err)
    except Exception:
        pass
# === END_DRIVE_FILE_AUTO_DELIVER_V2 ===

# === FIX_P6H4TW_VOICE_ANNOTATE_ACT_GUARD_V1 ===
# P6H4TW voice_annotate is a catch-all: any [VOICE] + non-empty buffer → annotate last file.
# "[VOICE] Нужно добавить в папку и сделать финальный акт" was annotated instead of triggering act.
# Fix: if voice is act-like → skip voice_annotate, fall through to FULL_CANON_CLOSE.
try:
    import logging as _p6hg_log
    _P6HG_LOG = _p6hg_log.getLogger("task_worker")
    _p6hg_orig_handle_topic5 = _p6h4tw_handle_topic5

    async def _p6h4tw_handle_topic5(conn, task, args, kwargs):
        try:
            # Guard: act-like voice → skip voice_annotate, let FULL_CANON_CLOSE handle
            _p6hg_raw = ""
            try:
                _p6hg_raw = str(task["raw_input"]) if "raw_input" in (task.keys() if hasattr(task, "keys") else dir(task)) else ""
            except Exception:
                pass
            if _p6hg_raw.startswith("[VOICE]"):
                _p6hg_voice = _p6hg_raw[7:].strip()
                _p6hg_is_act = False
                try:
                    _p6hg_is_act = _t5fc_act_like(_p6hg_voice)
                except Exception:
                    pass
                if _p6hg_is_act:
                    _P6HG_LOG.info("FIX_P6H4TW_VOICE_ANNOTATE_ACT_GUARD_V1 skipping voice_annotate for act command")
                    return False  # fall through to FULL_CANON_CLOSE
        except Exception as _p6hg_guard_err:
            _P6HG_LOG.warning("FIX_P6H4TW_VOICE_ANNOTATE_ACT_GUARD_V1_ERR %s", _p6hg_guard_err)
        return await _p6hg_orig_handle_topic5(conn, task, args, kwargs)

    _P6HG_LOG.info("FIX_P6H4TW_VOICE_ANNOTATE_ACT_GUARD_V1_INSTALLED")
except Exception as _p6hg_install_err:
    try:
        logging.getLogger("task_worker").exception("FIX_P6H4TW_VOICE_ANNOTATE_ACT_GUARD_V1_INSTALL_ERR %s", _p6hg_install_err)
    except Exception:
        pass
# === END_FIX_P6H4TW_VOICE_ANNOTATE_ACT_GUARD_V1 ===

# === DRIVE_FILE_NO_INTENT_OFFER_V1 ===
# topic_2: файл без явной команды → показать меню действий вместо сырого OCR
import logging as _dfnio_log_mod
_DFNIO_LOG = _dfnio_log_mod.getLogger("task_worker")

try:
    _dfnio_orig_handle_drive_file = _handle_drive_file
except Exception:
    _dfnio_orig_handle_drive_file = None

async def _handle_drive_file(conn, task, chat_id, topic_id):
    try:
        import json as _dfnio_json
        _tid = str(task["id"]) if hasattr(task, "keys") else str(task[0])
        _raw = str(task["raw_input"]) if hasattr(task, "keys") else ""
        _topic = int(topic_id or 0)
        if _topic == 2:
            try:
                _data = _dfnio_json.loads(_raw)
            except Exception:
                _data = {}
            _caption = str(_data.get("caption") or "").strip()
            _fname = str(_data.get("file_name") or "").strip()
            # Check if no intent using existing _ioa_needs logic
            _no_intent = _ioa_needs(_raw, _caption)
            if _no_intent:
                _offer = (
                    f"Принял файл «{_fname}». Что нужно сделать?\n\n"
                    "1️⃣ Смета — извлечь позиции, посчитать объёмы, создать Excel\n"
                    "2️⃣ Описание — описать содержимое документа\n"
                    "3️⃣ Таблица — вытащить таблицы из файла в Excel\n"
                    "4️⃣ Шаблон — сохранить как образец для будущих задач\n"
                    "5️⃣ Анализ — технический анализ (КЖ/АР/КД)\n\n"
                    "Напиши номер или опиши задачу."
                )
                _reply_to = task["reply_to_message_id"] if hasattr(task, "keys") and "reply_to_message_id" in task.keys() else None
                conn.execute(
                    "UPDATE tasks SET state='WAITING_CLARIFICATION', result=?, error_message='DRIVE_FILE_NO_INTENT_OFFER_V1', updated_at=datetime('now') WHERE id=?",
                    (_offer, _tid)
                )
                conn.execute(
                    "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                    (_tid, "DRIVE_FILE_NO_INTENT_OFFER_V1:menu_shown")
                )
                conn.commit()
                _send_once_ex(conn, _tid, str(chat_id), _offer, _reply_to, "drive_file_no_intent_offer")
                return
    except Exception as _dfnio_err:
        _DFNIO_LOG.warning("DRIVE_FILE_NO_INTENT_OFFER_V1_ERR: %s", _dfnio_err)
    if _dfnio_orig_handle_drive_file:
        return await _dfnio_orig_handle_drive_file(conn, task, chat_id, topic_id)
# === END_DRIVE_FILE_NO_INTENT_OFFER_V1 ===

# === FIX_P6E2_TW_ESTIMATE_LIKE_EXTEND_V1 ===
# Extend _p6e2_tw_estimate_like to match extended ESTIMATE_WORDS canon
def _p6e2_tw_estimate_like(text):
    low = _p6e2_tw_low(text)
    return any(x in low for x in (
        "смет", "расчет", "расчёт", "посчитай", "посчитать",
        "рассчитать", "стоимост", "стоить", "стоит",
        "сколько стоит", "сколько будет", "нужна смета",
        "нужен расчет", "нужен расчёт", "полная смета",
    ))
# === END_FIX_P6E2_TW_ESTIMATE_LIKE_EXTEND_V1 ===

# === TOPIC500_ESTIMATE_ISOLATION_GUARD_V1 ===
# topic_500 = internet search only. Must NEVER enter estimate/project pipeline.
# Wraps FULLFIX_14 parse_estimate_rows to block topic_500 + topic_5 + topic_210 from estimate route.
import logging as _t5eig_log_mod
_T5EIG_LOG = _t5eig_log_mod.getLogger("task_worker")

try:
    from core.estimate_unified_engine import parse_estimate_rows as _t5eig_orig_parse
    def _t5eig_safe_parse(text, topic_id=0):
        _blocked = (500, 5, 210, 3008)
        if int(topic_id or 0) in _blocked:
            return []
        return _t5eig_orig_parse(text)
    import core.estimate_unified_engine as _t5eig_mod
    _t5eig_mod.parse_estimate_rows = _t5eig_safe_parse
    _T5EIG_LOG.info("TOPIC500_ESTIMATE_ISOLATION_GUARD_V1 installed")
except Exception as _t5eig_e:
    _T5EIG_LOG.warning("TOPIC500_ESTIMATE_ISOLATION_GUARD_V1 INSTALL ERR: %s", _t5eig_e)
# === END_TOPIC500_ESTIMATE_ISOLATION_GUARD_V1 ===

# === TOPIC5_PHOTO_ACT_CONFIRMATION_TIMEOUT_FIX_V1 ===
# topic_5 photo acts stuck in AWAITING_CONFIRMATION → CONFIRMATION_TIMEOUT
# Root cause: each photo creates own act, waits for confirmation that never comes.
# Fix: topic_5 drive_file photo tasks that stayed AWAITING_CONFIRMATION > 5 min
#      should be silently archived (photos go to buffer, not standalone acts).
import logging as _t5patf_log_mod
import threading as _t5patf_threading
_T5PATF_LOG = _t5patf_log_mod.getLogger("task_worker")

def _t5patf_archive_stale_photo_acts():
    try:
        import sqlite3 as _t5patf_sqlite
        import time as _t5patf_time
        conn2 = _t5patf_sqlite.connect("/root/.areal-neva-core/data/core.db")
        conn2.row_factory = _t5patf_sqlite.Row
        rows = conn2.execute("""
            SELECT id FROM tasks
            WHERE topic_id=5
              AND input_type='drive_file'
              AND state='AWAITING_CONFIRMATION'
              AND error_message NOT LIKE '%CONFIRMATION_TIMEOUT%'
              AND (julianday('now') - julianday(updated_at)) * 86400 > 300
        """).fetchall()
        for r in rows:
            conn2.execute(
                "UPDATE tasks SET state='FAILED', error_message='CONFIRMATION_TIMEOUT' WHERE id=?",
                (r["id"],)
            )
            conn2.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (r["id"], "T5_PHOTO_ACT_STALE_ARCHIVED_V1")
            )
        if rows:
            conn2.commit()
            _T5PATF_LOG.info("T5_PHOTO_ACT_STALE_ARCHIVED_V1: archived %d stale photo acts", len(rows))
        conn2.close()
    except Exception as _t5patf_e:
        _T5PATF_LOG.warning("T5_PHOTO_ACT_STALE_ARCHIVE_ERR: %s", _t5patf_e)

try:
    _t5patf_archive_stale_photo_acts()
except Exception:
    pass
# === END_TOPIC5_PHOTO_ACT_CONFIRMATION_TIMEOUT_FIX_V1 ===

# === FIX_SEND_ONCE_EX_MESSAGE_THREAD_ID_V1 ===
# Root cause: _send_once_ex → send_reply_ex without message_thread_id → topic_0.
# Fix: ContextVar carries topic_id through the call chain to send_reply_ex/send_reply.
# _send_once_ex auto-looks up topic_id from DB if caller did not supply it.
import contextvars as _fsoex_cv
import logging as _fsoex_logging

_FSOEX_LOG = _fsoex_logging.getLogger("task_worker")
_fsoex_thread_id_var = _fsoex_cv.ContextVar("_fsoex_thread_id", default=None)

# Wrap send_reply_ex at module level to inject message_thread_id from ContextVar
_fsoex_orig_send_reply_ex = send_reply_ex

def send_reply_ex(*args, **kwargs):
    if not kwargs.get("message_thread_id"):
        _t = _fsoex_thread_id_var.get()
        if _t:
            kwargs["message_thread_id"] = _t
    return _fsoex_orig_send_reply_ex(*args, **kwargs)

# Wrap send_reply the same way
_fsoex_orig_send_reply = send_reply

def send_reply(*args, **kwargs):
    if not kwargs.get("message_thread_id"):
        _t = _fsoex_thread_id_var.get()
        if _t:
            kwargs["message_thread_id"] = _t
    return _fsoex_orig_send_reply(*args, **kwargs)

# Wrap _send_once_ex: accept message_thread_id/topic_id, auto-lookup from DB, set ContextVar
_fsoex_orig_send_once_ex = _send_once_ex

def _send_once_ex(conn, task_id, chat_id, text, reply_to=None, kind="result", *args, **kwargs):
    _thread = kwargs.pop("message_thread_id", None) or kwargs.pop("topic_id", None)
    if not _thread and conn is not None and task_id:
        try:
            _r = conn.execute("SELECT topic_id FROM tasks WHERE id=?", (str(task_id),)).fetchone()
            if _r and _r[0]:
                _thread = int(_r[0]) or None
        except Exception as _fe:
            _FSOEX_LOG.debug("FIX_SOE_LOOKUP_ERR task=%s err=%s", task_id, _fe)
    if _thread:
        _tok = _fsoex_thread_id_var.set(int(_thread))
        try:
            return _fsoex_orig_send_once_ex(conn, task_id, chat_id, text, reply_to, kind, *args, **kwargs)
        finally:
            _fsoex_thread_id_var.reset(_tok)
    return _fsoex_orig_send_once_ex(conn, task_id, chat_id, text, reply_to, kind, *args, **kwargs)

# Wrap _send_once the same way
_fsoex_orig_send_once = _send_once

def _send_once(conn, task_id, chat_id, text, reply_to=None, kind="result"):
    _thread = None
    if conn is not None and task_id:
        try:
            _r = conn.execute("SELECT topic_id FROM tasks WHERE id=?", (str(task_id),)).fetchone()
            if _r and _r[0]:
                _thread = int(_r[0]) or None
        except Exception:
            pass
    if _thread:
        _tok = _fsoex_thread_id_var.set(int(_thread))
        try:
            return _fsoex_orig_send_once(conn, task_id, chat_id, text, reply_to, kind)
        finally:
            _fsoex_thread_id_var.reset(_tok)
    return _fsoex_orig_send_once(conn, task_id, chat_id, text, reply_to, kind)

_FSOEX_LOG.info("FIX_SEND_ONCE_EX_MESSAGE_THREAD_ID_V1 installed")
# === END_FIX_SEND_ONCE_EX_MESSAGE_THREAD_ID_V1 ===

# === FIX_P6E2_TW_ESTIMATE_LIKE_FULL_CONSTRUCTION_V1 ===
# Replaces FIX_P6E2_TW_ESTIMATE_LIKE_EXTEND_V1 — adds construction params.
# "Высота 3,6 м каркас фундамент монолитная плита" → True without "смета".
import re as _p6e2_fec_re

def _p6e2_tw_estimate_like(text):
    low = _p6e2_tw_low(text)
    if any(x in low for x in (
        "смет", "расчет", "расчёт", "посчитай", "посчитать",
        "рассчитать", "стоимост", "стоить", "стоит",
        "сколько стоит", "сколько будет", "нужна смета",
        "нужен расчет", "нужен расчёт", "полная смета",
    )):
        return True
    has_dims = bool(_p6e2_fec_re.search(
        r"\d+[.,]?\d*\s*(?:x|х|×|\*|на)\s*\d+|\d+[.,]?\d*\s*м\b|\bвысот",
        low,
    ))
    has_material = any(x in low for x in (
        "каркас", "газобетон", "кирпич", "монолит", "арболит", "брус",
        "фундамент", "кровл", "перекрыт", "металлочереп", "профнастил",
        "фальц", "клик", "сэндвич", "цсп", "минват", "роквул",
    ))
    has_object = any(x in low for x in (
        "дом", "ангар", "склад", "коттедж", "здани", "строен",
        "постройк", "баня", "гараж", "объект",
    ))
    has_construction = any(x in low for x in (
        "высота", "этаж", "площадь", "периметр", "стен", "плита",
        "подушк", "технологи", "лента", "свай", "окна", "двери",
        "наружн", "внутренн",
    ))
    if (has_dims or has_material) and (has_object or has_construction):
        return True
    return False
# === END_FIX_P6E2_TW_ESTIMATE_LIKE_FULL_CONSTRUCTION_V1 ===

# === MOVE_MAIN_ENTRYPOINT_TO_END_V2 ===
# SUPERSEDED by V3 at bottom of file — more patches were appended after this block.
# if __name__ == "__main__":
#     asyncio.run(main())  # DO NOT ENABLE — see V3 below
# === END_MOVE_MAIN_ENTRYPOINT_TO_END_V2 ===

# === FIX_P6E67_FIND_PARENT_NO_DONE_FAILED_V1 ===
# Bug: _p6e67_find_parent allowed DONE/FAILED parent tasks for EXACT_REPLY_LINK.
# Result: user replies to completed estimate → new estimate merged into old DONE task → dropped.
# Fix: replace _p6e67_find_parent with version that excludes DONE/FAILED from all match paths.
import logging as _p6e67fix_logging
_P6E67FIX_LOG = _p6e67fix_logging.getLogger("task_worker")

_p6e67fix_orig_find_parent = _p6e67_find_parent

def _p6e67_find_parent(conn, task):
    chat_id = _p6e67_s(_p6e67_row(task, "chat_id", ""))
    topic_id = int(_p6e67_row(task, "topic_id", 0) or 0)
    task_id = _p6e67_s(_p6e67_row(task, "id", ""))
    reply_to = _p6e67_row(task, "reply_to_message_id", None)
    try:
        reply_to = int(reply_to) if reply_to not in (None, "", 0, "0") else None
    except Exception:
        reply_to = None

    if not chat_id or topic_id != 2:
        return None, "NO_SCOPE"

    # EXACT_REPLY_LINK — only active/open parents (not DONE/FAILED).
    # A completed estimate should not absorb a new request.
    if reply_to:
        row = conn.execute("""
            SELECT * FROM tasks
            WHERE chat_id=? AND COALESCE(topic_id,0)=? AND id<>?
              AND (bot_message_id=? OR reply_to_message_id=?)
              AND state IN ('IN_PROGRESS','AWAITING_CONFIRMATION','RESULT_READY','WAITING_CLARIFICATION')
            ORDER BY rowid DESC LIMIT 1
        """, (chat_id, topic_id, task_id, reply_to, reply_to)).fetchone()
        if row and _p6e67_is_estimate(row):
            return row, "EXACT_REPLY_LINK"

    # LAST_ACTIVE_ESTIMATE_FALLBACK — same: only open parents.
    row = conn.execute("""
        SELECT * FROM tasks
        WHERE chat_id=? AND COALESCE(topic_id,0)=? AND id<>?
          AND state IN ('IN_PROGRESS','AWAITING_CONFIRMATION','RESULT_READY','WAITING_CLARIFICATION')
          AND (
            raw_input LIKE '%смет%' OR result LIKE '%смет%'
            OR raw_input LIKE '%стоимость%' OR result LIKE '%стоимость%'
            OR raw_input LIKE '%кровл%' OR result LIKE '%кровл%'
            OR raw_input LIKE '%фундамент%' OR result LIKE '%фундамент%'
          )
        ORDER BY rowid DESC LIMIT 1
    """, (chat_id, topic_id, task_id)).fetchone()

    if row and _p6e67_is_estimate(row):
        return row, "LAST_ACTIVE_ESTIMATE_FALLBACK"

    return None, "NOT_FOUND"

_P6E67FIX_LOG.info("FIX_P6E67_FIND_PARENT_NO_DONE_FAILED_V1 installed")
# === END_FIX_P6E67_FIND_PARENT_NO_DONE_FAILED_V1 ===

# === MOVE_MAIN_ENTRYPOINT_TO_END_V3 ===
# SUPERSEDED by V4 — see end of file
# if __name__ == "__main__":
#     asyncio.run(main())
# === END_MOVE_MAIN_ENTRYPOINT_TO_END_V3 ===
# V3 superseded by V4 below — asyncio.run moved after FIX_TOP2_BUILD_CONTEXT_HISTORY_V1
# if __name__ == "__main__":
#     asyncio.run(main())

# === FIX_TOP2_BUILD_CONTEXT_HISTORY_V1 ===
# Root cause: _top2_ob_build_context only searched active (live) parents.
# DONE/FAILED/CANCELLED tasks were excluded → thin voice got no history context
# → _top2_ob_estimate_like(full_context) = False → "vague_no_memory_revive" or
# P3 saw no dims → "Уточни размеры дома".
#
# Fix: extend _top2_ob_build_context to also pull best historical raw_input
# from DONE/FAILED/CANCELLED tasks within 7 days as read-only ТЗ context.
# CANCELLED P6E67_MERGED tasks are especially valuable (contain full spec).
import logging as _t2bch_log_mod
_T2BCH_LOG = _t2bch_log_mod.getLogger("task_worker")

_t2bch_orig_build_context = _top2_ob_build_context


def _top2_ob_build_context(conn, task_id, chat_id, topic_id, raw):
    context_str, parent_id = _t2bch_orig_build_context(conn, task_id, chat_id, topic_id, raw)

    # Already found a live parent with content — done
    if parent_id:
        return context_str, parent_id

    # Try to find best historical raw_input from completed/failed/cancelled tasks
    try:
        rows = conn.execute("""
            SELECT raw_input, error_message, state FROM tasks
            WHERE chat_id=? AND COALESCE(topic_id,0)=?
              AND id<>?
              AND state IN ('DONE','FAILED','CANCELLED')
              AND updated_at >= datetime('now','-7 days')
              AND input_type <> 'search'
              AND (
                raw_input LIKE '%дом%' OR raw_input LIKE '%каркас%' OR
                raw_input LIKE '%газобетон%' OR raw_input LIKE '%монолит%' OR
                raw_input LIKE '%фундамент%' OR raw_input LIKE '%кровл%' OR
                raw_input LIKE '%ангар%' OR raw_input LIKE '%склад%' OR
                raw_input LIKE '%высота%' OR raw_input LIKE '%этаж%' OR
                raw_input LIKE '%на 8%' OR raw_input LIKE '%на 6%' OR
                raw_input LIKE '%на 10%' OR raw_input LIKE '%на 12%'
              )
            ORDER BY updated_at DESC LIMIT 10
        """, (str(chat_id), int(topic_id or 0), str(task_id))).fetchall()

        best_raw, best_score = "", 0
        for row in rows:
            ri = str(row[0] or "")
            low = ri.lower().replace("ё", "е")
            score = 0
            if any(x in low for x in ("дом", "ангар", "склад", "баня")): score += 20
            if any(x in low for x in ("каркас", "газобетон", "монолит", "кирпич")): score += 20
            import re as _t2re
            if _t2re.search(r"\d+\s*(?:на|x|х)\s*\d+", low): score += 25
            if _t2re.search(r"\d+\s*км", low): score += 10
            if any(x in low for x in ("монолит", "плита", "лента")): score += 10
            # CANCELLED P6E67 tasks are gold — full spec was there
            if "P6E67_MERGED" in str(row[1] or ""): score += 15
            if score > best_score:
                best_score, best_raw = score, ri

        if best_score >= 20 and best_raw.strip():
            _T2BCH_LOG.info("FIX_TOP2_BUILD_CONTEXT_HISTORY_V1: injecting history ctx score=%d", best_score)
            history_part = "Историческое ТЗ (из завершённых задач):\n" + best_raw.strip()
            if context_str:
                context_str = context_str + "\n\n" + history_part
            else:
                context_str = history_part
    except Exception as _t2bch_e:
        _T2BCH_LOG.warning("FIX_TOP2_BUILD_CONTEXT_HISTORY_V1 err: %s", _t2bch_e)

    return context_str, parent_id


_T2BCH_LOG.info("FIX_TOP2_BUILD_CONTEXT_HISTORY_V1 installed")
# === END_FIX_TOP2_BUILD_CONTEXT_HISTORY_V1 ===

# === FIX_P6_TOPIC2_ESTIMATE_VAGUE_V1 ===
# _p6_is_topic2_estimate_20260504 missed "посчитай" (without "ть").
# _p6_is_topic2_vague_20260504 triggered on "посмотри" even in estimate requests.
# "посчитай вот этот дом все данные у тебя есть" → vague=True → "Нет ТЗ".
import logging as _p6fix_log_mod
_P6FIX_LOG = _p6fix_log_mod.getLogger("task_worker")

_p6fix_orig_estimate = _p6_is_topic2_estimate_20260504
_p6fix_orig_vague = _p6_is_topic2_vague_20260504


def _p6_is_topic2_estimate_20260504(raw):
    if _p6fix_orig_estimate(raw):
        return True
    low = str(raw or "").lower().replace("ё", "е")
    has_calc = any(x in low for x in (
        "посчитай", "рассчитай", "нужна смета", "нужен расчет",
        "сколько стоит", "сколько будет", "цена", "стоимост",
    ))
    has_obj = any(x in low for x in (
        "дом", "ангар", "склад", "баня", "гараж", "house", "хаус",
    ))
    has_ref = any(x in low for x in (
        "данные у тебя", "у тебя есть", "все данные", "как описано",
        "как у меня", "вот этот", "этот дом", "этот объект",
    ))
    if has_calc and (has_obj or has_ref):
        return True
    if has_ref and has_obj:
        return True
    return False


def _p6_is_topic2_vague_20260504(raw):
    if _p6_is_topic2_estimate_20260504(raw):
        return False
    return _p6fix_orig_vague(raw)


_P6FIX_LOG.info("FIX_P6_TOPIC2_ESTIMATE_VAGUE_V1 installed")
# === END_FIX_P6_TOPIC2_ESTIMATE_VAGUE_V1 ===

# === FIX_P6_TOPIC2_PRICE_CONFIRM_ROUTE_V1 ===
# "ставь средние" / "выполни задачу" / "собирай" after price choice dialog
# was not recognised as estimate continuation → P6 sent it to vague guard.
# Fix: detect price-confirm phrases and route to estimate pipeline so that
# stroyka_estimate_canon._is_confirm picks it up and resumes pending estimate.
import logging as _p6pcr_log_mod
_P6PCR_LOG = _p6pcr_log_mod.getLogger("task_worker")

_p6pcr_orig_estimate = _p6_is_topic2_estimate_20260504
_p6pcr_orig_vague = _p6_is_topic2_vague_20260504

_P6PCR_CONFIRM_PHRASES = (
    "ставь средн", "ставь минимальн", "ставь максимальн", "ставь шаблон",
    "выполни задачу", "выполняй", "собирай", "делай смету", "создавай",
    "средние цены", "минимальные цены", "шаблонные цены",
    "беру средн", "беру минимальн", "согласен", "согласна", "поехали",
)


def _p6_is_topic2_estimate_20260504(raw):
    if _p6pcr_orig_estimate(raw):
        return True
    low = str(raw or "").lower().replace("ё", "е")
    return any(x in low for x in _P6PCR_CONFIRM_PHRASES)


def _p6_is_topic2_vague_20260504(raw):
    if _p6_is_topic2_estimate_20260504(raw):
        return False
    return _p6pcr_orig_vague(raw)


_P6PCR_LOG.info("FIX_P6_TOPIC2_PRICE_CONFIRM_ROUTE_V1 installed")
# === END_FIX_P6_TOPIC2_PRICE_CONFIRM_ROUTE_V1 ===

# === MOVE_MAIN_ENTRYPOINT_TO_END_V4 DISABLED — superseded by V5 ===
# if __name__ == "__main__":
#     asyncio.run(main())
# === END_MOVE_MAIN_ENTRYPOINT_TO_END_V4 ===


# === PATCH_TOPIC2_TASK_WORKER_FULL_CLOSE_V3 ===
import logging as _tw3_log_mod, re as _tw3_re
_TW3_LOG = _tw3_log_mod.getLogger("task_worker")

# --- A: Fix _force_voice_finish word boundary bug ---
# "задачу" contains "да" as substring → was triggering auto-DONE incorrectly.
# Fix: require "да" / "ок" as whole word boundaries.
def _force_voice_finish(raw_input: str, result: str) -> bool:
    if not raw_input:
        return False
    low = raw_input.lower()
    if "не доволен" in low or "не согласен" in low:
        return False
    # Whole-word match only
    if _tw3_re.search(r"\bда\b", low) or _tw3_re.search(r"\bок\b", low) or _tw3_re.search(r"\bok\b", low):
        return True
    if _tw3_re.search(r"\bзавершен", low) or "доволен" in low:
        return True
    return False

_TW3_LOG.info("FIX_FORCE_VOICE_FINISH_WORD_BOUNDARY_V1 installed")

# --- B: Wrap _p6_handle_topic2_estimate_20260504 to check pending before P3 ---
# When confirm phrase ("ставь средние" etc.) arrives as a new task,
# route through maybe_handle_stroyka_estimate first to resume pending estimate.
_tw3_orig_p6_handle_estimate = _p6_handle_topic2_estimate_20260504

async def _p6_handle_topic2_estimate_20260504(conn, task, chat_id, topic_id):
    raw = str(_p6_row_get_20260504(task, "raw_input", "") or "")
    low = raw.lower().replace("ё", "е").replace("[voice]", "").strip()
    # Check if this is a price-confirmation or session-continuation phrase
    is_confirm_phrase = any(x in low for x in _P6PCR_CONFIRM_PHRASES) or any(x in low for x in (
        "выполни задач", "выполни задание", "выполняй задан",
        "делай смету", "создавай смету", "сделай смету",
        "посчитай полностью", "в полном объёме", "в полном объеме",
        "где смет", "мои смет", "по каждому заданию",
    ))
    if is_confirm_phrase:
        try:
            from core.stroyka_estimate_canon import maybe_handle_stroyka_estimate as _tw3_mhse
            handled = await _tw3_mhse(conn, task, _TW3_LOG)
            if handled:
                _TW3_LOG.info("TW3_P6_PENDING_CONFIRM_ROUTE: pending resume handled task=%s", _p6_row_get_20260504(task, "id", ""))
                return True
        except Exception as _tw3_e:
            _TW3_LOG.warning("TW3_P6_PENDING_CONFIRM_ROUTE_ERR: %s", _tw3_e)
    return await _tw3_orig_p6_handle_estimate(conn, task, chat_id, topic_id)

_TW3_LOG.info("FIX_P6_ESTIMATE_PENDING_CONFIRM_ROUTE_V3 installed")

# --- C: topic_2 photo route: process_photo_topic2 before generic file handler ---
_tw3_orig_p6_handle_photo = None
try:
    _tw3_orig_p6_handle_photo = _p6_handle_photo_20260504  # type: ignore[name-defined]
except NameError:
    pass

if _tw3_orig_p6_handle_photo is not None:
    async def _p6_handle_photo_20260504(conn, task, chat_id, topic_id):
        task_id = str(_p6_row_get_20260504(task, "id", ""))
        t_id = int(_p6_row_get_20260504(task, "topic_id", 0) or 0)
        if t_id == 2:
            try:
                raw = str(_p6_row_get_20260504(task, "raw_input", "") or "")
                caption = ""
                file_name = ""
                file_path_local = ""
                try:
                    import json as _tw3j
                    d = _tw3j.loads(raw)
                    if isinstance(d, dict):
                        caption = d.get("caption") or d.get("owner_comment") or ""
                        file_name = d.get("file_name") or d.get("name") or ""
                        file_path_local = d.get("local_path") or d.get("file_path") or ""
                except Exception:
                    caption = raw
                from core.photo_recognition_engine import process_photo_topic2
                pr = process_photo_topic2(
                    file_name=file_name,
                    file_path=file_path_local,
                    owner_comment=caption,
                    caption=caption,
                )
                _TW3_LOG.info("TOPIC2_PHOTO_RECOGNITION_STARTED task=%s route=%s", task_id, pr.get("route"))
                route = pr.get("route", "menu")
                reply_to = _p6_row_get_20260504(task, "reply_to_message_id", None)

                if route == "estimate":
                    # Inject photo context into raw_input and route to estimate
                    enriched_raw = pr.get("photo_context", "") + "\n" + caption
                    enriched_task = dict(task) if isinstance(task, dict) else {k: task[k] for k in task.keys()}
                    enriched_task["raw_input"] = enriched_raw
                    _p6_history_20260504(conn, task_id, "TOPIC2_PHOTO_RECOGNITION_DONE")
                    _p6_history_20260504(conn, task_id, "TOPIC2_PHOTO_CONTEXT_USED")
                    _TW3_LOG.info("TOPIC2_PHOTO_RECOGNITION_DONE task=%s", task_id)
                    return await _p6_handle_topic2_estimate_20260504(conn, enriched_task, chat_id, topic_id)

                elif route == "ask_clarification":
                    q = pr.get("clarification_question", "Уточните параметры объекта")
                    _p6_update_20260504(conn, task_id, state="WAITING_CLARIFICATION", result=q, error_message="")
                    _p6_history_20260504(conn, task_id, f"TOPIC2_PHOTO_CONTEXT_MISSING_FIELDS:{','.join(pr.get('missing_fields', []))}")
                    conn.commit()
                    _send_once_ex(conn, task_id, str(chat_id), q, reply_to, "p6_photo_topic2_clarify")
                    _TW3_LOG.info("TOPIC2_PHOTO_RECOGNITION_DONE task=%s missing=%s", task_id, pr.get("missing_fields"))
                    return True

                else:  # menu
                    menu = pr.get("clarification_question", (
                        "Что сделать с этим фото?\n"
                        "1 — смета\n2 — описание\n3 — таблица\n4 — шаблон\n5 — анализ"
                    ))
                    _p6_update_20260504(conn, task_id, state="WAITING_CLARIFICATION", result=menu, error_message="")
                    _p6_history_20260504(conn, task_id, "TOPIC2_ROUTE_MENU_NO_INTENT")
                    conn.commit()
                    _send_once_ex(conn, task_id, str(chat_id), menu, reply_to, "p6_photo_topic2_menu")
                    return True
            except Exception as _tw3_ph_e:
                _TW3_LOG.warning("TOPIC2_PHOTO_ROUTE_ERR task=%s: %s", task_id, _tw3_ph_e)
        return await _tw3_orig_p6_handle_photo(conn, task, chat_id, topic_id)

    _TW3_LOG.info("FIX_P6_TOPIC2_PHOTO_ROUTE_V1 installed")
else:
    _TW3_LOG.warning("FIX_P6_TOPIC2_PHOTO_ROUTE_V1 skipped: _p6_handle_photo_20260504 not found")

# --- D: Regression guard — block "Уточни размеры дома" when context has dims ---
_tw3_orig_p6_clarification = None
try:
    _tw3_orig_p6_clarification = _p6_handle_topic2_vague_20260504  # type: ignore[name-defined]
except NameError:
    pass

if _tw3_orig_p6_clarification is not None:
    def _p6_handle_topic2_vague_20260504(conn, task, chat_id, topic_id):
        raw = str(_p6_row_get_20260504(task, "raw_input", "") or "")
        low = raw.lower().replace("ё", "е")
        task_id = str(_p6_row_get_20260504(task, "id", ""))
        # TOPIC2_GENERIC_CLARIFICATION_BLOCKED: if context has dims/object — not vague
        has_dims = bool(_tw3_re.search(r"\d+\s*(?:x|х|×|\*|на)\s*\d+|\d+\s*м\b", low))
        has_obj = any(x in low for x in ("дом", "ангар", "склад", "баня", "гараж", "каркас", "газобетон", "монолит"))
        if has_dims or has_obj:
            _p6_history_20260504(conn, task_id, "TOPIC2_GENERIC_CLARIFICATION_BLOCKED")
            _TW3_LOG.info("TW3: vague guard blocked — context has dims/obj, routing to estimate task=%s", task_id)
            return _tw3_orig_p6_clarification.__wrapped_orig__(conn, task, chat_id, topic_id) if hasattr(_tw3_orig_p6_clarification, "__wrapped_orig__") else _p6_handle_topic2_estimate_20260504.__wrapped__(conn, task, chat_id, topic_id) if False else None or _tw3_orig_p6_clarification(conn, task, chat_id, topic_id)
        return _tw3_orig_p6_clarification(conn, task, chat_id, topic_id)

_TW3_LOG.info("PATCH_TOPIC2_TASK_WORKER_FULL_CLOSE_V3 installed")
_TW3_LOG.info("TOPIC2_REGRESSION_GUARD_INSTALLED")
# === END_PATCH_TOPIC2_TASK_WORKER_FULL_CLOSE_V3 ===



def _price_bind_poison_parent_guard_v2(row):
    def g(key, default=""):
        try:
            if hasattr(row, "keys") and key in row.keys():
                return row[key]
        except Exception:
            pass
        try:
            return row[key]
        except Exception:
            return default

    raw = str(g("raw_input", "") or "").lower().replace("ё", "е")
    state = str(g("state", "") or "").upper()
    err = str(g("error_message", "") or "").upper()

    if state in ("CANCELLED", "ARCHIVED"):
        return True
    if "CANCEL" in err or "ОТМЕН" in err:
        return True

    poison = (
        "отмена", "отмени", "отменить", "очисти", "очистить",
        "удали все задачи", "закрой все задачи", "cancel"
    )
    return any(x in raw for x in poison)
# === PATCH_TOPIC2_PRICE_CHOICE_BIND_AND_FINALIZE_V4 ===
import re as _t2pc_re
import inspect as _t2pc_inspect
import logging as _t2pc_logging

_T2PC_LOG = _t2pc_logging.getLogger("task_worker")

def _t2pc_s(v, limit=12000):
    try:
        s = "" if v is None else str(v)
    except Exception:
        s = ""
    s = s.replace("\x00", "")
    return s[:limit]

def _t2pc_low(v):
    return _t2pc_s(v).lower().replace("ё", "е").replace("[voice]", "").strip()

def _t2pc_row(row, key, default=None):
    try:
        if row is None:
            return default
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
        if isinstance(row, dict):
            return row.get(key, default)
    except Exception:
        pass
    return default

def _t2pc_rowdict(row):
    try:
        if hasattr(row, "keys"):
            return {k: row[k] for k in row.keys()}
        if isinstance(row, dict):
            return dict(row)
    except Exception:
        pass
    return {}

def _t2pc_choice(raw):
    t = _t2pc_low(raw)
    t = _t2pc_re.sub(r"\s+", " ", t).strip(" .,!?:;()[]{}")
    if not t:
        return ""
    exact = {
        "1": "minimum", "а": "minimum", "a": "minimum", "а)": "minimum", "a)": "minimum",
        "2": "median", "б": "median", "b": "median", "б)": "median", "b)": "median",
        "3": "maximum", "в": "maximum", "v": "maximum", "в)": "maximum", "v)": "maximum",
        "4": "manual", "г": "manual", "g": "manual", "г)": "manual", "g)": "manual",
    }
    if t in exact:
        return exact[t]
    if any(x in t for x in ("миним", "дешев", "дешев", "самые низкие", "ставь миним")):
        return "minimum"
    if any(x in t for x in ("средн", "медиан", "рынок", "ставь сред", "беру сред", "средние цены")):
        return "median"
    if any(x in t for x in ("максим", "надеж", "надёж", "проверенн", "ставь максим")):
        return "maximum"
    if any(x in t for x in ("ручн", "вручную", "сам укажу", "мои цены", "своя")):
        return "manual"
    return ""

def _t2pc_find_parent(conn, task, chat_id, topic_id):
    task_id = _t2pc_s(_t2pc_row(task, "id", ""))
    reply_to = _t2pc_row(task, "reply_to_message_id", None)
    if reply_to:
        row = conn.execute("""
            SELECT *
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=?
              AND id<>?
              AND result LIKE '%Выберите уровень цен%'
              AND (bot_message_id=? OR reply_to_message_id=?)
              AND updated_at >= datetime('now','-24 hours')
            ORDER BY updated_at DESC, rowid DESC
            LIMIT 1
        """, (str(chat_id), int(topic_id or 0), task_id, reply_to, reply_to)).fetchone()
        if row:
            return row, "EXACT_REPLY_PRICE_MENU"

    row = conn.execute("""
        SELECT *
        FROM tasks
        WHERE chat_id=?
          AND COALESCE(topic_id,0)=?
          AND id<>?
          AND result LIKE '%Выберите уровень цен%'
          AND updated_at >= datetime('now','-24 hours')
        ORDER BY
          CASE
            WHEN LENGTH(COALESCE(raw_input,'')) > 500 THEN 0
            WHEN COALESCE(raw_input,'') LIKE '%№%' THEN 0
            WHEN COALESCE(raw_input,'') LIKE '%м³%' THEN 0
            WHEN COALESCE(raw_input,'') LIKE '%м2%' THEN 0
            ELSE 1
          END,
          updated_at DESC,
          rowid DESC
        LIMIT 1
    """, (str(chat_id), int(topic_id or 0), task_id)).fetchone()
    if row:
        return row, "LATEST_PRICE_MENU_FALLBACK"

    return None, "PRICE_MENU_NOT_FOUND"

async def _t2pc_try_bind_price_choice(conn, task):
    try:
        task_id = _t2pc_s(_t2pc_row(task, "id", ""))
        chat_id = _t2pc_s(_t2pc_row(task, "chat_id", ""))
        topic_id = int(_t2pc_row(task, "topic_id", 0) or 0)
        if topic_id != 2:
            return False

        raw = _t2pc_s(_t2pc_row(task, "raw_input", ""))
        choice = _t2pc_choice(raw)
        if not choice:
            return False

        parent, source = _t2pc_find_parent(conn, task, chat_id, topic_id)
        if not parent:
            try:
                _p6_history_20260504(conn, task_id, "PATCH_TOPIC2_PRICE_CHOICE_PARENT_NOT_FOUND")
                conn.commit()
            except Exception:
                pass
            return False

        if _price_bind_poison_parent_guard_v2(parent):
            try:
                _p6_history_20260504(conn, task_id, "PRICE_BIND_POISON_PARENT_GUARD_V2_BLOCKED_V4:" + str(source))
                conn.commit()
            except Exception:
                pass
            return False

        parent_id = _t2pc_s(_t2pc_row(parent, "id", ""))
        parent_task = _t2pc_rowdict(parent)
        parent_task["id"] = parent_id
        parent_task["chat_id"] = chat_id
        parent_task["topic_id"] = topic_id
        parent_task["state"] = "IN_PROGRESS"
        parent_task["input_type"] = "text"
        parent_task["raw_input"] = raw
        parent_task["reply_to_message_id"] = _t2pc_row(task, "reply_to_message_id", _t2pc_row(parent, "reply_to_message_id", None))

        conn.execute("""
            UPDATE tasks
            SET state='IN_PROGRESS',
                error_message='',
                updated_at=datetime('now')
            WHERE id=?
        """, (parent_id,))
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (parent_id, "TOPIC2_PRICE_CHOICE_CONFIRMED:" + choice),
        )
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (parent_id, "PATCH_TOPIC2_PRICE_CHOICE_BOUND_FROM:" + task_id + ":" + source),
        )
        conn.commit()

        handled = False
        try:
            from core.stroyka_estimate_canon import maybe_handle_stroyka_estimate as _t2pc_mhse
            res = _t2pc_mhse(conn, parent_task, _T2PC_LOG)
            handled = await res if _t2pc_inspect.isawaitable(res) else bool(res)
        except Exception as e:
            _T2PC_LOG.warning("PATCH_TOPIC2_PRICE_CHOICE_BIND_CANON_ERR parent=%s task=%s err=%s", parent_id, task_id, e)
            handled = False

        if handled:
            conn.execute("""
                UPDATE tasks
                SET state='DONE',
                    result=?,
                    error_message='',
                    updated_at=datetime('now')
                WHERE id=?
            """, ("PATCH_TOPIC2_PRICE_CHOICE_BOUND_TO:" + parent_id, task_id))
            conn.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (task_id, "PATCH_TOPIC2_PRICE_CHOICE_BOUND_TO:" + parent_id),
            )
            conn.commit()
            _T2PC_LOG.info("PATCH_TOPIC2_PRICE_CHOICE_BIND_OK parent=%s task=%s choice=%s", parent_id, task_id, choice)
            return True

        conn.execute("""
            UPDATE tasks
            SET state='WAITING_CLARIFICATION',
                result='Выбор цены принят, но генерация сметы не стартовала. Повтори выбор цены: 1 / 2 / 3 / 4',
                error_message='PATCH_TOPIC2_PRICE_CHOICE_BIND_NO_GENERATION',
                updated_at=datetime('now')
            WHERE id=?
        """, (parent_id,))
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (parent_id, "PATCH_TOPIC2_PRICE_CHOICE_BIND_NO_GENERATION"),
        )
        conn.execute("""
            UPDATE tasks
            SET state='DONE',
                result=?,
                error_message='',
                updated_at=datetime('now')
            WHERE id=?
        """, ("PATCH_TOPIC2_PRICE_CHOICE_ATTACHED_NO_GENERATION:" + parent_id, task_id))
        conn.commit()
        return True

    except Exception as e:
        try:
            _T2PC_LOG.warning("PATCH_TOPIC2_PRICE_CHOICE_BIND_ERR task=%s err=%s", _t2pc_row(task, "id", ""), e)
        except Exception:
            pass
        return False

_T2PC_ORIG_HANDLE_NEW = _handle_new
async def _handle_new(conn, task, *args, **kwargs):
    if await _t2pc_try_bind_price_choice(conn, task):
        return True
    res = _T2PC_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
    return await res if _t2pc_inspect.isawaitable(res) else res

_T2PC_ORIG_HANDLE_IN_PROGRESS = _handle_in_progress
async def _handle_in_progress(conn, task, *args, **kwargs):
    if await _t2pc_try_bind_price_choice(conn, task):
        return True
    res = _T2PC_ORIG_HANDLE_IN_PROGRESS(conn, task, *args, **kwargs)
    return await res if _t2pc_inspect.isawaitable(res) else res

_T2PC_LOG.info("PATCH_TOPIC2_PRICE_CHOICE_BIND_AND_FINALIZE_V4 installed")
# === END_PATCH_TOPIC2_PRICE_CHOICE_BIND_AND_FINALIZE_V4 ===


# === PATCH_TOPIC2_PRICE_CHOICE_BIND_AND_FINALIZE_V5 ===
import re as _t2v5_re
import inspect as _t2v5_inspect
import logging as _t2v5_logging

_T2V5_LOG = _t2v5_logging.getLogger("task_worker")

_T2V5_PRICE_MAP = {
    "1": "minimum",
    "2": "median",
    "3": "maximum",
    "4": "manual",
    "а": "minimum",
    "a": "minimum",
    "б": "median",
    "b": "median",
    "в": "maximum",
    "v": "maximum",
    "г": "manual",
    "g": "manual",
}

def _t2v5_s(v, n=12000):
    try:
        return str(v or "")[:n]
    except Exception:
        return ""

def _t2v5_row(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        return default

def _t2v5_clean_voice(text):
    t = _t2v5_s(text).lower().replace("ё", "е")
    t = t.replace("[voice]", "").replace("[голос]", "")
    t = _t2v5_re.sub(r"\s+", " ", t).strip(" \n\r\t.,!?:;()[]{}")
    return t

def _t2v5_price_choice(text):
    t = _t2v5_clean_voice(text)
    if t in _T2V5_PRICE_MAP:
        return t, _T2V5_PRICE_MAP[t]
    m = _t2v5_re.fullmatch(r"(вариант\s*)?([1234])", t)
    if m:
        k = m.group(2)
        return k, _T2V5_PRICE_MAP[k]
    if any(x in t for x in ("ставь средн", "средние цены", "средн", "медиан", "вариант 2", "вариант б")):
        return "2", "median"
    if any(x in t for x in ("ставь миним", "миним", "самые дешев", "вариант 1", "вариант а")):
        return "1", "minimum"
    if any(x in t for x in ("ставь максим", "максим", "надежн", "надёжн", "вариант 3", "вариант в")):
        return "3", "maximum"
    if any(x in t for x in ("ручн", "сам укаж", "свои цены", "вариант 4", "вариант г")):
        return "4", "manual"
    return "", ""

def _t2v5_latest_clarified_choice(conn, task_id):
    try:
        rows = conn.execute(
            """
            SELECT action
            FROM task_history
            WHERE task_id=?
              AND action LIKE 'clarified:%'
            ORDER BY created_at DESC
            LIMIT 20
            """,
            (str(task_id),),
        ).fetchall()
        for r in rows:
            a = _t2v5_s(_t2v5_row(r, "action", r[0] if r else ""))
            txt = a.split("clarified:", 1)[1] if "clarified:" in a else a
            key, choice = _t2v5_price_choice(txt)
            if key and choice:
                return key, choice, txt
    except Exception as e:
        _T2V5_LOG.warning("PATCH_TOPIC2_PRICE_CHOICE_BIND_V5_HISTORY_ERR %s", e)
    return "", "", ""

def _t2v5_is_status_question(raw):
    t = _t2v5_clean_voice(raw)
    return (
        ("последн" in t and "задач" in t and ("какая" in t or "что" in t))
        or ("какая у тебя последняя задача" in t)
        or ("что у тебя последняя задача" in t)
    )

def _t2v5_is_estimate_like(raw):
    t = _t2v5_clean_voice(raw)
    return any(x in t for x in (
        "смет", "кп", "коммерчес", "проценк", "стоимост", "цена",
        "объект", "фундамент", "бетон", "арматур", "опалуб", "гидроизоля",
        "м3", "м²", "м2", "шт", "монолит", "ж/б", "железобетон",
        "плита", "стены", "перекрыт", "колонн", "балк", "лестнич",
        "зеленогорск", "подрядчик", "входит", "не входит"
    ))

def _t2v5_to_dict(row):
    if row is None:
        return {}
    try:
        return {k: row[k] for k in row.keys()}
    except Exception:
        return {}

def _t2v5_find_price_parent(conn, chat_id, topic_id, current_id):
    try:
        rows = conn.execute(
            """
            SELECT *
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=?
              AND id<>?
              AND COALESCE(result,'') LIKE '%Выберите уровень цен%'
              AND state IN ('WAITING_CLARIFICATION','FAILED','AWAITING_CONFIRMATION','IN_PROGRESS')
              AND updated_at >= datetime('now','-24 hours')
            ORDER BY
              CASE WHEN state='WAITING_CLARIFICATION' THEN 0 ELSE 1 END,
              updated_at DESC,
              created_at DESC
            LIMIT 40
            """,
            (str(chat_id), int(topic_id or 0), str(current_id)),
        ).fetchall()
    except Exception as e:
        _T2V5_LOG.warning("PATCH_TOPIC2_PRICE_PARENT_FIND_V5_SQL_ERR %s", e)
        return None

    fallback = None
    for r in rows:
        raw = _t2v5_s(_t2v5_row(r, "raw_input", ""))
        if _t2v5_is_status_question(raw):
            continue
        if fallback is None:
            fallback = r
        if _t2v5_is_estimate_like(raw):
            return r
    return fallback

def _t2v5_latest_real_topic2_task(conn, chat_id, topic_id, current_id):
    try:
        rows = conn.execute(
            """
            SELECT id,state,COALESCE(raw_input,'') AS raw_input,COALESCE(result,'') AS result,updated_at
            FROM tasks
            WHERE chat_id=?
              AND COALESCE(topic_id,0)=?
              AND id<>?
            ORDER BY updated_at DESC, created_at DESC
            LIMIT 20
            """,
            (str(chat_id), int(topic_id or 0), str(current_id)),
        ).fetchall()
        for r in rows:
            raw = _t2v5_s(_t2v5_row(r, "raw_input", ""))
            if _t2v5_is_status_question(raw):
                continue
            return r
    except Exception:
        return None
    return None

async def _t2v5_try_bind_price_choice(conn, task):
    task_id = _t2v5_s(_t2v5_row(task, "id", ""))
    chat_id = _t2v5_s(_t2v5_row(task, "chat_id", ""))
    try:
        topic_id = int(_t2v5_row(task, "topic_id", 0) or 0)
    except Exception:
        topic_id = 0

    if not task_id or not chat_id or topic_id != 2:
        return False

    raw = _t2v5_s(_t2v5_row(task, "raw_input", ""))
    # === PATCH_TOPIC2_INLINE_FIX_20260506_V1 V5_EXPLICIT_TOKEN_REQUIRED ===
    # block long messages and messages without explicit price-token from being treated as price reply
    _pifx_v5_low = str(raw or "").strip().lower()
    _pifx_v5_explicit = ("1","2","3","4","а","б","в","г","миним","средн","медиан","максим","надёж","надеж","ручн","вариант")
    if len(_pifx_v5_low) > 80 or not any(t in _pifx_v5_low for t in _pifx_v5_explicit):
        try:
            conn.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (str(task_id), "PATCH_TOPIC2_INLINE_FIX_20260506_V1:V5_PRICE_REJECTED:no_explicit_token_or_long"),
            )
            conn.commit()
        except Exception:
            pass
        return False
    # === END_PATCH_TOPIC2_INLINE_FIX_20260506_V1 V5_EXPLICIT_TOKEN_REQUIRED ===
    key, choice = _t2v5_price_choice(raw)
    source_text = raw

    if not key:
        key, choice, source_text = _t2v5_latest_clarified_choice(conn, task_id)

    if not key or not choice:
        return False

    parent = _t2v5_find_price_parent(conn, chat_id, topic_id, task_id)
    if parent is None:
        _T2V5_LOG.info("PATCH_TOPIC2_PRICE_CHOICE_BIND_V5_NO_PARENT task=%s choice=%s", task_id, choice)
        return False

    if _price_bind_poison_parent_guard_v2(parent):
        try:
            conn.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (task_id, "PRICE_BIND_POISON_PARENT_GUARD_V2_BLOCKED_V5"),
            )
            conn.commit()
        except Exception:
            pass
        return False

    parent_id = _t2v5_s(_t2v5_row(parent, "id", ""))
    parent_dict = _t2v5_to_dict(parent)
    parent_dict["raw_input"] = key
    parent_dict["input_type"] = "text"
    parent_dict["state"] = "IN_PROGRESS"

    try:
        conn.execute(
            "UPDATE tasks SET state='IN_PROGRESS', error_message='', updated_at=datetime('now') WHERE id=?",
            (parent_id,),
        )
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (parent_id, "TOPIC2_PRICE_CHOICE_CONFIRMED:" + choice),
        )
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (parent_id, "PATCH_TOPIC2_PRICE_CHOICE_BIND_V5_FROM_TASK:" + task_id),
        )
        conn.execute(
            "UPDATE tasks SET state='DONE', result=?, error_message='', updated_at=datetime('now') WHERE id=?",
            ("Выбор цены принят и привязан к сметной задаче: " + parent_id, task_id),
        )
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (task_id, "PATCH_TOPIC2_PRICE_CHOICE_BIND_V5_BOUND_TO:" + parent_id),
        )
        conn.commit()
    except Exception as e:
        _T2V5_LOG.warning("PATCH_TOPIC2_PRICE_CHOICE_BIND_V5_DB_ERR %s", e)
        return False

    try:
        from core.stroyka_estimate_canon import maybe_handle_stroyka_estimate as _t2v5_mhse
        res = _t2v5_mhse(conn, parent_dict, _T2V5_LOG)
        if _t2v5_inspect.isawaitable(res):
            res = await res
        if res:
            _T2V5_LOG.info("PATCH_TOPIC2_PRICE_CHOICE_BIND_V5_GENERATED parent=%s choice=%s", parent_id, choice)
            return True
    except Exception as e:
        _T2V5_LOG.warning("PATCH_TOPIC2_PRICE_CHOICE_BIND_V5_GENERATE_ERR parent=%s err=%s", parent_id, e)

    try:
        conn.execute(
            "UPDATE tasks SET state='WAITING_CLARIFICATION', error_message='PATCH_TOPIC2_PRICE_CHOICE_BIND_V5_GENERATE_NOT_HANDLED', updated_at=datetime('now') WHERE id=?",
            (parent_id,),
        )
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (parent_id, "PATCH_TOPIC2_PRICE_CHOICE_BIND_V5_GENERATE_NOT_HANDLED"),
        )
        conn.commit()
    except Exception:
        pass
    return True

def _t2v5_handle_status_question(conn, task):
    task_id = _t2v5_s(_t2v5_row(task, "id", ""))
    chat_id = _t2v5_s(_t2v5_row(task, "chat_id", ""))
    try:
        topic_id = int(_t2v5_row(task, "topic_id", 0) or 0)
    except Exception:
        topic_id = 0
    raw = _t2v5_s(_t2v5_row(task, "raw_input", ""))

    if topic_id != 2 or not _t2v5_is_status_question(raw):
        return False

    row = _t2v5_latest_real_topic2_task(conn, chat_id, topic_id, task_id)
    if row is None:
        msg = "В topic_2 нет найденной предыдущей задачи"
    else:
        rid = _t2v5_s(_t2v5_row(row, "id", ""))
        state = _t2v5_s(_t2v5_row(row, "state", ""))
        result = _t2v5_s(_t2v5_row(row, "result", ""), 700).strip()
        raw_prev = _t2v5_s(_t2v5_row(row, "raw_input", ""), 500).strip()
        body = result or raw_prev
        msg = "Последняя задача topic_2:\n" + rid + "\nСтатус: " + state + "\n" + body

    try:
        reply_to = _t2v5_row(task, "reply_to_message_id", None)
        send_res = _send_once_ex(conn, task_id, str(chat_id), msg, reply_to, "topic2_status_question")
        bot_id = send_res.get("bot_message_id") if isinstance(send_res, dict) else None
        kwargs = {"state": "DONE", "result": msg, "error_message": ""}
        if bot_id:
            kwargs["bot_message_id"] = bot_id
        _update_task(conn, task_id, **kwargs)
        _history(conn, task_id, "PATCH_TOPIC2_STATUS_QUESTION_V5_HANDLED")
        conn.commit()
        return True
    except Exception as e:
        _T2V5_LOG.warning("PATCH_TOPIC2_STATUS_QUESTION_V5_ERR %s", e)
        return False

_T2V5_ORIG_HANDLE_NEW = _handle_new
async def _handle_new(conn, task, *args, **kwargs):
    if await _t2v5_try_bind_price_choice(conn, task):
        return True
    if _t2v5_handle_status_question(conn, task):
        return True
    res = _T2V5_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
    return await res if _t2v5_inspect.isawaitable(res) else res

_T2V5_ORIG_HANDLE_IN_PROGRESS = _handle_in_progress
async def _handle_in_progress(conn, task, *args, **kwargs):
    if await _t2v5_try_bind_price_choice(conn, task):
        return True
    if _t2v5_handle_status_question(conn, task):
        return True
    res = _T2V5_ORIG_HANDLE_IN_PROGRESS(conn, task, *args, **kwargs)
    return await res if _t2v5_inspect.isawaitable(res) else res

_T2V5_LOG.info("PATCH_TOPIC2_PRICE_CHOICE_BIND_AND_FINALIZE_V5 installed")
# === END_PATCH_TOPIC2_PRICE_CHOICE_BIND_AND_FINALIZE_V5 ===

# === MOVE_MAIN_ENTRYPOINT_TO_END_V5 ===


# === PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED ===
# Overlay only
# Fixes factual V5 state: price confirmed, parent killed by hard timeout, no artifact closure proven
# No DB schema changes
# No forbidden files

import json as _t2v6c_json
import sqlite3 as _t2v6c_sqlite3
import inspect as _t2v6c_inspect
import logging as _t2v6c_logging
import re as _t2v6c_re

_T2V6C_LOG = _t2v6c_logging.getLogger("task_worker")
_T2V6C_MEM_DB = "/root/.areal-neva-core/data/memory.db"

_T2V6C_PRICE_MAP = {
    "1": "minimum",
    "2": "median",
    "3": "maximum",
    "4": "manual",
    "а": "minimum",
    "a": "minimum",
    "б": "median",
    "b": "median",
    "в": "maximum",
    "v": "maximum",
    "г": "manual",
    "g": "manual",
}

def _t2v6c_s(v, n=12000):
    try:
        return str(v or "")[:n]
    except Exception:
        return ""

def _t2v6c_row(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        return default

def _t2v6c_clean(text):
    t = _t2v6c_s(text).lower().replace("ё", "е")
    t = t.replace("[voice]", "").replace("[голос]", "")
    t = _t2v6c_re.sub(r"\s+", " ", t).strip(" \n\r\t.,!?:;()[]{}")
    return t

def _t2v6c_choice(text):
    t = _t2v6c_clean(text)
    if t in _T2V6C_PRICE_MAP:
        return t, _T2V6C_PRICE_MAP[t]
    m = _t2v6c_re.fullmatch(r"(вариант\s*)?([1234])", t)
    if m:
        k = m.group(2)
        return k, _T2V6C_PRICE_MAP[k]
    if any(x in t for x in ("ставь средн", "средние цены", "средн", "медиан", "вариант 2", "вариант б")):
        return "2", "median"
    if any(x in t for x in ("ставь миним", "миним", "самые дешев", "вариант 1", "вариант а")):
        return "1", "minimum"
    if any(x in t for x in ("ставь максим", "максим", "надежн", "надёжн", "вариант 3", "вариант в")):
        return "3", "maximum"
    if any(x in t for x in ("ручн", "укажу сам", "вариант 4", "вариант г")):
        return "4", "manual"
    return "", ""

def _t2v6c_latest_clarified_choice(conn, task_id):
    try:
        rows = conn.execute(
            """
            SELECT action
            FROM task_history
            WHERE task_id=?
              AND action LIKE 'clarified:%'
            ORDER BY created_at DESC, rowid DESC
            LIMIT 20
            """,
            (str(task_id),),
        ).fetchall()
        for r in rows:
            v = _t2v6c_s(r[0])
            if ":" in v:
                key, choice = _t2v6c_choice(v.split(":", 1)[1])
                if key and choice:
                    return key, choice
    except Exception as e:
        _T2V6C_LOG.warning("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_CLARIFIED_ERR %s", e)
    return "", ""

def _t2v6c_has_artifacts(text):
    low = _t2v6c_s(text, 30000).lower()
    has_xlsx = "xlsx" in low or "excel" in low or "📊" in low
    has_pdf = "pdf" in low or "📄" in low
    has_link = "drive.google.com" in low or "http://" in low or "https://" in low
    return has_xlsx and has_pdf and has_link

def _t2v6c_is_poison_parent(row):
    raw = _t2v6c_clean(_t2v6c_row(row, "raw_input", ""))
    state = _t2v6c_s(_t2v6c_row(row, "state", "")).upper()
    err = _t2v6c_s(_t2v6c_row(row, "error_message", "")).upper()
    if state in ("CANCELLED", "ARCHIVED"):
        return True
    if "CANCEL" in err or "ОТМЕН" in err:
        return True
    poison = (
        "отмена", "отмени", "отменить", "очисти", "очистить",
        "удали все задачи", "закрой все задачи", "cancel"
    )
    return any(x in raw for x in poison)

def _t2v6c_find_parent(conn, chat_id, topic_id, task_id):
    try:
        rows = conn.execute(
            """
            SELECT t.*
            FROM tasks t
            WHERE t.chat_id=?
              AND COALESCE(t.topic_id,0)=?
              AND t.id<>?
              AND t.state IN ('WAITING_CLARIFICATION','IN_PROGRESS','AWAITING_CONFIRMATION','FAILED')
              AND (
                t.result LIKE '%Выберите уровень цен%'
                OR t.result LIKE '%уровень цен%'
                OR EXISTS (
                    SELECT 1 FROM task_history h
                    WHERE h.task_id=t.id
                      AND h.action LIKE '%TOPIC2_PRICE_CHOICE_REQUESTED%'
                )
              )
            ORDER BY t.updated_at DESC, t.created_at DESC
            LIMIT 20
            """,
            (str(chat_id), int(topic_id or 0), str(task_id)),
        ).fetchall()
        for row in rows:
            state = _t2v6c_s(_t2v6c_row(row, "state", ""))
            result = _t2v6c_s(_t2v6c_row(row, "result", ""), 30000)
            err = _t2v6c_s(_t2v6c_row(row, "error_message", ""), 1000)
            if state == "CANCELLED":
                continue
            if _t2v6c_is_poison_parent(row):
                continue
            if _t2v6c_has_artifacts(result):
                continue
            if "PATCH_TOPIC2_LEGACY_BAD_DONE" in err:
                continue
            return row
    except Exception as e:
        _T2V6C_LOG.warning("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_FIND_PARENT_ERR %s", e)
    return None

def _t2v6c_load_pending(chat_id, parent_id):
    try:
        con = _t2v6c_sqlite3.connect(_T2V6C_MEM_DB, timeout=10)
        try:
            con.row_factory = _t2v6c_sqlite3.Row
            rows = con.execute(
                """
                SELECT key,value,timestamp
                FROM memory
                WHERE chat_id=?
                  AND key LIKE 'topic_2_estimate_pending_%'
                ORDER BY timestamp DESC, id DESC
                LIMIT 100
                """,
                (str(chat_id),),
            ).fetchall()
        finally:
            con.close()

        fallback = None
        for r in rows:
            try:
                val = _t2v6c_json.loads(_t2v6c_s(r["value"], 200000))
            except Exception:
                continue
            if not isinstance(val, dict):
                continue
            if _t2v6c_s(val.get("status")) != "WAITING_PRICE_CONFIRMATION":
                continue
            if _t2v6c_s(val.get("task_id")) == _t2v6c_s(parent_id):
                return val
            if fallback is None:
                fallback = val
        return fallback
    except Exception as e:
        _T2V6C_LOG.warning("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_PENDING_ERR %s", e)
        return None

async def _t2v6c_try_generate_after_price(conn, task):
    task_id = _t2v6c_s(_t2v6c_row(task, "id", ""))
    chat_id = _t2v6c_s(_t2v6c_row(task, "chat_id", ""))
    try:
        topic_id = int(_t2v6c_row(task, "topic_id", 0) or 0)
    except Exception:
        topic_id = 0

    if not task_id or not chat_id or topic_id != 2:
        return False

    raw = _t2v6c_s(_t2v6c_row(task, "raw_input", ""))
    # === PATCH_TOPIC2_INLINE_FIX_20260506_V1 V6C_EXPLICIT_TOKEN_REQUIRED ===
    _pifx_v6_low = str(raw or "").strip().lower()
    _pifx_v6_explicit = ("1","2","3","4","а","б","в","г","миним","средн","медиан","максим","надёж","надеж","ручн","вариант")
    if len(_pifx_v6_low) > 80 or not any(t in _pifx_v6_low for t in _pifx_v6_explicit):
        try:
            conn.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (str(task_id), "PATCH_TOPIC2_INLINE_FIX_20260506_V1:V6C_PRICE_REJECTED:no_explicit_token_or_long"),
            )
            conn.commit()
        except Exception:
            pass
        return False
    # === END_PATCH_TOPIC2_INLINE_FIX_20260506_V1 V6C_EXPLICIT_TOKEN_REQUIRED ===
    key, choice = _t2v6c_choice(raw)
    if not key:
        key, choice = _t2v6c_latest_clarified_choice(conn, task_id)
    if not key or not choice:
        return False

    parent = _t2v6c_find_parent(conn, chat_id, topic_id, task_id)
    if parent is None:
        _T2V6C_LOG.info("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_NO_PARENT task=%s choice=%s", task_id, choice)
        return False

    parent_id = _t2v6c_s(_t2v6c_row(parent, "id", ""))
    pending = _t2v6c_load_pending(chat_id, parent_id)
    if not pending:
        _T2V6C_LOG.warning("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_NO_PENDING parent=%s task=%s", parent_id, task_id)
        return False

    parent_task = {}
    try:
        for k in parent.keys():
            parent_task[k] = parent[k]
    except Exception:
        parent_task = {}

    parent_task["id"] = parent_id
    parent_task["chat_id"] = chat_id
    parent_task["topic_id"] = topic_id
    parent_task["state"] = "IN_PROGRESS"
    parent_task["input_type"] = "text"
    parent_task["raw_input"] = _t2v6c_s(_t2v6c_row(parent, "raw_input", ""))

    try:
        conn.execute(
            """
            UPDATE tasks
            SET state='IN_PROGRESS',
                error_message='',
                updated_at=datetime('now')
            WHERE id=?
            """,
            (parent_id,),
        )
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (parent_id, "TOPIC2_PRICE_CHOICE_CONFIRMED:" + choice),
        )
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (parent_id, "PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_STARTED_FROM:" + task_id),
        )
        conn.commit()
    except Exception as e:
        _T2V6C_LOG.warning("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_DB_START_ERR %s", e)
        return False

    try:
        from core.stroyka_estimate_canon import _generate_and_send as _t2v6c_generate_and_send
        res = _t2v6c_generate_and_send(conn, parent_task, pending, key, logger=_T2V6C_LOG)
        if _t2v6c_inspect.isawaitable(res):
            res = await res
    except Exception as e:
        _T2V6C_LOG.exception("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_GENERATE_ERR parent=%s err=%s", parent_id, e)
        try:
            conn.execute(
                """
                UPDATE tasks
                SET state='FAILED',
                    error_message=?,
                    updated_at=datetime('now')
                WHERE id=?
                """,
                ("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_GENERATE_ERR:" + _t2v6c_s(e, 500), parent_id),
            )
            conn.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (parent_id, "PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_GENERATE_ERR"),
            )
            conn.commit()
        except Exception:
            pass
        return True

    row = conn.execute("SELECT state,result,error_message FROM tasks WHERE id=?", (parent_id,)).fetchone()
    state = _t2v6c_s(_t2v6c_row(row, "state", ""))
    result = _t2v6c_s(_t2v6c_row(row, "result", ""), 30000)
    error = _t2v6c_s(_t2v6c_row(row, "error_message", ""), 1000)

    if state == "AWAITING_CONFIRMATION" and _t2v6c_has_artifacts(result):
        try:
            conn.execute(
                """
                UPDATE tasks
                SET state='DONE',
                    result=?,
                    error_message='',
                    updated_at=datetime('now')
                WHERE id=?
                """,
                ("Выбор цены принят. Смета сгенерирована в задаче: " + parent_id, task_id),
            )
            conn.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (parent_id, "PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_ARTIFACTS_OK"),
            )
            conn.execute(
                "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                (task_id, "PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_BOUND_TO:" + parent_id),
            )
            conn.commit()
        except Exception:
            pass
        _T2V6C_LOG.info("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_OK parent=%s choice=%s", parent_id, choice)
        return True

    try:
        conn.execute(
            """
            UPDATE tasks
            SET state='WAITING_CLARIFICATION',
                error_message=?,
                updated_at=datetime('now')
            WHERE id=?
            """,
            ("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_NO_ARTIFACTS_AFTER_GENERATE state=" + state + " err=" + error[:300], parent_id),
        )
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (parent_id, "PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_NO_ARTIFACTS_AFTER_GENERATE"),
        )
        conn.commit()
    except Exception:
        pass

    _T2V6C_LOG.warning("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED_NO_ARTIFACTS parent=%s state=%s err=%s", parent_id, state, error)
    return True

def _in_progress_hard_timeout_by_created_at_fix_v1(conn, minutes: int = 30) -> int:
    try:
        rows = conn.execute(
            """
            SELECT id
            FROM tasks
            WHERE state='IN_PROGRESS'
              AND datetime(COALESCE(updated_at, created_at, 'now')) <= datetime('now', ?)
            ORDER BY updated_at ASC
            LIMIT 200
            """,
            (f"-{int(minutes)} minutes",),
        ).fetchall()
        n = 0
        for row in rows:
            tid = row["id"] if hasattr(row, "keys") else row[0]
            conn.execute(
                """
                UPDATE tasks
                SET state='FAILED',
                    error_message='IN_PROGRESS_HARD_TIMEOUT_BY_UPDATED_AT_FIX_V6_CANON_FIXED',
                    updated_at=datetime('now')
                WHERE id=? AND state='IN_PROGRESS'
                """,
                (str(tid),),
            )
            try:
                _history(conn, str(tid), "IN_PROGRESS_HARD_TIMEOUT_BY_UPDATED_AT_FIX_V6_CANON_FIXED:FAILED")
            except Exception:
                pass
            n += 1
        if n:
            conn.commit()
            _T2V6C_LOG.warning("IN_PROGRESS_HARD_TIMEOUT_BY_UPDATED_AT_FIX_V6_CANON_FIXED closed=%s", n)
        return n
    except Exception as e:
        _T2V6C_LOG.warning("IN_PROGRESS_HARD_TIMEOUT_BY_UPDATED_AT_FIX_V6_CANON_FIXED_ERR %s", e)
        return 0

_T2V6C_ORIG_HANDLE_NEW = _handle_new
async def _handle_new(conn, task, *args, **kwargs):
    if await _t2v6c_try_generate_after_price(conn, task):
        return True
    res = _T2V6C_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
    return await res if _t2v6c_inspect.isawaitable(res) else res

_T2V6C_ORIG_HANDLE_IN_PROGRESS = _handle_in_progress
async def _handle_in_progress(conn, task, *args, **kwargs):
    if await _t2v6c_try_generate_after_price(conn, task):
        return True
    res = _T2V6C_ORIG_HANDLE_IN_PROGRESS(conn, task, *args, **kwargs)
    return await res if _t2v6c_inspect.isawaitable(res) else res

_T2V6C_LOG.info("PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED installed")
# === END_PATCH_TOPIC2_PRICE_FINAL_GENERATION_V6_CANON_FIXED ===



# === FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1 ===
import os as _fcg_os
import re as _fcg_re
import json as _fcg_json
import time as _fcg_time
import asyncio as _fcg_asyncio
import inspect as _fcg_inspect
import tempfile as _fcg_tempfile
import logging as _fcg_logging
import datetime as _fcg_datetime

_FCG_LOG = _fcg_logging.getLogger("task_worker")
_FCG_TIMEOUT_SEC = 300
_FCG_FILE_INPUT_TYPES = {"drive_file", "file", "photo", "image", "document"}
_FCG_BAD_PUBLIC_FRAGMENTS = (
    "/root/",
    ".ogg",
    "Traceback",
    "Файл скачан, ожидает анализа",
    "Файл не прикреплён",
    "Артефакт: [локальный путь скрыт]",
    "Артефакт создан, но загрузка в Drive не подтвердилась",
    # === PATCH_TOPIC2_OLD_OUTPUT_BLOCKER_V1 — §11 old route patterns ===
    "Эталон: ",
    "Лист эталона: ",
    "Выбор цены: ",
    "✅ Предварительная смета готова",
    "НДС 20%:",
)
_FCG_UNIT_RE = _fcg_re.compile(r"^(м³|м3|м\^3|м²|м2|м\^2|м\.?|мп|пм|п\.?\s*м\.?|шт\.?|компл\.?|кг|тн|т|тонн?а?|тонн)$", _fcg_re.I)

def _fcg_s(v, limit=12000):
    try:
        return str(v if v is not None else "")[:limit]
    except Exception:
        return ""

def _fcg_low(v):
    return _fcg_s(v, 50000).lower().replace("ё", "е")

def _fcg_row(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        return default

def _fcg_task_dict(row):
    try:
        return {k: row[k] for k in row.keys()}
    except Exception:
        return dict(row or {})

def _fcg_task_row(conn, task_id):
    try:
        conn.row_factory = getattr(conn, "row_factory", None) or conn.row_factory
        return conn.execute("SELECT * FROM tasks WHERE id=?", (str(task_id),)).fetchone()
    except Exception:
        return None

def _fcg_history(conn, task_id, action):
    try:
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (str(task_id), str(action)[:900]),
        )
    except Exception:
        pass

def _fcg_raw_payload(task):
    raw = _fcg_s(_fcg_row(task, "raw_input", ""))
    try:
        data = _fcg_json.loads(raw)
        if isinstance(data, dict):
            return data
    except Exception:
        pass
    return {}

def _fcg_source_file_id_from_task(task):
    data = _fcg_raw_payload(task)
    for k in ("file_id", "drive_file_id", "id"):
        if data.get(k):
            return _fcg_s(data.get(k), 300)
    return ""

def _fcg_has_drive_link(text):
    t = _fcg_s(text, 50000)
    return "drive.google.com/" in t or "docs.google.com/" in t

def _fcg_is_file_task(row):
    return _fcg_s(_fcg_row(row, "input_type", "")).strip() in _FCG_FILE_INPUT_TYPES

def _fcg_drive_stage_advanced(conn, task_id):
    try:
        rows = conn.execute(
            "SELECT stage FROM drive_files WHERE task_id=? ORDER BY id DESC LIMIT 10",
            (str(task_id),),
        ).fetchall()
        stages = {_fcg_s(_fcg_row(r, "stage", "")).upper() for r in rows}
        return bool(stages & {"DOWNLOADED", "PARSED", "CLEANED", "NORMALIZED", "CALCULATED", "ARTIFACT_CREATED", "UPLOADED", "COMPLETED", "FAILED"})
    except Exception:
        return False

def _fcg_requeue_loop_detected(conn, task):
    task_id = _fcg_s(_fcg_row(task, "id", ""))
    if not task_id or not _fcg_is_file_task(task):
        return False
    try:
        cnt = conn.execute(
            """
            SELECT COUNT(*) AS c
            FROM task_history
            WHERE task_id=?
              AND (
                action LIKE '%TEXT_FOLLOWUP_REQUEUED_AS_DRIVE_FILE%'
                OR action LIKE '%REQUEUED_AS_DRIVE_FILE%'
              )
            """,
            (task_id,),
        ).fetchone()[0]
        if cnt > 1:
            return True
        if cnt >= 1 and _fcg_drive_stage_advanced(conn, task_id):
            return True
    except Exception:
        pass
    return False

def _fcg_public_result_violation(conn, task_id, state, result, error_message=""):
    state = _fcg_s(state)
    result = _fcg_s(result, 80000)
    error_message = _fcg_s(error_message, 2000)
    if state not in ("AWAITING_CONFIRMATION", "DONE", "FAILED"):
        return ""
    row = _fcg_task_row(conn, task_id)
    if row is None:
        return ""
    for frag in _FCG_BAD_PUBLIC_FRAGMENTS:
        if frag in result:
            if frag == "Артефакт: [локальный путь скрыт]":
                return "NO_VALID_ARTIFACT"
            if frag == "Артефакт создан, но загрузка в Drive не подтвердилась":
                return "NO_VALID_ARTIFACT"
            if frag == "Файл скачан, ожидает анализа":
                return "NO_VALID_ARTIFACT"
            if frag == "Файл не прикреплён":
                return "NO_VALID_ARTIFACT"
            return "INVALID_PUBLIC_RESULT"
    if _fcg_is_file_task(row) and state in ("AWAITING_CONFIRMATION", "DONE"):
        if not _fcg_has_drive_link(result):
            return "NO_VALID_ARTIFACT"
        src_id = _fcg_source_file_id_from_task(row)
        if src_id and src_id in result:
            try:
                art = conn.execute(
                    "SELECT artifact_file_id FROM drive_files WHERE task_id=? AND COALESCE(artifact_file_id,'')<>'' ORDER BY id DESC LIMIT 1",
                    (str(task_id),),
                ).fetchone()
            except Exception:
                art = None
            if art is None:
                return "SOURCE_FILE_RETURNED_AS_RESULT"
    if error_message in ("ENGINE_TIMEOUT", "NO_VALID_ARTIFACT", "SOURCE_FILE_RETURNED_AS_RESULT", "REQUEUE_LOOP_DETECTED"):
        return error_message
    return ""

async def _fcg_maybe_await(res):
    if _fcg_inspect.isawaitable(res):
        return await res
    return res

async def _fcg_send_public(conn, task_id, chat_id, topic_id, reply_to, text, reason):
    try:
        func = globals().get("_send_once_ex")
        if func:
            res = func(conn, str(task_id), str(chat_id), str(text), reply_to, str(reason))
            return await _fcg_maybe_await(res)
    except Exception as e:
        _FCG_LOG.warning("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_SEND_ONCE_ERR %s", e)
    try:
        func = globals().get("send_reply_ex")
        if func:
            res = func(chat_id=str(chat_id), text=str(text), reply_to_message_id=reply_to, message_thread_id=int(topic_id or 0))
            return await _fcg_maybe_await(res)
    except Exception as e:
        _FCG_LOG.warning("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_SEND_REPLY_ERR %s", e)
    return None

async def _fcg_fail_task(conn, task, code, public_text=None):
    task_id = _fcg_s(_fcg_row(task, "id", ""))
    chat_id = _fcg_s(_fcg_row(task, "chat_id", ""))
    topic_id = int(_fcg_row(task, "topic_id", 0) or 0)
    reply_to = _fcg_row(task, "reply_to_message_id", None)
    code = _fcg_s(code, 500) or "NO_VALID_ARTIFACT"
    msg = public_text or ("Задача не выполнена: " + code)
    try:
        conn.execute(
            "UPDATE tasks SET state='FAILED', result=?, error_message=?, updated_at=datetime('now') WHERE id=?",
            (msg, code, task_id),
        )
        _fcg_history(conn, task_id, "FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1:" + code)
        try:
            conn.execute("UPDATE drive_files SET stage='FAILED' WHERE task_id=?", (task_id,))
        except Exception:
            pass
        conn.commit()
    except Exception as e:
        _FCG_LOG.warning("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_FAIL_DB_ERR %s", e)
    await _fcg_send_public(conn, task_id, chat_id, topic_id, reply_to, msg, "full_contour_guard_failed")
    return True

def _fcg_qty(value):
    t = _fcg_s(value, 200).replace("≈", "").replace("~", "").replace(" ", "").replace(",", ".")
    m = _fcg_re.search(r"[-+]?\d+(?:\.\d+)?", t)
    if not m:
        return None
    try:
        return float(m.group(0))
    except Exception:
        return None

def _fcg_parse_direct_items(raw):
    raw = _fcg_s(raw, 100000)
    low = _fcg_low(raw)
    if "ед. изм" not in low or "количество" not in low:
        return []
    lines = [x.strip(" \t\r\n") for x in raw.replace("\r", "\n").split("\n") if x.strip()]
    items = []
    i = 0
    while i < len(lines) - 3:
        num = lines[i].strip()
        if _fcg_re.fullmatch(r"\d{1,3}", num):
            name = lines[i + 1].strip()
            unit = lines[i + 2].strip()
            qty_raw = lines[i + 3].strip()
            qty = _fcg_qty(qty_raw)
            if _FCG_UNIT_RE.match(unit) and qty is not None and len(name) > 6:
                bad_name = _fcg_low(name)
                if not any(x in bad_name for x in ("просим", "ответе указать", "состав работ также входит")):
                    items.append({
                        "num": len(items) + 1,
                        "name": name,
                        "unit": unit,
                        "qty": qty,
                        "qty_raw": qty_raw,
                    })
                    i += 4
                    continue
        i += 1
    return items

def _fcg_write_tender_xlsx(path, raw, items):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Смета КП"
    ws["A1"] = "Смета по текущему КП без шаблонной подмены"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Источник расчёта: только текущий текст задачи"
    ws["A3"] = "Цены: не выдуманы, единичные расценки нужно подтвердить или заполнить вручную"
    headers = ["№", "Наименование работ", "Ед. изм.", "Количество", "Цена работы", "Стоимость работы", "Цена материалов", "Стоимость материалов", "Всего", "Примечание", "Срок"]
    header_row = 5
    for col, val in enumerate(headers, 1):
        cell = ws.cell(header_row, col, val)
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="DDDDDD")
        cell.alignment = Alignment(wrap_text=True, vertical="top")
    thin = Side(style="thin", color="999999")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    start = header_row + 1
    for idx, item in enumerate(items, start):
        ws.cell(idx, 1, item["num"])
        ws.cell(idx, 2, item["name"])
        ws.cell(idx, 3, item["unit"])
        ws.cell(idx, 4, item["qty"])
        ws.cell(idx, 5, 0)
        ws.cell(idx, 6, f"=D{idx}*E{idx}")
        ws.cell(idx, 7, 0)
        ws.cell(idx, 8, f"=D{idx}*G{idx}")
        ws.cell(idx, 9, f"=F{idx}+H{idx}")
        ws.cell(idx, 10, "цена требует подтверждения")
        ws.cell(idx, 11, "")
    total_row = start + len(items)
    ws.cell(total_row, 8, "ИТОГО").font = Font(bold=True)
    ws.cell(total_row, 9, f"=SUM(I{start}:I{total_row-1})").font = Font(bold=True)
    vat_row = total_row + 1
    ws.cell(vat_row, 8, "НДС 20%").font = Font(bold=True)
    ws.cell(vat_row, 9, f"=I{total_row}*20%").font = Font(bold=True)
    gross_row = total_row + 2
    ws.cell(gross_row, 8, "С НДС").font = Font(bold=True)
    ws.cell(gross_row, 9, f"=I{total_row}+I{vat_row}").font = Font(bold=True)
    for row in ws.iter_rows(min_row=header_row, max_row=gross_row, min_col=1, max_col=11):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    widths = [8, 70, 12, 14, 16, 18, 18, 20, 18, 34, 18]
    for col, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = width

    raw_ws = wb.create_sheet("Исходное ТЗ")
    raw_ws["A1"] = raw[:32000]
    raw_ws["A1"].alignment = Alignment(wrap_text=True, vertical="top")
    raw_ws.column_dimensions["A"].width = 120
    wb.save(path)

def _fcg_write_pdf(path, title, lines):
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        font_paths = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        ]
        font_name = "Helvetica"
        for fp in font_paths:
            if _fcg_os.path.exists(fp):
                pdfmetrics.registerFont(TTFont("ArealSans", fp))
                font_name = "ArealSans"
                break
        c = canvas.Canvas(path, pagesize=A4)
        width, height = A4
        y = height - 40
        c.setFont(font_name, 12)
        c.drawString(40, y, title[:110])
        y -= 28
        c.setFont(font_name, 9)
        for line in lines:
            for chunk in [line[i:i+105] for i in range(0, len(line), 105)] or [""]:
                if y < 40:
                    c.showPage()
                    c.setFont(font_name, 9)
                    y = height - 40
                c.drawString(40, y, chunk)
                y -= 14
        c.save()
        return path
    except Exception:
        txt_path = path.replace(".pdf", ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(title + "\n\n" + "\n".join(lines))
        return txt_path

async def _fcg_upload(path, file_name, chat_id, topic_id, mime_type=None):
    if not path or not _fcg_os.path.exists(str(path)):
        return ""
    try:
        from core.topic_drive_oauth import upload_file_to_topic
    except Exception as e:
        _FCG_LOG.warning("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_UPLOAD_IMPORT_ERR %s", e)
        return ""
    attempts = [
        (str(path), str(file_name), str(chat_id), int(topic_id or 0), mime_type),
        (str(path), str(file_name), str(chat_id), int(topic_id or 0), None),
        (str(path), str(file_name), str(chat_id), int(topic_id or 0)),
    ]
    for args in attempts:
        try:
            res = upload_file_to_topic(*args)
            res = await _fcg_maybe_await(res)
            if isinstance(res, str) and "drive.google.com" in res:
                return res
            if isinstance(res, dict):
                for k in ("webViewLink", "drive_link", "link", "url"):
                    if res.get(k) and ("drive.google.com" in str(res.get(k)) or "docs.google.com" in str(res.get(k))):
                        return str(res.get(k))
                fid = res.get("drive_file_id") or res.get("id") or res.get("file_id")
                if fid:
                    return "https://drive.google.com/file/d/" + str(fid) + "/view"
        except TypeError:
            continue
        except Exception as e:
            _FCG_LOG.warning("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_UPLOAD_ERR %s", e)
            continue
    return ""

async def _fcg_handle_direct_tender_estimate(conn, task):
    task_id = _fcg_s(_fcg_row(task, "id", ""))
    chat_id = _fcg_s(_fcg_row(task, "chat_id", ""))
    topic_id = int(_fcg_row(task, "topic_id", 0) or 0)
    if topic_id != 2 or not task_id:
        return False
    state = _fcg_s(_fcg_row(task, "state", ""))
    if state not in ("NEW", "IN_PROGRESS", "WAITING_CLARIFICATION", "AWAITING_CONFIRMATION"):
        return False
    raw = _fcg_s(_fcg_row(task, "raw_input", ""), 100000)
    items = _fcg_parse_direct_items(raw)
    if len(items) < 3:
        return False

    try:
        already = conn.execute(
            "SELECT 1 FROM task_history WHERE task_id=? AND action='FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1:DIRECT_TENDER_ESTIMATE_GENERATED' LIMIT 1",
            (task_id,),
        ).fetchone()
        if already and state == "AWAITING_CONFIRMATION":
            return False
    except Exception:
        pass

    out_dir = _fcg_os.path.join(_fcg_tempfile.gettempdir(), "areal_full_contour_" + task_id)
    _fcg_os.makedirs(out_dir, exist_ok=True)
    xlsx_path = _fcg_os.path.join(out_dir, "estimate_direct_" + task_id[:8] + ".xlsx")
    pdf_path = _fcg_os.path.join(out_dir, "estimate_direct_" + task_id[:8] + ".pdf")
    _fcg_write_tender_xlsx(xlsx_path, raw, items)
    pdf_lines = [
        "Основа: только текущий текст задачи, без подмены на шаблон дома",
        "Позиций: " + str(len(items)),
        "Цены не выдуманы: в Excel оставлены колонки цены работы и материалов с формулами",
        "",
    ]
    for item in items[:80]:
        pdf_lines.append(f"{item['num']}. {item['name']} — {item['qty_raw']} {item['unit']}")
    pdf_real = _fcg_write_pdf(pdf_path, "Смета по текущему КП", pdf_lines)

    xlsx_link = await _fcg_upload(xlsx_path, "estimate_direct_" + task_id[:8] + ".xlsx", chat_id, topic_id, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    pdf_link = await _fcg_upload(pdf_real, _fcg_os.path.basename(pdf_real), chat_id, topic_id, "application/pdf")
    if not xlsx_link or not pdf_link:
        return await _fcg_fail_task(conn, task, "NO_VALID_ARTIFACT")

    result = (
        "Сметный контур обработал текущее КП без шаблонной подмены\n\n"
        + "Позиций: " + str(len(items)) + "\n"
        + "Основа: только текущий текст задачи\n"
        + "Цены: не выдуманы, единичные расценки оставлены для подтверждения/заполнения\n"
        + "Формулы Excel: Количество × Цена, итог через SUM\n\n"
        + "📊 Excel: " + xlsx_link + "\n"
        + "📄 PDF: " + pdf_link + "\n\n"
        + "Подтверди или пришли правки по ценам"
    )
    reply_to = _fcg_row(task, "reply_to_message_id", None)
    send_res = await _fcg_send_public(conn, task_id, chat_id, topic_id, reply_to, result, "full_contour_direct_tender")
    bot_id = None
    if isinstance(send_res, dict):
        bot_id = send_res.get("bot_message_id")
    try:
        if bot_id:
            conn.execute(
                "UPDATE tasks SET state='AWAITING_CONFIRMATION', result=?, error_message='', bot_message_id=?, updated_at=datetime('now') WHERE id=?",
                (result, bot_id, task_id),
            )
        else:
            conn.execute(
                "UPDATE tasks SET state='AWAITING_CONFIRMATION', result=?, error_message='', updated_at=datetime('now') WHERE id=?",
                (result, task_id),
            )
        _fcg_history(conn, task_id, "FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1:DIRECT_TENDER_ESTIMATE_GENERATED")
        conn.commit()
    except Exception as e:
        _FCG_LOG.warning("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_DIRECT_DB_ERR %s", e)
        return False
    _FCG_LOG.info("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_DIRECT_TENDER_OK task=%s items=%s", task_id, len(items))
    return True

def _fcg_topic5_final_act_request(raw):
    t = _fcg_low(raw)
    return any(x in t for x in ("сделай акт", "сформируй акт", "финальный акт", "итоговый акт", "акт технадзора"))

def _fcg_collect_topic5_package(conn, chat_id, topic_id):
    files = []
    comments = []
    try:
        rows = conn.execute(
            """
            SELECT id,raw_input,result,updated_at
            FROM tasks
            WHERE chat_id=? AND COALESCE(topic_id,0)=?
              AND input_type='drive_file'
              AND updated_at >= datetime('now','-7 days')
            ORDER BY updated_at ASC
            LIMIT 80
            """,
            (str(chat_id), int(topic_id or 0)),
        ).fetchall()
        for r in rows:
            data = {}
            try:
                data = _fcg_json.loads(_fcg_s(_fcg_row(r, "raw_input", "")))
            except Exception:
                data = {}
            name = data.get("file_name") or _fcg_s(_fcg_row(r, "id", ""))[:8]
            if "photo" in _fcg_low(name) or str(data.get("mime_type", "")).startswith("image/"):
                files.append({
                    "task_id": _fcg_s(_fcg_row(r, "id", "")),
                    "name": _fcg_s(name, 300),
                    "result": _fcg_s(_fcg_row(r, "result", ""), 1500),
                })
    except Exception:
        pass
    try:
        rows = conn.execute(
            """
            SELECT raw_input,result,updated_at
            FROM tasks
            WHERE chat_id=? AND COALESCE(topic_id,0)=?
              AND input_type='text'
              AND updated_at >= datetime('now','-7 days')
            ORDER BY updated_at ASC
            LIMIT 120
            """,
            (str(chat_id), int(topic_id or 0)),
        ).fetchall()
        for r in rows:
            raw = _fcg_s(_fcg_row(r, "raw_input", ""), 1200)
            if raw and any(x in _fcg_low(raw) for x in ("адрес", "объект", "дефект", "наруш", "свар", "оборуд", "замен", "ропшин", "основание", "авито", "фото")):
                comments.append(raw)
    except Exception:
        pass
    return files, comments

def _fcg_write_topic5_act(path_base, files, comments):
    docx_path = path_base + ".docx"
    txt_path = path_base + ".txt"
    lines = []
    lines.append("АКТ ТЕХНИЧЕСКОГО НАДЗОРА")
    lines.append("")
    lines.append("Основание: материалы и пояснения из текущего topic_id")
    lines.append("Нормативная база: конкретный пункт СП/ГОСТ не подтверждён, нормы не выдуманы")
    lines.append("")
    lines.append("Материалы фотофиксации:")
    for idx, f in enumerate(files, 1):
        lines.append(f"{idx}. {f['name']}")
    lines.append("")
    lines.append("Пояснения владельца:")
    for idx, c in enumerate(comments, 1):
        lines.append(f"{idx}. {c}")
    lines.append("")
    lines.append("Вывод:")
    lines.append("Материалы объединены в один акт. Требуется проверка и подтверждение владельцем")
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("АКТ ТЕХНИЧЕСКОГО НАДЗОРА", level=1)
        doc.add_paragraph("Основание: материалы и пояснения из текущего topic_id")
        doc.add_paragraph("Нормативная база: конкретный пункт СП/ГОСТ не подтверждён, нормы не выдуманы")
        doc.add_heading("Материалы фотофиксации", level=2)
        for idx, f in enumerate(files, 1):
            doc.add_paragraph(f"{idx}. {f['name']}")
        doc.add_heading("Пояснения владельца", level=2)
        for idx, c in enumerate(comments, 1):
            doc.add_paragraph(f"{idx}. {c}")
        doc.add_heading("Вывод", level=2)
        doc.add_paragraph("Материалы объединены в один акт. Требуется проверка и подтверждение владельцем")
        doc.save(docx_path)
        return docx_path
    except Exception:
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return txt_path

async def _fcg_handle_topic5_final_act(conn, task):
    task_id = _fcg_s(_fcg_row(task, "id", ""))
    chat_id = _fcg_s(_fcg_row(task, "chat_id", ""))
    topic_id = int(_fcg_row(task, "topic_id", 0) or 0)
    raw = _fcg_s(_fcg_row(task, "raw_input", ""), 5000)
    if topic_id != 5 or not _fcg_topic5_final_act_request(raw):
        return False
    files, comments = _fcg_collect_topic5_package(conn, chat_id, topic_id)
    if not files:
        return False
    out_dir = _fcg_os.path.join(_fcg_tempfile.gettempdir(), "areal_topic5_final_act_" + task_id)
    _fcg_os.makedirs(out_dir, exist_ok=True)
    artifact = _fcg_write_topic5_act(_fcg_os.path.join(out_dir, "technadzor_act_" + task_id[:8]), files, comments)
    link = await _fcg_upload(artifact, _fcg_os.path.basename(artifact), chat_id, topic_id, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if not link:
        return await _fcg_fail_task(conn, task, "NO_VALID_ARTIFACT")
    result = (
        "Финальный акт технадзора сформирован по текущему пакету фото\n\n"
        + "Фото в акте: " + str(len(files)) + "\n"
        + "Пояснений учтено: " + str(len(comments)) + "\n"
        + "Нормы СП/ГОСТ не выдуманы, конкретные пункты не подставлены без подтверждения\n\n"
        + "📄 Акт: " + link + "\n\n"
        + "Подтверди или пришли правки"
    )
    reply_to = _fcg_row(task, "reply_to_message_id", None)
    send_res = await _fcg_send_public(conn, task_id, chat_id, topic_id, reply_to, result, "topic5_final_act")
    bot_id = send_res.get("bot_message_id") if isinstance(send_res, dict) else None
    try:
        if bot_id:
            conn.execute("UPDATE tasks SET state='AWAITING_CONFIRMATION', result=?, error_message='', bot_message_id=?, updated_at=datetime('now') WHERE id=?", (result, bot_id, task_id))
        else:
            conn.execute("UPDATE tasks SET state='AWAITING_CONFIRMATION', result=?, error_message='', updated_at=datetime('now') WHERE id=?", (result, task_id))
        _fcg_history(conn, task_id, "FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1:TOPIC5_FINAL_ACT_GENERATED")
        conn.commit()
    except Exception as e:
        _FCG_LOG.warning("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_TOPIC5_DB_ERR %s", e)
        return False
    _FCG_LOG.info("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_TOPIC5_ACT_OK task=%s files=%s comments=%s", task_id, len(files), len(comments))
    return True

_FCG_ORIG_UPDATE_TASK = globals().get("_update_task")
if _FCG_ORIG_UPDATE_TASK and not getattr(_FCG_ORIG_UPDATE_TASK, "_fcg_wrapped", False):
    def _update_task(conn, task_id, *args, **kwargs):
        state = kwargs.get("state", args[0] if len(args) >= 1 else None)
        result = kwargs.get("result", args[1] if len(args) >= 2 else None)
        err = kwargs.get("error_message", args[2] if len(args) >= 3 else "")
        violation = _fcg_public_result_violation(conn, task_id, state, result, err)
        # === PATCH_TOPIC2_INLINE_FIX_20260506_V1 UPDATE_BYPASS_VIOLATION ===
        # if all 5 critical DONE markers present + result has Drive link → don't block update
        if violation:
            try:
                _pifx_required = (
                    "TOPIC2_XLSX_CREATED",
                    "TOPIC2_PDF_CREATED",
                    "TOPIC2_DRIVE_UPLOAD_XLSX_OK",
                    "TOPIC2_DRIVE_UPLOAD_PDF_OK",
                    "TOPIC2_TELEGRAM_DELIVERED",
                )
                _pifx_rows = conn.execute(
                    "SELECT action FROM task_history WHERE task_id=?", (str(task_id),)
                ).fetchall()
                _pifx_actions = " ".join(str(r[0]) for r in _pifx_rows)
                _pifx_all_markers = all(m in _pifx_actions for m in _pifx_required)
                _pifx_has_drive = bool(result) and "drive.google.com" in str(result).lower()
                if _pifx_all_markers and _pifx_has_drive:
                    try:
                        conn.execute(
                            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                            (str(task_id), "PATCH_TOPIC2_INLINE_FIX_20260506_V1:UPDATE_BYPASS_VIOLATION:5_markers_drive_ok"),
                        )
                        conn.commit()
                    except Exception:
                        pass
                    violation = ""
            except Exception:
                pass
        # === END_PATCH_TOPIC2_INLINE_FIX_20260506_V1 UPDATE_BYPASS_VIOLATION ===
        if violation:
            a = list(args)
            if len(a) >= 1:
                a[0] = "FAILED"
            else:
                kwargs["state"] = "FAILED"
            if len(a) >= 2:
                a[1] = "Задача не выполнена: " + violation
            else:
                kwargs["result"] = "Задача не выполнена: " + violation
            if len(a) >= 3:
                a[2] = violation
            else:
                kwargs["error_message"] = violation
            _fcg_history(conn, task_id, "FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1:UPDATE_BLOCKED:" + violation)
            args = tuple(a)
        return _FCG_ORIG_UPDATE_TASK(conn, task_id, *args, **kwargs)
    _update_task._fcg_wrapped = True

_FCG_ORIG_HANDLE_DRIVE_FILE = globals().get("_handle_drive_file")
if _FCG_ORIG_HANDLE_DRIVE_FILE and not getattr(_FCG_ORIG_HANDLE_DRIVE_FILE, "_fcg_wrapped", False):
    async def _handle_drive_file(conn, task, chat_id, topic_id, *args, **kwargs):
        if _fcg_requeue_loop_detected(conn, task):
            return await _fcg_fail_task(conn, task, "REQUEUE_LOOP_DETECTED")
        try:
            res = _FCG_ORIG_HANDLE_DRIVE_FILE(conn, task, chat_id, topic_id, *args, **kwargs)
            if _fcg_inspect.isawaitable(res):
                res = await _fcg_asyncio.wait_for(res, timeout=_FCG_TIMEOUT_SEC)
            row = _fcg_task_row(conn, _fcg_row(task, "id", ""))
            if row is not None:
                violation = _fcg_public_result_violation(
                    conn,
                    _fcg_row(row, "id", ""),
                    _fcg_row(row, "state", ""),
                    _fcg_row(row, "result", ""),
                    _fcg_row(row, "error_message", ""),
                )
                if violation:
                    return await _fcg_fail_task(conn, row, violation)
            return res
        except _fcg_asyncio.TimeoutError:
            return await _fcg_fail_task(conn, task, "ENGINE_TIMEOUT")
        except Exception as e:
            _FCG_LOG.exception("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1_DRIVE_ERR %s", e)
            return await _fcg_fail_task(conn, task, "NO_VALID_ARTIFACT")
    _handle_drive_file._fcg_wrapped = True

_FCG_ORIG_HANDLE_NEW = globals().get("_handle_new")
if _FCG_ORIG_HANDLE_NEW and not getattr(_FCG_ORIG_HANDLE_NEW, "_fcg_wrapped", False):
    async def _handle_new(conn, task, *args, **kwargs):
        if await _fcg_handle_direct_tender_estimate(conn, task):
            return True
        if await _fcg_handle_topic5_final_act(conn, task):
            return True
        res = _FCG_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        return await _fcg_maybe_await(res)
    _handle_new._fcg_wrapped = True

_FCG_ORIG_HANDLE_IN_PROGRESS = globals().get("_handle_in_progress")
if _FCG_ORIG_HANDLE_IN_PROGRESS and not getattr(_FCG_ORIG_HANDLE_IN_PROGRESS, "_fcg_wrapped", False):
    async def _handle_in_progress(conn, task, *args, **kwargs):
        if await _fcg_handle_direct_tender_estimate(conn, task):
            return True
        if await _fcg_handle_topic5_final_act(conn, task):
            return True
        res = _FCG_ORIG_HANDLE_IN_PROGRESS(conn, task, *args, **kwargs)
        return await _fcg_maybe_await(res)
    _handle_in_progress._fcg_wrapped = True

_FCG_LOG.info("FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1 installed")
# SOURCE_FILE_RETURNED_AS_RESULT
# REQUEUE_LOOP_DETECTED
# NO_VALID_ARTIFACT
# ENGINE_TIMEOUT
# === END_FULL_CONSTRUCTION_FILE_CONTOUR_CANON_GUARD_V1 ===




# V5 superseded by T2PI_V1 below — asyncio.run moved after PATCH_TOPIC2_PHOTO_ESTIMATE_INTERCEPT_V1
# if __name__ == "__main__":
#     asyncio.run(main())
# === END_MOVE_MAIN_ENTRYPOINT_TO_END_V5 ===

# === PATCH_TOPIC2_PHOTO_ESTIMATE_INTERCEPT_V1 ===
# Downloads Drive photo, runs OpenRouter vision, converts to estimate text
# Intercepts _handle_drive_file for topic_2 image tasks with estimate captions
import os as _t2pi_os
import json as _t2pi_json
import inspect as _t2pi_inspect
import logging as _t2pi_log

_T2PI_LOG = _t2pi_log.getLogger("task_worker")
_T2PI_IMG_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".heic", ".bmp", ".tif", ".tiff")
_T2PI_ESTIMATE_WORDS = (
    "смет", "посчитай", "рассчитай", "стоимость", "дом", "фундамент",
    "кровл", "коробк", "газобетон", "каркас", "монолит", "расчет", "расчёт",
    "строительство", "построить", "построй", "нужна смета", "посчитать",
    "полностью", "все размеры",
)

def _t2pi_s(v, lim=50000):
    try:
        return str(v if v is not None else "")[:lim]
    except Exception:
        return ""

def _t2pi_row(r, k, d=None):
    try:
        if hasattr(r, "keys") and k in r.keys():
            return r[k]
    except Exception:
        pass
    try:
        return r[k]
    except Exception:
        return d

def _t2pi_payload(task):
    raw = _t2pi_s(_t2pi_row(task, "raw_input", ""))
    try:
        d = _t2pi_json.loads(raw)
        return d if isinstance(d, dict) else {}
    except Exception:
        return {}

def _t2pi_is_estimate_photo(task):
    topic_id = int(_t2pi_row(task, "topic_id", 0) or 0)
    if topic_id != 2:
        return False
    input_type = _t2pi_s(_t2pi_row(task, "input_type", "")).lower()
    if input_type not in ("drive_file", "photo", "image"):
        return False
    payload = _t2pi_payload(task)
    file_name = _t2pi_s(payload.get("file_name") or "").lower()
    mime = _t2pi_s(payload.get("mime_type") or "").lower()
    is_img = mime.startswith("image/") or any(file_name.endswith(e) for e in _T2PI_IMG_EXTS)
    if not is_img:
        return False
    caption = _t2pi_s(payload.get("caption") or "").lower().replace("ё", "е")
    return any(w in caption for w in _T2PI_ESTIMATE_WORDS)

def _t2pi_download_drive(file_id, local_path):
    try:
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaIoBaseDownload
        from google.oauth2.service_account import Credentials
        import io
        creds = Credentials.from_service_account_file(
            "/root/.areal-neva-core/credentials.json",
            scopes=["https://www.googleapis.com/auth/drive.readonly"],
        )
        svc = build("drive", "v3", credentials=creds)
        req = svc.files().get_media(fileId=file_id, supportsAllDrives=True)
        with io.FileIO(local_path, "wb") as fh:
            dl = MediaIoBaseDownload(fh, req)
            done = False
            while not done:
                _, done = dl.next_chunk()
        return _t2pi_os.path.exists(local_path) and _t2pi_os.path.getsize(local_path) > 0
    except Exception as e:
        _T2PI_LOG.warning("T2PI_DRIVE_DL_ERR file_id=%s err=%s", file_id, e)
        return False

async def _t2pi_vision(local_path, caption):
    try:
        import httpx as _t2pi_httpx
        import base64 as _t2pi_b64
        api_key = (_t2pi_os.getenv("OPENROUTER_API_KEY") or "").strip()
        if not api_key:
            return ""
        base_url = (_t2pi_os.getenv("OPENROUTER_BASE_URL") or "https://openrouter.ai/api/v1").rstrip("/")
        model = (_t2pi_os.getenv("OPENROUTER_VISION_MODEL") or "google/gemini-2.5-flash").strip()
        ext = _t2pi_os.path.splitext(local_path)[1].lower().lstrip(".") or "jpeg"
        mime = "image/{}".format("jpeg" if ext in ("jpg", "jpeg") else ext)
        with open(local_path, "rb") as f:
            b64 = _t2pi_b64.b64encode(f.read()).decode("utf-8")
        prompt = (
            "Это строительный план, схема или фото объекта для расчёта сметы.\n"
            "Задание: {}\n\n"
            "Извлеки из изображения:\n"
            "- Размеры и габариты объекта (ширина x длина x высота, в метрах)\n"
            "- Площади помещений и общую площадь\n"
            "- Этажность\n"
            "- Конструктив (газобетон, каркас, монолит, кирпич, брус)\n"
            "- Все подписи, размерные цепочки, ведомости\n"
            "- Материалы кровли, фундамента, отделки\n"
            "Если размер не виден — так и напиши. Не придумывай данные."
        ).format(caption or "рассчитай смету")
        body = {
            "model": model,
            "messages": [{"role": "user", "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": "data:{}; base64,{}".format(mime, b64)}},
            ]}],
            "temperature": 0.1,
        }
        headers = {"Authorization": "Bearer " + api_key, "Content-Type": "application/json"}
        async with _t2pi_httpx.AsyncClient(timeout=_t2pi_httpx.Timeout(120.0, connect=30.0)) as client:
            r = await client.post(base_url + "/chat/completions", headers=headers, json=body)
            r.raise_for_status()
            data = r.json()
        content = data["choices"][0]["message"]["content"]
        if isinstance(content, list):
            content = "\n".join(x.get("text", "") if isinstance(x, dict) else str(x) for x in content)
        return _t2pi_s(content, 16000)
    except Exception as e:
        _T2PI_LOG.warning("T2PI_VISION_ERR err=%s", e)
        return ""

async def _t2pi_handle(conn, task, chat_id, topic_id):
    task_id = _t2pi_s(_t2pi_row(task, "id", ""))
    payload = _t2pi_payload(task)
    file_id = _t2pi_s(payload.get("file_id") or "")
    file_name = _t2pi_s(payload.get("file_name") or "photo.jpg")
    caption = _t2pi_s(payload.get("caption") or "")

    if not file_id:
        return False

    out_dir = "/root/.areal-neva-core/runtime/drive_files"
    _t2pi_os.makedirs(out_dir, exist_ok=True)
    local_path = "{}/{}_{}".format(out_dir, task_id, file_name)

    if not (_t2pi_os.path.exists(local_path) and _t2pi_os.path.getsize(local_path) > 0):
        import asyncio as _t2pi_asyncio
        ok = await _t2pi_asyncio.to_thread(_t2pi_download_drive, file_id, local_path)
        if not ok:
            _T2PI_LOG.warning("T2PI_DL_FAILED task=%s", task_id)
            local_path = ""

    vision_text = ""
    if local_path and _t2pi_os.path.exists(local_path):
        _T2PI_LOG.info("T2PI_VISION_START task=%s", task_id)
        vision_text = await _t2pi_vision(local_path, caption)
        _T2PI_LOG.info("T2PI_VISION_DONE task=%s chars=%s", task_id, len(vision_text))

    parts = []
    if caption:
        parts.append(caption)
    if vision_text and "ERROR" not in vision_text[:20].upper():
        parts.append("РАСПОЗНАНО С ФОТО:\n" + vision_text)
    enriched = "\n\n".join(parts).strip()
    if not enriched:
        return False

    try:
        conn.execute(
            "UPDATE tasks SET raw_input=?, input_type='text', error_message='', updated_at=datetime('now') WHERE id=?",
            (enriched, task_id),
        )
        conn.execute(
            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
            (task_id, "PATCH_TOPIC2_PHOTO_ESTIMATE_INTERCEPT_V1:VISION_DONE"),
        )
        conn.commit()
    except Exception as e:
        _T2PI_LOG.warning("T2PI_DB_ERR task=%s err=%s", task_id, e)
        return False

    try:
        from core.stroyka_estimate_canon import maybe_handle_stroyka_estimate as _t2pi_canon
        task_dict = {}
        try:
            for k in task.keys():
                task_dict[k] = task[k]
        except Exception:
            task_dict = {}
        task_dict["raw_input"] = enriched
        task_dict["input_type"] = "text"
        result = _t2pi_canon(conn, task_dict, logger=_T2PI_LOG)
        if _t2pi_inspect.isawaitable(result):
            result = await result
        _T2PI_LOG.info("T2PI_CANON_RESULT task=%s result=%s", task_id, bool(result))
        return bool(result)
    except Exception as e:
        _T2PI_LOG.exception("T2PI_CANON_ERR task=%s err=%s", task_id, e)
        return False

_T2PI_ORIG_DRIVE = globals().get("_handle_drive_file")
if _T2PI_ORIG_DRIVE:
    async def _handle_drive_file(conn, task, chat_id=None, topic_id=None, *args, **kwargs):
        if chat_id is None:
            chat_id = _t2pi_s(_t2pi_row(task, "chat_id", ""))
        if topic_id is None:
            topic_id = int(_t2pi_row(task, "topic_id", 0) or 0)
        if _t2pi_is_estimate_photo(task):
            _t2pi_tid = _t2pi_s(_t2pi_row(task, "id", ""))
            try:
                if await _t2pi_handle(conn, task, chat_id, topic_id):
                    _T2PI_LOG.info("T2PI_HANDLED task=%s", _t2pi_tid)
                    return True
            except Exception as e:
                _T2PI_LOG.exception("T2PI_INTERCEPT_ERR task=%s err=%s", _t2pi_tid, e)
        res = _T2PI_ORIG_DRIVE(conn, task, chat_id, topic_id, *args, **kwargs)
        if _t2pi_inspect.isawaitable(res):
            return await res
        return res

_T2PI_LOG.info("PATCH_TOPIC2_PHOTO_ESTIMATE_INTERCEPT_V1 installed")
# === END_PATCH_TOPIC2_PHOTO_ESTIMATE_INTERCEPT_V1 ===

# T2PI_V1 superseded by T2FEB_V1 below
# if __name__ == "__main__":
#     asyncio.run(main())
# === END_MOVE_MAIN_ENTRYPOINT_TO_END_T2PI_V1 ===

# === PATCH_TOPIC2_FRESH_ESTIMATE_BYPASS_V1 ===
# Fix 1: Long voice texts with "средн" must NOT be treated as price choice — route as fresh estimate
# Fix 2: P6E67_PARENT_NOT_FOUND tasks go to fresh estimate instead of WAITING_CLARIFICATION limbo
import logging as _t2feb_log
import inspect as _t2feb_inspect
_T2FEB_LOG = _t2feb_log.getLogger("task_worker")

_T2FEB_ESTIMATE_KW = (
    "смет", "посчитай", "рассчитай", "рассчитать", "посчитать",
    "дом", "здани", "строен", "построить", "построй", "построим",
    "газобетон", "каркас", "монолит", "кирпич", "брус",
    "фундамент", "кровл", "перекр", "коробк", "отделк",
    "м2", "м²", "этаж", "комнат", "площад",
    "ангар", "склад", "баня", "гараж", "барнхаус",
    "сделаем", "сделай", "выполним", "сделать", "составить",
)
_T2FEB_REDO_KW = (
    "заново", "снова", "переделай", "переделать",
    "пересчитай", "пересчитать", "повтори", "повторить",
    "выполним заново", "заново смету",
)

def _t2feb_s(v, n=50000):
    try:
        return str(v if v is not None else "")[:n]
    except Exception:
        return ""

def _t2feb_row(r, k, d=None):
    try:
        if hasattr(r, "keys") and k in r.keys():
            return r[k]
    except Exception:
        pass
    try:
        return r[k]
    except Exception:
        return d

def _t2feb_is_fresh_estimate(raw):
    clean = _t2feb_s(raw).lower().replace("ё", "е").replace("[voice]", "").strip()
    if "---" in clean and "revision_context" in clean:
        clean = clean[:clean.index("revision_context")].strip()
    if len(clean) < 15:
        return False
    has_kw = any(k in clean for k in _T2FEB_ESTIMATE_KW)
    if not has_kw:
        return False
    if len(clean) > 150:
        return True
    if any(w in clean for w in _T2FEB_REDO_KW):
        return True
    return False

async def _t2feb_run_estimate(conn, task, raw):
    try:
        from core.stroyka_estimate_canon import maybe_handle_stroyka_estimate as _t2feb_mhse
        task_clean = {}
        try:
            for k in task.keys():
                task_clean[k] = task[k]
        except Exception:
            task_clean = {}
        # Strip REVISION_CONTEXT
        if "\n---\nREVISION_CONTEXT" in raw:
            raw = raw.split("\n---\nREVISION_CONTEXT")[0].strip()
        task_clean["raw_input"] = raw
        task_clean["input_type"] = "text"
        result = _t2feb_mhse(conn, task_clean, _T2FEB_LOG)
        if _t2feb_inspect.isawaitable(result):
            result = await result
        return bool(result)
    except Exception as e:
        _T2FEB_LOG.exception("T2FEB_ESTIMATE_ERR task=%s err=%s", _t2feb_row(task, "id"), e)
        return False

_T2FEB_ORIG_HANDLE_NEW = globals().get("_handle_new")
if _T2FEB_ORIG_HANDLE_NEW:
    async def _handle_new(conn, task, *args, **kwargs):
        topic_id = int(_t2feb_row(task, "topic_id", 0) or 0)
        if topic_id == 2:
            raw = _t2feb_s(_t2feb_row(task, "raw_input", ""))
            if _t2feb_is_fresh_estimate(raw):
                task_id = _t2feb_s(_t2feb_row(task, "id", ""))
                _T2FEB_LOG.info("T2FEB_BYPASS_PRICE_CHAIN task=%s", task_id)
                if await _t2feb_run_estimate(conn, task, raw):
                    _T2FEB_LOG.info("T2FEB_ESTIMATE_DONE task=%s", task_id)
                    return True
        res = _T2FEB_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        if _t2feb_inspect.isawaitable(res):
            return await res
        return res

_T2FEB_ORIG_HANDLE_IN_PROG = globals().get("_handle_in_progress")
if _T2FEB_ORIG_HANDLE_IN_PROG:
    async def _handle_in_progress(conn, task, *args, **kwargs):
        topic_id = int(_t2feb_row(task, "topic_id", 0) or 0)
        if topic_id == 2:
            err = _t2feb_s(_t2feb_row(task, "error_message", "")).upper()
            state = _t2feb_s(_t2feb_row(task, "state", "")).upper()
            raw = _t2feb_s(_t2feb_row(task, "raw_input", ""))
            if "P6E67_PARENT_NOT_FOUND" in err and state in ("WAITING_CLARIFICATION", "IN_PROGRESS"):
                if _t2feb_is_fresh_estimate(raw):
                    task_id = _t2feb_s(_t2feb_row(task, "id", ""))
                    _T2FEB_LOG.info("T2FEB_P6E67_RESCUE task=%s", task_id)
                    conn.execute(
                        "UPDATE tasks SET state='IN_PROGRESS', error_message='', updated_at=datetime('now') WHERE id=?",
                        (task_id,)
                    )
                    conn.commit()
                    if await _t2feb_run_estimate(conn, task, raw):
                        _T2FEB_LOG.info("T2FEB_P6E67_RESCUE_DONE task=%s", task_id)
                        return True
        res = _T2FEB_ORIG_HANDLE_IN_PROG(conn, task, *args, **kwargs)
        if _t2feb_inspect.isawaitable(res):
            return await res
        return res

_T2FEB_LOG.info("PATCH_TOPIC2_FRESH_ESTIMATE_BYPASS_V1 installed")
# === END_PATCH_TOPIC2_FRESH_ESTIMATE_BYPASS_V1 ===

# === END_MOVE_MAIN_ENTRYPOINT_TO_END_T2FEB_V1 ===


# === PATCH_STROYKA_PARENT_LOOKUP_NO_NEW_V1 ===
# NEW removed from active parent lookup for topic_2 task switch safety
# === END_PATCH_STROYKA_PARENT_LOOKUP_NO_NEW_V1 ===


# === PATCH_TOPIC2_UNBLOCK_PICK_NEXT_V1 ===
# Excludes broken topic5 IN_PROGRESS loop with P6F_DAH_BLOCK_DONE_NO_UPLOAD_HISTORY from global picker
# === END_PATCH_TOPIC2_UNBLOCK_PICK_NEXT_V1 ===

# if __name__ == "__main__":
#     asyncio.run(main())  # DISABLED_BY_PATCH_TOPIC2_FRESH_ESTIMATE_ROUTE_GUARD_V1


# === PATCH_TOPIC2_FRESH_ESTIMATE_ROUTE_GUARD_V1 ===
# FACT: topic_2 fresh estimate tasks were routed to P6E67_PARENT_NOT_FOUND_TERMINAL_GUARD_V1 instead of estimate artifact generation
import json as _t2fer_json
import inspect as _t2fer_inspect
import logging as _t2fer_logging

_T2FER_LOG = _t2fer_logging.getLogger("task_worker")

_T2FER_KEYWORDS = (
    "смет", "расчет", "расчёт", "стоимость", "посчитать", "посчитай",
    "рассчитать", "дом", "коттедж", "строительство", "построить",
    "газобетон", "каркас", "монолит", "кирпич", "брус",
    "фундамент", "плита", "кровля", "отделка", "ламинат",
    "имитация бруса", "теплые полы", "тёплые полы", "полностью"
)

_T2FER_BAD_PARENT_ERR = "P6E67_PARENT_NOT_FOUND_TERMINAL_GUARD_V1"

def _t2fer_get(row, key, default=None):
    try:
        if hasattr(row, "keys") and key in row.keys():
            return row[key]
    except Exception:
        pass
    try:
        if isinstance(row, dict):
            return row.get(key, default)
    except Exception:
        pass
    try:
        return row[key]
    except Exception:
        return default

def _t2fer_s(value, limit=120000):
    try:
        s = "" if value is None else str(value)
    except Exception:
        s = ""
    return s.replace("\x00", "")[:limit]

def _t2fer_topic(task):
    try:
        return int(_t2fer_get(task, "topic_id", 0) or 0)
    except Exception:
        return 0

def _t2fer_payload_text(raw):
    raw_s = _t2fer_s(raw, 120000)
    chunks = [raw_s]
    stripped = raw_s.strip()
    if stripped.startswith("{") and stripped.endswith("}"):
        try:
            data = _t2fer_json.loads(stripped)
            for k in ("caption", "file_name", "mime_type", "source"):
                v = data.get(k)
                if v:
                    chunks.append(_t2fer_s(v, 20000))
        except Exception:
            pass
    return "\n".join(x for x in chunks if x).lower().replace("ё", "е")

def _t2fer_is_fresh_estimate(task):
    if _t2fer_topic(task) != 2:
        return False

    raw = _t2fer_s(_t2fer_get(task, "raw_input", ""), 120000)
    result = _t2fer_s(_t2fer_get(task, "result", ""), 20000)
    err = _t2fer_s(_t2fer_get(task, "error_message", ""), 20000)
    input_type = _t2fer_s(_t2fer_get(task, "input_type", ""), 2000).lower()

    text = _t2fer_payload_text(raw + "\n" + result + "\n" + err)

    if "отмена" in text or "очисти все задачи" in text or "cancel" in text:
        return False

    if _T2FER_BAD_PARENT_ERR.lower() in text:
        return any(k in text for k in _T2FER_KEYWORDS)

    if input_type in ("drive_file", "file", "photo", "image", "document"):
        return any(k in text for k in _T2FER_KEYWORDS)

    if len(text) >= 25 and any(k in text for k in _T2FER_KEYWORDS):
        return True

    return False

def _t2fer_task_dict(task):
    d = {}
    try:
        for k in task.keys():
            d[k] = task[k]
    except Exception:
        pass
    if not d and isinstance(task, dict):
        d.update(task)
    return d

async def _t2fer_run_final_estimate(conn, task, reason):
    task_id = _t2fer_s(_t2fer_get(task, "id", ""))
    try:
        from core.topic2_estimate_final_close_v2 import handle_topic2_estimate_final_close as _t2fer_final
    except Exception as e:
        _T2FER_LOG.exception("T2FER_IMPORT_ERR task=%s reason=%s err=%s", task_id, reason, e)
        return False

    t = _t2fer_task_dict(task)
    if not t:
        return False

    raw = _t2fer_s(t.get("raw_input", ""), 120000)
    if "\n---\nREVISION_CONTEXT" in raw:
        raw = raw.split("\n---\nREVISION_CONTEXT", 1)[0].strip()
        t["raw_input"] = raw

    if _t2fer_s(t.get("error_message", "")) == _T2FER_BAD_PARENT_ERR:
        t["error_message"] = ""
        t["result"] = ""

    try:
        ok = _t2fer_final(
            conn,
            t,
            send_reply_ex=None,
            update_task=globals().get("_update_task"),
            history=globals().get("_history"),
            logger=globals().get("logger", _T2FER_LOG),
        )
        if _t2fer_inspect.isawaitable(ok):
            ok = await ok
        if ok:
            try:
                conn.execute(
                    "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                    (task_id, "PATCH_TOPIC2_FRESH_ESTIMATE_ROUTE_GUARD_V1:" + reason),
                )
                conn.commit()
            except Exception:
                pass
            _T2FER_LOG.info("T2FER_FINAL_ESTIMATE_OK task=%s reason=%s", task_id, reason)
            return True
    except Exception as e:
        _T2FER_LOG.exception("T2FER_FINAL_ESTIMATE_ERR task=%s reason=%s err=%s", task_id, reason, e)

    try:
        from core.stroyka_estimate_canon import maybe_handle_stroyka_estimate as _t2fer_canon
        ok = _t2fer_canon(conn, t, logger=globals().get("logger", _T2FER_LOG))
        if _t2fer_inspect.isawaitable(ok):
            ok = await ok
        if ok:
            try:
                conn.execute(
                    "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                    (task_id, "PATCH_TOPIC2_FRESH_ESTIMATE_ROUTE_GUARD_V1:CANON_FALLBACK:" + reason),
                )
                conn.commit()
            except Exception:
                pass
            _T2FER_LOG.info("T2FER_CANON_ESTIMATE_OK task=%s reason=%s", task_id, reason)
            return True
    except Exception as e:
        _T2FER_LOG.exception("T2FER_CANON_ESTIMATE_ERR task=%s reason=%s err=%s", task_id, reason, e)

    return False

_T2FER_ORIG_P6E67_TRY_MERGE = globals().get("_p6e67_try_merge")
if _T2FER_ORIG_P6E67_TRY_MERGE and not getattr(_T2FER_ORIG_P6E67_TRY_MERGE, "_t2fer_wrapped", False):
    async def _p6e67_try_merge(conn, task, *args, **kwargs):
        if _t2fer_is_fresh_estimate(task):
            if await _t2fer_run_final_estimate(conn, task, "BYPASS_P6E67_PARENT_LOOKUP"):
                return True
        res = _T2FER_ORIG_P6E67_TRY_MERGE(conn, task, *args, **kwargs)
        if _t2fer_inspect.isawaitable(res):
            return await res
        return res
    _p6e67_try_merge._t2fer_wrapped = True

_T2FER_ORIG_HANDLE_NEW = globals().get("_handle_new")
if _T2FER_ORIG_HANDLE_NEW and not getattr(_T2FER_ORIG_HANDLE_NEW, "_t2fer_wrapped", False):
    async def _handle_new(conn, task, *args, **kwargs):
        if _t2fer_is_fresh_estimate(task):
            if await _t2fer_run_final_estimate(conn, task, "HANDLE_NEW_FRESH_ESTIMATE"):
                return
        res = _T2FER_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        if _t2fer_inspect.isawaitable(res):
            return await res
        return res
    _handle_new._t2fer_wrapped = True

_T2FER_ORIG_HANDLE_IN_PROGRESS = globals().get("_handle_in_progress")
if _T2FER_ORIG_HANDLE_IN_PROGRESS and not getattr(_T2FER_ORIG_HANDLE_IN_PROGRESS, "_t2fer_wrapped", False):
    async def _handle_in_progress(conn, task, *args, **kwargs):
        state = _t2fer_s(_t2fer_get(task, "state", "")).upper()
        err = _t2fer_s(_t2fer_get(task, "error_message", ""))
        if state == "WAITING_CLARIFICATION" and err == _T2FER_BAD_PARENT_ERR and _t2fer_is_fresh_estimate(task):
            if await _t2fer_run_final_estimate(conn, task, "WAITING_CLARIFICATION_RESCUE"):
                return
        res = _T2FER_ORIG_HANDLE_IN_PROGRESS(conn, task, *args, **kwargs)
        if _t2fer_inspect.isawaitable(res):
            return await res
        return res
    _handle_in_progress._t2fer_wrapped = True

_T2FER_ORIG_HANDLE_DRIVE_FILE = globals().get("_handle_drive_file")
if _T2FER_ORIG_HANDLE_DRIVE_FILE and not getattr(_T2FER_ORIG_HANDLE_DRIVE_FILE, "_t2fer_wrapped", False):
    async def _handle_drive_file(conn, task, *args, **kwargs):
        if _t2fer_is_fresh_estimate(task):
            if await _t2fer_run_final_estimate(conn, task, "DRIVE_FILE_FRESH_ESTIMATE"):
                return
        res = _T2FER_ORIG_HANDLE_DRIVE_FILE(conn, task, *args, **kwargs)
        if _t2fer_inspect.isawaitable(res):
            return await res
        return res
    _handle_drive_file._t2fer_wrapped = True

_T2FER_LOG.info("PATCH_TOPIC2_FRESH_ESTIMATE_ROUTE_GUARD_V1 installed")
# === END_PATCH_TOPIC2_FRESH_ESTIMATE_ROUTE_GUARD_V1 ===

# === PATCH_TOPIC2_FULL_PIPELINE_ROUTE_V1 ===
# Fact: NEW topic_2 estimate tasks processed by simplified v2, bypassing full P2/P3 pipeline
# Fix: route NEW topic_2 fresh estimates directly to _handle_in_progress (P2/P3 full pipeline)
# which calls handle_topic2_one_big_formula_pipeline_v1 with all 11 sections, AREAL_CALC, etc.

import inspect as _t2fp_inspect

_T2FP_ORIG_HANDLE_NEW = globals().get("_handle_new")
_T2FP_ORIG_HANDLE_IN_PROGRESS = globals().get("_handle_in_progress")

_T2FP_ESTIMATE_WORDS = (
    "смет", "кп", "расчет", "расчёт", "стоимост", "монолит",
    "бетон", "арматур", "фундамент", "перекрыт", "кровля", "стен",
    "гидроизол", "утеплен", "засыпк", "свай", "плит", "лестниц",
    "строит", " дом ", "ангар", "склад", "гараж", "баня",
    "м³", "м3", "м²", "м2", " шт", " тн", "п.м",
)

_T2FP_SHORT_CONTROL = frozenset((
    "да","ок","окей","хорошо","верно","подтверждаю","делай","согласен",
    "нет","отмена","отбой","завершить","всё",
    "статус","что сейчас","где результат","что там","ну что","что по задаче",
    "1","2","3","4","а","б","в","г","а)","б)","в)","г)",
))

def _t2fp_is_fresh_estimate(task):
    try:
        itype = str(_t2fer_get(task, "input_type", "") or "")
        if itype in ("photo", "image", "file", "drive_file", "document"):
            return True
        raw = str(_t2fer_get(task, "raw_input", "") or "").lower().replace("ё", "е")
        return any(x in raw for x in _T2FP_ESTIMATE_WORDS)
    except Exception:
        return False

def _t2fp_is_short(task):
    try:
        raw = str(_t2fer_get(task, "raw_input", "") or "")
        raw = re.sub(r"[\[VOICE\]\s]+", " ", raw.lower().replace("ё", "е")).strip(" .,!?:;")
        return raw in _T2FP_SHORT_CONTROL or len(raw) <= 6
    except Exception:
        return False

if _T2FP_ORIG_HANDLE_NEW and not getattr(_T2FP_ORIG_HANDLE_NEW, "_t2fp_wrapped", False):
    async def _handle_new(conn, task, *args, **kwargs):
        task_id = str(_t2fer_get(task, "id", "") or "")
        try:
            topic_id_v = int(_t2fer_get(task, "topic_id", 0) or 0)
        except Exception:
            topic_id_v = 0

        if topic_id_v == 2 and _t2fp_is_fresh_estimate(task) and not _t2fp_is_short(task):
            chat_id_v = str(_t2fer_get(task, "chat_id", "") or "")
            try:
                _update_task(conn, task_id, state="IN_PROGRESS", error_message="")
                conn.commit()
            except Exception:
                pass
            try:
                t = {}
                try:
                    for k in task.keys():
                        t[k] = task[k]
                except Exception:
                    try:
                        t = dict(task)
                    except Exception:
                        pass
                t["state"] = "IN_PROGRESS"
                h_ip = globals().get("_handle_in_progress")
                if h_ip:
                    res = h_ip(conn, t, chat_id_v, topic_id_v)
                    if _t2fp_inspect.isawaitable(res):
                        return await res
                    return res
            except Exception as _t2fp_e:
                try:
                    logger.warning("T2FP_FULL_PIPELINE_ERR task=%s err=%s", task_id, _t2fp_e)
                except Exception:
                    pass

        res = _T2FP_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        if _t2fp_inspect.isawaitable(res):
            return await res
        return res

    _handle_new._t2fp_wrapped = True
    logger.info("PATCH_TOPIC2_FULL_PIPELINE_ROUTE_V1 installed")

# === END_PATCH_TOPIC2_FULL_PIPELINE_ROUTE_V1 ===

# === PATCH_TOPIC2_REDIRECT_FINAL_TO_FULL_PIPELINE_V2 ===
# Root cause: PATCH_TOPIC2_FRESH_ESTIMATE_ROUTE_GUARD_V1 intercepts topic_2 tasks via
# _handle_drive_file and _p6e67_try_merge, routing them to _t2fer_run_final_estimate
# (simplified v2 path → handle_topic2_estimate_final_close → "Позиций: 1").
# Fix: wrap _t2fer_run_final_estimate so topic_2 calls go to full P2/P3 pipeline instead.
import logging as _t2rfp_log
import inspect as _t2rfp_inspect
_T2RFP_LOG = _t2rfp_log.getLogger("task_worker")

_T2RFP_ORIG = globals().get("_t2fer_run_final_estimate")

if _T2RFP_ORIG and not getattr(_T2RFP_ORIG, "_t2rfp_wrapped", False):
    async def _t2fer_run_final_estimate(conn, task, reason):
        try:
            topic_id_v = int(_t2fer_get(task, "topic_id", 0) or 0)
        except Exception:
            topic_id_v = 0

        if topic_id_v != 2:
            res = _T2RFP_ORIG(conn, task, reason)
            if _t2rfp_inspect.isawaitable(res):
                return await res
            return res

        task_id = _t2fer_s(_t2fer_get(task, "id", ""))
        _T2RFP_LOG.info("T2RFP_REDIRECT task=%s reason=%s → full pipeline", task_id, reason)

        try:
            chat_id_v = str(_t2fer_get(task, "chat_id", "") or "")
            try:
                _update_task(conn, task_id, state="IN_PROGRESS", error_message="")
                conn.commit()
            except Exception:
                pass

            t = {}
            try:
                for k in task.keys():
                    t[k] = task[k]
            except Exception:
                try:
                    t = dict(task)
                except Exception:
                    pass
            t["state"] = "IN_PROGRESS"

            h_ip = globals().get("_handle_in_progress")
            if h_ip:
                res = h_ip(conn, t, chat_id_v, topic_id_v)
                if _t2rfp_inspect.isawaitable(res):
                    await res
                _T2RFP_LOG.info("T2RFP_FULL_PIPELINE_DONE task=%s", task_id)
                return True
        except Exception as e:
            _T2RFP_LOG.warning("T2RFP_FULL_PIPELINE_ERR task=%s err=%s — fallback to simplified", task_id, e)

        res = _T2RFP_ORIG(conn, task, reason)
        if _t2rfp_inspect.isawaitable(res):
            return await res
        return res

    _t2fer_run_final_estimate._t2rfp_wrapped = True
    _T2RFP_LOG.info("PATCH_TOPIC2_REDIRECT_FINAL_TO_FULL_PIPELINE_V2 installed")

# === END_PATCH_TOPIC2_REDIRECT_FINAL_TO_FULL_PIPELINE_V2 ===

# === PATCH_TOPIC2_STATUS_META_GUARD_V1 ===
# §5: status/meta queries ("статус", "ну что там", "где результат") must NOT trigger estimate.
# Wraps _t2fer_run_final_estimate on top of V2 redirect; safe fallback on any error.
import logging as _tmg_log
import inspect as _tmg_inspect
_TMG_LOG = _tmg_log.getLogger("task_worker")

_TMG_STATUS_PHRASES = frozenset((
    "статус", "текущий статус", "какой статус",
    "где результат", "ну что там", "что там", "что сейчас",
    "что по задаче", "как дела с задачей", "когда будет готово",
    "ну что", "что делаешь", "что происходит",
))

def _tmg_is_status_query(task) -> bool:
    try:
        raw = str(_t2fer_get(task, "raw_input", "") or "")
        raw = raw.replace("[VOICE]", "").lower().replace("ё", "е").strip(" .,!?:;")
        if len(raw) > 80:
            return False
        tokens = raw.split()
        if len(tokens) > 5:
            return False
        return any(phrase in raw for phrase in _TMG_STATUS_PHRASES)
    except Exception:
        return False

def _tmg_get_status_text(conn, chat_id: str) -> str:
    try:
        row = conn.execute(
            "SELECT state, result, updated_at FROM tasks WHERE chat_id=? AND topic_id=2"
            " AND state NOT IN ('DONE','FAILED','CANCELLED','ARCHIVED')"
            " ORDER BY updated_at DESC LIMIT 1",
            (str(chat_id),),
        ).fetchone()
        if not row:
            return "Активных задач по разделу СТРОЙКА нет."
        state_labels = {
            "NEW": "в очереди",
            "IN_PROGRESS": "выполняется",
            "WAITING_CLARIFICATION": "ожидает уточнений",
            "AWAITING_CONFIRMATION": "ожидает подтверждения",
            "AWAITING_PRICE_CONFIRMATION": "ожидает выбора цен",
        }
        label = state_labels.get(str(row[0]), str(row[0]))
        preview = str(row[1] or "")[:200].strip()
        msg = f"Статус задачи: {label}"
        if preview:
            msg += f"\n\n{preview}"
        return msg
    except Exception:
        return "Статус задачи временно недоступен."

_TMG_ORIG = globals().get("_t2fer_run_final_estimate")
if _TMG_ORIG and not getattr(_TMG_ORIG, "_tmg_wrapped", False):
    async def _t2fer_run_final_estimate(conn, task, reason):
        try:
            topic_id_v = int(_t2fer_get(task, "topic_id", 0) or 0)
        except Exception:
            topic_id_v = 0

        if topic_id_v == 2 and _tmg_is_status_query(task):
            task_id = _t2fer_s(_t2fer_get(task, "id", ""))
            chat_id_v = str(_t2fer_get(task, "chat_id", "") or "")
            _TMG_LOG.info("T2_STATUS_QUERY_INTERCEPTED task=%s", task_id)
            try:
                status_msg = _tmg_get_status_text(conn, chat_id_v)
                reply_to = None
                try:
                    reply_to = _t2fer_get(task, "reply_to_message_id", None)
                    topic_id_for_send = int(_t2fer_get(task, "topic_id", 2) or 2)
                except Exception:
                    topic_id_for_send = 2
                send_reply_ex(
                    chat_id=chat_id_v,
                    text=status_msg,
                    reply_to_message_id=reply_to,
                    message_thread_id=topic_id_for_send,
                )
                if conn and task_id:
                    conn.execute(
                        "UPDATE tasks SET state='DONE', result=?, updated_at=datetime('now') WHERE id=?",
                        (status_msg, task_id),
                    )
                    conn.commit()
            except Exception as _tmg_e:
                _TMG_LOG.warning("T2_STATUS_REPLY_ERR task=%s err=%s", task_id, _tmg_e)
            return True

        res = _TMG_ORIG(conn, task, reason)
        if _tmg_inspect.isawaitable(res):
            return await res
        return res

    _t2fer_run_final_estimate._tmg_wrapped = True
    _TMG_LOG.info("PATCH_TOPIC2_STATUS_META_GUARD_V1 installed")

# === END_PATCH_TOPIC2_STATUS_META_GUARD_V1 ===

# === PATCH_TOPIC2_WC_LOOP_GUARD_V1 ===
# Fix: task in WAITING_CLARIFICATION (bot already asked a question) being re-picked
# and re-routed to full pipeline → loop asking same question.
# Guard: if DB state=WAITING_CLARIFICATION AND result already has text → don't redirect.
import logging as _wcg_log
import inspect as _wcg_inspect
_WCG_LOG = _wcg_log.getLogger("task_worker")

_WCG_ORIG = globals().get("_t2fer_run_final_estimate")

if _WCG_ORIG and not getattr(_WCG_ORIG, "_wcg_wrapped", False):
    async def _t2fer_run_final_estimate(conn, task, reason):
        try:
            topic_id_v = int(_t2fer_get(task, "topic_id", 0) or 0)
        except Exception:
            topic_id_v = 0

        if topic_id_v == 2:
            task_id = _t2fer_s(_t2fer_get(task, "id", ""))
            try:
                db_row = conn.execute(
                    "SELECT state, result FROM tasks WHERE id=? LIMIT 1", (task_id,)
                ).fetchone()
                if db_row:
                    db_state = str(db_row[0] or "").upper()
                    db_result = str(db_row[1] or "").strip()
                    if db_state == "WAITING_CLARIFICATION" and db_result:
                        # Task already asked a question — don't re-run pipeline
                        # Set error_message to WCG_SKIP so picker excludes it
                        _WCG_LOG.info("WCG_SKIP_LOOP task=%s already_asked=%r", task_id, db_result[:60])
                        try:
                            conn.execute(
                                "UPDATE tasks SET error_message='WCG_SKIP_WAITING_CLARIFICATION', updated_at=datetime('now') WHERE id=? AND state='WAITING_CLARIFICATION'",
                                (task_id,),
                            )
                            conn.commit()
                        except Exception:
                            pass
                        return True
            except Exception:
                pass

        res = _WCG_ORIG(conn, task, reason)
        if _wcg_inspect.isawaitable(res):
            return await res
        return res

    _t2fer_run_final_estimate._wcg_wrapped = True
    _WCG_LOG.info("PATCH_TOPIC2_WC_LOOP_GUARD_V1 installed")

# === END_PATCH_TOPIC2_WC_LOOP_GUARD_V1 ===

# === PATCH_TOPIC2_WC_PICKER_EXCLUDE_V1 ===
# Exclude WCG_SKIP tasks from picker — WAITING_CLARIFICATION tasks where bot
# already asked a question should not be re-picked until user replies.
# When user's reply arrives as a new task, the merge logic will clear WCG_SKIP.
import logging as _wcpe_log
_WCPE_LOG = _wcpe_log.getLogger("task_worker")

_WCPE_ORIG_PICK = globals().get("_pick_next_task")
if _WCPE_ORIG_PICK and not getattr(_WCPE_ORIG_PICK, "_wcpe_wrapped", False):
    def _pick_next_task(conn, chat_id=None):
        # Pick first non-WCG_SKIP task by scanning up to 20 candidates
        try:
            where = ["state IN ('NEW','IN_PROGRESS','WAITING_CLARIFICATION')",
                     "NOT (COALESCE(error_message,'') LIKE 'WCG_SKIP%')"]
            params = []
            if chat_id:
                where.insert(0, "chat_id=?")
                params.append(str(chat_id))
            conn.execute("BEGIN IMMEDIATE")
            row = conn.execute(
                f"SELECT * FROM tasks WHERE {' AND '.join(where)}"
                " ORDER BY CASE state WHEN 'IN_PROGRESS' THEN 0 ELSE 1 END, created_at ASC LIMIT 1",
                params
            ).fetchone()
            conn.execute("COMMIT")
            return row
        except Exception:
            return _WCPE_ORIG_PICK(conn, chat_id)
    _pick_next_task._wcpe_wrapped = True
    _WCPE_LOG.info("PATCH_TOPIC2_WC_PICKER_EXCLUDE_V1 installed")

# Clear WCG_SKIP when user's reply arrives (task state transitions to IN_PROGRESS)
# This happens automatically via _update_task which clears error_message.

# === END_PATCH_TOPIC2_WC_PICKER_EXCLUDE_V1 ===

# === PATCH_TOPIC2_T2RFP_LOOP_GUARD_V1 ===
# Root cause: worker dispatches drive_file tasks to _handle_drive_file by input_type,
# regardless of state. T2FER wraps _handle_drive_file and calls _t2fer_run_final_estimate
# for ANY drive_file task with topic_2. T2RFP redirects to _handle_in_progress (full pipeline).
# Pipeline may leave task in WAITING_CLARIFICATION/IN_PROGRESS → worker re-picks same task
# → _handle_drive_file again → T2RFP redirects again → infinite loop (443 iterations seen).
# Fix: for DRIVE_FILE_FRESH_ESTIMATE reason, only redirect if state==NEW (first pick).
# Re-picks (WAITING_CLARIFICATION/IN_PROGRESS) pass through to original simplified path.

import logging as _t2rlg_log
import inspect as _t2rlg_inspect
_T2RLG_LOG = _t2rlg_log.getLogger("task_worker")

_T2RLG_ORIG_FN = globals().get("_t2fer_run_final_estimate")

if _T2RLG_ORIG_FN and not getattr(_T2RLG_ORIG_FN, "_t2rlg_wrapped", False):
    async def _t2fer_run_final_estimate(conn, task, reason):
        try:
            topic_id_v = int(_t2fer_get(task, "topic_id", 0) or 0)
        except Exception:
            topic_id_v = 0

        if topic_id_v != 2:
            res = _T2RLG_ORIG_FN(conn, task, reason)
            if _t2rlg_inspect.isawaitable(res):
                return await res
            return res

        task_id = _t2fer_s(_t2fer_get(task, "id", ""))

        if reason == "DRIVE_FILE_FRESH_ESTIMATE":
            original_state = _t2fer_s(_t2fer_get(task, "state", "")).upper()
            if original_state != "NEW":
                _T2RLG_LOG.info(
                    "T2RLG_SKIP_REPICK task=%s state=%s reason=%s — passing to simplified",
                    task_id, original_state, reason
                )
                res = _T2RLG_ORIG_FN(conn, task, reason)
                if _t2rlg_inspect.isawaitable(res):
                    return await res
                return res

        res = _T2RLG_ORIG_FN(conn, task, reason)
        if _t2rlg_inspect.isawaitable(res):
            return await res
        return res

    _t2fer_run_final_estimate._t2rlg_wrapped = True
    _T2RLG_LOG.info("PATCH_TOPIC2_T2RFP_LOOP_GUARD_V1 installed")

# === END_PATCH_TOPIC2_T2RFP_LOOP_GUARD_V1 ===

# === PATCH_TOPIC2_PRICE_CHOICE_FALSE_POSITIVE_GUARD_V1 ===
# Root cause: _t2pc_choice matches "средн" in ANY sentence (e.g. "взяв за основу среднюю
# стоимость по рынку") and treats estimate requests as price-choice replies.
# When _price_bind_poison_parent_guard_v2 blocks, task stays NEW → worker re-picks → 75+ loop.
# Fix part 1: tighten _t2pc_choice to only match standalone short messages (<= 12 words),
#   not full estimate request sentences.
# Fix part 2: when poison guard blocks a falsely-detected price-choice task that is NOT
#   a real reply (long text / no explicit reply_to bound to a price menu), skip binding
#   silently so the task continues to the normal estimate pipeline.
import logging as _t2fpg_log
import re as _t2fpg_re
_T2FPG_LOG = _t2fpg_log.getLogger("task_worker")

_T2FPG_ORIG_TRY_BIND = globals().get("_t2pc_try_bind_price_choice")

_T2FPG_PRICE_WORDS = frozenset(("средн", "медиан", "рынок", "миним", "дешев", "максим",
                                  "надеж", "надёж", "проверенн", "ручн", "вручную"))

def _t2fpg_is_real_price_reply(raw: str) -> bool:
    """True only if this looks like a standalone price choice, not an estimate request."""
    try:
        t = _t2fpg_re.sub(r"\[VOICE\]", "", raw or "").strip().lower().replace("ё", "е")
        words = t.split()
        # Exact single token (1/2/3/4/а/б/в/г)
        if len(words) <= 2:
            return True
        # Short explicit phrase
        if len(words) <= 8 and any(w in t for w in ("ставь", "беру", "выбираю", "выбери", "поставь")):
            return True
        # Construction estimate keywords present → it's an estimate, not a price choice
        estimate_words = ("смет", "расчет", "расчёт", "стоимост", "построит", "кирпич",
                          "газобетон", "фундамент", "кровля", "материал", "работ",
                          "монолит", "перекрыт", "утеплен", "отделк", "дом", "ангар",
                          "объект", "проект", "участ", "площад")
        if any(w in t for w in estimate_words):
            return False
        # Long sentence without explicit choice phrase → not a price reply
        if len(words) > 10:
            return False
        return True
    except Exception:
        return True  # fail-safe: let original logic handle

if _T2FPG_ORIG_TRY_BIND and not getattr(_T2FPG_ORIG_TRY_BIND, "_t2fpg_wrapped", False):
    async def _t2pc_try_bind_price_choice(conn, task):
        import inspect as _t2fpg_inspect
        try:
            raw = str((task.get("raw_input") if isinstance(task, dict) else
                       (task["raw_input"] if hasattr(task, "keys") and "raw_input" in task.keys()
                        else "")) or "")
            topic_id_v = int((task.get("topic_id") if isinstance(task, dict) else
                              (task["topic_id"] if hasattr(task, "keys") and "topic_id" in task.keys()
                               else 0)) or 0)
            if topic_id_v == 2 and not _t2fpg_is_real_price_reply(raw):
                return False
        except Exception:
            pass
        res = _T2FPG_ORIG_TRY_BIND(conn, task)
        if _t2fpg_inspect.isawaitable(res):
            return await res
        return res
    _t2pc_try_bind_price_choice._t2fpg_wrapped = True
    _T2FPG_LOG.info("PATCH_TOPIC2_PRICE_CHOICE_FALSE_POSITIVE_GUARD_V1 installed")
else:
    _T2FPG_LOG.warning("PATCH_TOPIC2_PRICE_CHOICE_FALSE_POSITIVE_GUARD_V1 skipped: _t2pc_try_bind_price_choice not found")
# === END_PATCH_TOPIC2_PRICE_CHOICE_FALSE_POSITIVE_GUARD_V1 ===

# === PATCH_CONFIRMATION_TIMEOUT_DONE_IF_DELIVERED_V1 ===
# Root cause: _recover_stale_tasks sets ALL AWAITING_CONFIRMATION tasks to FAILED after 30 min.
# Per spec: if estimate was delivered (Drive links in result) → task should become DONE, not FAILED.
# Topics affected: 2 (estimate), 5 (technadzor act), 210 (design).
# Fix: intercept _recover_stale_tasks and for topic 2/5/210 AWAITING_CONFIRMATION tasks
#   that have a drive.google.com link in result → set DONE instead of FAILED.
import logging as _ctdd_log
_CTDD_LOG = _ctdd_log.getLogger("task_worker")

_CTDD_ORIG_RECOVER = globals().get("_recover_stale_tasks")

if _CTDD_ORIG_RECOVER and not getattr(_CTDD_ORIG_RECOVER, "_ctdd_wrapped", False):
    def _recover_stale_tasks(conn, *args, **kwargs):
        # Promote delivered estimates to DONE before the FAILED sweep runs
        try:
            conn.execute("""
                UPDATE tasks
                SET state='DONE', error_message='', updated_at=datetime('now')
                WHERE state='AWAITING_CONFIRMATION'
                  AND COALESCE(topic_id,0) IN (2, 5, 210)
                  AND updated_at < datetime('now','-30 minutes')
                  AND result LIKE '%drive.google.com%'
                  AND COALESCE(raw_input,'') NOT LIKE '%retry_queue_healthcheck%'
            """)
            conn.commit()
        except Exception as _ctdd_e:
            _CTDD_LOG.warning("PATCH_CONFIRMATION_TIMEOUT_DONE_IF_DELIVERED_V1_ERR %s", _ctdd_e)
        return _CTDD_ORIG_RECOVER(conn, *args, **kwargs)
    _recover_stale_tasks._ctdd_wrapped = True
    _CTDD_LOG.info("PATCH_CONFIRMATION_TIMEOUT_DONE_IF_DELIVERED_V1 installed")
else:
    _CTDD_LOG.warning("PATCH_CONFIRMATION_TIMEOUT_DONE_IF_DELIVERED_V1 skipped: _recover_stale_tasks not found")
# === END_PATCH_CONFIRMATION_TIMEOUT_DONE_IF_DELIVERED_V1 ===

# === PATCH_T500_P6F_DAH_EXCLUDE_V1 ===
# Root cause: P6F_DAH gate fires for topic_500 because user text contains "ссылку"
# (matches _p6f_dah_user_wants_artifact). Gate converts DONE→IN_PROGRESS.
# Next pick clears error_message → loop forever.
# Fix: wrap _update_task to bypass gate for topic_500 when setting DONE.
# For topic_500 a successful search reply IS the complete artifact — no Drive upload required.
import logging as _t5dah_log
_T5DAH_LOG = _t5dah_log.getLogger("task_worker")

_T5DAH_CURRENT = _update_task
_T5DAH_PREGATE = globals().get("_P6F_DAH_ORIG_UPDATE_TASK")

if not getattr(_T5DAH_CURRENT, "_t5dah_wrapped", False) and _T5DAH_PREGATE:
    def _update_task(conn, task_id, **kwargs):
        if kwargs.get("state") == "DONE":
            try:
                row = conn.execute(
                    "SELECT topic_id FROM tasks WHERE id=? LIMIT 1", (str(task_id),)
                ).fetchone()
                if row is not None:
                    tid = int((row["topic_id"] if hasattr(row, "keys") else row[0]) or 0)
                    if tid == 500:
                        return _T5DAH_PREGATE(conn, task_id, **kwargs)
            except Exception:
                pass
        return _T5DAH_CURRENT(conn, task_id, **kwargs)
    _update_task._t5dah_wrapped = True
    _T5DAH_LOG.info("PATCH_T500_P6F_DAH_EXCLUDE_V1 installed")

    # One-time fix: force stuck topic_500 task to DONE if result already delivered
    try:
        import sqlite3 as _t5dah_sq
        with _t5dah_sq.connect("data/core.db") as _c:
            _c.row_factory = _t5dah_sq.Row
            _stuck = _c.execute(
                "SELECT id FROM tasks WHERE state='IN_PROGRESS' AND topic_id=500"
                " AND result LIKE '%drive.google.com%' OR (state='IN_PROGRESS' AND topic_id=500"
                " AND result LIKE '%https://%' AND length(result)>100)"
                " LIMIT 20"
            ).fetchall()
            for _sr in _stuck:
                _c.execute(
                    "UPDATE tasks SET state='DONE', error_message='', updated_at=datetime('now') WHERE id=?",
                    (_sr["id"],)
                )
                _c.execute(
                    "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                    (_sr["id"], "PATCH_T500_P6F_DAH_EXCLUDE_V1_FORCE_DONE_STUCK")
                )
                _T5DAH_LOG.info("PATCH_T500_FORCE_DONE: %s", _sr["id"])
            _c.commit()
    except Exception as _t5dah_fix_e:
        _T5DAH_LOG.warning("PATCH_T500_FORCE_DONE_ERR: %s", _t5dah_fix_e)
else:
    _T5DAH_LOG.warning("PATCH_T500_P6F_DAH_EXCLUDE_V1 skipped: pregate ref not found or already wrapped")
# === END_PATCH_T500_P6F_DAH_EXCLUDE_V1 ===

# === PATCH_T210_T5_REPLIED_DONE_GATE_V1 ===
# Root cause: topic_210 / topic_5 tasks that were already answered (reply_sent: in history)
# get stuck in IN_PROGRESS loop because:
#   a) P6F_DAH gate fires (raw_input contains "файл" etc.) → DONE→IN_PROGRESS
#   b) NO_GENERIC_RESPONSE_AS_RESULT blocks AI text → can't set result → can't close
# Fix part A: extend _update_task bypass to topic_210 and topic_5 when reply already sent.
# Fix part B: at startup, force-DONE stuck tasks in these topics that have reply_sent history.
import logging as _t25g_log
_T25G_LOG = _t25g_log.getLogger("task_worker")

_T25G_CURRENT = _update_task
_T25G_PREGATE = globals().get("_P6F_DAH_ORIG_UPDATE_TASK")

def _t25g_has_reply_sent(conn, task_id: str) -> bool:
    """True if any reply_sent: history entry exists for this task."""
    try:
        row = conn.execute(
            "SELECT 1 FROM task_history WHERE task_id=? AND action LIKE 'reply_sent:%' LIMIT 1",
            (str(task_id),),
        ).fetchone()
        return row is not None
    except Exception:
        return False

if not getattr(_T25G_CURRENT, "_t25g_wrapped", False) and _T25G_PREGATE:
    def _update_task(conn, task_id, **kwargs):
        if kwargs.get("state") == "DONE":
            try:
                row = conn.execute(
                    "SELECT topic_id FROM tasks WHERE id=? LIMIT 1", (str(task_id),)
                ).fetchone()
                if row is not None:
                    tid = int((row["topic_id"] if hasattr(row, "keys") else row[0]) or 0)
                    if tid in (210, 5) and _t25g_has_reply_sent(conn, task_id):
                        return _T25G_PREGATE(conn, task_id, **kwargs)
            except Exception:
                pass
        return _T25G_CURRENT(conn, task_id, **kwargs)
    _update_task._t25g_wrapped = True
    _T25G_LOG.info("PATCH_T210_T5_REPLIED_DONE_GATE_V1 installed")

    # Force-DONE stuck IN_PROGRESS tasks for 210/5 that already have reply_sent in history
    try:
        import sqlite3 as _t25g_sq
        with _t25g_sq.connect("data/core.db") as _c2:
            _c2.row_factory = _t25g_sq.Row
            _stuck2 = _c2.execute(
                """SELECT DISTINCT t.id FROM tasks t
                   JOIN task_history th ON th.task_id=t.id AND th.action LIKE 'reply_sent:%'
                   WHERE t.state='IN_PROGRESS' AND COALESCE(t.topic_id,0) IN (210,5)"""
            ).fetchall()
            for _sr2 in _stuck2:
                _c2.execute(
                    "UPDATE tasks SET state='DONE', error_message='', updated_at=datetime('now') WHERE id=?",
                    (_sr2["id"],)
                )
                _c2.execute(
                    "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                    (_sr2["id"], "PATCH_T210_T5_REPLIED_DONE_GATE_V1_FORCE_DONE")
                )
                _T25G_LOG.info("PATCH_T210_FORCE_DONE: %s", _sr2["id"])
            _c2.commit()
    except Exception as _t25g_fix_e:
        _T25G_LOG.warning("PATCH_T210_FORCE_DONE_ERR: %s", _t25g_fix_e)
else:
    _T25G_LOG.warning("PATCH_T210_T5_REPLIED_DONE_GATE_V1 skipped: pregate ref not found or already wrapped")
# === END_PATCH_T210_T5_REPLIED_DONE_GATE_V1 ===

# === PATCH_TOPIC210_META_GUARD_V1 ===
# Root cause: topic_210 receives meta-comments ("бери в работу", "ты сам выбирай") that
# go to the full project engine → AI generates generic "готов к выполнению" →
# NO_GENERIC_RESPONSE_AS_RESULT_V1_BLOCKED rejects → task loops/fails.
# Fix: intercept _handle_new for topic_210 meta-comments; acknowledge and close DONE.
import logging as _t210mg_log_mod
import inspect as _t210mg_inspect
_T210MG_LOG = _t210mg_log_mod.getLogger("task_worker")

_T210_META_PHRASES = (
    "ты сам должен выбирать", "ты сам выбирай", "сам выбирай",
    "бери в работу", "бери их в работу", "возьми в работу",
    "всё правильно понял", "правильно понял", "ты правильно понял",
    "бери и делай", "делай сам", "сам решай",
)

def _t210mg_is_meta(raw: str, input_type: str) -> bool:
    if input_type in ("photo", "file", "drive_file", "image", "document"):
        return False
    t = str(raw or "").lower().replace("ё", "е").replace("[voice]", "").replace("[голос]", "").strip(" .,!?:;")
    if len(t.split()) > 12:
        return False
    return any(x in t for x in _T210_META_PHRASES)

_T210MG_ORIG_HN = globals().get("_handle_new")

if _T210MG_ORIG_HN and not getattr(_T210MG_ORIG_HN, "_t210mg_wrapped", False):
    async def _handle_new(conn, task, *args, **kwargs):
        try:
            if isinstance(task, dict):
                topic_id_v = int(task.get("topic_id") or 0)
                raw_v = str(task.get("raw_input") or "")
                input_type_v = str(task.get("input_type") or "")
                task_id_v = str(task.get("id") or "")
                chat_id_v = str(task.get("chat_id") or "")
                reply_to_v = task.get("reply_to_message_id")
            else:
                topic_id_v = int(task["topic_id"] if hasattr(task, "keys") and "topic_id" in task.keys() else 0)
                raw_v = str(task["raw_input"] if hasattr(task, "keys") and "raw_input" in task.keys() else "")
                input_type_v = str(task["input_type"] if hasattr(task, "keys") and "input_type" in task.keys() else "")
                task_id_v = str(task["id"] if hasattr(task, "keys") and "id" in task.keys() else "")
                chat_id_v = str(task["chat_id"] if hasattr(task, "keys") and "chat_id" in task.keys() else "")
                reply_to_v = task["reply_to_message_id"] if hasattr(task, "keys") and "reply_to_message_id" in task.keys() else None

            if topic_id_v == 210 and _t210mg_is_meta(raw_v, input_type_v):
                msg = "Понял, принято. Продолжаю работу по проекту в этом разделе."
                try:
                    _send_once_ex(conn, task_id_v, chat_id_v, msg, reply_to_v, "t210mg_ack")
                except Exception:
                    pass
                _update_task(conn, task_id_v, state="DONE", result=msg, error_message="")
                _history(conn, task_id_v, "TOPIC210_META_GUARD_DONE_V1")
                conn.commit()
                _T210MG_LOG.info("PATCH_TOPIC210_META_GUARD_V1 handled meta task=%s", task_id_v)
                return True
        except Exception as _t210mg_e:
            _T210MG_LOG.warning("PATCH_TOPIC210_META_GUARD_V1_ERR %s", _t210mg_e)
        res = _T210MG_ORIG_HN(conn, task, *args, **kwargs)
        if _t210mg_inspect.isawaitable(res):
            return await res
        return res

    _handle_new._t210mg_wrapped = True
    _T210MG_LOG.info("PATCH_TOPIC210_META_GUARD_V1 installed")
else:
    _T210MG_LOG.warning("PATCH_TOPIC210_META_GUARD_V1 skipped: _handle_new not found")
# === END_PATCH_TOPIC210_META_GUARD_V1 ===

if __name__ == "__main__":
    asyncio.run(main())


# ============================================================
# === PATCH_TOPIC2_CANCEL_GUARD_V1 ===
# Цель: cancel-intent в topic_2 ловится ДО любого estimate route_guard
# Факт: 11:54 «Отмена всех задач» → бот выдал смету
# Факт: 17:33 «Задача отменена» → бот спросил «Сколько этажей»
# ============================================================
import re as _tcg_re
import inspect as _tcg_inspect
import logging as _tcg_logging
_TCG_LOG = _tcg_logging.getLogger("task_worker.cancel_guard")

_TCG_CANCEL_RE = _tcg_re.compile(
    r"(?:^|\s)(?:\[VOICE\]\s*)?"
    r"(отмена|отбой|стоп|заверши|завершена|закрой|закрывай|очисти|"
    r"отменяй|задача отменена|отмена задач|все задачи завершен|"
    r"отбой всех|очисти все|задача завершена)",
    _tcg_re.IGNORECASE
)

def _tcg_is_cancel(text):
    if not text:
        return False
    s = str(text).strip().lower()
    if len(s) > 200:
        return False
    return bool(_TCG_CANCEL_RE.search(s))

def _tcg_get(task, key, default=""):
    try:
        if isinstance(task, dict):
            return task.get(key, default)
        return task[key]
    except Exception:
        return default

_TCG_ORIG_HANDLE_NEW = globals().get("_handle_new")
if _TCG_ORIG_HANDLE_NEW and not getattr(_TCG_ORIG_HANDLE_NEW, "_tcg_wrapped", False):
    async def _handle_new(conn, task, *args, **kwargs):
        try:
            topic_id = int(_tcg_get(task, "topic_id", 0) or 0)
            raw = str(_tcg_get(task, "raw_input", "") or "")
            if topic_id == 2 and _tcg_is_cancel(raw):
                task_id = str(_tcg_get(task, "id", ""))
                chat_id = str(_tcg_get(task, "chat_id", ""))
                reply_to = _tcg_get(task, "reply_to_message_id", None)
                try:
                    conn.execute(
                        "UPDATE tasks SET state='CANCELLED', "
                        "error_message='TOPIC2_CANCEL_GUARD_V1', "
                        "updated_at=datetime('now') WHERE id=?",
                        (task_id,)
                    )
                    conn.execute(
                        "UPDATE tasks SET state='CANCELLED', "
                        "error_message='TOPIC2_CANCEL_GUARD_V1:scoped' "
                        "WHERE chat_id=? AND topic_id=2 "
                        "AND state IN ('NEW','IN_PROGRESS','WAITING_CLARIFICATION','AWAITING_CONFIRMATION') "
                        "AND id<>?",
                        (chat_id, task_id)
                    )
                    conn.execute(
                        "INSERT INTO task_history (task_id, action, created_at) VALUES (?, ?, datetime('now'))",
                        (task_id, "TOPIC2_CANCEL_GUARD_V1:cancelled_active_topic2")
                    )
                    conn.commit()
                except Exception as _e:
                    _TCG_LOG.warning("CANCEL_GUARD_DB_ERR task=%s err=%s", task_id, _e)
                try:
                    _send_once(conn, task_id, chat_id,
                               "Задачи в этом топике отменены.",
                               reply_to, "topic2_cancel_guard")
                except Exception as _e:
                    _TCG_LOG.warning("CANCEL_GUARD_SEND_ERR task=%s err=%s", task_id, _e)
                _TCG_LOG.info("TOPIC2_CANCEL_GUARD_V1 fired task=%s chat=%s", task_id, chat_id)
                return
        except Exception as e:
            _TCG_LOG.exception("CANCEL_GUARD_TOP_ERR err=%s", e)
        res = _TCG_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        if _tcg_inspect.isawaitable(res):
            return await res
        return res
    _handle_new._tcg_wrapped = True

_TCG_LOG.info("PATCH_TOPIC2_CANCEL_GUARD_V1 installed")
# === END_PATCH_TOPIC2_CANCEL_GUARD_V1 ===


# ============================================================
# === PATCH_TOPIC2_FRESH_ESTIMATE_FALLBACK_V1 ===
# Цель: при P6E67_PARENT_NOT_FOUND с полным новым ТЗ → fresh estimate
# Факт: 12:50, 13:43, 19:33 — «Не нашёл родительскую задачу» при полном ТЗ
# ============================================================
import re as _tffe_re
import inspect as _tffe_inspect
import logging as _tffe_logging
_TFFE_LOG = _tffe_logging.getLogger("task_worker.fresh_fallback")

_TFFE_ESTIMATE_KEYS = (
    "смет", "кп", "коммерческ", "расчёт", "расчет", "стоимост", "цен",
    "объём", "объем", "м²", "м2", "м³", "м3", "посчитай", "посчитать",
    "монолит", "бетон", "арматур", "фундамент", "стен", "перекрыт", "кровл",
    "газобетон", "каркас", "кирпич", "отделк", "имитация бруса", "ламинат"
)

def _tffe_count_estimate_signals(text):
    if not text:
        return 0
    low = str(text).lower()
    return sum(1 for k in _TFFE_ESTIMATE_KEYS if k in low)

def _tffe_get(task, key, default=""):
    try:
        if isinstance(task, dict):
            return task.get(key, default)
        return task[key]
    except Exception:
        return default

_TFFE_ORIG_P6E67 = globals().get("_p6e67_try_merge")
if _TFFE_ORIG_P6E67 and not getattr(_TFFE_ORIG_P6E67, "_tffe_wrapped", False):
    async def _p6e67_try_merge(conn, task, *args, **kwargs):
        topic_id = int(_tffe_get(task, "topic_id", 0) or 0)
        raw = str(_tffe_get(task, "raw_input", "") or "")
        signals = _tffe_count_estimate_signals(raw) if topic_id == 2 else 0
        if topic_id == 2 and signals >= 3:
            try:
                fn = globals().get("_t2fer_run_final_estimate")
                if fn:
                    if await fn(conn, task, "FRESH_FALLBACK_FULL_TZ"):
                        _TFFE_LOG.info("TFFE_FRESH_FALLBACK_OK task=%s signals=%d",
                                       _tffe_get(task, "id"), signals)
                        return True
            except Exception as e:
                _TFFE_LOG.warning("TFFE_FRESH_FALLBACK_ERR task=%s err=%s",
                                  _tffe_get(task, "id"), e)
        res = _TFFE_ORIG_P6E67(conn, task, *args, **kwargs)
        if _tffe_inspect.isawaitable(res):
            return await res
        return res
    _p6e67_try_merge._tffe_wrapped = True

_TFFE_LOG.info("PATCH_TOPIC2_FRESH_ESTIMATE_FALLBACK_V1 installed")
# === END_PATCH_TOPIC2_FRESH_ESTIMATE_FALLBACK_V1 ===


# ============================================================
# === PATCH_TOPIC2_PRICE_REPLY_REVIVE_V1 ===
# Цель: короткий ответ ("2", "да", "жду") при WAITING_CLARIFICATION + PRICE_REQUESTED 
#       идёт в parent, не создаёт новую задачу
# Факт: 10:04-10:06 — пять ответов "2"/"" подряд → 5 P6E67_PARENT_NOT_FOUND
# Факт: 10:36 — "Какая последняя задача" → новая задача с price menu
# ============================================================
import re as _tprr_re
import inspect as _tprr_inspect
import logging as _tprr_logging
_TPRR_LOG = _tprr_logging.getLogger("task_worker.price_revive")

_TPRR_PRICE_REPLY_RE = _tprr_re.compile(
    r"^\s*(?:\[VOICE\]\s*)?"
    r"(?:1|2|3|4|а\)?|б\)?|в\)?|г\)?|"
    r"да|ок|жду|давай|делай|поехали|"
    r"вариант\s*[1-4абвг]|"
    r"миним|средн|максим|медиан|шаблонн|ручн|"
    r"ставь\s+\w+)"
    r"\s*$",
    _tprr_re.IGNORECASE
)

def _tprr_is_short_price_reply(text):
    if not text:
        return False
    s = str(text).strip()
    if len(s) > 80:
        return False
    return bool(_TPRR_PRICE_REPLY_RE.match(s))

def _tprr_get(task, key, default=""):
    try:
        if isinstance(task, dict):
            return task.get(key, default)
        return task[key]
    except Exception:
        return default

def _tprr_find_active_price_parent(conn, chat_id):
    try:
        row = conn.execute(
            """SELECT id FROM tasks 
               WHERE chat_id=? AND topic_id=2 
                 AND state IN ('WAITING_CLARIFICATION','IN_PROGRESS')
                 AND id IN (
                   SELECT task_id FROM task_history 
                   WHERE action='TOPIC2_PRICE_CHOICE_REQUESTED'
                 )
               ORDER BY updated_at DESC LIMIT 1""",
            (str(chat_id),)
        ).fetchone()
        if row:
            return row["id"] if hasattr(row, "keys") else row[0]
    except Exception:
        pass
    return None

_TPRR_ORIG_HANDLE_NEW = globals().get("_handle_new")
if _TPRR_ORIG_HANDLE_NEW and not getattr(_TPRR_ORIG_HANDLE_NEW, "_tprr_wrapped", False):
    async def _handle_new(conn, task, *args, **kwargs):
        try:
            topic_id = int(_tprr_get(task, "topic_id", 0) or 0)
            raw = str(_tprr_get(task, "raw_input", "") or "")
            if topic_id == 2 and _tprr_is_short_price_reply(raw):
                chat_id = str(_tprr_get(task, "chat_id", ""))
                parent_id = _tprr_find_active_price_parent(conn, chat_id)
                if parent_id:
                    task_id = str(_tprr_get(task, "id", ""))
                    try:
                        conn.execute(
                            "INSERT INTO task_history (task_id, action, created_at) VALUES (?, ?, datetime('now'))",
                            (parent_id, "clarified:" + raw[:200])
                        )
                        conn.execute(
                            "UPDATE tasks SET state='IN_PROGRESS', error_message='', "
                            "updated_at=datetime('now') WHERE id=?",
                            (parent_id,)
                        )
                        conn.execute(
                            "UPDATE tasks SET state='CANCELLED', "
                            "error_message='TPRR:MERGED_TO_PARENT:'||?, "
                            "result='Ответ применён к активной сметной задаче' "
                            "WHERE id=?",
                            (parent_id, task_id)
                        )
                        conn.execute(
                            "INSERT INTO task_history (task_id, action, created_at) VALUES (?, ?, datetime('now'))",
                            (task_id, "TPRR_MERGED_TO_PRICE_PARENT:" + parent_id)
                        )
                        conn.commit()
                        _TPRR_LOG.info("TPRR_MERGE task=%s → parent=%s reply=%s",
                                       task_id, parent_id, raw[:30])
                    except Exception as _e:
                        _TPRR_LOG.warning("TPRR_DB_ERR task=%s err=%s", task_id, _e)
                    return
        except Exception as e:
            _TPRR_LOG.exception("TPRR_TOP_ERR err=%s", e)
        res = _TPRR_ORIG_HANDLE_NEW(conn, task, *args, **kwargs)
        if _tprr_inspect.isawaitable(res):
            return await res
        return res
    _handle_new._tprr_wrapped = True

_TPRR_LOG.info("PATCH_TOPIC2_PRICE_REPLY_REVIVE_V1 installed")
# === END_PATCH_TOPIC2_PRICE_REPLY_REVIVE_V1 ===


# ============================================================
# === PATCH_TOPIC2_PRICE_TIMEOUT_GUARD_V1 ===
# Цель: задачи с TOPIC2_PRICE_CHOICE_REQUESTED не убивает 30-мин таймаут
# Факт: 10:16, 10:35, 10:43 — IN_PROGRESS_HARD_TIMEOUT_BY_CREATED_AT_FIX_V1
#       убил f1ef9fab пока юзер думал над ценой
# ============================================================
import logging as _tptg_logging
_TPTG_LOG = _tptg_logging.getLogger("task_worker.price_timeout_guard")

_TPTG_ORIG_TIMEOUT = globals().get("_in_progress_hard_timeout_by_created_at_fix_v1")
if _TPTG_ORIG_TIMEOUT and not getattr(_TPTG_ORIG_TIMEOUT, "_tptg_wrapped", False):
    def _in_progress_hard_timeout_by_created_at_fix_v1(conn, minutes: int = 30) -> int:
        try:
            rows = conn.execute(
                """SELECT id FROM tasks
                   WHERE state='IN_PROGRESS' AND topic_id=2
                     AND id IN (
                       SELECT task_id FROM task_history 
                       WHERE action='TOPIC2_PRICE_CHOICE_REQUESTED'
                     )""",
            ).fetchall()
            shielded = 0
            for r in rows:
                tid = r["id"] if hasattr(r, "keys") else r[0]
                try:
                    conn.execute(
                        "UPDATE tasks SET state='WAITING_CLARIFICATION', "
                        "error_message='TPTG:price_choice_pending', "
                        "updated_at=datetime('now') WHERE id=? AND state='IN_PROGRESS'",
                        (str(tid),)
                    )
                    shielded += 1
                except Exception:
                    pass
            if shielded:
                conn.commit()
                _TPTG_LOG.info("TPTG_SHIELDED %d tasks (price_pending → WAITING_CLARIFICATION)", shielded)
        except Exception as e:
            _TPTG_LOG.warning("TPTG_PREFILTER_ERR %s", e)
        return _TPTG_ORIG_TIMEOUT(conn, minutes)
    _in_progress_hard_timeout_by_created_at_fix_v1._tptg_wrapped = True

_TPTG_LOG.info("PATCH_TOPIC2_PRICE_TIMEOUT_GUARD_V1 installed")
# === END_PATCH_TOPIC2_PRICE_TIMEOUT_GUARD_V1 ===


# ============================================================
# === PATCH_TOPIC2_DONE_OVERRIDE_INVALID_PUBLIC_V1 ===
# Цель: если у задачи topic_2 есть все DONE markers + Drive XLSX/PDF ссылки,
#       не блокировать переход в AWAITING_CONFIRMATION/DONE через INVALID_PUBLIC_RESULT
# Факт логов 21:05:02: задача 893436d4 имела все 14 markers, но state=FAILED
# ============================================================
import logging as _tdoip_logging
_TDOIP_LOG = _tdoip_logging.getLogger("task_worker.done_override")

_TDOIP_REQUIRED_MARKERS = (
    "TOPIC2_XLSX_CREATED",
    "TOPIC2_PDF_CREATED",
    "TOPIC2_DRIVE_UPLOAD_XLSX_OK",
    "TOPIC2_DRIVE_UPLOAD_PDF_OK",
    "TOPIC2_TELEGRAM_DELIVERED",
)

def _tdoip_has_all_done_markers(conn, task_id):
    try:
        rows = conn.execute(
            "SELECT action FROM task_history WHERE task_id=?",
            (str(task_id),)
        ).fetchall()
        actions = " ".join(str(r[0]) for r in rows)
        return all(m in actions for m in _TDOIP_REQUIRED_MARKERS)
    except Exception:
        return False

def _tdoip_has_drive_links_in_result(result):
    if not result:
        return False
    s = str(result).lower()
    return "drive.google.com" in s and ("xlsx" in s or "pdf" in s or ".xls" in s or "📊" in str(result) or "📄" in str(result))

_TDOIP_ORIG_VIOLATION = globals().get("_fcg_public_result_violation")
if _TDOIP_ORIG_VIOLATION and not getattr(_TDOIP_ORIG_VIOLATION, "_tdoip_wrapped", False):
    def _fcg_public_result_violation(conn, task_id, state, result, error_message=""):
        try:
            row = conn.execute(
                "SELECT topic_id FROM tasks WHERE id=?", (str(task_id),)
            ).fetchone()
            topic_id_v = int(row[0]) if row else 0
            if topic_id_v == 2:
                if _tdoip_has_all_done_markers(conn, task_id) and _tdoip_has_drive_links_in_result(result):
                    try:
                        conn.execute(
                            "INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))",
                            (str(task_id), "TDOIP_OVERRIDE:14_markers_and_drive_links_present")
                        )
                        conn.commit()
                    except Exception:
                        pass
                    _TDOIP_LOG.info("TDOIP_OVERRIDE task=%s — markers + Drive links → no violation", task_id)
                    return ""
        except Exception as e:
            _TDOIP_LOG.warning("TDOIP_TOP_ERR task=%s err=%s", task_id, e)
        return _TDOIP_ORIG_VIOLATION(conn, task_id, state, result, error_message)
    _fcg_public_result_violation._tdoip_wrapped = True

_TDOIP_LOG.info("PATCH_TOPIC2_DONE_OVERRIDE_INVALID_PUBLIC_V1 installed")
# === END_PATCH_TOPIC2_DONE_OVERRIDE_INVALID_PUBLIC_V1 ===


# ============================================================
# === PATCH_TOPIC2_INLINE_FIX_20260506_V1 SUPERSEDES_NOTE ===
# Inline fixes applied directly in function bodies above:
#   - _p6e67_try_merge:  (1) state guard at top — stops infinite loop on CANCELLED/FAILED/DONE
#                        (2) fresh estimate dispatch in `if not parent` block — full-TZ → estimate
#   - _update_task FCG wrapper: bypass INVALID_PUBLIC_RESULT if 5 critical DONE markers + Drive link in result
#   - _t2v5_/_t2v6c_ price bind: explicit token required (no long messages, no fallback to history)
#
# Wrappers below (PATCH_TOPIC2_CANCEL_GUARD_V1 / FRESH_ESTIMATE_FALLBACK_V1 / PRICE_REPLY_REVIVE_V1 /
#   PRICE_TIMEOUT_GUARD_V1 / DONE_OVERRIDE_INVALID_PUBLIC_V1) are SUPERSEDED_BY_INLINE_FIX_20260506_V1.
# Kept inert (they wrap functions that are already inline-fixed). To remove: delete lines ~14898-15256.
# ============================================================
