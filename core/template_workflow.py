# === PROJECT_TEMPLATE_WORKFLOW_FULL_CLOSE_V1 ===
# === TECHNADZOR_ACT_TEMPLATE_WORKFLOW_V1 ===
# === TEMPLATE_SCOPE_ENFORCER_V1 ===
from __future__ import annotations

import json
import os
import re
import sqlite3
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Optional

BASE = Path("/root/.areal-neva-core")
TEMPLATE_DIR = BASE / "data/templates"
PROJECT_DIR = TEMPLATE_DIR / "project"
TECH_DIR = TEMPLATE_DIR / "technadzor"
INDEX = TEMPLATE_DIR / "index.json"

for d in (PROJECT_DIR, TECH_DIR):
    d.mkdir(parents=True, exist_ok=True)

SAVE_WORDS = ("образец", "шаблон", "пример", "возьми это", "сохрани это", "запомни это")
APPLY_WORDS = ("по образцу", "по шаблону", "как в образце", "как шаблон", "сделай так же")
PROJECT_WORDS = ("проект", "чертеж", "чертёж", "dwg", "dxf", "pdf", "кж", "км", "кмд", "ар", "проектирование")
TECH_WORDS = ("акт", "технадзор", "дефект", "замечан", "осмотр", "гост", "сп", "снип")

def _s(v: Any, limit: int = 10000) -> str:
    if v is None:
        return ""
    if isinstance(v, (dict, list)):
        try:
            v = json.dumps(v, ensure_ascii=False)
        except Exception:
            v = str(v)
    return str(v).strip()[:limit]

def _safe(v: Any) -> str:
    return re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", _s(v, 120)).strip("._") or "template"

def _load_index() -> Dict[str, Any]:
    # === TEMPLATE_INDEX_DICT_FIX_V1 ===
    if INDEX.exists():
        try:
            data = json.loads(INDEX.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                return data
            return {
                "_schema": "TEMPLATE_INDEX_DICT_FIX_V1",
                "_legacy_type": type(data).__name__,
                "_legacy_data": data,
            }
        except Exception as e:
            return {
                "_schema": "TEMPLATE_INDEX_DICT_FIX_V1",
                "_legacy_error": str(e),
            }
    return {"_schema": "TEMPLATE_INDEX_DICT_FIX_V1"}
    # === END_TEMPLATE_INDEX_DICT_FIX_V1 ===

def _save_index(idx: Dict[str, Any]) -> None:
    # === TEMPLATE_INDEX_SAVE_DICT_FIX_V1 ===
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    if not isinstance(idx, dict):
        idx = {
            "_schema": "TEMPLATE_INDEX_SAVE_DICT_FIX_V1",
            "_legacy_type": type(idx).__name__,
            "_legacy_data": idx,
        }
    INDEX.write_text(json.dumps(idx, ensure_ascii=False, indent=2), encoding="utf-8")
    # === END_TEMPLATE_INDEX_SAVE_DICT_FIX_V1 ===

def _detect_domain(text: str) -> Optional[str]:
    low = (text or "").lower()
    if any(x in low for x in TECH_WORDS):
        return "technadzor"
    if any(x in low for x in PROJECT_WORDS):
        return "project"
    return None

def _is_save_template(text: str) -> bool:
    low = (text or "").lower()
    return any(x in low for x in SAVE_WORDS)

def _is_apply_template(text: str) -> bool:
    low = (text or "").lower()
    return any(x in low for x in APPLY_WORDS)

def _last_relevant_task(conn: sqlite3.Connection, chat_id: str, topic_id: int, domain: str) -> Optional[Dict[str, Any]]:
    conn.row_factory = sqlite3.Row
    like = "%акт%" if domain == "technadzor" else "%проект%"
    row = conn.execute(
        """
        SELECT * FROM tasks
        WHERE chat_id=? AND COALESCE(topic_id,0)=?
          AND (
            input_type IN ('drive_file','file','document','photo','image')
            OR raw_input LIKE '%file_id%'
            OR raw_input LIKE '%file_name%'
            OR result LIKE ?
            OR raw_input LIKE ?
          )
        ORDER BY updated_at DESC, created_at DESC
        LIMIT 1
        """,
        (str(chat_id), int(topic_id or 0), like, like),
    ).fetchone()
    return {k: row[k] for k in row.keys()} if row else None

def _write_docx(title: str, body: str, out_name: str) -> str:
    out = Path(tempfile.gettempdir()) / out_name
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(title, level=1)
        for part in (body or "").split("\n"):
            doc.add_paragraph(part)
        doc.save(out)
        return str(out)
    except Exception:
        txt = out.with_suffix(".txt")
        txt.write_text(title + "\n\n" + body, encoding="utf-8")
        return str(txt)

def _write_xlsx(meta: Dict[str, Any], out_name: str) -> str:
    out = Path(tempfile.gettempdir()) / out_name
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Template"
        for i, (k, v) in enumerate(meta.items(), 1):
            ws.cell(i, 1, str(k))
            ws.cell(i, 2, json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else str(v))
        wb.save(out)
        wb.close()
        return str(out)
    except Exception:
        csv = out.with_suffix(".csv")
        csv.write_text("\n".join(f"{k};{v}" for k, v in meta.items()), encoding="utf-8")
        return str(csv)

def _zip(paths: list[str], name: str) -> str:
    out = Path(tempfile.gettempdir()) / f"{_safe(name)}.zip"
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if p and os.path.exists(p):
                z.write(p, arcname=os.path.basename(p))
    return str(out)

def save_template(conn: sqlite3.Connection, chat_id: str, topic_id: int, domain: str, task: sqlite3.Row, user_text: str) -> Dict[str, Any]:
    source = _last_relevant_task(conn, chat_id, topic_id, domain)
    source_payload = {}
    if source:
        source_payload = {
            "source_task_id": source.get("id"),
            "source_input_type": source.get("input_type"),
            "source_raw_input": _s(source.get("raw_input"), 3000),
            "source_result": _s(source.get("result"), 3000),
        }

    tid = _s(task["id"] if "id" in task.keys() else datetime.now(timezone.utc).timestamp())
    model = {
        "schema": f"{domain.upper()}_TEMPLATE_MODEL_V1",
        "template_id": tid,
        "domain": domain,
        "chat_id": str(chat_id),
        "topic_id": int(topic_id or 0),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "user_text": _s(user_text, 3000),
        **source_payload,
    }

    target_dir = PROJECT_DIR if domain == "project" else TECH_DIR
    path = target_dir / f"{_safe(tid)}.json"
    path.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")

    idx = _load_index()
    idx[f"topic_{int(topic_id or 0)}_active_{domain}_template"] = str(path)
    idx[f"chat_{chat_id}_topic_{int(topic_id or 0)}_active_{domain}_template"] = str(path)
    _save_index(idx)

    return {
        "handled": True,
        "state": "DONE",
        "result": f"Шаблон сохранён\nТип: {domain}\nTopic: {topic_id}\nTemplate: {path.name}",
        "event": f"{domain.upper()}_TEMPLATE_WORKFLOW_FULL_CLOSE_V1:SAVED",
    }

def _load_active_template(chat_id: str, topic_id: int, domain: str) -> Optional[Dict[str, Any]]:
    idx = _load_index()
    path = idx.get(f"chat_{chat_id}_topic_{int(topic_id or 0)}_active_{domain}_template") or idx.get(f"topic_{int(topic_id or 0)}_active_{domain}_template")
    if not path:
        return None
    p = Path(path)
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None

def apply_template(chat_id: str, topic_id: int, domain: str, user_text: str) -> Dict[str, Any]:
    tpl = _load_active_template(chat_id, topic_id, domain)
    if not tpl:
        return {
            "handled": True,
            "state": "WAITING_CLARIFICATION",
            "result": f"Активный шаблон {domain} в этом топике не найден. Пришли файл/акт/проект и напиши: возьми это как образец",
            "event": f"{domain.upper()}_TEMPLATE_WORKFLOW_FULL_CLOSE_V1:NO_TEMPLATE",
        }

    title = "Проект по сохранённому шаблону" if domain == "project" else "Акт по сохранённому шаблону"
    body = "\n".join([
        title,
        f"Template ID: {tpl.get('template_id')}",
        f"Source task: {tpl.get('source_task_id','')}",
        f"Запрос: {user_text}",
        "",
        "Источник шаблона:",
        _s(tpl.get("source_raw_input") or tpl.get("source_result"), 4000),
    ])

    docx = _write_docx(title, body, f"{domain}_template_result_{_safe(tpl.get('template_id'))}.docx")
    xlsx = _write_xlsx(tpl, f"{domain}_template_model_{_safe(tpl.get('template_id'))}.xlsx")
    model_path = Path(tempfile.gettempdir()) / f"{domain}_template_model_{_safe(tpl.get('template_id'))}.json"
    model_path.write_text(json.dumps(tpl, ensure_ascii=False, indent=2), encoding="utf-8")
    package = _zip([docx, xlsx, str(model_path)], f"{domain}_template_package_{_safe(tpl.get('template_id'))}")

    link = ""
    try:
        from core.artifact_upload_guard import upload_many_or_fail
        up = upload_many_or_fail([{"path": package, "kind": f"{domain}_template_package"}], f"{domain}_template_{tpl.get('template_id')}", int(topic_id or 0))
        link = ((up.get("links") or {}).get(package) or "")
    except Exception:
        link = ""

    result = "\n".join([
        f"{title} создан",
        f"Артефакт: {package}",
        f"Drive: {link or 'не подтверждён'}",
    ])

    return {
        "handled": True,
        "state": "AWAITING_CONFIRMATION",
        "result": result,
        "artifact_path": package,
        "drive_link": link,
        "event": f"{domain.upper()}_TEMPLATE_WORKFLOW_FULL_CLOSE_V1:APPLIED",
    }

def maybe_handle_template_workflow(conn: sqlite3.Connection, task: sqlite3.Row, chat_id: str, topic_id: int) -> Optional[Dict[str, Any]]:
    raw = _s(task["raw_input"] if "raw_input" in task.keys() else "")
    clean = re.sub(r"^\s*\[VOICE\]\s*", "", raw, flags=re.I).strip()
    low = clean.lower()

    domain = _detect_domain(low)
    if not domain:
        return None

    if _is_save_template(low):
        return save_template(conn, chat_id, topic_id, domain, task, clean)

    if _is_apply_template(low):
        return apply_template(chat_id, topic_id, domain, clean)

    return None
# === END_TEMPLATE_SCOPE_ENFORCER_V1 ===
# === END_TECHNADZOR_ACT_TEMPLATE_WORKFLOW_V1 ===
# === END_PROJECT_TEMPLATE_WORKFLOW_FULL_CLOSE_V1 ===


# === PROJECT_DWG_TEMPLATE_REUSE_V2 ===
async def async_apply_template(chat_id: str, topic_id: int, domain: str, user_text: str) -> Optional[Dict[str, Any]]:
    tpl = _load_active_template(chat_id, topic_id, domain)
    if not tpl:
        return {
            "handled": True,
            "state": "WAITING_CLARIFICATION",
            "result": f"Активный шаблон {domain} не найден. Пришли файл и напиши: возьми это как образец",
            "event": f"{domain.upper()}_TEMPLATE_APPLY:NO_TEMPLATE",
        }

    tid = _safe(tpl.get("template_id") or "tpl")
    source_text = _s(tpl.get("source_raw_input") or tpl.get("source_result"), 3000)

    if domain == "project":
        try:
            from core.project_engine import create_project_pdf_dxf_artifact
            combined = str(user_text or "") + " " + str(source_text or "")
            result = await create_project_pdf_dxf_artifact(
                raw_input=combined,
                task_id="tpl_" + tid,
                topic_id=int(topic_id or 0),
                template_hint=_s(tpl.get("source_raw_input"), 500),
                require_template=False,
            )
            if result and result.get("success") and result.get("artifact_path"):
                link = result.get("drive_link") or ""
                if not link:
                    try:
                        from core.artifact_upload_guard import upload_many_or_fail
                        up = upload_many_or_fail(
                            [{"path": result["artifact_path"], "kind": "project_template_package"}],
                            "tpl_" + tid,
                            int(topic_id or 0),
                        )
                        link = list((up.get("links") or {}).values())[0] if up.get("links") else ""
                    except Exception:
                        pass
                res_text = "Проект создан по сохранённому шаблону"
                res_text += "\nПакет: " + str(result.get("artifact_path", ""))
                res_text += "\nDrive: " + (link or "не подтверждён")
                if result.get("region_detected"):
                    res_text += "\nРегион нагрузок: " + str(result.get("region_detected"))
                return {
                    "handled": True,
                    "state": "AWAITING_CONFIRMATION",
                    "result": res_text,
                    "artifact_path": result.get("artifact_path"),
                    "drive_link": link,
                    "event": "PROJECT_DWG_TEMPLATE_REUSE_V2:APPLIED_REAL_ENGINE",
                }
            err = (result or {}).get("error") or "PROJECT_ENGINE_RETURNED_NO_ARTIFACT"
            return {
                "handled": True,
                "state": "FAILED",
                "result": "",
                "error": "PROJECT_DWG_TEMPLATE_REUSE_V2:ENGINE_FAILED:" + str(err)[:200],
                "event": "PROJECT_DWG_TEMPLATE_REUSE_V2:ENGINE_FAILED",
            }
        except Exception as e:
            import logging
            logging.getLogger(__name__).warning("PROJECT_DWG_TEMPLATE_REUSE_V2_ERR %s", e)
            return {
                "handled": True,
                "state": "FAILED",
                "result": "",
                "error": "PROJECT_DWG_TEMPLATE_REUSE_V2:EXCEPTION:" + str(e)[:200],
                "event": "PROJECT_DWG_TEMPLATE_REUSE_V2:EXCEPTION",
            }

    return apply_template(chat_id, topic_id, domain, user_text)

async def maybe_handle_template_workflow_async(conn: sqlite3.Connection, task: sqlite3.Row, chat_id: str, topic_id: int) -> Optional[Dict[str, Any]]:
    raw = _s(task["raw_input"] if "raw_input" in task.keys() else "")
    clean = re.sub(r"^\s*\[VOICE\]\s*", "", raw, flags=re.I).strip()
    low = clean.lower()
    domain = _detect_domain(low)
    if not domain:
        return None
    if _is_save_template(low):
        return save_template(conn, chat_id, topic_id, domain, task, clean)
    if _is_apply_template(low):
        return await async_apply_template(chat_id, topic_id, domain, clean)
    return None
# === END_PROJECT_DWG_TEMPLATE_REUSE_V2 ===
