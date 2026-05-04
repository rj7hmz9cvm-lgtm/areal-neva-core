from core.gemini_vision import analyze_image_file  # GEMINI_VISION_V43
import os, logging, asyncio
from typing import Dict, Any, Optional, List

logger = logging.getLogger(__name__)

INTENT_TRIGGERS = {
    "estimate": [
        "смет", "расчёт", "расчет", "стоимость", "цена", "бюджет",
        "сделай таблицу", "таблицу", "вытащи объёмы", "вытащи объемы",
        "посчитай объём", "посчитай объем", "объём", "объем",
        "excel", "эксель", "google таблицу", "google таблица", "google sheets",
        "sheets", "спецификац", "ведомост", "бетон", "арматур", "фундамент"
    ],
    "ocr": [
        "распознай таблицу", "распознать таблицу", "распозна", "ocr",
        "скан", "текст с фото", "фото в таблицу"
    ],
    "technadzor": [
        "проверь дефекты", "акт технадзора", "дефекты", "дефект",
        "нарушение", "нарушения", "косяк", "осмотр", "технадзор",
        "проверь фото", "проверь узел", "проверь"
    ],
    "dwg": ["чертёж", "чертеж", "dxf", "dwg", "проект", "автокад"],
    "template": ["сделай так же", "по образцу", "как в", "шаблон"],
    "vision": ["анализ фото", "анализ схемы", "разбери фото", "что на фото", "что на схеме"],
    "search": ["найди в сметах", "поиск по сметам", "искать в старых"],
}
FORMAT_TRIGGERS = {
    "excel": ["excel", "эксель", "xlsx", "иксель"],
    "sheets": ["google sheets", "google таблиц", "гугл таблиц", "sheets", "таблицу на диск", "google sheet"],
    "word": ["word", "ворд", "docx"],
    "docs": ["google docs", "google документ", "гугл документ", "docs"],
}

def detect_intent(text: str) -> Optional[str]:
    if not text: return None
    t = text.lower()
    for intent, triggers in INTENT_TRIGGERS.items():
        if any(tr in t for tr in triggers): return intent
    return None

def detect_format(text: str) -> str:
    t = (text or "").lower()
    if any(x in t for x in FORMAT_TRIGGERS["sheets"]): return "sheets"
    elif any(x in t for x in FORMAT_TRIGGERS["docs"]): return "docs"
    elif any(x in t for x in FORMAT_TRIGGERS["word"]): return "word"
    elif any(x in t for x in FORMAT_TRIGGERS["excel"]): return "excel"
    return "excel"

def format_priority(file_name: str, available_files: list = None) -> str:
    """FORMAT_PRIORITY_V1 — DWG/DXF > XLSX > DOCX > PDF > IMAGE"""
    files = available_files or [file_name]
    # Приоритет по канону §11.9
    for ext in [".dwg", ".dxf"]:
        if any(f.lower().endswith(ext) for f in files):
            return "dwg"
    for ext in [".xlsx", ".xls", ".csv"]:
        if any(f.lower().endswith(ext) for f in files):
            return "excel"
    for ext in [".docx", ".doc"]:
        if any(f.lower().endswith(ext) for f in files):
            return "word"
    for ext in [".pdf"]:
        if any(f.lower().endswith(ext) for f in files):
            return "pdf"
    for ext in [".jpg", ".jpeg", ".png", ".heic", ".webp", ".tiff", ".bmp"]:
        if any(f.lower().endswith(ext) for f in files):
            return "image"
    return "unknown"

def get_topic_role(topic_id: int) -> str:
    try:
        import sqlite3
        conn = sqlite3.connect('/root/.areal-neva-core/data/memory.db')
        cur = conn.cursor()
        cur.execute("SELECT value FROM memory WHERE key = ?", (f"topic_{topic_id}_role",))
        row = cur.fetchone(); conn.close()
        return row[0] if row else ""
    except: return ""

def get_clarification_message(file_name: str, topic_id: int) -> str:
    role = get_topic_role(topic_id)
    base = f"📎 Получил файл «{file_name}».\nЧто с ним сделать?"
    if topic_id == 2 or "СТРОЙКА" in role:
        return base + "\n• Смета / Расчёт\n• Проектирование / Расчёт нагрузок\n• Анализ фото / Схема\n• Распознать таблицу\n• Просто сохранить\n\nВ каком формате?\n• Excel\n• Google Sheets"
    elif topic_id == 5 or "ТЕХНАДЗОР" in role:
        return base + "\n• Проверить дефекты / Составить акт\n• Анализ фото / Схема\n• Распознать таблицу\n• Просто сохранить\n\nВ каком формате?\n• Word\n• Google Docs"
    else:
        return base + "\n• Анализ фото / Схема\n• Распознать таблицу\n• Проверить дефекты\n• Смета / Расчёт\n• Просто сохранить"

# === CANON_PASS_INTAKE_OFFER_SENT ===
_OFFER_SENT = {}
def mark_offer_sent(tid): _OFFER_SENT[tid] = True
def is_offer_sent(tid): return _OFFER_SENT.get(tid, False)
# === END CANON_PASS_INTAKE_OFFER_SENT ===

def should_ask_clarification(raw_input: str, has_file: bool, already_asked: bool = False) -> bool:
    # === ANTI_REPEAT_V1 — не спрашивать если уже спрашивали ===
    if already_asked:
        return False
    if not has_file: return False
    return detect_intent(raw_input) is None


ESTIMATE_FILENAME_TRIGGERS = ["кж","кд","спецификац","ведомост","смет","расход","арматур","бетон","плит","конструкци","фундамент","у1-","у2-","кр-"]

def detect_intent_from_filename(file_name: str) -> Optional[str]:
    fn = (file_name or "").lower()
    if any(t in fn for t in ESTIMATE_FILENAME_TRIGGERS):
        return "estimate"
    if any(t in fn for t in ["акт","дефект","осмотр","технадзор"]):
        return "technadzor"
    if any(t in fn for t in [".dwg",".dxf","чертеж","чертёж"]):
        return "dwg"
    return None

async def route_file(file_path: str, task_id: str, topic_id: int, intent: str, fmt: str = "excel") -> Optional[Dict[str, Any]]:
    try:
        from core.engine_base import detect_real_file_type
        real_type = detect_real_file_type(file_path)
        # === UNIVERSAL_FILE_HANDLER_V1_WIRED ===
        if real_type in ("zip", "rar", "7z", "dwg", "dxf", "video", "mp4", "audio", "text_fallback", "unknown") or intent == "dwg":
            try:
                from core.universal_file_handler import extract_text_from_file
                _ufh = extract_text_from_file(file_path, task_id, topic_id)
                _ufh_text = (_ufh.get("text") or "").strip()
                _ufh_rows = _ufh.get("rows") or []
                if _ufh.get("success") and (_ufh_text or _ufh_rows):
                    _summary = f"\u0424\u0430\u0439\u043b \u043e\u0431\u0440\u0430\u0431\u043e\u0442\u0430\u043d ({_ufh.get('type','unknown')}):\n"
                    if _ufh_rows:
                        _summary += f"\u0421\u0442\u0440\u043e\u043a \u0434\u0430\u043d\u043d\u044b\u0445: {len(_ufh_rows)}\n"
                        for row in _ufh_rows[:5]:
                            _summary += "  " + " | ".join(str(c) for c in row if c) + "\n"
                    if _ufh_text:
                        _summary += _ufh_text[:1500]
                    return {"success": True, "text": _summary, "type": _ufh.get("type")}
                else:
                    return {"success": False, "error": _ufh.get("error") or f"\u0424\u043e\u0440\u043c\u0430\u0442 {real_type} \u043d\u0435 \u043f\u043e\u0434\u0434\u0435\u0440\u0436\u0438\u0432\u0430\u0435\u0442\u0441\u044f"}
            except Exception as _ufh_e:
                return {"success": False, "error": f"UNIVERSAL_HANDLER_ERROR: {_ufh_e}"}
        # === END UNIVERSAL_FILE_HANDLER_V1_WIRED ===
        if real_type in ("csv", "txt") and intent not in ("estimate", "technadzor", "ocr"):
            intent = "estimate"
        if real_type == "dwg":
            intent = "dwg"
        if intent == "estimate":
            if fmt == "sheets":
                from core.sheets_generator import create_google_sheet
                from core.estimate_engine import process_estimate_to_excel
                data = await process_estimate_to_excel(file_path, task_id, topic_id)
                if data.get("excel_path"):
                    from openpyxl import load_workbook
                    wb = load_workbook(data["excel_path"]); ws = wb.active
                    rows = [[cell.value for cell in row] for row in ws.iter_rows()]; wb.close()
                    link = None
                    try:
                        link = create_google_sheet(f"Estimate_{task_id[:8]}", rows)
                    except Exception as e:
                        err_str = str(e)
                        if "403" in err_str or "permission" in err_str.lower() or "quota" in err_str.lower():
                            logger.warning(f"Google Sheets 403/quota -> fallback XLSX: {e}")
                        else:
                            logger.warning(f"create_google_sheet fallback to XLSX: {e}")
                        link = None
                    if link and "docs.google.com" in str(link):
                        return {"success": True, "drive_link": link, "artifact_path": data["excel_path"]}
                    # Fallback: return XLSX artifact
                    return {"success": True, "artifact_path": data["excel_path"], "excel_path": data["excel_path"]}
            else:
                from core.estimate_engine import process_estimate_to_excel
                return await process_estimate_to_excel(file_path, task_id, topic_id)
        elif intent == "ocr":
            from core.ocr_engine import process_image_to_excel
            return await process_image_to_excel(file_path, task_id, topic_id)
        elif intent == "technadzor":
            from core.technadzor_engine import process_defect_to_report
            data = await process_defect_to_report(file_path, task_id, topic_id)
            if not data or not data.get("success"):
                return {"success": False, "error": (data.get("error") if isinstance(data, dict) else None) or "TECHNADZOR_FAILED"}
            rp = data.get("report_path")
            if rp:
                import os as _os
                rp_size = _os.path.getsize(rp) if _os.path.exists(rp) else 0
                if rp_size < 1000:
                    return {"success": False, "error": "DOCUMENT_EMPTY_RESULT: DOCX too small"}
                from docx import Document as _Doc
                _doc = _Doc(rp)
                _has_content = any(p.text.strip() for p in _doc.paragraphs) or len(_doc.tables) > 0
                if not _has_content:
                    return {"success": False, "error": "DOCUMENT_EMPTY_RESULT: no paragraphs or tables"}
            if fmt == "docs" and rp:
                try:
                    from core.docs_generator import create_google_doc
                    from docx import Document
                    doc = Document(rp)
                    content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    link = create_google_doc(f"Defect_{task_id[:8]}", content)
                    if link:
                        return {"success": True, "drive_link": link, "artifact_path": rp}
                except Exception as _e:
                    logger.warning(f"create_google_doc fallback to DOCX: {_e}")
            return data
        elif intent == "dwg":
            from core.dwg_engine import process_dwg_to_excel
            return await process_dwg_to_excel(file_path, task_id, topic_id)
        elif intent == "template":
            from core.template_manager import apply_template, get_template
            tmpl = get_template(str(topic_id), topic_id, "estimate")
            if tmpl:
                out = f"/tmp/{task_id}_template.xlsx"
                # Extract real data from source file if possible
                real_rows = []
                try:
                    real_type_t = detect_real_file_type(file_path)
                    if real_type_t in ("xlsx", "zip_or_office"):
                        from openpyxl import load_workbook
                        wb_t = load_workbook(file_path, data_only=True)
                        ws_t = wb_t.active
                        for row_t in ws_t.iter_rows(min_row=2, values_only=True):
                            if row_t and any(v is not None for v in row_t):
                                real_rows.append(list(row_t))
                        wb_t.close()
                    elif real_type_t == "pdf":
                        try:
                            import pdfplumber
                            with pdfplumber.open(file_path) as pdf_t:
                                for page_t in pdf_t.pages:
                                    tables_t = page_t.extract_tables()
                                    for table_t in tables_t:
                                        for row_t in table_t[1:]:
                                            if row_t and any(v for v in row_t if v):
                                                real_rows.append(row_t)
                        except Exception:
                            pass
                except Exception as _te:
                    logger.warning(f"template data extraction failed: {_te}")
                if apply_template(tmpl, out, real_rows): return {"success": True, "excel_path": out}
                elif apply_template(tmpl, out, []): return {"success": True, "excel_path": out}
        elif intent == "vision":
            result = await analyze_image_file(file_path, None)
            return {"success": True, "engine": "gemini_vision", "result_text": result, "text_result": result}
        elif intent == "search":
            from core.search_engine import search_in_estimates
            results = search_in_estimates(file_path, topic_id)
            return {"success": True, "text_result": "\n".join(results[:10])}
        return {"success": False, "error": f"PIPELINE_NOT_EXECUTED: no handler for intent={intent}"}
    except Exception as e:
        logger.error(f"route_file: {e}", exc_info=True)
        return {"success": False, "error": str(e)}

async def handle_multiple_files(file_paths: List[str], task_id: str, topic_id: int, intent: str, fmt: str = "excel") -> Optional[Dict[str, Any]]:
    if intent == "estimate":
        from core.multi_file_orchestrator import merge_estimate_results
        results = []
        for fp in file_paths:
            from core.estimate_engine import process_estimate_to_excel
            r = await process_estimate_to_excel(fp, task_id, topic_id)
            if r: results.append(r)
        return merge_estimate_results(results)
    return None

# === ARK_KZH_KD_TRIGGERS ===
ESTIMATE_FILENAME_TRIGGERS = list(set(ESTIMATE_FILENAME_TRIGGERS + [
    'арк', 'кж', 'кд', 'кр', 'ов', 'вк', 'эо', 'гп',
    'марка', 'раздел', 'спецификация', 'ведомость',
]))

_ARK_SECTION_MAP = {
    'арк': 'estimate', 'кж': 'estimate', 'кд': 'estimate',
    'кр': 'estimate', 'ов': 'estimate', 'вк': 'estimate',
    'эо': 'estimate', 'гп': 'estimate',
}

def detect_intent_from_filename_v2(file_name: str):
    fn = (file_name or '').lower()
    for key, intent in _ARK_SECTION_MAP.items():
        if key in fn:
            return intent
    return detect_intent_from_filename(file_name)
# === END ARK_KZH_KD_TRIGGERS ===

# === ALL_CONTOURS_ROUTE_FILE_V2 ===
try:
    _all_contours_orig_route_file=route_file
except Exception:
    _all_contours_orig_route_file=None
def _ac_words(*groups):
    return ["".join(chr(x) for x in g) for g in groups]
def _ac_has(text, words):
    t=str(text or "").lower()
    return any(w in t for w in words)
async def route_file(file_path, task_id, topic_id=0, *args, **kwargs):
    import inspect
    fp=str(file_path or "")
    raw=" ".join([fp,str(kwargs.get("raw_input") or ""),str(kwargs.get("prompt") or ""),str(kwargs.get("user_text") or "")]+[str(a or "") for a in args])
    template_words=_ac_words((1080,1089,1087,1086,1083,1100,1079,1091,1081,32,1082,1072,1082,32,1096,1072,1073,1083,1086,1085),(1089,1086,1093,1088,1072,1085,1080,32,1096,1072,1073,1083,1086,1085),(1101,1090,1086,32,1096,1072,1073,1083,1086,1085))
    if _ac_has(raw, template_words):
        from core.template_manager import save_template
        return {"success":True,"intent":"template","template_path":save_template(fp,int(topic_id or 0),"template")}
    kzh_words=_ac_words((1082,1078),(1082,1076),(1072,1088,1082),(1082,1088))
    detector=globals().get("detect_intent_from_filename_v2") or globals().get("detect_intent_from_filename")
    intent=detector(fp) if detector else "unknown"
    if intent=="estimate" and _ac_has(fp,kzh_words):
        from core.estimate_engine import process_kzh_pdf
        return await process_kzh_pdf(fp,str(task_id),int(topic_id or 0))
    if _all_contours_orig_route_file is None:
        return {"success":False,"error":"original_route_file_missing","intent":intent}
    res=_all_contours_orig_route_file(fp,task_id,topic_id,*args,**kwargs)
    if inspect.isawaitable(res):
        res=await res
    sheets_words=("google sheets","fmt=sheets","sheets")+tuple(_ac_words((1075,1091,1075,1083,32,1090,1072,1073,1083),(1075,1091,1075,1083,32,1090,1072,1073,1083,1080,1094),(1092,1086,1088,1084,1072,1090,32,115,104,101,101,116,115)))
    if isinstance(res,dict) and _ac_has(raw,sheets_words):
        xl=res.get("excel_path") or res.get("xlsx_path")
        if xl:
            try:
                from core.sheets_generator import create_google_sheet
                sh=create_google_sheet(xl,task_id=task_id,topic_id=topic_id)
                if inspect.isawaitable(sh):
                    sh=await sh
                if sh:
                    res["sheets_link"]=sh.get("url") if isinstance(sh,dict) else str(sh)
                    res["drive_link"]=res.get("drive_link") or res["sheets_link"]
            except Exception as e:
                res["sheets_error"]=str(e)[:300]
    return res
# === END_ALL_CONTOURS_ROUTE_FILE_V2 ===

# === FINAL_CODE_CONTOUR_FILE_INTAKE_V1 ===
try:
    ESTIMATE_FILENAME_TRIGGERS=list(set(ESTIMATE_FILENAME_TRIGGERS+[_f for _f in ["km","kmd"]]))
except Exception:
    pass
try:
    _ARK_SECTION_MAP.update({"km":"estimate","kmd":"estimate"})
except Exception:
    _ARK_SECTION_MAP={"km":"estimate","kmd":"estimate"}
try:
    _final_orig_route_file=route_file
except Exception:
    _final_orig_route_file=None
async def route_file(file_path, task_id, topic_id=0, *args, **kwargs):
    import inspect
    fp=str(file_path or "")
    detector=globals().get("detect_intent_from_filename_v2") or globals().get("detect_intent_from_filename")
    intent=detector(fp) if detector else "unknown"
    low=fp.lower()
    if intent=="estimate" and any(x in low for x in ["km","kmd","kzh","kd","ark","kr"]):
        from core.estimate_engine import process_kzh_pdf
        return await process_kzh_pdf(fp,str(task_id),int(topic_id or 0))
    if _final_orig_route_file is None:
        return {"success":False,"error":"original_route_file_missing","intent":intent}
    res=_final_orig_route_file(fp,task_id,topic_id,*args,**kwargs)
    if inspect.isawaitable(res):
        res=await res
    if isinstance(res,dict) and (kwargs.get("fmt")=="sheets" or str(kwargs.get("raw_input") or "").lower().find("sheets")>=0):
        xl=res.get("excel_path") or res.get("xlsx_path")
        if xl:
            from openpyxl import load_workbook
            wb=load_workbook(xl,data_only=False)
            ws=wb.active
            rows=[[cell.value for cell in row] for row in ws.iter_rows()]
            from core.sheets_generator import create_google_sheet
            sh=create_google_sheet(Path(xl).stem, rows)
            if inspect.isawaitable(sh): sh=await sh
            if sh:
                res["sheets_link"]=sh.get("url") if isinstance(sh,dict) else str(sh)
                res["drive_link"]=res.get("drive_link") or res["sheets_link"]
    return res
# === END_FINAL_CODE_CONTOUR_FILE_INTAKE_V1 ===

# === FILE_INTAKE_KM_V39 ===
try:
    ESTIMATE_FILENAME_TRIGGERS = list(set(list(ESTIMATE_FILENAME_TRIGGERS) + ["км","кмд","металл","конструкц"]))
except Exception:
    pass
try:
    _ARK_SECTION_MAP.update({"км": "estimate", "кмд": "estimate"})
except Exception:
    pass
# === END_FILE_INTAKE_KM_V39 ===

# === FILE_INTAKE_PROJECT_V41 ===

_PROJECT_V41_TRIGGERS = ("кж","км","кмд","ар","ов","вк","эом","сс","гп","пз","тх")

def _v41_project_section_hit(text):
    src = str(text or "").lower()
    import re
    for key in _PROJECT_V41_TRIGGERS:
        if re.search(r"(^|[^а-яa-z0-9])" + re.escape(key) + r"([^а-яa-z0-9]|$)", src, re.I):
            return True
    return False

try:
    _v41_orig_detect_intent = detect_intent
    def detect_intent(file_name: str, raw_input: str = "", *args, **kwargs):
        src = (str(file_name or "") + " " + str(raw_input or "")).lower()
        if _v41_project_section_hit(src):
            return "project"
        return _v41_orig_detect_intent(file_name)  # DETECT_INTENT_FIX_V1
except Exception:
    pass

if "route_file" in globals():
    _v41_orig_route_file = route_file

    async def route_file(file_path, task_id, topic_id=0, intent=None, fmt="excel", *args, **kwargs):
        raw_input = str(kwargs.get("raw_input") or "")
        file_name = str(file_path or "")
        final_intent = intent or detect_intent(file_name, raw_input)

        if final_intent == "project":
            try:
                from core.project_engine import process_project_file
                return await process_project_file(file_path, task_id, topic_id, raw_input)
            except Exception as e:
                return {"success": False, "error": "PROJECT_ENGINE_FAILED: " + str(e)[:300]}

        res = _v41_orig_route_file(file_path, task_id, topic_id, final_intent, fmt, *args, **kwargs)
        import inspect
        if inspect.isawaitable(res):
            res = await res

        if isinstance(res, dict) and res.get("success") is False:
            return res
        if not isinstance(res, dict):
            return {"success": False, "error": "FILE_RESULT_GUARD: route_file returned empty"}
        if not (res.get("drive_link") or res.get("sheets_link") or res.get("doc_link") or res.get("excel_path") or res.get("artifact_path") or res.get("text") or res.get("result")):
            return {"success": False, "error": "FILE_RESULT_GUARD: no usable output"}
        return res

# === END_FILE_INTAKE_PROJECT_V41 ===

# === FILE_INTAKE_PROJECT_SAFE_V42 ===

_PROJECT_V42_TRIGGERS = ("кж","км","кмд","ар","ов","вк","эом","сс","гп","пз","тх")

def _v42_project_choice(raw_input: str) -> bool:
    t = str(raw_input or "").lower()
    return "проектирование" in t or "расчёт нагрузок" in t or "расчет нагрузок" in t

if "route_file" in globals():
    _v42_orig_route_file = route_file

    async def route_file(file_path, task_id, topic_id=0, intent=None, fmt="excel", *args, **kwargs):
        import inspect
        raw_input = str(kwargs.get("raw_input") or "")

        if _v42_project_choice(raw_input):
            try:
                from core.project_engine import process_project_file
                res = await process_project_file(file_path, task_id, topic_id, raw_input)
                if not isinstance(res, dict):
                    return {"success": False, "error": "PROJECT_RESULT_GUARD: empty_payload"}
                if res.get("success") is False:
                    return res
                if not (res.get("drive_link") or res.get("excel_path") or res.get("docx_path") or res.get("pdf_path")):
                    return {"success": False, "error": "PROJECT_RESULT_GUARD: no_artifact"}
                return res
            except Exception as e:
                return {"success": False, "error": "PROJECT_ENGINE_FAILED: " + str(e)[:300]}

        if str(intent or "").lower() == "project" and not _v42_project_choice(raw_input):
            intent = "estimate"

        res = _v42_orig_route_file(file_path, task_id, topic_id, intent, fmt, *args, **kwargs)
        if inspect.isawaitable(res):
            res = await res

        if isinstance(res, dict) and res.get("success") is False:
            return res

        if not isinstance(res, dict):
            return {"success": False, "error": "FILE_RESULT_GUARD: route_file returned empty"}

        if not (res.get("drive_link") or res.get("sheets_link") or res.get("doc_link") or res.get("excel_path") or res.get("artifact_path") or res.get("text") or res.get("result")):
            return {"success": False, "error": "FILE_RESULT_GUARD: no usable output"}

        return res

# === END_FILE_INTAKE_PROJECT_SAFE_V42 ===

# === CODE_CLOSE_V43_FILE_INTAKE ===

def _v43_is_template_learn(raw_input):
    t = str(raw_input or "").lower()
    return "образец" in t or "шаблон" in t or "запомни структуру" in t

def _v43_is_project_choice(raw_input):
    t = str(raw_input or "").lower()
    return "проектирование" in t or "расчёт нагрузок" in t or "расчет нагрузок" in t

if "route_file" in globals():
    _v43_orig_route_file = route_file

    async def route_file(file_path, task_id, topic_id=0, intent=None, fmt="excel", *args, **kwargs):
        import inspect
        raw_input = str(kwargs.get("raw_input") or "")

        if _v43_is_template_learn(raw_input):
            try:
                from core.template_manager import template_learn_v43
                ok = template_learn_v43(file_path, topic_id, intent or "project")
                return {"success": bool(ok), "text": "Шаблон сохранён для топика" if ok else "TEMPLATE_LEARN_FAILED"}
            except Exception as e:
                return {"success": False, "error": "TEMPLATE_LEARN_FAILED: " + str(e)[:300]}

        if _v43_is_project_choice(raw_input):
            try:
                from core.template_manager import template_priority_v43
                template = template_priority_v43(topic_id, "project")
                from core.project_engine import process_project_file
                res = await process_project_file(file_path, task_id, topic_id, raw_input)
                if template:
                    res["template_used"] = template
                return res
            except Exception as e:
                return {"success": False, "error": "PROJECT_ENGINE_FAILED: " + str(e)[:300]}

        if str(intent or "").lower() == "project" and not _v43_is_project_choice(raw_input):
            intent = "estimate"

        res = _v43_orig_route_file(file_path, task_id, topic_id, intent, fmt, *args, **kwargs)
        if inspect.isawaitable(res):
            res = await res

        if isinstance(res, dict) and res.get("success") is False:
            return res
        if not isinstance(res, dict):
            return {"success": False, "error": "FILE_RESULT_GUARD: empty_route_result"}
        if not (res.get("drive_link") or res.get("sheets_link") or res.get("doc_link") or res.get("excel_path") or res.get("artifact_path") or res.get("text") or res.get("result")):
            return {"success": False, "error": "FILE_RESULT_GUARD: no_output"}
        return res

# === END_CODE_CLOSE_V43_FILE_INTAKE ===


# === FILE_INTAKE_KZH_INTENT_FIX_V1 ===
# Bare KЖ/КД/project files are project-context files, not estimate files.
# File upload without explicit user instruction must ask what to do.

import json as _fik_json
import os as _fik_os
import re as _fik_re
from typing import Optional as _FIKOptional

ESTIMATE_FILENAME_TRIGGERS = [
    "смет", "расход", "ведомост", "спецификац",
    "у1-", "у2-", "кр-",
]

PROJECT_FILENAME_TRIGGERS = [
    "кж", "кд", "кмд", "км", "ар",
    "плит", "фундамент", "конструкци", "конструкция",
    "узел", "цоколь", "разрез", "армирован", "чертеж", "чертёж",
]

def _fik_norm_text(value) -> str:
    return str(value or "").lower().replace("ё", "е").strip()

def _fik_word_hit(text: str, words) -> bool:
    t = _fik_norm_text(text)
    for w in words:
        ww = _fik_norm_text(w)
        if not ww:
            continue
        if len(ww) <= 3:
            if _fik_re.search(r"(^|[^а-яa-z0-9])" + _fik_re.escape(ww) + r"([^а-яa-z0-9]|$)", t, _fik_re.I):
                return True
        elif ww in t:
            return True
    return False

def _fik_extract_user_instruction(raw_input: str) -> str:
    """
    Extract only explicit user instruction.
    Ignore file_name/file_id/mime/source metadata, because filename alone must not auto-run pipeline.
    """
    raw = str(raw_input or "").strip()
    if not raw:
        return ""
    try:
        obj = _fik_json.loads(raw)
        if isinstance(obj, dict):
            for key in ("caption", "user_text", "text", "prompt", "comment", "message"):
                val = obj.get(key)
                if isinstance(val, str) and val.strip():
                    return val.strip()
            return ""
    except Exception:
        pass
    if raw.startswith("{") and ("file_name" in raw or "file_id" in raw or "mime_type" in raw):
        return ""
    return raw

def detect_intent_from_filename(file_name: str) -> _FIKOptional[str]:
    fn = _fik_norm_text(_fik_os.path.basename(str(file_name or "")))

    if _fik_word_hit(fn, PROJECT_FILENAME_TRIGGERS):
        return "project"

    if _fik_word_hit(fn, ESTIMATE_FILENAME_TRIGGERS):
        return "estimate"

    if _fik_word_hit(fn, ("акт", "дефект", "осмотр", "технадзор")):
        return "technadzor"

    if _fik_word_hit(fn, (".dwg", ".dxf", "dwg", "dxf", "чертеж", "чертёж")):
        return "dwg"

    return None

def detect_intent_from_filename_v2(file_name: str) -> _FIKOptional[str]:
    return detect_intent_from_filename(file_name)

def should_ask_clarification(raw_input: str, has_file: bool, already_asked: bool = False) -> bool:
    if already_asked:
        return False
    if not has_file:
        return False

    instruction = _fik_extract_user_instruction(raw_input)
    low = _fik_norm_text(instruction)

    service_words = (
        "tmp", "healthcheck", "service_file", "служебный", "синхронизации"
    )
    if _fik_word_hit(low, service_words):
        return False

    explicit_action_words = (
        "сделай", "делай", "создай", "сформируй", "подготовь", "разработай",
        "посчитай", "рассчитай", "проверь", "проанализируй", "выгрузи",
        "сохрани", "возьми", "прими", "используй", "распознай", "обработай",
        "шаблон", "образец", "по образцу", "по шаблону",
        "смет", "проект", "кж", "кд", "км", "кмд", "акт", "технадзор",
        "таблиц", "excel", "xlsx", "pdf", "dxf", "dwg", "ocr",
        "проектирование", "расчет нагрузок", "расчёт нагрузок",
    )

    return not _fik_word_hit(low, explicit_action_words)

try:
    _fik_orig_route_file = route_file
except Exception:
    _fik_orig_route_file = None

async def route_file(file_path, task_id, topic_id=0, intent=None, fmt="excel", *args, **kwargs):
    """
    Final guard:
    - KЖ/KД/project filenames never become estimate automatically
    - if caller passes estimate for project file without explicit estimate command, downgrade to project
    - raw file with no user instruction must return clarification payload
    """
    import inspect as _fik_inspect

    raw_input = str(kwargs.get("raw_input") or kwargs.get("prompt") or kwargs.get("user_text") or "")
    instruction = _fik_extract_user_instruction(raw_input)
    filename_intent = detect_intent_from_filename(str(file_path or ""))

    if not instruction and filename_intent in ("project", "estimate", "technadzor", "dwg"):
        try:
            msg = get_clarification_message(_fik_os.path.basename(str(file_path or "")), int(topic_id or 0))
        except Exception:
            msg = (
                "Файл принят.\n"
                "Что сделать с ним?\n\n"
                "1. Взять как шаблон проекта\n"
                "2. Взять как шаблон сметы\n"
                "3. Взять как шаблон технадзора\n"
                "4. Обработать как обычный файл\n\n"
                "Ответь одним сообщением"
            )
        return {
            "success": False,
            "needs_clarification": True,
            "state": "WAITING_CLARIFICATION",
            "intent": filename_intent,
            "result_text": msg,
            "error": "FILE_INTAKE_NEEDS_CONTEXT",
        }

    low_instruction = _fik_norm_text(instruction)
    explicit_estimate = _fik_word_hit(low_instruction, ("смет", "расход", "ведомост", "спецификац", "расчет", "расчёт", "стоимость", "цена", "объем", "объём"))

    if filename_intent == "project" and str(intent or "").lower() == "estimate" and not explicit_estimate:
        intent = "project"

    if intent is None:
        intent = filename_intent

    if _fik_orig_route_file is None:
        return {"success": False, "error": "original_route_file_missing", "intent": intent}

    res = _fik_orig_route_file(file_path, task_id, topic_id, intent, fmt, *args, **kwargs)
    if _fik_inspect.isawaitable(res):
        res = await res
    return res

# === END_FILE_INTAKE_KZH_INTENT_FIX_V1 ===


# === CONTEXT_AWARE_FILE_INTAKE_V1_ROUTER_WRAPPER ===
try:
    _ca_fir_orig_route_file = route_file
except Exception:
    _ca_fir_orig_route_file = None

async def route_file(file_path, task_id, topic_id=0, intent=None, fmt="excel", *args, **kwargs):
    import inspect
    raw_input = str(kwargs.get("raw_input") or kwargs.get("prompt") or kwargs.get("user_text") or "")
    if not raw_input:
        try:
            from core.file_context_intake import latest_pending_instruction_for_topic
            chat_id = str(kwargs.get("chat_id") or "")
            pending = latest_pending_instruction_for_topic(int(topic_id or 0), chat_id)
            if pending:
                kwargs["raw_input"] = pending
                raw_input = pending
        except Exception:
            pass

    if _ca_fir_orig_route_file is None:
        return {"success": False, "error": "CONTEXT_AWARE_FILE_INTAKE_V1: original_route_file_missing"}

    res = _ca_fir_orig_route_file(file_path, task_id, topic_id, intent, fmt, *args, **kwargs)
    if inspect.isawaitable(res):
        res = await res
    return res

# === END_CONTEXT_AWARE_FILE_INTAKE_V1_ROUTER_WRAPPER ===



# === FILE_INTAKE_SUPPORTED_FORMATS_V1 ===
SUPPORTED_IMAGE_FORMATS = {".jpg", ".jpeg", ".png", ".heic", ".heif", ".webp", ".bmp", ".tiff", ".tif", ".gif", ".svg"}
SUPPORTED_DOCUMENT_FORMATS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".txt", ".rtf"}
SUPPORTED_CAD_FORMATS = {".dwg", ".dxf", ".ifc", ".pln", ".rvt"}
SUPPORTED_AUDIO_FORMATS = {".ogg", ".mp3", ".wav", ".m4a", ".flac", ".aac"}
SUPPORTED_ARCHIVE_FORMATS = {".zip", ".rar", ".7z"}

def get_supported_file_format_v1(file_name: str) -> str:
    import os
    ext = os.path.splitext(str(file_name or "").lower())[1]
    if ext in SUPPORTED_IMAGE_FORMATS:
        return "image_ocr_vision"
    if ext in SUPPORTED_DOCUMENT_FORMATS:
        return "document"
    if ext in SUPPORTED_CAD_FORMATS:
        if ext == ".pln":
            return "pln_metadata_only"
        if ext == ".rvt":
            return "rvt_metadata_only"
        if ext == ".ifc":
            return "ifc_ifcopenshell"
        return "cad_ezdxf"
    if ext in SUPPORTED_AUDIO_FORMATS:
        return "audio_groq_stt"
    if ext in SUPPORTED_ARCHIVE_FORMATS:
        return "archive_recursive"
    return "unsupported"

def unsupported_format_message_v1(file_name: str) -> str:
    return "Формат файла не читается напрямую. Пришли PDF/DOCX/XLSX или экспортированный чертёж"

try:
    ESTIMATE_FILENAME_TRIGGERS = list(set(ESTIMATE_FILENAME_TRIGGERS + [
        "jpg", "jpeg", "png", "heic", "heif", "webp", "bmp", "tiff", "gif", "svg",
        "dwg", "dxf", "ifc", "pln", "rvt", "zip", "rar", "7z",
        "км", "кмд", "ов", "вк", "эо", "эм", "эос"
    ]))
except Exception:
    pass
# === END_FILE_INTAKE_SUPPORTED_FORMATS_V1 ===


# === CONTEXT_AWARE_FILE_INTAKE_V1_DB_LOOKUP ===
def _context_aware_file_intake_lookup_v1(chat_id: str = "", topic_id: int = 0) -> str:
    import sqlite3
    try:
        conn = sqlite3.connect("/root/.areal-neva-core/data/core.db")
        try:
            if chat_id:
                rows = conn.execute(
                    "SELECT raw_input, input_type, state, result FROM tasks WHERE chat_id=? AND COALESCE(topic_id,0)=? AND state IN ('DONE','AWAITING_CONFIRMATION','IN_PROGRESS','WAITING_CLARIFICATION') ORDER BY updated_at DESC LIMIT 5",
                    (str(chat_id), int(topic_id or 0)),
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT raw_input, input_type, state, result FROM tasks WHERE COALESCE(topic_id,0)=? AND state IN ('DONE','AWAITING_CONFIRMATION','IN_PROGRESS','WAITING_CLARIFICATION') ORDER BY updated_at DESC LIMIT 5",
                    (int(topic_id or 0),),
                ).fetchall()
        finally:
            conn.close()
    except Exception:
        rows = []

    for r in rows or []:
        combined = " ".join(str(x or "") for x in r).lower()
        if any(x in combined for x in ("смет", "estimate", "расцен", "стоимость", "цена")):
            return "estimate"
        if any(x in combined for x in ("проект", "кж", "кд", "ар", "эскиз", "км", "кмд", "ов", "вк", "эо")):
            return "project"
        if any(x in combined for x in ("акт", "технадзор", "дефект", "нарушение")):
            return "technadzor"
        if any(x in combined for x in ("образец", "шаблон", "эталон")):
            return "template"
    return ""

try:
    _context_aware_orig_route_file_v1 = route_file
    async def route_file(*args, **kwargs):
        import inspect
        lst = list(args)
        raw_input = str(kwargs.get("raw_input") or kwargs.get("caption") or kwargs.get("user_text") or "")
        chat_id = str(kwargs.get("chat_id") or "")
        topic_id = kwargs.get("topic_id", lst[2] if len(lst) >= 3 else 0)
        current_intent = kwargs.get("intent", lst[3] if len(lst) >= 4 else None)
        final_intent = current_intent

        if not final_intent or str(final_intent).lower() in ("unknown", "none", ""):
            if not raw_input.strip():
                final_intent = _context_aware_file_intake_lookup_v1(chat_id=chat_id, topic_id=int(topic_id or 0))

        if final_intent and final_intent != current_intent:
            if "intent" in kwargs:
                kwargs["intent"] = final_intent
            elif len(lst) >= 4:
                lst[3] = final_intent
            else:
                kwargs["intent"] = final_intent

        res = _context_aware_orig_route_file_v1(*lst, **kwargs)
        if inspect.isawaitable(res):
            res = await res
        return res
except Exception:
    pass
# === END_CONTEXT_AWARE_FILE_INTAKE_V1_DB_LOOKUP ===


# === P6D_FILE_INTAKE_IMAGE_ESTIMATE_KWARGS_CLOSE_20260504_V1 ===
try:
    _p6d_orig_route_file_20260504 = route_file
except Exception:
    _p6d_orig_route_file_20260504 = None

def _p6d_fi_s(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p6d_fi_low(v):
    return _p6d_fi_s(v).lower().replace("ё", "е")

def _p6d_fi_is_image(path, mime=""):
    low = _p6d_fi_low(str(path) + " " + str(mime))
    return low.startswith("image/") or any(x in low for x in (".jpg", ".jpeg", ".png", ".webp", ".heic", ".tif", ".tiff", ".bmp"))

def _p6d_fi_is_estimate_text(raw):
    low = _p6d_fi_low(raw)
    return any(x in low for x in (
        "смет", "стоимость", "расчет", "расчёт", "полная смета", "дом",
        "фундамент", "плита", "каркас", "кровля", "стены", "отделка", "санузел"
    ))

async def route_file(file_path, task_id, topic_id=0, intent=None, fmt="excel", *args, **kwargs):
    import inspect
    fp = _p6d_fi_s(file_path, 3000)
    raw_input = _p6d_fi_s(kwargs.get("raw_input") or kwargs.get("caption") or kwargs.get("user_text") or kwargs.get("prompt") or "", 12000)
    mime_type = _p6d_fi_s(kwargs.get("mime_type") or "", 500)
    final_intent = _p6d_fi_s(intent or kwargs.get("intent") or "", 100)

    if int(topic_id or 0) == 2 and _p6d_fi_is_image(fp, mime_type) and _p6d_fi_is_estimate_text(raw_input):
        try:
            from core import sample_template_engine as _ste
            fake_task = {"id": str(task_id), "raw_input": raw_input, "topic_id": int(topic_id or 0), "input_type": "drive_file"}
            res = _ste.handle_topic2_image_estimate_pipeline_p6d(
                conn=kwargs.get("conn"),
                task=fake_task,
                chat_id=kwargs.get("chat_id"),
                topic_id=int(topic_id or 0),
                raw_input=raw_input,
                local_path=fp,
                full_context=raw_input,
            )
            if inspect.isawaitable(res):
                res = await res
            if res:
                return {"success": True, "intent": "estimate", "engine": "P6D_IMAGE_ESTIMATE_FROM_PHOTO_FULL_CLOSE_20260504_V1", "text": "image estimate handled"}
        except Exception as e:
            return {"success": False, "error": "P6D_IMAGE_ESTIMATE_ROUTE_FAILED:" + str(e)[:500]}

    if _p6d_orig_route_file_20260504 is None:
        return {"success": False, "error": "P6D_ORIGINAL_ROUTE_FILE_MISSING"}

    try:
        res = _p6d_orig_route_file_20260504(fp, task_id, topic_id, final_intent or intent, fmt, *args, **kwargs)
    except TypeError:
        clean_kwargs = {k: v for k, v in kwargs.items() if k not in ("raw_input", "caption", "user_text", "prompt", "mime_type", "conn", "chat_id")}
        res = _p6d_orig_route_file_20260504(fp, task_id, topic_id, final_intent or intent, fmt, *args, **clean_kwargs)

    if inspect.isawaitable(res):
        res = await res
    return res
# === END_P6D_FILE_INTAKE_IMAGE_ESTIMATE_KWARGS_CLOSE_20260504_V1 ===

# === P6E2_FILE_INTAKE_ROUTE_FILE_KWARGS_AND_IMAGE_ESTIMATE_20260504_V1 ===
try:
    _P6E2_ORIG_ROUTE_FILE_20260504 = route_file
except Exception:
    _P6E2_ORIG_ROUTE_FILE_20260504 = None

def _p6e2_fi_s(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p6e2_fi_low(v):
    return _p6e2_fi_s(v).lower().replace("ё", "е")

def _p6e2_fi_is_image(path="", mime_type="", file_name=""):
    low = _p6e2_fi_low(" ".join([path or "", mime_type or "", file_name or ""]))
    return low.startswith("image/") or any(x in low for x in (".jpg", ".jpeg", ".png", ".webp", ".heic", ".tif", ".tiff", ".bmp"))

def _p6e2_fi_estimate_like(text):
    low = _p6e2_fi_low(text)
    return any(x in low for x in ("смет", "расчет", "расчёт", "посчитай", "стоимость", "полная смета"))

async def route_file(file_path, task_id, topic_id=0, intent=None, fmt="excel", *args, **kwargs):
    raw_input = _p6e2_fi_s(kwargs.get("raw_input") or kwargs.get("caption") or kwargs.get("user_text") or kwargs.get("prompt") or "", 100000)
    mime_type = _p6e2_fi_s(kwargs.get("mime_type") or "")
    file_name = _p6e2_fi_s(kwargs.get("file_name") or kwargs.get("name") or "")
    conn = kwargs.get("conn")
    chat_id = kwargs.get("chat_id")
    if int(topic_id or 0) == 2 and _p6e2_fi_is_image(file_path, mime_type, file_name) and _p6e2_fi_estimate_like(raw_input):
        try:
            from core.sample_template_engine import handle_topic2_image_estimate_p6e2
            if conn is not None:
                fake_task = {"id": str(task_id), "raw_input": raw_input, "input_type": "drive_file", "topic_id": int(topic_id or 0), "chat_id": chat_id}
                ok = await handle_topic2_image_estimate_p6e2(conn=conn, task=fake_task, chat_id=chat_id, topic_id=topic_id, raw_input=raw_input, local_path=file_path, file_name=file_name, mime_type=mime_type)
                if ok:
                    return {"success": True, "intent": "estimate", "result_text": "Смета по фото сформирована"}
        except Exception as e:
            return {"success": False, "error": "P6E2_IMAGE_ESTIMATE_ROUTE_FAILED:" + str(e)[:500]}
    if _P6E2_ORIG_ROUTE_FILE_20260504:
        clean = dict(kwargs)
        for k in ("raw_input", "caption", "user_text", "prompt", "mime_type", "conn", "chat_id", "file_name", "name"):
            clean.pop(k, None)
        return await _P6E2_ORIG_ROUTE_FILE_20260504(file_path, task_id, topic_id, intent, fmt, *args, **clean)
    return {"success": False, "error": "P6E2_ORIGINAL_ROUTE_FILE_MISSING"}
# === END_P6E2_FILE_INTAKE_ROUTE_FILE_KWARGS_AND_IMAGE_ESTIMATE_20260504_V1 ===
