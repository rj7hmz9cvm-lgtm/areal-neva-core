# === UNIVERSAL_FILE_HANDLER_V1 ===
import os, logging, tempfile, subprocess, csv, zipfile, json
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

# --- Magic bytes detection ---
_MAGIC = {
    b"%PDF": "pdf",
    b"PK\x03\x04": "xlsx_or_zip",
    b"\xd0\xcf\x11\xe0": "doc_or_xls",
    b"\xff\xd8\xff": "jpg",
    b"\x89PNG": "png",
    b"GIF8": "gif",
    b"BM": "bmp",
    b"II\x2a\x00": "tiff",
    b"MM\x00\x2a": "tiff",
    b"RIFF": "webp_or_avi",
    b"ftyp": "mp4",
    b"ID3": "mp3",
    b"AC10": "dwg",
    b"AC12": "dwg",
    b"AC14": "dwg",
    b"AC15": "dwg",
    b"AC18": "dwg",
    b"AC21": "dwg",
    b"AC24": "dwg",
    b"AC27": "dwg",
    b"  0\r\nSECTION": "dxf",
}

EXT_MAP = {
    ".pdf": "pdf", ".docx": "docx", ".doc": "doc_old",
    ".xlsx": "xlsx", ".xls": "xls_old", ".csv": "csv",
    ".txt": "text", ".md": "text", ".json": "json", ".xml": "xml",
    ".jpg": "image", ".jpeg": "image", ".png": "image",
    ".heic": "image", ".webp": "image", ".bmp": "image", ".tiff": "image",
    ".dwg": "dwg", ".dxf": "dxf", ".dgn": "dgn",
    ".zip": "zip", ".rar": "rar", ".7z": "7z",
    ".mp4": "video", ".avi": "video", ".mov": "video",
    ".mp3": "audio", ".ogg": "audio", ".wav": "audio",
    ".odt": "odt", ".ods": "ods", ".rtf": "rtf",
}

def detect_type(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    try:
        with open(file_path, "rb") as f:
            header = f.read(16)
        for magic, ftype in _MAGIC.items():
            if header[:len(magic)] == magic:
                # PK magic = ZIP or XLSX — уточняем по расширению
                if ftype == "xlsx_or_zip":
                    return "xlsx" if ext in (".xlsx", ".xlsm", ".xltx") else "zip"
                # RIFF = webp or avi — уточняем по расширению
                if ftype == "webp_or_avi":
                    return "image" if ext == ".webp" else "video"
                return ftype
    except Exception:
        pass
    return EXT_MAP.get(ext, "unknown")


def extract_text_from_file(file_path: str, task_id: str = "", topic_id: int = 0) -> Dict[str, Any]:
    """
    Универсальный экстрактор текста/данных из любого файла.
    Маркер: UNIVERSAL_FILE_HANDLER_V1
    Возвращает: {"success": bool, "type": str, "text": str, "rows": list, "error": str}
    """
    result = {"success": False, "type": "unknown", "text": "", "rows": [], "error": ""}
    ftype = detect_type(file_path)
    result["type"] = ftype
    logger.info("UNIVERSAL_FILE_HANDLER type=%s file=%s", ftype, os.path.basename(file_path))

    try:
        # --- PDF ---
        if ftype == "pdf":
            import pdfplumber, re
            with pdfplumber.open(file_path) as pdf:
                parts = []
                rows = []
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    t = re.sub(r'\(cid:\d+\)', '', t)
                    if t.strip():
                        parts.append(t)
                    for tbl in (page.extract_tables() or []):
                        rows.extend(tbl)
                result["text"] = "\n".join(parts)
                result["rows"] = rows
                result["success"] = True

        # --- DOCX / ODT ---
        elif ftype in ("docx", "odt"):
            import docx as _docx
            doc = _docx.Document(file_path)
            result["text"] = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            result["rows"] = [[c.text for c in row.cells] for tbl in doc.tables for row in tbl.rows]
            result["success"] = True

        # --- XLSX / ODS ---
        elif ftype in ("xlsx_or_zip", "xlsx"):
            from openpyxl import load_workbook
            wb = load_workbook(file_path, data_only=True)
            rows = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    if any(c is not None for c in row):
                        rows.append([str(c) if c is not None else "" for c in row])
            result["rows"] = rows
            result["text"] = "\n".join("\t".join(r) for r in rows[:50])
            result["success"] = True

        # --- CSV ---
        elif ftype == "csv":
            rows = []
            enc = "utf-8"
            try:
                import chardet
                with open(file_path, "rb") as f:
                    enc = chardet.detect(f.read(4096)).get("encoding", "utf-8") or "utf-8"
            except Exception:
                pass
            with open(file_path, encoding=enc, errors="replace") as f:
                for row in csv.reader(f):
                    rows.append(row)
            result["rows"] = rows
            result["text"] = "\n".join("\t".join(r) for r in rows[:50])
            result["success"] = True

        # --- TEXT / JSON / XML / MD ---
        elif ftype in ("text", "json", "xml", "rtf"):
            enc = "utf-8"
            try:
                import chardet
                with open(file_path, "rb") as f:
                    enc = chardet.detect(f.read(4096)).get("encoding", "utf-8") or "utf-8"
            except Exception:
                pass
            with open(file_path, encoding=enc, errors="replace") as f:
                result["text"] = f.read(50000)
            result["success"] = True

        # --- ИЗОБРАЖЕНИЯ (JPG/PNG/HEIC/BMP/TIFF/WEBP/GIF) ---
        elif ftype in ("jpg", "png", "image", "gif", "bmp", "tiff", "webp_or_avi"):
            import pytesseract
            from PIL import Image
            try:
                from pillow_heif import register_heif_opener
                register_heif_opener()
            except Exception:
                pass
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img, lang="rus+eng")
            result["text"] = text.strip()
            result["success"] = True
            result["type"] = "image"

        # --- DWG → конвертация в DXF → ezdxf ---
        elif ftype == "dwg":
            result = _handle_dwg(file_path, result)

        # --- DXF ---
        elif ftype == "dxf":
            result = _handle_dxf(file_path, result)

        # --- ZIP ---
        elif ftype == "zip":
            result = _handle_zip(file_path, task_id, topic_id, result)

        # --- RAR ---
        elif ftype == "rar":
            try:
                import rarfile
                tmp = tempfile.mkdtemp()
                with rarfile.RarFile(file_path) as rf:
                    rf.extractall(tmp)
                texts = []
                for fn in os.listdir(tmp)[:5]:
                    sub = extract_text_from_file(os.path.join(tmp, fn), task_id, topic_id)
                    if sub["success"]:
                        texts.append(f"[{fn}]\n{sub['text']}")
                result["text"] = "\n\n".join(texts)
                result["success"] = True
                result["type"] = "rar"
            except Exception as e:
                result["error"] = f"RAR: {e}"

        # --- 7Z ---
        elif ftype == "7z":
            try:
                import py7zr
                tmp = tempfile.mkdtemp()
                with py7zr.SevenZipFile(file_path) as sz:
                    sz.extractall(tmp)
                texts = []
                for fn in os.listdir(tmp)[:5]:
                    sub = extract_text_from_file(os.path.join(tmp, fn), task_id, topic_id)
                    if sub["success"]:
                        texts.append(f"[{fn}]\n{sub['text']}")
                result["text"] = "\n\n".join(texts)
                result["success"] = True
                result["type"] = "7z"
            except Exception as e:
                result["error"] = f"7Z: {e}"

        # --- ВИДЕО/АУДИО — метаданные через ffmpeg ---
        elif ftype in ("mp4", "video", "mp3", "audio"):
            try:
                out = subprocess.check_output(
                    ["ffmpeg", "-i", file_path],
                    stderr=subprocess.STDOUT, timeout=10
                ).decode(errors="replace")
            except subprocess.CalledProcessError as e:
                out = e.output.decode(errors="replace")
            result["text"] = out[:2000]
            result["success"] = True

        # --- UNKNOWN — попытка открыть как текст ---
        else:
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    txt = f.read(10000)
                if len(txt.strip()) > 20:
                    result["text"] = txt
                    result["success"] = True
                    result["type"] = "text_fallback"
                else:
                    result["error"] = f"Формат не поддерживается: {os.path.splitext(file_path)[1]}"
            except Exception as e:
                result["error"] = f"Неизвестный формат: {e}"

    except Exception as e:
        logger.error("UNIVERSAL_FILE_HANDLER_ERROR type=%s err=%s", ftype, e)
        result["error"] = str(e)

    return result


def _handle_dwg(file_path: str, result: dict) -> dict:
    """DWG: конвертация через dwg2dxf (libredwg), fallback через imagemagick preview"""
    dxf_path = file_path.replace(".dwg", ".dxf").replace(".DWG", ".dxf")
    if not dxf_path.endswith(".dxf"):
        dxf_path = file_path + ".dxf"

    # Попытка 1: dwg2dxf
    try:
        subprocess.run(["dwg2dxf", file_path, "-o", dxf_path],
                       timeout=30, capture_output=True, check=True)
        if os.path.exists(dxf_path):
            logger.info("DWG→DXF conversion OK: %s", dxf_path)
            return _handle_dxf(dxf_path, result)
    except Exception as e:
        logger.warning("dwg2dxf failed: %s", e)

    # Попытка 2: imagemagick — превью в PNG + OCR
    try:
        png_path = file_path + "_preview.png"
        subprocess.run(
            ["convert", "-density", "150", file_path + "[0]", png_path],
            timeout=30, capture_output=True, check=True
        )
        if os.path.exists(png_path):
            import pytesseract
            from PIL import Image
            text = pytesseract.image_to_string(Image.open(png_path), lang="rus+eng")
            result["text"] = f"[DWG файл — превью через OCR]\n{text.strip()}"
            result["success"] = True
            result["type"] = "dwg_ocr_preview"
            return result
    except Exception as e:
        logger.warning("DWG imagemagick fallback failed: %s", e)

    result["error"] = "DWG: конвертация не удалась. Пришли файл в формате .dxf"
    result["text"] = "Файл формата DWG получен. Для полной обработки конвертируй в DXF."
    result["success"] = False
    return result


def _handle_dxf(file_path: str, result: dict) -> dict:
    """DXF через ezdxf"""
    try:
        import ezdxf
        doc = ezdxf.readfile(file_path)
        msp = doc.modelspace()
        counts = {}
        texts = []
        for e in msp:
            t = e.dxftype()
            counts[t] = counts.get(t, 0) + 1
            if t in ("TEXT", "MTEXT") and hasattr(e.dxf, "text"):
                txt = str(e.dxf.text or "").strip()
                if txt:
                    texts.append(txt)
        summary = "DXF элементы:\n"
        for k, v in sorted(counts.items(), key=lambda x: -x[1])[:15]:
            summary += f"  {k}: {v}\n"
        if texts:
            summary += "\nТексты в чертеже:\n" + "\n".join(texts[:30])
        result["text"] = summary
        result["rows"] = [[k, str(v)] for k, v in counts.items()]
        result["success"] = True
        result["type"] = "dxf"
    except Exception as e:
        result["error"] = f"DXF: {e}"
    return result


def _handle_zip(file_path: str, task_id: str, topic_id: int, result: dict) -> dict:
    """ZIP — распаковка и рекурсивная обработка"""
    try:
        tmp = tempfile.mkdtemp()
        with zipfile.ZipFile(file_path) as zf:
            names = zf.namelist()[:20]
            zf.extractall(tmp)
        texts = []
        all_rows = []
        for fn in names:
            fp = os.path.join(tmp, fn)
            if not os.path.isfile(fp):
                continue
            sub = extract_text_from_file(fp, task_id, topic_id)
            if sub["success"]:
                texts.append(f"[{fn}]\n{sub['text'][:1000]}")
                all_rows.extend(sub.get("rows", []))
        result["text"] = f"ZIP архив ({len(names)} файлов):\n\n" + "\n\n".join(texts)
        result["rows"] = all_rows
        result["success"] = True
        result["type"] = "zip"
    except Exception as e:
        result["error"] = f"ZIP: {e}"
    return result
# === END UNIVERSAL_FILE_HANDLER_V1 ===
