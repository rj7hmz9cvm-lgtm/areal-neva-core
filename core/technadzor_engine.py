# === FINAL_CLOSURE_BLOCKER_FIX_V1_TECHNADZOR_ENGINE ===
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict

BASE = Path("/root/.areal-neva-core")
OUT = BASE / "outputs" / "technadzor"
OUT.mkdir(parents=True, exist_ok=True)


def is_technadzor_intent(text: str = "", file_name: str = "") -> bool:
    t = f"{text} {file_name}".lower().replace("ё", "е")
    return bool(re.search(r"\b(акт|технадзор|техническ.*надзор|дефект|замечан|нарушен|освидетельств|стройконтроль|сп|гост|снип)\b", t))


def _norm_refs(text: str) -> str:
    refs = []
    for m in re.findall(r"\b(сп\s*\d+[.\d]*|гост\s*\d+[.\d-]*|снип\s*[\w.\-]+)\b", text or "", flags=re.I):
        refs.append(m.upper().replace("  ", " "))
    return ", ".join(sorted(set(refs))) if refs else "Норма не подтверждена"


def process_technadzor(text: str = "", task_id: str = "", chat_id: str = "", topic_id: int = 0, file_path: str = "", file_name: str = "") -> Dict[str, Any]:
    if not is_technadzor_intent(text, file_name):
        return {"ok": False, "handled": False, "reason": "NOT_TECHNADZOR"}

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = f"TECHNADZOR_ACT__{task_id[:8] or ts}"
    txt_path = OUT / f"{stem}.txt"

    body = [
        "АКТ ТЕХНИЧЕСКОГО НАДЗОРА",
        "",
        f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Задача: {task_id}",
        f"Топик: {topic_id}",
    ]

    if file_name:
        body.append(f"Файл: {file_name}")

    body.extend(
        [
            "",
            "Исходное описание:",
            (text or "").strip() or "UNKNOWN",
            "",
            "Нормативная база:",
            _norm_refs(text),
            "",
            "Вывод:",
            "Черновик акта создан. Если норматив не подтверждён источником, в акте указано: Норма не подтверждена",
        ]
    )

    txt_path.write_text("\n".join(body) + "\n", encoding="utf-8")

    return {
        "ok": True,
        "handled": True,
        "kind": "technadzor_act",
        "state": "DONE",
        "artifact_path": str(txt_path),
        "message": "Технадзорный акт подготовлен",  # TECHNADZOR_PUBLIC_MESSAGE_NO_LOCAL_PATH_V1
        "history": "FINAL_CLOSURE_BLOCKER_FIX_V1:TECHNADZOR_ACT_CREATED",
    }


# === END_FINAL_CLOSURE_BLOCKER_FIX_V1_TECHNADZOR_ENGINE ===


# === P6_TECHNADZOR_TEMPLATE_AND_ARTIFACT_CLOSE_20260504_V1 ===
# Scope:
# - technadzor sample/template files can be saved as active reference per chat/topic
# - future technadzor acts use active reference metadata
# - produces TXT and DOCX when python-docx exists; no DB schema changes

import json as _p6tz_json
import re as _p6tz_re
from datetime import datetime as _p6tz_datetime
from pathlib import Path as _p6tz_Path

_P6TZ_BASE = _p6tz_Path("/root/.areal-neva-core")
_P6TZ_TEMPLATE_DIR = _P6TZ_BASE / "data/templates/technadzor"
_P6TZ_OUT = _P6TZ_BASE / "outputs/technadzor"
_P6TZ_TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
_P6TZ_OUT.mkdir(parents=True, exist_ok=True)

def _p6tz_s(v, limit=12000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p6tz_low(v):
    return _p6tz_s(v).lower().replace("ё", "е")

def _p6tz_template_path(chat_id, topic_id):
    safe_chat = _p6tz_re.sub(r"[^0-9a-zA-Z_-]+", "_", str(chat_id or "unknown"))
    return _P6TZ_TEMPLATE_DIR / f"ACTIVE__chat_{safe_chat}__topic_{int(topic_id or 0)}.json"

def _p6tz_is_template_intent(text="", file_name=""):
    low = _p6tz_low(str(text) + " " + str(file_name))
    return any(x in low for x in ("образец", "шаблон", "пример", "как образец", "как шаблон", "возьми его как образец", "сохрани как образец")) and any(x in low for x in ("технадзор", "акт", "замечан", "дефект", "строительный контроль", "стройконтроль"))

def _p6tz_save_template(text="", task_id="", chat_id="", topic_id=0, file_path="", file_name=""):
    meta = {
        "engine": "P6_TECHNADZOR_TEMPLATE_AND_ARTIFACT_CLOSE_20260504_V1",
        "kind": "technadzor_template",
        "status": "active",
        "chat_id": str(chat_id or ""),
        "topic_id": int(topic_id or 0),
        "source_task_id": str(task_id or ""),
        "source_file_path": str(file_path or ""),
        "source_file_name": str(file_name or ""),
        "raw_user_instruction": _p6tz_s(text, 4000),
        "usage_rule": "Use this file as formatting/sample reference for future technadzor acts in same chat/topic",
        "saved_at": _p6tz_datetime.now().isoformat(),
    }
    path = _p6tz_template_path(chat_id, topic_id)
    path.write_text(_p6tz_json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return path

def _p6tz_load_template(chat_id, topic_id):
    path = _p6tz_template_path(chat_id, topic_id)
    if not path.exists():
        return {}
    try:
        return _p6tz_json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}

def _p6tz_refs(text):
    refs = []
    for m in _p6tz_re.findall(r"\b(сп\s*\d+[.\d]*|гост\s*\d+[.\d-]*|снип\s*[\w.\-]+)\b", text or "", flags=_p6tz_re.I):
        refs.append(m.upper().replace("  ", " "))
    return ", ".join(sorted(set(refs))) if refs else "Норма не подтверждена"

def _p6tz_make_docx(path, lines):
    try:
        from docx import Document
        doc = Document()
        for i, line in enumerate(lines):
            if i == 0:
                doc.add_heading(line, level=1)
            elif line == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)
        doc.save(str(path))
        return str(path)
    except Exception:
        return ""

try:
    _p6tz_orig_is_intent = is_technadzor_intent
    def is_technadzor_intent(text: str = "", file_name: str = "") -> bool:
        low = _p6tz_low(str(text) + " " + str(file_name))
        if _p6tz_is_template_intent(text, file_name):
            return True
        if any(x in low for x in ("технадзор", "акт", "замечан", "дефект", "нарушен", "освидетельств", "стройконтроль", "строительный контроль", "сп ", "гост", "снип")):
            return True
        return _p6tz_orig_is_intent(text, file_name)
except Exception:
    pass

try:
    _p6tz_orig_process = process_technadzor
    def process_technadzor(text: str = "", task_id: str = "", chat_id: str = "", topic_id: int = 0, file_path: str = "", file_name: str = ""):
        if _p6tz_is_template_intent(text, file_name):
            meta_path = _p6tz_save_template(text=text, task_id=task_id, chat_id=chat_id, topic_id=topic_id, file_path=file_path, file_name=file_name)
            return {
                "ok": True,
                "handled": True,
                "kind": "technadzor_template_saved",
                "state": "DONE",
                "artifact_path": str(meta_path),
                "message": "Образец технадзора сохранён для этого топика",
                "history": "P6_TECHNADZOR_TEMPLATE_SAVED",
            }

        if not is_technadzor_intent(text, file_name):
            return {"ok": False, "handled": False, "reason": "NOT_TECHNADZOR"}

        tpl = _p6tz_load_template(chat_id, topic_id)
        ts = _p6tz_datetime.now().strftime("%Y%m%d_%H%M%S")
        safe = str(task_id or ts)[:8] or ts
        stem = f"TECHNADZOR_ACT__{safe}"
        txt_path = _P6TZ_OUT / f"{stem}.txt"
        docx_path = _P6TZ_OUT / f"{stem}.docx"

        lines = [
            "АКТ ТЕХНИЧЕСКОГО НАДЗОРА",
            "",
            f"Дата: {_p6tz_datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Задача: {task_id}",
            f"Топик: {topic_id}",
        ]
        if file_name:
            lines.append(f"Файл: {file_name}")
        if tpl:
            lines.append(f"Образец: {tpl.get('source_file_name') or tpl.get('source_file_path') or 'активный шаблон топика'}")
        lines += [
            "",
            "Исходное описание:",
            _p6tz_s(text, 6000) or "UNKNOWN",
            "",
            "Нормативная база:",
            _p6tz_refs(text),
            "",
            "Выявленные замечания:",
            "1. Требуется заполнение по присланным фото/файлам и описанию",
            "",
            "Требуемые действия:",
            "1. Устранить замечания",
            "2. Предоставить фотофиксацию устранения",
            "3. Повторно предъявить участок работ техническому надзору",
            "",
            "Статус:",
            "Черновик подготовлен по текущим данным",
        ]

        txt_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        docx_created = _p6tz_make_docx(docx_path, lines)
        return {
            "ok": True,
            "handled": True,
            "kind": "technadzor_act",
            "state": "DONE",
            "artifact_path": docx_created or str(txt_path),
            "extra_artifact_path": str(txt_path),
            "message": "Технадзорный акт подготовлен",
            "history": "P6_TECHNADZOR_ACT_CREATED",
        }
except Exception:
    pass

# === END_P6_TECHNADZOR_TEMPLATE_AND_ARTIFACT_CLOSE_20260504_V1 ===

# === P6C_TECHNADZOR_CONN_COMPAT_TEMPLATE_ACT_CLOSE_20260504_V1 ===
import json as _p6c_te_json
import re as _p6c_te_re
from pathlib import Path as _p6c_te_Path
from datetime import datetime as _p6c_te_datetime

_P6C_TE_BASE = _p6c_te_Path("/root/.areal-neva-core")
_P6C_TE_OUT = _P6C_TE_BASE / "outputs" / "technadzor"
_P6C_TE_TPL = _P6C_TE_BASE / "data" / "templates" / "technadzor"
_P6C_TE_OUT.mkdir(parents=True, exist_ok=True)
_P6C_TE_TPL.mkdir(parents=True, exist_ok=True)

def _p6c_te_s(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p6c_te_low(v):
    return _p6c_te_s(v).lower().replace("ё", "е")

def is_technadzor_intent(text: str = "", file_name: str = "") -> bool:
    low = _p6c_te_low(f"{text} {file_name}")
    return any(x in low for x in (
        "технадзор", "акт", "осмотр", "выезд", "дефект", "замечан",
        "образец написания", "образец акта", "как образец"
    ))

def _p6c_te_is_template(text, file_name):
    low = _p6c_te_low(f"{text} {file_name}")
    return any(x in low for x in (
        "образец", "как образец", "прими это как факт", "образец написания",
        "возьми это как образец", "шаблон"
    ))

def _p6c_te_read_text(file_path):
    fp = _p6c_te_Path(_p6c_te_s(file_path))
    if not fp.exists():
        return ""
    suf = fp.suffix.lower()
    try:
        if suf == ".txt":
            return fp.read_text(encoding="utf-8", errors="ignore")[:30000]
        if suf == ".pdf":
            try:
                import fitz
                doc = fitz.open(str(fp))
                return "\n".join(page.get_text("text") for page in doc)[:30000]
            except Exception:
                return ""
        if suf == ".docx":
            try:
                import docx
                d = docx.Document(str(fp))
                return "\n".join(p.text for p in d.paragraphs)[:30000]
            except Exception:
                return ""
    except Exception:
        return ""
    return ""

def _p6c_te_save_template(chat_id, topic_id, task_id, text, file_name, file_path):
    data = {
        "engine": "P6C_TECHNADZOR_CONN_COMPAT_TEMPLATE_ACT_CLOSE_20260504_V1",
        "chat_id": str(chat_id or ""),
        "topic_id": int(topic_id or 0),
        "task_id": str(task_id or ""),
        "file_name": _p6c_te_s(file_name),
        "file_path": _p6c_te_s(file_path),
        "saved_at": _p6c_te_datetime.utcnow().isoformat() + "Z",
        "template_text": _p6c_te_s(text, 25000),
    }
    out = _P6C_TE_TPL / f"ACTIVE__chat_{chat_id}__topic_{int(topic_id or 0)}.json"
    out.write_text(_p6c_te_json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(out)

def _p6c_te_load_template(chat_id, topic_id):
    p = _P6C_TE_TPL / f"ACTIVE__chat_{chat_id}__topic_{int(topic_id or 0)}.json"
    if not p.exists():
        return {}
    try:
        return _p6c_te_json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}

def _p6c_te_write_act(task_id, chat_id, topic_id, body):
    out = _P6C_TE_OUT / f"TECHNADZOR_ACT__{str(task_id)[:8]}.txt"
    out.write_text(body, encoding="utf-8")
    return str(out)

def process_technadzor(text: str = "", task_id: str = "", chat_id: str = "", topic_id: int = 0, file_path: str = "", file_name: str = "", **kwargs):
    conn = kwargs.get("conn")
    raw_text = _p6c_te_s(text, 50000)
    file_name = _p6c_te_s(file_name or kwargs.get("name") or "")
    file_path = _p6c_te_s(file_path or kwargs.get("local_path") or "")
    task_id = _p6c_te_s(task_id or kwargs.get("id") or kwargs.get("task_id") or "technadzor")
    chat_id = _p6c_te_s(chat_id or kwargs.get("chat_id") or "")
    try:
        topic_id = int(topic_id or kwargs.get("topic_id") or 0)
    except Exception:
        topic_id = 0

    extracted = _p6c_te_read_text(file_path)
    combined = "\n".join(x for x in [raw_text, extracted] if x).strip()

    if _p6c_te_is_template(raw_text, file_name):
        tpl_path = _p6c_te_save_template(chat_id, topic_id, task_id, combined, file_name, file_path)
        return {
            "ok": True,
            "status": "DONE",
            "result_text": f"Образец технадзора принят и сохранён\nФайл: {file_name}\nШаблон: активен для topic_{topic_id}",
            "artifact_path": tpl_path,
            "history": "P6C_TECHNADZOR_TEMPLATE_SAVED",
        }

    tpl = _p6c_te_load_template(chat_id, topic_id)
    tpl_note = "Использован сохранённый образец" if tpl else "Сохранённый образец не найден"

    body = "\n".join([
        "АКТ ОСМОТРА ПО ФАКТУ ВЫЕЗДА",
        "",
        f"Дата формирования: {_p6c_te_datetime.now().strftime('%d.%m.%Y %H:%M')}",
        f"Файл: {file_name or 'без файла'}",
        f"Источник: topic_{topic_id}",
        f"Шаблон: {tpl_note}",
        "",
        "Исходные данные:",
        combined[:12000] if combined else "Данные из файла не извлечены автоматически",
        "",
        "Вывод технического надзора:",
        "Документ сформирован как рабочий черновик акта. Требуется проверка владельцем перед выдачей заказчику.",
    ])
    artifact = _p6c_te_write_act(task_id, chat_id, topic_id, body)

    return {
        "ok": True,
        "status": "DONE",
        "result_text": f"Акт технадзора сформирован\nФайл: {file_name or 'без файла'}\nАртефакт: {artifact}",
        "artifact_path": artifact,
        "history": "P6C_TECHNADZOR_ACT_CREATED",
    }
# === END_P6C_TECHNADZOR_CONN_COMPAT_TEMPLATE_ACT_CLOSE_20260504_V1 ===

# === P6E2_TECHNADZOR_FOLDER_AWARE_NO_DRIVE_POLLUTION_20260504_V1 ===
import json as _p6e2_te_json
import re as _p6e2_te_re
from pathlib import Path as _p6e2_te_Path
from datetime import datetime as _p6e2_te_datetime

_P6E2_TE_BASE = _p6e2_te_Path("/root/.areal-neva-core")
_P6E2_TE_OUT = _P6E2_TE_BASE / "outputs" / "technadzor"
_P6E2_TE_TPL = _P6E2_TE_BASE / "data" / "templates" / "technadzor"
_P6E2_TE_OUT.mkdir(parents=True, exist_ok=True)
_P6E2_TE_TPL.mkdir(parents=True, exist_ok=True)

def _p6e2_te_s(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _p6e2_te_low(v):
    return _p6e2_te_s(v).lower().replace("ё", "е")

def _p6e2_te_is_template(text, file_name=""):
    low = _p6e2_te_low(text + " " + file_name)
    return any(x in low for x in ("как образец", "образец написания", "прими это как факт", "возьми это как образец", "шаблон акта"))

def _p6e2_te_extract_links(text):
    return _p6e2_te_re.findall(r"https?://\S+", _p6e2_te_s(text, 50000))

def _p6e2_te_read_file(path):
    p = _p6e2_te_Path(_p6e2_te_s(path))
    if not p.exists():
        return ""
    try:
        if p.suffix.lower() == ".txt":
            return p.read_text(encoding="utf-8", errors="ignore")[:50000]
        if p.suffix.lower() == ".docx":
            import docx
            d = docx.Document(str(p))
            return "\n".join(x.text for x in d.paragraphs)[:50000]
        if p.suffix.lower() == ".pdf":
            try:
                import fitz
                doc = fitz.open(str(p))
                return "\n".join(page.get_text("text") for page in doc)[:50000]
            except Exception:
                return ""
    except Exception:
        return ""
    return ""

def _p6e2_te_tpl_path(chat_id, topic_id):
    return _P6E2_TE_TPL / f"ACTIVE__chat_{chat_id}__topic_{int(topic_id or 0)}.json"

def _p6e2_te_save_template(chat_id, topic_id, task_id, body, file_name, file_path, links):
    data = {
        "engine": "P6E2_TECHNADZOR_FOLDER_AWARE_NO_DRIVE_POLLUTION_20260504_V1",
        "chat_id": str(chat_id),
        "topic_id": int(topic_id or 0),
        "task_id": str(task_id),
        "file_name": _p6e2_te_s(file_name),
        "file_path": _p6e2_te_s(file_path),
        "drive_links_seen": links,
        "saved_at": _p6e2_te_datetime.utcnow().isoformat() + "Z",
        "template_text": _p6e2_te_s(body, 30000),
    }
    p = _p6e2_te_tpl_path(chat_id, topic_id)
    p.write_text(_p6e2_te_json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(p)

def _p6e2_te_load_template(chat_id, topic_id):
    p = _p6e2_te_tpl_path(chat_id, topic_id)
    if not p.exists():
        return {}
    try:
        return _p6e2_te_json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return {}

def is_technadzor_intent(text: str = "", file_name: str = "") -> bool:
    low = _p6e2_te_low(text + " " + file_name)
    return any(x in low for x in ("технадзор", "акт", "осмотр", "выезд", "дефект", "замечан", "нарушен", "образец написания", "как образец", "прими это как факт"))

def process_technadzor(text: str = "", task_id: str = "", chat_id: str = "", topic_id: int = 0, file_path: str = "", file_name: str = "", **kwargs):
    raw = _p6e2_te_s(text, 50000)
    extracted = _p6e2_te_read_file(file_path)
    combined = "\n".join(x for x in (raw, extracted) if x).strip()
    links = _p6e2_te_extract_links(combined)
    task_id = _p6e2_te_s(task_id or kwargs.get("task_id") or kwargs.get("id") or "technadzor")
    chat_id = _p6e2_te_s(chat_id or kwargs.get("chat_id") or "")
    try:
        topic_id = int(topic_id or kwargs.get("topic_id") or 0)
    except Exception:
        topic_id = 0
    file_name = _p6e2_te_s(file_name or kwargs.get("file_name") or kwargs.get("name") or "")
    file_path = _p6e2_te_s(file_path or kwargs.get("local_path") or "")

    if _p6e2_te_is_template(raw, file_name):
        tpl = _p6e2_te_save_template(chat_id, topic_id, task_id, combined, file_name, file_path, links)
        return {
            "ok": True,
            "handled": True,
            "status": "DONE",
            "state": "DONE",
            "result_text": f"Образец технадзора принят и сохранён\nФайл: {file_name or 'без файла'}\nШаблон: active topic_{topic_id}\nDrive-ссылки учтены: {len(links)}",
            "message": "Образец технадзора принят и сохранён",
            "artifact_path": tpl,
            "history": "P6E2_TECHNADZOR_TEMPLATE_SAVED",
        }

    if not is_technadzor_intent(combined, file_name):
        return {"ok": False, "handled": False, "reason": "NOT_TECHNADZOR"}

    tpl = _p6e2_te_load_template(chat_id, topic_id)
    tpl_note = "использован сохранённый образец" if tpl else "сохранённый образец не найден"
    out = _P6E2_TE_OUT / f"TECHNADZOR_ACT__{str(task_id)[:8]}.txt"
    body = "\n".join([
        "АКТ ОСМОТРА ПО ФАКТУ ВЫЕЗДА",
        "",
        f"Дата формирования: {_p6e2_te_datetime.now().strftime('%d.%m.%Y %H:%M')}",
        f"Файл: {file_name or 'без файла'}",
        f"Источник: topic_{topic_id}",
        f"Шаблон: {tpl_note}",
        f"Drive-ссылки в исходных данных: {len(links)}",
        "",
        "Исходные данные:",
        combined[:15000] if combined else "UNKNOWN",
        "",
        "Вывод технического надзора:",
        "Документ сформирован по текущим данным и сохранён локально без записи мусора в Google Drive",
    ])
    out.write_text(body + "\n", encoding="utf-8")
    return {
        "ok": True,
        "handled": True,
        "status": "DONE",
        "state": "DONE",
        "result_text": f"Акт технадзора сформирован\nФайл: {file_name or 'без файла'}\nАртефакт: {out}",
        "message": "Акт технадзора сформирован",
        "artifact_path": str(out),
        "history": "P6E2_TECHNADZOR_ACT_CREATED",
    }
# === END_P6E2_TECHNADZOR_FOLDER_AWARE_NO_DRIVE_POLLUTION_20260504_V1 ===


# === P6F_TECHNADZOR_CLEAN_OUTPUT_AND_NORM_GATE_V1 ===
# FACT: wraps process_technadzor to ensure clean user-facing output.
# Uses core.normative_engine for confirmed references; if no confidence
# >= PARTIAL — explicitly states "норма не подтверждена" (per canon).
# Forbids JSON output to user.
import json as _p6f_tnz_json
import logging as _p6f_tnz_logging

_P6F_TNZ_LOG = _p6f_tnz_logging.getLogger("technadzor_engine")

def _p6f_tnz_clean_for_user(text):
    if not text:
        return ""
    s = str(text).strip()
    if (s.startswith("{") and s.rstrip().endswith("}")) or (s.startswith("[") and s.rstrip().endswith("]")):
        try:
            obj = _p6f_tnz_json.loads(s)
            if isinstance(obj, dict):
                summary = str(obj.get("summary") or obj.get("message") or "").strip()
                if summary:
                    return summary
                lines = []
                for k in ("kind", "state", "artifact_path", "history"):
                    if k in obj:
                        v = obj[k]
                        if k == "artifact_path":
                            lines.append("Артефакт: создан")
                        else:
                            lines.append(str(k) + ": " + str(v))
                if lines:
                    return "\n".join(lines)
                return "Технадзорный результат подготовлен"
        except Exception:
            pass
    return s

def _p6f_tnz_norm_block(text):
    try:
        from core.normative_engine import search_norms_sync, format_norms_for_act
    except Exception:
        return ""
    try:
        norms = search_norms_sync(text or "", limit=3)
    except Exception:
        return ""
    if not norms:
        return "Норма не подтверждена. Акт оформлен без ссылки на конкретный пункт СП/ГОСТ"
    confirmed_or_partial = [n for n in norms if str(n.get("confidence", "")).upper() in ("CONFIRMED", "PARTIAL")]
    if not confirmed_or_partial:
        return "Норма не подтверждена. Источник не найден"
    try:
        return format_norms_for_act(confirmed_or_partial)
    except Exception:
        out = []
        for n in confirmed_or_partial:
            nid = str(n.get("norm_id", "")).strip()
            sec = str(n.get("section", "")).strip()
            req = str(n.get("requirement", "")).strip()
            conf = str(n.get("confidence", "")).strip()
            if nid:
                out.append(f"- {nid} | {sec} | {req} | confidence={conf}")
        return "\n".join(out)

try:
    _P6F_TNZ_ORIG_PROCESS = process_technadzor
    if not getattr(_P6F_TNZ_ORIG_PROCESS, "_p6f_tnz_wrapped", False):
        def process_technadzor(text="", task_id="", chat_id="", topic_id=0, file_path="", file_name="", **kwargs):
            try:
                res = _P6F_TNZ_ORIG_PROCESS(text=text, task_id=task_id, chat_id=chat_id, topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs)
            except TypeError:
                res = _P6F_TNZ_ORIG_PROCESS(text=text, task_id=task_id, chat_id=chat_id, topic_id=topic_id, file_path=file_path, file_name=file_name)
            try:
                if isinstance(res, dict) and res.get("ok") and res.get("handled"):
                    raw_msg = res.get("message") or ""
                    clean = _p6f_tnz_clean_for_user(raw_msg)
                    norm_block = _p6f_tnz_norm_block(text)
                    artifact_line = ""
                    if res.get("artifact_path"):
                        artifact_line = "Артефакт акта: создан"
                    parts = [clean]
                    if norm_block:
                        parts.append("\nНормативная база:\n" + norm_block)
                    if artifact_line:
                        parts.append("\n" + artifact_line)
                    res["message"] = "\n".join([p for p in parts if p]).strip()
                    res["history"] = (res.get("history") or "") + ";P6F_TNZ_CLEANED_OUTPUT"
            except Exception as _e:
                try:
                    _P6F_TNZ_LOG.warning("P6F_TNZ_WRAP_ERR %s", _e)
                except Exception:
                    pass
            return res
        process_technadzor._p6f_tnz_wrapped = True
        _P6F_TNZ_LOG.info("P6F_TECHNADZOR_CLEAN_OUTPUT_AND_NORM_GATE_INSTALLED")
except Exception as _e:
    try:
        _P6F_TNZ_LOG.exception("P6F_TNZ_INSTALL_ERR %s", _e)
    except Exception:
        pass
# === END_P6F_TECHNADZOR_CLEAN_OUTPUT_AND_NORM_GATE_V1 ===


# === P6F_TECHNADZOR_PHOTO_TO_DOCX_REAL_V1 ===
# FACT: real photo defect → Vision (OpenRouter) → DOCX акт →
# Drive upload via topic_drive_oauth.upload_file_to_topic.
# No direct Google API. Norms only via core.normative_engine
# (no invented references).
import os as _p6f_tnz_os
import base64 as _p6f_tnz_base64
import asyncio as _p6f_tnz_asyncio
import json as _p6f_tnz_json2
import logging as _p6f_tnz_logging2

_P6F_TNZ_REAL_LOG = _p6f_tnz_logging2.getLogger("technadzor_engine")

_P6F_TNZ_REAL_PROMPT = (
    "Ты эксперт технического надзора в строительстве. На фото фиксация состояния "
    "конструкций или дефекта. Верни СТРОГО JSON без пояснений со схемой:\n"
    "{\n"
    "  \"summary\": \"кратко что видно\",\n"
    "  \"defects\": [{\n"
    "    \"title\": \"короткое название\",\n"
    "    \"location\": \"место\",\n"
    "    \"severity\": \"low|medium|high|critical\",\n"
    "    \"description\": \"что не так\",\n"
    "    \"recommendation\": \"что делать\"\n"
    "  }],\n"
    "  \"confidence\": \"HIGH|MEDIUM|LOW\"\n"
    "}\n"
    "Если на фото нет дефектов — defects=[]. "
    "Не выдумывай нормы СП/ГОСТ — это не твоя задача."
)

def _p6f_tnz_is_image_path(path):
    if not path:
        return False
    p = str(path).lower()
    return any(p.endswith(e) for e in (".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif", ".bmp"))

async def _p6f_tnz_vision_via_openrouter(local_path):
    if not _p6f_tnz_is_image_path(local_path) or not _p6f_tnz_os.path.exists(str(local_path)):
        return None, "PATH_MISSING"
    api_key = (_p6f_tnz_os.getenv("OPENROUTER_API_KEY") or "").strip()
    if not api_key:
        return None, "NO_OPENROUTER_KEY"
    base_url = (_p6f_tnz_os.getenv("OPENROUTER_BASE_URL") or "https://openrouter.ai/api/v1").strip().rstrip("/")
    model = (_p6f_tnz_os.getenv("OPENROUTER_VISION_MODEL") or "google/gemini-2.5-flash").strip()
    ext = _p6f_tnz_os.path.splitext(str(local_path))[1].lower().lstrip(".") or "jpeg"
    mime = "image/jpeg" if ext in ("jpg", "jpeg") else "image/{}".format(ext)
    try:
        with open(str(local_path), "rb") as f:
            b64 = _p6f_tnz_base64.b64encode(f.read()).decode("utf-8")
    except Exception as e:
        return None, "READ_ERR:{}".format(e)

    body = {
        "model": model,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "text", "text": _P6F_TNZ_REAL_PROMPT},
                {"type": "image_url", "image_url": {"url": "data:" + mime + ";base64," + b64}},
            ],
        }],
        "temperature": 0.1,
    }
    headers = {"Authorization": "Bearer " + api_key, "Content-Type": "application/json"}
    try:
        import httpx
        async with httpx.AsyncClient(timeout=httpx.Timeout(180.0, connect=30.0)) as client:
            r = await client.post(base_url + "/chat/completions", headers=headers, json=body)
            r.raise_for_status()
            data = r.json()
        content = data["choices"][0]["message"]["content"]
        if isinstance(content, list):
            content = "\n".join(x.get("text", "") if isinstance(x, dict) else str(x) for x in content)
    except Exception as e:
        return None, "OPENROUTER_CALL_ERR:{}".format(type(e).__name__)

    s = str(content).strip()
    if s.startswith("```"):
        import re as _re
        s = _re.sub(r"^```(?:json)?\s*", "", s)
        s = _re.sub(r"\s*```\s*$", "", s)
    try:
        return _p6f_tnz_json2.loads(s), "OK"
    except Exception:
        return {"summary": s[:2000], "defects": [], "confidence": "LOW"}, "PARTIAL"

def _p6f_tnz_norms_block(text_for_search):
    try:
        from core.normative_engine import search_norms_sync
    except Exception:
        return [], ""
    try:
        norms = search_norms_sync(text_for_search or "", limit=3)
    except Exception:
        return [], ""
    confirmed = [n for n in norms if str(n.get("confidence", "")).upper() in ("CONFIRMED", "PARTIAL")]
    if not confirmed:
        return [], "Норма не подтверждена"
    return confirmed, ""

def _p6f_tnz_build_docx_lines(vision_result, norms, file_name, task_id):
    from datetime import datetime as _dt
    lines = [
        "АКТ ТЕХНИЧЕСКОГО НАДЗОРА",
        "",
        "Дата: " + _dt.now().strftime("%Y-%m-%d %H:%M:%S"),
        "Задача: " + str(task_id),
        "",
    ]
    if file_name:
        lines.append("Источник: фото " + str(file_name))
        lines.append("")
    summary = (vision_result.get("summary") or "").strip() if isinstance(vision_result, dict) else ""
    if summary:
        lines.extend(["Сводка по фото:", summary, ""])
    defects = vision_result.get("defects") or [] if isinstance(vision_result, dict) else []
    if defects:
        lines.append("Выявленные дефекты:")
        for i, d in enumerate(defects, 1):
            lines.append("{}. {}".format(i, d.get("title", "Дефект")))
            if d.get("location"):
                lines.append("   Место: " + str(d["location"]))
            if d.get("severity"):
                lines.append("   Степень: " + str(d["severity"]))
            if d.get("description"):
                lines.append("   Описание: " + str(d["description"]))
            if d.get("recommendation"):
                lines.append("   Рекомендация: " + str(d["recommendation"]))
            lines.append("")
    else:
        lines.extend(["Дефекты на фото не выявлены или фото недостаточно информативно", ""])

    if norms:
        lines.append("Нормативная база:")
        for n in norms:
            lines.append("- " + str(n.get("norm_id", "")) + " — " + str(n.get("section", "")))
            req = str(n.get("requirement", "")).strip()
            if req:
                lines.append("  " + req)
            lines.append("  confidence=" + str(n.get("confidence", "")))
        lines.append("")
    else:
        lines.extend(["Нормативная база: норма не подтверждена", ""])

    confidence = vision_result.get("confidence", "LOW") if isinstance(vision_result, dict) else "LOW"
    lines.append("Источник анализа: OpenRouter Vision (model=google/gemini-2.5-flash, confidence={})".format(confidence))
    return lines

async def _p6f_tnz_upload_to_topic(local_path, file_name, chat_id, topic_id):
    try:
        from core.topic_drive_oauth import upload_file_to_topic
    except Exception as e:
        _P6F_TNZ_REAL_LOG.warning("P6F_TNZ_UPLOAD_IMPORT_ERR %s", e)
        return None, "NO_UPLOADER"
    try:
        result = await upload_file_to_topic(
            file_path=str(local_path),
            file_name=str(file_name),
            chat_id=str(chat_id),
            topic_id=int(topic_id or 5),
        )
        if isinstance(result, dict) and result.get("ok"):
            file_id = result.get("drive_file_id") or result.get("id") or ""
            if file_id:
                return "https://drive.google.com/file/d/" + str(file_id) + "/view", "OK"
            link = result.get("link") or result.get("web_view_link") or ""
            return link, "OK_NO_ID"
        return None, "UPLOAD_FAIL:" + str(result)[:200]
    except Exception as e:
        return None, "UPLOAD_ERR:{}".format(type(e).__name__)

async def p6f_tnz_handle_photo_act_real(file_path, file_name, task_id, chat_id, topic_id, user_text=""):
    """
    Real entry point: photo → Vision → DOCX → Drive upload → return dict.
    Used by topic_5 photo-act flow.
    """
    vision, vstatus = await _p6f_tnz_vision_via_openrouter(file_path)
    if vision is None:
        return {
            "ok": False, "handled": True, "kind": "technadzor_photo_act",
            "state": "WAITING_CLARIFICATION",
            "message": "Не удалось проанализировать фото через Vision ({}). Пришли фото крупнее или текстовое описание дефекта".format(vstatus),
            "history": "P6F_TNZ_VISION_FAIL:{}".format(vstatus),
        }
    norms_text = (vision.get("summary", "") or "") + " " + " ".join(
        str(d.get("title", "")) + " " + str(d.get("description", ""))
        for d in (vision.get("defects") or [])
    ) + " " + str(user_text or "")
    confirmed_norms, _ = _p6f_tnz_norms_block(norms_text)
    docx_lines = _p6f_tnz_build_docx_lines(vision, confirmed_norms, file_name, task_id)

    from datetime import datetime as _dt
    ts = _dt.now().strftime("%Y%m%d_%H%M%S")
    safe = (str(task_id or ts)[:8] or ts).replace("/", "_").replace("\\", "_")
    out_dir = "/root/.areal-neva-core/outputs/technadzor_acts"
    _p6f_tnz_os.makedirs(out_dir, exist_ok=True)
    docx_path = "{}/TECHNADZOR_ACT_PHOTO__{}_{}.docx".format(out_dir, safe, ts)

    try:
        from core.technadzor_engine import _p6tz_make_docx as _make_docx
    except Exception:
        _make_docx = None
    if _make_docx is None:
        return {
            "ok": False, "handled": True, "kind": "technadzor_photo_act",
            "state": "FAILED",
            "message": "DOCX-генератор недоступен (python-docx не установлен)",
            "history": "P6F_TNZ_DOCX_GEN_MISSING",
        }
    written = _make_docx(docx_path, docx_lines)
    if not written:
        return {
            "ok": False, "handled": True, "kind": "technadzor_photo_act",
            "state": "FAILED",
            "message": "Ошибка создания DOCX акта",
            "history": "P6F_TNZ_DOCX_WRITE_FAIL",
        }

    drive_link, ustatus = await _p6f_tnz_upload_to_topic(
        docx_path, _p6f_tnz_os.path.basename(docx_path), chat_id or "-1003725299009", topic_id or 5
    )
    confidence = vision.get("confidence", "LOW")
    summary = vision.get("summary", "") or ""
    defects_count = len(vision.get("defects") or [])
    norms_count = len(confirmed_norms)

    public_lines = [
        "Акт технадзора по фото готов",
        "Файл: " + str(file_name or "photo"),
        "Уверенность Vision: " + str(confidence),
        "Дефектов на фото: " + str(defects_count),
        "Нормативных ссылок: " + str(norms_count) + (" (confidence=PARTIAL/CONFIRMED)" if norms_count else " — норма не подтверждена"),
    ]
    if drive_link:
        public_lines.append("Drive DOCX: " + str(drive_link))
    else:
        public_lines.append("Drive upload: не выполнен (" + str(ustatus) + "). DOCX лежит локально, доставка через Telegram fallback в следующей итерации")
    if summary:
        public_lines.append("")
        public_lines.append("Краткое описание: " + summary[:600])

    history_marker = "P6F_TNZ_PHOTO_ACT_DONE_DEFECTS_{}_NORMS_{}_DRIVE_{}".format(
        defects_count, norms_count, "OK" if drive_link else "FAIL"
    )
    return {
        "ok": True,
        "handled": True,
        "kind": "technadzor_photo_act",
        "state": "DONE" if drive_link else "AWAITING_CONFIRMATION",
        "artifact_path": docx_path,
        "drive_link": drive_link or "",
        "message": "\n".join(public_lines),
        "history": history_marker,
    }

_P6F_TNZ_REAL_LOG.info("P6F_TECHNADZOR_PHOTO_TO_DOCX_REAL_V1_INSTALLED")
# === END_P6F_TECHNADZOR_PHOTO_TO_DOCX_REAL_V1 ===


# === P6H_TOPIC5_USE_EXISTING_TEMPLATES_PHOTO_TO_TECH_REPORT_20260504_V1 ===
# Append-only wrapper.
# - Auto-indexes Drive topic_5 contents (PRIMARY_PDF_STYLE, SECONDARY_DOCX_REFERENCE, etc.)
#   on first photo / "акт" request per chat — without manual "образец" command.
# - On photo: Vision (existing _p6f_tnz_vision_via_openrouter) → section classifier
#   → clean Telegram text (no JSON, no /root, no internal paths).
# - On "сделай акт" / "оформи акт": same + DOCX (service folder _drafts) +
#   client-grade PDF A4 with cyrillic + clickable hyperlinks.
# - DOCX always lands in topic_5/_drafts/ (system).  PDF lands in topic root by default;
#   if user explicitly named a client folder, drop PDF there (only PDF allowed in client
#   folders per spec).
# - Telegram fallback when Drive upload fails — handled by existing P6F path
#   (we keep returning local artifact_path in that case so caller can retry).
import logging as _p6h_logging
import os as _p6h_os
import asyncio as _p6h_asyncio
from pathlib import Path as _P6H_Path
from datetime import datetime as _p6h_dt

_P6H_LOG = _p6h_logging.getLogger("task_worker")

# Eager import so drive_index install marker fires at worker startup
try:
    from core import technadzor_drive_index as _p6h_tdi  # noqa: F401
except Exception as _e_imp:
    _P6H_LOG.warning("P6H_DRIVE_INDEX_IMPORT_FAIL: %s", _e_imp)
_P6H_BASE = _P6H_Path(__file__).resolve().parent.parent
_P6H_OUTDIR = _P6H_BASE / "outputs" / "technadzor_p6h"
_P6H_OUTDIR.mkdir(parents=True, exist_ok=True)

_P6H_DEJAVU = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
_P6H_DEJAVU_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

# Section classifier — keyword (lowercase substring) → canonical section title.
_P6H_SECTIONS = [
    ("Опорные узлы колонн",                        ["опорн", "анкерн", "колонн", "подлив", "опорная плита", "узел опор"]),
    ("Сварные соединения металлоконструкций",      ["сварн", "сварка", "шов", "провар", "наплыв", "сварные стык"]),
    ("Антикоррозионная защита",                    ["окрас", "лакокрас", "корроз", "защитн", "ржавчин", "антикорроз"]),
    ("Состояние основания и прилегающей территории", ["грунт", "основан", "замачив", "размыв", "просадк", "водоотвод", "канав", "лужа"]),
    ("Перекрытия",                                  ["перекрыт", "ригел", "балк перекрыт", "плита перекрыт"]),
    ("Узлы пересечения укосин",                     ["укосин", "связи", "диагональн", "жесткост", "пространств"]),
    ("Узлы крепления элементов покрытия",           ["крепл", "примыкан", "покрыт", "узел кров"]),
    ("Бетонные / железобетонные конструкции",       ["бетон", "железобетон", "арматур", "ж/б", "опалуб", "монолит"]),
    ("Гидроизоляция",                               ["гидроизол", "пароизол", "мембран"]),
    ("Кровля",                                       ["кровл", "крыш", "водосток", "желоб", "конек"]),
    ("Фасад",                                        ["фасад", "облиц", "сайдинг"]),
    ("Общие обзорные материалы",                    ["обзор", "общ вид", "общий вид"]),
]


def _p6h_classify_defect(d):
    text_pool = " ".join([
        str(d.get("title", "") or ""),
        str(d.get("description", "") or ""),
        str(d.get("section_hint", "") or ""),
        str(d.get("category", "") or ""),
    ]).lower()
    for section, kws in _P6H_SECTIONS:
        for kw in kws:
            if kw in text_pool:
                return section
    return "Прочие выявленные замечания"


def _p6h_group_defects_by_section(defects):
    groups = {}
    for d in defects or []:
        sec = _p6h_classify_defect(d)
        groups.setdefault(sec, []).append(d)
    ordered = []
    for sec, _ in _P6H_SECTIONS:
        if sec in groups:
            ordered.append((sec, groups[sec]))
    if "Прочие выявленные замечания" in groups:
        ordered.append(("Прочие выявленные замечания", groups["Прочие выявленные замечания"]))
    return ordered


def _p6h_clean_text(s, limit=4000):
    """Strip JSON/system markers, /root paths, traceback patterns, internal markers
    so result text is safe to send to Telegram."""
    if not s:
        return ""
    txt = str(s)
    # Strip lines starting with /root or containing internal markers
    bad_substrings = ["/root/", "task_id=", "TRACEBACK", "Traceback (most recent",
                       "P6F_", "P6G_", "P6H_", "P6E", "INSTALLED",
                       "MARKER:", "DEBUG:", "DEBUG ", "MANIFEST", "raw_input"]
    cleaned_lines = []
    for ln in txt.splitlines():
        keep = True
        for bad in bad_substrings:
            if bad in ln:
                keep = False
                break
        if keep:
            cleaned_lines.append(ln)
    out = "\n".join(cleaned_lines).strip()
    # Collapse triple+ blank lines
    while "\n\n\n" in out:
        out = out.replace("\n\n\n", "\n\n")
    return out[:limit]


def _p6h_norms_for_haystack(text):
    try:
        from core.normative_engine import search_norms_sync
        return search_norms_sync(text or "", limit=8)
    except Exception:
        return []


def _p6h_norms_for_section(section_title, defect_texts):
    haystack = (section_title + " " + " ".join(defect_texts)).strip()
    return _p6h_norms_for_haystack(haystack)


def _p6h_human_act_number(task_id):
    """Pretty act number 12-03/26 style — fall back to short task_id if missing."""
    today = _p6h_dt.now()
    n = today.strftime("%d-%m") + "/" + today.strftime("%y")
    if task_id:
        suffix = str(task_id)[:6]
        return f"{n}-{suffix}"
    return n


# ────────────────────────────────────────────────────────────────────
# DOCX builder — service-side, draft, with clickable hyperlinks
# ────────────────────────────────────────────────────────────────────
def _p6h_docx_add_hyperlink(paragraph, url, text):
    try:
        from docx.oxml.shared import OxmlElement, qn
    except Exception:
        paragraph.add_run(text + " (" + url + ")")
        return None
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink


def _p6h_build_docx_act(payload, dst_path):
    """payload = {
        act_number, date_str, place, object_descr, method, performer, specialist,
        photos_link, general_purpose, sections=[(title, defects=[{title,description,norm_id,section_norms}], norms=[...], photos_block=[...])],
        recommendations=[str], consequences=[str], violations_table=[(violation, norm_id, photo)]
    }"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(f"АКТ ОСМОТРА ОБЪЕКТА № {payload.get('act_number','')}")
    r.bold = True
    r.font.size = Pt(14)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Методом визуального неразрушающего контроля").italic = True

    doc.add_paragraph(f"Дата осмотра: {payload.get('date_str','')}")
    doc.add_paragraph(f"Место осмотра: {payload.get('place','')}")
    doc.add_paragraph(f"Объект осмотра: {payload.get('object_descr','')}")
    doc.add_paragraph(f"Метод обследования: {payload.get('method','визуальный неразрушающий контроль с выездом на объект')}")
    if payload.get("performer"):
        doc.add_paragraph(f"Представитель подрядчика: {payload['performer']}")
    doc.add_paragraph(f"Технический специалист: {payload.get('specialist','Кузнецов Илья Владимирович')}")

    if payload.get("photos_link"):
        p = doc.add_paragraph("Ссылка на фотоматериалы: ")
        _p6h_docx_add_hyperlink(p, payload["photos_link"], payload["photos_link"])

    # 1. Общие сведения
    doc.add_paragraph()
    h1 = doc.add_paragraph()
    h1.add_run("1. Общие сведения").bold = True
    doc.add_paragraph(payload.get("general_purpose",
        "Осмотр выполнен методом визуального неразрушающего контроля с выездом на "
        "объект. Цель осмотра — выявление фактически наблюдаемых дефектов, определение "
        "рекомендаций к устранению и возможных последствий для заказчика при сохранении "
        "текущего состояния объекта"))

    # 2. Установлено по факту осмотра
    doc.add_paragraph()
    h2 = doc.add_paragraph()
    h2.add_run("2. Установлено по факту осмотра").bold = True
    sections = payload.get("sections") or []
    for i, sec in enumerate(sections, 1):
        ph = doc.add_paragraph()
        ph.add_run(f"2.{i} {sec.get('title','')}").bold = True
        # facts
        for d in sec.get("defects") or []:
            line = (d.get("title") or "").strip()
            descr = (d.get("description") or "").strip()
            if line and descr:
                doc.add_paragraph(f"— {line}: {descr}")
            elif line:
                doc.add_paragraph(f"— {line}")
            elif descr:
                doc.add_paragraph(f"— {descr}")
        # norm refs for section
        norms = sec.get("norms") or []
        if norms:
            np = doc.add_paragraph()
            np.add_run("Нормативная отсылка: ").italic = True
            np.add_run("; ".join(f"{n.get('norm_id','')} — {n.get('section','')}" for n in norms if n.get("norm_id")))
        else:
            np = doc.add_paragraph()
            np.add_run("Нормативная отсылка: ").italic = True
            np.add_run("норма не подтверждена").italic = True
        # photos
        photos = sec.get("photos_block") or []
        if photos:
            pp = doc.add_paragraph()
            pp.add_run("Фотоматериалы: ").italic = True
            pp.add_run(", ".join(photos))

    # Рекомендации
    if payload.get("recommendations"):
        doc.add_paragraph()
        rh = doc.add_paragraph()
        rh.add_run("3. Рекомендовано к устранению").bold = True
        for i, line in enumerate(payload["recommendations"], 1):
            doc.add_paragraph(f"{i}. {line}")

    # Возможные последствия
    if payload.get("consequences"):
        doc.add_paragraph()
        ch = doc.add_paragraph()
        ch.add_run("4. Возможные последствия при отсутствии устранения").bold = True
        for line in payload["consequences"]:
            doc.add_paragraph(f"— {line}")

    # Таблица: Нарушение / Норматив / Фото
    if payload.get("violations_table"):
        doc.add_paragraph()
        th = doc.add_paragraph()
        th.add_run("5. Сводная таблица нарушений").bold = True
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Light Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Нарушение"
        hdr[1].text = "Норматив"
        hdr[2].text = "Фото"
        for v, n, ph in payload["violations_table"]:
            row = tbl.add_row().cells
            row[0].text = str(v or "")
            row[1].text = str(n or "норма не подтверждена")
            row[2].text = str(ph or "")

    # Подпись
    doc.add_paragraph()
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run(f"Технический специалист: {payload.get('specialist','Кузнецов Илья Владимирович')}").bold = True
    doc.add_paragraph(f"Дата: {payload.get('date_str','')}")

    doc.save(str(dst_path))
    return str(dst_path)


# ────────────────────────────────────────────────────────────────────
# PDF builder — A4, cyrillic, clickable hyperlinks (reportlab platypus)
# ────────────────────────────────────────────────────────────────────
def _p6h_register_fonts():
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        if "DejaVuSans" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("DejaVuSans", _P6H_DEJAVU))
        if "DejaVuSans-Bold" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", _P6H_DEJAVU_BOLD))
        return True
    except Exception:
        return False


def _p6h_build_pdf_act(payload, dst_path):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
    from reportlab.lib.units import cm
    from reportlab.lib import colors

    _p6h_register_fonts()

    styles = getSampleStyleSheet()
    base_font = "DejaVuSans"
    bold_font = "DejaVuSans-Bold"

    sty_title = ParagraphStyle("title", parent=styles["Title"], fontName=bold_font, fontSize=14, alignment=1, spaceAfter=4)
    sty_subtitle = ParagraphStyle("subtitle", parent=styles["Normal"], fontName=base_font, fontSize=10, alignment=1, textColor=colors.grey, spaceAfter=10)
    sty_h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontName=bold_font, fontSize=12, spaceBefore=8, spaceAfter=4)
    sty_h3 = ParagraphStyle("h3", parent=styles["Heading3"], fontName=bold_font, fontSize=11, spaceBefore=6, spaceAfter=2)
    sty_body = ParagraphStyle("body", parent=styles["Normal"], fontName=base_font, fontSize=10, leading=13, spaceAfter=2)
    sty_italic = ParagraphStyle("italic", parent=sty_body, textColor=colors.grey)
    sty_small = ParagraphStyle("small", parent=sty_body, fontSize=9)

    flow = []
    flow.append(Paragraph(f"АКТ ОСМОТРА ОБЪЕКТА № {payload.get('act_number','')}", sty_title))
    flow.append(Paragraph("Методом визуального неразрушающего контроля", sty_subtitle))
    flow.append(Paragraph(f"<b>Дата осмотра:</b> {payload.get('date_str','')}", sty_body))
    flow.append(Paragraph(f"<b>Место осмотра:</b> {payload.get('place','')}", sty_body))
    flow.append(Paragraph(f"<b>Объект осмотра:</b> {payload.get('object_descr','')}", sty_body))
    flow.append(Paragraph(f"<b>Метод обследования:</b> {payload.get('method','визуальный неразрушающий контроль с выездом на объект')}", sty_body))
    if payload.get("performer"):
        flow.append(Paragraph(f"<b>Представитель подрядчика:</b> {payload['performer']}", sty_body))
    flow.append(Paragraph(f"<b>Технический специалист:</b> {payload.get('specialist','Кузнецов Илья Владимирович')}", sty_body))
    if payload.get("photos_link"):
        flow.append(Paragraph(
            f'<b>Ссылка на фотоматериалы:</b> <link href="{payload["photos_link"]}"><font color="#0563C1"><u>{payload["photos_link"]}</u></font></link>',
            sty_body,
        ))

    flow.append(Spacer(1, 8))
    flow.append(Paragraph("1. Общие сведения", sty_h2))
    flow.append(Paragraph(payload.get("general_purpose",
        "Осмотр выполнен методом визуального неразрушающего контроля с выездом на "
        "объект. Цель осмотра — выявление фактически наблюдаемых дефектов, определение "
        "рекомендаций к устранению и возможных последствий для заказчика."), sty_body))

    flow.append(Paragraph("2. Установлено по факту осмотра", sty_h2))
    sections = payload.get("sections") or []
    for i, sec in enumerate(sections, 1):
        flow.append(Paragraph(f"2.{i} {sec.get('title','')}", sty_h3))
        for d in sec.get("defects") or []:
            line = (d.get("title") or "").strip()
            descr = (d.get("description") or "").strip()
            txt = (line + ((": " + descr) if descr and line else descr)).strip()
            if txt:
                flow.append(Paragraph("— " + txt, sty_body))
        norms = sec.get("norms") or []
        if norms:
            ns = "; ".join(f"{n.get('norm_id','')} — {n.get('section','')}" for n in norms if n.get("norm_id"))
            flow.append(Paragraph(f"<i>Нормативная отсылка:</i> {ns}", sty_italic))
        else:
            flow.append(Paragraph("<i>Нормативная отсылка: норма не подтверждена</i>", sty_italic))
        photos = sec.get("photos_block") or []
        if photos:
            flow.append(Paragraph(f"<i>Фотоматериалы:</i> {', '.join(photos)}", sty_italic))

    if payload.get("recommendations"):
        flow.append(Spacer(1, 4))
        flow.append(Paragraph("3. Рекомендовано к устранению", sty_h2))
        for i, line in enumerate(payload["recommendations"], 1):
            flow.append(Paragraph(f"{i}. {line}", sty_body))

    if payload.get("consequences"):
        flow.append(Spacer(1, 4))
        flow.append(Paragraph("4. Возможные последствия при отсутствии устранения", sty_h2))
        for line in payload["consequences"]:
            flow.append(Paragraph("— " + line, sty_body))

    if payload.get("violations_table"):
        flow.append(Spacer(1, 6))
        flow.append(Paragraph("5. Сводная таблица нарушений", sty_h2))
        rows = [["Нарушение", "Норматив", "Фото"]]
        for v, n, ph in payload["violations_table"]:
            rows.append([Paragraph(str(v or ""), sty_small),
                         Paragraph(str(n or "норма не подтверждена"), sty_small),
                         Paragraph(str(ph or ""), sty_small)])
        tbl = Table(rows, colWidths=[7*cm, 6*cm, 4*cm], repeatRows=1)
        tbl.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), base_font),
            ("FONTNAME", (0, 0), (-1, 0), bold_font),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        flow.append(tbl)

    flow.append(Spacer(1, 16))
    flow.append(Paragraph(
        f"<b>Технический специалист:</b> {payload.get('specialist','Кузнецов Илья Владимирович')}",
        sty_body,
    ))
    flow.append(Paragraph(f"<b>Дата:</b> {payload.get('date_str','')}", sty_body))

    doc = SimpleDocTemplate(
        str(dst_path), pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=f"Акт осмотра № {payload.get('act_number','')}",
    )
    doc.build(flow)
    return str(dst_path)


# ────────────────────────────────────────────────────────────────────
# Async pipeline: photo → Vision → sections → response (and optionally DOCX+PDF)
# ────────────────────────────────────────────────────────────────────
async def _p6h_process_photo_async(file_path, file_name, task_id, chat_id, topic_id, user_text="", make_act=False, place="", object_descr=""):
    # Ensure Drive index is built (silently, best-effort)
    try:
        from core import technadzor_drive_index as _tdi
        idx = _tdi.build_technadzor_template_index(str(chat_id), int(topic_id), force=False)
    except Exception:
        idx = {}

    # Vision
    vision, vstatus = await _p6f_tnz_vision_via_openrouter(file_path)
    if vstatus == "FAIL" or vision is None:
        return {
            "ok": False, "handled": True, "kind": "technadzor_p6h_photo",
            "state": "WAITING_CLARIFICATION",
            "message": "Не удалось проанализировать фото через Vision. Пришли крупнее или текстовое описание дефекта",
            "history": "P6H_VISION_FAIL",
        }

    summary = (vision.get("summary") or "").strip() if isinstance(vision, dict) else ""
    defects = (vision.get("defects") or []) if isinstance(vision, dict) else []
    confidence = vision.get("confidence", "LOW") if isinstance(vision, dict) else "LOW"

    grouped = _p6h_group_defects_by_section(defects)

    # Norms per section + global
    section_norms = []
    all_haystack = summary + " " + (user_text or "") + " " + " ".join(
        str(d.get("title", "")) + " " + str(d.get("description", "")) for d in defects
    )
    global_norms = _p6h_norms_for_haystack(all_haystack)

    # Build sections payload (used by both DOCX and PDF, and partly by Telegram)
    sections_payload = []
    for sec_title, ds in grouped:
        defect_texts = [str(d.get("title", "")) + " " + str(d.get("description", "")) for d in ds]
        snorms = _p6h_norms_for_section(sec_title, defect_texts)
        section_norms.append((sec_title, snorms))
        sections_payload.append({
            "title": sec_title,
            "defects": ds,
            "norms": snorms,
            "photos_block": [str(file_name or "")] if file_name else [],
        })

    # Topic folder link from index for "Ссылка на фотоматериалы"
    topic_folder_link = (idx or {}).get("topic_folder_link", "")

    # Build clean Telegram text response.
    # If make_act=False — photo-only response (template per spec).
    # If make_act=True  — short summary + DOCX/PDF links.
    if not make_act:
        out_lines = ["Технический осмотр по фото", ""]
        # 1. Что видно
        out_lines.append("1. Что видно:")
        if summary:
            out_lines.append(_p6h_clean_text(summary, 700))
        else:
            out_lines.append("Описание не сформировано")
        out_lines.append("")
        # 2. Замечания
        out_lines.append("2. Обнаруженные замечания:")
        if grouped:
            for sec_title, ds in grouped:
                titles = []
                for d in ds:
                    t = str(d.get("title") or d.get("description") or "").strip()
                    if t:
                        titles.append(t[:120])
                line = f"— {sec_title}: " + ("; ".join(titles) if titles else "замечания зафиксированы")
                out_lines.append(_p6h_clean_text(line, 500))
        else:
            out_lines.append("Дефектов на фото не выявлено")
        out_lines.append("")
        # 3. Почему плохо — берём severity/why из Vision если есть, иначе общий
        out_lines.append("3. Почему это плохо:")
        why_lines = []
        for d in defects:
            w = str(d.get("why") or d.get("severity") or d.get("impact") or "").strip()
            if w:
                why_lines.append("— " + w[:200])
        if why_lines:
            out_lines.extend(why_lines[:6])
        elif grouped:
            out_lines.append("Зафиксированные отклонения снижают эксплуатационную надёжность и/или несущую способность конструкции и требуют проверки и устранения")
        else:
            out_lines.append("Отклонений не выявлено")
        out_lines.append("")
        # 4. Как исправить
        out_lines.append("4. Как исправить:")
        fix_lines = []
        for d in defects:
            f = str(d.get("fix") or d.get("recommendation") or "").strip()
            if f:
                fix_lines.append("— " + f[:200])
        if fix_lines:
            out_lines.extend(fix_lines[:6])
        else:
            out_lines.append("Привести узлы и покрытия к нормативному состоянию по соответствующим СП/ГОСТ. Уточнить решения по проектной документации")
        out_lines.append("")
        # 5. Что проверить на объекте
        out_lines.append("5. Что проверить на объекте:")
        check_lines = []
        for d in defects:
            c = str(d.get("verify") or d.get("check") or "").strip()
            if c:
                check_lines.append("— " + c[:200])
        if check_lines:
            out_lines.extend(check_lines[:6])
        else:
            out_lines.append("Состояние конструкции в полном объёме, наличие и соответствие исполнительной документации, ранее выданные предписания и их устранение")
        out_lines.append("")
        # 6. Норма
        out_lines.append("6. Нормативная отсылка:")
        if global_norms:
            for n in global_norms[:5]:
                out_lines.append("— " + str(n.get("norm_id","")) + ": " + str(n.get("section","")) + f" [{n.get('confidence','PARTIAL')}]")
        else:
            out_lines.append("норма не подтверждена")
        out_lines.append("")
        # 7. Акт
        out_lines.append("7. Акт:")
        out_lines.append("Могу оформить акт по текущим фото — напишите «сделай акт» / «оформи акт»")

        return {
            "ok": True, "handled": True, "kind": "technadzor_p6h_photo",
            "state": "DONE",
            "message": _p6h_clean_text("\n".join(out_lines), 6000),
            "history": "P6H_PHOTO_REPORT_DEFECTS_{}_NORMS_{}".format(len(defects), len(global_norms)),
        }

    # ── Build DOCX (service _drafts) + PDF (client topic root) ──
    ts = _p6h_dt.now().strftime("%Y%m%d_%H%M%S")
    safe_tid = (str(task_id or ts)[:8] or ts).replace("/", "_").replace("\\", "_")
    docx_local = _P6H_OUTDIR / f"P6H_TNZ_ACT_DRAFT__{safe_tid}_{ts}.docx"
    pdf_local = _P6H_OUTDIR / f"АКТ_ОСМОТРА__{safe_tid}_{ts}.pdf"

    # Recommendations / consequences — pulled from defects
    recs = []
    cons = []
    for d in defects:
        r = str(d.get("fix") or d.get("recommendation") or "").strip()
        if r:
            recs.append(r[:300])
        c = str(d.get("consequence") or d.get("why") or "").strip()
        if c:
            cons.append(c[:300])
    # Violation table rows
    vtable = []
    for sec_title, ds in grouped:
        for d in ds:
            v = str(d.get("title") or d.get("description") or sec_title)[:200]
            sn = ""
            for n in section_norms:
                if n[0] == sec_title and n[1]:
                    sn = n[1][0].get("norm_id", "")
                    break
            ph = str(file_name or "")
            vtable.append((v, sn or "норма не подтверждена", ph))

    payload = {
        "act_number": _p6h_human_act_number(task_id),
        "date_str": _p6h_dt.now().strftime("%d.%m.%Y"),
        "place": place or "место уточняется по запросу владельца",
        "object_descr": object_descr or "объект уточняется по запросу владельца",
        "method": "визуальный неразрушающий контроль с выездом на объект",
        "performer": "",
        "specialist": "Кузнецов Илья Владимирович",
        "photos_link": topic_folder_link or "",
        "general_purpose": (
            "Осмотр выполнен методом визуального неразрушающего контроля. "
            "Цель осмотра — выявление фактически наблюдаемых дефектов, определение "
            "рекомендаций к устранению и возможных последствий для заказчика при "
            "сохранении текущего состояния объекта."
        ),
        "sections": sections_payload,
        "recommendations": recs[:20] if recs else ["Привести выявленные узлы и покрытия к нормативному состоянию по соответствующим СП/ГОСТ"],
        "consequences": cons[:10] if cons else ["Снижение несущей способности и эксплуатационной надёжности конструкций"],
        "violations_table": vtable[:30],
    }

    docx_ok = False
    pdf_ok = False
    try:
        _p6h_build_docx_act(payload, docx_local)
        docx_ok = True
    except Exception:
        _P6H_LOG.exception("P6H_DOCX_BUILD_FAIL")
    try:
        _p6h_build_pdf_act(payload, pdf_local)
        pdf_ok = True
    except Exception:
        _P6H_LOG.exception("P6H_PDF_BUILD_FAIL")

    drive_docx = None
    drive_pdf = None
    if docx_ok or pdf_ok:
        try:
            from core import technadzor_drive_index as _tdi
            if docx_ok:
                drive_docx = _tdi.upload_to_service_subfolder(
                    docx_local, docx_local.name, str(chat_id), int(topic_id), subfolder="_drafts",
                )
            if pdf_ok:
                drive_pdf = _tdi.upload_client_pdf_to_folder(
                    pdf_local, pdf_local.name, str(chat_id), int(topic_id), target_folder_name=None,
                )
        except Exception:
            _P6H_LOG.exception("P6H_UPLOAD_FAIL")

    pdf_link = (drive_pdf or {}).get("link", "") if drive_pdf else ""
    docx_link = (drive_docx or {}).get("link", "") if drive_docx else ""

    msg_lines = ["Акт сформирован"]
    if pdf_link:
        msg_lines.append(f"PDF: {pdf_link}")
    elif pdf_ok:
        msg_lines.append("PDF: подготовлен локально, загрузка на Drive не выполнена — повторная попытка через retry queue")
    else:
        msg_lines.append("PDF: ошибка генерации")
    if docx_link:
        msg_lines.append(f"DOCX (черновик, служебно): {docx_link}")
    if topic_folder_link:
        msg_lines.append(f"Фото: {topic_folder_link}")
    msg_lines.append("Норма: " + ("подтверждена" if global_norms else "не подтверждена"))

    return {
        "ok": True if (pdf_ok or docx_ok) else False,
        "handled": True,
        "kind": "technadzor_p6h_act",
        "state": "DONE" if pdf_link else "AWAITING_CONFIRMATION",
        "artifact_path": str(pdf_local if pdf_ok else docx_local),
        "extra_artifact_path": str(docx_local if docx_ok else ""),
        "drive_link": pdf_link or docx_link or "",
        "message": _p6h_clean_text("\n".join(msg_lines), 4000),
        "history": "P6H_ACT_DOCX_{}_PDF_{}_DRIVE_PDF_{}".format(
            "OK" if docx_ok else "FAIL",
            "OK" if pdf_ok else "FAIL",
            "OK" if pdf_link else "FAIL",
        ),
    }


# ────────────────────────────────────────────────────────────────────
# Sync wrapper around process_technadzor — only intercepts topic_5
# ────────────────────────────────────────────────────────────────────
def _p6h_is_image_path(path):
    if not path:
        return False
    p = str(path).lower()
    return p.endswith((".jpg", ".jpeg", ".png", ".webp", ".bmp", ".heic"))


_P6H_ACT_TRIGGERS = (
    "сделай акт", "оформи акт", "выдай акт", "финальн", "акт по фото",
    "сформируй акт", "подготовь акт", "сделать акт", "оформить акт",
)


def _p6h_should_handle(topic_id, file_path, user_text):
    if int(topic_id or 0) != 5:
        return False
    if file_path and _p6h_is_image_path(file_path):
        return True
    low = (user_text or "").lower()
    if any(t in low for t in _P6H_ACT_TRIGGERS):
        return True
    return False


def _p6h_run_async(coro):
    """Run coro from sync context. Mimics codebase pattern."""
    try:
        return _p6h_asyncio.run(coro)
    except RuntimeError:
        # Already in a loop — schedule in a worker thread with its own loop
        import threading
        result = {"v": None, "exc": None}
        def _runner():
            new_loop = _p6h_asyncio.new_event_loop()
            try:
                _p6h_asyncio.set_event_loop(new_loop)
                result["v"] = new_loop.run_until_complete(coro)
            except Exception as e:
                result["exc"] = e
            finally:
                new_loop.close()
        t = threading.Thread(target=_runner, daemon=True)
        t.start()
        t.join()
        if result["exc"]:
            raise result["exc"]
        return result["v"]


try:
    _p6h_orig_process = process_technadzor
    if not getattr(_p6h_orig_process, "_p6h_wrapped", False):
        def process_technadzor(text="", task_id="", chat_id="", topic_id=0, file_path="", file_name="", **kwargs):
            try:
                if _p6h_should_handle(topic_id, file_path, text):
                    is_image = _p6h_is_image_path(file_path)
                    low = (text or "").lower()
                    make_act = any(t in low for t in _P6H_ACT_TRIGGERS)
                    if is_image or make_act:
                        return _p6h_run_async(
                            _p6h_process_photo_async(
                                file_path=file_path,
                                file_name=file_name,
                                task_id=task_id,
                                chat_id=chat_id,
                                topic_id=topic_id,
                                user_text=text,
                                make_act=make_act,
                            )
                        )
            except Exception:
                _P6H_LOG.exception("P6H_WRAPPER_FAIL — falling back to original")
            return _p6h_orig_process(
                text=text, task_id=task_id, chat_id=chat_id, topic_id=topic_id,
                file_path=file_path, file_name=file_name, **kwargs,
            )
        process_technadzor._p6h_wrapped = True
        _p6h_orig_process._p6h_wrapped = True
    _P6H_LOG.info("P6H_TOPIC5_USE_EXISTING_TEMPLATES_PHOTO_TO_TECH_REPORT_20260504_V1_INSTALLED")
except Exception as _exc:
    _P6H_LOG.exception("P6H_TOPIC5_INSTALL_FAIL: %s", _exc)
# === END_P6H_TOPIC5_USE_EXISTING_TEMPLATES_PHOTO_TO_TECH_REPORT_20260504_V1 ===


# === P6H_PART_2: PHOTO_NUMBER_DEFECT_NORM_CLARIFICATION_LOGIC + VOICE_LIVE_DIALOG_CLARIFICATION_GATE_20260504 ===
# Adds on top of P6H_PART_1:
# - voice transcript parser (extracts object_hint, folder_hint, visit_date_hint, client_facing flag, raw user-stated defects)
# - clarification gate (returns WAITING_CLARIFICATION with concrete questions, never «что строим?»)
# - photo-numbered Telegram output («Фото №N — <file>»)
# - 8-column violations_table for акты
# - memory summary write after DONE (key topic_5_technadzor_photo_report_summary)
# - replaces _p6h_process_photo_async with enhanced version
import re as _p6h2_re
import json as _p6h2_json

_P6H2_LOG = _p6h_logging.getLogger("task_worker")

_P6H2_FOLDER_HINT_PATTERNS = [
    _p6h2_re.compile(r"папк[аеу]?\s+[«\"\'']?([^«»\"\'\.,;\n]{2,80})[»\"\'']?", _p6h2_re.IGNORECASE),
    _p6h2_re.compile(r"директор[ияю]?\s+[«\"\'']?([^«»\"\'\.,;\n]{2,80})[»\"\'']?", _p6h2_re.IGNORECASE),
]
_P6H2_OBJECT_HINTS = [
    "ангар", "склад", "цех", "корпус", "коттедж", "дом", "здание", "сооружение",
    "фундамент", "кровл", "фасад", "перекрыт", "колонн", "балк", "ферм", "плита",
    "паркинг", "гараж", "трибун", "купол", "теплиц",
]
_P6H2_DATE_RE = _p6h2_re.compile(
    r"(?P<d>\d{1,2})[\s\-./](?P<m>\d{1,2}|январ|феврал|март|апрел|ма[яй]|июн|июл|август|сентябр|октябр|ноябр|декабр)[\s\-./а-я]*?(?P<y>\d{2,4})?",
    _p6h2_re.IGNORECASE,
)
_P6H2_CLIENT_HINTS = ("заказчик", "клиент", "генподряд", "застройщ", "владельцу объекта")
_P6H2_INTERNAL_HINTS = ("служебн", "не для клиента", "черновик", "внутрен", "тестов", "smoke")


def _p6h_parse_voice_instruction(raw_input):
    """Parse [VOICE] transcript into structured TZ context.

    Returns dict with keys:
      object_hint, folder_hint, visit_date_hint, client_facing,
      explicit_include, explicit_exclude, output_kind, requires_act
    All values are conservative — present only when transcript explicitly mentioned them.
    """
    text = (raw_input or "").strip()
    if text.startswith("[VOICE]"):
        text_body = text[len("[VOICE]"):].strip()
    else:
        text_body = text
    low = text_body.lower()
    ctx = {
        "is_voice": text.startswith("[VOICE]"),
        "transcript": text_body,
        "object_hint": "",
        "folder_hint": "",
        "visit_date_hint": "",
        "client_facing": None,  # None=unknown, True/False=explicit
        "explicit_include": [],
        "explicit_exclude": [],
        "output_kind": "",  # text|act|pdf|docx|description
        "requires_act": False,
    }
    if not text_body:
        return ctx

    # Object
    for h in _P6H2_OBJECT_HINTS:
        if h in low:
            ctx["object_hint"] = h
            break

    # Folder hint
    for pat in _P6H2_FOLDER_HINT_PATTERNS:
        m = pat.search(text_body)
        if m:
            ctx["folder_hint"] = m.group(1).strip()
            break

    # Date hint
    m = _P6H2_DATE_RE.search(text_body)
    if m:
        ctx["visit_date_hint"] = m.group(0).strip()

    # Client-facing
    if any(h in low for h in _P6H2_CLIENT_HINTS):
        ctx["client_facing"] = True
    if any(h in low for h in _P6H2_INTERNAL_HINTS):
        ctx["client_facing"] = False

    # Output
    if any(t in low for t in _P6H_ACT_TRIGGERS):
        ctx["requires_act"] = True
        ctx["output_kind"] = "act"
    elif "pdf" in low or "пдф" in low:
        ctx["output_kind"] = "pdf"
    elif "docx" in low or "ворд" in low or "word" in low:
        ctx["output_kind"] = "docx"
    elif "описан" in low or "опиши" in low or "посмотри" in low:
        ctx["output_kind"] = "description"

    # Explicit include/exclude lists
    for marker, key in (("включ", "explicit_include"), ("не включ", "explicit_exclude"),
                         ("исключ", "explicit_exclude")):
        i = low.find(marker)
        if i >= 0:
            tail = text_body[i:i + 240]
            ctx[key].append(tail.strip())
    return ctx


def _p6h_should_wait_for_clarification(vision, defects, voice_ctx, drive_idx):
    """Returns (should_wait: bool, questions: [str]).

    Triggers:
      • Vision confidence=LOW AND no defects → ask what's on the photo, before/after, side
      • voice_ctx folder_hint set but folder not in Drive index
      • voice_ctx client_facing is None AND folder_hint set (need to know if it's customer-visible)
    Never asks "что строим?" / "что это?" / "пришлите шаблон".
    """
    questions = []
    confidence = (vision or {}).get("confidence", "LOW") if isinstance(vision, dict) else "LOW"
    nd = len(defects or [])

    # 1. Folder named by owner — check it exists
    fh = (voice_ctx or {}).get("folder_hint", "").strip()
    if fh:
        folders = []
        for f in (drive_idx or {}).get("folders_client", []) + (drive_idx or {}).get("folders_system", []):
            folders.append(f.get("name", ""))
        if not any(fh.lower() == n.lower() or fh.lower() in n.lower() for n in folders if n):
            questions.append(
                f"Не нашёл папку «{fh}» в Drive topic_5. "
                "Уточни точное имя или создай её перед загрузкой фото"
            )

    # 2. Folder client_facing flag uncertain
    if fh and (voice_ctx or {}).get("client_facing") is None:
        questions.append(
            f"Папку «{fh}» считать клиентской (туда складывать только фото и чистовой PDF) "
            "или служебной (тогда туда можно DOCX-черновик)?"
        )

    # 3. Vision low-confidence + no defects on a photo
    if confidence == "LOW" and nd == 0:
        questions.append(
            "Фото не позволяет однозначно определить дефект. "
            "Это до или после исправления? С какой стороны конструкции снято? "
            "Что именно нужно зафиксировать на этом фото?"
        )

    # Never trigger generic clarifier
    return (bool(questions), questions)


def _p6h_save_summary_to_memory(chat_id, topic_id, summary_dict):
    """Persist a compact summary to memory.db under topic_5_technadzor_photo_report_summary.

    Uses task_worker._save_memory_safe if exposed; falls back to direct memory_client write.
    Stores ONLY: folder, object, date, owner_directives, defect_brief, pdf_link, status.
    """
    try:
        payload = {
            "folder": str(summary_dict.get("folder") or "")[:200],
            "object": str(summary_dict.get("object") or "")[:200],
            "date": str(summary_dict.get("date") or "")[:40],
            "owner_directives": (summary_dict.get("owner_directives") or [])[:6],
            "defect_brief": (summary_dict.get("defect_brief") or [])[:12],
            "pdf_link": str(summary_dict.get("pdf_link") or "")[:400],
            "docx_link": str(summary_dict.get("docx_link") or "")[:400],
            "status": str(summary_dict.get("status") or "")[:60],
            "ts": int(_p6h_dt.now().timestamp()),
        }
        body = _p6h2_json.dumps(payload, ensure_ascii=False)[:4000]
    except Exception:
        return False
    try:
        from core.memory_client import save_memory  # type: ignore
        key = f"topic_{int(topic_id or 5)}_technadzor_photo_report_summary"
        save_memory(chat_id=str(chat_id), key=key, value=body)
        return True
    except Exception:
        try:
            import sqlite3 as _sql
            db = _P6H_BASE / "data" / "memory.db"
            con = _sql.connect(str(db))
            cur = con.cursor()
            cur.execute(
                "INSERT INTO memory(chat_id,key,value,timestamp) VALUES (?,?,?,strftime('%s','now'))",
                (str(chat_id), f"topic_{int(topic_id or 5)}_technadzor_photo_report_summary", body),
            )
            con.commit()
            con.close()
            return True
        except Exception:
            _P6H2_LOG.exception("P6H2_MEMORY_SAVE_FAIL")
            return False


def _p6h_format_photo_numbered_response(vision, defects, grouped, global_norms, file_name, photo_no=1, voice_ctx=None):
    """Build per-photo numbered Telegram text per spec section 'Формат текстового ответа без акта'."""
    fname = str(file_name or "photo")
    out = ["Технический осмотр по фото", ""]
    if voice_ctx and (voice_ctx.get("folder_hint") or voice_ctx.get("object_hint")):
        out.append("Объект / папка:")
        parts = []
        if voice_ctx.get("object_hint"):
            parts.append(voice_ctx["object_hint"])
        if voice_ctx.get("folder_hint"):
            parts.append(f"папка «{voice_ctx['folder_hint']}»")
        if voice_ctx.get("visit_date_hint"):
            parts.append(f"дата {voice_ctx['visit_date_hint']}")
        out.append(" / ".join(parts))
        out.append("")

    out.append(f"Фото №{photo_no} — {fname}")

    summary = (vision or {}).get("summary", "") if isinstance(vision, dict) else ""
    out.append("Что видно:")
    if summary:
        out.append(_p6h_clean_text(summary, 600))
    else:
        out.append("по фото однозначно не определяется")

    out.append("")
    out.append("Замечание:")
    if grouped:
        per_section = []
        for sec_title, ds in grouped:
            tt = []
            for d in ds:
                t = str(d.get("title") or d.get("description") or "").strip()
                if t:
                    tt.append(t[:120])
            per_section.append(f"— {sec_title}: " + ("; ".join(tt) if tt else "замечания зафиксированы"))
        out.extend(per_section)
    else:
        out.append("Дефектов на фото не выявлено")

    out.append("")
    out.append("Чем опасно:")
    risks = []
    for d in defects or []:
        r = str(d.get("risk") or d.get("why") or d.get("severity") or d.get("impact") or "").strip()
        if r:
            risks.append("— " + r[:200])
    if risks:
        out.extend(risks[:6])
    elif grouped:
        out.append("Зафиксированные отклонения снижают эксплуатационную надёжность и/или несущую способность конструкции")
    else:
        out.append("Отклонений не выявлено")

    out.append("")
    out.append("Как исправить:")
    fixes = []
    for d in defects or []:
        f = str(d.get("recommended_fix") or d.get("fix") or d.get("recommendation") or "").strip()
        if f:
            fixes.append("— " + f[:200])
    if fixes:
        out.extend(fixes[:6])
    else:
        out.append("Привести к нормативному состоянию по соответствующим СП/ГОСТ; уточнить решения по проектной документации")

    out.append("")
    out.append("Что проверить:")
    checks = []
    for d in defects or []:
        c = d.get("site_checks") or d.get("verify") or d.get("check")
        if isinstance(c, list):
            for x in c[:3]:
                xs = str(x).strip()
                if xs:
                    checks.append("— " + xs[:200])
        elif c:
            checks.append("— " + str(c)[:200])
    if checks:
        out.extend(checks[:6])
    else:
        out.append("Состояние конструкции в полном объёме, исполнительная документация, ранее выданные предписания")

    out.append("")
    out.append("Нормативная отсылка:")
    if global_norms:
        for n in global_norms[:5]:
            out.append("— " + str(n.get("norm_id", "")) + ": " + str(n.get("section", "")) + f" [{n.get('confidence', 'PARTIAL')}]")
    else:
        out.append("норма не подтверждена")

    out.append("")
    out.append("Итог:")
    crit = [d for d in (defects or []) if str(d.get("severity", "")).lower() in ("high", "critical", "критическ", "high_risk")]
    out.append(f"— критичные замечания: {len(crit)}")
    out.append(f"— рабочие замечания: {len(defects or []) - len(crit)}")
    out.append("— нужен ли акт: могу оформить акт по текущим фото — напиши «сделай акт»")

    return _p6h_clean_text("\n".join(out), 6000)


# Override _p6h_process_photo_async with the enhanced version
async def _p6h_process_photo_async(file_path, file_name, task_id, chat_id, topic_id, user_text="", make_act=False, place="", object_descr=""):
    # 1. Drive index — best-effort eager build
    try:
        from core import technadzor_drive_index as _tdi
        idx = _tdi.build_technadzor_template_index(str(chat_id), int(topic_id), force=False)
    except Exception:
        idx = {}

    # 2. Voice transcript parsing
    voice_ctx = _p6h_parse_voice_instruction(user_text or "")

    # 3. Vision
    vision, vstatus = await _p6f_tnz_vision_via_openrouter(file_path)
    if vstatus == "FAIL" or vision is None:
        return {
            "ok": False, "handled": True, "kind": "technadzor_p6h_photo",
            "state": "WAITING_CLARIFICATION",
            "message": "Не удалось проанализировать фото через Vision. "
                       "Пришли крупнее или короткое описание дефекта текстом",
            "history": "P6H_VISION_FAIL",
        }

    summary = (vision.get("summary") or "").strip() if isinstance(vision, dict) else ""
    defects = (vision.get("defects") or []) if isinstance(vision, dict) else []

    # Tag defects with photo_no/file_name (single-photo task = #1)
    photo_no = 1
    for d in defects:
        if isinstance(d, dict):
            d.setdefault("photo_no", photo_no)
            d.setdefault("file_name", file_name or "")

    grouped = _p6h_group_defects_by_section(defects)

    # 4. Clarification gate
    should_wait, questions = _p6h_should_wait_for_clarification(vision, defects, voice_ctx, idx)
    if should_wait and not make_act:
        return {
            "ok": True, "handled": True, "kind": "technadzor_p6h_clarify",
            "state": "WAITING_CLARIFICATION",
            "message": _p6h_clean_text(
                "Технадзор topic_5 — нужны уточнения перед разбором:\n\n" + "\n".join(f"— {q}" for q in questions),
                3000,
            ),
            "history": "P6H_CLARIFY:{}".format(len(questions)),
        }

    # 5. Norms
    haystack = " ".join([
        summary, voice_ctx.get("transcript", "") or "",
    ] + [str(d.get("title", "")) + " " + str(d.get("description", "")) for d in defects])
    global_norms = _p6h_norms_for_haystack(haystack)

    # 6. Per-section payload
    sections_payload = []
    for sec_title, ds in grouped:
        defect_texts = [str(d.get("title", "")) + " " + str(d.get("description", "")) for d in ds]
        snorms = _p6h_norms_for_section(sec_title, defect_texts)
        sections_payload.append({
            "title": sec_title,
            "defects": ds,
            "norms": snorms,
            "photos_block": [str(file_name or "")] if file_name else [],
        })

    topic_folder_link = (idx or {}).get("topic_folder_link", "")

    # 7a. Photo-only response (numbered format)
    if not make_act:
        msg = _p6h_format_photo_numbered_response(
            vision, defects, grouped, global_norms, file_name or "photo",
            photo_no=photo_no, voice_ctx=voice_ctx,
        )
        # Save compact summary to memory
        _p6h_save_summary_to_memory(chat_id, topic_id, {
            "folder": voice_ctx.get("folder_hint") or "",
            "object": voice_ctx.get("object_hint") or "",
            "date": voice_ctx.get("visit_date_hint") or "",
            "owner_directives": (voice_ctx.get("explicit_include") or []) + (voice_ctx.get("explicit_exclude") or []),
            "defect_brief": [str(d.get("title") or d.get("description") or "")[:200] for d in defects][:8],
            "pdf_link": "",
            "docx_link": "",
            "status": "PHOTO_REPORT_DONE",
        })
        return {
            "ok": True, "handled": True, "kind": "technadzor_p6h_photo",
            "state": "DONE",
            "message": msg,
            "history": "P6H_PHOTO_REPORT_PHOTO{}_DEFECTS_{}_NORMS_{}".format(photo_no, len(defects), len(global_norms)),
        }

    # 7b. Build act: DOCX (service) + PDF (client topic root) + upload
    ts = _p6h_dt.now().strftime("%Y%m%d_%H%M%S")
    safe_tid = (str(task_id or ts)[:8] or ts).replace("/", "_").replace("\\", "_")
    docx_local = _P6H_OUTDIR / f"P6H_TNZ_ACT_DRAFT__{safe_tid}_{ts}.docx"
    pdf_local = _P6H_OUTDIR / f"АКТ_ОСМОТРА__{safe_tid}_{ts}.pdf"

    recs = []
    cons = []
    for d in defects:
        r = str(d.get("recommended_fix") or d.get("fix") or d.get("recommendation") or "").strip()
        if r:
            recs.append(r[:300])
        c = str(d.get("consequence") or d.get("risk") or d.get("why") or "").strip()
        if c:
            cons.append(c[:300])

    # 8-column violations table per spec
    violations_8 = []
    for sec_title, ds in grouped:
        for d in ds:
            num = len(violations_8) + 1
            ph = str(d.get("file_name") or file_name or "")
            place_node = sec_title
            violation = str(d.get("title") or d.get("description") or sec_title)[:200]
            consequence = str(d.get("consequence") or d.get("risk") or "")[:200]
            fix = str(d.get("recommended_fix") or d.get("fix") or "")[:200]
            # find norm for this defect's section
            norm_id = ""
            for s in sections_payload:
                if s["title"] == sec_title and s["norms"]:
                    norm_id = s["norms"][0].get("norm_id", "") or ""
                    break
            status = ""
            conf = (vision or {}).get("confidence", "LOW")
            if conf == "HIGH" and norm_id:
                status = "CONFIRMED_BY_PHOTO"
            elif conf in ("HIGH", "MEDIUM") and not norm_id:
                status = "NORM_NOT_CONFIRMED"
            elif conf == "MEDIUM":
                status = "PARTIAL_BY_PHOTO"
            else:
                status = "NEEDS_OWNER_CLARIFICATION"
            violations_8.append((num, ph, place_node, violation, consequence, fix, norm_id or "норма не подтверждена", status))

    payload = {
        "act_number": _p6h_human_act_number(task_id),
        "date_str": _p6h_dt.now().strftime("%d.%m.%Y"),
        "place": place or (voice_ctx.get("folder_hint") or "место уточняется по запросу владельца"),
        "object_descr": object_descr or (voice_ctx.get("object_hint") or "объект уточняется по запросу владельца"),
        "method": "визуальный неразрушающий контроль с выездом на объект",
        "performer": "",
        "specialist": "Кузнецов Илья Владимирович",
        "photos_link": topic_folder_link or "",
        "general_purpose": (
            "Осмотр выполнен методом визуального неразрушающего контроля. "
            "Цель осмотра — выявление фактически наблюдаемых дефектов, определение "
            "рекомендаций к устранению и возможных последствий для заказчика."
        ),
        "sections": sections_payload,
        "recommendations": (recs[:20] if recs else
                             ["Привести выявленные узлы и покрытия к нормативному состоянию по соответствующим СП/ГОСТ"]),
        "consequences": (cons[:10] if cons else
                          ["Снижение несущей способности и эксплуатационной надёжности конструкций"]),
        # 8-col rich table
        "violations_table_8col": violations_8[:30],
        # 3-col simple table (back-compat for DOCX/PDF builders that use violations_table)
        "violations_table": [(v, n, p) for (_no, p, _pl, v, _co, _fi, n, _st) in violations_8[:30]],
    }

    docx_ok = False
    pdf_ok = False
    try:
        _p6h_build_docx_act(payload, docx_local)
        docx_ok = True
    except Exception:
        _P6H2_LOG.exception("P6H2_DOCX_BUILD_FAIL")
    try:
        _p6h_build_pdf_act(payload, pdf_local)
        pdf_ok = True
    except Exception:
        _P6H2_LOG.exception("P6H2_PDF_BUILD_FAIL")

    drive_docx = None
    drive_pdf = None
    if docx_ok or pdf_ok:
        try:
            from core import technadzor_drive_index as _tdi2
            if docx_ok:
                drive_docx = _tdi2.upload_to_service_subfolder(
                    docx_local, docx_local.name, str(chat_id), int(topic_id), subfolder="_drafts",
                )
            if pdf_ok:
                target_folder = None
                # Owner explicitly named a client folder for the final PDF?
                fh = voice_ctx.get("folder_hint", "") or ""
                if fh:
                    # Only allow target placement if folder is explicitly client-facing per voice
                    if voice_ctx.get("client_facing") is True:
                        from core.technadzor_drive_index import is_system_folder as _is_sys
                        if not _is_sys(fh):
                            target_folder = fh
                drive_pdf = _tdi2.upload_client_pdf_to_folder(
                    pdf_local, pdf_local.name, str(chat_id), int(topic_id),
                    target_folder_name=target_folder,
                )
        except Exception:
            _P6H2_LOG.exception("P6H2_UPLOAD_FAIL")

    pdf_link = (drive_pdf or {}).get("link", "") if drive_pdf else ""
    docx_link = (drive_docx or {}).get("link", "") if drive_docx else ""

    msg_lines = ["Акт сформирован"]
    if pdf_link:
        msg_lines.append(f"PDF: {pdf_link}")
    elif pdf_ok:
        msg_lines.append("PDF: подготовлен локально, загрузка на Drive не выполнена — Telegram fallback в следующей итерации")
    else:
        msg_lines.append("PDF: ошибка генерации — повторите позже")
    if docx_link:
        msg_lines.append(f"DOCX (черновик, служебно): {docx_link}")
    if topic_folder_link:
        msg_lines.append(f"Фото: {topic_folder_link}")
    msg_lines.append("Норма: " + ("подтверждена" if global_norms else "не подтверждена"))

    # Memory save
    _p6h_save_summary_to_memory(chat_id, topic_id, {
        "folder": voice_ctx.get("folder_hint") or "",
        "object": voice_ctx.get("object_hint") or "",
        "date": voice_ctx.get("visit_date_hint") or "",
        "owner_directives": (voice_ctx.get("explicit_include") or []) + (voice_ctx.get("explicit_exclude") or []),
        "defect_brief": [str(d.get("title") or d.get("description") or "")[:200] for d in defects][:8],
        "pdf_link": pdf_link,
        "docx_link": docx_link,
        "status": "ACT_DONE" if pdf_link else ("ACT_PARTIAL" if (pdf_ok or docx_ok) else "ACT_FAIL"),
    })

    return {
        "ok": True if (pdf_ok or docx_ok) else False,
        "handled": True,
        "kind": "technadzor_p6h_act",
        "state": "DONE" if pdf_link else "AWAITING_CONFIRMATION",
        "artifact_path": str(pdf_local if pdf_ok else docx_local),
        "extra_artifact_path": str(docx_local if docx_ok else ""),
        "drive_link": pdf_link or docx_link or "",
        "message": _p6h_clean_text("\n".join(msg_lines), 4000),
        "history": "P6H_ACT_DOCX_{}_PDF_{}_DRIVE_PDF_{}".format(
            "OK" if docx_ok else "FAIL",
            "OK" if pdf_ok else "FAIL",
            "OK" if pdf_link else "FAIL",
        ),
    }


_P6H2_LOG.info("P6H_TOPIC5_PHOTO_NUMBER_DEFECT_NORM_CLARIFICATION_LOGIC_20260504_INSTALLED")
_P6H2_LOG.info("P6H_TOPIC5_VOICE_LIVE_DIALOG_CLARIFICATION_GATE_20260504_INSTALLED")
# === END_P6H_PART_2 ===


# === P6H_PART_3: TOPIC5_OBJECT_REGISTRY_INSPECTION_CHAIN_20260504 ===
# Final wiring of object registry into the photo pipeline.
# Replaces _p6h_process_photo_async with a registry-aware version that:
#   • Identifies object_id from voice/Drive/file_name signals.
#   • Loads existing card and prior inspection_chain.
#   • Detects visit_mode (initial / repeat / extension / description_only).
#   • Adds clarification questions when object_id is ambiguous OR when voice/Vision conflict.
#   • For follow-up acts: carries forward open_items with statuses
#     (УСТРАНЕНО / УСТРАНЕНО ЧАСТИЧНО / НЕ УСТРАНЕНО / ТРЕБУЕТ УТОЧНЕНИЯ / ...).
#   • After the act is built, appends an inspection record to the chain
#     (server JSON + memory + timeline).
#   • Never writes registry/system files into client-facing folders.

try:
    from core import technadzor_object_registry as _p6h3_reg
except Exception as _exc_reg:
    _p6h3_reg = None
    _P6H2_LOG.warning("P6H3_REGISTRY_IMPORT_FAIL: %s", _exc_reg)


async def _p6h_process_photo_async(file_path, file_name, task_id, chat_id, topic_id, user_text="", make_act=False, place="", object_descr=""):
    # 1. Drive index
    try:
        from core import technadzor_drive_index as _tdi
        idx = _tdi.build_technadzor_template_index(str(chat_id), int(topic_id), force=False)
    except Exception:
        idx = {}

    # 2. Voice context
    voice_ctx = _p6h_parse_voice_instruction(user_text or "")

    # 3. Object registry — identify object_id
    object_id = ""
    object_card = None
    visit_mode = "initial"
    derive_sources = {}
    if _p6h3_reg is not None:
        try:
            object_id, derive_sources = _p6h3_reg.derive_object_id_from_context(
                voice_ctx, idx, file_path or "", file_name or "",
            )
            if object_id:
                object_card = _p6h3_reg.load_object(object_id)
            visit_mode = _p6h3_reg.detect_visit_mode(object_card, voice_ctx)
        except Exception:
            _P6H2_LOG.exception("P6H3_REGISTRY_DERIVE_FAIL")

    # 4. Vision
    vision, vstatus = await _p6f_tnz_vision_via_openrouter(file_path)
    if vstatus == "FAIL" or vision is None:
        return {
            "ok": False, "handled": True, "kind": "technadzor_p6h_photo",
            "state": "WAITING_CLARIFICATION",
            "message": "Не удалось проанализировать фото через Vision. "
                       "Пришли крупнее или короткое описание дефекта текстом",
            "history": "P6H_VISION_FAIL",
        }

    summary = (vision.get("summary") or "").strip() if isinstance(vision, dict) else ""
    defects = (vision.get("defects") or []) if isinstance(vision, dict) else []
    photo_no = 1
    for d in defects:
        if isinstance(d, dict):
            d.setdefault("photo_no", photo_no)
            d.setdefault("file_name", file_name or "")

    grouped = _p6h_group_defects_by_section(defects)

    # 5. Clarification gate (registry + Vision + voice/Vision conflict)
    should_wait_basic, basic_questions = _p6h_should_wait_for_clarification(vision, defects, voice_ctx, idx)
    questions = list(basic_questions)

    # Object identity question — only if no folder/object hints AND no object_id derived
    if not object_id and not voice_ctx.get("folder_hint") and not voice_ctx.get("object_hint"):
        existing_summaries = []
        if _p6h3_reg is not None:
            try:
                existing_summaries = _p6h3_reg.list_object_summaries()
            except Exception:
                existing_summaries = []
        if existing_summaries:
            names = ", ".join(
                f"«{e.get('object_name') or e.get('object_id')}»"
                for e in existing_summaries[:5]
            )
            questions.append(
                f"Это новый объект или продолжение одного из существующих ({names})? "
                "Уточни — иначе не смогу привязать к истории объекта"
            )
        else:
            questions.append(
                "Назови объект (адрес / папка / имя), чтобы я завёл карточку и привязал акт"
            )

    # Voice/Vision conflict
    if _p6h3_reg is not None and voice_ctx.get("transcript"):
        try:
            conflict_flags = _p6h3_reg.detect_voice_vision_conflict(voice_ctx, grouped)
        except Exception:
            conflict_flags = []
        for cf in conflict_flags:
            if cf not in questions:
                questions.append(cf)
    else:
        conflict_flags = []

    # If photo-only mode AND we have questions — return WAITING_CLARIFICATION
    if questions and not make_act:
        return {
            "ok": True, "handled": True, "kind": "technadzor_p6h_clarify",
            "state": "WAITING_CLARIFICATION",
            "message": _p6h_clean_text(
                "Технадзор topic_5 — нужны уточнения перед разбором:\n\n"
                + "\n".join(f"— {q}" for q in questions),
                3000,
            ),
            "history": "P6H_CLARIFY_WITH_REGISTRY:{}_visit_{}".format(len(questions), visit_mode),
        }

    # 6. Norms
    haystack = " ".join([
        summary, voice_ctx.get("transcript", "") or "",
    ] + [str(d.get("title", "")) + " " + str(d.get("description", "")) for d in defects])
    global_norms = _p6h_norms_for_haystack(haystack)

    # 7. Sections payload
    sections_payload = []
    for sec_title, ds in grouped:
        defect_texts = [str(d.get("title", "")) + " " + str(d.get("description", "")) for d in ds]
        snorms = _p6h_norms_for_section(sec_title, defect_texts)
        sections_payload.append({
            "title": sec_title,
            "defects": ds,
            "norms": snorms,
            "photos_block": [str(file_name or "")] if file_name else [],
        })

    topic_folder_link = (idx or {}).get("topic_folder_link", "")

    # 7a. Photo-only response (numbered) ── append registry context if known
    if not make_act:
        msg = _p6h_format_photo_numbered_response(
            vision, defects, grouped, global_norms, file_name or "photo",
            photo_no=photo_no, voice_ctx=voice_ctx,
        )
        if object_id and object_card:
            registry_tail = (
                "\n\nКарточка объекта: "
                f"{object_card.get('object_name') or object_id}"
                f" (осмотров в истории: {len(object_card.get('inspection_chain') or [])})"
            )
            msg = (msg + registry_tail).strip()

        # Update memory summary
        try:
            _p6h_save_summary_to_memory(chat_id, topic_id, {
                "folder": voice_ctx.get("folder_hint") or "",
                "object": voice_ctx.get("object_hint") or object_id or "",
                "date": voice_ctx.get("visit_date_hint") or "",
                "owner_directives": (voice_ctx.get("explicit_include") or []) + (voice_ctx.get("explicit_exclude") or []),
                "defect_brief": [str(d.get("title") or d.get("description") or "")[:200] for d in defects][:8],
                "pdf_link": "",
                "docx_link": "",
                "status": f"PHOTO_REPORT_DONE:visit={visit_mode}",
            })
        except Exception:
            pass

        return {
            "ok": True, "handled": True, "kind": "technadzor_p6h_photo",
            "state": "DONE",
            "message": msg,
            "history": "P6H_PHOTO_REPORT_PHOTO{}_DEFECTS_{}_NORMS_{}_VISIT_{}".format(
                photo_no, len(defects), len(global_norms), visit_mode,
            ),
        }

    # 7b. Act build — DOCX (service _drafts) + PDF (client topic root or named client folder)
    ts = _p6h_dt.now().strftime("%Y%m%d_%H%M%S")
    safe_tid = (str(task_id or ts)[:8] or ts).replace("/", "_").replace("\\", "_")
    docx_local = _P6H_OUTDIR / f"P6H_TNZ_ACT_DRAFT__{safe_tid}_{ts}.docx"
    pdf_local = _P6H_OUTDIR / f"АКТ_ОСМОТРА__{safe_tid}_{ts}.pdf"

    # Recommendations / consequences
    recs, cons = [], []
    for d in defects:
        r = str(d.get("recommended_fix") or d.get("fix") or d.get("recommendation") or "").strip()
        if r:
            recs.append(r[:300])
        c = str(d.get("consequence") or d.get("risk") or d.get("why") or "").strip()
        if c:
            cons.append(c[:300])

    # Carry forward open_items for follow-up acts
    carried_open_items = []
    if _p6h3_reg is not None and visit_mode in ("repeat", "extension"):
        try:
            carried_open_items = _p6h3_reg.carry_forward_open_items(object_card, defects)
        except Exception:
            carried_open_items = []

    # Build follow-up section payload (added before regular sections in act)
    if carried_open_items:
        follow_section = {
            "title": "Состояние ранее выданных замечаний (повторный осмотр)",
            "defects": [
                {
                    "title": f"[{it.get('status','?')}] {it.get('title','') or it.get('description','')}",
                    "description": (it.get("description", "") or "")[:300]
                                    + (f" (из акта № {it.get('from_act_no')})" if it.get("from_act_no") else ""),
                }
                for it in carried_open_items
            ],
            "norms": [],
            "photos_block": [],
        }
        sections_payload.insert(0, follow_section)

    # 8-column violations table
    violations_8 = []
    for sec_title, ds in grouped:
        for d in ds:
            num = len(violations_8) + 1
            ph = str(d.get("file_name") or file_name or "")
            place_node = sec_title
            violation = str(d.get("title") or d.get("description") or sec_title)[:200]
            consequence = str(d.get("consequence") or d.get("risk") or "")[:200]
            fix = str(d.get("recommended_fix") or d.get("fix") or "")[:200]
            norm_id = ""
            for s in sections_payload:
                if s["title"] == sec_title and s["norms"]:
                    norm_id = s["norms"][0].get("norm_id", "") or ""
                    break
            conf = (vision or {}).get("confidence", "LOW")
            if conf == "HIGH" and norm_id:
                status = "CONFIRMED_BY_PHOTO"
            elif conf in ("HIGH", "MEDIUM") and not norm_id:
                status = "NORM_NOT_CONFIRMED"
            elif conf == "MEDIUM":
                status = "PARTIAL_BY_PHOTO"
            else:
                status = "NEEDS_OWNER_CLARIFICATION"
            violations_8.append((num, ph, place_node, violation, consequence, fix, norm_id or "норма не подтверждена", status))

    # Determine general_purpose by visit_mode
    if visit_mode == "repeat":
        gen_purpose = (
            "Повторный осмотр выполнен в развитие предыдущих актов по объекту. "
            "Цель — фиксация выполненных исправлений и неустранённых замечаний, "
            "определение рекомендаций к доведению конструктивных решений до нормативного состояния."
        )
    elif visit_mode == "extension":
        gen_purpose = (
            "Дополнение к предыдущему акту по объекту. "
            "Фиксируются дополнительные материалы и замечания, обнаруженные после основного осмотра."
        )
    else:
        gen_purpose = (
            "Осмотр выполнен методом визуального неразрушающего контроля. "
            "Цель осмотра — выявление фактически наблюдаемых дефектов, определение "
            "рекомендаций к устранению и возможных последствий для заказчика."
        )

    # Act number — for follow-up, hint at parent
    act_number = _p6h_human_act_number(task_id)
    if visit_mode == "repeat" and object_card and (object_card.get("last_act_no") or ""):
        act_number = act_number + f" (в развитие акта № {object_card['last_act_no']})"

    payload = {
        "act_number": act_number,
        "date_str": _p6h_dt.now().strftime("%d.%m.%Y"),
        "place": place or (voice_ctx.get("folder_hint") or "место уточняется по запросу владельца"),
        "object_descr": object_descr or (voice_ctx.get("object_hint") or
                                          (object_card.get("object_name") if object_card else "")
                                          or "объект уточняется по запросу владельца"),
        "method": "визуальный неразрушающий контроль с выездом на объект",
        "performer": "",
        "specialist": "Кузнецов Илья Владимирович",
        "photos_link": topic_folder_link or "",
        "general_purpose": gen_purpose,
        "sections": sections_payload,
        "recommendations": (recs[:20] if recs else
                             ["Привести выявленные узлы и покрытия к нормативному состоянию по соответствующим СП/ГОСТ"]),
        "consequences": (cons[:10] if cons else
                          ["Снижение несущей способности и эксплуатационной надёжности конструкций"]),
        "violations_table_8col": violations_8[:30],
        "violations_table": [(v, n, p) for (_no, p, _pl, v, _co, _fi, n, _st) in violations_8[:30]],
    }

    # Build files
    docx_ok = pdf_ok = False
    try:
        _p6h_build_docx_act(payload, docx_local)
        docx_ok = True
    except Exception:
        _P6H2_LOG.exception("P6H3_DOCX_BUILD_FAIL")
    try:
        _p6h_build_pdf_act(payload, pdf_local)
        pdf_ok = True
    except Exception:
        _P6H2_LOG.exception("P6H3_PDF_BUILD_FAIL")

    # Upload — DOCX → _drafts (system).  PDF → client folder if explicitly named, else topic root.
    drive_docx = drive_pdf = None
    if docx_ok or pdf_ok:
        try:
            from core import technadzor_drive_index as _tdi2
            if docx_ok:
                drive_docx = _tdi2.upload_to_service_subfolder(
                    docx_local, docx_local.name, str(chat_id), int(topic_id), subfolder="_drafts",
                )
            if pdf_ok:
                target_folder = None
                fh = voice_ctx.get("folder_hint", "") or ""
                if fh and voice_ctx.get("client_facing") is True:
                    from core.technadzor_drive_index import is_system_folder as _is_sys
                    if not _is_sys(fh):
                        target_folder = fh
                drive_pdf = _tdi2.upload_client_pdf_to_folder(
                    pdf_local, pdf_local.name, str(chat_id), int(topic_id),
                    target_folder_name=target_folder,
                )
        except Exception:
            _P6H2_LOG.exception("P6H3_UPLOAD_FAIL")

    pdf_link = (drive_pdf or {}).get("link", "") if drive_pdf else ""
    docx_link = (drive_docx or {}).get("link", "") if drive_docx else ""

    # Record inspection in chain
    if _p6h3_reg is not None and (object_id or pdf_ok or docx_ok):
        try:
            # If we still don't have object_id by here (e.g., act forced through),
            # create a synthetic stable id from file/date so chain still records.
            if not object_id:
                object_id = _p6h3_reg._slug(
                    voice_ctx.get("object_hint") or
                    voice_ctx.get("folder_hint") or
                    (file_name.rsplit(".", 1)[0] if file_name else "") or
                    f"obj_{ts}"
                )
            new_open = []
            for sec_title, ds in grouped:
                for d in ds:
                    new_open.append({
                        "title": str(d.get("title") or "")[:200],
                        "description": str(d.get("description") or "")[:300],
                        "section": sec_title,
                        "act_no": payload["act_number"],
                    })
            _p6h3_reg.record_inspection(
                object_id, str(chat_id),
                act_no=payload["act_number"],
                date_str=payload["date_str"],
                mode=visit_mode,
                pdf_link=pdf_link,
                docx_link=docx_link,
                source_photo_folder=voice_ctx.get("folder_hint") or "",
                findings=[{"section": s["title"],
                            "defects": s["defects"][:10],
                            "norms": [n.get("norm_id") for n in (s.get("norms") or [])]}
                           for s in sections_payload],
                open_items=new_open,
                closed_items=[
                    {"title": it.get("title"), "from_act_no": it.get("from_act_no")}
                    for it in carried_open_items if it.get("status") == "УСТРАНЕНО"
                ],
                new_items=new_open,
                owner_observation=voice_ctx.get("transcript", "")[:1000],
                conflict_flags=conflict_flags or [],
                object_name=(voice_ctx.get("object_hint") or
                              (object_card.get("object_name") if object_card else "") or ""),
                object_folder_url=topic_folder_link or "",
            )
        except Exception:
            _P6H2_LOG.exception("P6H3_RECORD_INSPECTION_FAIL")

    msg_lines = ["Акт сформирован"]
    if pdf_link:
        msg_lines.append(f"PDF: {pdf_link}")
    elif pdf_ok:
        msg_lines.append("PDF: подготовлен локально, загрузка на Drive не выполнена — Telegram fallback")
    else:
        msg_lines.append("PDF: ошибка генерации — повторите позже")
    if docx_link:
        msg_lines.append(f"DOCX (черновик, служебно): {docx_link}")
    if topic_folder_link:
        msg_lines.append(f"Фото: {topic_folder_link}")
    msg_lines.append("Норма: " + ("подтверждена" if global_norms else "не подтверждена"))
    if visit_mode == "repeat":
        msg_lines.append(f"Тип осмотра: повторный (история объекта: {len((object_card or {}).get('inspection_chain') or []) + 1} записей)")
    elif visit_mode == "extension":
        msg_lines.append("Тип осмотра: дополнение к предыдущему акту")
    else:
        msg_lines.append("Тип осмотра: первичный")

    # Memory summary
    try:
        _p6h_save_summary_to_memory(chat_id, topic_id, {
            "folder": voice_ctx.get("folder_hint") or "",
            "object": voice_ctx.get("object_hint") or object_id or "",
            "date": payload["date_str"],
            "owner_directives": (voice_ctx.get("explicit_include") or []) + (voice_ctx.get("explicit_exclude") or []),
            "defect_brief": [str(d.get("title") or d.get("description") or "")[:200] for d in defects][:8],
            "pdf_link": pdf_link,
            "docx_link": docx_link,
            "status": ("ACT_DONE" if pdf_link else ("ACT_PARTIAL" if (pdf_ok or docx_ok) else "ACT_FAIL"))
                       + f":visit={visit_mode}",
        })
    except Exception:
        pass

    return {
        "ok": True if (pdf_ok or docx_ok) else False,
        "handled": True,
        "kind": "technadzor_p6h_act",
        "state": "DONE" if pdf_link else "AWAITING_CONFIRMATION",
        "artifact_path": str(pdf_local if pdf_ok else docx_local),
        "extra_artifact_path": str(docx_local if docx_ok else ""),
        "drive_link": pdf_link or docx_link or "",
        "message": _p6h_clean_text("\n".join(msg_lines), 4000),
        "history": "P6H_ACT_VISIT_{}_DOCX_{}_PDF_{}_DRIVE_PDF_{}_OPEN_CARRIED_{}".format(
            visit_mode,
            "OK" if docx_ok else "FAIL",
            "OK" if pdf_ok else "FAIL",
            "OK" if pdf_link else "FAIL",
            len(carried_open_items),
        ),
    }


_P6H2_LOG.info("P6H_TOPIC5_OBJECT_REGISTRY_INSPECTION_CHAIN_20260504_INSTALLED")
# === END_P6H_PART_3 ===

# ─── P6H_EXTERNAL_VISION_GUARD_V1 ──────────────────────────────────────────
# CANON: TECHNADZOR_DOMAIN_LOGIC_CANON_V2 §33
# EXTERNAL_PHOTO_ANALYSIS_ALLOWED = False by default
# Vision запускается только после явного разрешения владельца

_P6H_EXTERNAL_VISION_ALLOWED = False

_p6h_vision_orig = _p6f_tnz_vision_via_openrouter  # сохраняем оригинал

async def _p6f_tnz_vision_via_openrouter(local_path):  # noqa: F811
    if not _P6H_EXTERNAL_VISION_ALLOWED:
        _P6H2_LOG.warning("EXTERNAL_VISION_BLOCKED path=%s EXTERNAL_PHOTO_ANALYSIS_ALLOWED=False", local_path)
        return {}, "EXTERNAL_PHOTO_ANALYSIS_BLOCKED"
    return await _p6h_vision_orig(local_path)

def _p6h_allow_external_vision():
    global _P6H_EXTERNAL_VISION_ALLOWED
    _P6H_EXTERNAL_VISION_ALLOWED = True
    _P6H2_LOG.info("EXTERNAL_VISION_ALLOWED_SET owner_approved=True")

_P6H2_LOG.info("P6H_EXTERNAL_VISION_GUARD_V1_INSTALLED allowed=%s", _P6H_EXTERNAL_VISION_ALLOWED)
# ─── END P6H_EXTERNAL_VISION_GUARD_V1 ──────────────────────────────────────

# ─── P6H_PART_4_VISIT_BUFFER_V1 ────────────────────────────────────────────
# CANON: TECHNADZOR_DOMAIN_LOGIC_CANON_V2
# ActiveTechnadzorFolder / VisitMaterial / visit_buffer_add / visit_buffer_flush

import os as _p6h4_os
import json as _p6h4_json
import time as _p6h4_time
import logging as _p6h4_logging

_P6H4_LOG = _p6h4_logging.getLogger("task_worker")

_P6H4_BASE = _p6h4_os.path.join(
    _p6h4_os.path.dirname(_p6h4_os.path.dirname(_p6h4_os.path.abspath(__file__))),
    "data", "technadzor"
)


def _p6h4_ensure():
    _p6h4_os.makedirs(_P6H4_BASE, exist_ok=True)


def _p6h4_buf_path(chat_id, topic_id):
    _p6h4_ensure()
    return _p6h4_os.path.join(_P6H4_BASE, f"buf_{chat_id}_{topic_id}.json")


def _p6h4_folder_path(chat_id, topic_id):
    _p6h4_ensure()
    return _p6h4_os.path.join(_P6H4_BASE, f"active_folder_{chat_id}_{topic_id}.json")


def visit_buffer_add(chat_id, topic_id, material: dict) -> int:
    """Append VisitMaterial to persistent buffer. Returns new total count."""
    path = _p6h4_buf_path(str(chat_id), int(topic_id))
    try:
        with open(path, "r", encoding="utf-8") as _f:
            buf = _p6h4_json.load(_f)
    except Exception:
        buf = {"materials": [], "created_at": _p6h4_time.time()}
    if "material_id" not in material:
        material["material_id"] = f"{int(_p6h4_time.time() * 1000)}"
    material.setdefault("added_at", _p6h4_time.time())
    buf["materials"].append(material)
    buf["updated_at"] = _p6h4_time.time()
    with open(path, "w", encoding="utf-8") as _f:
        _p6h4_json.dump(buf, _f, ensure_ascii=False, indent=2)
    count = len(buf["materials"])
    _P6H4_LOG.info("P6H4_VISIT_BUFFER_ADD chat=%s topic=%s count=%s", chat_id, topic_id, count)
    return count


def visit_buffer_flush(chat_id, topic_id) -> list:
    """Return all buffered VisitMaterials and clear buffer."""
    path = _p6h4_buf_path(str(chat_id), int(topic_id))
    try:
        with open(path, "r", encoding="utf-8") as _f:
            buf = _p6h4_json.load(_f)
        materials = buf.get("materials", [])
        _p6h4_os.remove(path)
        _P6H4_LOG.info("P6H4_VISIT_BUFFER_FLUSH chat=%s topic=%s count=%s", chat_id, topic_id, len(materials))
        return materials
    except Exception:
        _P6H4_LOG.info("P6H4_VISIT_BUFFER_FLUSH_EMPTY chat=%s topic=%s", chat_id, topic_id)
        return []


def visit_buffer_count(chat_id, topic_id) -> int:
    path = _p6h4_buf_path(str(chat_id), int(topic_id))
    try:
        with open(path, "r", encoding="utf-8") as _f:
            buf = _p6h4_json.load(_f)
        return len(buf.get("materials", []))
    except Exception:
        return 0


def set_active_folder(chat_id, topic_id, folder_data: dict):
    path = _p6h4_folder_path(str(chat_id), int(topic_id))
    folder_data["set_at"] = _p6h4_time.time()
    with open(path, "w", encoding="utf-8") as _f:
        _p6h4_json.dump(folder_data, _f, ensure_ascii=False, indent=2)
    _P6H4_LOG.info(
        "P6H4_ACTIVE_FOLDER_SET chat=%s topic=%s name=%s",
        chat_id, topic_id, folder_data.get("folder_name", "?"),
    )


def get_active_folder(chat_id, topic_id) -> dict:
    path = _p6h4_folder_path(str(chat_id), int(topic_id))
    try:
        with open(path, "r", encoding="utf-8") as _f:
            return _p6h4_json.load(_f)
    except Exception:
        return {}


def process_drive_folder_batch(chat_id, topic_id, folder_id: str, folder_name: str = "") -> int:
    """Scan Drive folder, add all files as VisitMaterials. Returns count added."""
    added = 0
    try:
        from core.topic_drive_oauth import get_drive_service as _p6h4_get_drive
        svc = _p6h4_get_drive(chat_id=str(chat_id), topic_id=int(topic_id))
        items = svc.files().list(
            q=f"'{folder_id}' in parents and trashed=false",
            fields="files(id,name,mimeType,webViewLink)",
            pageSize=200,
        ).execute().get("files", [])
        for item in items:
            mime = item.get("mimeType", "")
            if "folder" in mime:
                continue
            ftype = "PHOTO" if mime.startswith("image/") else "PDF" if "pdf" in mime else "OTHER"
            mat = {
                "source": "DRIVE",
                "file_type": ftype,
                "file_name": item.get("name", ""),
                "drive_url": item.get("webViewLink", ""),
                "drive_file_id": item.get("id", ""),
                "include_in_act": True,
                "include_in_report": True,
                "group_label": folder_name or "",
            }
            visit_buffer_add(str(chat_id), int(topic_id), mat)
            added += 1
        _P6H4_LOG.info("P6H4_DRIVE_FOLDER_BATCH chat=%s topic=%s folder=%s added=%s", chat_id, topic_id, folder_id, added)
    except Exception as _p6h4_batch_err:
        _P6H4_LOG.warning("P6H4_DRIVE_FOLDER_BATCH_ERR %s", _p6h4_batch_err)
    return added


_P6H4_LOG.info("P6H_PART_4_VISIT_BUFFER_V1_INSTALLED")
# ─── END P6H_PART_4_VISIT_BUFFER_V1 ─────────────────────────────────────────


# === P6H4TW_BATCH_TRIGGER_V1 ===
# FIX: original P6H_PART_4 hook in task_worker.py is after asyncio.run() and never fires.
# This wrapper hooks process_technadzor here (before asyncio.run()), intercepting all topic_5 calls.
# Handles: photo/file buffering, Drive folder batch load, visit buffer flush to process_technadzor.
# EXTERNAL_PHOTO_ANALYSIS_ALLOWED=False: no Vision without explicit owner permission.
import logging as _p6h4tw_v1_log_mod
import os as _p6h4tw_v1_os
import re as _p6h4tw_v1_re

_P6H4TW_V1_LOG = _p6h4tw_v1_log_mod.getLogger("task_worker")
_P6H4TW_V1_DRIVE_RE = _p6h4tw_v1_re.compile(
    r"https://drive\.google\.com/drive/folders/([A-Za-z0-9_-]+)"
)
_P6H4TW_V1_BATCH_TRIGGERS = (
    "загрузи все файлы из папки", "загрузи все файлы",
    "возьми файлы из папки", "прочитай папку",
    "обработай папку", "сделай разбор по папке", "сделай акт по папке",
    "разбор по папке", "акт по папке",
    "загрузи папку", "возьми из папки",
)
_P6H4TW_V1_BATCH_AND_FLUSH = (
    "сделай разбор по папке", "сделай акт по папке",
    "разбор по папке", "акт по папке",
)
_P6H4TW_V1_FLUSH_TRIGGERS = (
    "сделай акт", "собери акт", "сделай разбор", "сделай анализ",
    "собери разбор", "разберись", "сделай отчет", "сделай отчёт",
    "начни анализ", "сформируй акт",
)
_P6H4TW_V1_ACTIVE_FOLDER_TRIGGERS = (
    "работаем по этой папке", "установи папку", "активная папка это",
    "drive.google.com/drive/folders/",
)
_P6H4TW_V1_SHOW_FOLDER_TRIGGERS = (
    "покажи активную папку", "какая активная папка", "какая папка",
    "текущая папка", "покажи папку",
)


def _p6h4tw_v1_low(v):
    return str(v or "").lower().replace("ё", "е")


try:
    _p6h4tw_v1_orig = process_technadzor
    if not getattr(_p6h4tw_v1_orig, "_p6h4tw_v1_wrapped", False):

        def process_technadzor(text="", task_id="", chat_id="", topic_id=0, file_path="", file_name="", **kwargs):  # noqa: F811
            if int(topic_id or 0) != 5:
                return _p6h4tw_v1_orig(
                    text=text, task_id=task_id, chat_id=chat_id,
                    topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs
                )

            chat_str = str(chat_id)
            txt_low = _p6h4tw_v1_low(text)

            # ── Photo / Drive file → buffer (no Vision) ──────────────────────
            if file_path or file_name:
                fn = file_name or _p6h4tw_v1_os.path.basename(file_path or "")
                fn_low = fn.lower()
                is_photo = fn_low.endswith((".jpg", ".jpeg", ".png", ".webp", ".heic"))
                ftype = "PHOTO" if is_photo else "PDF" if fn_low.endswith(".pdf") else "DOCUMENT"
                material = {
                    "source": "telegram",
                    "file_type": ftype,
                    "file_name": fn,
                    "drive_file_id": "",
                    "drive_url": file_path or "",
                    "caption": text or "",
                    "include_in_act": True,
                    "include_in_report": True,
                }
                try:
                    count = visit_buffer_add(chat_str, 5, material)
                    _P6H4TW_V1_LOG.info(
                        "P6H4TW_V1_PHOTO_BUFFERED chat=%s count=%s fn=%s", chat_str, count, fn
                    )
                    return {
                        "ok": True,
                        "result_text": f"Добавлено в пакет ({count} шт.). Когда готово — скажи «сделай разбор».",
                        "history": "P6H4TW_V1_PHOTO_BUFFERED",
                    }
                except Exception as _e:
                    _P6H4TW_V1_LOG.warning("P6H4TW_V1_BUF_ADD_ERR %s", _e)
                    return _p6h4tw_v1_orig(
                        text=text, task_id=task_id, chat_id=chat_id,
                        topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs
                    )

            # ── Drive folder URL / set active folder ─────────────────────────
            m_drive = _P6H4TW_V1_DRIVE_RE.search(text or "")
            if m_drive or any(t in txt_low for t in _P6H4TW_V1_ACTIVE_FOLDER_TRIGGERS):
                folder_id = m_drive.group(1) if m_drive else ""
                folder_name = ""
                if folder_id:
                    try:
                        from core.topic_drive_oauth import get_drive_service as _p6h4tw_v1_gds
                        svc = _p6h4tw_v1_gds(chat_id=chat_str, topic_id=5)
                        folder_name = (
                            svc.files().get(fileId=folder_id, fields="name").execute().get("name", "")
                        )
                    except Exception:
                        pass
                try:
                    set_active_folder(chat_str, 5, {
                        "folder_id": folder_id,
                        "folder_name": folder_name,
                        "source_text": (text or "")[:500],
                    })
                    return {
                        "ok": True,
                        "result_text": f"Активная папка установлена: {folder_name or folder_id or '(из текста)'}.",
                        "history": "P6H4TW_V1_ACTIVE_FOLDER_SET",
                    }
                except Exception as _e:
                    _P6H4TW_V1_LOG.warning("P6H4TW_V1_SET_FOLDER_ERR %s", _e)

            # ── Show active folder ────────────────────────────────────────────
            if any(t in txt_low for t in _P6H4TW_V1_SHOW_FOLDER_TRIGGERS):
                try:
                    af = get_active_folder(chat_str, 5)
                    if af:
                        name = af.get("folder_name") or af.get("folder_id", "(нет имени)")
                        fid = af.get("folder_id", "")
                        link = f"https://drive.google.com/drive/folders/{fid}" if fid else "—"
                        msg = f"Активная папка: {name}\n{link}"
                    else:
                        msg = "Активная папка не установлена. Пришли ссылку на папку."
                    return {"ok": True, "result_text": msg, "history": "P6H4TW_V1_SHOW_FOLDER"}
                except Exception as _e:
                    _P6H4TW_V1_LOG.warning("P6H4TW_V1_SHOW_FOLDER_ERR %s", _e)

            # ── Drive folder batch load ───────────────────────────────────────
            if any(t in txt_low for t in _P6H4TW_V1_BATCH_TRIGGERS):
                try:
                    m_drive2 = _P6H4TW_V1_DRIVE_RE.search(text or "")
                    if m_drive2:
                        batch_fid = m_drive2.group(1)
                        batch_fname = ""
                        try:
                            from core.topic_drive_oauth import get_drive_service as _p6h4tw_v1_gds2
                            svc2 = _p6h4tw_v1_gds2(chat_id=chat_str, topic_id=5)
                            batch_fname = (
                                svc2.files().get(fileId=batch_fid, fields="name").execute().get("name", "")
                            )
                        except Exception:
                            pass
                    else:
                        af2 = get_active_folder(chat_str, 5) or {}
                        batch_fid = af2.get("folder_id", "")
                        batch_fname = af2.get("folder_name", "")
                    if not batch_fid:
                        return {
                            "ok": True,
                            "result_text": "Не найдена активная папка. Пришли ссылку на папку Drive.",
                            "history": "P6H4TW_V1_BATCH_NO_FOLDER",
                        }
                    added = process_drive_folder_batch(chat_str, 5, batch_fid, batch_fname)
                    _P6H4TW_V1_LOG.info(
                        "P6H4TW_V1_BATCH_LOADED chat=%s folder=%s added=%s", chat_str, batch_fid, added
                    )
                    do_flush = any(t in txt_low for t in _P6H4TW_V1_BATCH_AND_FLUSH)
                    if not do_flush:
                        return {
                            "ok": True,
                            "result_text": (
                                f"Принял. Файлы из папки добавлены в пакет выезда: {added} шт. "
                                "Vision не запускаю без разрешения владельца. Скажи «сделай разбор» когда готово."
                            ),
                            "history": "P6H4TW_V1_BATCH_LOADED",
                        }
                    txt_low = "сделай разбор"  # fall through to flush
                except Exception as _e:
                    _P6H4TW_V1_LOG.warning("P6H4TW_V1_BATCH_ERR %s", _e)

            # ── Flush buffer → process_technadzor ────────────────────────────
            if any(t in txt_low for t in _P6H4TW_V1_FLUSH_TRIGGERS):
                try:
                    count = visit_buffer_count(chat_str, 5)
                    if count == 0:
                        return {
                            "ok": True,
                            "result_text": "Буфер пуст — сначала пришли фото или файлы.",
                            "history": "P6H4TW_V1_FLUSH_EMPTY",
                        }
                    materials = visit_buffer_flush(chat_str, 5)
                    lines = ["Технический надзор. Акт по материалам выезда:", "VISIT_PACKAGE:"]
                    for i, m in enumerate(materials, 1):
                        fn2 = m.get("file_name", f"файл {i}")
                        url2 = m.get("drive_url", "")
                        note2 = (m.get("caption", "") or m.get("voice_comment", "") or "").strip()
                        line = f"  {i}. {fn2}"
                        if url2:
                            line += f" {url2}"
                        if note2:
                            line += f" — {note2}"
                        lines.append(line)
                    package_text = "\n".join(lines)
                    _P6H4TW_V1_LOG.info(
                        "P6H4TW_V1_FLUSH chat=%s count=%s", chat_str, len(materials)
                    )
                    return _p6h4tw_v1_orig(
                        text=package_text, task_id=task_id, chat_id=chat_id,
                        topic_id=topic_id, file_path="", file_name="", **kwargs
                    )
                except Exception as _e:
                    _P6H4TW_V1_LOG.warning("P6H4TW_V1_FLUSH_ERR %s", _e)

            # ── Default pass-through ──────────────────────────────────────────
            return _p6h4tw_v1_orig(
                text=text, task_id=task_id, chat_id=chat_id,
                topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs
            )

        process_technadzor._p6h4tw_v1_wrapped = True
        _p6h4tw_v1_orig._p6h4tw_v1_wrapped = True
        _P6H4TW_V1_LOG.info("P6H4TW_BATCH_TRIGGER_V1_INSTALLED")
except Exception as _p6h4tw_v1_err:
    _P6H4TW_V1_LOG.exception("P6H4TW_V1_INSTALL_ERR %s", _p6h4tw_v1_err)
# === END_P6H4TW_BATCH_TRIGGER_V1 ===

# === P6H4FD_FOLDER_DISCOVERY_V1 ===
import re as _p6h4fd_re
import logging as _p6h4fd_log_mod

_P6H4FD_LOG = _p6h4fd_log_mod.getLogger("technadzor_engine.p6h4fd")

_P6H4FD_FOLDER_INTENTS = (
    "папка",
    "новая папка",
    "создана папка",
    "создал папку",
    "обнаружь папку",
    "найди папку",
    "папка называется",
    "работаем по папке",
    "текущая папка",
    "прими папку",
    "туда складывать",
    "туда загружать",
    "все материалы туда",
)

# Matches "папк[у/а/и] <name>" — captures up to 60 chars until punctuation or end
_P6H4FD_NAME_RE = _p6h4fd_re.compile(
    r"(?:папк[уаи]|папка)\s+([А-Яа-яёЁA-Za-z0-9][А-Яа-яёЁA-Za-z0-9 \-_]{0,60}?)(?:[,.\n!?]|$)",
    _p6h4fd_re.IGNORECASE,
)


def _p6h4fd_extract_name(text: str) -> str:
    m = _P6H4FD_NAME_RE.search(text)
    if m:
        return m.group(1).strip()
    return ""


def _p6h4fd_norm(s: str) -> str:
    return " ".join(s.lower().replace("ё", "е").split())


def _p6h4fd_match_score(candidate: str, folder_name: str) -> int:
    c = _p6h4fd_norm(candidate)
    f = _p6h4fd_norm(folder_name)
    if not c:
        return 0
    if c == f:
        return 100
    if c in f or f in c:
        return 80
    cw = set(c.split())
    fw = set(f.split())
    overlap = len(cw & fw)
    return overlap * 10 if overlap else 0


try:
    _p6h4fd_orig_pt = process_technadzor
    if not getattr(_p6h4fd_orig_pt, "_p6h4fd_wrapped", False):

        def process_technadzor(  # noqa: F811
            text="", task_id="", chat_id="", topic_id=0,
            file_path="", file_name="", **kwargs
        ):
            if int(topic_id or 0) != 5:
                return _p6h4fd_orig_pt(
                    text=text, task_id=task_id, chat_id=chat_id,
                    topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs
                )

            txt_low = (text or "").lower().replace("ё", "е")
            if not any(t in txt_low for t in _P6H4FD_FOLDER_INTENTS):
                return _p6h4fd_orig_pt(
                    text=text, task_id=task_id, chat_id=chat_id,
                    topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs
                )

            # — folder/context intent: fresh Drive lookup by name —
            try:
                candidate = _p6h4fd_extract_name(text or "")
                chat_str = str(chat_id or "")
                _P6H4FD_LOG.info(
                    "P6H4FD_DISCOVERY_START candidate=%r chat=%s", candidate, chat_str
                )

                from core.technadzor_drive_index import _service as _p6h4fd_svc

                # system/container names — never a valid active folder result
                _P6H4FD_NEVER_RESULT = frozenset({
                    "technadzor", "технадзор", "topic_5", "_orchestra_work",
                    "_system", "_tmp", "_archive", "_drafts", "_templates", "_manifests",
                })

                def _p6h4fd_is_container(name):
                    return (name or "").strip().lower() in _P6H4FD_NEVER_RESULT

                svc = _p6h4fd_svc()

                def _p6h4fd_list_subfolders(parent_fid):
                    r = svc.files().list(
                        q=(
                            f"'{parent_fid}' in parents"
                            " and mimeType='application/vnd.google-apps.folder'"
                            " and trashed=false"
                        ),
                        fields="files(id,name,createdTime,modifiedTime)",
                        orderBy="createdTime desc",
                        pageSize=50,
                    ).execute()
                    return r.get("files", [])

                # Step A: search inside correct user ТЕХНАДЗОР root
                _TECHNADZOR_ROOT_FID = "1s2y5l2mJFTb7P90XVokErXYVzmoH-VtD"
                raw_a = _p6h4fd_list_subfolders(_TECHNADZOR_ROOT_FID)
                folders = [f for f in raw_a if not _p6h4fd_is_container(f.get("name", ""))]
                _P6H4FD_LOG.info("P6H4FD_ROOT_SEARCH count=%s", len(folders))

                # Step D: fallback — Drive-wide exact name search
                if not folders and candidate:
                    _safe = candidate.replace("'", "\\'")
                    _gr = svc.files().list(
                        q=(
                            f"name='{_safe}'"
                            " and mimeType='application/vnd.google-apps.folder'"
                            " and trashed=false"
                        ),
                        fields="files(id,name,createdTime,modifiedTime)",
                        orderBy="createdTime desc",
                        pageSize=10,
                    ).execute()
                    folders = [
                        f for f in _gr.get("files", [])
                        if not _p6h4fd_is_container(f.get("name", ""))
                    ]
                    _P6H4FD_LOG.info("P6H4FD_GLOBAL_SEARCH count=%s candidate=%r", len(folders), candidate)

                if not folders:
                    msg_nf = (
                        "Папку не нашёл"
                        + (f" «{candidate}»" if candidate else "")
                        + ". Укажи точное название или пришли ссылку."
                    )
                    return {
                        "ok": True,
                        "handled": True,
                        "state": "DONE",
                        "result_text": msg_nf,
                        "message": msg_nf,
                        "history": "P6H4FD_V1:NO_USER_FOLDERS",
                    }

                # exact match first, then fuzzy, then newest
                best = None
                best_score = 0
                if candidate:
                    for f in folders:
                        score = _p6h4fd_match_score(candidate, f.get("name", ""))
                        if score > best_score:
                            best_score = score
                            best = f

                if best is None or best_score == 0:
                    best = folders[0]

                fid = best["id"]
                fname = best.get("name", "")
                link = f"https://drive.google.com/drive/folders/{fid}"

                set_active_folder(chat_str, 5, {
                    "folder_id": fid,
                    "folder_name": fname,
                    "source_text": (text or "")[:500],
                })
                msg = f"Нашёл папку «{fname}» и установил её как активную.\n{link}"
                _P6H4FD_LOG.info(
                    "P6H4FD_SET_ACTIVE folder_id=%s name=%r score=%s chat=%s",
                    fid, fname, best_score, chat_str,
                )
                return {
                    "ok": True,
                    "handled": True,
                    "result_text": msg,
                    "message": msg,
                    "history": f"P6H4FD_V1:SET_ACTIVE:{fid}",
                }

            except Exception as _e:
                _P6H4FD_LOG.warning("P6H4FD_ERR %s", _e)
                # on error fall through to lower wrapper
                return _p6h4fd_orig_pt(
                    text=text, task_id=task_id, chat_id=chat_id,
                    topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs
                )

        process_technadzor._p6h4fd_wrapped = True
        _p6h4fd_orig_pt._p6h4fd_wrapped = True
        _P6H4FD_LOG.info("P6H4FD_FOLDER_DISCOVERY_V1_INSTALLED")
except Exception as _p6h4fd_err:
    _P6H4FD_LOG.exception("P6H4FD_INSTALL_ERR %s", _p6h4fd_err)
# === END_P6H4FD_FOLDER_DISCOVERY_V1 ===

# === PHOTO_RECOGNITION_TOPIC5_RUNTIME_BINDING_V1 ===
try:
    _photo_t5_orig_process_technadzor = process_technadzor

    def process_technadzor(
        text: str = "",
        task_id: str = "",
        chat_id: str = "",
        topic_id: int = 0,
        file_path: str = "",
        file_name: str = "",
        **kwargs,
    ):
        from core.photo_recognition_engine import is_image_file, process_photo_recognition

        clean_kwargs = dict(kwargs)
        for key in (
            "text", "raw_input", "task_id", "id", "chat_id", "topic_id",
            "file_path", "local_path", "file_name", "name",
        ):
            clean_kwargs.pop(key, None)

        raw_text = str(text or kwargs.get("raw_input") or "")
        fp = str(file_path or kwargs.get("local_path") or "")
        fn = str(file_name or kwargs.get("file_name") or kwargs.get("name") or "")
        resolved_task_id = str(task_id or kwargs.get("task_id") or kwargs.get("id") or "")
        resolved_chat_id = str(chat_id or kwargs.get("chat_id") or "")

        try:
            tid = int(topic_id or kwargs.get("topic_id") or 0)
        except Exception:
            tid = 0

        photo_result = None
        if tid == 5 and is_image_file(file_name=fn, file_path=fp):
            photo_result = process_photo_recognition(
                topic_id=5,
                file_name=fn,
                file_path=fp,
                owner_comment=raw_text,
                source="TELEGRAM",
            )

            if not raw_text.strip():
                return {
                    "ok": True,
                    "handled": True,
                    "status": "WAITING_CLARIFICATION",
                    "state": "WAITING_CLARIFICATION",
                    "kind": "technadzor_photo_material",
                    "message": "Фото принято как материал технадзора. Укажи, к какому замечанию или разделу его отнести",
                    "result_text": "Фото принято как материал технадзора. Укажи, к какому замечанию или разделу его отнести",
                    "photo_recognition": photo_result,
                    "history": "PHOTO_RECOGNITION_TOPIC5_RUNTIME_BINDING_V1:WAITING_OWNER_COMMENT",
                }

        result = _photo_t5_orig_process_technadzor(
            text=raw_text,
            task_id=resolved_task_id,
            chat_id=resolved_chat_id,
            topic_id=tid,
            file_path=fp,
            file_name=fn,
            **clean_kwargs,
        )

        if photo_result and isinstance(result, dict):
            result["photo_recognition"] = photo_result
            result["photo_recognition_status"] = photo_result.get("status")
            result["history"] = str(result.get("history") or "") + "|PHOTO_RECOGNITION_TOPIC5_RUNTIME_BINDING_V1"
            if "message" not in result and "result_text" not in result:
                result["message"] = "Фото принято и связано с технадзорным материалом"

        return result
except Exception:
    pass
# === END_PHOTO_RECOGNITION_TOPIC5_RUNTIME_BINDING_V1 ===


# === P7_TOPIC5_REPLY_VOICE_BINDING_V1 ===
# Binds Telegram text/voice reply to the VisitMaterial created from the replied photo/file.
import json as _p7_t5_json
import re as _p7_t5_re
import time as _p7_t5_time
from pathlib import Path as _p7_t5_Path

_P7_T5_ORIG_PROCESS_TECHNADZOR = process_technadzor
_P7_T5_DATA = _p7_t5_Path("/root/.areal-neva-core/data/technadzor")

def _p7_t5_s(v, limit=20000):
    try:
        return "" if v is None else str(v).strip()[:limit]
    except Exception:
        return ""

def _p7_t5_low(v):
    return _p7_t5_s(v).lower().replace("ё", "е")

def _p7_t5_task_get(task, key, default=None):
    if task is None:
        return default
    try:
        if isinstance(task, dict):
            return task.get(key, default)
        return task[key]
    except Exception:
        return getattr(task, key, default)

def _p7_t5_parse_payload(text):
    raw = _p7_t5_s(text, 50000)
    try:
        obj = _p7_t5_json.loads(raw)
        return obj if isinstance(obj, dict) else {"text": raw}
    except Exception:
        return {"text": raw}

def _p7_t5_buf_path(chat_id):
    _P7_T5_DATA.mkdir(parents=True, exist_ok=True)
    return _P7_T5_DATA / f"buf_{chat_id}_5.json"

def _p7_t5_active_folder(chat_id):
    p = _P7_T5_DATA / f"active_folder_{chat_id}_5.json"
    try:
        obj = _p7_t5_json.loads(p.read_text(encoding="utf-8"))
        return obj if isinstance(obj, dict) else {}
    except Exception:
        return {}

def _p7_t5_load_buf(chat_id):
    p = _p7_t5_buf_path(chat_id)
    try:
        obj = _p7_t5_json.loads(p.read_text(encoding="utf-8"))
        if isinstance(obj, dict):
            obj.setdefault("materials", [])
            return obj
    except Exception:
        pass
    return {"materials": [], "created_at": _p7_t5_time.time()}

def _p7_t5_save_buf(chat_id, buf):
    p = _p7_t5_buf_path(chat_id)
    buf["updated_at"] = _p7_t5_time.time()
    p.write_text(_p7_t5_json.dumps(buf, ensure_ascii=False, indent=2), encoding="utf-8")

def _p7_t5_msg_id_from_name(name):
    m = _p7_t5_re.search(r"_([0-9]{3,})\.[A-Za-z0-9]+$", _p7_t5_s(name))
    return m.group(1) if m else ""

def _p7_t5_add_material(chat_id, material):
    buf = _p7_t5_load_buf(chat_id)
    mid = _p7_t5_s(material.get("telegram_message_id"))
    fname = _p7_t5_s(material.get("file_name"))
    for old in buf.get("materials", []):
        if mid and _p7_t5_s(old.get("telegram_message_id")) == mid:
            old.update({k: v for k, v in material.items() if v not in ("", None)})
            _p7_t5_save_buf(chat_id, buf)
            return len(buf.get("materials", [])), old
        if fname and _p7_t5_s(old.get("file_name")) == fname:
            old.update({k: v for k, v in material.items() if v not in ("", None)})
            _p7_t5_save_buf(chat_id, buf)
            return len(buf.get("materials", [])), old
    buf["materials"].append(material)
    _p7_t5_save_buf(chat_id, buf)
    return len(buf.get("materials", [])), material

def _p7_t5_bind_comment(chat_id, reply_to_message_id, comment, is_voice=False):
    reply_id = _p7_t5_s(reply_to_message_id)
    if not reply_id:
        return None
    comment = _p7_t5_s(comment, 8000)
    if not comment:
        return None

    buf = _p7_t5_load_buf(chat_id)
    for m in buf.get("materials", []):
        ids = {
            _p7_t5_s(m.get("telegram_message_id")),
            _p7_t5_s(m.get("reply_to_message_id")),
            _p7_t5_msg_id_from_name(m.get("file_name")),
        }
        if reply_id in ids:
            field = "voice_comment" if is_voice else "owner_comment"
            prev = _p7_t5_s(m.get(field))
            m[field] = (prev + "\n" + comment).strip() if prev and comment not in prev else comment
            m["status"] = "LINKED"
            m["linked_reply_to_message_id"] = reply_id
            m["updated_at"] = _p7_t5_time.time()
            _p7_t5_save_buf(chat_id, buf)
            return m
    return None

def _p7_t5_is_flush_command(text):
    low = _p7_t5_low(text)
    return any(x in low for x in (
        "сделай акт", "собери акт", "сделай разбор", "сделай анализ",
        "собери разбор", "сформируй акт", "сделай отчет", "сделай отчёт"
    ))

def process_technadzor(text: str = "", task_id: str = "", chat_id: str = "", topic_id: int = 0, file_path: str = "", file_name: str = "", **kwargs):
    if int(topic_id or 0) != 5:
        return _P7_T5_ORIG_PROCESS_TECHNADZOR(text=text, task_id=task_id, chat_id=chat_id, topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs)

    chat = _p7_t5_s(chat_id or _p7_t5_task_get(kwargs.get("task"), "chat_id", ""))
    payload = _p7_t5_parse_payload(text)
    task = kwargs.get("task")
    reply_to = payload.get("telegram_reply_to_message_id") or _p7_t5_task_get(task, "reply_to_message_id", "")
    current_msg = payload.get("telegram_message_id") or _p7_t5_msg_id_from_name(file_name)
    comment = payload.get("transcript") or payload.get("text") or text
    is_voice = _p7_t5_low(comment).startswith("[voice]") or _p7_t5_low(payload.get("input_type", "")) == "voice"
    comment_clean = _p7_t5_re.sub(r"^\s*\[VOICE\]\s*", "", _p7_t5_s(comment), flags=_p7_t5_re.I)

    if (file_path or file_name) and not _p7_t5_is_flush_command(comment_clean):
        af = _p7_t5_active_folder(chat)
        fn = _p7_t5_s(file_name or payload.get("file_name") or payload.get("name") or "")
        low_fn = fn.lower()
        ftype = "PHOTO" if low_fn.endswith((".jpg", ".jpeg", ".png", ".webp", ".heic")) else "PDF" if low_fn.endswith(".pdf") else "DOCUMENT"
        material = {
            "source": "TELEGRAM",
            "file_type": ftype,
            "file_name": fn,
            "telegram_message_id": _p7_t5_s(current_msg),
            "reply_to_message_id": _p7_t5_s(reply_to),
            "drive_file_id": _p7_t5_s(payload.get("drive_file_id")),
            "drive_url": _p7_t5_s(payload.get("drive_url") or payload.get("webViewLink") or file_path),
            "active_folder_id": _p7_t5_s(af.get("folder_id")),
            "active_folder_name": _p7_t5_s(af.get("folder_name")),
            "caption": _p7_t5_s(payload.get("caption") or comment_clean),
            "include_in_act": True,
            "include_in_report": True,
            "status": "PENDING",
            "added_at": _p7_t5_time.time(),
        }
        count, _ = _p7_t5_add_material(chat, material)
        return {
            "ok": True,
            "handled": True,
            "state": "DONE",
            "result_text": f"Фото/файл принят в пакет выезда: {count} шт. Активная папка: {material.get('active_folder_name') or material.get('active_folder_id')}.",
            "message": "Фото/файл принят в пакет выезда",
            "history": "P7_TOPIC5_MATERIAL_BUFFERED_WITH_ACTIVE_FOLDER",
        }

    if reply_to and comment_clean and not _p7_t5_is_flush_command(comment_clean):
        linked = _p7_t5_bind_comment(chat, reply_to, comment_clean, is_voice=is_voice)
        if linked:
            return {
                "ok": True,
                "handled": True,
                "state": "DONE",
                "result_text": f"Пояснение привязано к фото: {linked.get('file_name', '')}",
                "message": "Пояснение привязано к фото",
                "history": "P7_TOPIC5_REPLY_VOICE_BOUND_TO_MATERIAL",
            }

    return _P7_T5_ORIG_PROCESS_TECHNADZOR(text=text, task_id=task_id, chat_id=chat_id, topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs)
# === END_P7_TOPIC5_REPLY_VOICE_BINDING_V1 ===

# === FULLFIX_TOPIC5_TECHNADZOR_CANON_CONTOUR_V2_TECHNADZOR ===
import json as _t5v2_json
import sqlite3 as _t5v2_sqlite3
import time as _t5v2_time
import uuid as _t5v2_uuid
from pathlib import Path as _t5v2_Path

_T5V2_ORIG_PROCESS_TECHNADZOR = process_technadzor
_T5V2_DB = "/root/.areal-neva-core/data/core.db"
_T5V2_DATA = _t5v2_Path("/root/.areal-neva-core/data/technadzor")
_T5V2_DATA.mkdir(parents=True, exist_ok=True)

def _t5v2_s(v, limit=50000):
    try:
        if v is None:
            return ""
        return str(v).strip()[:limit]
    except Exception:
        return ""

def _t5v2_low(v):
    return _t5v2_s(v).lower().replace("ё", "е")

def _t5v2_json_load(raw):
    try:
        d = _t5v2_json.loads(_t5v2_s(raw))
        return d if isinstance(d, dict) else {}
    except Exception:
        return {}

def _t5v2_get(obj, key, default=""):
    try:
        if isinstance(obj, dict):
            return obj.get(key, default)
        return obj[key]
    except Exception:
        return getattr(obj, key, default)

def _t5v2_task(task_id, kwargs):
    task = kwargs.get("task")
    if task:
        return {
            "id": _t5v2_s(_t5v2_get(task, "id", task_id)),
            "chat_id": _t5v2_s(_t5v2_get(task, "chat_id", kwargs.get("chat_id", ""))),
            "topic_id": int(_t5v2_get(task, "topic_id", kwargs.get("topic_id", 0)) or 0),
            "reply_to_message_id": _t5v2_s(_t5v2_get(task, "reply_to_message_id", "")),
            "bot_message_id": _t5v2_s(_t5v2_get(task, "bot_message_id", "")),
            "input_type": _t5v2_s(_t5v2_get(task, "input_type", "")),
            "raw_input": _t5v2_s(_t5v2_get(task, "raw_input", "")),
        }

    if not task_id:
        return {}

    con = _t5v2_sqlite3.connect(_T5V2_DB)
    try:
        r = con.execute(
            "SELECT id,chat_id,topic_id,reply_to_message_id,bot_message_id,input_type,raw_input FROM tasks WHERE id=? LIMIT 1",
            (_t5v2_s(task_id),)
        ).fetchone()
    finally:
        con.close()

    if not r:
        return {}

    return {
        "id": _t5v2_s(r[0]),
        "chat_id": _t5v2_s(r[1]),
        "topic_id": int(r[2] or 0),
        "reply_to_message_id": _t5v2_s(r[3]),
        "bot_message_id": _t5v2_s(r[4]),
        "input_type": _t5v2_s(r[5]),
        "raw_input": _t5v2_s(r[6]),
    }

def _t5v2_is_photo_meta(meta):
    fn = _t5v2_s(meta.get("file_name") or meta.get("name"))
    mt = _t5v2_s(meta.get("mime_type"))
    return fn.lower().endswith((".jpg", ".jpeg", ".png", ".webp", ".heic")) or mt.startswith("image/")

def _t5v2_msg_id(meta):
    for k in ("telegram_message_id", "_reply_to_message_id", "reply_to_message_id"):
        v = _t5v2_s(meta.get(k))
        if v:
            return v
    fn = _t5v2_s(meta.get("file_name"))
    import re as _re
    m = _re.search(r"_(\d+)\.(?:jpg|jpeg|png|webp|heic)$", fn, _re.I)
    return m.group(1) if m else ""

def _t5v2_photo_rows(chat_id, topic_id=5):
    con = _t5v2_sqlite3.connect(_T5V2_DB)
    try:
        rows = con.execute(
            """
            SELECT rowid,id,raw_input,reply_to_message_id,created_at
            FROM tasks
            WHERE chat_id=?
              AND topic_id=?
              AND input_type='drive_file'
            ORDER BY rowid DESC
            LIMIT 300
            """,
            (_t5v2_s(chat_id), int(topic_id or 0))
        ).fetchall()
    finally:
        con.close()

    out = []
    for rowid, tid, raw, reply_to, created_at in rows:
        meta = _t5v2_json_load(raw)
        if not _t5v2_is_photo_meta(meta):
            continue
        meta["_rowid"] = int(rowid)
        meta["_task_id"] = _t5v2_s(tid)
        meta["_reply_to_message_id"] = _t5v2_s(reply_to)
        meta["_created_at"] = _t5v2_s(created_at)
        out.append(meta)
    return out

def _t5v2_parent_reply_by_bot(chat_id, topic_id, bot_message_id):
    if not bot_message_id:
        return ""
    con = _t5v2_sqlite3.connect(_T5V2_DB)
    try:
        r = con.execute(
            """
            SELECT reply_to_message_id
            FROM tasks
            WHERE chat_id=?
              AND topic_id=?
              AND CAST(bot_message_id AS TEXT)=?
            ORDER BY rowid DESC
            LIMIT 1
            """,
            (_t5v2_s(chat_id), int(topic_id or 0), _t5v2_s(bot_message_id))
        ).fetchone()
    finally:
        con.close()
    return _t5v2_s(r[0]) if r else ""

def _t5v2_find_anchor_photo(chat_id, topic_id, reply_to_message_id):
    rid = _t5v2_s(reply_to_message_id)
    if not rid:
        return {}

    rows = _t5v2_photo_rows(chat_id, topic_id)

    for meta in rows:
        ids = {
            _t5v2_s(meta.get("_reply_to_message_id")),
            _t5v2_s(meta.get("telegram_message_id")),
            _t5v2_msg_id(meta),
        }
        if rid in ids:
            return meta

    parent_reply = _t5v2_parent_reply_by_bot(chat_id, topic_id, rid)
    if parent_reply and parent_reply != rid:
        for meta in rows:
            ids = {
                _t5v2_s(meta.get("_reply_to_message_id")),
                _t5v2_s(meta.get("telegram_message_id")),
                _t5v2_msg_id(meta),
            }
            if parent_reply in ids:
                return meta

    return {}

def _t5v2_group_requested(text):
    low = _t5v2_low(text)
    return any(x in low for x in (
        "этими фото",
        "эти фото",
        "этих фото",
        "все фото",
        "всеми фото",
        "несколько фото",
        "три фото",
        "фотографии",
        "с ними",
        "по ним",
        "их",
        "фото",
        "пакет",
    ))

def _t5v2_select_photo_group(chat_id, topic_id, anchor, text):
    if not anchor:
        return []

    if not _t5v2_group_requested(text):
        return [anchor]

    rows = _t5v2_photo_rows(chat_id, topic_id)

    try:
        anchor_rowid = int(anchor.get("_rowid") or 0)
    except Exception:
        anchor_rowid = 0

    try:
        anchor_msg = int(_t5v2_msg_id(anchor) or 0)
    except Exception:
        anchor_msg = 0

    selected = []
    for meta in rows:
        try:
            rowid = int(meta.get("_rowid") or 0)
        except Exception:
            rowid = 0

        try:
            msg = int(_t5v2_msg_id(meta) or 0)
        except Exception:
            msg = 0

        same_row_cluster = bool(anchor_rowid and abs(rowid - anchor_rowid) <= 10)
        same_msg_cluster = bool(anchor_msg and msg and abs(msg - anchor_msg) <= 20)

        if same_row_cluster and same_msg_cluster:
            selected.append(meta)

    selected = sorted(selected, key=lambda m: int(m.get("_rowid") or 0))
    return selected or [anchor]

def _t5v2_active_folder(chat_id):
    try:
        if "get_active_folder" in globals():
            af = get_active_folder(str(chat_id), 5)
            if isinstance(af, dict):
                return af
    except Exception:
        pass

    try:
        p = _T5V2_DATA / f"active_folder_{chat_id}_5.json"
        d = _t5v2_json.loads(p.read_text(encoding="utf-8"))
        return d if isinstance(d, dict) else {}
    except Exception:
        return {}

def _t5v2_buf_path(chat_id):
    return _T5V2_DATA / f"buf_{chat_id}_5.json"

def _t5v2_load_buf(chat_id):
    p = _t5v2_buf_path(chat_id)
    try:
        d = _t5v2_json.loads(p.read_text(encoding="utf-8"))
        if isinstance(d, dict):
            d.setdefault("materials", [])
            return d
    except Exception:
        pass
    return {"source": "topic5_visit_buffer", "materials": [], "created_at": _t5v2_time.time()}

def _t5v2_save_buf(chat_id, buf):
    buf["updated_at"] = _t5v2_time.time()
    _t5v2_buf_path(chat_id).write_text(_t5v2_json.dumps(buf, ensure_ascii=False, indent=2), encoding="utf-8")

def _t5v2_material(chat_id, meta, comment=""):
    af = _t5v2_active_folder(chat_id)
    fid = _t5v2_s(meta.get("drive_file_id") or meta.get("file_id") or meta.get("id"))
    fn = _t5v2_s(meta.get("file_name") or meta.get("name"))
    mid = _t5v2_msg_id(meta)
    clean = _t5v2_s(comment, 20000)

    if clean.upper().startswith("[VOICE]"):
        clean = clean[7:].strip()

    return {
        "material_id": str(_t5v2_uuid.uuid4()),
        "source": "TELEGRAM",
        "file_type": "PHOTO",
        "file_name": fn,
        "drive_file_id": fid,
        "drive_url": _t5v2_s(meta.get("drive_url") or meta.get("webViewLink") or (f"https://drive.google.com/file/d/{fid}/view?usp=drivesdk" if fid else "")),
        "telegram_message_id": mid,
        "reply_to_message_id": mid,
        "source_task_id": _t5v2_s(meta.get("_task_id")),
        "active_folder_id": _t5v2_s(af.get("folder_id")),
        "active_folder_name": _t5v2_s(af.get("folder_name")),
        "include_in_report": True,
        "include_in_act": True,
        "status": "LINKED" if clean else "PENDING",
        "voice_comment": clean,
        "added_at": _t5v2_time.time(),
        "updated_at": _t5v2_time.time(),
    }

def _t5v2_upsert_material(chat_id, material):
    buf = _t5v2_load_buf(chat_id)
    mid = _t5v2_s(material.get("telegram_message_id"))
    fn = _t5v2_s(material.get("file_name"))

    target = None
    for old in buf.get("materials", []):
        if (mid and _t5v2_s(old.get("telegram_message_id")) == mid) or (fn and _t5v2_s(old.get("file_name")) == fn):
            target = old
            break

    if target is None:
        buf["materials"].append(material)
    else:
        old_comment = _t5v2_s(target.get("voice_comment"), 20000)
        new_comment = _t5v2_s(material.get("voice_comment"), 20000)
        target.update({k: v for k, v in material.items() if v not in ("", None)})

        if old_comment and new_comment and new_comment not in old_comment:
            target["voice_comment"] = old_comment + "\n" + new_comment
        elif old_comment and not new_comment:
            target["voice_comment"] = old_comment

    _t5v2_save_buf(chat_id, buf)
    return len(buf.get("materials", []))

def _t5v2_bind_photos(chat_id, photos, comment):
    count = 0
    for meta in photos:
        count = _t5v2_upsert_material(chat_id, _t5v2_material(chat_id, meta, comment))
    return count

def _t5v2_positive_act(text):
    low = _t5v2_low(text)

    negated = any(x in low for x in (
        "не делай акт",
        "не надо акт",
        "не нужно акт",
        "не формируй акт",
        "не должен был сделать акт",
        "не должен делать акт",
        "акт не для каждого",
        "не для каждого из",
        "принять к сведению",
        "принять это к сведению",
        "прими к сведению",
        "прими это к сведению",
    ))

    positive = any(x in low for x in (
        "сделай акт",
        "сформируй акт",
        "собери акт",
        "готовь акт",
        "акт по этим фото",
        "сделай разбор",
        "сформируй документ",
    ))

    return positive and not negated

def _t5v2_buffer_summary(chat_id):
    buf = _t5v2_load_buf(chat_id)
    mats = buf.get("materials", [])

    lines = []
    for i, m in enumerate(mats, 1):
        lines.append(f"Фото №{i}: {m.get('file_name','')}")
        if m.get("voice_comment"):
            lines.append(f"Пояснение: {m.get('voice_comment')}")
        if m.get("drive_url"):
            lines.append(f"Ссылка: {m.get('drive_url')}")

    return "\n".join(lines), len(mats)

def process_technadzor(text: str = "", task_id: str = "", chat_id: str = "", topic_id: int = 0, file_path: str = "", file_name: str = "", **kwargs):
    task = _t5v2_task(task_id, kwargs)

    try:
        tid = int(topic_id or task.get("topic_id") or kwargs.get("topic_id") or 0)
    except Exception:
        tid = 0

    if tid != 5:
        return _T5V2_ORIG_PROCESS_TECHNADZOR(text=text, task_id=task_id, chat_id=chat_id, topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs)

    chat = _t5v2_s(chat_id or task.get("chat_id") or kwargs.get("chat_id"))
    raw = _t5v2_s(text or task.get("raw_input"))
    input_type = _t5v2_s(task.get("input_type"))
    reply_to = _t5v2_s(task.get("reply_to_message_id"))

    if input_type == "drive_file":
        meta = _t5v2_json_load(raw)
        if _t5v2_is_photo_meta(meta):
            count = _t5v2_upsert_material(chat, _t5v2_material(chat, meta, ""))
            return {
                "ok": True,
                "handled": True,
                "state": "DONE",
                "status": "DONE",
                "result_text": f"Фото принято в пакет технадзора: {count} шт. Акт не формирую без отдельной команды.",
                "message": "Фото принято в пакет технадзора",
                "history": "FULLFIX_TOPIC5_PHOTO_TO_VISITBUFFER",
            }

    if input_type in ("text", "voice", "") and reply_to and raw and not _t5v2_positive_act(raw):
        anchor = _t5v2_find_anchor_photo(chat, 5, reply_to)
        if anchor:
            photos = _t5v2_select_photo_group(chat, 5, anchor, raw)
            count = _t5v2_bind_photos(chat, photos, raw)
            names = ", ".join(_t5v2_s(p.get("file_name")) for p in photos if p.get("file_name"))
            return {
                "ok": True,
                "handled": True,
                "state": "DONE",
                "status": "DONE",
                "result_text": f"Пояснение принято к фото: {len(photos)} шт. В пакете технадзора: {count} шт. Акт не формирую без отдельной команды.\nФайлы: {names}",
                "message": "Пояснение принято к фото",
                "history": "FULLFIX_TOPIC5_REPLY_TO_PHOTO_BOUND",
            }

    if input_type in ("text", "voice", "") and _t5v2_positive_act(raw):
        summary, n = _t5v2_buffer_summary(chat)
        if n <= 0:
            return {
                "ok": True,
                "handled": True,
                "state": "DONE",
                "status": "DONE",
                "result_text": "В пакете технадзора нет фото. Сначала пришли фото или ответь голосом на фото.",
                "message": "В пакете технадзора нет фото",
                "history": "FULLFIX_TOPIC5_ACT_NO_MATERIALS",
            }

        enriched = raw + "\n\nПакет фото технадзора:\n" + summary
        return _T5V2_ORIG_PROCESS_TECHNADZOR(text=enriched, task_id=task_id, chat_id=chat, topic_id=5, file_path=file_path, file_name=file_name, **kwargs)

    return _T5V2_ORIG_PROCESS_TECHNADZOR(text=text, task_id=task_id, chat_id=chat_id, topic_id=topic_id, file_path=file_path, file_name=file_name, **kwargs)
# === END_FULLFIX_TOPIC5_TECHNADZOR_CANON_CONTOUR_V2_TECHNADZOR ===
