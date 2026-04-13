import os
import logging
from typing import Dict, Any

logger = logging.getLogger(__name__)

def parse_document(path: str) -> Dict[str, Any]:
    if not os.path.exists(path):
        return {"text": "", "tables": [], "metadata": {}, "error": f"File not found: {path}"}
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _parse_pdf(path)
        elif ext == ".docx":
            return _parse_docx(path)
        elif ext in (".xlsx", ".xls"):
            return _parse_excel(path)
        elif ext == ".csv":
            return _parse_csv(path)
        else:
            return {"text": "", "tables": [], "metadata": {}, "error": f"Unsupported: {ext}"}
    except Exception as e:
        logger.error(f"parse_document error: {e}")
        return {"text": "", "tables": [], "metadata": {}, "error": str(e)}

def extract_text_from_document(path: str) -> str:
    return parse_document(path).get("text", "")

def extract_tables_from_document(path: str) -> list:
    return parse_document(path).get("tables", [])

def _parse_pdf(path: str) -> Dict[str, Any]:
    import pdfplumber
    text_parts = []
    tables = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t: text_parts.append(t)
            for tbl in page.extract_tables():
                if tbl: tables.append(tbl)
    return {"text": "\n".join(text_parts), "tables": tables, "metadata": {"pages": len(pdf.pages)}, "error": ""}

def _parse_docx(path: str) -> Dict[str, Any]:
    from docx import Document
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text)
    tables = [[[cell.text for cell in row.cells] for row in t.rows] for t in doc.tables]
    return {"text": text, "tables": tables, "metadata": {"source": path}, "error": ""}

def _parse_excel(path: str) -> Dict[str, Any]:
    import pandas as pd
    sheets = pd.read_excel(path, sheet_name=None)
    text, tables = [], []
    for name, df in sheets.items():
        text.append(f"=== {name} ===\n{df.to_string(max_rows=50)}")
        tables.append({"sheet": name, "data": df.fillna("").to_dict(orient="records")})
    return {"text": "\n".join(text), "tables": tables, "metadata": {"sheets": list(sheets.keys())}, "error": ""}

def _parse_csv(path: str) -> Dict[str, Any]:
    import pandas as pd
    df = pd.read_csv(path)
    text = df.to_string(max_rows=50)
    tables = [{"data": df.fillna("").to_dict(orient="records")}]
    return {"text": text, "tables": tables, "metadata": {"source": path}, "error": ""}
