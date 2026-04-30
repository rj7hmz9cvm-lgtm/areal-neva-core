# === PROJECT_ENGINE_V1 ===
"""
core/project_engine.py
Разработка проектной документации по нормам ГОСТ/СНиП/СП
на основе шаблонов пользователя.
Разрешение на создание получено: 29.04.2026
"""
import os, re, logging, tempfile
from pathlib import Path
from typing import Dict, Any, Optional, List

logger = logging.getLogger(__name__)

SECTION_MAP = {
    "кж":  "КЖ — Конструкции железобетонные",
    "км":  "КМ — Конструкции металлические",
    "кмд": "КМД — Конструкции металлические деталировочные",
    "ар":  "АР — Архитектурные решения",
    "ов":  "ОВ — Отопление и вентиляция",
    "вк":  "ВК — Водоснабжение и канализация",
    "эом": "ЭОМ — Электроосвещение",
    "сс":  "СС — Слаботочные системы",
    "гп":  "ГП — Генеральный план",
    "пз":  "ПЗ — Пояснительная записка",
    "см":  "СМ — Смета",
    "тх":  "ТХ — Технологические решения",
}

SECTION_STRUCTURE = {
    "кж":  ["Армирование", "Схемы", "Спецификация арматуры", "Спецификация материалов"],
    "км":  ["Нагрузки", "Узлы сопряжений", "Спецификация металла"],
    "кмд": ["Деталировка", "Узлы", "Спецификация"],
    "ар":  ["Планы этажей", "Фасады", "Разрезы", "Экспликация помещений"],
    "ов":  ["Схема системы", "Расчёт нагрузок", "Спецификация оборудования"],
    "вк":  ["Схема водоснабжения", "Схема канализации", "Спецификация"],
    "эом": ["Однолинейная схема", "Расчёт нагрузок", "Спецификация"],
}

NORMS_MAP = {
    "кж":  ["СП 63.13330.2018", "ГОСТ 34028-2016", "СП 20.13330.2017"],
    "км":  ["СП 16.13330.2017", "ГОСТ 27772-2015", "СП 20.13330.2017"],
    "ар":  ["СП 118.13330.2022", "ГОСТ 21.501-2018"],
    "ов":  ["СП 60.13330.2020", "ГОСТ 30494-2011"],
    "вк":  ["СП 30.13330.2020", "СП 31.13330.2021"],
    "эом": ["СП 256.1325800.2016", "ПУЭ-7"],
}

SNOW_LOADS = {1: 0.8, 2: 1.2, 3: 1.8, 4: 2.4, 5: 3.2, 6: 4.0, 7: 4.8, 8: 5.6}
WIND_LOADS = {1: 0.17, 2: 0.23, 3: 0.30, 4: 0.38, 5: 0.48, 6: 0.60, 7: 0.73, 8: 0.85}

SPEC_HEADERS = ["№", "Наименование", "Марка/Обозначение", "Ед. изм.", "Кол-во", "Примечание"]
UNITS = {"мм", "м", "м2", "м3", "кг", "т", "шт", "пог.м"}


def detect_section(file_name: str, text: str = "") -> Optional[str]:
    # FULLFIX_02_B1: filename-first section priority
    fn = (file_name or "").lower()
    for key in SECTION_MAP:
        if key in fn:
            return key
    src = ((file_name or "") + " " + (text or "")).lower()
    for key in SECTION_MAP:
        if key in src:
            return key
    return None


def calc_loads(region: int = 3) -> Dict[str, float]:
    return {
        "snow_kPa":  SNOW_LOADS.get(region, 1.8),
        "wind_kPa":  WIND_LOADS.get(region, 0.30),
        "region":    region,
        "note":      f"СП 20.13330.2017 — район {region}",
    }


def normalize_unit(unit: str) -> str:
    u = str(unit or "").strip().lower()
    mapping = {"м2": "м2", "м²": "м2", "м3": "м3", "м³": "м3", "кг": "кг", "т": "т", "шт": "шт", "м": "м", "мм": "мм"}
    return mapping.get(u, u)


def build_specification(items: List[Dict]) -> List[List]:
    rows = [SPEC_HEADERS]
    for i, item in enumerate(items, 1):
        rows.append([
            i,
            item.get("name", ""),
            item.get("mark", ""),
            normalize_unit(item.get("unit", "")),
            item.get("qty", ""),
            item.get("note", ""),
        ])
    return rows


def _write_project_xlsx(section: str, items: List[Dict], loads: Dict, task_id: str) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    wb = Workbook()
    ws = wb.active
    ws.title = section.upper()

    ws.merge_cells("A1:F1")
    ws["A1"] = SECTION_MAP.get(section, section.upper())
    ws["A1"].font = Font(bold=True, size=13)
    ws["A1"].alignment = Alignment(horizontal="center")

    norms = NORMS_MAP.get(section, [])
    ws["A2"] = "Нормы: " + ", ".join(norms) if norms else ""

    if section in ("кж", "км", "кмд"):
        ws["A3"] = f"Снег: {loads['snow_kPa']} кПа | Ветер: {loads['wind_kPa']} кПа | {loads['note']}"

    spec = build_specification(items)
    start_row = 5
    for r_idx, row in enumerate(spec, start_row):
        for c_idx, val in enumerate(row, 1):
            cell = ws.cell(r_idx, c_idx, value=val)
            if r_idx == start_row:
                cell.font = Font(bold=True)
                cell.fill = PatternFill("solid", fgColor="DDEEFF")

    struct = SECTION_STRUCTURE.get(section, [])
    if struct:
        ws.cell(start_row + len(spec) + 2, 1, "Состав раздела:")
        for i, s in enumerate(struct, 1):
            ws.cell(start_row + len(spec) + 2 + i, 1, f"{i}. {s}")

    tmp = os.path.join(tempfile.gettempdir(), f"project_{section}_{task_id}.xlsx")
    wb.save(tmp)
    return tmp


async def generate_project_section(section: str, items: List[Dict], task_id: str, topic_id: int, region: int = 3) -> Dict[str, Any]:
    res = {"success": False, "excel_path": None, "drive_link": None, "section": section, "error": None}
    try:
        loads = calc_loads(region)
        xl = _write_project_xlsx(section, items, loads, task_id)
        res["excel_path"] = xl

        from core.engine_base import upload_artifact_to_drive, quality_gate
        qg = quality_gate(xl, task_id, "excel")
        if not qg["passed"]:
            res["error"] = f"QualityGate: {qg['errors']}"
            return res

        link = upload_artifact_to_drive(xl, task_id, topic_id)
        if link:
            res["drive_link"] = link
            res["success"] = True
        else:
            res["error"] = "UPLOAD_FAILED"
    except Exception as e:
        logger.error(f"project_engine: {e}", exc_info=True)
        res["error"] = str(e)[:300]
    return res


def project_result_guard(result: Dict) -> Dict:
    if not result.get("success"):
        return result
    if not result.get("excel_path") and not result.get("drive_link"):
        result["success"] = False
        result["error"] = "PROJECT_RESULT_GUARD: нет артефакта"
    return result


async def process_project_file(file_path: str, task_id: str, topic_id: int, raw_input: str = "") -> Dict[str, Any]:
    section = detect_section(file_path, raw_input) or "кж"
    items = []

    try:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        for row in (table or []):
                            if row and any(row):
                                items.append({
                                    "name": str(row[0] or ""),
                                    "mark": str(row[1] or "") if len(row) > 1 else "",
                                    "unit": str(row[2] or "") if len(row) > 2 else "",
                                    "qty":  str(row[3] or "") if len(row) > 3 else "",
                                })
        elif ext in (".xlsx", ".xls"):
            from openpyxl import load_workbook
            wb = load_workbook(file_path, data_only=True)
            ws = wb.active
            for row in ws.iter_rows(min_row=2, values_only=True):
                if row and any(v for v in row if v):
                    items.append({
                        "name": str(row[0] or ""),
                        "mark": str(row[1] or "") if len(row) > 1 else "",
                        "unit": str(row[2] or "") if len(row) > 2 else "",
                        "qty":  str(row[3] or "") if len(row) > 3 else "",
                    })
            wb.close()
    except Exception as e:
        logger.warning(f"project extract: {e}")

    result = await generate_project_section(section, items, task_id, topic_id)
    return project_result_guard(result)
# === END_PROJECT_ENGINE_V1 ===

# === CODE_CLOSE_V43_PROJECT_ENGINE ===

def normative_search_engine_v43(section: str, query: str = ""):
    base = NORMS_MAP.get(section, [])
    if base:
        return {"success": True, "norms": base, "source": "local_norms_map"}
    return {"success": False, "norms": [], "error": "норма не подтверждена"}

def project_validator_v43(result):
    if not isinstance(result, dict):
        return False, "PROJECT_VALIDATOR: empty"
    if result.get("success") is False:
        return False, str(result.get("error") or "PROJECT_VALIDATOR: failed")
    if not (result.get("drive_link") or result.get("excel_path") or result.get("docx_path") or result.get("pdf_path")):
        return False, "PROJECT_VALIDATOR: no_artifact"
    return True, ""

def metal_structure_engine_v43(items, region=3):
    loads = calc_loads(region)
    spec = []
    for item in items or []:
        name = str(item.get("name") or "")
        if any(x in name.lower() for x in ("колонна","балка","ферма","связь","прогон")):
            spec.append(item)
    return {"loads": loads, "items": spec, "norms": NORMS_MAP.get("км", [])}

def project_result_guard_v43(result):
    ok, reason = project_validator_v43(result)
    if not ok:
        result = result if isinstance(result, dict) else {}
        result["success"] = False
        result["error"] = reason
    return result

try:
    _v43_orig_generate_project_section = generate_project_section
    async def generate_project_section(section, items, task_id, topic_id, region=3):
        res = await _v43_orig_generate_project_section(section, items, task_id, topic_id, region)
        res["normative_search"] = normative_search_engine_v43(section)
        if section in ("км","кмд"):
            res["metal_structure"] = metal_structure_engine_v43(items, region)
        return project_result_guard_v43(res)
except Exception:
    pass

# === END_CODE_CLOSE_V43_PROJECT_ENGINE ===

# === PROJECT_CLOSE_V44 ===

def normative_search_engine_v44(section: str, query: str = ""):
    norms = NORMS_MAP.get(section, [])
    if norms:
        return {"success": True, "norms": norms, "source": "NORMS_MAP"}
    return {"success": False, "norms": [], "error": "норма не подтверждена"}

def project_validator_v44(result):
    if not isinstance(result, dict):
        return False, "PROJECT_VALIDATOR_EMPTY"
    if result.get("success") is False:
        return False, str(result.get("error") or "PROJECT_VALIDATOR_FAILED")
    if not result.get("section"):
        return False, "PROJECT_VALIDATOR_NO_SECTION"
    if not (result.get("drive_link") or result.get("excel_path") or result.get("docx_path") or result.get("pdf_path")):
        return False, "PROJECT_VALIDATOR_NO_ARTIFACT"
    return True, ""

def metal_structure_engine_v44(items, region=3):
    spec = []
    for item in items or []:
        name = str(item.get("name") or "").lower()
        if any(x in name for x in ("колонна","балка","ферма","связь","прогон","рама","ангар")):
            spec.append(item)
    return {"success": True, "loads": calc_loads(region), "items": spec, "norms": NORMS_MAP.get("км", [])}

def _write_project_docx_v44(section, items, loads, task_id):
    import os, tempfile
    path = os.path.join(tempfile.gettempdir(), f"project_{section}_{task_id}.docx")
    try:
        from docx import Document
        doc = Document()
        doc.add_heading(SECTION_MAP.get(section, section.upper()), level=1)
        doc.add_paragraph("Нормы: " + ", ".join(NORMS_MAP.get(section, [])))
        doc.add_paragraph(f"Снег: {loads.get('snow_kPa')} кПа | Ветер: {loads.get('wind_kPa')} кПа")
        table = doc.add_table(rows=1, cols=6)
        for i, h in enumerate(SPEC_HEADERS):
            table.rows[0].cells[i].text = h
        for n, item in enumerate(items or [], 1):
            row = table.add_row().cells
            row[0].text = str(n)
            row[1].text = str(item.get("name",""))
            row[2].text = str(item.get("mark",""))
            row[3].text = normalize_unit(item.get("unit",""))
            row[4].text = str(item.get("qty",""))
            row[5].text = str(item.get("note",""))
        doc.save(path)
    except Exception:
        with open(path, "w", encoding="utf-8") as f:
            f.write(SECTION_MAP.get(section, section.upper()) + "\n")
            f.write("Нормы: " + ", ".join(NORMS_MAP.get(section, [])) + "\n")
            f.write(str(items or []))
    return path

try:
    _v44_orig_generate_project_section = generate_project_section

    async def generate_project_section(section, items, task_id, topic_id, region=3):
        res = await _v44_orig_generate_project_section(section, items, task_id, topic_id, region)
        loads = calc_loads(region)
        res["normative_search"] = normative_search_engine_v44(section)
        if section in ("км", "кмд"):
            res["metal_structure"] = metal_structure_engine_v44(items, region)
        try:
            docx_path = _write_project_docx_v44(section, items, loads, task_id)
            res["docx_path"] = docx_path
            from core.engine_base import upload_artifact_to_drive
            docx_link = upload_artifact_to_drive(docx_path, task_id, topic_id)
            if docx_link:
                res["docx_link"] = docx_link
        except Exception as e:
            res["docx_error"] = str(e)[:300]
        ok, reason = project_validator_v44(res)
        if not ok:
            res["success"] = False
            res["error"] = reason
        return res
except Exception:
    pass

# === END_PROJECT_CLOSE_V44 ===


# === PATCH_TEMPLATE_MODEL_EXTRACTOR_V1 ===
import re as _re_pte

def extract_template_model_from_text(text: str, file_name: str = "", user_text: str = "") -> dict:
    lines = [l.strip() for l in (text or "").replace("\r","\n").split("\n") if l.strip()]
    # FULLFIX_02_B2: filename-first project type priority, КД/КЖ before АР
    _MARKS_PRI = ("КД","КЖ","КМД","КМ","КР","АР","ОВ","ВК","ЭОМ","СС","ГП","ПЗ","ТХ","СМ")
    project_type = "UNKNOWN"
    for _pt_src in ((file_name or ""), (user_text or ""), (text or "")[:500]):
        _pt_upper = _pt_src.upper()
        for mark in _MARKS_PRI:
            if _re_pte.search(rf"(^|[^А-ЯA-Z0-9]){_re_pte.escape(mark)}([^А-ЯA-Z0-9]|$)", _pt_upper):
                project_type = mark
                break
        if project_type != "UNKNOWN":
            break
    src = f"{file_name} {user_text} {text[:3000]}".upper()
    sheets = []
    for line in lines:
        m = _re_pte.search(r"(АР|КЖ|КД|КР|КМ|КМД|ОВ|ВК|ЭОМ)[\s\-]*(\d+[А-ЯA-Z0-9\-\.]*)\s+(.{4,120})", line, _re_pte.I)
        if m and len(sheets) < 80:
            sheets.append({"mark": m.group(1), "number": m.group(2), "title": m.group(3).strip()})
    section_keys = ("общие данные","исходные данные","расч","план","фасад","разрез","узел","схема","спецификация","ведомость","конструктив","материал")
    sections = []
    seen = set()
    for line in lines:
        if any(k in line.lower() for k in section_keys) and line.lower() not in seen:
            seen.add(line.lower())
            sections.append(line[:160])
        if len(sections) >= 60:
            break
    # FULLFIX_02_B3: sheet_register fallback from extracted structure when explicit sheet marks are absent
    if not sheets and sections:
        _sf_keys = ("общие данные","ведомость","план","фасады","фасад","разрез","узел","спецификация","схема","расч","конструктив")
        _sf_seen = set()
        _sf_seq = 1
        for _sf_line in sections:
            _sf_low = _sf_line.lower()
            if any(k in _sf_low for k in _sf_keys):
                _sf_title = _sf_line[:120].strip()
                if _sf_title and _sf_title.lower() not in _sf_seen:
                    _sf_seen.add(_sf_title.lower())
                    sheets.append({"mark": project_type, "number": str(_sf_seq), "title": _sf_title})
                    _sf_seq += 1
            if _sf_seq > 30:
                break
    axes_letters = sorted(set(_re_pte.findall(r"(?<![А-ЯA-Z])([А-ЯA-Z])(?=\s*[-–]\s*[А-ЯA-Z])", text)))
    axes_numbers = sorted(set(_re_pte.findall(r"(?<!\d)(\d{1,2})(?=\s*[-–]\s*\d{1,2})(?!\d)", text)), key=lambda x: int(x))
    dims = []
    for x in _re_pte.findall(r"(?<!\d)(\d{3,5})(?!\d)", text):
        try:
            v = int(x)
            if 300 <= v <= 50000 and v not in dims:
                dims.append(v)
        except Exception:
            pass
    levels = []
    for v in _re_pte.findall(r"[-+]?\d+[,.]\d{2,3}", text):
        try:
            f = float(v.replace(",","."))
            s = str(f)
            if -20 <= f <= 100 and s not in levels:
                levels.append(s)
        except Exception:
            pass
    mat_keys = ("бетон","арматур","a500","b25","в25","доска","брус","фанера","утепл","профлист","металл","кирпич","газобетон","сталь","с255")
    materials = []
    mat_seen = set()
    for line in lines:
        if any(k in line.lower() for k in mat_keys) and line.lower() not in mat_seen:
            mat_seen.add(line.lower())
            materials.append(line[:180])
        if len(materials) >= 60:
            break
    stamp = {}
    for m in _re_pte.finditer(r"((?:Адрес|По адресу)[:\s]+)([^\n]{5,180})", text, _re_pte.I):
        stamp["address"] = m.group(2).strip()[:200]
        break
    for m in _re_pte.finditer(r"((?:ООО|ОАО|ЗАО|ИП)[^\n]{3,180})", text):
        stamp["developer"] = m.group(1).strip()[:200]
        break
    for m in _re_pte.finditer(r"\b(20\d{2})\b", text):
        stamp["year"] = m.group(1)
        break
    model = {
        "schema": "PROJECT_TEMPLATE_MODEL_V1",
        "project_type": project_type,
        "source_files": [file_name] if file_name else [],
        "sheet_register": sheets,
        "marks": [project_type] if project_type != "UNKNOWN" else [],
        "sections": sections,
        "axes_grid": {"axes_letters": axes_letters[:30], "axes_numbers": axes_numbers[:30]},
        "dimensions": dims[:80],
        "levels": levels[:40],
        "nodes": [x for x in sections if "узел" in x.lower()][:30],
        "specifications": [x for x in sections if any(k in x.lower() for k in ("спецификац","ведомость"))][:30],
        "materials": materials,
        "stamp_fields": stamp,
        "variable_parameters": ["project_name","address","customer","area","floors","axes_grid","dimensions","materials","sheet_register"],
        "output_documents": ["DOCX_PROJECT_TEMPLATE_SUMMARY","JSON_PROJECT_TEMPLATE_MODEL","XLSX_SPECIFICATION_DRAFT"],
        "quality": {
            "has_sheet_register": bool(sheets),
            "has_sections": bool(sections),
            "has_axes_or_dimensions": bool(axes_letters or axes_numbers or dims),
            "has_materials": bool(materials),
            "text_chars": len(text or ""),
            "lines": len(lines),
        }
    }
    return model


def is_valid_project_template_model(model: dict) -> bool:
    if not isinstance(model, dict):
        return False
    q = model.get("quality") or {}
    return bool(
        model.get("schema") == "PROJECT_TEMPLATE_MODEL_V1"
        and q.get("text_chars", 0) > 200
        and (q.get("has_sheet_register") or q.get("has_sections") or q.get("has_axes_or_dimensions") or q.get("has_materials"))
    )


def model_to_text_report(model: dict) -> str:
    lines = ["PROJECT_TEMPLATE_MODEL создан"]
    lines.append(f"Раздел: {model.get('project_type','UNKNOWN')}")
    lines.append("")
    sheets = model.get("sheet_register") or []
    lines.append(f"Состав листов ({len(sheets)}):")
    for s in sheets[:30]:
        lines.append(f"  {s.get('mark','')} {s.get('number','')} {s.get('title','')}".strip())
    if not sheets:
        lines.append("  не извлечён явно")
    lines.append("")
    lines.append("Структура/разделы:")
    for s in (model.get("sections") or [])[:20]:
        lines.append(f"  {s}")
    lines.append("")
    ag = model.get("axes_grid") or {}
    lines.append(f"Оси буквенные: {', '.join(ag.get('axes_letters',[]) or []) or 'не извлечены'}")
    lines.append(f"Оси цифровые: {', '.join(ag.get('axes_numbers',[]) or []) or 'не извлечены'}")
    dims = model.get("dimensions") or []
    lines.append(f"Размеры мм: {', '.join(map(str,dims[:20])) if dims else 'не извлечены'}")
    lines.append("")
    mats = model.get("materials") or []
    lines.append(f"Материалы ({len(mats)}):")
    for m in mats[:15]:
        lines.append(f"  {m}")
    return "\n".join(lines).strip()

# === END PATCH_TEMPLATE_MODEL_EXTRACTOR_V1 ===


# === FULLFIX_03_PROJECT_ARTIFACT_GENERATOR ===
def _ff3_latest_project_template_model(topic_id: int = 0) -> dict:
    import json, os, glob
    base = "/root/.areal-neva-core/data/project_templates"
    files = sorted(glob.glob(os.path.join(base, "PROJECT_TEMPLATE_MODEL__*.json")), key=os.path.getmtime, reverse=True)
    if not files:
        return {}
    topic_id = int(topic_id or 0)
    best = None
    for p in files:
        try:
            data = json.load(open(p, encoding="utf-8"))
            if topic_id and int(data.get("topic_id", 0) or 0) == topic_id:
                return data
            if best is None:
                best = data
        except Exception:
            pass
    return best or {}


def _ff3_extract_project_params(user_text: str) -> dict:
    import re
    txt = str(user_text or "")
    low = txt.lower()

    params = {}
    if "фундамент" in low or "плита" in low:
        params["project_name"] = "Проект фундаментной плиты"
        params["section"] = "КЖ"
    elif "кров" in low or "строп" in low:
        params["project_name"] = "Проект кровли"
        params["section"] = "КД"
    else:
        params["project_name"] = "Проект по образцу"
        params["section"] = ""

    m = re.search(r"(\d+(?:[,.]\d+)?)\s*[xх×]\s*(\d+(?:[,.]\d+)?)\s*м", low)
    if m:
        params["size"] = f"{m.group(1).replace(',', '.')} x {m.group(2).replace(',', '.')} м"

    for label, key in (
        ("толщина", "thickness"),
        ("песчан", "sand"),
        ("щеб", "gravel"),
        ("бетон", "concrete"),
        ("арматур", "rebar"),
    ):
        mm = re.search(label + r"[^0-9]{0,30}(\d{2,4})\s*мм", low)
        if mm:
            params[key] = mm.group(1) + " мм"

    return params


def _ff3_safe_docx_text(value) -> str:
    return str(value or "").replace("\x00", " ").strip()


def create_project_artifact_from_latest_template(user_text: str, task_id: str, topic_id: int = 0) -> dict:
    """
    Создаёт реальный DOCX + XLSX проектный артефакт по последней PROJECT_TEMPLATE_MODEL
    Возвращает пути и Drive-ссылки
    """
    import os, tempfile, json
    from datetime import datetime, timezone

    result = {
        "success": False,
        "error": "",
        "docx_path": "",
        "xlsx_path": "",
        "docx_link": "",
        "xlsx_link": "",
        "template_found": False,
        "project_type": "UNKNOWN",
    }

    model = _ff3_latest_project_template_model(topic_id)
    params = _ff3_extract_project_params(user_text)
    if not model:
        result["error"] = "PROJECT_TEMPLATE_MODEL_NOT_FOUND"
        return result

    result["template_found"] = True
    project_type = params.get("section") or model.get("project_type") or "UNKNOWN"
    result["project_type"] = project_type

    safe_task = str(task_id or "manual")[:8]
    out_dir = tempfile.gettempdir()
    docx_path = os.path.join(out_dir, f"project_{project_type}_{safe_task}.docx")
    xlsx_path = os.path.join(out_dir, f"project_{project_type}_{safe_task}.xlsx")

    sheets = model.get("sheet_register") or []
    if not sheets:
        for i, sec in enumerate(model.get("sections") or [], 1):
            sheets.append({"mark": project_type, "number": str(i), "title": str(sec)[:120]})

    if not sheets:
        sheets = [
            {"mark": project_type, "number": "1", "title": "Титульный лист"},
            {"mark": project_type, "number": "2", "title": "Общие данные"},
            {"mark": project_type, "number": "3", "title": "План"},
            {"mark": project_type, "number": "4", "title": "Разрезы"},
            {"mark": project_type, "number": "5", "title": "Спецификация"},
        ]

    try:
        from docx import Document
        doc = Document()
        doc.add_heading(_ff3_safe_docx_text(params.get("project_name") or "Проект по образцу"), level=1)
        doc.add_paragraph("Сформировано AREAL-NEVA ORCHESTRA по сохранённой PROJECT_TEMPLATE_MODEL")
        doc.add_paragraph(f"Раздел: {project_type}")
        doc.add_paragraph(f"Дата: {datetime.now(timezone.utc).isoformat()}")

        doc.add_heading("Параметры задания", level=2)
        if params:
            for k, v in params.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            doc.add_paragraph(_ff3_safe_docx_text(user_text))

        doc.add_heading("Состав проекта по образцу", level=2)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.rows[0].cells[0].text = "Марка"
        tbl.rows[0].cells[1].text = "Лист"
        tbl.rows[0].cells[2].text = "Наименование"
        for sh in sheets:
            row = tbl.add_row().cells
            row[0].text = _ff3_safe_docx_text(sh.get("mark") or project_type)
            row[1].text = _ff3_safe_docx_text(sh.get("number") or "")
            row[2].text = _ff3_safe_docx_text(sh.get("title") or "")

        doc.add_heading("Оси и размеры", level=2)
        ag = model.get("axes_grid") or {}
        doc.add_paragraph("Оси буквенные: " + (", ".join(ag.get("axes_letters") or []) or "не извлечены"))
        doc.add_paragraph("Оси цифровые: " + (", ".join(ag.get("axes_numbers") or []) or "не извлечены"))
        dims = model.get("dimensions") or []
        doc.add_paragraph("Размеры мм: " + (", ".join(map(str, dims[:60])) if dims else "не извлечены"))

        doc.add_heading("Материалы из образца", level=2)
        mats = model.get("materials") or []
        if mats:
            for m in mats[:60]:
                doc.add_paragraph(_ff3_safe_docx_text(m))
        else:
            doc.add_paragraph("Материалы не извлечены из образца")

        doc.add_heading("Техническая структура", level=2)
        for sec in (model.get("sections") or [])[:80]:
            doc.add_paragraph(_ff3_safe_docx_text(sec))

        doc.save(docx_path)
    except Exception as e:
        result["error"] = "DOCX_CREATE_FAILED: " + str(e)[:250]
        return result

    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Состав проекта"
        headers = ["№", "Марка", "Лист", "Наименование", "Источник"]
        for c, h in enumerate(headers, 1):
            ws.cell(1, c, h)
        for i, sh in enumerate(sheets, 2):
            ws.cell(i, 1, i - 1)
            ws.cell(i, 2, sh.get("mark") or project_type)
            ws.cell(i, 3, sh.get("number") or "")
            ws.cell(i, 4, sh.get("title") or "")
            ws.cell(i, 5, ",".join(model.get("source_files") or []))
        ws2 = wb.create_sheet("Параметры")
        ws2.cell(1, 1, "Параметр")
        ws2.cell(1, 2, "Значение")
        for r, (k, v) in enumerate(params.items(), 2):
            ws2.cell(r, 1, k)
            ws2.cell(r, 2, v)
        ws.column_dimensions["D"].width = 70
        ws.column_dimensions["E"].width = 50
        ws2.column_dimensions["A"].width = 30
        ws2.column_dimensions["B"].width = 70
        wb.save(xlsx_path)
        wb.close()
    except Exception as e:
        result["error"] = "XLSX_CREATE_FAILED: " + str(e)[:250]
        return result

    result["docx_path"] = docx_path
    result["xlsx_path"] = xlsx_path

    try:
        from core.engine_base import upload_artifact_to_drive
        docx_link = upload_artifact_to_drive(docx_path, task_id, int(topic_id or 0))
        xlsx_link = upload_artifact_to_drive(xlsx_path, task_id, int(topic_id or 0))
        result["docx_link"] = docx_link or ""
        result["xlsx_link"] = xlsx_link or ""
    except Exception as e:
        result["upload_error"] = str(e)[:250]

    result["success"] = bool(os.path.exists(docx_path) and os.path.getsize(docx_path) > 1000)
    if not result["success"]:
        result["error"] = "PROJECT_ARTIFACT_EMPTY"
    return result

# === END FULLFIX_03_PROJECT_ARTIFACT_GENERATOR ===


# === FULLFIX_05_REAL_PROJECT_ENGINE ===
import os as _os_ff05
import re as _re_ff05
import json as _json_ff05
import math as _math_ff05
import tempfile as _tempfile_ff05
from pathlib import Path as _Path_ff05
from datetime import datetime as _dt_ff05

def _ff05_float(v, default=0.0):
    try:
        return float(str(v).replace(",", "."))
    except Exception:
        return float(default)

def _ff05_int(v, default=0):
    try:
        return int(float(str(v).replace(",", ".")))
    except Exception:
        return int(default)

def _ff05_latest_template(section: str = "КЖ") -> dict:
    base = _Path_ff05("/root/.areal-neva-core/data/project_templates")
    if not base.exists():
        return {}
    files = sorted(base.glob("PROJECT_TEMPLATE_MODEL__*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    section_u = str(section or "").upper()
    fallback = {}
    for p in files[:50]:
        try:
            data = _json_ff05.loads(p.read_text(encoding="utf-8"))
            data["_template_file"] = str(p)
            pt = str(data.get("project_type") or "").upper()
            if not fallback:
                fallback = data
            if pt == section_u:
                return data
        except Exception:
            continue
    return fallback

def _ff05_parse_project_request(raw_input: str, template_hint: str = "") -> dict:
    text = str(raw_input or "")
    low = text.lower()

    section = "КЖ"
    if any(x in low for x in ("кд", "деревян", "стропил", "кровл")):
        section = "КД"
    if any(x in low for x in ("кж", "фундамент", "плит")):
        section = "КЖ"

    length_m = 10.0
    width_m = 10.0
    slab_mm = 200
    sand_mm = 300
    gravel_mm = 100
    concrete_class = "B25"
    rebar_class = "A500C"
    rebar_step_mm = 200
    rebar_diam_mm = 12
    cover_mm = 40

    m = _re_ff05.search(r"(\d+(?:[,.]\d+)?)\s*[xх×]\s*(\d+(?:[,.]\d+)?)\s*м?", low)
    if m:
        length_m = _ff05_float(m.group(1), length_m)
        width_m = _ff05_float(m.group(2), width_m)

    def find_mm(keys, default):
        for key in keys:
            m2 = _re_ff05.search(key + r".{0,40}?(\d{2,4})\s*мм", low)
            if m2:
                return _ff05_int(m2.group(1), default)
        return default

    slab_mm = find_mm(("толщин", "плит", "бетон"), slab_mm)
    sand_mm = find_mm(("песчан", "песок"), sand_mm)
    gravel_mm = find_mm(("щеб", "основан"), gravel_mm)

    m = _re_ff05.search(r"(b|в)\s?(\d{2,3})", low)
    if m:
        concrete_class = "B" + m.group(2)

    m = _re_ff05.search(r"(a|а)\s?500", low)
    if m:
        rebar_class = "A500C"

    m = _re_ff05.search(r"(?:шаг|ячейк).{0,30}?(\d{2,4})\s*мм", low)
    if m:
        rebar_step_mm = _ff05_int(m.group(1), rebar_step_mm)

    m = _re_ff05.search(r"(?:арматур|ø|ф|диаметр).{0,30}?(\d{1,2})\s*мм", low)
    if m:
        rebar_diam_mm = _ff05_int(m.group(1), rebar_diam_mm)

    template = _ff05_latest_template(section)

    return {
        "project_name": "Проект фундаментной плиты" if section == "КЖ" else "Проект КД",
        "section": section,
        "length_m": length_m,
        "width_m": width_m,
        "slab_mm": slab_mm,
        "sand_mm": sand_mm,
        "gravel_mm": gravel_mm,
        "concrete_class": concrete_class,
        "rebar_class": rebar_class,
        "rebar_step_mm": rebar_step_mm,
        "rebar_diam_mm": rebar_diam_mm,
        "cover_mm": cover_mm,
        "template": {
            "project_type": template.get("project_type"),
            "source_files": template.get("source_files") or [],
            "template_file": template.get("_template_file"),
            "sections": template.get("sections") or [],
            "sheet_register": template.get("sheet_register") or [],
            "materials": template.get("materials") or [],
        },
    }

def _ff05_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for fp in candidates:
        if _os_ff05.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont("AREALFONT", fp))
                return "AREALFONT"
            except Exception:
                pass
    return "Helvetica"

def _ff05_draw_frame(c, page_w, page_h, title, sheet_no, sheet_total, font):
    from reportlab.lib.units import mm
    c.setLineWidth(0.7)
    c.rect(12*mm, 10*mm, page_w - 24*mm, page_h - 20*mm)
    c.line(12*mm, 28*mm, page_w - 12*mm, 28*mm)
    c.line(page_w - 95*mm, 10*mm, page_w - 95*mm, 28*mm)
    c.line(page_w - 55*mm, 10*mm, page_w - 55*mm, 28*mm)
    c.line(page_w - 25*mm, 10*mm, page_w - 25*mm, 28*mm)

    c.setFont(font, 9)
    c.drawString(15*mm, 18*mm, "AREAL-NEVA")
    c.drawString(page_w - 92*mm, 18*mm, "Стадия: П")
    c.drawString(page_w - 52*mm, 18*mm, f"Лист: {sheet_no}")
    c.drawString(page_w - 22*mm, 18*mm, f"Листов: {sheet_total}")

    c.setFont(font, 13)
    c.drawString(18*mm, page_h - 20*mm, title)
    c.setFont(font, 8)
    c.drawString(18*mm, page_h - 27*mm, "Комплект создан автоматически по задаче пользователя и сохранён как PDF/DXF артефакт")

def _ff05_draw_text(c, x, y, text, font, size=9):
    from reportlab.lib.units import mm
    c.setFont(font, size)
    c.drawString(x*mm, y*mm, str(text))

def _ff05_write_project_pdf(path: str, data: dict) -> str:
    from reportlab.lib.pagesizes import A3, landscape
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas

    font = _ff05_font()
    page_w, page_h = landscape(A3)
    c = canvas.Canvas(path, pagesize=landscape(A3))

    L = float(data["length_m"])
    W = float(data["width_m"])
    slab = int(data["slab_mm"])
    sand = int(data["sand_mm"])
    gravel = int(data["gravel_mm"])
    rebar_d = int(data["rebar_diam_mm"])
    rebar_step = int(data["rebar_step_mm"])
    cover = int(data["cover_mm"])
    concrete = data["concrete_class"]
    rebar_class = data["rebar_class"]
    sheet_total = 6

    # 1 title / general data
    _ff05_draw_frame(c, page_w, page_h, "Общие данные", 1, sheet_total, font)
    lines = [
        f"Наименование: {data['project_name']}",
        f"Раздел: {data['section']}",
        f"Габарит плиты: {L:g} x {W:g} м",
        f"Толщина плиты: {slab} мм",
        f"Подготовка: песок {sand} мм, щебень {gravel} мм",
        f"Бетон: {concrete}",
        f"Арматура: {rebar_class} Ø{rebar_d} с шагом {rebar_step} мм, защитный слой {cover} мм",
        "Выходные файлы: PDF-комплект + DXF-чертёж со слоями",
    ]
    y = 255
    for line in lines:
        _ff05_draw_text(c, 25, y, line, font, 11)
        y -= 9

    tpl = data.get("template") or {}
    _ff05_draw_text(c, 25, y - 5, "Источник структуры шаблона:", font, 11)
    y -= 15
    src_files = tpl.get("source_files") or []
    _ff05_draw_text(c, 30, y, ", ".join(src_files) if src_files else "не найден сохранённый исходный PDF-шаблон", font, 9)
    y -= 10
    c.showPage()

    # 2 plan
    _ff05_draw_frame(c, page_w, page_h, "План фундаментной плиты", 2, sheet_total, font)
    x0, y0 = 70*mm, 55*mm
    max_w, max_h = 260*mm, 150*mm
    scale = min(max_w/(L*1000), max_h/(W*1000))
    rw, rh = L*1000*scale, W*1000*scale
    c.setLineWidth(1.2)
    c.rect(x0, y0, rw, rh)
    c.setDash(4, 3)
    c.line(x0, y0+rh/2, x0+rw, y0+rh/2)
    c.line(x0+rw/2, y0, x0+rw/2, y0+rh)
    c.setDash()

    # rebar grid
    c.setLineWidth(0.25)
    step_draw = max(0.6*mm, rebar_step*scale)
    xx = x0 + step_draw
    while xx < x0 + rw:
        c.line(xx, y0, xx, y0+rh)
        xx += step_draw
    yy = y0 + step_draw
    while yy < y0 + rh:
        c.line(x0, yy, x0+rw, yy)
        yy += step_draw

    c.setFont(font, 9)
    c.drawString(x0, y0 - 8*mm, f"{L:g} м")
    c.saveState()
    c.translate(x0 - 10*mm, y0)
    c.rotate(90)
    c.drawString(0, 0, f"{W:g} м")
    c.restoreState()
    _ff05_draw_text(c, 25, 240, f"Армирование: {rebar_class} Ø{rebar_d} шаг {rebar_step} мм в двух направлениях", font, 10)
    _ff05_draw_text(c, 25, 230, f"Защитный слой бетона: {cover} мм", font, 10)
    c.showPage()

    # 3 section
    _ff05_draw_frame(c, page_w, page_h, "Разрез 1-1", 3, sheet_total, font)
    bx, by = 60*mm, 70*mm
    total = slab + gravel + sand
    k = 95*mm / total
    layers = [
        ("Фундаментная плита", slab, "Бетон " + concrete),
        ("Щебёночное основание", gravel, "Щебень"),
        ("Песчаная подушка", sand, "Песок"),
    ]
    ycur = by
    c.setLineWidth(0.8)
    for name, thick, note in layers:
        hh = thick * k
        c.rect(bx, ycur, 210*mm, hh)
        c.setFont(font, 10)
        c.drawString(bx + 5*mm, ycur + hh/2, f"{name}: {thick} мм — {note}")
        ycur += hh
    _ff05_draw_text(c, 25, 230, f"Общая конструктивная толщина: {total} мм", font, 10)
    c.showPage()

    # 4 reinforcement
    _ff05_draw_frame(c, page_w, page_h, "Схема армирования", 4, sheet_total, font)
    x0, y0 = 70*mm, 55*mm
    c.rect(x0, y0, rw, rh)
    c.setLineWidth(0.35)
    step_draw = max(1.0*mm, rebar_step*scale)
    xx = x0 + step_draw
    while xx < x0 + rw:
        c.line(xx, y0, xx, y0+rh)
        xx += step_draw
    yy = y0 + step_draw
    while yy < y0 + rh:
        c.line(x0, yy, x0+rw, yy)
        yy += step_draw
    _ff05_draw_text(c, 25, 240, f"Нижняя сетка: {rebar_class} Ø{rebar_d} шаг {rebar_step} мм", font, 10)
    _ff05_draw_text(c, 25, 230, f"Верхняя сетка: {rebar_class} Ø{rebar_d} шаг {rebar_step} мм", font, 10)
    _ff05_draw_text(c, 25, 220, "Выпуски, усиления, проёмы и закладные требуют отдельного задания", font, 9)
    c.showPage()

    # 5 specification
    _ff05_draw_frame(c, page_w, page_h, "Спецификация материалов", 5, sheet_total, font)
    area = L * W
    concrete_m3 = round(area * slab / 1000, 3)
    sand_m3 = round(area * sand / 1000, 3)
    gravel_m3 = round(area * gravel / 1000, 3)
    bars_x = _math_ff05.floor((W*1000 - 2*cover) / rebar_step) + 1
    bars_y = _math_ff05.floor((L*1000 - 2*cover) / rebar_step) + 1
    rebar_m = round((bars_x * L + bars_y * W) * 2, 1)
    spec = [
        ("1", f"Бетон {concrete}", "м3", concrete_m3),
        ("2", "Песчаная подушка", "м3", sand_m3),
        ("3", "Щебёночное основание", "м3", gravel_m3),
        ("4", f"Арматура {rebar_class} Ø{rebar_d}", "п.м", rebar_m),
    ]
    y = 240
    _ff05_draw_text(c, 25, y, "№   Наименование                              Ед.     Кол-во", font, 11)
    y -= 10
    for row in spec:
        _ff05_draw_text(c, 25, y, f"{row[0]:<3} {row[1]:<40} {row[2]:<6} {row[3]}", font, 10)
        y -= 9
    c.showPage()

    # 6 sheet list
    _ff05_draw_frame(c, page_w, page_h, "Ведомость листов", 6, sheet_total, font)
    sheets = [
        ("1", "Общие данные"),
        ("2", "План фундаментной плиты"),
        ("3", "Разрез 1-1"),
        ("4", "Схема армирования"),
        ("5", "Спецификация материалов"),
        ("6", "Ведомость листов"),
    ]
    y = 240
    for n, title in sheets:
        _ff05_draw_text(c, 25, y, f"{n}. {title}", font, 11)
        y -= 9
    c.save()

    return path

def _ff05_write_project_dxf(path: str, data: dict) -> str:
    import ezdxf

    L = float(data["length_m"]) * 1000.0
    W = float(data["width_m"]) * 1000.0
    slab = int(data["slab_mm"])
    sand = int(data["sand_mm"])
    gravel = int(data["gravel_mm"])
    step = int(data["rebar_step_mm"])
    rebar_d = int(data["rebar_diam_mm"])
    concrete = str(data["concrete_class"])
    rebar_class = str(data["rebar_class"])

    doc = ezdxf.new("R2010")
    msp = doc.modelspace()

    layers = {
        "KJ-SLAB": 7,
        "KJ-AXIS": 1,
        "KJ-REBAR": 3,
        "KJ-DIMS": 5,
        "KJ-TEXT": 2,
        "KJ-SECTION": 4,
    }
    for name, color in layers.items():
        if name not in doc.layers:
            doc.layers.new(name=name, dxfattribs={"color": color})

    # plan outline
    pts = [(0, 0), (L, 0), (L, W), (0, W), (0, 0)]
    msp.add_lwpolyline(pts, dxfattribs={"layer": "KJ-SLAB", "closed": True})

    # axes
    msp.add_line((L/2, -500), (L/2, W+500), dxfattribs={"layer": "KJ-AXIS"})
    msp.add_line((-500, W/2), (L+500, W/2), dxfattribs={"layer": "KJ-AXIS"})

    # rebar grid
    x = step
    while x < L:
        msp.add_line((x, 0), (x, W), dxfattribs={"layer": "KJ-REBAR"})
        x += step
    y = step
    while y < W:
        msp.add_line((0, y), (L, y), dxfattribs={"layer": "KJ-REBAR"})
        y += step

    # section block at right side
    sx = L + 2500
    sy = 0
    total = slab + gravel + sand
    msp.add_lwpolyline([(sx, sy), (sx+5000, sy), (sx+5000, sy+total), (sx, sy+total), (sx, sy)], dxfattribs={"layer": "KJ-SECTION", "closed": True})
    y1 = sy
    for name, th in [("SAND", sand), ("GRAVEL", gravel), ("SLAB", slab)]:
        msp.add_line((sx, y1), (sx+5000, y1), dxfattribs={"layer": "KJ-SECTION"})
        msp.add_text(f"{name} {th}mm", dxfattribs={"height": 250, "layer": "KJ-TEXT"}).set_placement((sx+5200, y1 + th/2))
        y1 += th
    msp.add_line((sx, y1), (sx+5000, y1), dxfattribs={"layer": "KJ-SECTION"})

    # notes
    msp.add_text(f"FOUNDATION SLAB {L/1000:g}x{W/1000:g}m", dxfattribs={"height": 350, "layer": "KJ-TEXT"}).set_placement((0, -1200))
    msp.add_text(f"SLAB {slab}mm / {concrete}", dxfattribs={"height": 250, "layer": "KJ-TEXT"}).set_placement((0, -1700))
    msp.add_text(f"REBAR {rebar_class} D{rebar_d} STEP {step}mm", dxfattribs={"height": 250, "layer": "KJ-TEXT"}).set_placement((0, -2100))
    msp.add_text(f"SAND {sand}mm / GRAVEL {gravel}mm", dxfattribs={"height": 250, "layer": "KJ-TEXT"}).set_placement((0, -2500))

    doc.saveas(path)
    return path

async def create_project_pdf_dxf_artifact(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "") -> dict:
    data = _ff05_parse_project_request(raw_input, template_hint)
    stamp = _dt_ff05.utcnow().strftime("%Y%m%d_%H%M%S")
    safe_task = _re_ff05.sub(r"[^A-Za-z0-9_-]+", "_", str(task_id or "manual"))[:20]
    out_dir = _Path_ff05(_tempfile_ff05.gettempdir()) / f"areal_project_{safe_task}_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = str(out_dir / f"{data['section']}_foundation_slab_{safe_task}.pdf")
    dxf_path = str(out_dir / f"{data['section']}_foundation_slab_{safe_task}.dxf")
    manifest_path = str(out_dir / f"{data['section']}_foundation_slab_{safe_task}.manifest.json")

    res = {
        "success": False,
        "section": data["section"],
        "pdf_path": pdf_path,
        "dxf_path": dxf_path,
        "manifest_path": manifest_path,
        "pdf_link": None,
        "dxf_link": None,
        "manifest_link": None,
        "error": None,
        "data": data,
    }

    try:
        _ff05_write_project_pdf(pdf_path, data)
        _ff05_write_project_dxf(dxf_path, data)

        if not _os_ff05.path.exists(pdf_path) or _os_ff05.path.getsize(pdf_path) < 3000:
            res["error"] = "PDF_NOT_CREATED_OR_TOO_SMALL"
            return res
        if not _os_ff05.path.exists(dxf_path) or _os_ff05.path.getsize(dxf_path) < 1500:
            res["error"] = "DXF_NOT_CREATED_OR_TOO_SMALL"
            return res

        manifest = {
            "schema": "AREAL_PROJECT_ARTIFACT_V1",
            "created_at": _dt_ff05.utcnow().isoformat() + "Z",
            "task_id": task_id,
            "topic_id": topic_id,
            "engine": "FULLFIX_05_REAL_PROJECT_ENGINE",
            "input": raw_input,
            "data": data,
            "files": {
                "pdf": pdf_path,
                "dxf": dxf_path,
            },
        }
        _Path_ff05(manifest_path).write_text(_json_ff05.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

        from core.engine_base import upload_artifact_to_drive
        pdf_link = upload_artifact_to_drive(pdf_path, task_id, topic_id)
        dxf_link = upload_artifact_to_drive(dxf_path, task_id, topic_id)
        manifest_link = upload_artifact_to_drive(manifest_path, task_id, topic_id)

        if not pdf_link:
            res["error"] = "PDF_UPLOAD_FAILED"
            return res
        if not dxf_link:
            res["error"] = "DXF_UPLOAD_FAILED"
            return res

        res.update({
            "success": True,
            "pdf_link": str(pdf_link),
            "dxf_link": str(dxf_link),
            "manifest_link": str(manifest_link or ""),
        })
        return res

    except Exception as e:
        res["error"] = str(e)[:500]
        return res

# === END FULLFIX_05_REAL_PROJECT_ENGINE ===


# === FULLFIX_06_FINAL_PROJECT_TEMPLATE_REPLAY ===
import os as _os_ff06
import re as _re_ff06
import json as _json_ff06
import glob as _glob_ff06
import tempfile as _tempfile_ff06
from pathlib import Path as _Path_ff06
from datetime import datetime as _dt_ff06

def _ff06_clean_text(v, limit=5000):
    return str(v or "").replace("\x00", " ").strip()[:limit]

def _ff06_find_font():
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for x in candidates:
        if _os_ff06.path.exists(x):
            return x
    return ""

def _ff06_latest_template(topic_id=0, preferred_section=""):
    base = "/root/.areal-neva-core/data/project_templates"
    files = sorted(
        _glob_ff06.glob(_os_ff06.path.join(base, "PROJECT_TEMPLATE_MODEL__*.json")),
        key=lambda x: _os_ff06.path.getmtime(x),
        reverse=True,
    )
    if not files:
        return {}
    topic_id = int(topic_id or 0)
    preferred_section = _ff06_clean_text(preferred_section).upper()
    best = None
    for fp in files:
        try:
            data = _json_ff06.loads(_Path_ff06(fp).read_text(encoding="utf-8"))
            data["_template_file"] = fp
            pt = _ff06_clean_text(data.get("project_type")).upper()
            dtid = int(data.get("topic_id", 0) or 0)
            if topic_id and dtid == topic_id and preferred_section and pt == preferred_section:
                return data
            if topic_id and dtid == topic_id and not best:
                best = data
            if preferred_section and pt == preferred_section and not best:
                best = data
            if not best:
                best = data
        except Exception:
            continue
    return best or {}

def _ff06_parse_request(raw_input, template_hint=""):
    text = _ff06_clean_text(raw_input, 10000)
    low = text.lower()
    out = {
        "section": "КЖ",
        "project_name": "Проект фундаментной плиты",
        "length_m": 10.0,
        "width_m": 10.0,
        "slab_mm": 200,
        "sand_mm": 300,
        "gravel_mm": 100,
        "concrete_class": "B25",
        "rebar_class": "A500",
        "rebar_diam_mm": 12,
        "rebar_step_mm": 200,
        "raw_input": text,
        "template_hint": template_hint or "",
    }

    if "кд" in low or "кров" in low or "строп" in low:
        out["section"] = "КД"
        out["project_name"] = "Проект кровли"
    if "ар" in low and "фундамент" not in low and "кров" not in low:
        out["section"] = "АР"
        out["project_name"] = "Архитектурный раздел"
    if "кж" in low or "фундамент" in low or "плит" in low:
        out["section"] = "КЖ"
        out["project_name"] = "Проект фундаментной плиты"

    m = _re_ff06.search(r"(\d+(?:[,.]\d+)?)\s*[xх×]\s*(\d+(?:[,.]\d+)?)\s*(?:м|m)?", low)
    if m:
        out["length_m"] = float(m.group(1).replace(",", "."))
        out["width_m"] = float(m.group(2).replace(",", "."))

    def mm(keys, default):
        for key in keys:
            mmx = _re_ff06.search(key + r".{0,40}?(\d{2,4})\s*мм", low)
            if mmx:
                return int(mmx.group(1))
        return default

    out["slab_mm"] = mm(["толщина", "плита", "бетон"], out["slab_mm"])
    out["sand_mm"] = mm(["песчан", "песок"], out["sand_mm"])
    out["gravel_mm"] = mm(["щеб", "основан"], out["gravel_mm"])
    out["rebar_step_mm"] = mm(["шаг"], out["rebar_step_mm"])

    md = _re_ff06.search(r"(?:ø|Ø|ф|диаметр)\s*(\d{1,2})", low)
    if md:
        out["rebar_diam_mm"] = int(md.group(1))
    mc = _re_ff06.search(r"\b[вbВB]\s*([123456789]\d)\b", text)
    if mc:
        out["concrete_class"] = "B" + mc.group(1)
    ma = _re_ff06.search(r"\bA\s*([245]\d{2})\b", text, _re_ff06.I)
    if ma:
        out["rebar_class"] = "A" + ma.group(1)

    return out

def _ff06_sheet_rows(sheet_register, section):
    rows = []
    for i, sh in enumerate(sheet_register or [], 1):
        if isinstance(sh, dict):
            mark = _ff06_clean_text(sh.get("mark") or section, 20) or section
            num = _ff06_clean_text(sh.get("number") or str(i), 30) or str(i)
            title = _ff06_clean_text(sh.get("title") or sh.get("name") or "", 180)
            if not title:
                title = f"Лист {num}"
            rows.append({"mark": mark, "number": num, "title": title})
        else:
            raw = _ff06_clean_text(sh, 180)
            if not raw:
                continue
            rows.append({"mark": section, "number": str(i), "title": raw})
    return rows

def _ff06_material_rows(materials, data):
    rows = []
    for x in materials or []:
        if isinstance(x, dict):
            name = _ff06_clean_text(x.get("name") or x.get("material") or x.get("title") or "Материал", 180)
            unit = _ff06_clean_text(x.get("unit") or x.get("ед") or "шт", 30)
            qty = x.get("quantity", x.get("qty", x.get("count", "-")))
            note = _ff06_clean_text(x.get("note") or "", 120)
            rows.append((name, unit, qty, note))
        else:
            name = _ff06_clean_text(x, 180)
            if name:
                rows.append((name, "по проекту", "-", "из шаблона"))
    if rows:
        return rows[:80]

    L = float(data["length_m"])
    W = float(data["width_m"])
    area = L * W
    slab = int(data["slab_mm"])
    sand = int(data["sand_mm"])
    gravel = int(data["gravel_mm"])
    step = max(int(data["rebar_step_mm"]), 50)
    bars_x = int((W * 1000) / step) + 1
    bars_y = int((L * 1000) / step) + 1
    rebar_m = round((bars_x * L + bars_y * W) * 2, 1)

    return [
        (f"Бетон {data['concrete_class']}", "м3", round(area * slab / 1000, 3), "фундаментная плита"),
        ("Песчаная подушка", "м3", round(area * sand / 1000, 3), ""),
        ("Щебёночное основание", "м3", round(area * gravel / 1000, 3), ""),
        (f"Арматура {data['rebar_class']} Ø{data['rebar_diam_mm']} шаг {step}", "п.м", rebar_m, "верхняя и нижняя сетка"),
    ]

def _ff06_register_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    font_path = _ff06_find_font()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("ArealDejaVu", font_path))
            return "ArealDejaVu"
        except Exception:
            pass
    return "Helvetica"

def _ff06_draw_frame(c, page_w, page_h, sheet, sheet_no, sheet_total, font):
    from reportlab.lib.units import mm
    c.setLineWidth(0.7)
    c.rect(10*mm, 10*mm, page_w - 20*mm, page_h - 20*mm)
    c.rect(page_w - 190*mm, 10*mm, 180*mm, 40*mm)
    c.line(page_w - 190*mm, 30*mm, page_w - 10*mm, 30*mm)
    c.line(page_w - 80*mm, 10*mm, page_w - 80*mm, 50*mm)
    c.line(page_w - 45*mm, 10*mm, page_w - 45*mm, 50*mm)
    c.setFont(font, 8)
    c.drawString(page_w - 187*mm, 42*mm, "AREAL-NEVA")
    c.drawString(page_w - 187*mm, 34*mm, _ff06_clean_text(sheet.get("title"), 80))
    c.drawString(page_w - 77*mm, 34*mm, f"Лист {sheet_no}")
    c.drawString(page_w - 42*mm, 34*mm, f"Листов {sheet_total}")
    c.setFont(font, 14)
    c.drawString(18*mm, page_h - 22*mm, f"{sheet.get('mark')} {sheet.get('number')} — {sheet.get('title')}")

def _ff06_write_pdf(path, data, template, sheets, material_rows):
    from reportlab.lib.pagesizes import A3, landscape
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas

    font = _ff06_register_font()
    page_w, page_h = landscape(A3)
    c = canvas.Canvas(path, pagesize=landscape(A3))

    L = float(data["length_m"])
    W = float(data["width_m"])
    slab = int(data["slab_mm"])
    sand = int(data["sand_mm"])
    gravel = int(data["gravel_mm"])
    step = int(data["rebar_step_mm"])
    rd = int(data["rebar_diam_mm"])
    section = data["section"]

    for idx, sheet in enumerate(sheets, 1):
        _ff06_draw_frame(c, page_w, page_h, sheet, idx, len(sheets), font)
        title_low = (sheet.get("title") or "").lower()

        if any(x in title_low for x in ("общие", "данные", "пояснит", "исходн")):
            y = 260
            lines = [
                f"Проект: {data['project_name']}",
                f"Раздел: {section}",
                f"Основа генерации: сохранённый PROJECT_TEMPLATE_MODEL",
                f"Шаблон: {template.get('_template_file', '')}",
                f"Плита: {L:g} x {W:g} м, толщина {slab} мм",
                f"Основание: щебень {gravel} мм, песчаная подушка {sand} мм",
                f"Бетон: {data['concrete_class']}",
                f"Арматура: {data['rebar_class']} Ø{rd}, шаг {step} мм",
            ]
            c.setFont(font, 10)
            for line in lines:
                c.drawString(25*mm, y*mm, line)
                y -= 8

        elif any(x in title_low for x in ("ведомость лист", "состав лист", "листов")):
            y = 260
            c.setFont(font, 10)
            for j, sh in enumerate(sheets, 1):
                c.drawString(25*mm, y*mm, f"{j}. {sh['mark']} {sh['number']} — {sh['title']}")
                y -= 7

        elif any(x in title_low for x in ("план", "плит", "схема")):
            x0, y0 = 70*mm, 60*mm
            scale = min((page_w - 150*mm) / (L * 1000), 115*mm / (W * 1000))
            rw = L * 1000 * scale
            rh = W * 1000 * scale
            c.setLineWidth(1.1)
            c.rect(x0, y0, rw, rh)
            c.setDash(4, 3)
            c.line(x0, y0 + rh / 2, x0 + rw, y0 + rh / 2)
            c.line(x0 + rw / 2, y0, x0 + rw / 2, y0 + rh)
            c.setDash()
            step_draw = max(0.7*mm, step * scale)
            c.setLineWidth(0.25)
            xx = x0 + step_draw
            while xx < x0 + rw:
                c.line(xx, y0, xx, y0 + rh)
                xx += step_draw
            yy = y0 + step_draw
            while yy < y0 + rh:
                c.line(x0, yy, x0 + rw, yy)
                yy += step_draw
            c.setFont(font, 9)
            c.drawString(x0, y0 - 8*mm, f"{L:g} м")
            c.saveState()
            c.translate(x0 - 8*mm, y0)
            c.rotate(90)
            c.drawString(0, 0, f"{W:g} м")
            c.restoreState()
            c.setFont(font, 10)
            c.drawString(25*mm, 260*mm, f"Армирование: {data['rebar_class']} Ø{rd}, шаг {step} мм")

        elif any(x in title_low for x in ("разрез", "сечени", "1-1")):
            bx, by = 55*mm, 70*mm
            total = slab + gravel + sand
            k = 105*mm / total
            ycur = by
            c.setLineWidth(0.9)
            for name, th, note in [
                ("Фундаментная плита", slab, f"Бетон {data['concrete_class']}"),
                ("Щебёночное основание", gravel, "Щебень"),
                ("Песчаная подушка", sand, "Песок"),
            ]:
                hh = th * k
                c.rect(bx, ycur, 230*mm, hh)
                c.setFont(font, 10)
                c.drawString(bx + 5*mm, ycur + hh/2, f"{name}: {th} мм — {note}")
                ycur += hh
            c.setFont(font, 10)
            c.drawString(25*mm, 240*mm, "Разрез 1-1")

        elif any(x in title_low for x in ("армир", "арматур", "сетка")):
            x0, y0 = 70*mm, 60*mm
            scale = min((page_w - 150*mm) / (L * 1000), 115*mm / (W * 1000))
            rw = L * 1000 * scale
            rh = W * 1000 * scale
            c.setLineWidth(1.0)
            c.rect(x0, y0, rw, rh)
            step_draw = max(1.0*mm, step * scale)
            c.setLineWidth(0.35)
            xx = x0 + step_draw
            while xx < x0 + rw:
                c.line(xx, y0, xx, y0 + rh)
                xx += step_draw
            yy = y0 + step_draw
            while yy < y0 + rh:
                c.line(x0, yy, x0 + rw, yy)
                yy += step_draw
            c.setFont(font, 10)
            c.drawString(25*mm, 260*mm, f"Верхняя и нижняя сетки: {data['rebar_class']} Ø{rd}, шаг {step} мм")

        elif any(x in title_low for x in ("специф", "материал", "ведомость объем", "ведомость объём")):
            y = 260
            c.setFont(font, 10)
            c.drawString(25*mm, y*mm, "№")
            c.drawString(40*mm, y*mm, "Наименование")
            c.drawString(170*mm, y*mm, "Ед")
            c.drawString(195*mm, y*mm, "Кол-во")
            c.drawString(230*mm, y*mm, "Примечание")
            y -= 8
            for j, row in enumerate(material_rows[:28], 1):
                c.drawString(25*mm, y*mm, str(j))
                c.drawString(40*mm, y*mm, _ff06_clean_text(row[0], 70))
                c.drawString(170*mm, y*mm, _ff06_clean_text(row[1], 12))
                c.drawString(195*mm, y*mm, _ff06_clean_text(row[2], 18))
                c.drawString(230*mm, y*mm, _ff06_clean_text(row[3], 50))
                y -= 7

        else:
            y = 250
            c.setFont(font, 10)
            c.drawString(25*mm, y*mm, f"Лист выполнен по структуре шаблона: {sheet['title']}")
            y -= 10
            c.drawString(25*mm, y*mm, f"Раздел: {section}")
            y -= 8
            c.drawString(25*mm, y*mm, f"Параметры: {L:g} x {W:g} м, бетон {data['concrete_class']}")

        c.showPage()

    c.save()
    return path

def _ff06_write_dxf(path, data, sheets, material_rows):
    import ezdxf
    doc = ezdxf.new("R2010")
    doc.header["$INSUNITS"] = 4
    msp = doc.modelspace()

    for layer, color in [("KJ_PLAN", 7), ("KJ_AXES", 2), ("KJ_REBAR", 3), ("KJ_TEXT", 1), ("KJ_SECTION", 5)]:
        if layer not in doc.layers:
            doc.layers.add(layer, color=color)

    L = float(data["length_m"]) * 1000
    W = float(data["width_m"]) * 1000
    slab = int(data["slab_mm"])
    sand = int(data["sand_mm"])
    gravel = int(data["gravel_mm"])
    step = int(data["rebar_step_mm"])

    msp.add_lwpolyline([(0,0), (L,0), (L,W), (0,W), (0,0)], dxfattribs={"layer": "KJ_PLAN"})
    msp.add_line((L/2, 0), (L/2, W), dxfattribs={"layer": "KJ_AXES"})
    msp.add_line((0, W/2), (L, W/2), dxfattribs={"layer": "KJ_AXES"})

    x = step
    while x < L:
        msp.add_line((x, 0), (x, W), dxfattribs={"layer": "KJ_REBAR"})
        x += step
    y = step
    while y < W:
        msp.add_line((0, y), (L, y), dxfattribs={"layer": "KJ_REBAR"})
        y += step

    msp.add_text(f"FOUNDATION SLAB {L/1000:g} x {W/1000:g} m", dxfattribs={"height": 250, "layer": "KJ_TEXT"}).set_placement((0, -700))
    msp.add_text(f"SLAB {slab} mm / GRAVEL {gravel} mm / SAND {sand} mm", dxfattribs={"height": 250, "layer": "KJ_TEXT"}).set_placement((0, -1100))
    msp.add_text(f"REBAR {data['rebar_class']} D{data['rebar_diam_mm']} STEP {step} mm", dxfattribs={"height": 250, "layer": "KJ_TEXT"}).set_placement((0, -1500))

    sx, sy = 0, -3000
    widths = [slab, gravel, sand]
    names = ["SLAB", "GRAVEL", "SAND"]
    cur = sy
    for name, th in zip(names, widths):
        msp.add_lwpolyline([(sx,cur), (sx+L,cur), (sx+L,cur-th), (sx,cur-th), (sx,cur)], dxfattribs={"layer": "KJ_SECTION"})
        msp.add_text(f"{name} {th} mm", dxfattribs={"height": 220, "layer": "KJ_TEXT"}).set_placement((sx + 300, cur - th/2))
        cur -= th

    doc.saveas(path)
    return path

def _ff06_write_xlsx(path, data, sheets, material_rows, template):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment
    wb = Workbook()

    ws = wb.active
    ws.title = "Ведомость листов"
    ws.append(["№", "Марка", "Лист", "Наименование"])
    for c in ws[1]:
        c.font = Font(bold=True)
        c.alignment = Alignment(horizontal="center")
    for i, sh in enumerate(sheets, 1):
        ws.append([i, sh["mark"], sh["number"], sh["title"]])
    ws.column_dimensions["D"].width = 80

    ws2 = wb.create_sheet("Спецификация")
    ws2.append(["№", "Наименование", "Ед. изм", "Кол-во", "Примечание"])
    for c in ws2[1]:
        c.font = Font(bold=True)
    for i, row in enumerate(material_rows, 1):
        ws2.append([i, row[0], row[1], row[2], row[3]])
    ws2.column_dimensions["B"].width = 70
    ws2.column_dimensions["E"].width = 50

    ws3 = wb.create_sheet("Параметры")
    for k in ["project_name", "section", "length_m", "width_m", "slab_mm", "sand_mm", "gravel_mm", "concrete_class", "rebar_class", "rebar_diam_mm", "rebar_step_mm"]:
        ws3.append([k, data.get(k)])
    ws3.append(["template_file", template.get("_template_file", "")])
    ws3.column_dimensions["A"].width = 25
    ws3.column_dimensions["B"].width = 80

    wb.save(path)
    return path

async def create_project_pdf_dxf_artifact(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", require_template: bool = True) -> dict:
    data = _ff06_parse_request(raw_input, template_hint)
    template = _ff06_latest_template(topic_id, data.get("section"))

    res = {
        "success": False,
        "engine": "FULLFIX_06_FINAL_PROJECT_TEMPLATE_REPLAY",
        "section": data["section"],
        "pdf_path": "",
        "dxf_path": "",
        "xlsx_path": "",
        "manifest_path": "",
        "pdf_link": "",
        "dxf_link": "",
        "xlsx_link": "",
        "manifest_link": "",
        "template_file": "",
        "sheet_count": 0,
        "error": None,
        "data": data,
    }

    if require_template and not template:
        res["error"] = "PROJECT_TEMPLATE_MODEL_NOT_FOUND"
        return res

    sheets = _ff06_sheet_rows((template or {}).get("sheet_register") or [], data["section"])
    if require_template and not sheets:
        res["error"] = "PROJECT_TEMPLATE_MODEL_HAS_EMPTY_SHEET_REGISTER"
        res["template_file"] = (template or {}).get("_template_file", "")
        return res

    if not sheets:
        res["error"] = "NO_TEMPLATE_SHEETS_NO_PROJECT"
        return res

    material_rows = _ff06_material_rows((template or {}).get("materials") or [], data)

    stamp = _dt_ff06.utcnow().strftime("%Y%m%d_%H%M%S")
    safe_task = _re_ff06.sub(r"[^A-Za-z0-9_-]+", "_", str(task_id or "manual"))[:24]
    out_dir = _Path_ff06(_tempfile_ff06.gettempdir()) / f"areal_project_ff06_{safe_task}_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = str(out_dir / f"{data['section']}_PROJECT_{safe_task}.pdf")
    dxf_path = str(out_dir / f"{data['section']}_PROJECT_{safe_task}.dxf")
    xlsx_path = str(out_dir / f"{data['section']}_SPEC_{safe_task}.xlsx")
    manifest_path = str(out_dir / f"{data['section']}_MANIFEST_{safe_task}.json")

    try:
        _ff06_write_pdf(pdf_path, data, template, sheets, material_rows)
        _ff06_write_dxf(dxf_path, data, sheets, material_rows)
        _ff06_write_xlsx(xlsx_path, data, sheets, material_rows, template)

        checks = [
            ("PDF_NOT_CREATED", pdf_path, 2500),
            ("DXF_NOT_CREATED", dxf_path, 500),
            ("XLSX_NOT_CREATED", xlsx_path, 1000),
        ]
        for err, fp, min_size in checks:
            if not _os_ff06.path.exists(fp) or _os_ff06.path.getsize(fp) < min_size:
                res["error"] = err
                return res

        manifest = {
            "schema": "AREAL_PROJECT_ARTIFACT_V3",
            "engine": "FULLFIX_06_FINAL_PROJECT_TEMPLATE_REPLAY",
            "created_at": _dt_ff06.utcnow().isoformat() + "Z",
            "task_id": task_id,
            "topic_id": topic_id,
            "input": raw_input,
            "template_file": template.get("_template_file", ""),
            "sheet_count": len(sheets),
            "sheets": sheets,
            "data": data,
            "artifacts": {
                "pdf_path": pdf_path,
                "dxf_path": dxf_path,
                "xlsx_path": xlsx_path,
            },
        }
        _Path_ff06(manifest_path).write_text(_json_ff06.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

        from core.engine_base import upload_artifact_to_drive
        pdf_link = upload_artifact_to_drive(pdf_path, task_id, topic_id)
        dxf_link = upload_artifact_to_drive(dxf_path, task_id, topic_id)
        xlsx_link = upload_artifact_to_drive(xlsx_path, task_id, topic_id)
        manifest_link = upload_artifact_to_drive(manifest_path, task_id, topic_id)

        if not pdf_link:
            res["error"] = "PDF_UPLOAD_FAILED"
            return res
        if not dxf_link:
            res["error"] = "DXF_UPLOAD_FAILED"
            return res
        if not xlsx_link:
            res["error"] = "XLSX_UPLOAD_FAILED"
            return res

        res.update({
            "success": True,
            "pdf_path": pdf_path,
            "dxf_path": dxf_path,
            "xlsx_path": xlsx_path,
            "manifest_path": manifest_path,
            "pdf_link": str(pdf_link),
            "dxf_link": str(dxf_link),
            "xlsx_link": str(xlsx_link),
            "manifest_link": str(manifest_link or ""),
            "template_file": template.get("_template_file", ""),
            "sheet_count": len(sheets),
        })
        return res

    except Exception as e:
        res["error"] = str(e)[:500]
        return res

# === END FULLFIX_06_FINAL_PROJECT_TEMPLATE_REPLAY ===

# === FULLFIX_07_PROJECT_DESIGN_CLOSURE ===
import os as _os_ff07
import re as _re_ff07
import json as _json_ff07
import glob as _glob_ff07
import math as _math_ff07
import tempfile as _tempfile_ff07
from pathlib import Path as _Path_ff07
from datetime import datetime as _dt_ff07

def _ff07_clean(v, limit=5000):
    return str(v or "").replace("\x00", " ").strip()[:limit]

def _ff07_template_files():
    base = "/root/.areal-neva-core/data/project_templates"
    return sorted(
        _glob_ff07.glob(_os_ff07.path.join(base, "PROJECT_TEMPLATE_MODEL__*.json")),
        key=lambda x: _os_ff07.path.getmtime(x),
        reverse=True,
    )

def _ff07_load_latest_template(topic_id=0, preferred_section=""):
    files = _ff07_template_files()
    best = {}
    preferred_section = _ff07_clean(preferred_section).upper()
    topic_id = int(topic_id or 0)

    for fp in files:
        try:
            data = _json_ff07.loads(_Path_ff07(fp).read_text(encoding="utf-8"))
            if not isinstance(data, dict):
                continue
            data["_template_file"] = fp
            pt = _ff07_clean(data.get("project_type")).upper()
            dtid = int(data.get("topic_id", 0) or 0)

            if topic_id and dtid == topic_id and preferred_section and pt == preferred_section:
                return data
            if preferred_section and pt == preferred_section and not best:
                best = data
            if topic_id and dtid == topic_id and not best:
                best = data
            if not best:
                best = data
        except Exception:
            continue

    return best or {}

def _ff07_parse_request(raw_input, template_hint=""):
    text = _ff07_clean(raw_input, 12000)
    low = text.lower()

    data = {
        "section": "КЖ",
        "project_name": "Проект фундаментной плиты",
        "length_m": 10.0,
        "width_m": 10.0,
        "slab_mm": 200,
        "sand_mm": 300,
        "gravel_mm": 100,
        "concrete_class": "B25",
        "rebar_class": "A500",
        "rebar_diam_mm": 12,
        "rebar_step_mm": 200,
        "raw_input": text,
        "template_hint": template_hint or "",
    }

    if any(x in low for x in ("кров", "строп", "кд")) and not any(x in low for x in ("фундамент", "плит")):
        data["section"] = "КД"
        data["project_name"] = "Проект кровли"
    if any(x in low for x in ("архитектур", " ар ", "раздел ар")) and not any(x in low for x in ("фундамент", "плит", "кров")):
        data["section"] = "АР"
        data["project_name"] = "Архитектурный раздел"
    if any(x in low for x in ("фундамент", "плит", "кж")):
        data["section"] = "КЖ"
        data["project_name"] = "Проект фундаментной плиты"

    m = _re_ff07.search(r"(\d+(?:[,.]\d+)?)\s*[xх×]\s*(\d+(?:[,.]\d+)?)\s*(?:м|m)?", low)
    if m:
        data["length_m"] = float(m.group(1).replace(",", "."))
        data["width_m"] = float(m.group(2).replace(",", "."))

    def find_mm(keys, default):
        for key in keys:
            mm = _re_ff07.search(key + r".{0,45}?(\d{2,4})\s*мм", low)
            if mm:
                return int(mm.group(1))
        return default

    data["slab_mm"] = find_mm(("толщина", "плита", "бетон"), data["slab_mm"])
    data["sand_mm"] = find_mm(("песчан", "песок"), data["sand_mm"])
    data["gravel_mm"] = find_mm(("щеб", "основан"), data["gravel_mm"])
    data["rebar_step_mm"] = find_mm(("шаг",), data["rebar_step_mm"])

    md = _re_ff07.search(r"(?:ø|Ø|ф|диаметр)\s*(\d{1,2})", text, _re_ff07.I)
    if md:
        data["rebar_diam_mm"] = int(md.group(1))

    mc = _re_ff07.search(r"\b[вbВB]\s*([123456789]\d)\b", text)
    if mc:
        data["concrete_class"] = "B" + mc.group(1)

    ma = _re_ff07.search(r"\b[аaАA]\s*([245]\d{2})\b", text)
    if ma:
        data["rebar_class"] = "A" + ma.group(1)

    return data

def _ff07_dedup_titles(rows):
    out, seen = [], set()
    for row in rows:
        title = _ff07_clean(row.get("title"), 160)
        key = title.lower()
        if not title or key in seen:
            continue
        seen.add(key)
        out.append(row)
    return out

def _ff07_sheet_rows_from_template(template, section, data):
    raw = template.get("sheet_register") or []
    rows = []

    for i, sh in enumerate(raw, 1):
        if isinstance(sh, dict):
            rows.append({
                "mark": _ff07_clean(sh.get("mark") or section, 20) or section,
                "number": _ff07_clean(sh.get("number") or str(i), 30) or str(i),
                "title": _ff07_clean(sh.get("title") or sh.get("name") or "", 180),
                "source": "sheet_register",
            })
        else:
            rows.append({
                "mark": section,
                "number": str(i),
                "title": _ff07_clean(sh, 180),
                "source": "sheet_register_raw",
            })

    rows = _ff07_dedup_titles(rows)

    if len(rows) >= 8:
        return rows

    sections = []
    for x in template.get("sections") or []:
        if isinstance(x, dict):
            t = _ff07_clean(x.get("title") or x.get("name") or x.get("text"), 180)
        else:
            t = _ff07_clean(x, 180)
        if t:
            sections.append(t)

    canonical = []
    if section == "КЖ":
        canonical = [
            "Общие данные",
            "Ведомость листов",
            "План фундаментной плиты",
            "Разрез 1-1",
            "Схема армирования нижней сетки",
            "Схема армирования верхней сетки",
            "Узлы армирования и защитные слои",
            "Спецификация материалов",
            "Ведомость расхода стали",
            "Пояснительная записка",
        ]
    elif section == "КД":
        canonical = [
            "Общие данные",
            "Ведомость листов",
            "План кровли",
            "План стропильной системы",
            "Разрезы кровли",
            "Узлы кровли",
            "Схема обрешётки",
            "Спецификация пиломатериалов",
            "Ведомость элементов",
            "Пояснительная записка",
        ]
    elif section == "АР":
        canonical = [
            "Общие данные",
            "Ведомость листов",
            "План этажа",
            "Фасады",
            "Разрезы",
            "План кровли",
            "Экспликация помещений",
            "Спецификация заполнения проёмов",
            "Узлы",
            "Пояснительная записка",
        ]
    else:
        canonical = [
            "Общие данные",
            "Ведомость листов",
            "План",
            "Разрез",
            "Схема",
            "Узлы",
            "Спецификация материалов",
            "Пояснительная записка",
        ]

    for title in canonical:
        rows.append({
            "mark": section,
            "number": str(len(rows) + 1),
            "title": title,
            "source": "canonical_required",
        })

    for sec in sections:
        low = sec.lower()
        if any(k in low for k in ("общие данные", "ведомость", "план", "фасад", "разрез", "узел", "схема", "спецификац", "расчет", "расчёт", "конструктив")):
            rows.append({
                "mark": section,
                "number": str(len(rows) + 1),
                "title": sec[:150],
                "source": "template_sections",
            })

    rows = _ff07_dedup_titles(rows)

    renumbered = []
    for i, row in enumerate(rows[:24], 1):
        row = dict(row)
        row["mark"] = row.get("mark") or section
        row["number"] = str(i)
        renumbered.append(row)

    return renumbered

def _ff07_material_rows(template, data):
    rows = []
    for x in template.get("materials") or []:
        if isinstance(x, dict):
            rows.append({
                "name": _ff07_clean(x.get("name") or x.get("material") or x.get("title") or "Материал", 180),
                "unit": _ff07_clean(x.get("unit") or "шт", 30),
                "qty": x.get("qty", x.get("quantity", "-")),
                "note": _ff07_clean(x.get("note") or "из шаблона", 120),
            })
        else:
            t = _ff07_clean(x, 180)
            if t:
                rows.append({"name": t, "unit": "по проекту", "qty": "-", "note": "из шаблона"})

    L = float(data["length_m"])
    W = float(data["width_m"])
    area = L * W
    slab = int(data["slab_mm"])
    sand = int(data["sand_mm"])
    gravel = int(data["gravel_mm"])
    step = max(int(data["rebar_step_mm"]), 50)
    bars_x = int((W * 1000) / step) + 1
    bars_y = int((L * 1000) / step) + 1
    rebar_m = round((bars_x * L + bars_y * W) * 2, 1)

    calc_rows = [
        {"name": f"Бетон {data['concrete_class']}", "unit": "м3", "qty": round(area * slab / 1000, 3), "note": "фундаментная плита"},
        {"name": "Песчаная подушка", "unit": "м3", "qty": round(area * sand / 1000, 3), "note": ""},
        {"name": "Щебёночное основание", "unit": "м3", "qty": round(area * gravel / 1000, 3), "note": ""},
        {"name": f"Арматура {data['rebar_class']} Ø{data['rebar_diam_mm']} шаг {step}", "unit": "п.м", "qty": rebar_m, "note": "верхняя и нижняя сетка"},
    ]

    names = {r["name"].lower() for r in rows}
    for r in calc_rows:
        if r["name"].lower() not in names:
            rows.append(r)

    return rows[:80]

def _ff07_register_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for fp in candidates:
        if _os_ff07.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont("ArealDejaVu", fp))
                return "ArealDejaVu"
            except Exception:
                pass
    return "Helvetica"

def _ff07_draw_frame(c, page_w, page_h, sheet, sheet_no, sheet_total, font):
    from reportlab.lib.units import mm

    c.setLineWidth(0.7)
    c.rect(10 * mm, 10 * mm, page_w - 20 * mm, page_h - 20 * mm)
    c.rect(page_w - 190 * mm, 10 * mm, 180 * mm, 42 * mm)
    c.line(page_w - 190 * mm, 32 * mm, page_w - 10 * mm, 32 * mm)
    c.line(page_w - 80 * mm, 10 * mm, page_w - 80 * mm, 52 * mm)
    c.line(page_w - 45 * mm, 10 * mm, page_w - 45 * mm, 52 * mm)

    c.setFont(font, 8)
    c.drawString(page_w - 187 * mm, 44 * mm, "AREAL-NEVA")
    c.drawString(page_w - 187 * mm, 36 * mm, _ff07_clean(sheet.get("title"), 80))
    c.drawString(page_w - 77 * mm, 36 * mm, f"Лист {sheet_no}")
    c.drawString(page_w - 42 * mm, 36 * mm, f"Листов {sheet_total}")

    c.setFont(font, 14)
    c.drawString(18 * mm, page_h - 22 * mm, f"{sheet.get('mark')} {sheet.get('number')} — {sheet.get('title')}")

def _ff07_write_pdf(path, data, template, sheets, materials):
    from reportlab.lib.pagesizes import A3, landscape
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas

    font = _ff07_register_font()
    page_w, page_h = landscape(A3)
    c = canvas.Canvas(path, pagesize=landscape(A3))

    L = float(data["length_m"])
    W = float(data["width_m"])
    slab = int(data["slab_mm"])
    sand = int(data["sand_mm"])
    gravel = int(data["gravel_mm"])
    step = int(data["rebar_step_mm"])
    rd = int(data["rebar_diam_mm"])
    section = data["section"]

    for idx, sheet in enumerate(sheets, 1):
        title = sheet.get("title") or ""
        low = title.lower()
        _ff07_draw_frame(c, page_w, page_h, sheet, idx, len(sheets), font)

        if any(x in low for x in ("общие", "данные", "пояснит", "исходн")):
            y = 260
            lines = [
                f"Проект: {data['project_name']}",
                f"Раздел: {section}",
                "Тип выдачи: PDF + DXF + XLSX + MANIFEST",
                f"Шаблон: {template.get('_template_file', '')}",
                f"Плита: {L:g} x {W:g} м, толщина {slab} мм",
                f"Основание: щебень {gravel} мм, песчаная подушка {sand} мм",
                f"Бетон: {data['concrete_class']}",
                f"Арматура: {data['rebar_class']} Ø{rd}, шаг {step} мм",
                f"Листов: {len(sheets)}",
            ]
            c.setFont(font, 10)
            for line in lines:
                c.drawString(25 * mm, y * mm, line)
                y -= 8

        elif any(x in low for x in ("ведомость лист", "состав лист", "листов")):
            y = 260
            c.setFont(font, 10)
            for j, sh in enumerate(sheets, 1):
                c.drawString(25 * mm, y * mm, f"{j}. {sh['mark']} {sh['number']} — {sh['title']}")
                y -= 7
                if y < 35:
                    c.showPage()
                    _ff07_draw_frame(c, page_w, page_h, sheet, idx, len(sheets), font)
                    y = 260

        elif any(x in low for x in ("план", "плит", "схема")) and not any(x in low for x in ("армир", "арматур")):
            x0, y0 = 70 * mm, 60 * mm
            scale = min((page_w - 150 * mm) / (L * 1000), 115 * mm / (W * 1000))
            rw = L * 1000 * scale
            rh = W * 1000 * scale

            c.setLineWidth(1.1)
            c.rect(x0, y0, rw, rh)
            c.setDash(4, 3)
            c.line(x0, y0 + rh / 2, x0 + rw, y0 + rh / 2)
            c.line(x0 + rw / 2, y0, x0 + rw / 2, y0 + rh)
            c.setDash()

            step_draw = max(0.7 * mm, step * scale)
            c.setLineWidth(0.25)

            xx = x0 + step_draw
            while xx < x0 + rw:
                c.line(xx, y0, xx, y0 + rh)
                xx += step_draw

            yy = y0 + step_draw
            while yy < y0 + rh:
                c.line(x0, yy, x0 + rw, yy)
                yy += step_draw

            c.setFont(font, 9)
            c.drawString(x0, y0 - 8 * mm, f"{L:g} м")
            c.saveState()
            c.translate(x0 - 8 * mm, y0)
            c.rotate(90)
            c.drawString(0, 0, f"{W:g} м")
            c.restoreState()

            c.setFont(font, 10)
            c.drawString(25 * mm, 260 * mm, f"Армирование: {data['rebar_class']} Ø{rd}, шаг {step} мм")

        elif any(x in low for x in ("разрез", "сечени", "1-1")):
            bx, by = 55 * mm, 70 * mm
            total = slab + gravel + sand
            k = 105 * mm / total
            ycur = by

            c.setLineWidth(0.9)
            for name, th, note in [
                ("Фундаментная плита", slab, f"Бетон {data['concrete_class']}"),
                ("Щебёночное основание", gravel, "Щебень"),
                ("Песчаная подушка", sand, "Песок"),
            ]:
                hh = th * k
                c.rect(bx, ycur, 230 * mm, hh)
                c.setFont(font, 10)
                c.drawString(bx + 5 * mm, ycur + hh / 2, f"{name}: {th} мм — {note}")
                ycur += hh

            c.setFont(font, 10)
            c.drawString(25 * mm, 240 * mm, "Разрез 1-1")

        elif any(x in low for x in ("армир", "арматур", "сетка")):
            x0, y0 = 70 * mm, 60 * mm
            scale = min((page_w - 150 * mm) / (L * 1000), 115 * mm / (W * 1000))
            rw = L * 1000 * scale
            rh = W * 1000 * scale

            c.setLineWidth(1.0)
            c.rect(x0, y0, rw, rh)

            step_draw = max(1.0 * mm, step * scale)
            c.setLineWidth(0.35)

            xx = x0 + step_draw
            while xx < x0 + rw:
                c.line(xx, y0, xx, y0 + rh)
                xx += step_draw

            yy = y0 + step_draw
            while yy < y0 + rh:
                c.line(x0, yy, x0 + rw, yy)
                yy += step_draw

            c.setFont(font, 10)
            c.drawString(25 * mm, 260 * mm, f"Верхняя и нижняя сетки: {data['rebar_class']} Ø{rd}, шаг {step} мм")

        elif any(x in low for x in ("специф", "материал", "ведомость расход", "ведомость объем", "ведомость объём", "стали")):
            y = 260
            c.setFont(font, 10)
            c.drawString(25 * mm, y * mm, "№")
            c.drawString(40 * mm, y * mm, "Наименование")
            c.drawString(175 * mm, y * mm, "Ед")
            c.drawString(200 * mm, y * mm, "Кол-во")
            c.drawString(235 * mm, y * mm, "Примечание")
            y -= 8

            for j, row in enumerate(materials[:32], 1):
                c.drawString(25 * mm, y * mm, str(j))
                c.drawString(40 * mm, y * mm, _ff07_clean(row["name"], 70))
                c.drawString(175 * mm, y * mm, _ff07_clean(row["unit"], 12))
                c.drawString(200 * mm, y * mm, _ff07_clean(row["qty"], 18))
                c.drawString(235 * mm, y * mm, _ff07_clean(row["note"], 50))
                y -= 7

        elif any(x in low for x in ("узел", "защитн", "слои")):
            c.setFont(font, 10)
            y = 250
            for line in [
                f"Защитный слой бетона принят по СП 63.13330.2018",
                f"Арматура {data['rebar_class']} Ø{rd}, шаг {step} мм",
                "Стыковка и нахлёсты выполнять по рабочей документации",
                "Геометрия узлов уточняется по месту и исполнительным размерам",
            ]:
                c.drawString(25 * mm, y * mm, line)
                y -= 9

        else:
            c.setFont(font, 10)
            y = 250
            for line in [
                f"Лист выполнен в составе комплекта: {title}",
                f"Раздел: {section}",
                f"Параметры: {L:g} x {W:g} м",
                f"Бетон: {data['concrete_class']}",
                f"Шаблон: {template.get('_template_file', '')}",
            ]:
                c.drawString(25 * mm, y * mm, line)
                y -= 8

        c.showPage()

    c.save()
    return path

def _ff07_write_dxf(path, data, sheets, materials):
    import ezdxf

    doc = ezdxf.new("R2010")
    doc.header["$INSUNITS"] = 4
    msp = doc.modelspace()

    for layer, color in [
        ("KJ_PLAN", 7),
        ("KJ_AXES", 2),
        ("KJ_REBAR", 3),
        ("KJ_TEXT", 1),
        ("KJ_SECTION", 5),
        ("KJ_SHEETS", 4),
    ]:
        if layer not in doc.layers:
            doc.layers.add(layer, color=color)

    L = float(data["length_m"]) * 1000
    W = float(data["width_m"]) * 1000
    slab = int(data["slab_mm"])
    sand = int(data["sand_mm"])
    gravel = int(data["gravel_mm"])
    step = max(int(data["rebar_step_mm"]), 50)

    msp.add_lwpolyline([(0, 0), (L, 0), (L, W), (0, W), (0, 0)], dxfattribs={"layer": "KJ_PLAN"})
    msp.add_line((L / 2, 0), (L / 2, W), dxfattribs={"layer": "KJ_AXES"})
    msp.add_line((0, W / 2), (L, W / 2), dxfattribs={"layer": "KJ_AXES"})

    x = step
    while x < L:
        msp.add_line((x, 0), (x, W), dxfattribs={"layer": "KJ_REBAR"})
        x += step

    y = step
    while y < W:
        msp.add_line((0, y), (L, y), dxfattribs={"layer": "KJ_REBAR"})
        y += step

    msp.add_text(
        f"FOUNDATION SLAB {L/1000:g} x {W/1000:g} m",
        dxfattribs={"height": 250, "layer": "KJ_TEXT"},
    ).set_placement((0, -700))

    msp.add_text(
        f"SLAB {slab} mm / GRAVEL {gravel} mm / SAND {sand} mm",
        dxfattribs={"height": 250, "layer": "KJ_TEXT"},
    ).set_placement((0, -1100))

    msp.add_text(
        f"REBAR {data['rebar_class']} D{data['rebar_diam_mm']} STEP {step} mm",
        dxfattribs={"height": 250, "layer": "KJ_TEXT"},
    ).set_placement((0, -1500))

    sx, sy = 0, -3000
    cur = sy
    for name, th in [("SLAB", slab), ("GRAVEL", gravel), ("SAND", sand)]:
        msp.add_lwpolyline([(sx, cur), (sx + L, cur), (sx + L, cur - th), (sx, cur - th), (sx, cur)], dxfattribs={"layer": "KJ_SECTION"})
        msp.add_text(f"{name} {th} mm", dxfattribs={"height": 220, "layer": "KJ_TEXT"}).set_placement((sx + 300, cur - th / 2))
        cur -= th

    sheet_x = L + 2000
    sheet_y = 0
    for i, sh in enumerate(sheets, 1):
        msp.add_text(
            f"{i}. {sh['mark']} {sh['number']} {sh['title']}",
            dxfattribs={"height": 220, "layer": "KJ_SHEETS"},
        ).set_placement((sheet_x, sheet_y - i * 350))

    doc.saveas(path)
    return path

def _ff07_write_xlsx(path, data, sheets, materials, template):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    wb = Workbook()

    ws = wb.active
    ws.title = "Ведомость листов"
    ws.append(["№", "Марка", "Лист", "Наименование", "Источник"])
    for c in ws[1]:
        c.font = Font(bold=True)
        c.alignment = Alignment(horizontal="center")

    for i, sh in enumerate(sheets, 1):
        ws.append([i, sh["mark"], sh["number"], sh["title"], sh.get("source", "")])

    ws.column_dimensions["D"].width = 80
    ws.column_dimensions["E"].width = 25

    ws2 = wb.create_sheet("Спецификация")
    ws2.append(["№", "Наименование", "Ед. изм", "Кол-во", "Примечание"])
    for c in ws2[1]:
        c.font = Font(bold=True)

    for i, row in enumerate(materials, 1):
        ws2.append([i, row["name"], row["unit"], row["qty"], row["note"]])

    ws2.column_dimensions["B"].width = 70
    ws2.column_dimensions["E"].width = 50

    ws3 = wb.create_sheet("Параметры")
    for k in ["project_name", "section", "length_m", "width_m", "slab_mm", "sand_mm", "gravel_mm", "concrete_class", "rebar_class", "rebar_diam_mm", "rebar_step_mm"]:
        ws3.append([k, data.get(k)])
    ws3.append(["template_file", template.get("_template_file", "")])
    ws3.append(["engine", "FULLFIX_07_PROJECT_DESIGN_CLOSURE"])
    ws3.column_dimensions["A"].width = 28
    ws3.column_dimensions["B"].width = 90

    wb.save(path)
    return path

def _ff07_verify_pdf_pages(path, min_pages):
    try:
        from pypdf import PdfReader
        return len(PdfReader(path).pages) >= int(min_pages)
    except Exception:
        return _os_ff07.path.exists(path) and _os_ff07.path.getsize(path) > 3000

async def create_project_pdf_dxf_artifact(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", require_template: bool = True) -> dict:
    data = _ff07_parse_request(raw_input, template_hint)
    template = _ff07_load_latest_template(topic_id, data.get("section"))

    res = {
        "success": False,
        "engine": "FULLFIX_07_PROJECT_DESIGN_CLOSURE",
        "section": data["section"],
        "pdf_path": "",
        "dxf_path": "",
        "xlsx_path": "",
        "manifest_path": "",
        "pdf_link": "",
        "dxf_link": "",
        "xlsx_link": "",
        "manifest_link": "",
        "template_file": "",
        "sheet_count": 0,
        "error": None,
        "data": data,
    }

    if require_template and not template:
        res["error"] = "PROJECT_TEMPLATE_MODEL_NOT_FOUND"
        return res

    sheets = _ff07_sheet_rows_from_template(template or {}, data["section"], data)
    materials = _ff07_material_rows(template or {}, data)

    if len(sheets) < 8:
        res["error"] = f"SHEET_REGISTER_TOO_SMALL:{len(sheets)}"
        res["template_file"] = (template or {}).get("_template_file", "")
        res["sheet_count"] = len(sheets)
        return res

    stamp = _dt_ff07.utcnow().strftime("%Y%m%d_%H%M%S")
    safe_task = _re_ff07.sub(r"[^A-Za-z0-9_-]+", "_", str(task_id or "manual"))[:24]
    out_dir = _Path_ff07(_tempfile_ff07.gettempdir()) / f"areal_project_ff07_{safe_task}_{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    pdf_path = str(out_dir / f"{data['section']}_PROJECT_{safe_task}.pdf")
    dxf_path = str(out_dir / f"{data['section']}_PROJECT_{safe_task}.dxf")
    xlsx_path = str(out_dir / f"{data['section']}_SPEC_{safe_task}.xlsx")
    manifest_path = str(out_dir / f"{data['section']}_MANIFEST_{safe_task}.json")

    try:
        _ff07_write_pdf(pdf_path, data, template, sheets, materials)
        _ff07_write_dxf(dxf_path, data, sheets, materials)
        _ff07_write_xlsx(xlsx_path, data, sheets, materials, template)

        checks = [
            ("PDF_NOT_CREATED", pdf_path, 3000),
            ("DXF_NOT_CREATED", dxf_path, 500),
            ("XLSX_NOT_CREATED", xlsx_path, 1000),
        ]
        for err, fp, min_size in checks:
            if not _os_ff07.path.exists(fp) or _os_ff07.path.getsize(fp) < min_size:
                res["error"] = err
                return res

        if not _ff07_verify_pdf_pages(pdf_path, len(sheets)):
            res["error"] = "PDF_PAGE_COUNT_INVALID"
            return res

        manifest = {
            "schema": "AREAL_PROJECT_ARTIFACT_V4",
            "engine": "FULLFIX_07_PROJECT_DESIGN_CLOSURE",
            "created_at": _dt_ff07.utcnow().isoformat() + "Z",
            "task_id": task_id,
            "topic_id": topic_id,
            "input": raw_input,
            "template_file": template.get("_template_file", ""),
            "template_project_type": template.get("project_type", ""),
            "sheet_count": len(sheets),
            "sheets": sheets,
            "materials": materials,
            "data": data,
            "artifacts": {
                "pdf_path": pdf_path,
                "dxf_path": dxf_path,
                "xlsx_path": xlsx_path,
            },
        }
        _Path_ff07(manifest_path).write_text(_json_ff07.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

        from core.engine_base import upload_artifact_to_drive

        pdf_link = upload_artifact_to_drive(pdf_path, task_id, topic_id)
        dxf_link = upload_artifact_to_drive(dxf_path, task_id, topic_id)
        xlsx_link = upload_artifact_to_drive(xlsx_path, task_id, topic_id)
        manifest_link = upload_artifact_to_drive(manifest_path, task_id, topic_id)

        if not pdf_link:
            res["error"] = "PDF_UPLOAD_FAILED"
            return res
        if not dxf_link:
            res["error"] = "DXF_UPLOAD_FAILED"
            return res
        if not xlsx_link:
            res["error"] = "XLSX_UPLOAD_FAILED"
            return res

        res.update({
            "success": True,
            "pdf_path": pdf_path,
            "dxf_path": dxf_path,
            "xlsx_path": xlsx_path,
            "manifest_path": manifest_path,
            "pdf_link": str(pdf_link),
            "dxf_link": str(dxf_link),
            "xlsx_link": str(xlsx_link),
            "manifest_link": str(manifest_link or ""),
            "template_file": template.get("_template_file", ""),
            "sheet_count": len(sheets),
            "data": {**data, "template_file": template.get("_template_file", "")},
        })
        return res

    except Exception as e:
        res["error"] = str(e)[:700]
        return res

# === END FULLFIX_07_PROJECT_DESIGN_CLOSURE ===


# === FULLFIX_07_PROJECT_ENGINE_OVERRIDE ===
try:
    from core.cad_project_engine import (
        create_project_pdf_dxf_artifact,
        create_full_project_package,
        is_project_design_request,
        format_project_result_message,
    )
except Exception:
    pass
# === END FULLFIX_07_PROJECT_ENGINE_OVERRIDE ===


# === FULLFIX_08_PROJECT_SIGNATURE_COMPAT_OVERRIDE ===
# Final public project API used by task_worker
# Accepts any old/new call signature and delegates to CAD closure engine

try:
    from core.cad_project_engine import (
        create_project_pdf_dxf_artifact as _ff08_cad_create_project_pdf_dxf_artifact,
        create_full_project_documentation as _ff08_cad_create_full_project_documentation,
    )

    async def create_project_pdf_dxf_artifact(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", *args, **kwargs) -> dict:
        return await _ff08_cad_create_project_pdf_dxf_artifact(
            raw_input,
            task_id,
            int(topic_id or 0),
            str(template_hint or ""),
            *args,
            **kwargs
        )

    async def create_full_project_documentation(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", *args, **kwargs) -> dict:
        return await _ff08_cad_create_full_project_documentation(
            raw_input,
            task_id,
            int(topic_id or 0),
            str(template_hint or ""),
            *args,
            **kwargs
        )

except Exception as _ff08_import_error:
    async def create_project_pdf_dxf_artifact(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", *args, **kwargs) -> dict:
        return {
            "success": False,
            "engine": "FULLFIX_08_PROJECT_SIGNATURE_COMPAT_OVERRIDE",
            "error": "CAD_PROJECT_ENGINE_IMPORT_FAILED: " + str(_ff08_import_error)[:300],
        }

    async def create_full_project_documentation(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", *args, **kwargs) -> dict:
        return await create_project_pdf_dxf_artifact(raw_input, task_id, topic_id, template_hint, *args, **kwargs)

# === END FULLFIX_08_PROJECT_SIGNATURE_COMPAT_OVERRIDE ===



# === FULLFIX_09_PROJECT_TEMPLATE_REGISTER_PUBLIC_OVERRIDE ===
try:
    from core.cad_project_engine import (
        create_project_pdf_dxf_artifact as _ff09_create_project_pdf_dxf_artifact,
        create_full_project_documentation as _ff09_create_full_project_documentation,
    )

    async def create_project_pdf_dxf_artifact(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", *args, **kwargs) -> dict:
        return await _ff09_create_project_pdf_dxf_artifact(raw_input, task_id, topic_id, template_hint, *args, **kwargs)

    async def create_full_project_documentation(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", *args, **kwargs) -> dict:
        return await _ff09_create_full_project_documentation(raw_input, task_id, topic_id, template_hint, *args, **kwargs)

except Exception as _ff09_public_e:
    async def create_project_pdf_dxf_artifact(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", *args, **kwargs) -> dict:
        return {
            "success": False,
            "engine": "FULLFIX_09_PROJECT_TEMPLATE_REGISTER_PUBLIC_OVERRIDE",
            "error": str(_ff09_public_e)[:500],
        }

    async def create_full_project_documentation(raw_input: str, task_id: str, topic_id: int = 0, template_hint: str = "", *args, **kwargs) -> dict:
        return await create_project_pdf_dxf_artifact(raw_input, task_id, topic_id, template_hint, *args, **kwargs)

# === END FULLFIX_09_PROJECT_TEMPLATE_REGISTER_PUBLIC_OVERRIDE ===

