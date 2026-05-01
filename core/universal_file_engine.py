# === UNIVERSAL_FILE_ENGINE_V1 ===
from __future__ import annotations

import json
import os
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List

from core.format_registry import classify_file

def _clean(v: Any, limit: int = 20000) -> str:
    s = "" if v is None else str(v)
    s = s.replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()[:limit]

def _safe(v: Any, fallback: str = "file") -> str:
    s = re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", _clean(v, 120)).strip("._")
    return s or fallback

def _try_extract_text(path: str, file_name: str = "") -> str:
    ext = Path(file_name or path).suffix.lower()
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            return _clean("\n".join((p.extract_text() or "") for p in reader.pages[:50]), 50000)
        except Exception as e:
            return f"PDF_PARSE_ERROR: {e}"
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            return _clean("\n".join(p.text for p in doc.paragraphs if p.text), 50000)
        except Exception as e:
            return f"DOCX_PARSE_ERROR: {e}"
    if ext in (".txt", ".md", ".csv", ".json", ".xml", ".html", ".htm", ".yaml", ".yml"):
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return _clean(f.read(), 50000)
        except Exception as e:
            return f"TEXT_PARSE_ERROR: {e}"
    return ""

def _write_docx(model: Dict[str, Any], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"universal_file_report_{_safe(task_id)}.docx"
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("UNIVERSAL FILE REPORT", level=1)
        doc.add_paragraph(f"Файл: {model.get('file_name')}")
        doc.add_paragraph(f"Тип: {model.get('kind')}")
        doc.add_paragraph(f"Домен: {model.get('domain')}")
        doc.add_paragraph(f"Расширение: {model.get('extension')}")
        doc.add_paragraph(f"Размер: {model.get('size_bytes')} bytes")
        doc.add_paragraph(f"Engine hint: {model.get('engine_hint')}")
        doc.add_heading("Текст/превью", level=2)
        doc.add_paragraph(_clean(model.get("text_preview"), 12000) or "Текст не извлечён")
        doc.add_heading("Статус", level=2)
        doc.add_paragraph(model.get("status") or "INDEXED_METADATA")
        doc.save(out)
        return str(out)
    except Exception:
        txt = out.with_suffix(".txt")
        txt.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(txt)

def _write_json(model: Dict[str, Any], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"universal_file_model_{_safe(task_id)}.json"
    out.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(out)

def _write_xlsx(model: Dict[str, Any], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"universal_file_register_{_safe(task_id)}.xlsx"
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "File"
        for row, (k, v) in enumerate(model.items(), 1):
            ws.cell(row=row, column=1, value=str(k))
            ws.cell(row=row, column=2, value=json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else str(v))
        wb.save(out)
        wb.close()
        return str(out)
    except Exception:
        csv = out.with_suffix(".csv")
        csv.write_text("key,value\n" + "\n".join(f"{k},{v}" for k, v in model.items()), encoding="utf-8")
        return str(csv)

def _zip(paths: List[str], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"universal_file_package_{_safe(task_id)}.zip"
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if p and os.path.exists(p):
                z.write(p, arcname=os.path.basename(p))
        z.writestr("manifest.json", json.dumps({
            "engine": "UNIVERSAL_FILE_ENGINE_V1",
            "task_id": task_id,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "files": [os.path.basename(p) for p in paths if p and os.path.exists(p)],
        }, ensure_ascii=False, indent=2))
    return str(out)

def process_universal_file(
    local_path: str,
    file_name: str = "",
    mime_type: str = "",
    user_text: str = "",
    topic_role: str = "",
    task_id: str = "universal_file",
    topic_id: int = 0,
) -> Dict[str, Any]:
    if not local_path or not os.path.exists(local_path):
        return {"success": False, "error": "FILE_NOT_FOUND", "summary": "Файл не найден"}

    cls = classify_file(file_name or os.path.basename(local_path), mime_type, user_text, topic_role)
    size = os.path.getsize(local_path)
    text = _try_extract_text(local_path, file_name or local_path)

    model = {
        "schema": "UNIVERSAL_FILE_MODEL_V1",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "file_name": file_name or os.path.basename(local_path),
        "local_path": local_path,
        "mime_type": mime_type,
        "topic_id": topic_id,
        "user_text": user_text,
        "topic_role": topic_role,
        "size_bytes": size,
        "text_preview": _clean(text, 5000),
        **cls,
        "status": "INDEXED_WITH_TEXT" if text else "INDEXED_METADATA_ONLY",
    }

    docx = _write_docx(model, task_id)
    xlsx = _write_xlsx(model, task_id)
    js = _write_json(model, task_id)
    package = _zip([docx, xlsx, js], task_id)

    summary = "\n".join([
        "Универсальный файловый контур отработал",
        f"Файл: {model['file_name']}",
        f"Тип: {model['kind']}",
        f"Домен: {model['domain']}",
        f"Статус: {model['status']}",
        "Артефакты: DOCX + XLSX + JSON + ZIP",
    ])

    return {
        "success": True,
        "engine": "UNIVERSAL_FILE_ENGINE_V1",
        "summary": summary,
        "artifact_path": package,
        "artifact_name": f"{Path(model['file_name']).stem}_universal_file_package.zip",
        "extra_artifacts": [docx, xlsx, js],
        "model": model,
    }
# === END_UNIVERSAL_FILE_ENGINE_V1 ===
