# === DWG_DXF_PROJECT_CLOSE_V1 ===
from __future__ import annotations

import json
import math
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

SECTION_MAP = {
    "кж": "КЖ — Конструкции железобетонные",
    "км": "КМ — Конструкции металлические",
    "кмд": "КМД — Конструкции металлические деталировочные",
    "кр": "КР — Конструктивные решения",
    "ар": "АР — Архитектурные решения",
    "ов": "ОВ — Отопление и вентиляция",
    "вк": "ВК — Водоснабжение и канализация",
    "эом": "ЭОМ — Электрооборудование",
    "гп": "ГП — Генеральный план",
    "пз": "ПЗ — Пояснительная записка",
}

NORMS_MAP = {
    "кж": ["СП 63.13330.2018", "СП 20.13330.2016/2017", "ГОСТ 34028-2016", "ГОСТ 21.501-2018"],
    "км": ["СП 16.13330.2017", "СП 20.13330.2016/2017", "ГОСТ 27772-2015", "ГОСТ 21.502-2016"],
    "кмд": ["СП 16.13330.2017", "ГОСТ 23118-2019", "ГОСТ 21.502-2016"],
    "кр": ["СП 20.13330.2016/2017", "ГОСТ 21.501-2018"],
    "ар": ["ГОСТ 21.501-2018", "ГОСТ 21.101-2020", "СП 55.13330.2016"],
    "ов": ["СП 60.13330.2020", "ГОСТ 21.602-2016"],
    "вк": ["СП 30.13330.2020", "ГОСТ 21.601-2011"],
    "эом": ["ПУЭ-7", "СП 256.1325800.2016", "ГОСТ 21.608-2014"],
    "гп": ["СП 42.13330.2016", "ГОСТ 21.508-2020"],
}

ENTITY_PROJECT_HINTS = {
    "LINE": "линейная геометрия",
    "LWPOLYLINE": "полилинии/контуры",
    "POLYLINE": "полилинии/контуры",
    "CIRCLE": "окружности/отверстия",
    "ARC": "дуги",
    "TEXT": "текстовые подписи",
    "MTEXT": "многострочные подписи",
    "DIMENSION": "размеры",
    "INSERT": "блоки/узлы",
    "HATCH": "штриховки",
}

def _clean(v: Any, limit: int = 2000) -> str:
    if v is None:
        return ""
    s = str(v).replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()[:limit]

def _safe_name(v: Any, fallback: str = "drawing") -> str:
    s = re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", _clean(v, 120)).strip("._")
    return s or fallback

def _detect_section(file_name: str = "", user_text: str = "", drawing_text: str = "") -> str:
    hay_sources = [file_name or "", user_text or "", drawing_text[:2000] if drawing_text else ""]
    priority = ("кж", "кмд", "км", "кр", "ар", "ов", "вк", "эом", "гп", "пз")
    for src in hay_sources:
        low = src.lower()
        up = src.upper()
        for key in priority:
            if re.search(rf"(^|[^А-ЯA-Z0-9]){re.escape(key.upper())}([^А-ЯA-Z0-9]|$)", up):
                return key
            if key in low:
                return key
    return "кр"

def _read_bytes_head(path: str, n: int = 64) -> bytes:
    try:
        with open(path, "rb") as f:
            return f.read(n)
    except Exception:
        return b""

def _try_read_text(path: str, limit: int = 4_000_000) -> str:
    data = b""
    try:
        with open(path, "rb") as f:
            data = f.read(limit)
    except Exception:
        return ""
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(enc, errors="ignore")
        except Exception:
            pass
    return data.decode("latin-1", errors="ignore")

def _file_signature(path: str) -> str:
    head = _read_bytes_head(path, 32)
    if head.startswith(b"AC10"):
        return head[:12].decode("latin-1", errors="ignore")
    txt = head.decode("latin-1", errors="ignore")
    if "SECTION" in _try_read_text(path, 4096).upper():
        return "ASCII_DXF"
    return head[:16].hex()

def _try_convert_dwg_to_dxf(path: str) -> Optional[str]:
    src = Path(path)
    if src.suffix.lower() != ".dwg":
        return None

    tmp = Path(tempfile.mkdtemp(prefix="dwg_convert_"))
    in_dir = tmp / "in"
    out_dir = tmp / "out"
    in_dir.mkdir(parents=True, exist_ok=True)
    out_dir.mkdir(parents=True, exist_ok=True)

    safe = in_dir / src.name
    shutil.copy2(src, safe)

    converters = [
        shutil.which("dwg2dxf"),
        shutil.which("ODAFileConverter"),
        shutil.which("ODAFileConverter.exe"),
    ]

    for conv in [c for c in converters if c]:
        try:
            name = os.path.basename(conv).lower()
            if "dwg2dxf" in name:
                out = out_dir / (src.stem + ".dxf")
                subprocess.run([conv, str(safe), str(out)], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)
                if out.exists() and out.stat().st_size > 100:
                    return str(out)
            else:
                subprocess.run(
                    [conv, str(in_dir), str(out_dir), "ACAD2018", "DXF", "0", "1"],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=180,
                )
                found = list(out_dir.rglob("*.dxf"))
                if found:
                    found.sort(key=lambda p: p.stat().st_size, reverse=True)
                    if found[0].stat().st_size > 100:
                        return str(found[0])
        except Exception:
            continue

    return None

def _parse_ascii_dxf(path: str) -> Dict[str, Any]:
    text = _try_read_text(path)
    lines = [x.rstrip("\n") for x in text.splitlines()]
    pairs: List[Tuple[str, str]] = []
    i = 0
    while i + 1 < len(lines):
        code = lines[i].strip()
        value = lines[i + 1].strip()
        pairs.append((code, value))
        i += 2

    entities = []
    current: Optional[Dict[str, Any]] = None
    in_entities = False

    for code, value in pairs:
        if code == "0" and value.upper() == "SECTION":
            current = None
            continue
        if code == "2" and value.upper() == "ENTITIES":
            in_entities = True
            continue
        if code == "0" and value.upper() == "ENDSEC":
            if current:
                entities.append(current)
                current = None
            in_entities = False
            continue
        if not in_entities:
            continue

        if code == "0":
            if current:
                entities.append(current)
            current = {"type": value.upper(), "layer": "", "points": [], "texts": [], "raw": {}}
            continue

        if current is None:
            continue

        current["raw"].setdefault(code, []).append(value)
        if code == "8":
            current["layer"] = value
        elif code in ("1", "2", "3"):
            if value and len(value) <= 500:
                current["texts"].append(value)
        elif code in ("10", "20", "30", "11", "21", "31", "12", "22", "32"):
            try:
                current["points"].append((code, float(str(value).replace(",", "."))))
            except Exception:
                pass

    if current:
        entities.append(current)

    entity_counts = Counter(e.get("type") or "UNKNOWN" for e in entities)
    layer_counts = Counter(e.get("layer") or "0" for e in entities)

    texts = []
    for e in entities:
        for t in e.get("texts") or []:
            if t and t not in texts:
                texts.append(t)
            if len(texts) >= 80:
                break

    x_vals, y_vals = [], []
    for e in entities:
        coords = e.get("points") or []
        for code, val in coords:
            if code in ("10", "11", "12"):
                x_vals.append(val)
            elif code in ("20", "21", "22"):
                y_vals.append(val)

    extents = {}
    if x_vals and y_vals:
        extents = {
            "min_x": min(x_vals),
            "max_x": max(x_vals),
            "min_y": min(y_vals),
            "max_y": max(y_vals),
            "width": max(x_vals) - min(x_vals),
            "height": max(y_vals) - min(y_vals),
        }

    dims = []
    for t in texts:
        for m in re.findall(r"(?<!\d)(\d{2,6})(?!\d)", t):
            try:
                v = int(m)
                if 10 <= v <= 100000:
                    dims.append(v)
            except Exception:
                pass
    dims = sorted(set(dims))[:120]

    return {
        "parse_status": "DXF_PARSED",
        "raw_text_chars": len(text),
        "entities_total": len(entities),
        "entity_counts": dict(entity_counts.most_common(60)),
        "layers": dict(layer_counts.most_common(80)),
        "texts": texts[:80],
        "dimensions_detected": dims,
        "extents": extents,
    }

def _parse_dwg_metadata(path: str) -> Dict[str, Any]:
    p = Path(path)
    sig = _file_signature(path)
    return {
        "parse_status": "DWG_BINARY_METADATA_ONLY",
        "signature": sig,
        "file_size": p.stat().st_size if p.exists() else 0,
        "note": "DWG binary parsed as metadata only. For geometry extraction install ODAFileConverter or dwg2dxf on server; DXF is parsed directly",
    }

def _build_model(local_path: str, file_name: str, mime_type: str = "", user_text: str = "", topic_role: str = "") -> Dict[str, Any]:
    p = Path(local_path)
    ext = p.suffix.lower()
    source_path = str(p)
    converted_from_dwg = False

    if ext == ".dwg":
        converted = _try_convert_dwg_to_dxf(local_path)
        if converted:
            source_path = converted
            ext = ".dxf"
            converted_from_dwg = True

    if ext == ".dxf":
        parsed = _parse_ascii_dxf(source_path)
    elif p.suffix.lower() == ".dwg":
        parsed = _parse_dwg_metadata(local_path)
    else:
        parsed = {
            "parse_status": "UNSUPPORTED_DRAWING_FORMAT",
            "signature": _file_signature(local_path),
            "file_size": p.stat().st_size if p.exists() else 0,
        }

    drawing_text = "\n".join(parsed.get("texts") or [])
    section = _detect_section(file_name, user_text, drawing_text)
    entity_counts = parsed.get("entity_counts") or {}
    layers = parsed.get("layers") or {}
    texts = parsed.get("texts") or []

    output_documents = [
        "DOCX_DWG_DXF_ANALYSIS_REPORT",
        "XLSX_DWG_DXF_ENTITY_REGISTER",
        "ZIP_PROJECT_PACKAGE",
    ]

    if section in ("кж", "км", "кмд", "кр"):
        output_documents.extend(["SPECIFICATION_DRAFT", "STRUCTURAL_DRAWING_REGISTER"])
    if section == "ар":
        output_documents.extend(["ARCHITECTURAL_SHEET_REGISTER", "ROOM_PLAN_REVIEW"])

    risk_flags = []
    if parsed.get("parse_status") == "DWG_BINARY_METADATA_ONLY":
        risk_flags.append("DWG_GEOMETRY_NOT_EXTRACTED_WITHOUT_CONVERTER")
    if not layers:
        risk_flags.append("LAYERS_NOT_FOUND")
    if not entity_counts:
        risk_flags.append("ENTITIES_NOT_FOUND")
    if not texts:
        risk_flags.append("TEXT_LABELS_NOT_FOUND")

    model = {
        "schema": "DWG_DXF_PROJECT_MODEL_V1",
        "source_file": file_name or p.name,
        "source_path": local_path,
        "mime_type": mime_type,
        "topic_role": topic_role,
        "user_text": user_text,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "section": section,
        "section_title": SECTION_MAP.get(section, section.upper()),
        "norms": NORMS_MAP.get(section, []),
        "converted_from_dwg": converted_from_dwg,
        "parse": parsed,
        "layers": layers,
        "entity_counts": entity_counts,
        "texts": texts[:80],
        "dimensions_detected": parsed.get("dimensions_detected") or [],
        "extents": parsed.get("extents") or {},
        "output_documents": output_documents,
        "risk_flags": risk_flags,
        "status": "PARTIAL" if risk_flags else "CONFIRMED",
    }
    return model

def _summary(model: Dict[str, Any]) -> str:
    parse = model.get("parse") or {}
    entity_counts = model.get("entity_counts") or {}
    layers = model.get("layers") or {}
    risk_flags = model.get("risk_flags") or []
    lines = [
        "DWG/DXF проектный контур отработал",
        f"Файл: {model.get('source_file')}",
        f"Раздел: {model.get('section_title')}",
        f"Статус: {model.get('status')}",
        f"Parse: {parse.get('parse_status')}",
        f"Слоёв: {len(layers)}",
        f"Сущностей: {sum(int(v) for v in entity_counts.values()) if entity_counts else 0}",
    ]
    if entity_counts:
        lines.append("Типы сущностей: " + ", ".join(f"{k}:{v}" for k, v in list(entity_counts.items())[:12]))
    if layers:
        lines.append("Слои: " + ", ".join(list(layers.keys())[:20]))
    if model.get("dimensions_detected"):
        lines.append("Размеры/числа из подписей: " + ", ".join(map(str, model.get("dimensions_detected")[:30])))
    if model.get("norms"):
        lines.append("Нормы: " + ", ".join(model.get("norms")))
    if risk_flags:
        lines.append("Ограничения: " + ", ".join(risk_flags))
    lines.append("Артефакты: DOCX отчёт + XLSX реестр + ZIP пакет")
    return "\n".join(lines).strip()

def _write_docx(model: Dict[str, Any], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"dwg_dxf_report_{_safe_name(task_id, 'manual')}.docx"
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("DWG/DXF PROJECT MODEL", level=1)
        doc.add_paragraph(f"Файл: {model.get('source_file')}")
        doc.add_paragraph(f"Раздел: {model.get('section_title')}")
        doc.add_paragraph(f"Статус: {model.get('status')}")
        doc.add_paragraph(f"Parse: {(model.get('parse') or {}).get('parse_status')}")

        doc.add_heading("Нормативная база", level=2)
        norms = model.get("norms") or []
        if norms:
            for n in norms:
                doc.add_paragraph(f"• {n}")
        else:
            doc.add_paragraph("Норма не подтверждена")

        doc.add_heading("Слои", level=2)
        layers = model.get("layers") or {}
        if layers:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Слой"
            table.rows[0].cells[1].text = "Кол-во"
            for name, cnt in list(layers.items())[:80]:
                row = table.add_row().cells
                row[0].text = str(name)
                row[1].text = str(cnt)
        else:
            doc.add_paragraph("Слои не извлечены")

        doc.add_heading("Сущности", level=2)
        ents = model.get("entity_counts") or {}
        if ents:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Тип"
            table.rows[0].cells[1].text = "Кол-во"
            table.rows[0].cells[2].text = "Назначение"
            for name, cnt in list(ents.items())[:80]:
                row = table.add_row().cells
                row[0].text = str(name)
                row[1].text = str(cnt)
                row[2].text = ENTITY_PROJECT_HINTS.get(str(name), "")
        else:
            doc.add_paragraph("Сущности не извлечены")

        doc.add_heading("Текстовые подписи", level=2)
        texts = model.get("texts") or []
        if texts:
            for t in texts[:60]:
                doc.add_paragraph(f"• {t}")
        else:
            doc.add_paragraph("Текстовые подписи не извлечены")

        doc.add_heading("Проектные выходные документы", level=2)
        for x in model.get("output_documents") or []:
            doc.add_paragraph(f"• {x}")

        doc.add_heading("Ограничения", level=2)
        risks = model.get("risk_flags") or []
        if risks:
            for r in risks:
                doc.add_paragraph(f"• {r}")
        else:
            doc.add_paragraph("Критичных ограничений не выявлено")

        doc.save(out)
        return str(out)
    except Exception:
        txt = out.with_suffix(".txt")
        txt.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(txt)

def _write_xlsx(model: Dict[str, Any], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"dwg_dxf_register_{_safe_name(task_id, 'manual')}.xlsx"
    try:
        from openpyxl import Workbook
        wb = Workbook()

        ws = wb.active
        ws.title = "Summary"
        rows = [
            ("Файл", model.get("source_file")),
            ("Раздел", model.get("section_title")),
            ("Статус", model.get("status")),
            ("Parse", (model.get("parse") or {}).get("parse_status")),
            ("Нормы", ", ".join(model.get("norms") or [])),
            ("Ограничения", ", ".join(model.get("risk_flags") or [])),
        ]
        for i, (k, v) in enumerate(rows, 1):
            ws.cell(i, 1, k)
            ws.cell(i, 2, v)

        ws2 = wb.create_sheet("Layers")
        ws2.append(["Слой", "Кол-во"])
        for k, v in (model.get("layers") or {}).items():
            ws2.append([k, v])

        ws3 = wb.create_sheet("Entities")
        ws3.append(["Тип", "Кол-во", "Назначение"])
        for k, v in (model.get("entity_counts") or {}).items():
            ws3.append([k, v, ENTITY_PROJECT_HINTS.get(str(k), "")])

        ws4 = wb.create_sheet("Texts")
        ws4.append(["№", "Текст"])
        for i, t in enumerate(model.get("texts") or [], 1):
            ws4.append([i, t])

        ws5 = wb.create_sheet("ModelJSON")
        raw = json.dumps(model, ensure_ascii=False, indent=2)
        for i, line in enumerate(raw.splitlines(), 1):
            ws5.cell(i, 1, line)

        wb.save(out)
        wb.close()
        return str(out)
    except Exception:
        csv = out.with_suffix(".csv")
        lines = ["key,value"]
        lines.append(f"file,{model.get('source_file')}")
        lines.append(f"section,{model.get('section_title')}")
        lines.append(f"status,{model.get('status')}")
        csv.write_text("\n".join(lines), encoding="utf-8")
        return str(csv)

def _write_json(model: Dict[str, Any], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"dwg_dxf_model_{_safe_name(task_id, 'manual')}.json"
    out.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(out)

def _zip_artifacts(paths: List[str], task_id: str, source_file: str = "") -> str:
    out = Path(tempfile.gettempdir()) / f"dwg_dxf_project_package_{_safe_name(task_id, 'manual')}.zip"
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if p and os.path.exists(p):
                z.write(p, arcname=os.path.basename(p))
        manifest = {
            "created_at": datetime.now(timezone.utc).isoformat(),
            "task_id": task_id,
            "source_file": source_file,
            "files": [os.path.basename(p) for p in paths if p and os.path.exists(p)],
            "engine": "DWG_DXF_PROJECT_CLOSE_V1",
        }
        z.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
    return str(out)

def process_drawing_file(
    local_path: str,
    file_name: str = "",
    mime_type: str = "",
    user_text: str = "",
    topic_role: str = "",
    task_id: str = "artifact",
    topic_id: int = 0,
) -> Dict[str, Any]:
    if not local_path or not os.path.exists(local_path):
        return {
            "success": False,
            "error": "DRAWING_FILE_NOT_FOUND",
            "summary": "DWG/DXF файл не найден",
        }

    model = _build_model(local_path, file_name or os.path.basename(local_path), mime_type, user_text, topic_role)
    docx = _write_docx(model, task_id)
    xlsx = _write_xlsx(model, task_id)
    js = _write_json(model, task_id)
    package = _zip_artifacts([docx, xlsx, js], task_id, model.get("source_file") or file_name)

    return {
        "success": True,
        "engine": "DWG_DXF_PROJECT_CLOSE_V1",
        "summary": _summary(model),
        "model": model,
        "docx_path": docx,
        "xlsx_path": xlsx,
        "json_path": js,
        "artifact_path": package,
        "artifact_name": f"{Path(file_name or local_path).stem}_dwg_dxf_project_package.zip",
        "extra_artifacts": [docx, xlsx, js],
        "status": model.get("status"),
    }

async def process_drawing_file_async(*args, **kwargs) -> Dict[str, Any]:
    return process_drawing_file(*args, **kwargs)

# === END_DWG_DXF_PROJECT_CLOSE_V1 ===
