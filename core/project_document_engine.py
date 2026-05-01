# === PROJECT_DOCUMENT_ENGINE_V1 ===
from __future__ import annotations

import json
import os
import re
import tempfile
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List

SECTION_MAP = {
    "кж": "КЖ — Конструкции железобетонные",
    "км": "КМ — Конструкции металлические",
    "кмд": "КМД — Конструкции металлические деталировочные",
    "кр": "КР — Конструктивные решения",
    "ар": "АР — Архитектурные решения",
    "ов": "ОВ — Отопление и вентиляция",
    "вк": "ВК — Водоснабжение и канализация",
    "эом": "ЭОМ — Электрооборудование",
    "гп": "ГП — Генеральный план",
    "пз": "ПЗ — Пояснительная записка",
}

NORMS_MAP = {
    "кж": ["СП 63.13330.2018", "СП 20.13330.2016/2017", "ГОСТ 21.501-2018", "ГОСТ 34028-2016"],
    "км": ["СП 16.13330.2017", "ГОСТ 23118-2019", "ГОСТ 21.502-2016"],
    "кмд": ["СП 16.13330.2017", "ГОСТ 21.502-2016", "ГОСТ 23118-2019"],
    "кр": ["СП 20.13330.2016/2017", "ГОСТ 21.501-2018"],
    "ар": ["ГОСТ 21.101-2020", "ГОСТ 21.501-2018", "СП 55.13330.2016"],
    "ов": ["СП 60.13330.2020", "ГОСТ 21.602-2016"],
    "вк": ["СП 30.13330.2020", "ГОСТ 21.601-2011"],
    "эом": ["ПУЭ-7", "СП 256.1325800.2016", "ГОСТ 21.608-2014"],
    "гп": ["СП 42.13330.2016", "ГОСТ 21.508-2020"],
}

def _clean(v: Any, limit: int = 20000) -> str:
    s = "" if v is None else str(v)
    s = s.replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()[:limit]

def _safe(v: Any, fallback: str = "project_document") -> str:
    s = re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", _clean(v, 160)).strip("._")
    return s or fallback

def _extract_text(path: str, file_name: str = "") -> str:
    ext = Path(file_name or path).suffix.lower()
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            parts = []
            for page in reader.pages[:80]:
                try:
                    parts.append(page.extract_text() or "")
                except Exception:
                    pass
            return _clean("\n".join(parts), 50000)
        except Exception as e:
            return f"PDF_PARSE_ERROR: {e}"
    if ext == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            return _clean("\n".join(p.text for p in doc.paragraphs if p.text), 50000)
        except Exception as e:
            return f"DOCX_PARSE_ERROR: {e}"
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return _clean(f.read(), 50000)
    except Exception as e:
        return f"TEXT_PARSE_ERROR: {e}"

def _detect_section(file_name: str, user_text: str, text: str) -> str:
    hay = f"{file_name}\n{user_text}\n{text[:5000]}".lower()
    up = hay.upper()
    for key in ("кж", "кмд", "км", "кр", "ар", "ов", "вк", "эом", "гп", "пз"):
        if key in hay or re.search(rf"(^|[^А-ЯA-Z0-9]){key.upper()}([^А-ЯA-Z0-9]|$)", up):
            return key
    if any(x in hay for x in ("фундамент", "плита", "бетон", "арматур", "монолит")):
        return "кж"
    if any(x in hay for x in ("архитектур", "планиров", "фасад", "разрез")):
        return "ар"
    if any(x in hay for x in ("отоплен", "вентиляц")):
        return "ов"
    if any(x in hay for x in ("водоснаб", "канализац", "вк")):
        return "вк"
    return "кр"

def _extract_design_items(text: str) -> List[Dict[str, str]]:
    patterns = [
        ("бетон", r"\bB\d{2,3}\b|В\d{2,3}\b"),
        ("арматура", r"\bA\d{3}\b|А\d{3}\b|Ø\s*\d+|Ф\s*\d+|\b\d{1,2}\s*мм\b"),
        ("сталь", r"\bC\d{3}\b|С\d{3}\b|09Г2С|С245|С255|С345"),
        ("лист", r"\b\d+[,.]?\d*\s*мм\b"),
        ("размер", r"\b\d{3,6}\s*[xх×]\s*\d{3,6}\b|\b\d{3,6}\s*мм\b"),
    ]
    out = []
    for name, pat in patterns:
        found = sorted(set(re.findall(pat, text, flags=re.I)))
        for v in found[:30]:
            out.append({"type": name, "value": str(v), "note": ""})
    return out[:200]

def _build_model(path: str, file_name: str, user_text: str = "", topic_role: str = "") -> Dict[str, Any]:
    text = _extract_text(path, file_name)
    section = _detect_section(file_name, user_text, text)
    items = _extract_design_items(text)
    return {
        "schema": "PROJECT_DOCUMENT_MODEL_V1",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "source_file": file_name or os.path.basename(path),
        "source_path": path,
        "section": section,
        "section_title": SECTION_MAP.get(section, section.upper()),
        "norms": NORMS_MAP.get(section, []),
        "topic_role": topic_role,
        "user_text": user_text,
        "text_chars": len(text or ""),
        "text_preview": _clean(text, 5000),
        "items": items,
        "output_documents": [
            "DOCX_PROJECT_REVIEW",
            "XLSX_PROJECT_REGISTER",
            "JSON_PROJECT_MODEL",
            "ZIP_PROJECT_PACKAGE",
        ],
        "status": "CONFIRMED" if text and not text.endswith("_ERROR") else "PARTIAL",
    }

def _write_docx(model: Dict[str, Any], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"project_document_report_{_safe(task_id)}.docx"
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("PROJECT DOCUMENT MODEL", level=1)
        doc.add_paragraph(f"Файл: {model.get('source_file')}")
        doc.add_paragraph(f"Раздел: {model.get('section_title')}")
        doc.add_paragraph(f"Статус: {model.get('status')}")
        doc.add_heading("Нормативная база", level=2)
        for n in model.get("norms") or []:
            doc.add_paragraph(f"• {n}")
        doc.add_heading("Выделенные проектные параметры", level=2)
        items = model.get("items") or []
        if items:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Тип"
            table.rows[0].cells[1].text = "Значение"
            table.rows[0].cells[2].text = "Примечание"
            for it in items[:120]:
                row = table.add_row().cells
                row[0].text = str(it.get("type") or "")
                row[1].text = str(it.get("value") or "")
                row[2].text = str(it.get("note") or "")
        else:
            doc.add_paragraph("Проектные параметры не выделены")
        doc.add_heading("Текстовая сводка", level=2)
        doc.add_paragraph(_clean(model.get("text_preview"), 12000) or "Текст не извлечён")
        doc.save(out)
        return str(out)
    except Exception:
        out_txt = out.with_suffix(".txt")
        out_txt.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
        return str(out_txt)

def _write_xlsx(model: Dict[str, Any], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"project_document_register_{_safe(task_id)}.xlsx"
    try:
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Summary"
        rows = [
            ("Файл", model.get("source_file")),
            ("Раздел", model.get("section_title")),
            ("Статус", model.get("status")),
            ("Нормы", ", ".join(model.get("norms") or [])),
            ("Символов текста", model.get("text_chars")),
        ]
        for i, (k, v) in enumerate(rows, 1):
            ws.cell(i, 1, k)
            ws.cell(i, 2, v)
        ws2 = wb.create_sheet("Items")
        ws2.append(["Тип", "Значение", "Примечание"])
        for it in model.get("items") or []:
            ws2.append([it.get("type"), it.get("value"), it.get("note")])
        ws3 = wb.create_sheet("ModelJSON")
        for i, line in enumerate(json.dumps(model, ensure_ascii=False, indent=2).splitlines(), 1):
            ws3.cell(i, 1, line)
        wb.save(out)
        wb.close()
        return str(out)
    except Exception:
        out_csv = out.with_suffix(".csv")
        out_csv.write_text("type,value,note\n", encoding="utf-8")
        return str(out_csv)

def _write_json(model: Dict[str, Any], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"project_document_model_{_safe(task_id)}.json"
    out.write_text(json.dumps(model, ensure_ascii=False, indent=2), encoding="utf-8")
    return str(out)

def _zip(paths: List[str], task_id: str) -> str:
    out = Path(tempfile.gettempdir()) / f"project_document_package_{_safe(task_id)}.zip"
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for p in paths:
            if p and os.path.exists(p):
                z.write(p, arcname=os.path.basename(p))
        z.writestr("manifest.json", json.dumps({
            "engine": "PROJECT_DOCUMENT_ENGINE_V1",
            "task_id": task_id,
            "files": [os.path.basename(p) for p in paths if p and os.path.exists(p)],
            "created_at": datetime.now(timezone.utc).isoformat(),
        }, ensure_ascii=False, indent=2))
    return str(out)

async def process_project_document(
    file_path: str,
    file_name: str = "",
    user_text: str = "",
    topic_role: str = "",
    task_id: str = "artifact",
    topic_id: int = 0,
) -> Dict[str, Any]:
    if not file_path or not os.path.exists(file_path):
        return {"success": False, "error": "PROJECT_DOCUMENT_FILE_NOT_FOUND"}
    model = _build_model(file_path, file_name or os.path.basename(file_path), user_text, topic_role)
    docx = _write_docx(model, task_id)
    xlsx = _write_xlsx(model, task_id)
    js = _write_json(model, task_id)
    package = _zip([docx, xlsx, js], task_id)
    summary = "\n".join([
        "Проектный документ обработан",
        f"Файл: {model.get('source_file')}",
        f"Раздел: {model.get('section_title')}",
        f"Статус: {model.get('status')}",
        f"Нормы: {', '.join(model.get('norms') or [])}",
        f"Проектных параметров: {len(model.get('items') or [])}",
        "Артефакты: DOCX отчёт + XLSX реестр + JSON модель + ZIP пакет",
    ])
    return {
        "success": True,
        "engine": "PROJECT_DOCUMENT_ENGINE_V1",
        "summary": summary,
        "model": model,
        "artifact_path": package,
        "artifact_name": f"{Path(file_name or file_path).stem}_project_document_package.zip",
        "docx_path": docx,
        "xlsx_path": xlsx,
        "json_path": js,
        "extra_artifacts": [docx, xlsx, js],
    }

# === END_PROJECT_DOCUMENT_ENGINE_V1 ===
