# === FULLFIX_15_DEFECT_ACT ===
import os, logging
from datetime import date
logger = logging.getLogger(__name__)
ENGINE = "FULLFIX_15_DEFECT_ACT"
RUNTIME_DIR = "/root/.areal-neva-core/runtime"
os.makedirs(RUNTIME_DIR, exist_ok=True)
ACT_PHRASES = ["акт", "дефект", "осмотр", "сделай акт", "по фото", "технадзор", "нарушение"]

def is_defect_act_intent(text, mime_type=""):
    t = (text or "").lower()
    return "image" in (mime_type or "") and any(p in t for p in ACT_PHRASES)

def generate_act_docx(task_id, caption, file_name, object_name="UNKNOWN"):
    from docx import Document
    path = os.path.join(RUNTIME_DIR, "act_" + task_id[:8] + ".docx")
    doc = Document()
    doc.add_heading("АКТ ОСМОТРА / ДЕФЕКТНАЯ ВЕДОМОСТЬ", 0)
    today = date.today().strftime("%d.%m.%Y")
    doc.add_paragraph("Дата: " + today)
    doc.add_paragraph("Объект: " + (object_name or "UNKNOWN"))
    doc.add_paragraph("Основание: фото — " + (file_name or ""))
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for i, h in enumerate(["№", "Фото/файл", "Описание дефекта", "Локация", "Рекомендация", "Статус"]):
        table.rows[0].cells[i].text = h
    row = table.add_row().cells
    row[0].text = "1"; row[1].text = file_name or ""; row[2].text = caption or "требует уточнения"
    row[3].text = "-"; row[4].text = "Устранить"; row[5].text = "Открыт"
    doc.add_paragraph("Заключение: зафиксированы дефекты, требующие устранения.")
    doc.add_paragraph("Составил: ________________________  Дата: " + today)
    doc.save(path)
    return path

def generate_act_pdf(task_id, caption, file_name, object_name="UNKNOWN"):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Spacer
    from reportlab.lib import colors
    from core.pdf_cyrillic import register_cyrillic_fonts, make_styles, make_paragraph, clean_pdf_text, FONT_REGULAR, FONT_BOLD
    path = os.path.join(RUNTIME_DIR, "act_" + task_id[:8] + ".pdf")
    register_cyrillic_fonts()
    styles = make_styles()
    today = date.today().strftime("%d.%m.%Y")
    doc = SimpleDocTemplate(path, pagesize=A4, topMargin=20, bottomMargin=20, leftMargin=20, rightMargin=20)
    story = [
        make_paragraph("АКТ ОСМОТРА / ДЕФЕКТНАЯ ВЕДОМОСТЬ", "header", styles), Spacer(1,8),
        make_paragraph("Дата: " + today, "normal", styles),
        make_paragraph("Объект: " + (object_name or "UNKNOWN"), "normal", styles),
        make_paragraph("Основание: " + (file_name or ""), "normal", styles),
        Spacer(1,10),
    ]
    data = [
        [make_paragraph(h, "bold", styles) for h in ["№", "Файл", "Описание", "Локация", "Рекомендация", "Статус"]],
        [make_paragraph(x, "normal", styles) for x in ["1", clean_pdf_text(file_name or ""), clean_pdf_text(caption or "требует уточнения"), "-", "Устранить", "Открыт"]],
    ]
    tbl = Table(data, colWidths=[22, 70, 150, 55, 80, 50])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#444444")),
        ("TEXTCOLOR",(0,0),(-1,0),colors.white),
        ("FONTNAME",(0,0),(-1,0),FONT_BOLD),
        ("FONTSIZE",(0,0),(-1,-1),8),
        ("GRID",(0,0),(-1,-1),0.4,colors.black),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
    ]))
    story.append(tbl); story.append(Spacer(1,12))
    story.append(make_paragraph("Заключение: зафиксированы дефекты, требующие устранения.", "normal", styles))
    story.append(make_paragraph("Составил: ________________________  Дата: " + today, "normal", styles))
    doc.build(story)
    return path

def process_defect_act_sync(conn, task_id, chat_id, topic_id, raw_input, file_name="", local_path=""):
    from core.artifact_upload_guard import upload_many_or_fail
    from core.reply_sender import send_reply_ex
    # === FULLFIX_20_GEMINI_DEFECT_SYNC ===
    _ff20_vision_text = ""
    try:
        if local_path and any(str(local_path).lower().endswith(ext) for ext in (".jpg",".jpeg",".png",".webp",".heic")):
            import asyncio as _ff20_aio
            from core.gemini_vision import analyze_image_file as _ff20_gif
            try:
                _ff20_aio.get_running_loop()
            except RuntimeError:
                _ff20_vision_text = _ff20_aio.run(
                    _ff20_gif(local_path, prompt="\u041e\u043f\u0438\u0448\u0438 \u0434\u0435\u0444\u0435\u043a\u0442 \u0434\u043b\u044f \u0430\u043a\u0442\u0430", timeout=60)
                ) or ""
            logger.info("FF20_GEMINI_DEFECT_SYNC len=%s", len(_ff20_vision_text))
    except Exception as _ff20_ve:
        logger.warning("FF20_GEMINI_DEFECT_SYNC_ERR=%s", _ff20_ve)
    if _ff20_vision_text:
        raw_input = str(raw_input or "") + "\n\n\u041e\u043f\u0438\u0441\u0430\u043d\u0438\u0435 \u0434\u0435\u0444\u0435\u043a\u0442\u0430: " + str(_ff20_vision_text)
    # === END FULLFIX_20_GEMINI_DEFECT_SYNC ===
    # === NORMATIVE_DB_V1_WIRED ===
    try:
        import asyncio as _norm_aio
        from core.normative_db import search_norms as _search_norms
        _norm_desc = str(raw_input or "") + " " + str(_ff20_vision_text or "")
        try:
            _norm_loop = _norm_aio.get_running_loop()
            _norm_results = []
        except RuntimeError:
            _norm_results = _norm_aio.run(_search_norms(_norm_desc))
        if _norm_results:
            _norm_lines = ["\n\nНормативные требования:"]
            for _n in _norm_results:
                _norm_lines.append(f"  {_n['norm_id']}: {_n['requirement'][:200]}")
            raw_input = str(raw_input or "") + "\n".join(_norm_lines)
    except Exception as _ne:
        logger.warning("NORMATIVE_DB_V1_WIRED err=%s", _ne)
    # === END NORMATIVE_DB_V1_WIRED ===


    try:
        caption = raw_input or file_name
        docx_path = generate_act_docx(task_id, caption, file_name)
        pdf_path = generate_act_pdf(task_id, caption, file_name)
        files = [{"path": pdf_path, "kind": "act_pdf"}, {"path": docx_path, "kind": "act_docx"}]
        up = upload_many_or_fail(files, task_id, topic_id)
        pdf_r = up["results"].get(pdf_path, {}); docx_r = up["results"].get(docx_path, {})
        lines = ["Акт осмотра готов."]
        if pdf_r.get("success") and pdf_r.get("link"): lines.append("PDF: " + pdf_r["link"])
        if docx_r.get("success") and docx_r.get("link"): lines.append("DOCX: " + docx_r["link"])
        if len(lines) == 1: lines.append("Drive недоступен. Файл: " + (file_name or ""))
        result_text = "\n".join(lines)
        conn.execute("UPDATE tasks SET state='AWAITING_CONFIRMATION',result=?,updated_at=datetime('now') WHERE id=?", (result_text, task_id))
        conn.execute("INSERT INTO task_history(task_id,action,created_at) VALUES(?,?,datetime('now'))", (task_id, "state:AWAITING_CONFIRMATION"))
        conn.commit()
        try:
            _br = send_reply_ex(chat_id=str(chat_id), text=result_text, reply_to_message_id=None)
            _bmid = None
            if isinstance(_br, dict): _bmid = _br.get("bot_message_id") or _br.get("message_id")
            elif _br and hasattr(_br, "message_id"): _bmid = _br.message_id
            if _bmid:
                conn.execute("UPDATE tasks SET bot_message_id=? WHERE id=?", (str(_bmid), task_id))
                conn.commit()
        except Exception as _se:
            logger.error("ACT_SEND_ERR task=%s err=%s", task_id, _se)
        return True
    except Exception as e:
        logger.error("DEFECT_ACT_ERROR task=%s err=%s", task_id, e)
        return False

async def process_defect_act(conn, task_id, chat_id, topic_id, raw_input, file_name="", local_path=""):
    import asyncio
    return await asyncio.get_event_loop().run_in_executor(
        None, process_defect_act_sync, conn, task_id, chat_id, topic_id, raw_input, file_name, local_path
    )
# === END FULLFIX_15_DEFECT_ACT ===
