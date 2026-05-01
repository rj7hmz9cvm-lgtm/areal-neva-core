import os
import re
import csv
import json
import base64
import tempfile
from typing import Any, Dict, List, Optional

from dotenv import load_dotenv

BASE = "/root/.areal-neva-core"
load_dotenv(f"{BASE}/.env", override=True)

def _s(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, str):
        return v.strip()
    try:
        return json.dumps(v, ensure_ascii=False)
    except Exception:
        return str(v).strip()

def _clean(text: str, limit: int = 12000) -> str:
    text = (text or "").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:limit]

def _kind(file_name: str, mime_type: str = "") -> str:
    # === UNIVERSAL_FORMAT_REGISTRY_V1_KIND ===
    # === DWG_DXF_KIND_FIX_V1_ARTIFACT_PIPELINE ===
    try:
        from core.format_registry import classify_file
        return classify_file(file_name, mime_type).get("kind") or "binary"
    except Exception:
        ext = os.path.splitext((file_name or "").lower())[1]
        mime = (mime_type or "").lower()

        # drawing first: mimetypes may classify .dwg/.dxf as image/*
        if ext in (".dwg", ".dxf", ".ifc", ".rvt", ".rfa", ".skp", ".stl", ".obj", ".step", ".stp", ".iges", ".igs") or any(x in mime for x in ("dxf", "dwg", "ifc", "cad", "step", "stp", "iges", "igs")):
            return "drawing"
        if ext in (".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif", ".tif", ".tiff", ".bmp", ".gif") or mime.startswith("image/"):
            return "image"
        if ext in (".xlsx", ".xls", ".xlsm", ".csv", ".ods", ".tsv") or "spreadsheet" in mime or mime in ("text/csv", "application/vnd.ms-excel"):
            return "table"
        if ext in (".pdf", ".docx", ".doc", ".txt", ".md", ".rtf", ".odt", ".html", ".htm", ".xml", ".json", ".yaml", ".yml") or mime in ("application/pdf", "text/plain"):
            return "document"
        if ext in (".ppt", ".pptx", ".odp", ".key"):
            return "presentation"
        if ext in (".zip", ".7z", ".rar", ".tar", ".gz", ".tgz"):
            return "archive"
        if ext in (".mp4", ".mov", ".avi", ".mkv", ".mp3", ".wav", ".m4a", ".ogg"):
            return "media"
        return "binary"
    # === END_DWG_DXF_KIND_FIX_V1_ARTIFACT_PIPELINE ===
    # === END_UNIVERSAL_FORMAT_REGISTRY_V1_KIND ===


# === DOMAIN_CONTOUR_ROUTER_V1 ===
def _artifact_task_id(file_name: str, engine: str = "artifact") -> str:
    base = re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", os.path.splitext(os.path.basename(file_name or engine))[0]).strip("._")
    return (base or engine)[:80]

def _domain_flags(file_name: str, mime_type: str = "", user_text: str = "", topic_role: str = "") -> Dict[str, bool]:
    hay = f"{file_name}\n{mime_type}\n{user_text}\n{topic_role}".lower()
    estimate = any(x in hay for x in ("смет", "расчёт", "расчет", "вор", "ведомость объем", "ведомость объём", "estimate", "xlsx", "xls", "csv"))
    tech = any(x in hay for x in ("технадзор", "дефект", "акт", "осмотр", "нарушен", "предписан", "гост", "снип", "сп ", "фотофиксац", "трещин", "протеч", "скол"))
    project = any(x in hay for x in ("проект", "проектирован", "кж", "кмд", "км", "кр", "ар", "ов", "вк", "эом", "гп", "пз", "чертеж", "чертёж", "dxf", "dwg"))
    return {"estimate": estimate, "tech": tech, "project": project}


# === ESTIMATE_PDF_PACKAGE_V2 ===
def _pdf_escape_v2(text: str) -> str:
    return str(text or "").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

def _write_simple_pdf_v2(path: str, title: str, lines: List[str]) -> str:
    out = os.path.join(tempfile.gettempdir(), os.path.splitext(os.path.basename(path))[0] + "_estimate_summary.pdf")
    safe_lines = [_pdf_escape_v2(title or "Estimate summary")]
    safe_lines += [_pdf_escape_v2(x) for x in (lines or [])[:40]]

    stream_lines = ["BT", "/F1 11 Tf", "50 790 Td"]
    first = True
    for line in safe_lines:
        if not first:
            stream_lines.append("0 -16 Td")
        first = False
        stream_lines.append(f"({line[:105]}) Tj")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("utf-8", errors="ignore")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
    ]

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, 1):
        offsets.append(len(pdf))
        pdf.extend(f"{i} 0 obj\n".encode())
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")
    xref = len(pdf)
    pdf.extend(f"xref\n0 {len(objects)+1}\n".encode())
    pdf.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        pdf.extend(f"{off:010d} 00000 n \n".encode())
    pdf.extend(f"trailer << /Size {len(objects)+1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode())

    with open(out, "wb") as f:
        f.write(pdf)
    return out

def _zip_files_v2(files: List[str], name: str) -> str:
    import zipfile
    out = os.path.join(tempfile.gettempdir(), re.sub(r"[^A-Za-zА-Яа-я0-9_.-]+", "_", name or "estimate_package").strip("._") + ".zip")
    with zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for f in files or []:
            if f and os.path.exists(f):
                z.write(f, arcname=os.path.basename(f))
    return out
# === END_ESTIMATE_PDF_PACKAGE_V2 ===

async def _domain_estimate_artifact(local_path: str, file_name: str, mime_type: str, user_text: str, topic_role: str) -> Optional[Dict[str, Any]]:
    # === DOMAIN_ESTIMATE_PDF_XLSX_PACKAGE_V2 ===
    try:
        from core.estimate_engine import process_estimate_to_excel
        tid = _artifact_task_id(file_name, "estimate_artifact")
        res = await process_estimate_to_excel(local_path, tid, 0)

        if res and (res.get("success") or res.get("excel_path")) and res.get("excel_path"):
            excel_path = res.get("excel_path")
            link = res.get("drive_link") or ""
            lines = [
                f"Файл: {file_name}",
                "Engine: DOMAIN_ESTIMATE_ENGINE_V1",
                f"Drive: {link or 'не подтвержден'}",
                f"Status: {'OK' if excel_path else 'PARTIAL'}",
            ]
            if res.get("error"):
                lines.append(f"Ограничение: {res.get('error')}")

            pdf_path = _write_simple_pdf_v2(excel_path or local_path, "Сметный результат", lines)
            package = _zip_files_v2([excel_path, pdf_path], os.path.splitext(os.path.basename(file_name or "estimate"))[0] + "_estimate_package")

            summary = "Сметный файл обработан\nАртефакты: XLSX + PDF"
            if link:
                summary += f"\nExcel: {link}"
            else:
                summary += "\nExcel создан локально, Drive ссылка не подтверждена"
            summary += "\nPDF включён в ZIP пакет"

            return {
                "summary": summary,
                "artifact_path": package,
                "artifact_name": f"{os.path.splitext(os.path.basename(file_name))[0]}_estimate_package.zip",
                "engine": "DOMAIN_ESTIMATE_ENGINE_V1",
                "drive_link": link,
                "extra_artifacts": [excel_path, pdf_path],
            }

        reason = (res or {}).get("error") or "ESTIMATE_ENGINE_NO_ARTIFACT"
        pdf_path = _write_simple_pdf_v2(local_path, "Сметный файл принят без расчётного результата", [
            f"Файл: {file_name}",
            f"Причина: {reason}",
            "Расчётные строки не подтверждены",
        ])
        package = _zip_files_v2([pdf_path], os.path.splitext(os.path.basename(file_name or "estimate"))[0] + "_estimate_diagnostic_package")
        return {
            "summary": f"Сметный файл принят, но расчётные строки не подтверждены\nПричина: {reason}\nСоздан диагностический PDF пакет",
            "artifact_path": package,
            "artifact_name": f"{os.path.splitext(os.path.basename(file_name))[0]}_estimate_diagnostic_package.zip",
            "engine": "DOMAIN_ESTIMATE_ENGINE_V1",
            "error": str(reason)[:300],
            "extra_artifacts": [pdf_path],
        }
    except Exception as e:
        pdf_path = _write_simple_pdf_v2(local_path, "Сметный engine недоступен", [
            f"Файл: {file_name}",
            f"Ошибка: {e}",
            "Расчётные строки не подтверждены",
        ])
        package = _zip_files_v2([pdf_path], os.path.splitext(os.path.basename(file_name or "estimate"))[0] + "_estimate_error_package")
        return {
            "summary": f"Сметный engine недоступен: {e}\nСоздан диагностический PDF пакет",
            "artifact_path": package,
            "artifact_name": f"{os.path.splitext(os.path.basename(file_name))[0]}_estimate_error_package.zip",
            "engine": "DOMAIN_ESTIMATE_ENGINE_V1",
            "error": str(e)[:300],
            "extra_artifacts": [pdf_path],
        }
    # === END_DOMAIN_ESTIMATE_PDF_XLSX_PACKAGE_V2 ===

async def _domain_technadzor_artifact(local_path: str, file_name: str, mime_type: str, user_text: str, topic_role: str, extracted_text: str = "") -> Optional[Dict[str, Any]]:
    try:
        from core.technadzor_engine import process_technadzor
        tid = _artifact_task_id(file_name, "technadzor_artifact")
        raw = "\n".join(x for x in [user_text or "", extracted_text or "", topic_role or ""] if x).strip()
        res = process_technadzor(
            conn=None,
            task_id=tid,
            chat_id="artifact_pipeline",
            topic_id=0,
            raw_input=raw or "Технический осмотр файла",
            file_name=file_name,
            local_path=local_path,
        )
        if res and res.get("ok"):
            art = res.get("artifact") or {}
            path = art.get("path") or ""
            link = art.get("drive_link") or ""
            summary = _clean(res.get("result_text") or "Акт технического осмотра сформирован", 6000)
            if link and link not in summary:
                summary += f"\n\nДокумент: {link}"
            return {
                "summary": summary,
                "artifact_path": path,
                "artifact_name": f"{os.path.splitext(os.path.basename(file_name))[0]}_technadzor_act.docx",
                "engine": "DOMAIN_TECHNADZOR_ENGINE_V1",
                "drive_link": link,
            }
    except Exception as e:
        return {
            "summary": f"Технадзор engine недоступен: {e}",
            "artifact_path": "",
            "artifact_name": "",
            "engine": "DOMAIN_TECHNADZOR_ENGINE_V1",
            "error": str(e)[:300],
        }
    return None

async def _domain_project_document_artifact(local_path: str, file_name: str, mime_type: str, user_text: str, topic_role: str) -> Optional[Dict[str, Any]]:
    try:
        from core.project_document_engine import process_project_document
        tid = _artifact_task_id(file_name, "project_document")
        res = await process_project_document(
            file_path=local_path,
            file_name=file_name,
            user_text=user_text,
            topic_role=topic_role,
            task_id=tid,
            topic_id=0,
        )
        if res and res.get("success"):
            return {
                "summary": _clean(res.get("summary") or "Проектный документ обработан", 6000),
                "artifact_path": res.get("artifact_path"),
                "artifact_name": res.get("artifact_name") or f"{os.path.splitext(os.path.basename(file_name))[0]}_project_document_package.zip",
                "extra_artifacts": res.get("extra_artifacts") or [],
                "engine": "PROJECT_DOCUMENT_ENGINE_V1",
                "model": res.get("model") or {},
            }
    except Exception as e:
        return {
            "summary": f"Project document engine недоступен: {e}",
            "artifact_path": "",
            "artifact_name": "",
            "engine": "PROJECT_DOCUMENT_ENGINE_V1",
            "error": str(e)[:300],
        }
    return None
# === END_DOMAIN_CONTOUR_ROUTER_V1 ===

def _build_word(title: str, summary: str, defects: List[Dict[str, Any]], recommendations: List[str], sources: List[str]) -> str:
    from docx import Document

    fd, out = tempfile.mkstemp(prefix="artifact_", suffix=".docx", dir="/tmp")
    os.close(fd)

    doc = Document()
    doc.add_heading(title or "Результат обработки", level=1)

    if summary:
        doc.add_paragraph(_clean(summary, 12000))

    if sources:
        doc.add_heading("Источники", level=2)
        for s in sources:
            doc.add_paragraph(_s(s))

    doc.add_heading("Замечания", level=2)
    if defects:
        for idx, item in enumerate(defects, 1):
            p = doc.add_paragraph()
            p.add_run(f"{idx}. ").bold = True
            p.add_run(_s(item.get("title")) or "Замечание")
            sev = _s(item.get("severity"))
            if sev:
                p.add_run(f" [{sev}]")
            desc = _s(item.get("description"))
            if desc:
                doc.add_paragraph(desc)
    else:
        doc.add_paragraph("Замечания не выделены")

    doc.add_heading("Рекомендации", level=2)
    if recommendations:
        for r in recommendations:
            doc.add_paragraph(_s(r))
    else:
        doc.add_paragraph("Рекомендации не сформированы")

    doc.save(out)
    return out

def _build_excel(title: str, items: List[Dict[str, str]], summary: str, sources: List[str]) -> str:
    from openpyxl import Workbook

    fd, out = tempfile.mkstemp(prefix="artifact_", suffix=".xlsx", dir="/tmp")
    os.close(fd)

    wb = Workbook()
    ws = wb.active
    ws.title = "Результат"
    ws["A1"] = title or "Табличный результат"
    ws["A2"] = _clean(summary, 1000)

    row = 4
    headers = ["№", "Наименование", "Ед", "Кол-во", "Примечание"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=row, column=col, value=h)
    row += 1

    for idx, item in enumerate(items, 1):
        ws.cell(row=row, column=1, value=idx)
        ws.cell(row=row, column=2, value=_s(item.get("name")))
        ws.cell(row=row, column=3, value=_s(item.get("unit")))
        ws.cell(row=row, column=4, value=_s(item.get("qty")))
        ws.cell(row=row, column=5, value=_s(item.get("note")))
        row += 1

    src = wb.create_sheet("Источники")
    for idx, s in enumerate(sources, 1):
        src.cell(row=idx, column=1, value=_s(s))

    wb.save(out)
    return out

def _extract_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                pass
        return _clean("\n".join(parts), 12000)
    except Exception as e:
        return f"PDF_PARSE_ERROR: {e}"

def _extract_docx(path: str) -> str:
    try:
        from docx import Document
        doc = Document(path)
        return _clean("\n".join(p.text for p in doc.paragraphs if p.text), 12000)
    except Exception as e:
        return f"DOCX_PARSE_ERROR: {e}"

def _extract_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return _clean(f.read(), 12000)
    except Exception as e:
        return f"TXT_PARSE_ERROR: {e}"

def _extract_table_items(path: str, file_name: str) -> List[Dict[str, str]]:
    rows: List[List[str]] = []
    ext = os.path.splitext((file_name or "").lower())[1]

    try:
        if ext == ".csv":
            with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
                reader = csv.reader(f)
                for idx, row in enumerate(reader):
                    rows.append([_s(x) for x in row])
                    if idx >= 300:
                        break
        else:
            from openpyxl import load_workbook
            wb = load_workbook(path, data_only=True, read_only=True)
            for ws in wb.worksheets[:3]:
                rows.append([f"__SHEET__:{ws.title}"])
                for idx, row in enumerate(ws.iter_rows(values_only=True)):
                    rows.append([_s(x) for x in row])
                    if idx >= 300:
                        break
    except Exception as e:
        rows.append([f"TABLE_PARSE_ERROR: {e}"])

    items: List[Dict[str, str]] = []
    unit_re = re.compile(r"\b(м2|м3|м\.п\.|п\.м\.|шт|кг|тн|т|м)\b", re.I)
    qty_re = re.compile(r"^\d+[.,]?\d*$")

    for row in rows:
        if not row:
            continue
        if len(row) == 1 and _s(row[0]).startswith("__SHEET__:"):
            continue

        cleaned = [_s(x) for x in row if _s(x)]
        if not cleaned:
            continue

        name = cleaned[0]
        unit = ""
        qty = ""
        note = ""

        for cell in cleaned[1:]:
            if not unit:
                m = unit_re.search(cell)
                if m:
                    unit = m.group(1)
                    continue
            if not qty and qty_re.match(cell.replace(" ", "")):
                qty = cell.replace(" ", "")
                continue
            note = (note + " | " + cell).strip(" |") if note else cell

        if len(name) < 2:
            continue

        items.append({
            "name": name[:500],
            "unit": unit[:32],
            "qty": qty[:64],
            "note": note[:500],
        })

        if len(items) >= 500:
            break

    return items

async def _vision_image(path: str, user_text: str, topic_role: str) -> Optional[Dict[str, Any]]:
    api_key = (os.getenv("OPENROUTER_API_KEY") or "").strip()
    if not api_key:
        return None

    base_url = (os.getenv("OPENROUTER_BASE_URL") or "https://openrouter.ai/api/v1").strip().rstrip("/")
    model = (os.getenv("OPENROUTER_VISION_MODEL") or "google/gemini-2.5-flash").strip()

    ext = os.path.splitext(path)[1].lower().lstrip(".") or "jpeg"
    mime = f"image/{'jpeg' if ext in ('jpg', 'jpeg') else ext}"

    try:
        import httpx
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")

        prompt = (
            "Ты анализируешь строительную фотофиксацию\n"
            f"Роль чата: {topic_role or 'технадзор'}\n"
            f"Задача пользователя: {user_text or 'проанализируй фото'}\n\n"
            "Верни только JSON вида:\n"
            "{\n"
            '  "summary": "краткое резюме",\n'
            '  "defects": [{"title":"...", "description":"...", "severity":"low|medium|high"}],\n'
            '  "recommendations": ["...", "..."]\n'
            "}"
        )

        body = {
            "model": model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    ],
                }
            ],
            "temperature": 0.1,
        }
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }

        async with httpx.AsyncClient(timeout=httpx.Timeout(180.0, connect=30.0)) as client:
            r = await client.post(f"{base_url}/chat/completions", headers=headers, json=body)
            r.raise_for_status()
            data = r.json()

        content = data["choices"][0]["message"]["content"]
        if isinstance(content, list):
            content = "\n".join(x.get("text", "") if isinstance(x, dict) else str(x) for x in content)
        content = _clean(_s(content), 12000)

        try:
            parsed = json.loads(content)
            if isinstance(parsed, dict):
                return parsed
        except Exception:
            return {
                "summary": content[:3000],
                "defects": [],
                "recommendations": [],
            }

    except Exception:
        return None

    return None

async def analyze_downloaded_file(local_path: str, file_name: str, mime_type: str = "", user_text: str = "", topic_role: str = "") -> Optional[Dict[str, Any]]:
    kind = _kind(file_name, mime_type)
    sources = [file_name]

    if kind == "image":
        # === DOMAIN_TECHNADZOR_IMAGE_ASYNC_VISION_V2 ===
        flags = _domain_flags(file_name, mime_type, user_text, topic_role)
        vision_text = ""
        analysis = None

        if flags.get("tech") or not flags.get("estimate"):
            analysis = await _vision_image(local_path, user_text, topic_role)
            if isinstance(analysis, dict):
                vision_text = json.dumps(analysis, ensure_ascii=False)
            elif analysis:
                vision_text = str(analysis)

            routed = await _domain_technadzor_artifact(local_path, file_name, mime_type, user_text, topic_role, vision_text)
            if routed and (routed.get("artifact_path") or routed.get("summary")):
                routed["engine"] = routed.get("engine") or "DOMAIN_TECHNADZOR_IMAGE_ASYNC_VISION_V2"
                return routed

        if analysis is None:
            analysis = await _vision_image(local_path, user_text, topic_role)

        if not analysis:
            routed = await _domain_technadzor_artifact(local_path, file_name, mime_type, user_text, topic_role, "Фото принято без vision-анализа")
            if routed and (routed.get("artifact_path") or routed.get("summary")):
                return routed
            return None

        summary = _s(analysis.get("summary")) if isinstance(analysis, dict) else _s(analysis)
        summary = summary or "Фото проанализировано"
        defects = analysis.get("defects") if isinstance(analysis, dict) and isinstance(analysis.get("defects"), list) else []
        recommendations = analysis.get("recommendations") if isinstance(analysis, dict) and isinstance(analysis.get("recommendations"), list) else []
        artifact_path = _build_word("Акт замечаний по фотофиксации", summary, defects, [_s(x) for x in recommendations], sources)
        return {
            "summary": summary,
            "artifact_path": artifact_path,
            "artifact_name": f"{os.path.splitext(os.path.basename(file_name))[0]}_photo_report.docx",
            "engine": "DOMAIN_TECHNADZOR_IMAGE_ASYNC_VISION_V2",
        }
        # === END_DOMAIN_TECHNADZOR_IMAGE_ASYNC_VISION_V2 ===

    if kind == "table":
        flags = _domain_flags(file_name, mime_type, user_text, topic_role)
        if flags.get("estimate") or not flags.get("project"):
            routed = await _domain_estimate_artifact(local_path, file_name, mime_type, user_text, topic_role)
            if routed and (routed.get("artifact_path") or routed.get("summary")):
                return routed

        items = _extract_table_items(local_path, file_name)
        summary = f"Нормализовано позиций: {len(items)}"
        artifact_path = _build_excel("Сметный/табличный результат", items, summary, sources)
        return {
            "summary": summary,
            "artifact_path": artifact_path,
            "artifact_name": f"{os.path.splitext(os.path.basename(file_name))[0]}_estimate.xlsx",
            "engine": "TABLE_FALLBACK_ENGINE",
        }

    if kind == "drawing":
        try:
            from core.dwg_engine import process_drawing_file
            data = process_drawing_file(
                local_path=local_path,
                file_name=file_name,
                mime_type=mime_type,
                user_text=user_text,
                topic_role=topic_role,
                task_id="artifact",
                topic_id=0,
            )
            if data and data.get("success"):
                return {
                    "summary": _clean(data.get("summary") or "DWG/DXF файл обработан", 4000),
                    "artifact_path": data.get("artifact_path"),
                    "artifact_name": data.get("artifact_name") or f"{os.path.splitext(os.path.basename(file_name))[0]}_dwg_dxf_project_package.zip",
                    "extra_artifacts": data.get("extra_artifacts") or [],
                    "engine": "DWG_DXF_PROJECT_CLOSE_V1",
                    "model": data.get("model") or {},
                }
            return {
                "summary": _clean((data or {}).get("summary") or (data or {}).get("error") or "DWG/DXF файл не обработан", 3000),
                "artifact_path": "",
                "artifact_name": "",
                "engine": "DWG_DXF_PROJECT_CLOSE_V1",
                "error": (data or {}).get("error") or "DRAWING_PROCESS_FAILED",
            }
        except Exception as e:
            return {
                "summary": f"DWG/DXF обработка завершилась ошибкой: {e}",
                "artifact_path": "",
                "artifact_name": "",
                "engine": "DWG_DXF_PROJECT_CLOSE_V1",
                "error": str(e)[:300],
            }

    if kind == "document":
        flags = _domain_flags(file_name, mime_type, user_text, topic_role)
        ext = os.path.splitext((file_name or "").lower())[1]
        if ext == ".pdf":
            domain_text = _extract_pdf(local_path)
        elif ext == ".docx":
            domain_text = _extract_docx(local_path)
        else:
            domain_text = _extract_txt(local_path)

        if flags.get("tech"):
            routed = await _domain_technadzor_artifact(local_path, file_name, mime_type, user_text, topic_role, domain_text)
            if routed and (routed.get("artifact_path") or routed.get("summary")):
                return routed

        if flags.get("estimate"):
            routed = await _domain_estimate_artifact(local_path, file_name, mime_type, user_text, topic_role)
            if routed and (routed.get("artifact_path") or routed.get("summary")):
                return routed

        if flags.get("project"):
            routed = await _domain_project_document_artifact(local_path, file_name, mime_type, user_text, topic_role)
            if routed and (routed.get("artifact_path") or routed.get("summary")):
                return routed

        summary = _clean(domain_text, 3000) if domain_text else "Документ обработан"
        artifact_path = _build_word("Сводка по документу", summary, [], [], sources)
        return {
            "summary": summary,
            "artifact_path": artifact_path,
            "artifact_name": f"{os.path.splitext(os.path.basename(file_name))[0]}_document_summary.docx",
            "engine": "DOCUMENT_FALLBACK_ENGINE",
        }

    # === UNIVERSAL_FILE_ENGINE_FALLBACK_V1 ===
    try:
        from core.universal_file_engine import process_universal_file
        data = process_universal_file(
            local_path=local_path,
            file_name=file_name,
            mime_type=mime_type,
            user_text=user_text,
            topic_role=topic_role,
            task_id=_artifact_task_id(file_name, "universal_file"),
            topic_id=0,
        )
        if data and data.get("success"):
            return {
                "summary": _clean(data.get("summary") or "Файл обработан универсальным контуром", 6000),
                "artifact_path": data.get("artifact_path"),
                "artifact_name": data.get("artifact_name") or f"{os.path.splitext(os.path.basename(file_name))[0]}_universal_file_package.zip",
                "extra_artifacts": data.get("extra_artifacts") or [],
                "engine": "UNIVERSAL_FILE_ENGINE_V1",
                "model": data.get("model") or {},
            }
    except Exception as e:
        return {
            "summary": f"Универсальный файловый контур завершился ошибкой: {e}",
            "artifact_path": "",
            "artifact_name": "",
            "engine": "UNIVERSAL_FILE_ENGINE_V1",
            "error": str(e)[:300],
        }
    return None
    # === END_UNIVERSAL_FILE_ENGINE_FALLBACK_V1 ===
