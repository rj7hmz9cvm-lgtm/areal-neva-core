import logging
import math
from pathlib import Path
from typing import Any, Dict, List

import pandas as pd
import pdfplumber
from docx import Document

try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
except Exception:
    pytesseract = None
    convert_from_path = None
    Image = None

logger = logging.getLogger("document_engine")

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"}


class DocumentEngine:
    @staticmethod
    def _clean(val: Any) -> str:
        if val is None:
            return ""
        if isinstance(val, float) and (math.isnan(val) or math.isinf(val)):
            return ""
        return str(val).strip()

    @staticmethod
    def _clean_table(table_data: List[List[Any]]) -> List[List[str]]:
        cleaned = []
        for row in table_data or []:
            if row is None:
                continue
            c_row = [DocumentEngine._clean(c) for c in row]
            if any(cell != "" for cell in c_row):
                cleaned.append(c_row)
        return cleaned

    @staticmethod
    def process(local_path: str) -> Dict[str, Any]:
        path = Path(local_path)
        ext = path.suffix.lower()

        res = {
            "status": "ok",
            "error": "",
            "file_path": str(path),
            "text": "",
            "tables": [],
            "meta": {
                "size": path.stat().st_size if path.exists() else 0,
                "file_name": path.name,
                "method": "native",
            },
        }

        try:
            if ext == ".pdf":
                DocumentEngine._parse_pdf(path, res)
                if (
                    not res["text"].strip()
                    and not res["tables"]
                    and pytesseract is not None
                    and convert_from_path is not None
                ):
                    logger.info("OCR triggered for PDF %s", path.name)
                    images = convert_from_path(str(path))
                    texts = []
                    for img in images:
                        texts.append(pytesseract.image_to_string(img, lang="rus+eng"))
                    res["text"] = "\n".join(texts).strip()
                    res["meta"]["method"] = "ocr_pdf"

            elif ext in [".xlsx", ".xls"]:
                DocumentEngine._parse_excel(path, res)

            elif ext == ".docx":
                DocumentEngine._parse_docx(path, res)

            elif ext == ".csv":
                DocumentEngine._parse_csv(path, res)

            elif ext in IMAGE_EXTS:
                DocumentEngine._parse_image_ocr(path, res)

            else:
                res["status"] = "failed"
                res["error"] = f"Unsupported extension: {ext}"
                return res

        except Exception as e:
            res["status"] = "failed"
            res["error"] = f"{type(e).__name__}: {e}"

        return res

    @staticmethod
    def _parse_pdf(path: Path, res: Dict[str, Any]) -> None:
        with pdfplumber.open(path) as pdf:
            res["meta"]["pages"] = len(pdf.pages)
            txt = []
            for i, page in enumerate(pdf.pages, 1):
                txt.append(page.extract_text() or "")
                for j, table in enumerate(page.extract_tables() or [], 1):
                    clean = DocumentEngine._clean_table(table)
                    if clean:
                        res["tables"].append({"title": f"P{i}_T{j}", "data": clean})
            res["text"] = "\n".join(txt).strip()

    @staticmethod
    def _parse_excel(path: Path, res: Dict[str, Any]) -> None:
        engine = "xlrd" if path.suffix.lower() == ".xls" else None
        with pd.ExcelFile(path, engine=engine) as xls:
            sheet_names = xls.sheet_names
            res["meta"]["sheets"] = sheet_names
            for sheet in sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet, dtype=object)
                clean = DocumentEngine._clean_table([df.columns.tolist()] + df.values.tolist())
                if clean:
                    res["tables"].append({"title": sheet, "data": clean})
            res["text"] = f"Sheets: {', '.join(sheet_names)}"

    @staticmethod
    def _parse_docx(path: Path, res: Dict[str, Any]) -> None:
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        res["text"] = "\n".join(paragraphs)

        for i, table in enumerate(doc.tables, 1):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            clean = DocumentEngine._clean_table(rows)
            if clean:
                res["tables"].append({"title": f"T{i}", "data": clean})

    @staticmethod
    def _parse_csv(path: Path, res: Dict[str, Any]) -> None:
        last_err = None
        for enc in ("utf-8", "cp1251", "utf-8-sig"):
            try:
                df = pd.read_csv(path, encoding=enc, sep=None, engine="python", dtype=object)
                clean = DocumentEngine._clean_table([df.columns.tolist()] + df.values.tolist())
                if clean:
                    res["tables"].append({"title": path.name, "data": clean})
                res["text"] = f"CSV Content: {path.name}"
                return
            except Exception as e:
                last_err = e
        raise RuntimeError(f"CSV parse failed: {last_err}")

    @staticmethod
    def _parse_image_ocr(path: Path, res: Dict[str, Any]) -> None:
        if pytesseract is None or Image is None:
            raise RuntimeError("OCR dependencies are not available")
        img = Image.open(path)
        res["text"] = pytesseract.image_to_string(img, lang="rus+eng").strip()
        res["meta"]["method"] = "ocr_image"
